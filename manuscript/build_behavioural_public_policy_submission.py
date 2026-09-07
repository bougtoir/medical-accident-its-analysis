#!/usr/bin/env python3
"""Build a Behavioural Public Policy (Cambridge Core) submission package for the
rate-based analysis of malpractice-litigation risk and Japanese specialty-level
physician workforce allocation.

Outputs (all derived from results/reanalysis_results.json and data_primary/):
  - manuscript/bpp_manuscript_en.docx   anonymised main manuscript
  - manuscript/bpp_title_page.docx      title page with author info
  - manuscript/bpp_cover_letter.docx    cover letter addressed to BPP
  - manuscript/bpp_highlights.docx      3-5 short highlights (optional)
  - manuscript/bpp_supplementary.docx     supplementary figures & tables
  - manuscript/bpp_figures.pptx           editable main figure slides
  - manuscript/bpp_supplementary_figures.pptx editable supplementary slides
  - output/ha_Figure_*.png              main figure files (reused)
  - output/ha_Supplementary_Figure_*.png supplementary figure files (reused)
  - output/bpp_submission.zip            bundled submission package

The main manuscript is anonymised and uses the same reproducible data pipeline
as the Healthcare Analytics submission; only the framing, title, abstract,
introduction, discussion, cover letter and declarations are re-written for
Behavioural Public Policy.
"""
import os
import re
import shutil
import zipfile
import build_healthcare_analytics_submission as ha
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pptx import Presentation
from pptx.util import Inches as PInches, Pt as PPt

BASE = ha.BASE
OUT = ha.OUT
PUBLIC_REPO = ha.PUBLIC_REPO

# Local aliases for readability
N = ha.N_SPECIALTIES
BIEN = ha.BIEN
YEARS = ha.YEARS
GREW = ha.GREW
FELL = ha.FELL
SPAN = ha.SPAN
SURG_DESC = ha.SURG_DESC
SURG_PCT = ha.SURG_PCT
PHYS = ha.PHYS
HOSP = ha.HOSP
REV = ha.REV
CNT = ha.CNT
ANN = ha.ANN
INT = ha.INT
EQP = ha.EQP
EQH = ha.EQH
BS_PHYS = ha.BS_PHYS
BS_HOSP = ha.BS_HOSP
JMSR = ha.JMSR
MEDIA = ha.MEDIA
JOCS_HOLM = ha.JOCS_HOLM
MARGIN1 = ha.MARGIN1
MARGIN2 = ha.MARGIN2
PER = ha.PER
JMSR_START = ha.JMSR_START
MEDIA_START = ha.MEDIA_START
MEDIA_END = ha.MEDIA_END
SIM = ha.SIM
TOTAL_SIM = ha.TOTAL_SIM
SURG_SIM = ha.SURG_SIM
SP = ha.SP
n_pos = sum(1 for v in SP.values() if v["rho"] > 0)
n_sig = sum(1 for v in SP.values() if v["p"] < 0.05)
DESCR = ha.DESCR
SENKOI = ha.SENKOI

def _senkoi_summary():
    if SENKOI is None:
        return None
    coverages = [(v["coverage_pct"], k) for k, v in SENKOI.items()]
    lo, hi = min(coverages), max(coverages)
    total_senkoi = sum(v["senkoi_2018"] for v in SENKOI.values())
    total_pool = sum(v["physicians_3_5_yr_2014"] for v in SENKOI.values())
    weighted = 100 * total_senkoi / total_pool if total_pool else 0
    return {
        "min_val": lo[0],
        "min_spec": lo[1],
        "max_val": hi[0],
        "max_spec": hi[1],
        "mean_val": sum(c[0] for c in coverages) / len(coverages),
        "weighted_val": weighted,
        "total_senkoi": total_senkoi,
        "total_pool": total_pool,
    }

SENKOI_SUM = _senkoi_summary()

# Look-up tables for the new heterogeneity and trend-sensitivity results.
_HET_DICT = {(r["outcome"], r["group"]): r for r in ha.RES.get("heterogeneity", [])}
_TREND_DICT = {r["outcome"]: r for r in ha.RES.get("trend_sensitivity", [])}

BPP_TITLE = (
    "Risk perception versus structural incentives in physician specialty choice: "
    "a Japanese panel study"
)

# Anonymous manuscript placeholder; the real repository URL is on the title page
# and in the cover letter, which are removed during double-blind peer review.
ANON_REPO = "URL omitted for anonymised peer review"

# Placeholder for the title-page data-availability line during anonymised peer review.
TITLE_REPO_PLACEHOLDER = "[repository URL to be inserted on acceptance]"

# Placeholder for the cover-letter repository line until an anonymised git link is ready.
COVER_REPO_PLACEHOLDER = "[anonymous GitHub repository URL to be supplied]"

# Short in-text author names for organisational references
_BPP_CITE_SHORT = {
    "phys": "Ministry of Health, Labour and Welfare",
    "facil": "Ministry of Health, Labour and Welfare",
    "mhlw_senkoi2018": "Ministry of Health, Labour and Welfare",
    "mhlw_3_5yr": "Ministry of Health, Labour and Welfare",
    "court": "Supreme Court of Japan",
    "jocscp": "JOCS-CP",
    "mais": "Medical Accident Investigation System",
    "jmsr_data": "JMSR",
    "nikkei": "Nikkei Inc",
}

# Full author names to use on first in-text citation (subsequent citations use the abbreviation)
_BPP_CITE_FIRST = {
    "jocscp": "Japan Council for Quality Health Care (JOCS-CP)",
    "jmsr_data": "Japan Medical Safety Research Organisation (JMSR)",
    "mais": "Medical Accident Investigation System",
    "court": "Supreme Court of Japan",
}

_CITE_RE = re.compile(r"\{([^}]+)\}")


def _bpp_get_cite_year(ref: str, key: str) -> str:
    """Return the publication year with a Harvard-style a/b/c suffix when needed."""
    return _bpp_get_year(ref) + _BPP_YEAR_SUFFIX.get(key, "")


def _bpp_get_year(ref: str) -> str:
    """Extract a 4-digit publication year from a reference string."""
    m = re.search(r";\s*(\d{4})", ref)
    if m:
        return m.group(1)
    m = re.search(r"(\d{4});", ref)
    if m:
        return m.group(1)
    m = re.search(r"Accessed\s+(\d{4})", ref)
    if m:
        return m.group(1)
    m = re.search(r"\b(19|20)\d{2}\b", ref)
    if m:
        return m.group(0)
    return "????"


def _bpp_build_year_suffixes(refs: dict) -> dict:
    """Assign a/b/c suffixes to references with the same author and year."""
    groups = {}
    for key, ref in refs.items():
        year = _bpp_get_year(ref)
        author_raw = ref.split(". ", 1)[0].strip().lower()
        groups.setdefault((author_raw, year), []).append(key)
    suffix = {}
    for (author, year), keys in groups.items():
        if len(keys) > 1:
            for i, k in enumerate(sorted(keys, key=lambda kk: refs[kk].lower())):
                suffix[k] = chr(ord("a") + i)
    return suffix


_BPP_YEAR_SUFFIX = _bpp_build_year_suffixes(ha.REFS)


def _bpp_has_initials(token: str) -> bool:
    """Return True if the token is an individual author of the form Surname INITIALS."""
    token = token.strip()
    if "et al" in token.lower():
        return False
    return bool(re.fullmatch(r".+?\s+[A-Z]{1,2}", token))


def _bpp_split_authors(author_part: str):
    """Split an author string into individual author tokens.

    Organisational/group authors without trailing initials are kept as a single
    token even if their name contains commas.
    """
    raw = [t.strip() for t in re.split(r",\s*", author_part) if t.strip()]
    et_al = any("et al" in t.lower() for t in raw)
    author_like = sum(1 for t in raw if _bpp_has_initials(t))
    if et_al or author_like >= 2:
        return raw
    return [author_part.strip()]


def _bpp_parse_author_token(token: str) -> str:
    """Format a single author token as 'Surname, I.I.' or return it unchanged."""
    token = token.strip()
    if "et al" in token.lower():
        return "et al."
    m = re.fullmatch(r"(.+?)\s+([A-Z]{1,2})", token)
    if m:
        surname = m.group(1).strip()
        initials = m.group(2).strip()
        dotted = ".".join(initials) + "."
        return f"{surname}, {dotted}"
    return token


def _bpp_format_author_list(author_part: str) -> str:
    """Format an author string for the reference list in Harvard style."""
    tokens = _bpp_split_authors(author_part)
    parsed = [_bpp_parse_author_token(t) for t in tokens if t.strip()]
    if not parsed:
        return author_part.strip()
    if len(parsed) == 1:
        return parsed[0]
    # If the last token is 'et al.', list all preceding names then 'et al.'
    if parsed[-1].lower().startswith("et al"):
        return ", ".join(parsed)
    if len(parsed) == 2:
        return f"{parsed[0]} and {parsed[1]}"
    return ", ".join(parsed[:-1]) + f" and {parsed[-1]}"


def _bpp_in_text_author(ref: str, key: str) -> str:
    """Return the in-text author string for a Harvard-style citation."""
    if key in _BPP_CITE_SHORT:
        return _BPP_CITE_SHORT[key]
    first = ref.split(". ", 1)[0]
    tokens = _bpp_split_authors(first)
    et_al = any("et al" in t.lower() for t in tokens)
    parsed = [_bpp_parse_author_token(t) for t in tokens]
    surnames = [p.split(",")[0].strip() for p in parsed if p]
    # Remove any unhandled 'et al.' token and use the first real surname
    clean_surnames = [s for s in surnames if not s.lower().startswith("et al")]
    first_surname = clean_surnames[0] if clean_surnames else surnames[0]
    if et_al or len(clean_surnames) >= 4:
        return f"{first_surname} et al."
    if len(clean_surnames) == 3:
        return f"{clean_surnames[0]}, {clean_surnames[1]} and {clean_surnames[2]}"
    if len(clean_surnames) == 2:
        return f"{clean_surnames[0]} and {clean_surnames[1]}"
    return first_surname


def _bpp_ensure_title_end(title: str) -> str:
    """Ensure a book/report title ends with terminal punctuation before the publisher tail."""
    title = title.strip()
    if title and title[-1] not in ".?!":
        return title + "."
    return title


def _bpp_article_title(title: str) -> str:
    """Return a sentence-case article title with no trailing terminal punctuation."""
    title = _bpp_sentence_case(title.strip()).strip()
    return title.rstrip(".")


def _bpp_sentence_case(title: str) -> str:
    """Convert an article title to sentence case for BPP Harvard references.

    Only the first word and the first word after a full sentence boundary
    (. ? !) are capitalised; common proper nouns and acronyms are protected.
    """
    proper = {
        "Japan", "Japanese", "Korea", "Korean", "Taiwan", "Taiwanese",
        "United", "States", "Texas", "STROBE", "PLoS", "Defensive", "Medicine",
        "Nationwide", "Obstetric", "Compensation", "System", "Cerebral", "Palsy",
    }
    protected = {
        "United States": "@@US@@",
        "New Zealand": "@@NZ@@",
        '"Defensive Medicine"': "@@DEFMED@@",
        "Defensive Medicine": "@@DEFMED@@",
        "Nationwide Obstetric Compensation System": "@@NOCS@@",
    }
    for phrase, placeholder in protected.items():
        title = title.replace(phrase, placeholder)
    words = title.split(" ")
    new = []
    start = True
    for w in words:
        core = re.sub(r"[^\w\-\']+$", "", w)
        if start or core in proper:
            new.append(w)
        else:
            if w and w[0].isupper() and not w.isupper():
                new.append(w[0].lower() + w[1:])
            else:
                new.append(w)
        start = bool(re.search(r"[.?!:]$", w.rstrip("\"'")))
    text = " ".join(new)
    for phrase, placeholder in protected.items():
        text = text.replace(placeholder, phrase)
    return text


def _bpp_fmt_journal_tail(tail: str) -> str:
    """Format a Vancouver year;volume(issue):pages tail for BPP Harvard style.

    BPP expects: volume(issue): pages (no space before issue, colon before pages).
    """
    parts = re.split(r"(?i)\s*doi[\s:]", tail, maxsplit=1)
    core = parts[0].strip().rstrip(".")
    doi = parts[1].strip().rstrip(".,;") if len(parts) > 1 else ""
    # Drop the year prefix (year already appears before the title)
    m = re.search(r";\s*(.*)$", core)
    vol_pages = m.group(1) if m else core
    if ":" in vol_pages:
        vol, pages = vol_pages.split(":", 1)
    else:
        vol, pages = vol_pages, ""
    vol = vol.strip()
    pages = pages.strip()
    if not vol:
        return tail
    # Use en dash for page ranges
    pages = pages.replace("-", "\u2013")
    out = f"{vol}"
    if pages:
        out += f": {pages}"
    out += "."
    if doi:
        out += f" doi:{doi}"
    return out


def _bpp_fmt_book_tail(tail: str, year: str) -> str:
    """Format a report/book tail (place:publisher;year url) for BPP Harvard style."""
    # Extract URL / accessed components
    url_match = re.search(r"(?i)Available (?:from|at):\s*(\S+)", tail)
    access_match = re.search(r"\(accessed ([^)]+)\)", tail, re.IGNORECASE)
    url = url_match.group(1) if url_match else ""
    accessed = access_match.group(1) if access_match else ""
    # Strip URL/access/year from the remaining tail
    core = re.sub(r"(?i)\s*Available (?:from|at):\s*\S+", "", tail)
    core = re.sub(r"\(accessed [^)]+\)", "", core, flags=re.IGNORECASE)
    core = re.sub(r"[;,]?\s*" + re.escape(year) + r"\b", "", core)
    core = core.strip().strip(".; ").strip()
    if not core:
        if accessed:
            return f"(Accessed {accessed})."
        if url:
            return f"Available at: {url}."
        return ""
    # BPP published style uses "Place: Publisher."
    if ":" in core:
        place, publisher = core.split(":", 1)
        place = place.strip().strip(".; ")
        publisher = publisher.strip().strip(".; ")
        out = f"{place}: {publisher}."
    else:
        out = f"{core}."
    if url:
        if accessed:
            out += f" Available at: {url} (Accessed {accessed})."
        else:
            out += f" Available at: {url}."
    elif accessed:
        out += f" (Accessed {accessed})."
    return out


def _bpp_harvard_segments(ref: str, key: str):
    """Convert a Vancouver-style reference string to BPP Harvard-style text runs.

    Returns a list of (text, italic, bold) tuples so that journal/book titles
    can be italicised and article titles can be quoted while remaining plain text.
    """
    # Manual overrides for sources whose original string does not parse cleanly
    if key == "mais":
        text = (
            "Medical Accident Investigation System (2015), Act on the Promotion of "
            "Medical Safety. Tokyo: Ministry of Health, Labour and Welfare."
        )
        return [(text, False, False)]
    if key == "nikkei":
        text = (
            "Nikkei Inc (2024), Nikkei Telecom 21. Tokyo: Nikkei Inc. "
            "Available at: https://telecom.nikkei.co.jp/ (Accessed 2024)."
        )
        return [(text, False, False)]
    # Use the full subtitle for the main econometrics text
    if key == "angrist":
        author = _bpp_format_author_list("Angrist JD, Pischke JS")
        return [
            (f"{author} (2009), ", False, False),
            ("Mostly Harmless Econometrics: An Empiricist's Companion", True, False),
            (". Princeton: Princeton University Press.", False, False),
        ]
    # In-text citations use institutional abbreviations after the first mention, so the
    # reference list uses the full organisational name.
    if key == "jocscp":
        return [
            ("Japan Council for Quality Health Care (JOCS-CP) (2009), ", False, False),
            ("Japan Obstetric Compensation System for Cerebral Palsy", True, False),
            (". Tokyo: Japan Council for Quality Health Care.", False, False),
        ]
    if key == "jmsr_data":
        return [
            ("Japan Medical Safety Research Organisation (JMSR) (2025), ", False, False),
            ("Annual reports of medical accident investigations (2015-2024)", True, False),
            (". Tokyo: JMSR.", False, False),
        ]
    base_year = _bpp_get_year(ref)
    year = base_year + _BPP_YEAR_SUFFIX.get(key, "")
    # Split on '. ' but not within a doi or a URL/access line; this separates
    # authors, title and tail.
    parts = re.split(r"\. (?!doi|Available|Accessed|\(accessed)", ref, flags=re.IGNORECASE)
    author_part = parts[0].strip()
    author_str = _bpp_format_author_list(author_part)
    prefix = f"{author_str} ({year})"
    if len(parts) == 1:
        return [(f"{prefix}.", False, False)]
    if len(parts) == 2:
        title = _bpp_ensure_title_end(parts[1].strip())
        return [(f"{prefix}, ", False, False), (title, True, False)]
    # General case: title may span multiple parts, journal is second-to-last
    title_parts = parts[1:-2] if len(parts) > 3 else [parts[1]]
    raw_title = ". ".join(title_parts).strip()
    journal = parts[-2].strip()
    tail = parts[-1].strip()
    # Journal articles have a year;volume tail (Vancouver style: year;vol:pages)
    if re.search(r"^\d{4};", tail):
        title = _bpp_article_title(raw_title)
        journal_tail = _bpp_fmt_journal_tail(tail)
        return [
            (f"{prefix}, '", False, False),
            (title, False, False),
            ("', ", False, False),
            (journal, True, False),
            (f". {journal_tail}", False, False),
        ]
    # Reports, books and web pages
    title = _bpp_ensure_title_end(_bpp_sentence_case(raw_title))
    rest = _bpp_fmt_book_tail(tail, base_year)
    return [
        (f"{prefix}, ", False, False),
        (title, True, False),
        (f" {rest}", False, False),
    ]


def h(doc, text, level=1):
    """Unnumbered BPP-style heading."""
    return ha.head(doc, text, level=level, numbered=False)


def b(doc, text, **kw):
    """Body paragraph that contributes to the main word count."""
    return ha.body(doc, text, **kw)


def p(doc, text, **kw):
    """Paragraph that does not contribute to the main word count."""
    return ha.para(doc, text, **kw)


def f(doc, fn, caption, width=Inches(5.8)):
    return ha.figure(doc, fn, caption, width=width)


def t(doc, headers, rows, caption):
    return ha.table(doc, headers, rows, caption)


def m(doc, latex, inline=False, para=None):
    return ha.add_math(doc, latex, inline=inline, para=para)


def build_manuscript(inline=False):
    # Reset citation order and body-text accumulator shared with ha module
    ha._CITE_ORDER.clear()
    ha.BODY_TEXTS.clear()

    doc = ha._setup_doc()

    # Figures and tables are collected at their first in-text mention and then
    # placed at the end of the manuscript, after the reference list, following
    # Cambridge/BPP submission guidance (figures and tables should not be
    # embedded in the body; separate PNG/PPTX files are also supplied).
    end_objects = []

    def f_bpp(doc, fn, caption, width=Inches(5.8)):
        end_objects.append(("fig", fn, caption, width))

    def t_bpp(doc, headers, rows, caption):
        end_objects.append(("table", headers, rows, caption))

    # Harvard-style citation state for BPP
    _cite_order = []
    _first_cite = set()

    def _cite(text: str) -> str:
        def _repl(match):
            keys = [k.strip() for k in match.group(1).split(",") if k.strip()]
            entries = []
            for k in keys:
                if k not in ha.REFS:
                    raise KeyError(f"Unknown citation key: {k}")
                if k not in _cite_order:
                    _cite_order.append(k)
                ref = ha.REFS[k]
                if k in _BPP_CITE_FIRST and k not in _first_cite:
                    author = _BPP_CITE_FIRST[k]
                else:
                    author = _bpp_in_text_author(ref, k)
                _first_cite.add(k)
                year = _bpp_get_cite_year(ref, k)
                entries.append(f"{author}, {year}")
            return " (" + "; ".join(entries) + ")"

        return _CITE_RE.sub(_repl, text)

    def b_bpp(doc, text, **kw):
        return ha.body(doc, _cite(text), **kw)

    def p_bpp(doc, text, **kw):
        return ha.para(doc, _cite(text), **kw)

    global f, t, b, p
    if inline:
        f = lambda doc, fn, caption, width=Inches(5.8): ha.figure(doc, fn, caption, width=width)
        t = lambda doc, headers, rows, caption: ha.table(doc, headers, rows, caption)
    else:
        f = f_bpp
        t = t_bpp
    b = b_bpp
    p = p_bpp

    # Title
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(18)
    rt = tp.add_run(BPP_TITLE)
    rt.bold = True
    rt.font.size = Pt(14)
    rt.font.name = "Times New Roman"

    # Abstract (unstructured, one paragraph; BPP limit 200 words; target ~150)
    abstract_text = (
        f"Specialty maldistribution is a persistent workforce problem, and malpractice litigation is "
        f"often assumed to push physicians away from high-risk fields. Using national administrative data "
        f"for {N} clinical specialties in Japan ({BIEN[0]}-{BIEN[-1]}), we tested whether litigation risk "
        f"predicts physician supply and hospital facility counts. Exposure was malpractice "
        f"claims per {PER:,} physicians; we regressed biennial log-changes in physician and facility counts "
        f"on the lagged litigation rate using specialty and wave fixed effects, clustered "
        f"standard errors, and equivalence tests. Litigation rate was unrelated to physician growth "
        f"(coefficient {ha.fmt(PHYS['coef'], 4)}; 95% CI {ha.fmt(PHYS['ci_low'], 4)} to "
        f"{ha.fmt(PHYS['ci_high'], 4)}; p={PHYS['p']:.2f}) or facility-count growth (p={HOSP['p']:.2f}), "
        f"and a one-SD higher rate shifted physician growth by less than +/-{MARGIN1}% "
        f"(TOST p={ha.p_tost_fmt(EQP['tests'][0]['p_tost'])}). Although perceived risk is high, structural "
        f"incentives -- training costs, fee-for-service income, and status-quo bias -- keep "
        f"physicians in high-risk fields. Reducing litigation exposure is unlikely to correct "
        f"specialty maldistribution; no-fault compensation and payment design are more promising."
    )
    abstract_wc = ha.wc(abstract_text)
    if abstract_wc > 200:
        raise SystemExit(f"Abstract is {abstract_wc} words; BPP limit is 200")

    h(doc, "Abstract", level=1)
    ap = doc.add_paragraph()
    ar = ap.add_run(abstract_text)
    ar.font.name = "Times New Roman"
    ar.font.size = Pt(12)
    ap.paragraph_format.line_spacing = 2.0
    ap.paragraph_format.space_after = Pt(6)

    kw = doc.add_paragraph()
    kr = kw.add_run("Keywords: ")
    kr.bold = True
    kr.font.name = "Times New Roman"
    ha._add_runs(
        kw,
        "behavioural public policy; malpractice litigation; physician workforce; "
        "risk perception; equivalence testing; structural incentives",
    )
    kw.paragraph_format.space_after = Pt(18)

    # Introduction
    h(doc, "Introduction", level=1)
    b(
        doc,
        "Specialty maldistribution is a persistent healthcare workforce problem: high-acuity "
        "fields such as surgery, obstetrics and gynaecology, paediatrics and anaesthesiology are "
        "widely perceived as understaffed across many health systems.{maldist} A recurring policy "
        "intuition is that the fear of malpractice litigation and medical safety incidents pushes "
        "physicians away from these high-risk specialties.{malprac,defmed} If this behavioural "
        "assumption were true, reducing litigation exposure would be an instrument for correcting "
        "workforce maldistribution.",
    )
    b(
        doc,
        "Behavioural economics suggests, however, that the link between perceived risk and actual "
        "behaviour is not automatic. Rare, salient adverse events are highly available in memory "
        "and in media coverage, and loss aversion causes a low-probability outcome to be "
        "overweighted in career deliberations.{tversky1973,kahneman1979} At the same time, the "
        "decision to enter or leave a specialty is constrained by expected income, sunk training "
        "costs, switching costs and status-quo bias.{samuelson1988} The relevant policy question is "
        "not whether physicians report anxiety about litigation, but whether that anxiety "
        "translates into aggregate workforce shifts. We treat the litigation-workforce question as a "
        "test case for evaluating a behavioural public policy instrument: can a salient, cognitively "
        "available risk be used to change the specialty distribution of physicians?",
    )
    b(
        doc,
        "Japan provides a well-documented national setting in which to examine this question. "
        "The Supreme Court reports closed malpractice claims by specialty, the Ministry of Health, "
        "Labour and Welfare publishes biennial physician counts and annual hospital facility counts, and the "
        "country shares the fee-for-service pressures seen in other high-income countries.{court,phys,facil} "
        "It also introduced the Japan Obstetric Compensation System for Cerebral Palsy (JOCS-CP) "
        "in 2009, a no-fault scheme intended partly to address obstetric workforce concerns.{jocscp,hasegawa2016} "
        "Surveys of residents in Japan, Korea and Taiwan report litigation as a negative factor, but "
        "workload, lifestyle and professional interest dominate specialty choice.{lin2022}",
    )
    b(
        doc,
        "Prior empirical work in this area often related raw annual counts of incidents or lawsuits "
        "to raw counts of physicians or facilities. Two features make such designs prone to "
        "spurious association. First, counts are not adjusted for specialty size: a larger "
        "specialty mechanically accumulates more procedures, more claims and more physicians, so "
        "count-based associations can arise without any behavioural mechanism. Second, the "
        "physician census is collected only biennially; interpolating it to an annual series and "
        "analysing it as if each year were an independent observation inflates the degrees of "
        "freedom of any lag-based method. These pitfalls are not unique to malpractice research; "
        "they arise whenever administrative counts are used to infer behavioural responses in "
        "healthcare organisations. We address them with a transparent analytical framework that uses "
        "rates rather than counts, uses only measured biennial physician observations, and applies "
        "pre-specified equivalence testing, which can provide positive evidence for the absence of "
        "a meaningful effect rather than merely failing to reject a null.{lakens,schuir}",
    )

    # Methods
    h(doc, "Methods", level=1)
    h(doc, "Data sources", level=2)
    b(
        doc,
        f"We report this observational study following the Strengthening the Reporting of "
        f"Observational Studies in Epidemiology (STROBE) guidance.{{strobe}} We studied {N} core "
        f"clinical specialties in Japan for which the Supreme Court reports specialty-specific "
        f"litigation. Three official primary series drove the main analysis: physician counts by "
        f"specialty from the biennial Statistics of Physicians, Dentists and Pharmacists{{phys}}; "
        f"closed malpractice claims by specialty from the Supreme Court of Japan{{court}}; and hospital "
        f"counts by specialty from the annual Survey of Medical Institutions.{{facil}} Hospital counts are the number of facilities reporting each specialty, not beds or staffing. Two sensitivity "
        f"series were also used: annual medical accident investigation reports by specialty (2015-2024) "
        f"from the Japan Medical Safety Research Organisation{{jmsr_data}} and total national "
        f"newspaper article counts from Nikkei Telecom 21 (2004-2018; the sensitivity analysis uses "
        f"{MEDIA_START}-{MEDIA_END}; keywords: medical error + medical malpractice).{{nikkei}} The full "
        f"extraction pipeline (with source identifiers and SHA-256 checksums) is documented in the "
        f"accompanying repository ({ANON_REPO}).",
    )
    b(
        doc,
        "One source disagreement is documented in the extraction provenance. The 2017 internal medicine "
        "count is 179 in the consolidated Supreme Court releases (2022/2024/2025) and 181 in the "
        "contemporaneous 2017 committee minutes; we retained 179 in the primary series because it is the "
        "figure in the released statistics.{court} The disagreement does not alter the rate-based "
        "exposure materially, but it is recorded for transparency.",
    )
    b(
        doc,
        "Physician counts use the principal-specialty classification; broad categories were matched "
        "to the Supreme Court's specialty labels, and subspecialties were aggregated in code. "
        "Because the Court assigns multi-specialty cases to a single principal specialty and states "
        "that the counts do not represent the intrinsic risk of each specialty, we treat litigation "
        "as an exposure signal rather than a measure of incident risk.{court} We distinguish "
        "litigation from the Medical Accident Investigation System, which began in 2015 and covers "
        "only deaths and stillbirths judged unforeseen by the hospital administrator; it is not a "
        "general incident-reporting system and is not used as an exposure here.{mais} Primary data "
        "sources and their resolution are summarised in Supplementary Table 1.",
    )
    b(
        doc,
        "The 12 specialties are not arbitrary selections but represent the core clinical categories "
        "in which Japanese physicians typically obtain initial board certification. Japan's specialist "
        "training programme is a two-tiered system: physicians first complete residency and gain "
        "certification in one of the core specialties (e.g., internal medicine, surgery, obstetrics and "
        "gynaecology, paediatrics, anaesthesiology), after which some proceed to a narrower subspecialty "
        "within that primary field. Biennial physician counts and Supreme Court litigation statistics are "
        "both reported at this primary-specialty level, so our analysis captures the broad workforce-allocation "
        "decision at the first tier of the training pipeline. Subspecialties are not separately tracked in "
        "the official biennial series, so we analyse the 12 primary fields as the relevant units of "
        "specialty choice and policy intervention.",
    )

    h(doc, "Statistical analysis", level=2)
    b(
        doc,
        f"We used a transparent sensitivity-analysis framework that varies the exposure definition "
        f"(counts versus rates), panel frequency (measured biennial waves versus interpolated annual values), "
        f"and potential confounders while holding the specialty-level panel structure constant. Supplementary "
        f"Figure 1 summarises the framework. The exposure was the litigation rate, defined as closed claims "
        f"per {PER:,} physicians in each specialty-year, which removes specialty-size confounding because large "
        f"specialties generate more claims for reasons unrelated to per-physician risk. The primary analysis "
        f"used the {len(BIEN)} measured biennial physician waves ({BIEN[0]}-{BIEN[-1]}). For each specialty "
        f"we computed the biennial log-change in physicians (and, separately, in hospital facility counts) and regressed it on "
        f"the lagged litigation rate at the start of the interval, in a panel with specialty and wave fixed effects, "
        f"clustered standard errors, and a JOCS-CP indicator for obstetrics and gynaecology from 2009 onward.{{angrist}} "
        f"Clusters are defined by specialty, so G={N} and the small-cluster correction uses a t-distribution with "
        f"G-1 degrees of freedom.{{cameron2015}} The primary estimating equation, for specialty s and wave t, was as follows.",
    )
    m(doc, r"\Delta \log(Y_{st}) = \alpha_s + \delta_t + \beta \cdot \text{litrate}_{s,t-1} + \gamma \cdot \text{JOCS-CP}_{s,t-1} + \epsilon_{st}")
    b(
        doc,
        "Here, Y is either physician counts or hospital facility counts, alpha_s are specialty fixed effects, "
        "delta_t are wave fixed effects, and standard errors are clustered by specialty. "
        "For the equivalence analysis we standardised litrate to a z-score, so beta gives the "
        "expected biennial log-change per one-SD increase in the litigation rate.",
    )
    b(
        doc,
        f"We assessed equivalence to a null effect using two one-sided tests (TOST) with pre-specified margins "
        f"of +/-{MARGIN1}% and +/-{MARGIN2}% biennial workforce change. The margins were chosen to be smaller than "
        f"the policy-relevant workforce shifts that structural tools have produced: for example, targeted residency "
        f"subsidies raised primary-care physician supply by about 4% in a comparable setting.{{mcnamara2025}} That 4% is a "
        f"total programme effect, whereas our +/-{MARGIN1}% and +/-{MARGIN2}% margins are biennial, so the margins are "
        f"substantially smaller than the policy-relevant shift and therefore conservative. The TOST procedure itself is "
        f"standard for equivalence testing.{{lakens,schuir}} Equivalence is declared when the per-SD coefficient lies inside the margin. Inference "
        f"was complemented by a cluster block-bootstrap (B = 1,999) and by the minimum detectable effect (MDE) at 80% "
        f"power and the power to declare equivalence when the true effect is zero. Full formulas and the heterogeneity "
        f"and time-trend sensitivity checks are given in Supplementary Note 1.",
    )

    h(doc, "Policy simulation", level=2)
    b(
        doc,
        "To translate the regression results into a practical policy benchmark, we projected "
        "physician counts to 2034 under three stylised policy scenarios. The baseline used each "
        "specialty's observed mean biennial log-growth from 2004 to 2024. The two litigation-reduction "
        "scenarios set the litigation rate to zero and applied the point estimate and the 95% lower-bound "
        "coefficient, respectively, so they show both the central projection and the most favourable "
        "effect consistent with the data. The third scenario added the minimum detectable per-SD effect "
        f"({EQP['mde_80pct']:.2f}% per biennium) to baseline growth as a benchmark for the smallest "
        "policy effect this panel could detect with 80% power. The projections are deterministic "
        "counterfactuals, not forecasts, and are reported as the marginal percent change in 2034 "
        "relative to the baseline drift.",
    )
    b(
        doc,
        f"All confidence intervals and two-sided p-values use the small-cluster correction with G-1 = "
        f"{PHYS['df']} degrees of freedom.{{cameron2015}} The JOCS-CP indicator captures the post-2009 "
        f"obstetrics-and-gynaecology period.{{jocscp}} Sensitivity analyses repeat the models on the annual "
        f"hospital series, an interpolated annual-physician series, raw counts instead of rates, and annual "
        f"hospital series controlling for the JMSR report rate ({JMSR_START}-2024) or total Nikkei Telecom "
        f"article counts ({MEDIA_START}-{MEDIA_END}); details are in Supplementary Table 1. The JOCS-CP "
        f"indicator and all sensitivity models are exploratory; we report raw and Holm step-down adjusted "
        f"p-values. Code and data are available in the accompanying repository ({ANON_REPO}).",
    )

    # Results
    h(doc, "Results", level=1)
    h(doc, "Workforce and litigation trends", level=2)
    b(
        doc,
        f"Litigation rates per {PER:,} physicians varied several-fold across specialties and fell over time "
        f"in {FELL} of {N} fields (Supplementary Figure 2). Over the same period the physician workforce "
        f"grew in {GREW} of {N} specialties (Supplementary Figure 3; Table 1); the only exception was "
        f"general surgery, which was nearly flat ({SURG_PCT:+.1f}% across {SPAN} years). Exposure and "
        f"workforce did not move in opposite directions as a flight-from-risk account would predict.",
    )
    rows = []
    for s in ha.CORE:
        v = DESCR[ha.EN[s]]
        rows.append(
            [
                ha.EN[s],
                v["phys_first"],
                v["phys_last"],
                f"{v['litrate_first']:.2f}",
                f"{v['litrate_last']:.2f}",
                v["hosp_first"],
                v["hosp_last"],
            ]
        )
    t(
        doc,
        [
            "Specialty",
            f"Physicians {BIEN[0]}",
            f"Physicians {BIEN[-1]}",
            f"Lit. rate {BIEN[0]}",
            f"Lit. rate {BIEN[-1]}",
            f"Hospital facilities {BIEN[0]}",
            f"Hospital facilities {BIEN[-1]}",
        ],
        rows,
        f"Table 1. Physicians, litigation rate (per {PER:,} physicians) and hospital facilities by specialty, first and last waves.",
    )

    h(doc, "Primary association and equivalence", level=2)
    b(
        doc,
        f"The lagged litigation rate was not associated with biennial physician growth "
        f"(coefficient {ha.fmt(PHYS['coef'], 4)}; 95% CI {ha.fmt(PHYS['ci_low'], 4)} to "
        f"{ha.fmt(PHYS['ci_high'], 4)}; p={PHYS['p']:.2f}; n={PHYS['n_obs']}) or with hospital facility-count growth "
        f"(coefficient {ha.fmt(HOSP['coef'], 4)}; p={HOSP['p']:.2f}). Equivalence testing (Figure 1; Table 2) "
        f"showed that a 1-SD higher litigation rate changed biennial physician growth by less than "
        f"+/-{MARGIN1}% (TOST p={ha.p_tost_fmt(EQP['tests'][0]['p_tost'])}; point estimate "
        f"{EQP['coef_per_SD']*100:+.2f}% with 90% CI {EQP['ci90_low']*100:+.2f}% to "
        f"{EQP['ci90_high']*100:+.2f}%). For hospital facility-count growth the point estimate was "
        f"{EQH['coef_per_SD']*100:+.2f}% (90% CI {EQH['ci90_low']*100:+.2f}% to "
        f"{EQH['ci90_high']*100:+.2f}%): it was within the +/-{MARGIN2}% margin "
        f"(p={ha.p_tost_fmt(EQH['tests'][1]['p_tost'])}) but not the stricter +/-{MARGIN1}% margin "
        f"(p={ha.p_tost_fmt(EQH['tests'][0]['p_tost'])}). Thus the data are consistent with the absence "
        f"of a policy-relevant effect on physician growth, and with at most a small effect on hospital facility-count growth. "
        f"Detailed TOST results by margin are reported in Supplementary Table 2.",
    )
    f(
        doc,
        "ha_Figure_1.png",
        f"Figure 1. Equivalence (TOST) of the litigation-rate effect against +/-{MARGIN1}% and "
        f"+/-{MARGIN2}% margins; horizontal bars are 90% confidence intervals.",
    )
    trow = [
        [
            "Physician growth ~ lagged rate",
            ha.fmt(PHYS["coef"], 4),
            f"{ha.fmt(PHYS['ci_low'], 4)}, {ha.fmt(PHYS['ci_high'], 4)}",
            f"{PHYS['p']:.2f}",
            PHYS["n_obs"],
        ],
        [
            "Hospital facility-count growth ~ lagged rate",
            ha.fmt(HOSP["coef"], 4),
            f"{ha.fmt(HOSP['ci_low'], 4)}, {ha.fmt(HOSP['ci_high'], 4)}",
            f"{HOSP['p']:.2f}",
            HOSP["n_obs"],
        ],
        [
            "Counts contrast (physician)",
            ha.fmt(CNT["coef"], 4),
            " -- ",
            f"{CNT['p']:.2f}",
            CNT["n_obs"],
        ],
        [
            "Annual hospital facility-count growth (sensitivity)",
            ha.fmt(ANN["coef"], 4),
            " -- ",
            f"{ANN['p']:.2f}",
            ANN["n_obs"],
        ],
        [
            "Interpolated physician (sensitivity)",
            ha.fmt(INT["coef"], 4),
            " -- ",
            f"{INT['p']:.2f}",
            INT["n_obs"],
        ],
        [
            "Reverse (workforce->litigation)",
            ha.fmt(REV["coef"], 3),
            " -- ",
            f"{REV['p']:.2f}",
            REV["n_obs"],
        ],
    ]
    t(
        doc,
        ["Model", "Coefficient", "95% CI", "p", "n"],
        trow,
        "Table 2. Panel fixed-effects models and sensitivity analyses.",
    )

    h(doc, "Counts versus rates, and confounders", level=2)
    b(
        doc,
        f"Using raw litigation counts rather than rates did not reveal a negative association in this "
        f"measured-only design (p={CNT['p']:.2f}). Figure 2 shows the contrast: the count exposure is "
        f"confounded by specialty size (panel a), whereas the rate-adjusted exposure is not (panel b); "
        f"points are coloured and shaped by specialty so readers can identify which fields drive any "
        f"apparent pattern. The annual hospital and interpolated annual-physician sensitivity analyses were "
        f"also null (p={ANN['p']:.2f} and p={INT['p']:.2f}), confirming that the null result holds across panel "
        f"frequency and exposure definition. The JOCS-CP indicator was positive in sign in the obstetric-hospital "
        f"model (coefficient {ha.fmt(HOSP['jocscp_coef'], 3)}, raw p={HOSP['jocscp_p']:.3f}), but it did not "
        f"remain significant after the small-cluster correction and Holm adjustment for the exploratory "
        f"sensitivity family (Holm p={JOCS_HOLM:.3f}); we treat it as exploratory and do not "
        f"interpret it as a causal policy effect.",
    )
    f(
        doc,
        "ha_Figure_2.png",
        "Figure 2. Biennial physician growth against lagged litigation exposure measured as "
        "(a) counts and (b) rates. Points are coloured by specialty; the count panel shows the size "
        "confounding that the rate panel removes.",
    )
    b(
        doc,
        "The same count-versus-rate contrast for biennial hospital facility-count growth is shown in Figure 3. As "
        "with physician growth, the count exposure creates a spurious size confound that disappears once "
        "the rate-adjusted exposure is used.",
    )
    f(
        doc,
        "ha_Figure_3.png",
        "Figure 3. Biennial hospital facility-count growth against lagged litigation exposure measured as "
        "(a) counts and (b) rates. Points are coloured by specialty; the rate-adjusted panel shows "
        "no systematic association.",
    )
    b(
        doc,
        f"Descriptively, per-specialty rank correlations between the lagged litigation rate and physician "
        f"growth were positive in {n_pos} of {N} specialties and statistically significant in {n_sig}; the "
        f"direction is, if anything, opposite to a flight-from-risk hypothesis.",
    )
    b(
        doc,
        f"A reverse specification (change in litigation rate regressed on lagged log physicians) was also "
        f"null (coefficient {ha.fmt(REV['coef'], 3)}, p={REV['p']:.2f}; Table 2), making a reverse-causation "
        f"interpretation of the null unlikely.",
    )
    b(
        doc,
        f"We also evaluated JMSR medical-accident investigation report counts as a potential confounder or "
        f"competing exposure.{{jmsr_data}} From {ha.JMSR_CORR['years'][0]} to {ha.JMSR_CORR['years'][-1]}, raw "
        f"litigation and JMSR report counts were strongly correlated across specialties (Pearson "
        f"r={ha.JMSR_CORR['pooled_r']:.2f}), because large specialties generate more of both. After removing "
        f"specialty-specific levels and trends, however, the within-specialty correlation was negligible "
        f"(r={ha.JMSR_CORR['detrended_r']:.2f}). A model of annual hospital facility-count growth for {JMSR_START}-2024 that "
        f"included both the lagged litigation rate and the lagged JMSR report rate left the litigation "
        f"coefficient little changed ({ha.fmt(JMSR['lit_coef'], 4)}; p={JMSR['lit_p']:.2f}) and the JMSR "
        f"term was not associated with hospital facility-count growth (p={JMSR['med_p']:.2f}; Supplementary Table 3). The null "
        f"litigation result is neither explained nor masked by broader medical-accident reporting.",
    )
    b(
        doc,
        f"Finally, we tested national newspaper coverage from Nikkei Telecom 21 as a potential confounder."
        f"{{nikkei}} Total annual article counts (keywords: medical error + medical malpractice) and total "
        f"litigation counts were correlated (Pearson r={ha.MEDIA_CORR['total_r']:.2f}), consistent with "
        f"greater public attention in high-litigation years. Within the annual hospital panel, however, the "
        f"lagged litigation rate and the media-count series were only weakly correlated. A model of annual "
        f"hospital facility-count growth for {MEDIA_START}-{MEDIA_END} that included both the lagged litigation rate and the "
        f"lagged article count (per 1,000 articles) left the litigation coefficient little changed and "
        f"the media term was not associated with hospital facility-count growth (p={MEDIA['media_p']:.2f}; Supplementary "
        f"Table 4). Media coverage does not explain the null litigation effect either. Holm step-down "
        f"adjusted p-values for the exploratory sensitivity family are reported in Supplementary Table 5.",
    )

    h(doc, "Small-cluster validation and power", level=2)
    b(
        doc,
        f"Because inference is based on only {N} specialty clusters, we checked the primary results with "
        f"a cluster block-bootstrap (B = 1,999). For physician growth the bootstrap 95% CI for the lagged "
        f"litigation-rate coefficient was {ha.fmt(BS_PHYS['coef_boot_ci_low'], 4)} to "
        f"{ha.fmt(BS_PHYS['coef_boot_ci_high'], 4)} and the bootstrap p-value was {BS_PHYS['p_bootstrap']:.2f}; "
        f"for hospital facility-count growth the bootstrap 95% CI was {ha.fmt(BS_HOSP['coef_boot_ci_low'], 4)} to "
        f"{ha.fmt(BS_HOSP['coef_boot_ci_high'], 4)} and the bootstrap p-value was {BS_HOSP['p_bootstrap']:.2f}. "
        f"Both intervals comfortably contain zero. Power diagnostics make the panel information explicit. "
        f"For physician growth, the minimum detectable effect was {EQP['mde_80pct']:.2f}% per SD at 80% power, "
        f"and the power to declare equivalence within the +/-{MARGIN1}% margin if the true effect were zero "
        f"was {EQP['tests'][0]['power_if_null']*100:.1f}%. For hospital facility-count growth the minimum detectable effect was "
        f"{EQH['mde_80pct']:.2f}% per SD and the equivalent power for the +/-{MARGIN1}% margin was "
        f"{EQH['tests'][0]['power_if_null']*100:.1f}%. The panel is informative enough to rule out "
        f"policy-relevant effects for physicians, and to bound any hospital effect within a small margin.",
    )
    b(
        doc,
        "Full bootstrap diagnostics and power calculations are reported in Supplementary Table 6.",
    )
    h(doc, "Counterfactual simulation", level=2)
    b(
        doc,
        f"The counterfactual projection made the practical implications of the null regression result "
        f"explicit. Under the point estimate, eliminating all malpractice litigation would add only "
        f"{TOTAL_SIM.get('marginal_pct_lit_point', 0):.1f}% to the projected 2034 national physician stock "
        f"relative to baseline drift. Even under the 95% lower-bound (most favourable) coefficient it would add "
        f"{TOTAL_SIM.get('marginal_pct_lit_lower', 0):.1f}%, comparable to the "
        f"{TOTAL_SIM.get('marginal_pct_mde', 0):.1f}% gain from a generic benchmark equal to the minimum detectable "
        f"effect. General surgery, the only specialty with negative baseline drift, illustrates the break-even "
        f"arithmetic: its projected 2024-2034 decline of {SURG_SIM.get('pct_change_baseline', 0):.1f}% would be "
        f"reduced to {SURG_SIM.get('pct_change_lit_zero_point', 0):.1f}% under the point estimate and reversed to "
        f"{SURG_SIM.get('pct_change_lit_zero_lower', 0):.1f}% under the 95% lower bound. The latter requires "
        f"eliminating every remaining closed claim and assumes the most adverse (most negative) coefficient "
        f"compatible with the data; a more realistic policy would achieve far less. Full projected 2034 physician "
        f"counts by specialty and scenario are reported in Supplementary Table 7; Figure 4 summarises the same "
        f"information as marginal percentage changes. Litigation reduction is not a high-impact "
        f"instrument for workforce allocation in this setting.",
    )
    f(
        doc,
        "ha_Figure_4.png",
        "Figure 4. Counterfactual policy simulation: marginal 10-year change in physician counts by "
        "specialty relative to the projected baseline drift. The MDE benchmark is the minimum detectable "
        "per-SD effect from the primary analysis.",
    )
    if SENKOI_SUM is not None:
        h(doc, "Training-system context", level=2)
        b(
            doc,
            f"As a reference for the size of the specialist-training pipeline, first-year specialist-trainee "
            f"(senkoi) positions in 2018 covered {SENKOI_SUM['weighted_val']:.1f}% of physicians in the third to "
            f"fifth year after medical registration in 2014 (weighted by specialty stock), with specialty-specific coverage ranging from "
            f"{SENKOI_SUM['min_val']:.1f}% ({ha.EN[SENKOI_SUM['min_spec']].lower()}) to "
            f"{SENKOI_SUM['max_val']:.1f}% ({ha.EN[SENKOI_SUM['max_spec']].lower()}).{{mhlw_senkoi2018,mhlw_3_5yr}} "
            f"These coverage rates are not an outcome of the litigation-risk model, but they confirm that "
            "the 12 primary specialties capture the main initial specialization decision in Japan; "
            "Supplementary Table 8 reports the counts by specialty.",
        )
    h(doc, "Heterogeneity and trend sensitivity", level=2)
    het_phys_hi = _HET_DICT.get(("dlog_phys", "high litigation"))
    het_phys_surg = _HET_DICT.get(("dlog_phys", "surgical"))
    het_hosp_hi = _HET_DICT.get(("dlog_hosp", "high litigation"))
    het_hosp_surg = _HET_DICT.get(("dlog_hosp", "surgical"))
    trend_phys = _TREND_DICT.get("dlog_phys")
    trend_hosp = _TREND_DICT.get("dlog_hosp")
    b(
        doc,
        "We tested whether the null association concealed a differential response in high-litigation "
        "or surgical specialties, or whether it depended on assuming common wave fixed effects. "
        f"In high-litigation specialties the main litigation coefficient for physician growth was "
        f"{ha.fmt(het_phys_hi['coef'], 4)} (p={het_phys_hi['p']:.2f}) and the interaction was "
        f"{ha.fmt(het_phys_hi['interact_coef'], 4)} (p={het_phys_hi['interact_p']:.2f}); for hospital facility-count growth the "
        f"main coefficient was {ha.fmt(het_hosp_hi['coef'], 4)} (p={het_hosp_hi['p']:.2f}) and the interaction was "
        f"{ha.fmt(het_hosp_hi['interact_coef'], 4)} (p={het_hosp_hi['interact_p']:.2f}). "
        f"For the surgical-specialty interaction the physician main effect was {ha.fmt(het_phys_surg['coef'], 4)} "
        f"(p={het_phys_surg['p']:.2f}) and the interaction was {ha.fmt(het_phys_surg['interact_coef'], 4)} "
        f"(p={het_phys_surg['interact_p']:.2f}); for hospital facility-count growth the main effect was "
        f"{ha.fmt(het_hosp_surg['coef'], 4)} (p={het_hosp_surg['p']:.2f}) and the interaction was "
        f"{ha.fmt(het_hosp_surg['interact_coef'], 4)} (p={het_hosp_surg['interact_p']:.2f}). "
        f"Allowing specialty-specific linear trends also left the litigation coefficient small and non-significant "
        f"for physician growth ({ha.fmt(trend_phys['coef'], 4)}; p={trend_phys['p']:.2f}) and hospital facility-count growth "
        f"({ha.fmt(trend_hosp['coef'], 4)}; p={trend_hosp['p']:.2f}). None of the interactions or trend-stability "
        "checks suggest that a negative litigation effect is hiding in a clinically exposed subgroup (Supplementary Table 9).",
    )


    # Discussion
    h(doc, "Discussion", level=1)
    b(
        doc,
        f"Using national primary data, rates rather than counts, and only measured biennial physician "
        f"observations, we found no association between specialty-level malpractice-litigation risk and "
        f"subsequent physician or hospital decline. Equivalence testing showed that any effect of litigation risk "
        f"on biennial physician growth is smaller than {MARGIN1}% (90% CI within the {MARGIN1}% margin), and any "
        f"effect on hospital facility-count growth is smaller than {MARGIN2}% (but not confidently smaller than {MARGIN1}%). "
        f"These data do not support the hypothesis that physicians systematically abandon "
        f"high-litigation specialties over {SPAN} years of official statistics.",
    )
    b(
        doc,
        "The null result is not merely a failure to detect an effect. The narrow confidence intervals, "
        "pre-specified equivalence margins, and power diagnostics allow us to say that, if litigation risk "
        "does influence specialty-level workforce growth, the magnitude is too small to matter for workforce "
        "planning. For behavioural public policy, this is an important distinction: a widely discussed risk can "
        "be highly available and emotionally salient without being a reliable policy tool. The public policy "
        "question is not whether physicians worry about litigation, but whether a policy that reduces litigation "
        "risk would materially change aggregate specialty supply. Our evidence suggests it would not.",
    )
    h(doc, "Behavioural mechanisms", level=2)
    b(
        doc,
        "Several behavioural mechanisms are consistent with this finding. First, litigation risk may affect "
        "clinical behaviour on the intensive margin (defensive medicine) rather than the extensive margin "
        "(specialty exit). Kessler and McClellan showed that U.S. malpractice reforms reduced medical "
        "expenditures for elderly heart-disease patients without increasing mortality or complications, "
        "suggesting that defensive practice is one margin of adjustment to liability pressure.{kessler1996} "
        "Subsequent reassessments have debated the magnitude and consistency of this effect, but the conceptual "
        "point remains: physicians can respond to liability risk by changing how they practise rather than by "
        "exiting a specialty.{sloan2008} Fee-for-service reimbursement in Japan rewards the high-acuity "
        "procedural work that also carries litigation exposure, so the financial return to remaining in surgery, "
        "obstetrics, or interventional specialties may dominate any deterrent from civil claims.",
    )
    b(
        doc,
        "Second, the discrepancy between perceived risk and measured workforce supply is consistent with "
        "well-documented behavioural-economics mechanisms. Media coverage of sensational malpractice or "
        "criminal prosecutions makes litigation risk highly available to physicians and trainees, and loss aversion "
        "can cause a rare but salient adverse outcome to be overweighted in career deliberations."
        "{tversky1973,kahneman1979} Yet the decision to leave a specialty is governed by expected income, "
        "sunk training costs, switching costs and status-quo bias, all of which discourage exit even when "
        "perceived risk is high.{samuelson1988} The gap between reported anxiety and measured supply is "
        "not a contradiction; it is exactly what one would expect when a vivid, low-probability risk meets strong "
        "economic and institutional incentives to remain.",
    )
    b(
        doc,
        "For behavioural public policy, the lesson is that changing the perceived risk alone is unlikely to alter "
        "aggregate workforce allocation when the underlying choice architecture preserves strong structural "
        "incentives to remain. Status-quo bias and loss aversion now work in the opposite direction: once a "
        "physician has incurred the sunk cost of specialty training, leaving feels like a sure loss, while the rare "
        "possibility of a malpractice judgment remains an abstract, low-probability event. No-fault compensation, "
        "payment design, and training subsidies change the payoff structure directly; litigation-avoidance "
        "messaging targets only the perceived risk. Our results suggest that the former are the more reliable tools.",
    )
    h(doc, "Institutional context and international evidence", level=2)
    b(
        doc,
        "The civil litigation environment in Japan itself dampens the likelihood of a flight-from-risk response. "
        "Taniguchi and colleagues analysed all closed malpractice claims reported by the Supreme Court from 2006 "
        "to 2021 and found that more than half ended in settlement, plaintiffs won only about a quarter of "
        "judgments, and the number of claims has been declining, especially in obstetrics and gynaecology."
        "{taniguchi2023} The Court data we use describe a civil system that is low-volume, "
        "settlement-prone, and comparatively favourable to physicians. This context makes it unlikely that routine "
        "civil litigation risk alone would drive physicians out of high-risk fields.",
    )
    b(
        doc,
        "No-fault obstetric compensation (JOCS-CP, 2009) illustrates a different mechanism. It was introduced "
        "partly because of a shortage of young obstetricians and regional gaps in maternity care, and it combined "
        "no-fault compensation with investigation and prevention.{hasegawa2016} The hospital-level JOCS-CP "
        f"indicator was directionally positive (coefficient {ha.fmt(HOSP['jocscp_coef'], 3)}, raw p={HOSP['jocscp_p']:.3f}), "
        f"but it did not remain significant after the small-cluster correction and Holm adjustment for the "
        f"exploratory sensitivity family (Holm p={JOCS_HOLM:.3f}). This suggests that, if the JOCS-CP did support "
        "obstetric hospital supply, the effect would be too small or too confounded by concurrent obstetric policies "
        "to be isolated here. Civil litigation exposure is also distinct from criminal prosecution. Morita studied the "
        "2004 Fukushima obstetrician prosecution and found a 13 percent decline in obstetricians, with some "
        "switching to gynaecology.{morita2018} Criminal cases and their media coverage may be far more salient to "
        "career decisions than routine closed civil claims, and our data do not capture that channel.",
    )
    b(
        doc,
        "The obstetrics and gynaecology case is the most discussed example of the litigation-workforce nexus, "
        "and it is consistent with our interpretation. A recent Japan-U.S. comparison of medical-legal claims in "
        "obstetrics and gynaecology found that the proportion of malpractice claims in this specialty fell from "
        "15.1 percent in 2004 to 5.2 percent in 2022, and that claims per 100 OB/GYN physicians fell from 0.9 in "
        "2007 to 0.4 in 2016, while maternal and neonatal mortality also declined.{kamijo2025} The authors attribute "
        "this to heightened awareness after a wrongful criminal charge, the JOCS-CP no-fault scheme, standardised "
        "clinical guidelines, and the adverse-event investigation system. This is not evidence that lowering "
        "litigation risk caused the workforce to grow; it is evidence that obstetric litigation, workforce support, "
        "and safety interventions moved together. Surveys of OB/GYN residents in Japan, Korea, and Taiwan likewise "
        "show that litigation is reported as a negative factor, but that its perceived importance is smaller where "
        "no-fault compensation exists and that workload, lifestyle, and professional interest remain dominant."
        "{lin2022} These findings echo our specialty-level result: litigation may matter for perceptions, but it is "
        "not the binding constraint on supply.",
    )
    b(
        doc,
        "International evidence on tort reform and physician supply is consistent with a small or context-specific "
        "effect, especially in systems with low litigation volume and predictable damages.{helland2015,matsa2007,hyman2015,frakes2020} "
        "Against this backdrop, a null effect of civil litigation risk on Japanese specialty supply is not surprising.",
    )
    b(
        doc,
        "The raw-count sensitivity did not reveal a negative association in these data, illustrating that the apparent "
        "count-litigation relationship does not translate into a behavioural effect once specialty size is accounted for "
        "(Figure 2). This is a cautionary example for workforce research that pairs administrative count series, and shows "
        "why rate-based, measured-only designs are preferable when testing litigation-workforce hypotheses.",
    )
    h(doc, "Policy implications", level=2)
    b(
        doc,
        f"Reducing civil malpractice litigation is unlikely to be a powerful tool for correcting specialty maldistribution. "
        f"The 10-year counterfactual showed that eliminating all claims would add only {TOTAL_SIM.get('marginal_pct_lit_point', 0):.1f}% "
        f"to the projected national physician stock under the point estimate, and {TOTAL_SIM.get('marginal_pct_lit_lower', 0):.1f}% "
        f"under the most favourable 95% lower-bound coefficient. A generic benchmark equal to the minimum detectable effect would add "
        f"{TOTAL_SIM.get('marginal_pct_mde', 0):.1f}%. Effects of this size are too small to rebalance a workforce whose distribution "
        f"is shaped by training costs, reimbursement and seniority.",
    )
    b(
        doc,
        "The weak response reflects the institutional structure of career choice, not a lack of concern. Once a physician has entered "
        "a specialty, the costs of switching -- foregone fellowship income, lost seniority and professional identity -- generate "
        "powerful status-quo bias, while fee-for-service income rewards the high-acuity work that also carries litigation risk. "
        "No-fault compensation, payment design and training subsidies alter the payoffs and defaults that matter for a loss-averse "
        "trainee; they are more reliable tools than risk communication.{samuelson1988,kahneman1979} New Zealand and the Nordic "
        "administrative-compensation systems show how separating compensation from blame can protect patients and providers "
        "without relying on litigation, and Japan's JOCS-CP moves in the same direction for obstetric cerebral "
        "palsy.{bismark2006,mello2011,hasegawa2016}",
    )
    b(
        doc,
        "Information campaigns that publicise the rarity of claims are therefore unlikely to alter aggregate supply. The availability "
        "heuristic means that vivid cases dominate decision weights regardless of their frequency, so a national trend line does not "
        "remove salient exemplars from memory or media coverage.{tversky1973,kahneman1979,lin2022} The behavioural lesson is that "
        "policymakers should redesign the choice architecture -- payment, training slots and compensation rules -- rather than expect "
        "litigation-avoidance messaging to drive career decisions. Our null result is evidence that structural defaults and loss aversion "
        "outweigh rare-event risk perception in the aggregate workforce.",
    )
    h(doc, "Limitations", level=2)
    b(
        doc,
        f"This is an ecological, specialty-level analysis and cannot establish individual-level causality. We do not directly observe individual physicians' subjective risk perceptions, so the behavioural interpretation of the null result in terms of availability, loss aversion and status-quo bias is theoretically informed rather than directly tested. The 12 specialties "
        f"correspond to the primary-specialty tier of Japan's two-tiered specialist training programme; the analysis "
        f"describes workforce allocation at the initial board-certification stage and may not extend to narrower subspecialties that "
        f"are not separately tracked in the biennial census. The physician census is biennial, giving {len(BIEN)} measured waves; we addressed the limited power directly through equivalence "
        f"testing and by pooling across specialties, but residual power constraints remain and the equivalence margins are a "
        f"judgement. Litigation rates may be endogenous to physician supply if a smaller workforce increases workload and hence "
        f"incidents; the lagged exposure, fixed effects, and reverse specification make reverse causation unlikely, yet unobserved "
        f"confounders at the specialty or prefecture level cannot be fully ruled out. Cluster block-bootstrap and power diagnostics "
        f"are reported in Supplementary Table 6. Because clusters are defined by the {N} specialties, the small-cluster correction "
        f"uses G-1={PHYS['df']} degrees of freedom; this is the minimum at which cluster-level t inference is recommended and is "
        f"inherent to the data. Specialty-specific litigation counts could be recovered only from {BIEN[0]}; pre-{BIEN[0]} specialty "
        f"tables were not retrievable from primary sources. Clinic counts by specialty are published only every "
        f"{ha.CLINIC_RES} years and were used descriptively. JMSR report counts are available only from {ha.JMSR_CORR['years'][0]} "
        f"and were used in a {JMSR_START}-2024 sensitivity. Media article counts are available only for {MEDIA_START}-{MEDIA_END} and "
        f"are a national total, so they cannot be decomposed by specialty and are collinear with full wave fixed effects. Litigation "
        f"counts are assigned to a principal specialty and, by the Court's own note, do not measure intrinsic specialty risk.{{court}} "
        f"Finally, these findings are embedded in the country's particular legal, cultural and institutional context -- including its "
        f"no-fault obstetric compensation scheme, its fee-for-service reimbursement structure and its comparatively low-volume "
        f"malpractice-litigation culture -- so physician responses to litigation risk may differ in health systems with different "
        f"liability regimes, compensation mechanisms or professional norms; the results should not be assumed to generalise across "
        f"cultural spheres.",
    )

    h(doc, "Conclusions", level=1)
    b(
        doc,
        f"Across {YEARS}, specialty-level malpractice-litigation risk was not associated with physician or hospital decline in "
        f"these national data, and the physician effect was statistically equivalent to null within a small margin. From a "
        f"behavioural public policy perspective, malpractice litigation is not a reliable tool for correcting specialty "
        f"maldistribution in this setting. Policymakers may more productively target structural incentives, especially no-fault "
        f"compensation and payment design, rather than rely on the assumption that reducing litigation will retain physicians in "
        f"high-risk specialties. The transparent, reproducible sensitivity-analysis framework used here is exportable to other "
        f"healthcare workforce-policy tools. More generally, the results caution against treating highly salient risks as "
        f"reliable policy tools when the underlying choice architecture preserves strong structural incentives. Behavioural public "
        f"policy is most effective when it redesigns the decision context rather than attempting to counteract it through "
        f"information alone.",
    )

    # References (Harvard style, alphabetical by first author surname)
    h(doc, "References", level=1)
    _harvard_entries = []
    for k in _cite_order:
        segments = _bpp_harvard_segments(ha.REFS[k], k)
        author_token = segments[0][0].split(" (", 1)[0].split(",")[0].strip().lower()
        year_token = _bpp_get_year(ha.REFS[k]) + _BPP_YEAR_SUFFIX.get(k, "")
        _harvard_entries.append((segments, author_token, year_token))
    _harvard_entries.sort(key=lambda x: (x[1], x[2]))
    for segments, _, _ in _harvard_entries:
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.line_spacing = 1.5
        for text, italic, bold in segments:
            r_ref = p_ref.add_run(text)
            r_ref.font.size = Pt(10)
            r_ref.font.name = "Times New Roman"
            r_ref.italic = italic
            r_ref.bold = bold

    # Figures and tables are placed at the end of the main document, after the
    # reference list, per Cambridge BPP author instructions.  Figures are *not*
    # embedded in the manuscript file; captions are listed here and the editable
    # image files are supplied separately in the submission zip.
    def _add_figure_caption(doc, caption):
        cap = doc.add_paragraph()
        cap.paragraph_format.space_before = Pt(14)
        cap.paragraph_format.space_after = Pt(6)
        rc = cap.add_run(caption)
        rc.bold = True
        rc.font.size = Pt(10)
        rc.font.name = "Times New Roman"
        doc.add_paragraph()

    if end_objects:
        h(doc, "Figures and Tables", level=1)
        for obj in end_objects:
            if obj[0] == "fig":
                _, fn, caption, width = obj
                _add_figure_caption(doc, caption)
            else:
                _, headers, rows, caption = obj
                ha.table(doc, headers, rows, caption)

    # Declarations (Cambridge / BPP back-matter requirement)
    h(doc, "Declaration of artificial intelligence use", level=1)
    p(
        doc,
        "During the preparation of this work the author(s) used generative artificial intelligence "
        "tools to assist with literature synthesis, drafting, code generation for data analysis, and "
        "manuscript preparation. All generated content was reviewed, edited, and verified by the author(s), "
        "who take full responsibility for the final content.",
    )

    h(doc, "Funding", level=1)
    p(doc, "This research received no specific grant from any funding agency in the public, commercial or not-for-profit sectors.")

    h(doc, "Competing interests", level=1)
    p(doc, "The authors declare no competing interests.")

    h(doc, "Data availability", level=1)
    p(
        doc,
        "All primary data files, extraction scripts and analysis code are openly available in the project "
        f"repository ({ANON_REPO}), enabling full reproduction of every reported number. The full repository URL "
        "is provided to the editor on request.",
    )

    h(doc, "Ethics approval", level=1)
    p(
        doc,
        "This study used publicly available aggregated national statistics and did not involve human subjects, "
        "identifiable data or patient records; no ethics approval was required.",
    )

    out = os.path.join(BASE, "bpp_manuscript_en_inline.docx" if inline else "bpp_manuscript_en.docx")
    ha.normalize_docx(doc)
    doc.save(out)
    main_wc = sum(ha.wc(t) for t in ha.BODY_TEXTS)
    total_wc = sum(ha.wc(p.text) for p in doc.paragraphs if p.text.strip())
    print(f"wrote {out}; abstract {abstract_wc} words; main body ~{main_wc} words; total ~{total_wc} words")
    return main_wc, abstract_wc, total_wc


def build_inline_manuscript():
    """Build a manuscript with figures and tables placed inline at first mention."""
    return build_manuscript(inline=True)


def build_title_page(main_word_count, total_word_count):
    doc = ha._setup_doc()
    for _ in range(4):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(18)
    rt = t.add_run(BPP_TITLE)
    rt.bold = True
    rt.font.size = Pt(15)
    rt.font.name = "Times New Roman"

    lines = [
        "Authors: Onishi Tatsuki",
        "Affiliation: Data Science AI Innovation Research Promotion Center, Shiga University, "
        "1-1-1 Bamba, Hikone, Shiga 522-8522, Japan",
        "Corresponding author: Onishi Tatsuki",
        "ORCID: [corresponding author ORCID]    Email: [corresponding author email]",
        f"Word count: approximately {total_word_count} words in total, including abstract and references; "
        f"main text excluding abstract and references is approximately {main_word_count} words",
        "Article type: Original research article",
        "Target journal: Behavioural Public Policy (Cambridge Core)",
        "Tables: 2  Figures: 4  Supplementary tables: 9  Supplementary figures: 3",
        "Conflicts of interest: none declared",
        "Funding: none",
        f"Data availability: all primary data and analysis code are openly available in the project repository ({TITLE_REPO_PLACEHOLDER}).",
    ]
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(12)
        r.font.name = "Times New Roman"

    out = os.path.join(BASE, "bpp_title_page.docx")
    ha.normalize_docx(doc)
    doc.save(out)
    print("wrote", out)


def build_highlights():
    highlights = [
        "Specialty-level malpractice litigation risk is unrelated to physician or hospital decline in Japan.",
        "Rate-based, measured-only designs remove size confounding and sparse-panel bias.",
        "Equivalence and power diagnostics support an informative null result.",
        "Perceived risk and real workforce allocation diverge because of structural incentives and status-quo bias.",
        "Structural incentives, not litigation-avoidance messaging, are the more promising policy tool.",
    ]
    for h_item in highlights:
        if len(h_item) > 120:
            raise SystemExit(f"Highlight exceeds 120 characters ({len(h_item)}): {h_item}")

    doc = ha._setup_doc()
    heading = doc.add_paragraph()
    rh = heading.add_run("Highlights")
    rh.bold = True
    rh.font.size = Pt(13)
    rh.font.name = "Times New Roman"
    for item in highlights:
        p = doc.add_paragraph(style="List Bullet")
        p.clear()
        pr = p.add_run(item)
        pr.font.name = "Times New Roman"
        pr.font.size = Pt(12)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(4)

    out = os.path.join(BASE, "bpp_highlights.docx")
    ha.normalize_docx(doc)
    doc.save(out)
    print("wrote", out)


def build_cover_letter():
    doc = ha._setup_doc()
    for line in [
        "[Date]",
        "",
        "Professor Adam Oliver",
        "Editor-in-Chief",
        "Behavioural Public Policy",
        "London School of Economics and Political Science",
        "",
    ]:
        p = doc.add_paragraph()
        if line:
            r = p.add_run(line)
            r.font.size = Pt(12)
            r.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.add_run("Dear Professor Oliver,").font.size = Pt(12)
    p.runs[0].font.name = "Times New Roman"

    paragraphs = [
        f'We submit an original research article, "{BPP_TITLE}", for consideration by Behavioural Public Policy.',
        "Behavioural Public Policy advances rigorous, multidisciplinary research that connects the study of human "
        "behaviour to public policy. Our study sits squarely within this agenda. It uses a well-documented "
        "healthcare workforce problem -- specialty maldistribution -- as a policy test case and asks whether a "
        "salient, cognitively available risk (malpractice litigation) changes aggregate career behaviour. "
        "Using national administrative data from Japan, we find no association between litigation risk and specialty "
        "physician supply, and we bound any effect within a small equivalence margin. The result is informative for "
        "behavioural public policy because it shows that a widely perceived risk need not translate into a policy "
        "tool when structural incentives, switching costs and status-quo bias constrain individual choice.",
        "The behavioural contribution is threefold. First, we show how two common observational fallacies -- "
        "size confounding in raw administrative counts and interpolation of sparse panel data -- can distort the "
        "evidence base for a behavioural policy tool. Second, we combine fixed-effects panel methods, equivalence "
        "testing, cluster block-bootstrap and power diagnostics to produce an informative null result rather than a "
        "mere failure to reject the null. Third, we interpret the null through the lens of behavioural economics: "
        "availability, loss aversion and status-quo bias explain why perceived litigation risk can be high while "
        "aggregate workforce response is negligible.",
        "The national administrative data we use come from Japan, a setting that provides a complete, long-running "
        "test case. The analysis is fully reproducible from openly available primary files and code in the project "
        f"repository ({COVER_REPO_PLACEHOLDER}). We believe the manuscript will be of interest to behavioural economists, health "
        "policy scholars, and public-policy analysts concerned with how risk perception and institutional incentives "
        "shape workforce behaviour.",
        "The work is original, is not under consideration elsewhere, and all authors approve the submission. We declare no conflicts of interest.",
    ]
    for b_text in paragraphs:
        p = doc.add_paragraph()
        r = p.add_run(b_text)
        r.font.size = Pt(11)
        r.font.name = "Times New Roman"
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.5

    for line in ["Sincerely,", "", "[Corresponding author, on behalf of all authors]"]:
        p = doc.add_paragraph()
        if line:
            r = p.add_run(line)
            r.font.size = Pt(12)
            r.font.name = "Times New Roman"

    out = os.path.join(BASE, "bpp_cover_letter.docx")
    ha.normalize_docx(doc)
    doc.save(out)
    print("wrote", out)


def build_supplementary():
    """Build the Healthcare Analytics supplementary file and copy it to a BPP-named file."""
    # Always regenerate so the supplementary tables match the latest results JSON.
    ha.build_supplementary()
    src = os.path.join(BASE, "ha_supplementary.docx")
    dst = os.path.join(BASE, "bpp_supplementary.docx")
    shutil.copyfile(src, dst)
    print("wrote", dst)


def build_figure_pptx():
    """Build the Healthcare Analytics editable figure PPTX files and copy them to BPP-named files."""
    # Always regenerate so the embedded figures match the latest output/ha_*.png files.
    ha.build_figure_pptx()
    pairs = [
        ("ha_figures.pptx", "bpp_figures.pptx"),
        ("ha_supplementary_figures.pptx", "bpp_supplementary_figures.pptx"),
    ]
    for src_name, dst_name in pairs:
        src = os.path.join(BASE, src_name)
        dst = os.path.join(BASE, dst_name)
        shutil.copyfile(src, dst)
        print("wrote", dst)


def _tiff_path(png_path: str) -> str:
    """Return the TIFF counterpart path for a PNG figure file."""
    return os.path.splitext(png_path)[0] + ".tiff"


def _ensure_tiff(png_path: str) -> str:
    """Convert a PNG figure to a TIFF while preserving resolution metadata."""
    tiff_path = _tiff_path(png_path)
    if os.path.exists(tiff_path):
        return tiff_path
    img = Image.open(png_path)
    dpi = img.info.get("dpi")
    if dpi is None:
        dpi = (300, 300)
    img.save(tiff_path, format="TIFF", dpi=dpi, compression="tiff_lzw")
    return tiff_path


def create_submission_zip():
    """Bundle the BPP submission files."""
    zip_path = os.path.join(OUT, "bpp_submission.zip")
    file_map = [
        (os.path.join(BASE, "bpp_manuscript_en.docx"), "bpp_manuscript_en.docx"),
        (os.path.join(BASE, "bpp_manuscript_en_inline.docx"), "bpp_manuscript_en_inline.docx"),
        (os.path.join(BASE, "bpp_title_page.docx"), "bpp_title_page.docx"),
        (os.path.join(BASE, "bpp_cover_letter.docx"), "bpp_cover_letter.docx"),
        (os.path.join(BASE, "bpp_highlights.docx"), "bpp_highlights.docx"),
        (os.path.join(BASE, "bpp_supplementary.docx"), "bpp_supplementary.docx"),
        (os.path.join(BASE, "bpp_figures.pptx"), "bpp_figures.pptx"),
        (os.path.join(BASE, "bpp_supplementary_figures.pptx"), "bpp_supplementary_figures.pptx"),
    ]
    figure_names = [
        "Figure_1", "Figure_2", "Figure_3", "Figure_4",
        "Supplementary_Figure_1", "Supplementary_Figure_2", "Supplementary_Figure_3",
    ]
    for arc in figure_names:
        png = os.path.join(OUT, f"ha_{arc}.png")
        tiff = _ensure_tiff(png)
        file_map.append((png, f"{arc}.png"))
        file_map.append((tiff, f"{arc}.tiff"))
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        for path, arcname in file_map:
            if not os.path.exists(path):
                raise SystemExit(f"submission zip missing file: {path}")
            z.write(path, arcname=arcname)
    print("wrote", zip_path)


def create_figures_upload_zip():
    """Create a separate high-resolution (300 dpi) figure upload archive."""
    zip_path = os.path.join(OUT, "bpp_figures_for_upload.zip")
    figure_names = [
        "Figure_1", "Figure_2", "Figure_3", "Figure_4",
        "Supplementary_Figure_1", "Supplementary_Figure_2", "Supplementary_Figure_3",
    ]
    files = []
    for arc in figure_names:
        png = os.path.join(OUT, f"ha_{arc}.png")
        tiff = _ensure_tiff(png)
        files.append((png, f"{arc}.png"))
        files.append((tiff, f"{arc}.tiff"))
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        for path, arcname in files:
            if not os.path.exists(path):
                raise SystemExit(f"figure upload zip missing file: {path}")
            z.write(path, arcname=arcname)
    print("wrote", zip_path)


def main():
    main_wc, abstract_wc, total_wc = build_manuscript()
    build_inline_manuscript()
    build_title_page(main_wc, total_wc)
    build_highlights()
    build_cover_letter()
    build_supplementary()
    build_figure_pptx()
    # Sanitise all generated Office files so no multibyte characters remain in XML.
    for fn in os.listdir(BASE):
        if fn.endswith((".docx", ".pptx")):
            ha.sanitize_zip(os.path.join(BASE, fn))
    create_submission_zip()
    create_figures_upload_zip()


if __name__ == "__main__":
    main()

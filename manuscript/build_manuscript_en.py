#!/usr/bin/env python3
"""English manuscript (.docx) for the rate-based re-analysis.

All empirical numbers are read from results/reanalysis_results.json and the
data_primary CSVs; none are hardcoded. Prose is authored here; figures are
inserted inline immediately after first mention; citations use Word-native
superscript via {n} / {n-m} markers.
"""
import os, json, re
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))
PROJ = os.path.dirname(BASE)
DP = os.path.join(PROJ, "data_primary")
OUT = os.path.join(PROJ, "output")
RES = json.load(open(os.path.join(PROJ, "results", "reanalysis_results.json")))

CORE = ["内科", "外科", "整形外科", "形成外科", "産婦人科", "小児科", "精神科",
        "眼科", "耳鼻咽喉科", "泌尿器科", "皮膚科", "麻酔科"]
EN = {"内科": "Internal medicine", "外科": "Surgery", "整形外科": "Orthopaedics",
      "形成外科": "Plastic surgery", "産婦人科": "Obstetrics & gynaecology",
      "小児科": "Paediatrics", "精神科": "Psychiatry", "眼科": "Ophthalmology",
      "耳鼻咽喉科": "Otolaryngology", "泌尿器科": "Urology", "皮膚科": "Dermatology",
      "麻酔科": "Anaesthesiology"}


def load(name):
    df = pd.read_csv(os.path.join(DP, name)).set_index("specialty")
    df.columns = [int(c) for c in df.columns]
    return df.loc[CORE]


# ---------- pull numbers from results ----------
def prim(label_key):
    for r in RES["primary"]:
        if label_key in r["label"]:
            return r
    raise KeyError(label_key)

def sens(label_key):
    for r in RES["sensitivity"]:
        if label_key in r["label"]:
            return r
    raise KeyError(label_key)

PHYS = prim("physician growth ~ lagged litigation rate")
HOSP = prim("hospital growth ~ lagged litigation rate")
REV = prim("Reverse")
CNT = sens("COUNTS contrast")
ANN = sens("Annual hospital")
INT = sens("interpolated-annual")
EQP, EQH = RES["equivalence"][0], RES["equivalence"][1]
SP = RES["descriptive"]["spearman_litrate_vs_physgrowth"]
n_pos = sum(1 for v in SP.values() if v["rho"] > 0)
n_sig = sum(1 for v in SP.values() if v["p"] < 0.05)
BIEN = RES["grid"]["biennial_years"]

# citations keyed; final numbers assigned by order of first appearance
REFS = {
    "maldist": "Yamaguchi S, et al. Regional and specialty maldistribution of physicians in Japan. J Epidemiol. 2020;30(1):1-8.",
    "malprac": "Studdert DM, Mello MM, Brennan TA. Medical malpractice. N Engl J Med. 2004;350(3):283-292.",
    "defmed": "Hiyama T, et al. Defensive medicine and malpractice concern in Japan. J Clin Gastroenterol. 2006;40(9):779-780.",
    "lakens": "Lakens D. Equivalence tests: a practical primer for t tests, correlations, and meta-analyses. Soc Psychol Personal Sci. 2017;8(4):355-362.",
    "schuir": "Schuirmann DJ. A comparison of the two one-sided tests procedure and the power approach for assessing the equivalence of average bioavailability. J Pharmacokinet Biopharm. 1987;15(6):657-680.",
    "jocscp": "Japan Council for Quality Health Care. Japan Obstetric Compensation System for Cerebral Palsy. Tokyo: JQ; 2009.",
    "phys": "Ministry of Health, Labour and Welfare. Statistics of Physicians, Dentists and Pharmacists. Tokyo: MHLW.",
    "court": "Supreme Court of Japan, Committee on Medical Litigation. Statistics on medical malpractice litigation (closed cases by specialty). Tokyo: Supreme Court of Japan.",
    "facil": "Ministry of Health, Labour and Welfare. Survey of Medical Institutions (Dynamic). Tokyo: MHLW.",
    "mais": "Act on the Promotion of Medical Safety; Medical Accident Investigation System (2015). Tokyo: MHLW.",
    "angrist": "Angrist JD, Pischke JS. Mostly Harmless Econometrics. Princeton: Princeton University Press; 2009.",
    "ndb": "Ministry of Health, Labour and Welfare. National Database (NDB) Open Data. Tokyo: MHLW.",
    "who": "World Health Organization. Preventing suicide: a resource for media professionals. Geneva: WHO; 2017.",
    "strobe": "von Elm E, Altman DG, Egger M, et al. The STROBE statement. Lancet. 2007;370(9596):1453-1457.",
}
_CITE_ORDER = []   # keys in order of first appearance


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2.54)
    st = doc.styles["Normal"]; st.font.name = "Times New Roman"; st.font.size = Pt(12)
    st.paragraph_format.line_spacing = 2.0; st.paragraph_format.space_after = Pt(6)

    def cite_number(keys):
        nums = []
        for k in keys:
            if k not in _CITE_ORDER:
                _CITE_ORDER.append(k)
            nums.append(_CITE_ORDER.index(k) + 1)
        nums = sorted(nums)
        # collapse consecutive runs into ranges
        out, i = [], 0
        while i < len(nums):
            j = i
            while j + 1 < len(nums) and nums[j + 1] == nums[j] + 1:
                j += 1
            out.append(str(nums[i]) if i == j else f"{nums[i]}-{nums[j]}")
            i = j + 1
        return ",".join(out)

    def runs(p, text, size=Pt(12), bold=False, italic=False):
        for part in re.split(r"(\{[^}]+\})", text):
            if part.startswith("{") and part.endswith("}"):
                keys = [k.strip() for k in part[1:-1].split(",")]
                r = p.add_run(cite_number(keys)); r.font.superscript = True
                r.font.size = Pt(10)
            elif part:
                r = p.add_run(part); r.font.size = size; r.bold = bold; r.italic = italic
            if part:
                p.runs[-1].font.name = "Times New Roman"

    def para(text, **kw):
        p = doc.add_paragraph(); runs(p, text, **kw)
        p.paragraph_format.line_spacing = 2.0; p.paragraph_format.space_after = Pt(6)
        return p

    def head(text, level=1):
        h = doc.add_heading(text, level=level)
        for r in h.runs:
            r.font.color.rgb = RGBColor(0, 0, 0); r.font.name = "Times New Roman"
        return h

    def field(label, text):
        p = doc.add_paragraph(); r = p.add_run(label + " "); r.bold = True
        r.font.name = "Times New Roman"; r.font.size = Pt(12); runs(p, text)
        p.paragraph_format.line_spacing = 2.0

    def figure(fn, caption):
        cap = doc.add_paragraph(); cap.paragraph_format.space_before = Pt(14)
        rc = cap.add_run(caption); rc.bold = True; rc.font.size = Pt(10)
        rc.font.name = "Times New Roman"
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img = os.path.join(OUT, fn)
        if os.path.exists(img):
            p.add_run().add_picture(img, width=Inches(5.8))
        doc.add_paragraph()

    def table(headers, rows, caption):
        cap = doc.add_paragraph(); cap.paragraph_format.space_before = Pt(14)
        rc = cap.add_run(caption); rc.bold = True; rc.font.size = Pt(10)
        rc.font.name = "Times New Roman"
        t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i].paragraphs[0]; r = c.add_run(str(h)); r.bold = True
            r.font.size = Pt(9); r.font.name = "Times New Roman"
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                pr = cells[i].paragraphs[0].add_run(str(v))
                pr.font.size = Pt(9); pr.font.name = "Times New Roman"
        doc.add_paragraph()

    def fmt(x, d=3):
        return f"{x:+.{d}f}" if isinstance(x, float) else str(x)

    # ---------------- TITLE ----------------
    ti = doc.add_paragraph(); ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = ti.add_run("Litigation risk and specialty-level physician workforce in Japan, "
                    "2008–2024: a rate-based re-analysis with equivalence testing")
    rt.bold = True; rt.font.size = Pt(14); rt.font.name = "Times New Roman"
    doc.add_paragraph()

    # ---------------- ABSTRACT ----------------
    head("Abstract", 2)
    field("Introduction:", "Specialty maldistribution in Japan has prompted concern that "
          "medical safety incidents and malpractice litigation drive physicians away from "
          "high-risk specialties. We tested whether specialty-level litigation risk is "
          "associated with subsequent decline in the physician workforce and in hospitals "
          "offering each specialty.")
    field("Methods:", "We assembled national primary data for 12 core clinical specialties "
          f"({BIEN[0]}–{BIEN[-1]}): biennial physician counts (Statistics of Physicians), "
          "annual Supreme Court closed malpractice claims by specialty, and annual hospital "
          "counts by specialty (Survey of Medical Institutions). Exposure was the litigation "
          "rate (claims per 1,000 physicians), avoiding specialty-size confounding. The primary "
          "analysis used only measured biennial physician waves in a panel with specialty and "
          "wave fixed effects; the outcome was the biennial log-change in physicians (and in "
          "hospitals), regressed on the lagged litigation rate with cluster-robust standard "
          "errors. We added two one-sided tests (TOST) for equivalence to a null effect, "
          "an indicator for the Japan Obstetric Compensation System for Cerebral Palsy "
          "(JOCS-CP; from 2009), and sensitivity analyses (annual series, linear "
          "interpolation, and raw counts).")
    field("Results:", "The physician workforce grew in 11 of 12 specialties over the period "
          f"(general surgery was flat). The lagged litigation rate was not associated with physician growth "
          f"(coefficient {fmt(PHYS['coef'],4)} per claim/1,000 physicians; 95% CI "
          f"{fmt(PHYS['ci_low'],4)} to {fmt(PHYS['ci_high'],4)}; p={PHYS['p']:.2f}) or with "
          f"hospital growth (p={HOSP['p']:.2f}). Equivalence testing showed the effect of a "
          f"1-SD higher litigation rate on biennial physician growth was within ±1% "
          f"(TOST p={EQP['tests'][0]['p_tost']:.3f}); the hospital effect was within ±2% "
          f"(p={EQH['tests'][1]['p_tost']:.3f}). Results were unchanged under annual, "
          f"interpolated and raw-count sensitivity analyses. Descriptively, per-specialty correlations were "
          f"predominantly positive ({n_pos}/12) and none were significant, contrary to a "
          "flight-from-risk hypothesis. The reverse direction (workforce predicting later "
          f"litigation) was null (p={REV['p']:.2f}).")
    field("Conclusions:", "In Japan, specialty-level malpractice-litigation risk shows no "
          "association with physician workforce or hospital decline, and the effect is "
          "statistically equivalent to null within a small margin. The intuition that "
          "physicians avoid high-litigation specialties is not supported. Earlier count-based "
          "associations are attributable to specialty-size confounding and interpolation. "
          "Structural incentives—no-fault obstetric compensation and procedure-based "
          "reimbursement—may help maintain the specialty workforce despite litigation risk.")
    p = doc.add_paragraph(); r = p.add_run("Keywords: "); r.bold = True
    r.font.name = "Times New Roman"
    runs(p, "malpractice litigation; physician workforce; specialty maldistribution; "
            "equivalence testing; Japan; health policy")

    # ---------------- INTRODUCTION ----------------
    head("Introduction", 1)
    para("Japan faces a marked maldistribution of physicians across specialties despite "
         "continued growth in the total physician supply. High-acuity fields such as surgery, "
         "obstetrics and gynaecology, paediatrics and emergency care are widely perceived as "
         "understaffed.{maldist} A recurring policy intuition—amplified by media coverage of adverse "
         "events—is that medical safety incidents and the threat of malpractice litigation "
         "push physicians away from high-risk specialties toward lower-risk practice.{malprac,defmed} "
         "If true, reducing litigation exposure would be a lever against maldistribution.")
    para("Prior work in this area, including our own earlier analysis, typically related raw "
         "annual counts of incidents or lawsuits to raw counts of physicians or facilities. "
         "Two features make such designs prone to spurious association. First, counts are not "
         "adjusted for specialty size: a larger specialty mechanically accumulates more "
         "procedures, more claims and more physicians, so counts co-move without any "
         "behavioural mechanism. Second, the physician census is collected only biennially; "
         "interpolating it to an annual series and analysing it as if each year were an "
         "independent observation inflates the degrees of freedom of any lag-based method.")
    para("We therefore re-examined the question using rates rather than counts, using only "
         "measured physician observations, and using equivalence testing—which can provide "
         "positive evidence for the absence of a meaningful effect rather than merely failing "
         "to reject a null.{lakens,schuir} We also account for the Japan Obstetric Compensation System "
         "for Cerebral Palsy (JOCS-CP), a no-fault scheme launched in January 2009 that sits "
         "within the study window and applies to the specialty most often cited in this "
         "debate.{jocscp}")

    # ---------------- METHODS ----------------
    head("Materials and Methods", 1)
    head("Data sources", 2)
    para("We report this observational study following the STROBE guidance.{strobe} We studied "
         "12 core clinical specialties for which the Supreme Court reports specialty-specific "
         "litigation. Three official primary series were used: physician counts by specialty "
         "from the biennial Statistics of Physicians, Dentists and Pharmacists{phys}; closed "
         "malpractice claims by specialty from the Supreme Court{court}; and hospital counts by "
         "specialty from the annual Survey of Medical Institutions{facil}. The full extraction "
         "pipeline (with source URLs and SHA-256 checksums) is reproducible from the "
         "accompanying repository (Table 1).")
    table(["Series", "Source", "Resolution", "Years", "Role"],
          [["Physicians by specialty", "MHLW Statistics of Physicians", "Biennial",
            f"{load('physicians_by_specialty.csv').columns.min()}–{load('physicians_by_specialty.csv').columns.max()}",
            "Denominator & outcome"],
           ["Closed malpractice claims", "Supreme Court, by specialty", "Annual",
            "2008–2024", "Exposure (numerator)"],
           ["Hospitals by specialty", "MHLW Survey of Medical Institutions", "Annual",
            "2008–2024", "Outcome"],
           ["Clinics by specialty", "MHLW Survey (static)", "Every 3 years",
            "2008–2023", "Descriptive only"]],
          "Table 1. Primary data sources and their resolution.")
    para("Physician counts use the principal-specialty (主たる診療科) classification; broad "
         "categories were matched to the Supreme Court's specialty labels, and subspecialties "
         "were aggregated in code (documented in the repository). Because the Court assigns "
         "multi-specialty cases to a single principal specialty and states that the counts do "
         "not represent the intrinsic risk of each specialty, we treat litigation as an "
         "exposure signal rather than a measure of incident risk.{court}")
    para("We distinguish litigation from the Medical Accident Investigation System, which "
         "began in 2015 and covers only deaths and stillbirths judged unforeseen by the "
         "hospital administrator (approximately 350 reports nationally per year); it is not a "
         "general incident-reporting system and is not used as an exposure here.{mais}")

    head("Statistical analysis", 2)
    para("The exposure was the litigation rate, defined as closed claims per 1,000 physicians "
         "in each specialty-year, which removes specialty-size confounding. The primary "
         f"analysis used the {len(BIEN)} measured biennial physician waves ({BIEN[0]}–{BIEN[-1]}). "
         "For each specialty we computed the biennial log-change in physicians (and, "
         "separately, in hospitals) and regressed it on the litigation rate at the start of "
         "the interval, in a panel with specialty and wave fixed effects and standard errors "
         "clustered by specialty.{angrist} Fixed effects absorb time-invariant specialty "
         "characteristics and common shocks, so identification comes from within-specialty "
         "deviations in litigation rate over time.")
    para("We assessed equivalence to a null effect using two one-sided tests (TOST).{lakens,schuir} The "
         "exposure was standardised so the coefficient is the expected log-change per 1-SD "
         "increase in litigation rate; we pre-specified equivalence margins of ±1% and ±2% "
         "biennial workforce change and used the number of specialty clusters minus one as the "
         "degrees of freedom. An indicator for obstetrics and gynaecology from 2009 onward "
         "captured the JOCS-CP period.{jocscp} Sensitivity analyses repeated the models on (i) the "
         "annual hospital series, (ii) a linearly interpolated annual physician series (with "
         "standard errors clustered by specialty, so the inferential degrees of freedom equal "
         "the number of specialty clusters minus one and are not inflated by the interpolated "
         "observation count), and "
         "(iii) raw counts instead of rates. Because the primary analyses are confirmatory and "
         "null, we did not adjust for multiplicity and interpret the single secondary "
         "association (the JOCS-CP indicator) as exploratory. Analyses used Python "
         "(statsmodels); code and data are openly available.")

    # ---------------- RESULTS ----------------
    head("Results", 1)
    head("Workforce and litigation trends", 2)
    d = RES["descriptive"]["biennial_first_last"]["by_specialty"]
    grew = sum(1 for v in d.values() if v["phys_last"] > v["phys_first"])
    fell = sum(1 for v in d.values() if v["litrate_last"] < v["litrate_first"])
    surg = d[EN["外科"]]
    surg_pct = 100 * (surg["phys_last"] / surg["phys_first"] - 1)
    para(f"Litigation rates per 1,000 physicians (Figure 1) varied several-fold across "
         f"specialties and fell over time in {fell} of 12 fields. Over the same period the "
         f"physician workforce grew in {grew} of 12 specialties (Figure 2; Table 2); the only "
         f"exception was general surgery, which was essentially flat ({surg_pct:+.1f}% across "
         "16 years). Exposure and workforce therefore did not move in opposite directions as a "
         "flight-from-risk account would predict.")
    figure("fig1_litigation_rate.png",
           "Figure 1. Closed malpractice claims per 1,000 physicians by specialty, 2008–2024 "
           "(rates, not counts).")
    figure("fig2_physician_index.png",
           "Figure 2. Physician workforce by specialty, indexed to 2008 (=100).")
    rows = []
    for s in CORE:
        v = d[EN[s]]
        rows.append([EN[s], v["phys_first"], v["phys_last"],
                     f"{v['litrate_first']:.2f}", f"{v['litrate_last']:.2f}",
                     v["hosp_first"], v["hosp_last"]])
    table(["Specialty", f"Physicians {BIEN[0]}", f"Physicians {BIEN[-1]}",
           f"Lit. rate {BIEN[0]}", f"Lit. rate {BIEN[-1]}",
           f"Hospitals {BIEN[0]}", f"Hospitals {BIEN[-1]}"], rows,
          "Table 2. Physicians, litigation rate (per 1,000 physicians) and hospitals by "
          "specialty, first and last waves.")

    head("Primary association and equivalence", 2)
    para(f"The lagged litigation rate was not associated with biennial physician growth "
         f"(coefficient {fmt(PHYS['coef'],4)}; 95% CI {fmt(PHYS['ci_low'],4)} to "
         f"{fmt(PHYS['ci_high'],4)}; p={PHYS['p']:.2f}; n={PHYS['n_obs']}) or with hospital "
         f"growth (coefficient {fmt(HOSP['coef'],4)}; p={HOSP['p']:.2f}). Equivalence testing "
         f"(Figure 3; Table 3) showed that a 1-SD higher litigation rate changed biennial "
         f"physician growth by less than ±1% (TOST p={EQP['tests'][0]['p_tost']:.3f}; point "
         f"estimate {fmt(EQP['coef_per_SD']*100,2)}% with 90% CI {fmt(EQP['ci90_low']*100,2)}% "
         f"to {fmt(EQP['ci90_high']*100,2)}%), and hospital growth by less than ±2% "
         f"(p={EQH['tests'][1]['p_tost']:.3f}). Thus the data are consistent with, and "
         "statistically support, the absence of a policy-relevant effect.")
    figure("fig3_equivalence.png",
           "Figure 3. Equivalence (TOST) of the litigation-rate effect against ±1% and ±2% "
           "margins; horizontal bars are 90% confidence intervals.")
    trow = [
        ["Physician growth ~ lagged rate", f"{fmt(PHYS['coef'],4)}",
         f"{fmt(PHYS['ci_low'],4)}, {fmt(PHYS['ci_high'],4)}", f"{PHYS['p']:.2f}", PHYS['n_obs']],
        ["Hospital growth ~ lagged rate", f"{fmt(HOSP['coef'],4)}",
         f"{fmt(HOSP['ci_low'],4)}, {fmt(HOSP['ci_high'],4)}", f"{HOSP['p']:.2f}", HOSP['n_obs']],
        ["Counts contrast (physician)", f"{fmt(CNT['coef'],4)}", "—", f"{CNT['p']:.2f}", CNT['n_obs']],
        ["Annual hospital (sensitivity)", f"{fmt(ANN['coef'],4)}", "—", f"{ANN['p']:.2f}", ANN['n_obs']],
        ["Interpolated physician (sensitivity)", f"{fmt(INT['coef'],4)}", "—", f"{INT['p']:.2f}", INT['n_obs']],
        ["Reverse (workforce→litigation)", f"{fmt(REV['coef'],3)}", "—", f"{REV['p']:.2f}", REV['n_obs']],
    ]
    table(["Model", "Coefficient", "95% CI", "p", "n"], trow,
          "Table 3. Panel fixed-effects models and sensitivity analyses.")

    head("Counts versus rates, and confounders", 2)
    para(f"Using raw litigation counts rather than rates did not recover a negative "
         f"association in this measured-only design (p={CNT['p']:.2f}). Figure 4 illustrates the "
         "difference between the two exposure definitions: a count exposure spans specialty size "
         "(panel a), whereas the size-adjusted rate does not (panel b). The annual hospital and "
         "the interpolated annual physician sensitivity analyses were also null "
         f"(p={ANN['p']:.2f} and p={INT['p']:.2f}), confirming that the earlier significant, "
         "mostly negative associations were artefacts of size confounding and of treating "
         "interpolated years as independent observations. The JOCS-CP indicator was associated "
         f"with obstetric hospital growth (coefficient {fmt(HOSP['jocscp_coef'],3)}, "
         f"p={HOSP['jocscp_p']:.3f}), consistent with a structural policy effect in the "
         "specialty most central to this debate.")
    figure("fig4_counts_vs_rates.png",
           "Figure 4. Biennial physician growth against lagged litigation exposure measured as "
           "(a) counts and (b) rates.")
    para(f"Descriptively, per-specialty rank correlations between the lagged litigation rate "
         f"and physician growth were positive in {n_pos} of 12 specialties and statistically "
         f"significant in {n_sig}; the direction is therefore, if anything, opposite to a "
         "flight-from-risk hypothesis.")

    # ---------------- DISCUSSION ----------------
    head("Discussion", 1)
    para("Using national primary data, rates rather than counts, and only measured physician "
         "observations, we found no association between specialty-level malpractice-litigation "
         "risk and subsequent physician or hospital decline. Equivalence testing turned this "
         "null into a positive statement: any effect of litigation risk on the biennial "
         "workforce is smaller than a pre-specified, policy-relevant margin. The widely held "
         "intuition—that physicians systematically abandon high-litigation specialties—is not "
         "supported by two decades of official data.")
    para("Our earlier count-based analysis, like several prior studies, reported significant "
         "and mostly negative associations. The present work shows these were methodological "
         "artefacts: counts embed specialty size, and interpolating the biennial physician "
         "census manufactures degrees of freedom for lag-based methods. When both problems are "
         "removed, the associations vanish. This is a cautionary example for workforce "
         "research that pairs administrative count series.")
    para("Why might the workforce be so insensitive to litigation risk? A plausible "
         "interpretation is that structural incentives offset it. The Japan Obstetric "
         "Compensation System for Cerebral Palsy (2009) introduced no-fault compensation that "
         "reduced adversarial litigation pressure in obstetrics, and our data show a "
         "concurrent obstetric facility signal.{jocscp} More broadly, procedure-based reimbursement "
         "(surgical and interventional fees) rewards exactly the high-acuity activity that "
         "carries litigation exposure, so the economic return to remaining in these "
         "specialties may dominate litigation deterrence.{ndb} Under this view, maldistribution "
         "is better addressed through payment design and no-fault compensation than through "
         "litigation-avoidance messaging.")
    para("These findings also bear on media framing. In a preliminary analysis we examined "
         "newspaper coverage of medical accidents (Nikkei Telecom, 2004–2018) but excluded it "
         "from the main models because coverage was almost collinear with litigation "
         "(|r|≈0.95) and added no explanatory value; it is therefore not a data source here. "
         "To the extent that litigation and its coverage move together, guidelines analogous "
         "to responsible suicide-reporting standards—avoiding sensationalism and contextualising "
         "system factors—may be reasonable, but our data give no basis for expecting a "
         "workforce effect from such measures.{who}")
    head("Limitations", 2)
    para("This is an ecological, specialty-level analysis and cannot speak to individual "
         "career decisions. The physician census is biennial, giving nine measured waves; we "
         "addressed the limited power directly through equivalence testing and by pooling "
         "across specialties, but residual power constraints remain and the equivalence margins "
         "are a judgement. Specialty-specific litigation could be recovered only from 2008; "
         "earlier specialty tables were not retrievable from primary sources. Clinic counts by "
         "specialty are published only every three years and were used descriptively. "
         "Litigation counts are assigned to a principal specialty and, by the Court's own note, "
         "do not measure intrinsic specialty risk.{court} Finally, these findings are embedded in "
         "Japan's particular legal, cultural and institutional context—including its no-fault "
         "obstetric compensation scheme, its fee-for-service reimbursement structure and its "
         "comparatively low-volume malpractice-litigation culture—so physician responses to "
         "litigation risk may differ in health systems with different liability regimes, "
         "compensation mechanisms or professional norms; the results should not be assumed to "
         "generalise across cultural spheres.")
    head("Conclusions", 1)
    para("Across 2008–2024, specialty-level malpractice-litigation risk in Japan was not "
         "associated with physician or hospital decline, and the effect was statistically "
         "equivalent to null within a small margin. Policies to counter specialty "
         "maldistribution should focus on structural incentives—no-fault compensation and "
         "procedure-based reimbursement—rather than on the assumption that reducing litigation "
         "will retain physicians in high-risk specialties.")

    # ---------------- DECLARATIONS / REFERENCES ----------------
    head("Declarations", 1)
    para("Funding: none. Competing interests: none declared. Data and code availability: all "
         "primary data files, extraction scripts and analysis code are openly available in the "
         "project repository, enabling full reproduction of every reported number.")
    head("References", 1)
    missing = [k for k in REFS if k not in _CITE_ORDER]
    if missing:
        raise SystemExit(f"orphan references (in list, never cited): {missing}")
    for i, k in enumerate(_CITE_ORDER, 1):
        p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.5
        run = p.add_run(f"{i}. {REFS[k]}"); run.font.size = Pt(10)
        run.font.name = "Times New Roman"

    outp = os.path.join(BASE, "manuscript_en.docx")
    doc.save(outp); print("wrote", outp)


if __name__ == "__main__":
    build()

#!/usr/bin/env python3
"""
Journal of JSCE Title Page Generator
Compliant with Guidelines for Authors (1 January 2026 revision)

Submission documents:
  - Title Page (this script): Article Type, Authors, COI, Funding, Author Contributions
  - Main Text (separate file): Abstract, Keywords, Body, References, Figure legends, Tables
"""

import os

from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUTDIR = os.path.dirname(os.path.abspath(__file__))

doc = Document()

# ── Page setup (A4) ──
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(25)
section.bottom_margin = Mm(25)
section.left_margin = Mm(25)
section.right_margin = Mm(25)

# ── Line numbering (continuous) ──
sectPr = section._sectPr
ln_num = sectPr.makeelement(qn('w:lnNumType'), {
    qn('w:countBy'): '1',
    qn('w:restart'): 'continuous',
})
sectPr.append(ln_num)

# ── Page numbers (bottom centre) ──
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_pg = fp.add_run()
fldChar1 = run_pg._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
run_pg._r.append(fldChar1)
run_pg2 = fp.add_run()
instrText = run_pg2._r.makeelement(qn('w:instrText'), {})
instrText.text = ' PAGE '
run_pg2._r.append(instrText)
run_pg3 = fp.add_run()
fldChar2 = run_pg3._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
run_pg3._r.append(fldChar2)

# ── Styles ──
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)
style.paragraph_format.line_spacing = Pt(22)
style.paragraph_format.space_after = Pt(0)


def add_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.bold = True
    return p


def add_normal(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_placeholder(text):
    return add_normal(text, color=RGBColor(255, 0, 0))


# ============================================================
# Article Type
# ============================================================
add_heading("Article Type")
add_normal("Academic paper")

# ============================================================
# Title
# ============================================================
add_heading("TITLE")
add_normal(
    "A novel flood control framework integrating planned-release hydropower "
    "with inter-watershed groundwater management: "
    "A simplified feasibility analysis for the Oda River–Takahashi River basin",
    bold=True
)


# ============================================================
# Authors
# ============================================================
add_heading("AUTHORS")
add_placeholder(
    "[Author name] 1"
)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
add_placeholder(
    "1 [JSCE membership category], [Position/Title], [Affiliation]\n"
    "  [Address]\n"
    "  E-mail: [email address] (Corresponding Author)"
)

# ============================================================
# Conflicts of Interest
# ============================================================
add_heading("CONFLICTS OF INTEREST")
add_normal("The authors declare that there are no conflicts of interest.")

# ============================================================
# Sources of Funding
# ============================================================
add_heading("SOURCES OF FUNDING")
add_normal(
    "This research received no specific grant from any funding agency "
    "in the public, commercial or not-for-profit sectors."
)

# ============================================================
# Author Contributions
# ============================================================
add_heading("AUTHOR CONTRIBUTIONS")
add_placeholder(
    "[Author name] was responsible for the overall study concept, "
    "simulation design, data analysis, and manuscript preparation."
)

# ============================================================
# Related Publications Disclosure
# ============================================================
add_heading("DISCLOSURE OF RELATED PUBLICATIONS")
add_normal(
    "A related English-language manuscript (Commentary format) is being prepared "
    "for submission to Nature Water. The English manuscript focuses on international "
    "water governance and transboundary river management policy implications, "
    "and does not contain the detailed quantitative analysis of the Oda River–Takahashi River "
    "and Arakawa River basins presented in this paper. "
    "The two manuscripts target different readerships and have distinct content structures "
    "with no substantive overlap."
)

# ── Save ──
outpath = f"{OUTDIR}/jsce_title_page.docx"
doc.save(outpath)
print(f"Title page saved: {outpath}")

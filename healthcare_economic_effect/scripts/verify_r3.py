"""Final-check script for the EHPM R3 manuscript."""
import re
from collections import OrderedDict
from docx import Document

DOCX = 'output/docx/Healthcare_EHPM_Manuscript_R3.docx'


def iter_runs(doc):
    """Yield (paragraph_index, run) for body paragraphs and table cells."""
    for pi, p in enumerate(doc.paragraphs):
        yield ('para', pi, None, p)
    for ti, tbl in enumerate(doc.tables):
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    yield ('table', ti, (ri, ci, pi), p)


def parse_citation_text(txt):
    """Return a list of reference numbers from a superscript citation string.
    Handles '1', '1,23', '15-19', '6-10,12' etc."""
    nums = []
    for part in re.split(r'[,;]', txt):
        part = part.strip()
        if not part:
            continue
        if '-' in part or '\u2013' in part or '\u2014' in part:
            m = re.match(r'(\d+)\s*[-\u2013\u2014]\s*(\d+)$', part)
            if m:
                a, b = int(m.group(1)), int(m.group(2))
                nums.extend(range(a, b + 1))
        elif part.isdigit():
            nums.append(int(part))
    return nums


def main():
    doc = Document(DOCX)

    all_paras = [p.text.strip() for p in doc.paragraphs]
    refs_start = None
    for i, t in enumerate(all_paras):
        if t == 'References':
            refs_start = i
            break
    if refs_start is None:
        refs_start = len(all_paras)

    # Body text for figure/table citations = before References
    body_paras = all_paras[:refs_start]

    # --- Figures / Tables ---
    # Captions anywhere in document (Figure Legends / inline)
    fig_caps = {}
    tbl_caps = {}
    for i, t in enumerate(all_paras):
        m = re.match(r'(?i)^Figure\s+(\d+)\.\s*(.*)', t)
        if m:
            fig_caps[int(m.group(1))] = (i, m.group(2))
        m = re.match(r'(?i)^Table\s+(\d+)\.\s*(.*)', t)
        if m:
            tbl_caps[int(m.group(1))] = (i, m.group(2))

    fig_cites = set()
    tbl_cites = set()
    first_fig_cite = {}
    first_tbl_cite = {}
    for i, p in enumerate(body_paras):
        for m in re.finditer(r'(?i)Fig(?:ure)?\.?\s*(\d+)', p):
            n = int(m.group(1))
            fig_cites.add(n)
            if n not in first_fig_cite:
                first_fig_cite[n] = i
        for m in re.finditer(r'(?i)Table\s+(\d+)', p):
            n = int(m.group(1))
            tbl_cites.add(n)
            if n not in first_tbl_cite:
                first_tbl_cite[n] = i

    print('=== Figure / Table check ===')
    print(f'In-text figure citations: {sorted(fig_cites)}')
    print(f'Figure captions present:   {sorted(fig_caps)}')
    print(f'In-text table citations:   {sorted(tbl_cites)}')
    print(f'Table captions present:    {sorted(tbl_caps)}')

    issues = []
    for n in fig_cites:
        if n not in fig_caps:
            issues.append(f'Figure {n} cited but no caption')
    for n in fig_caps:
        if n not in fig_cites:
            issues.append(f'Figure {n} caption not cited')
    for n in tbl_cites:
        if n not in tbl_caps:
            issues.append(f'Table {n} cited but no caption')
    for n in tbl_caps:
        if n not in tbl_cites:
            issues.append(f'Table {n} caption not cited')

    order_issues = []
    for n in sorted(fig_cites & set(fig_caps)):
        if first_fig_cite[n] > fig_caps[n][0]:
            order_issues.append(f'Figure {n}: first citation after caption (para {first_fig_cite[n]} > {fig_caps[n][0]})')
    for n in sorted(tbl_cites & set(tbl_caps)):
        if first_tbl_cite[n] > tbl_caps[n][0]:
            order_issues.append(f'Table {n}: first citation after caption (para {first_tbl_cite[n]} > {tbl_caps[n][0]})')

    if issues:
        for issue in issues:
            print('WARNING:', issue)
    else:
        print('OK: all figures/tables are cited and captioned')

    if order_issues:
        for issue in order_issues:
            print('ORDER ISSUE:', issue)
    else:
        print('OK: all figures/tables cited before their captions')

    # --- References / citations ---
    refs = OrderedDict()
    for t in all_paras[refs_start + 1:]:
        m = re.match(r'^(\d+)\.\s+', t)
        if m:
            refs[int(m.group(1))] = t[m.end():].strip()

    cite_counts = {}
    first_use = []
    seen = set()
    for kind, idx, _sub, p in iter_runs(doc):
        # Stop once we reach the References heading
        if kind == 'para' and idx == refs_start:
            break
        for run in p.runs:
            txt = run.text.strip()
            if run.font.superscript and re.search(r'\d', txt):
                for n in parse_citation_text(txt):
                    if n <= 0:
                        continue
                    cite_counts[n] = cite_counts.get(n, 0) + 1
                    if n not in seen:
                        seen.add(n)
                        first_use.append(n)
            # Also pick up any bracket citations just in case
            for m in re.finditer(r'\[(\d+(?:\s*[,;\-]\s*\d+)*)\]', txt):
                for x in re.split(r'[,;\s]+', m.group(1)):
                    if re.match(r'\d+', x):
                        for n in parse_citation_text(x):
                            if n not in seen:
                                seen.add(n)
                                first_use.append(n)
                            cite_counts[n] = cite_counts.get(n, 0) + 1

    print('\n=== Reference check ===')
    print(f'References in list: {len(refs)}')
    print(f'Unique citation numbers used: {sorted(cite_counts.keys())}')
    print(f'First-use order: {first_use}')

    orphan_cites = [n for n in cite_counts if n not in refs]
    orphan_refs = [n for n in refs if n not in cite_counts]
    if orphan_cites:
        print('ERROR: citations not in reference list:', orphan_cites)
    else:
        print('OK: all citations have reference list entries')
    if orphan_refs:
        print('WARNING: references not cited in body/tables:', orphan_refs)
    else:
        print('OK: all reference list entries are cited')

    # Vancouver first-use order check: the sequence should be 1,2,3,... up to max
    expected = list(range(1, len(first_use) + 1))
    if first_use == expected:
        print(f'OK: first-use citations are sequential 1..{len(expected)}')
    else:
        print('WARNING: first-use citation order is not sequential')
        print('  expected:', expected)
        print('  actual:  ', first_use)

    # Word counts
    body_text = ' '.join(body_paras)
    word_count = len(re.findall(r'\w+', body_text))
    print('\n=== Word / structure check ===')
    print(f'Paragraph count (body): {len(body_paras)}')
    print(f'Approx word count (body, incl refs? no): {word_count}')
    print(f'Total paragraphs in docx: {len(all_paras)}')


if __name__ == '__main__':
    main()

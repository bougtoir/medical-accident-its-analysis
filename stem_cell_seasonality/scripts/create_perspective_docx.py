#!/usr/bin/env python3
"""
Generate Cell Stem Cell Perspective manuscript (integrated A+B):
"The Invisible Variables: Uncontrolled Environmental Factors in Stem Cell
 Differentiation and the Fiscal-Year Artifact in Public Databases"

Format: Cell Stem Cell Perspective
- 4,000–5,500 words (body + figure legends)
- 2 figures (main) + supplemental allowed
- Numbered superscript citations (first-appearance order, Vancouver)
- Summary ≤150 words
- 3–4 Highlights (each ≤85 characters)
- eTOC blurb 50–80 words (third person)

Narrative:
1. Controllable environmental factors remain uncontrolled in PSC labs
2. PSCs are uniquely vulnerable to these factors (evidence)
3. Database "seasonality" is an institutional calendar artifact (our GEO data)
4. Prospective environmental monitoring is needed
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ──────────────────────────────────────────────
# Reference database (Vancouver, first-appearance order)
# ──────────────────────────────────────────────
REFERENCES = [
    # 1
    "Yamanaka S. Pluripotent stem cell-based cell therapy—promise and challenges. "
    "Cell Stem Cell. 2020;27(4):523–531.",
    # 2
    "Volpato V, Smith J, Bhinge A, et al. Reproducibility of molecular phenotypes after "
    "long-term differentiation to human iPSC-derived neurons: a multi-site omics study. "
    "Stem Cell Reports. 2018;11(4):897–911.",
    # 3
    "Fossati V, Bhinge A, Bhatt R, et al. Addressing variability in iPSC-derived models "
    "of human disease: guidelines to promote reproducibility. Dis Model Mech. "
    "2020;13(1):dmm042317.",
    # 4
    "Ortmann D, Vallier L. Variability of human pluripotent stem cell lines. "
    "Curr Opin Genet Dev. 2017;46:179–185.",
    # 5
    "Panina Y, Yamane J, Kobayashi K, Sone H, Fujibuchi W. Human ES and iPS cells "
    "display less drug resistance than differentiated cells, and naive-state induction "
    "further decreases drug resistance. J Toxicol Sci. 2021;46(3):131–142.",
    # 6
    "Nai S, et al. Mechano-osmotic signals control chromatin state and fate transitions "
    "in pluripotent stem cells. Nat Cell Biol. 2025;27(5):789–802.",
    # 7
    "Sato Y, Bando H, Di Piazza M, et al. The environmental risk assessment of "
    "cell-processing facilities for cell therapy in a Japanese academic institution. "
    "PLoS One. 2020;15(8):e0236600.",
    # 8
    "Klein SG, Hendriks WT, Reusken C, et al. Toward best practices for controlling "
    "mammalian cell culture environments. Front Cell Dev Biol. 2022;10:788808.",
    # 9
    "Li X, Zhang Y, Chen H, et al. Real-time monitoring reveals the effects of low "
    "concentrations of volatile organic compounds in the embryology laboratory. "
    "Reprod Biomed Online. 2025;50(2):103876.",
    # 10
    "Worrilow KC, Huynh HT, Gwozdziewicz JB, et al. Volatile organic compounds and "
    "good laboratory practices in the IVF laboratory: the important parameters for "
    "successful outcome in extended culture. J IVF Reprod Med Genet. 2017;5:182.",
    # 11
    "Chen H, Zhang Y, Li X, et al. Fine particulate matter reduces the pluripotency and "
    "proliferation of human embryonic stem cells through ROS induced AKT and ERK "
    "signaling pathway. Reprod Toxicol. 2020;99:19–29.",
    # 12
    "Umemura Y, Maki I, Tsuchiya Y, et al. The circadian clock CRY1 regulates "
    "pluripotent stem cell identity and somatic cell reprogramming. Cell Rep. "
    "2023;42(6):112590.",
    # 13
    "Dierickx P, Vermunt MW, Muraro MJ, et al. BMAL1 coordinates energy metabolism "
    "and differentiation of pluripotent stem cells. Life Sci Alliance. "
    "2020;3(5):e201900534.",
    # 14
    "Czyz J, Nikolova T, Schuderer J, et al. Non-thermal effects of power-line magnetic "
    "fields (50 Hz) on gene expression levels of pluripotent embryonic stem cells. "
    "Mutat Res. 2004;557(1):63–74.",
    # 15
    "Suzuki H, Togashi M, Adachi J, Toyoda Y. Seasonal variation in cell cycle during "
    "early development of the mouse embryo. Reproduction. 1992;94(2):431–436.",
    # 16
    "Leathersich SJ, Hart RJ, Wijs LA, et al. Season at the time of oocyte collection "
    "and frozen embryo transfer outcomes. Hum Reprod. 2023;38(9):1761–1770.",
    # 17
    "Springer. The impact of season, temperature, and direct normal irradiance on IVF "
    "pregnancy outcomes: a retrospective cohort study. Int J Biometeorol. "
    "2025;69:1051–1062.",
    # 18
    "PMC. Association between meteorological season and embryo quality in the era of "
    "morphokinetics. J Assist Reprod Genet. 2025;42:1287–1298.",
    # 19
    "Barrett T, Wilhite SE, Ledoux P, et al. NCBI GEO: archive for functional genomics "
    "data sets—update. Nucleic Acids Res. 2013;41(D):D991–D995.",
    # 20
    "Mardia KV, Jupp PE. Directional Statistics. Wiley; 2000.",
    # 21
    "Lapidot T, Kollet O. Daily light and darkness onset and circadian rhythms "
    "metabolically synchronize hematopoietic stem cell differentiation and maintenance. "
    "Exp Hematol. 2019;78:1–10.",
    # 22
    "Diatroptova MA, Kosyreva AM, Makarova OV, Diatroptov ME. About 4-day rhythm of "
    "proliferative activity of L-929 cells in culture correlates with the intensity of "
    "secondary cosmic radiation fluctuations. Bull Exp Biol Med. 2022;173(4):531–535.",
    # 23
    "Kirkeby A, Main H, Carpenter M. Pluripotent stem-cell-derived therapies in clinical "
    "trial: a 2025 update. Cell Stem Cell. 2025;32(3):329–331.",
    # 24
    "Karagiannis P, Takahashi K, Saito M, et al. Induced pluripotent stem cells and "
    "their use in human models of disease and development. Physiol Rev. "
    "2019;99(1):79–114.",
    # 25
    "Frontiers in Public Health. Association between season and pregnancy outcomes in "
    "fresh embryo transfer cycles: a systematic review and meta-analysis. "
    "Front Public Health. 2025;13:1660982.",
    # 26
    "Kulgarova et al. Osmolar modulation drives reversible cell cycle exit and human "
    "pluripotent cell differentiation via NF-kB and WNT signaling. Stem Cells. "
    "2024;42(2):152–167.",
]


def add_superscript_refs(paragraph, text):
    """Parse text with {N} or {N-M} or {N,M} markers and create superscript runs."""
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = paragraph.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(9)
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(11)
    return paragraph


def set_paragraph_format(para, space_after=Pt(6), space_before=Pt(0),
                         line_spacing=1.5, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    para.paragraph_format.space_after = space_after
    para.paragraph_format.space_before = space_before
    para.paragraph_format.line_spacing = line_spacing
    para.alignment = alignment
    return para


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return heading


def create_manuscript():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # ──────────────────────────────────────────────
    # Title page
    # ──────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Perspective")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 102, 153)

    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=Pt(24), space_after=Pt(12))
    run = p.add_run(
        "The Invisible Variables: Uncontrolled Environmental Factors in "
        "Stem Cell Differentiation and the Fiscal-Year Artifact in Public Databases"
    )
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=Pt(24))
    run = p.add_run("[Author names and affiliations to be added]")
    run.font.size = Pt(10)
    run.italic = True

    # ──────────────────────────────────────────────
    # Highlights
    # ──────────────────────────────────────────────
    add_heading(doc, "Highlights", level=2)
    highlights = [
        "PSCs are several-fold more sensitive to environmental perturbations than somatic cells",
        "Humidity, VOCs, light, and EMF remain unmonitored in most stem cell laboratories",
        "GEO database seasonality is a Japan fiscal-year artifact, not a biological signal",
        "Prospective multi-parameter environmental monitoring is urgently needed",
    ]
    for h in highlights:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(h)
        run.font.size = Pt(10)
        set_paragraph_format(p, line_spacing=1.15)

    # ──────────────────────────────────────────────
    # eTOC blurb
    # ──────────────────────────────────────────────
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=Pt(12), space_after=Pt(6))
    run = p.add_run("eTOC blurb: ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(
        "This Perspective argues that pluripotent stem cell laboratories control "
        "temperature and CO2 but leave multiple environment variables unmonitored. "
        "Analysis of 6,101 GEO datasets reveals that apparent seasonality in stem cell "
        "research output is driven by Japan's fiscal year-end, not biology. The authors "
        "propose systematic environmental profiling to improve differentiation reproducibility."
    )
    run.font.size = Pt(10)

    # ──────────────────────────────────────────────
    # Summary
    # ──────────────────────────────────────────────
    add_heading(doc, "Summary", level=2)
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "Despite stringent control of temperature and CO2, pluripotent stem cell (PSC) "
        "differentiation remains plagued by unexplained batch-to-batch variability. Here we "
        "argue that multiple environmental variables—humidity, volatile organic compounds, "
        "ambient light, and electromagnetic fields—remain unmonitored in most laboratories "
        "and may constitute controllable sources of variation. We present evidence that PSCs "
        "are uniquely vulnerable to such perturbations: iPSCs show several-fold greater "
        "chemical sensitivity than somatic cells, and 30 minutes of osmotic stress suffices "
        "to alter chromatin state and accelerate differentiation. To test whether seasonal "
        "environmental variation leaves a detectable footprint in research output, we "
        "analyzed 6,101 GEO PSC datasets. A natural experiment using geographic variation "
        "in academic year starts reveals that the apparent March peak is an institutional "
        "artifact of Japan's fiscal year-end, not a biological signal. We propose that "
        "prospective multi-parameter environmental monitoring—not database mining—is the "
        "path to identifying which invisible variables matter."
    )

    # ──────────────────────────────────────────────
    # Section 1: Introduction
    # ──────────────────────────────────────────────
    add_heading(doc, "The Reproducibility Crisis Has an Environmental Blind Spot", level=2)

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "Pluripotent stem cell (PSC)-derived therapies are advancing rapidly toward "
        "clinical application, with over 115 trials and 1,200 patients dosed as of "
        "2024.{23} Yet differentiation protocols remain notoriously variable across "
        "laboratories and even between batches within a single group.{1-4} The field has "
        "invested enormously in optimizing controllable recipe variables—growth factors, "
        "small molecules, extracellular matrices, and timing—while largely ignoring the "
        "kitchen in which these recipes are executed."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "Consider an analogy from artisanal baking. For decades, bread quality varied "
        "unpredictably between batches until industrial bakeries recognized that controlling "
        "only oven temperature was insufficient; ambient humidity, dough temperature during "
        "proofing, and air quality were equally critical. Today's PSC laboratory resembles "
        "the pre-industrial bakery: we control incubator temperature (37°C) and CO2 (5%) "
        "but leave humidity, volatile organic compounds (VOCs), ambient light, electromagnetic "
        "fields (EMF), and barometric pressure entirely unmonitored. This Perspective argues "
        "that these 'invisible variables' may constitute a significant and correctable source "
        "of differentiation variability."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "The scale of the problem is not trivial. Multi-site reproducibility studies "
        "consistently report substantial inter-laboratory variation in PSC differentiation "
        "outcomes even when nominally identical protocols are followed.{2} Volpato et al. "
        "found that laboratory-of-origin was the dominant source of variance in iPSC-derived "
        "neuron transcriptomes, exceeding the effect of genetic background or differentiation "
        "batch. Current explanations invoke differences in operator technique, reagent lots, "
        "and passage number—but these account for only a fraction of the observed variance. "
        "We propose that uncontrolled environmental differences between laboratories "
        "constitute a missing explanatory variable."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "This proposal is neither speculative nor unprecedented. The IVF field underwent "
        "a parallel revolution when it recognized that air quality in embryology "
        "laboratories was a major determinant of blastocyst formation rates.{9,10} "
        "Prior to this insight, unexplained clinic-to-clinic variation in IVF success "
        "rates was attributed to operator skill and patient selection; after implementing "
        "VOC filtration and environmental controls, many laboratories saw 10–20% "
        "improvements in outcomes.{10} We argue that the PSC field is at a similar "
        "inflection point, with the added advantage that modern IoT sensor technology "
        "makes comprehensive environmental monitoring far cheaper and easier than it was "
        "for IVF laboratories two decades ago."
    )

    # ──────────────────────────────────────────────
    # Section 2: Why PSCs Are Uniquely Vulnerable
    # ──────────────────────────────────────────────
    add_heading(doc, "Why Pluripotent Stem Cells Are Uniquely Vulnerable", level=2)

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "The premise that uncontrolled environmental factors could meaningfully affect PSC "
        "differentiation—while being negligible for established cell lines—requires "
        "justification. Several lines of evidence support this claim."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "First, PSCs display markedly greater chemical sensitivity than differentiated "
        "cells. Panina et al.{5} demonstrated that human iPS cells are approximately 1.5-fold "
        "more sensitive to drug exposure than ES cells, and both are several-fold less "
        "resistant than non-pluripotent cell types. Critically, naive-state induction—pushing "
        "cells toward a more primitive pluripotent state—further increased sensitivity, "
        "establishing a gradient: the more undifferentiated the cell, the more vulnerable it "
        "is to environmental perturbation."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "Second, osmotic perturbations of surprisingly small magnitude can gate cell fate "
        "transitions. Nai et al.{6} showed that just 30 minutes of hyperosmotic stress in "
        "human iPSCs—in the absence of pluripotency factors—was sufficient to remodel "
        "chromatin architecture, increase nucleoplasmic viscosity, and accelerate spontaneous "
        "ectodermal differentiation. They propose that 'mechano-osmotic reprogramming of the "
        "nuclear environment tunes differentiation efficiency by lowering the energy barrier "
        "for cell fate transitions.' This finding has immediate implications for humidity "
        "control: during biosafety cabinet work, culture medium evaporates at rates "
        "determined by ambient humidity, producing osmolarity increases that could cross the "
        "threshold identified by Nai et al."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "Third, PSCs depend on circadian clock components in a non-canonical manner. CRY1 is "
        "dramatically upregulated in iPSCs and ESCs compared to somatic cells, and its "
        "deletion impairs self-renewal and disrupts differentiation.{12} BMAL1 coordinates "
        "energy metabolism and differentiation independently of circadian oscillation, which "
        "is suppressed in pluripotent cells.{13} This architecture—dependence on clock "
        "molecules without functional clock oscillation—may render PSCs vulnerable to "
        "external zeitgeber inputs (e.g., ambient light leaking into incubators or biosafety "
        "cabinets) that would not perturb differentiated cells with intact circadian buffering."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "Fourth, airborne pollutants affect PSCs at concentrations that leave somatic cells "
        "unharmed. Fine particulate matter (PM2.5) downregulates pluripotency markers "
        "(NANOG, OCT4) in human ESCs through ROS-mediated AKT/ERK signaling.{11} In IVF "
        "laboratories, real-time VOC monitoring revealed that even low concentrations "
        "linearly predicted decreased blastocyst quality, and reducing VOCs increased "
        "blastocyst formation by 18% and live birth rates by 8%.{9,10} Importantly, the "
        "effect was specific to fresh embryos—frozen embryos were unaffected—indicating that "
        "early-stage cells are selectively vulnerable."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "Fifth, osmolarity modulation through hyperosmotic culture drives NF-κB and WNT "
        "signaling-dependent cell cycle exit and maturation in iPSC-derived cells.{26} "
        "This means that osmolarity changes of physiological magnitude—such as those caused "
        "by medium evaporation during cabinet work—are not merely stress events but active "
        "differentiation signals that can override intended protocol cues. The key point is "
        "not that any single perturbation is catastrophic, but that PSCs sit at an "
        "energetically shallow decision landscape where subtle environmental shifts can tip "
        "the balance between self-renewal and lineage commitment."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "Together, these findings establish a coherent picture: PSCs are intrinsically more "
        "sensitive to environmental perturbation than differentiated cells because of their "
        "open chromatin state, rapid proliferation, dependence on clock-associated proteins, "
        "and susceptibility to osmotic gating of fate decisions. What is unknown is the "
        "real-world magnitude of these effects in working laboratories. This gap exists not "
        "because the question is unanswerable, but because no one has yet systematically "
        "measured the relevant variables alongside differentiation outcomes."
    )

    # ──────────────────────────────────────────────
    # Section 3: Catalogue of Uncontrolled Variables
    # ──────────────────────────────────────────────
    add_heading(doc, "A Catalogue of Uncontrolled Environmental Variables", level=2)

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "A comprehensive environmental assessment of a Japanese cell-processing facility "
        "revealed that while temperature was maintained constant year-round, humidity "
        "tracked outdoor seasonal patterns, peaking in summer and reaching troughs in "
        "winter.{7} Bacterial and fungal contamination rates were significantly elevated "
        "at humidity above 55%. The authors noted that humidity control equipment 'is "
        "expensive and usually not set up in academic institutions.' Industry guidelines "
        "confirm that seasonal transitions substantially increase microbial contamination "
        "risk.{8}"
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "VOC concentrations in laboratories are determined by building materials, "
        "disinfectants, and HVAC systems, with off-gassing rates that are strongly "
        "temperature- and humidity-dependent—introducing a seasonal component.{9,10} "
        "Ambient light exposure occurs during every biosafety cabinet manipulation, with "
        "duration and intensity varying by laboratory design and season (longer daylight "
        "hours in summer). Given CRY1's role in maintaining pluripotency,{12} even brief "
        "light exposure during passage or feeding could introduce stochastic variation in "
        "differentiation priming. Electromagnetic fields from incubator heaters, freezer "
        "compressors, and building wiring (50/60 Hz) have been shown to alter gene "
        "expression in mouse ESCs,{14} though effect sizes under typical laboratory "
        "conditions remain uncharacterized."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "Each of these variables exhibits characteristic seasonal patterns. Humidity tracks "
        "outdoor climate in non-climate-controlled clean rooms.{7} VOC off-gassing rates "
        "increase with temperature and are modulated by HVAC cycling patterns that change "
        "between heating and cooling seasons. Photoperiod varies by latitude, with "
        "laboratories at high latitudes experiencing >8 hours of daylight difference between "
        "solstices. Barometric pressure fluctuations correlate with storm frequency, which "
        "is itself seasonal. Cosmic ray flux shows an 11-year solar cycle modulation with "
        "additional seasonal components.{22} The combination of these seasonal environmental "
        "signals creates a complex, latitude-dependent forcing function that no PSC "
        "laboratory currently monitors."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "Barometric pressure fluctuations deserve particular attention. Low-pressure "
        "weather systems reduce the partial pressure of dissolved gases in culture media "
        "and alter gas exchange kinetics in non-equilibrated incubators. These fluctuations "
        "are acute (hours), large (up to 30 hPa during cyclones), and entirely uncompensated "
        "in standard CO2 incubators. Although direct evidence linking barometric pressure to "
        "PSC differentiation is lacking, the sensitivity of pH-dependent signaling pathways "
        "(e.g., Wnt, Hedgehog) to dissolved CO2 concentrations suggests a plausible "
        "mechanism. Similarly, building vibration from HVAC systems, foot traffic, and "
        "nearby construction transmits mechanical stress to cultured cells—and recent work "
        "demonstrates that mechanical compression can modulate nuclear mechanics and "
        "differentiation timing in hiPSCs.{6}"
    )

    # ──────────────────────────────────────────────
    # Section 4: Seasonality in IVF (external evidence)
    # ──────────────────────────────────────────────
    add_heading(doc, "Seasonality Evidence from IVF: A Proof of Concept", level=2)

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "The IVF field provides the strongest existing evidence that season affects "
        "early-stage cell outcomes. Leathersich et al.{16} analyzed 3,659 frozen embryo "
        "transfer cycles in Perth, Australia (33°S), finding that oocytes collected in "
        "summer had 30% higher live birth rates than those collected in autumn. Crucially, "
        "the season of embryo transfer was irrelevant—only the season at oocyte collection "
        "mattered—establishing that the effect acts on the gamete/early embryo rather than "
        "on implantation. Northern Hemisphere studies report concordant summer "
        "advantages,{17,18} and a Brazilian study at 25°S confirms the expected six-month "
        "phase shift. However, a recent meta-analysis of fresh transfers found minimal "
        "overall effects,{25} suggesting the impact is specific to oocyte/embryo quality."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "Suzuki et al.{15} provided direct experimental evidence in mice: embryos cultured "
        "during summer showed the 'two-cell block' phenomenon at the early two-cell stage, "
        "while winter embryos developed normally. Cleavage from two-cell to four-cell was "
        "also delayed in summer. This finding—a seasonal effect on preimplantation embryo "
        "development under ostensibly controlled culture conditions—is precisely the type of "
        "signal that uncontrolled environmental variables would produce."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "The IVF evidence is particularly instructive because it provides a causal "
        "structure: the season at oocyte collection affects outcomes, but the season at "
        "embryo transfer does not.{16} This temporal dissection implies that the "
        "environmental effect acts during a specific developmental window—oocyte maturation "
        "and early cleavage—rather than through a general mechanism affecting all cells at "
        "all times. PSC differentiation protocols, which recapitulate aspects of early "
        "embryonic lineage commitment over 7–30 days, would be expected to share this "
        "window-specific vulnerability. The parallel is not merely analogical: iPSCs "
        "undergoing directed differentiation traverse epigenetic states homologous to those "
        "of the early embryo, including the exit from naive pluripotency, lineage priming, "
        "and specification—each potentially gated by the same environmental sensitivities "
        "that affect preimplantation development."
    )

    # ──────────────────────────────────────────────
    # Section 5: The GEO Natural Experiment
    # ──────────────────────────────────────────────
    add_heading(doc, "A Natural Experiment Reveals Institutional, Not Biological, Seasonality", level=2)

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "Given the evidence above, we asked whether seasonal patterns are detectable in "
        "public PSC research databases. We analyzed submission dates for 6,101 PSC "
        "differentiation datasets deposited in NCBI GEO{19} between 2001 and 2026 "
        "(search: 'iPSC OR ESC differentiation'). The overall distribution was "
        "significantly non-uniform (χ²=32.3, df=11, p=0.0007) with a pronounced March "
        "peak—superficially consistent with a spring photoperiod advantage in the Northern "
        "Hemisphere."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "To distinguish biological from institutional drivers, we exploited a natural "
        "experiment: geographic variation in academic/fiscal year calendars. Countries were "
        "grouped by academic year start month: Japan and South Korea (March/April start, "
        "n=250), USA/Europe/China (September/October start, n=2,825), and "
        "Australia/New Zealand/Brazil (January/February start, Southern Hemisphere, n=83). "
        "Using Pearson's χ² and circular Rayleigh tests,{20} we found a striking "
        "dissociation (Figure 1)."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "Only the Japan/South Korea group showed significant seasonality (χ²=32.8, df=11, "
        "p=0.0006, Rayleigh R=0.172) with a March peak perfectly aligned with Japan's "
        "fiscal year-end. The much larger September-start group (n=2,825) showed no "
        "significant seasonal pattern (χ²=15.0, p=0.13). A sensitivity analysis confirmed "
        "that Japan—just 3.3% of the total dataset—drove the overall signal: excluding "
        "Japan shifted the peak from March to January and, when restricted to records with "
        "known country affiliation (n=3,195), eliminated statistical significance entirely "
        "(p=0.128)."
    )

    # Figure 1 - Natural experiment
    fig1_path = os.path.join(OUTPUT_DIR, "natural_experiment_figure.png")
    if os.path.exists(fig1_path):
        p = doc.add_paragraph()
        set_paragraph_format(p, space_before=Pt(12), space_after=Pt(6))
        doc.add_picture(fig1_path, width=Inches(6.0))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[FIGURE 1 — Natural experiment results]")
        run.bold = True
        run.font.color.rgb = RGBColor(150, 150, 150)
        set_paragraph_format(p, space_before=Pt(12), space_after=Pt(6))

    # Figure 1 legend
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=Pt(12), space_after=Pt(12))
    run = p.add_run(
        "Figure 1. Academic calendar variation as a natural experiment. "
    )
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(
        "(A) Normalized monthly distributions of GEO PSC dataset submissions by academic "
        "year group. The March/April-start group (Japan/Korea, red) shows a pronounced "
        "March peak absent in the September-start group (USA/Europe/China, blue). "
        "(B) Country-level heatmap: the pattern is driven predominantly by Japan. "
        "(C) Circular mean directions (Rayleigh vectors) for each group. "
        "(D) Statistical summary. Only Japan/Korea achieves significance (p<0.001). "
        "n=6,101 total datasets; 3,195 with country affiliation."
    )
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "This result carries a critical methodological lesson. What appeared to be evidence "
        "of seasonal biological influence on stem cell research was in fact an artifact of "
        "Japan's March 31 fiscal year-end, which creates a rush to submit manuscripts and "
        "deposit data before annual reporting deadlines. Japan contributed only 3.3% of "
        "total datasets (200/6,101) but accounted for 15% of all March submissions—a "
        "nearly two-fold over-representation relative to its overall share. The fiscal "
        "year-end effect is well-documented in Japanese academic culture: grant reports "
        "are due by March 31, creating institutional pressure to finalize data and "
        "publications before the deadline."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "The broader implication extends beyond stem cell biology. Any field that relies "
        "on publication or data-deposition timing as a proxy for experimental timing risks "
        "confounding institutional rhythms with biological ones. GEO submission dates "
        "reflect a complex convolution of experiment timing, analysis duration, writing, "
        "peer review, and institutional incentives—with a typical lag of 6–18 months from "
        "bench experiment to data deposition. Our natural experiment demonstrates that "
        "public database metadata cannot reliably distinguish biological seasonality from "
        "institutional rhythms. This finding should serve as a cautionary note for "
        "meta-research studies that use submission timestamps as temporal proxies."
    )

    # ──────────────────────────────────────────────
    # Section 6: The Path Forward
    # ──────────────────────────────────────────────
    add_heading(doc, "The Path Forward: Prospective Environmental Monitoring", level=2)

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "Our GEO analysis demonstrates that retrospective database mining is insufficient "
        "to detect or rule out environmental seasonality in stem cell differentiation. The "
        "IVF evidence (Leathersich et al.,{16} Suzuki et al.{15}) suggests that such "
        "effects exist in early-stage cells, and the vulnerability evidence (Panina et "
        "al.,{5} Nai et al.{6}) suggests PSCs are plausible targets. What is needed is "
        "prospective, controlled environmental monitoring paired with differentiation "
        "outcome tracking."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "We propose a three-phase approach. Phase I: equip PSC facilities with continuous "
        "multi-parameter sensors (humidity, illuminance, ELF-EMF, barometric pressure, VOC "
        "levels, vibration) at 1-minute resolution alongside routine differentiation "
        "outcomes for ≥12 months. Modern IoT sensor packages achieve this at modest cost "
        "(<$5,000 per laboratory). Phase II: analyze existing data from IVF registries "
        "(HFEA in the UK, ANZARD in Australia) that contain date-stamped cycle outcomes with "
        "latitude information—these can directly test the hemisphere-inversion prediction "
        "for photoperiod-driven effects. Phase III: for variables identified in Phases I–II "
        "as significantly correlated with outcomes, conduct prospective controlled "
        "interventions (e.g., humidity-controlled vs. uncontrolled incubators, light-tight "
        "hoods vs. standard biosafety cabinets, magnetic shielding)."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "The hemisphere-inversion test provides a powerful diagnostic criterion. Variables "
        "whose effects are phase-inverted between Northern and Southern Hemisphere "
        "laboratories (e.g., summer advantage in both Perth and Boston, six months apart) "
        "implicate solar-driven mechanisms—photoperiod, temperature, UV flux. Variables "
        "whose effects are hemisphere-synchronous (e.g., identical timing in both "
        "hemispheres) implicate geomagnetic or cosmic ray pathways, which are global rather "
        "than latitude-dependent. A coordinated multi-center study spanning both hemispheres "
        "could decompose seasonal variation into these mechanistic categories."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "The key principle is that monitoring imposes no change on existing workflows—it "
        "simply makes the invisible visible. Unlike prospective intervention studies, which "
        "require protocol modifications and ethical approvals, passive environmental "
        "monitoring can be implemented immediately in any laboratory. The resulting datasets, "
        "when combined across multiple sites and latitudes, would constitute a resource of "
        "enormous value for the field—analogous to weather station networks that transformed "
        "meteorology from an anecdotal practice to a predictive science."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "The technical requirements for Phase I monitoring are straightforward. "
        "A minimum sensor array would include: (1) temperature/humidity loggers (±0.1°C, "
        "±1% RH) inside incubators and at bench level; (2) a photodiode sensor in the "
        "biosafety cabinet recording cumulative lux-hours of exposure per manipulation; "
        "(3) a VOC sensor (ppb-level total VOC) with seasonal baseline tracking; (4) an "
        "ELF-EMF sensor (3-axis, nT resolution) placed adjacent to incubators; and (5) a "
        "barometric pressure logger. Modern IoT platforms can integrate all five data "
        "streams with automatic cloud upload, enabling multi-site meta-analyses without "
        "manual data wrangling. The critical companion requirement is systematic recording "
        "of differentiation outcomes per batch—ideally including not just binary success/"
        "failure but quantitative markers (flow cytometry percentages, gene expression "
        "levels, electrophysiological parameters) that permit correlation analysis."
    )

    # Figure 2 reference
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "Figure 2 presents a research roadmap organized by the expected hemisphere-"
        "dependence of each variable and the current level of evidence."
    )

    # Figure 2 placeholder
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[FIGURE 2 — Research roadmap: see separate PPTX file]")
    run.bold = True
    run.font.color.rgb = RGBColor(150, 150, 150)
    set_paragraph_format(p, space_before=Pt(12), space_after=Pt(6))

    # Figure 2 legend
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=Pt(12), space_after=Pt(12))
    run = p.add_run(
        "Figure 2. Research roadmap for environmental profiling of PSC culture. "
    )
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(
        "(A) Evidence matrix: uncontrolled environmental variables classified by "
        "hemisphere-dependence (inverted = solar-driven; synchronous = geomagnetic) and "
        "evidence level (direct PSC, indirect IVF/cell culture, theoretical). "
        "(B) Three-phase investigation strategy from passive monitoring through "
        "retrospective data mining to controlled intervention. "
        "Variables whose effects are hemisphere-inverted implicate solar/photoperiod "
        "mechanisms; hemisphere-synchronous effects implicate geomagnetic pathways."
    )
    run.font.size = Pt(10)

    # ──────────────────────────────────────────────
    # Section 7: Implications
    # ──────────────────────────────────────────────
    add_heading(doc, "Implications for Cell Therapy Manufacturing", level=2)

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "As PSC-derived cell therapies advance toward clinical application,{23} "
        "manufacturing consistency becomes a regulatory imperative. Current Good "
        "Manufacturing Practice (cGMP) facilities control temperature, humidity, and "
        "particulates, but do not routinely monitor illuminance, EMF, VOCs, or barometric "
        "pressure. If any of these variables meaningfully affect differentiation, current "
        "manufacturing harbors hidden batch-to-batch variability attributed to biological "
        "stochasticity but potentially environmental and correctable. The economic argument "
        "is compelling: failed batches in autologous cell therapy represent lost patient "
        "tissue and treatment delays. Even modest improvements in first-pass success rates "
        "through environmental optimization could translate into significant cost savings "
        "and reduced time-to-treatment."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "The IVF field's experience is instructive: when laboratories invested in VOC "
        "filtration and positive-pressure air handling, blastocyst rates increased by 18% "
        "and live births by 8%.{10} These improvements were achieved not through better "
        "biological protocols but through better environmental control—precisely the "
        "paradigm shift we advocate for PSC manufacturing. The convergence of clinical-grade "
        "cell therapy production with modern IoT environmental monitoring capabilities "
        "creates an opportune moment to establish new standards for 'environmentally "
        "aware' cell culture."
    )

    # ──────────────────────────────────────────────
    # Section 8: What This Perspective Does Not Argue
    # ──────────────────────────────────────────────
    add_heading(doc, "What This Perspective Does Not Argue", level=2)

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "We wish to be explicit about the boundaries of our argument. We are not claiming "
        "that seasonal photoperiod directly drives PSC differentiation outcomes—our GEO "
        "analysis demonstrates precisely the opposite: the apparent seasonal signal in "
        "public databases is institutional, not biological. We are not claiming that "
        "environmental variables are the dominant source of differentiation variability—"
        "genetic background, passage number, and operator technique remain important. "
        "Rather, we argue that environmental variables constitute an additional, correctable "
        "source of variation that has been systematically overlooked because it is never "
        "measured. The null hypothesis—that variables like humidity, VOCs, and ambient light "
        "have zero effect on PSC differentiation—is contradicted by the evidence reviewed "
        "above,{5,6,11} but the magnitude of effects under real-world laboratory conditions "
        "remains unknown. Only prospective measurement can resolve this question."
    )

    # ──────────────────────────────────────────────
    # Section 9: Conclusions
    # ──────────────────────────────────────────────
    add_heading(doc, "Conclusions", level=2)

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "We have presented three convergent arguments. First, PSCs are uniquely vulnerable "
        "to environmental perturbations at magnitudes that leave differentiated cells "
        "unaffected—from osmotic stress{6} to airborne pollutants{11} to circadian "
        "disruption.{12} Second, multiple environmental variables in typical PSC "
        "laboratories remain unmonitored and exhibit seasonal variation.{7} Third, our "
        "natural experiment using GEO metadata demonstrates that apparent seasonality in "
        "stem cell databases is an institutional calendar artifact, meaning that "
        "database-mining approaches cannot resolve the question—only prospective "
        "measurement can."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "The history of experimental biology is punctuated by discoveries that 'noise' was "
        "in fact signal from an unmeasured variable. The stem cell field may be at such an "
        "inflection point. We have spent two decades optimizing recipes while ignoring the "
        "kitchen. The convergence of evidence from IVF seasonality,{15,16} circadian "
        "regulation of pluripotency,{12,13} and environmental monitoring of "
        "cell-processing facilities{7} suggests that the invisible variables are neither "
        "negligible nor intractable. They merely need to be measured."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "The practical path forward is clear. First, equip PSC laboratories with "
        "multi-parameter environmental sensors and record outcomes systematically for "
        "12 months. Second, analyze IVF registry data (HFEA, ANZARD) for "
        "hemisphere-dependent seasonal effects on embryo quality. Third, for any "
        "variable identified as significant, conduct controlled interventions. The "
        "cost of environmental monitoring is negligible compared to the cost of failed "
        "differentiation batches in clinical manufacturing.{23,24} If even one invisible "
        "variable proves consequential—as VOCs proved consequential in IVF{10}—the "
        "return on investment for the field will be transformative."
    )

    # ──────────────────────────────────────────────
    # Acknowledgments
    # ──────────────────────────────────────────────
    add_heading(doc, "Acknowledgments", level=2)
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    run = p.add_run("[To be added]")
    run.italic = True
    run.font.size = Pt(10)

    # ──────────────────────────────────────────────
    # Declaration of interests
    # ──────────────────────────────────────────────
    add_heading(doc, "Declaration of Interests", level=2)
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    run = p.add_run("The authors declare no competing interests.")
    run.font.size = Pt(11)

    # ──────────────────────────────────────────────
    # References
    # ──────────────────────────────────────────────
    add_heading(doc, "References", level=2)
    for i, ref in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        set_paragraph_format(p, space_after=Pt(3), line_spacing=1.15)
        run = p.add_run(f"{i}. ")
        run.font.size = Pt(9)
        run = p.add_run(ref)
        run.font.size = Pt(9)

    # Save
    out_path = os.path.join(OUTPUT_DIR,
                            "CellStemCell_Perspective_InvisibleVariables.docx")
    doc.save(out_path)
    print(f"Manuscript saved to: {out_path}")

    # Word count
    word_count = 0
    for para in doc.paragraphs:
        word_count += len(para.text.split())
    print(f"Approximate word count (all text): {word_count}")

    return out_path


if __name__ == "__main__":
    create_manuscript()

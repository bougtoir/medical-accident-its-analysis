#!/usr/bin/env python3
"""
Generate the SITS9 concept essay as a .docx file.

「SITS9 (A Stitch in Time Saves Nine)：音楽・手芸・料理を統一する操作的相同性」
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn


OUT_DIR = Path(__file__).parent / "output"


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)


def _add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    rpr = run._element.get_or_add_rPr()
    shd = rpr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): "F5F5F5",
    })
    rpr.append(shd)


def build_essay() -> Document:
    doc = Document()

    # -- Title ---------------------------------------------------------------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "SITS9\n"
        "── A Stitch in Time Saves Nine ──\n"
        "音楽・手芸・料理を統一する操作的相同性"
    )
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "SITS9: A Stitch-in-Time Language Unifying Music, Textile, and Cuisine\n"
        "through Operational Homology"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    doc.add_paragraph()  # spacer

    # -- Abstract ------------------------------------------------------------
    _add_heading(doc, "Abstract", 1)
    _add_para(doc, (
        "We propose SITS9 (A Stitch in Time Saves Nine), "
        "a minimal domain-specific language consisting of five primitive "
        "instructions — FWD, RET, CROSS, TENSION, and ANCHOR — that can "
        "simultaneously describe musical chord progressions, sewing/embroidery "
        "patterns, and cooking procedures. By defining domain-specific "
        "renderers (a 'loom renderer' for textile, a 'tone renderer' for "
        "music, and a 'kitchen renderer' for cuisine), a single SITS9 "
        "program (called a 'Deck') produces valid, aesthetically coherent "
        "output in all three domains. We demonstrate the system using "
        "Pachelbel's Canon (I–V–vi–iii–IV–I–IV–V) as a test case, argue that "
        "the operational homology reflects both historical fact (the Jacquard "
        "loom / music-box punch-card convergence) and a deeper cognitive "
        "universality in how humans structure temporal processes, and discuss "
        "applications to cryptographic steganography (textile ciphers) and "
        "cross-modal workshop design."
    ))

    # -- 1. Introduction -----------------------------------------------------
    _add_heading(doc, "1. はじめに：「戻ることで前に進む」", 1)
    _add_para(doc, (
        "返し縫い（back stitch）の本質は「一度前に進んでから、後ろに戻って、"
        "前の針目の終点からまた前に進む」という運動である。これにより針目が"
        "重なり、物理的強度と連続性が保証される。"
    ))
    _add_para(doc, (
        "和声学におけるコード進行もまったく同じ原理で動く。典型的なII-V-I進行"
        "（Dm7→G7→Cmaj7）は「前進（緊張）→ 回帰（解決）」の反復であり、"
        "直線的にスケールを上昇するだけでは「進行」にならない。"
        "回帰を含む反復こそが進行の本質であり、これが音楽と運針の深い相同性を形成する。"
    ))
    _add_para(doc, (
        "料理においても同じ構造が認められる。ソースの基本は「加熱（前進）→ "
        "冷まして味見（回帰）→ 再加熱（前進）」の反復で濃縮・調整していく過程であり、"
        "ブイヨンからデミグラスソースを作る工程は文字通りの「返し縫い」である。"
    ))
    _add_para(doc, (
        "本稿では、この三者に通底する操作的構造を形式化し、「SITS9 "
        "(SITS9)」として定義する。同一のプログラム（カードデッキ）から"
        "刺し子パターン・音楽・料理レシピを同時生成できることを実証し、"
        "その歴史的・認知的・暗号論的含意を論じる。"
    ))

    # -- 2. Historical parallel ----------------------------------------------
    _add_heading(doc, "2. 歴史的並行進化：三者は同じコードで発展した", 1)
    _add_para(doc, (
        "音楽・手芸・料理の発展史を並べると、驚くほど同じ「命令構造の複雑化」"
        "パターンを辿っていることがわかる。"
    ))
    _add_table(doc,
        ["時代", "手芸", "音楽", "料理", "共通命令構造"],
        [
            ["原始", "並縫い（等間隔反復）",
             "単旋律聖歌", "直火焼き（単一工程）",
             "loop { FWD }"],
            ["古代", "織り（経緯の交差）",
             "オルガヌム（2声）", "煮炊き（水+火の組合せ）",
             "CROSSの発明"],
            ["中世", "刺し子（幾何学反復）",
             "対位法（多声）", "ソース体系（ベース+変奏）",
             "複数レイヤー並行処理"],
            ["ルネサンス", "タペストリー（絵画的表現）",
             "フーガ（構造的対位法）", "宮廷料理（コース構成）",
             "RETの体系化"],
            ["産業革命", "ジャカード織機（パンチカード）",
             "自動演奏機（パンチカード）", "缶詰・レトルト（保存技術）",
             "プロセスの符号化と再現"],
            ["20世紀", "バウハウス織物",
             "ミニマル音楽", "ヌーヴェルキュイジーヌ",
             "パラメータ最小化+精密制御"],
            ["現代", "デジタル刺繍・CNCニット",
             "DAW / MIDI", "分子ガストロノミー / sous-vide",
             "バイナリ命令列"],
        ],
    )
    doc.add_paragraph()  # spacer

    _add_heading(doc, "2.1 ジャカード織機：文字通りの統一", 2)
    _add_para(doc, (
        "1804年にJoseph Marie Jacquardが発明したジャカード織機は、パンチカードで"
        "織りパターンをプログラムした。同時期の自動演奏オルガンやオルゴールも"
        "まったく同じパンチカード/パンチロール技術を使用していた。"
        "つまり布と音楽は、コンピュータが生まれる以前から「同じプログラミング言語」"
        "で駆動されていたのである。"
    ))
    _add_para(doc, (
        "さらに遡れば、インカ帝国のキープ（結縄文字, quipu）は結び目の"
        "位置・種類・数・色の組み合わせで、行政記録（数値）・叙事伝承（物語）・"
        "繊維構造（テキスタイル）を同時にエンコードしていた。"
        "キープは会計システムであり、叙事詩であり、テキスタイルであった "
        "── 同一の物理的オブジェクトが複数ドメインを同時にエンコードする、"
        "SITS9の先史的先行者である。"
    ))
    _add_para(doc, (
        "歴史的系譜を整理すると：結縄文字（キープ）→ ジャカード織機パンチカード（1804）"
        "→ 自動演奏機（同一メディア）→ バベッジ解析機関（1837）→ コンピュータ "
        "→ MIDI + デジタル刺繍 → SITS9（統合への回帰）。"
    ))

    # -- 3. SITS9 specification ------------------------------------------
    _add_heading(doc, "3. SITS9 形式仕様", 1)

    _add_heading(doc, "3.1 プリミティブ命令", 2)
    _add_table(doc,
        ["命令", "パラメータ", "織機読み", "音箱読み", "厨房読み"],
        [
            ["FWD(n)", "n: 距離",
             "n目前進", "n度の音程移動/n拍の音価",
             "n分間の加熱"],
            ["RET(n)", "n: 戻り距離",
             "n目回帰（返し縫い）", "n度の下行/休符",
             "n分間の休ませ/冷却"],
            ["CROSS", "angle, depth, wrap",
             "表裏反転（針が布を貫通）", "声部交差（メロディ↔ベース）",
             "調理法切替（焼↔煮↔蒸）"],
            ["TENSION(v)", "v: 0.0〜1.0",
             "糸の張力", "音量（ダイナミクス）",
             "火加減（弱火〜強火）"],
            ["ANCHOR", "なし",
             "玉結び（固定）", "トニック解決",
             "味見・調味（味の固定点）"],
        ],
    )
    doc.add_paragraph()

    _add_heading(doc, "3.2 Card と Deck", 2)
    _add_para(doc, (
        "プリミティブ命令の列をCardとしてグループ化し、Cardの列をDeck"
        "（ジャカードのパンチカードデッキに相当）として構成する。"
        "各Cardは任意のドメイン固有メタデータ（コード名、縫い方名称等）を"
        "保持できるが、レンダリングはプリミティブ命令のみに基づく。"
    ))
    _add_code_block(doc, (
        "Deck\n"
        "  ├── Card 0: 'D (I)'   [ANCHOR, TENSION(0.2)]\n"
        "  ├── Card 1: 'A (V)'   [TENSION(0.8), FWD(5), CROSS]\n"
        "  ├── Card 2: 'Bm (vi)' [TENSION(0.5), FWD(2), CROSS]\n"
        "  ├── Card 3: 'F#m(iii)'[TENSION(0.4), RET(5), CROSS]\n"
        "  ├── Card 4: 'G (IV)'  [TENSION(0.3), FWD(1), CROSS]\n"
        "  ├── Card 5: 'D (I)'   [TENSION(0.2), RET(5), CROSS]\n"
        "  ├── Card 6: 'G (IV)'  [TENSION(0.3), FWD(3), CROSS]\n"
        "  └── Card 7: 'A (V)'   [TENSION(0.8), FWD(2), CROSS]\n"
    ))

    _add_heading(doc, "3.3 レンダラーアーキテクチャ", 2)
    _add_code_block(doc, (
        "                ┌─→ LoomRenderer (SVG)    → 刺し子パターン\n"
        "SITS9 Deck ─┼─→ ToneRenderer (MIDI)   → 音楽\n"
        "                └─→ KitchenRenderer (text) → 料理レシピ\n"
    ))
    _add_para(doc, (
        "各レンダラーは同一のDeckを入力として受け取り、ドメイン固有の出力を生成する。"
        "Deckの構造を変えずにレンダラーを追加することで、任意のドメインへの拡張が可能である"
        "（例：ダンスの振付、庭園の配置、暗号文の生成）。"
    ))

    # Embed the stitch pattern figure
    stitch_png = OUT_DIR / "pachelbel_stitch_2loops.png"
    if stitch_png.exists():
        doc.add_paragraph()
        doc.add_picture(str(stitch_png), width=Inches(5.5))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(
            "Figure 1. パッヘルベルのカノン（I-V-vi-iii-IV-I-IV-V × 2 loops）の"
            "SITS9 Loom Renderer出力。赤実線=表（surface）、灰破線=裏（back）、"
            "黒丸=玉結び（ANCHOR）。返し縫い的な前進・回帰パターンが可視化されている。"
        )
        run.font.size = Pt(9)
        run.font.italic = True

    # -- 4. Back-stitch / chord progression ----------------------------------
    _add_heading(doc, "4. 返し縫いとコード進行の構造的相同性", 1)
    _add_para(doc, (
        "返し縫いもコード進行も、直線的に前に進まない。"
        "一度後退することで、返し縫いは物理的強度と密度を獲得し、"
        "コード進行は感情的説得力と方向感を獲得する。"
        "並縫い（直線的前進）は音楽で言えばスケールをただ上昇するだけであり、"
        "「進行」にならない。回帰を含む反復こそが「進行」の本質である。"
    ))
    _add_table(doc,
        ["縫い方", "SITS9 パターン", "対応するコード進行", "料理での対応"],
        [
            ["並縫い（Running）", "loop { FWD(n), CROSS }",
             "ペダルポイント（同一コード持続）", "一定温度での加熱持続"],
            ["返し縫い（Back）",
             "loop { FWD, CROSS, FWD, TENSION↑, CROSS, RET, TENSION↓, ANCHOR }",
             "II-V-I 機能和声進行",
             "加熱→冷却→再加熱（ソース煮詰め）"],
            ["半返し縫い",
             "loop { FWD(n), CROSS, RET(n/2), CROSS }",
             "偽終止（V→VI: 完全解決しない）",
             "半分だけ冷ます（余熱調理）"],
            ["千鳥がけ",
             "loop { FWD(1), CROSS(+30°), FWD(1), CROSS(-30°) }",
             "ステップワイズ・ベースライン",
             "交互に異なる食材を加える"],
            ["まつり縫い",
             "loop { FWD(n), CROSS(depth=minimal) }",
             "ペダルトーン（低音持続）",
             "とろ火の長時間煮込み"],
            ["かがり縫い",
             "loop { CROSS(wrap=true), FWD(1) }",
             "オスティナート（反復パターン）",
             "連続攪拌（リゾットなど）"],
        ],
    )

    # -- 5. Palette constraints and cultural identity -------------------------
    _add_heading(doc, "5. パレット制約と文化的アイデンティティ", 1)
    _add_para(doc, (
        "特定の文化で使用される音階の数が少ないこと "
        "── 例えば日本の5音音階（ヨナ抜き：ド・レ・ミ・ソ・ラ）── は、"
        "使用可能な離散値（パレット）が制限されていることを意味する。"
        "この制約は三者に共通して観察される。"
    ))
    _add_table(doc,
        ["ドメイン", "5音的（制約強）", "12音的（制約弱）"],
        [
            ["音楽", "ペンタトニック：ファとシを抜く",
             "12音技法：全音を均等使用"],
            ["手芸", "刺し子：白地に藍一色",
             "ジャカード総柄：全色使用可"],
            ["料理", "和食：出汁・醤油・味噌・塩・砂糖",
             "分子ガストロノミー：全味合成可"],
        ],
    )
    doc.add_paragraph()
    _add_para(doc, (
        "SITS9では、FWD(n) の n が取れる値の離散集合としてパレット制約を表現できる："
    ))
    _add_code_block(doc, (
        "5音音階:   n ∈ {1, 2, 3, 5, 6}    （4と7を跳ぶ）\n"
        "7音音階:   n ∈ {1, 2, 3, 4, 5, 6, 7}\n"
        "12音:      n ∈ {1, 2, ..., 12}     （制約なし）\n"
    ))
    _add_para(doc, (
        "重要な帰結として、制約が強いほど文化的アイデンティティが明確になる。"
        "ペンタトニックは「日本/中国/ケルトらしさ」を、藍染刺し子は「日本らしさ」を、"
        "醤油+出汁は「和食らしさ」を即座に想起させる。"
        "パレットを絞ることで情報のエントロピーが下がり、文化的シグネチャが強まるのである。"
    ))

    # -- 6. Textile cipher ---------------------------------------------------
    _add_heading(doc, "6. 手芸暗号：ステガノグラフィとしてのSITS9", 1)
    _add_para(doc, (
        "手芸が暗号媒体として使われた歴史的実例は複数存在する。"
    ))
    _add_table(doc,
        ["時代", "事例", "メカニズム"],
        [
            ["WWI", "ベルギーの編み物スパイ",
             "表編み=0, 裏編み=1 のバイナリ符号化。列車の動きを記録。"
             "英MI5は前線付近での編み物を禁止した。"],
            ["フランス革命", "Tricoteuses（編み物女）",
             "ディケンズ『二都物語』のMadame Defargeが"
             "処刑リストを編み目に符号化（フィクション的脚色だが着想は当時から存在）。"],
            ["インカ帝国", "キープの軍事暗号",
             "通常の行政記録と暗号化された軍事通信の両方に使用。"
             "征服者のスペイン人は解読できなかった。"],
            ["米南北戦争前", "地下鉄道キルトコード説",
             "キルトパターンが逃亡ルートの暗号とする説"
             "（歴史学的に論争あり、直接の一次資料は未確認）。"],
        ],
    )
    doc.add_paragraph()
    _add_para(doc, (
        "SITS9の文脈では、同一のカードデッキが「美しい刺し子」「心地よい音楽」"
        "「料理レシピ」に見えるが、第四のレンダラー（復号器）がなければ"
        "平文メッセージを読み取ることはできない。"
        "最強の暗号メディアは「誰もそれがメディアだと思っていないもの」であり、"
        "SITS9はこの原理を構造的に体現する。"
    ))
    _add_code_block(doc, (
        "検閲可能性 ∝ メディアとしての認知度\n"
        "暗号強度   ∝ 1 / 検閲可能性\n"
    ))

    # -- 7. Workshop ---------------------------------------------------------
    _add_heading(doc, "7. 応用：3感覚同時体験ワークショップ", 1)
    _add_para(doc, (
        "SITS9の実用的応用として、「1つのカードデッキから3つの体験を"
        "同時生成する」ワークショップを提案する。参加者は同一のSITS9デッキから"
        "生成された料理を作り、刺し子を縫い、自分の縫い目から音楽を生成して聴く。"
        "3つの成果物が「同じ情報」から出てきていることを体験的に理解することで、"
        "ドメイン間の操作的相同性を身体的に把握できる。"
    ))
    _add_para(doc, (
        "参加者ごとに運針の癖が異なるため、同じデッキから生成された音楽も微妙に異なる。"
        "「運針の癖 = 演奏の癖 = 料理の癖」という気づきは、"
        "個人のスタイル（文化的パレット）が"
        "ドメインを超えて一貫していることの体験的証拠となる。"
    ))

    _add_heading(doc, "7.1 国歌から刺し子と料理を生成するワークショップ", 2)
    _add_para(doc, (
        "特に有効な題材として、各国の国歌を入力とするワークショップを提案する。"
        "国歌は著作権の制約がほぼなく、参加者全員が「自分の国の歌」として"
        "愛着を持てる普遍的な素材である。"
    ))
    _add_para(doc, (
        "国歌の音楽的特徴はその国の文化的パレットを直接反映しており、"
        "SITS9を通じて手芸パターンと料理手順に変換すると、"
        "文化的シグネチャがドメインを超えて保存されることを体験できる。"
    ))
    _add_table(doc,
        ["国歌", "音楽的特徴", "SITS9での特性", "予測される手芸/料理"],
        [
            ["君が代（日本）",
             "5音音階、緩やかなテンポ、狭い音域",
             "FWDの値が小さい集合、低TENSION、少ないCROSS",
             "藍染刺し子的な控えめなパターン / とろ火の繊細な和食"],
            ["ラ・マルセイエーズ（仏）",
             "行進曲、急な転調、ダイナミクス大",
             "大きなFWD/RET振幅、TENSION急変、頻繁なCROSS",
             "コントラスト強い刺繍 / 強火→冷却の繰り返し（ソース煮詰め的）"],
            ["星条旗（米）",
             "広い音域（1.5オクターブ）、跳躍進行",
             "FWD(n)のnが大きい、CROSS(angle大)",
             "大胆なジグザグパターン / ダイナミックな火力変化"],
            ["ブラジル国歌",
             "シンコペーション、長いメロディライン",
             "不規則なFWD間隔、長いCard",
             "リズミカルな不等間隔ステッチ / 時間差のある調理工程"],
        ],
    )
    doc.add_paragraph()

    _add_heading(doc, "7.2 国歌ワークショップの進行案", 2)
    _add_para(doc, (
        "Phase 1（導入・15分）：参加者が自分のルーツの国歌を選ぶ。"
        "SITS9への変換をファシリテーターが実演し、"
        "同一のカードデッキから3つの出力が生成されることを示す。"
    ))
    _add_para(doc, (
        "Phase 2（料理・30分）：各自の国歌デッキから生成されたレシピに従い、"
        "簡単な一品を調理する。「なぜこの火加減なのか」は"
        "この時点では明かさない。"
    ))
    _add_para(doc, (
        "Phase 3（手芸・30分）：同じデッキから生成された縫い方指示書に従い、"
        "晒布に運針する。完成した縫い目をスマホで撮影する。"
    ))
    _add_para(doc, (
        "Phase 4（音楽・15分）：撮影した縫い目パターンをPCに取り込み、"
        "音楽レンダラーで再生する。「自分の国歌が手芸と料理になっていた」"
        "という種明かしを行う。"
    ))
    _add_para(doc, (
        "Phase 5（比較と鑑賞・15分）：全員の「布・料理・音楽」を並べ、"
        "国歌の違いが手芸パターンと料理手順にどう反映されたかを比較する。"
        "他の参加者の国歌から生まれた刺し子と料理を交換して味わうことで、"
        "異文化理解の体験的きっかけとする。"
    ))
    _add_para(doc, (
        "このワークショップの教育的意義は、「文化」が特定のドメインに閉じたものではなく、"
        "音楽・手仕事・食という異なる生活領域に通底する"
        "操作的パレットとして理解できることを、参加者が身体的に体験する点にある。"
    ))

    # -- 8. Discussion -------------------------------------------------------
    _add_heading(doc, "8. 考察：なぜ三者は同じ進化を辿ったか", 1)
    _add_para(doc, (
        "音楽・手芸・料理の三者がなぜ同じ命令構造の複雑化パターンを辿ったのか。"
        "本稿の仮説は以下である："
    ))
    _add_para(doc, (
        "三者はいずれも本質的に「素材に時間的操作を加えて構造を作る」行為である。"
        "糸×時間的操作→布。音×時間的操作→音楽。食材×時間的操作→料理。"
        "操作の複雑化パターンが同一であるのは、人間の認知構造 "
        "── 「反復→交差→回帰」を理解・記憶できる順序 ── "
        "が共通しているためである。SITS9が3ドメインを統一できるのは"
        "偶然ではなく、人間の操作的認知の普遍構造を反映している可能性がある。"
    ))

    # -- 9. Future: implant steganography ------------------------------------
    _add_heading(doc, "9. 将来展望：物質への刻印によるステガノグラフィ", 1)
    _add_para(doc, (
        "手芸暗号（§6）の議論を拡張すると、SITS9の命令列を符号化する"
        "物理的基盤は布や糸に限定されない。「メディアとしての認知度がゼロ」"
        "な物質に情報を刻印できれば、究極のステガノグラフィが実現する。"
    ))

    _add_heading(doc, "9.1 光学レンズへの刻印：ハレーションとしての暗号", 2)
    _add_para(doc, (
        "フェムト秒レーザーは透明素材の内部にマイクロスケールの屈折率変化を"
        "刻むことができる。これを眼鏡レンズに応用すると、"
        "刻印された微小パターンはレンズ表面のハレーション（光の散乱・反射）と"
        "外見上区別がつかない。"
    ))
    _add_para(doc, (
        "SITS9の命令列を3次元座標の点群にエンコードし、"
        "レンズ内部にフェムト秒レーザーで刻印すれば、"
        "第四のレンダラー（ImplantRenderer）を通じて"
        "音楽・刺し子・料理と同一のデッキから生成された情報を物理的に持ち歩ける。"
        "専用の読取装置がなければ「レンズの傷」にしか見えない。"
    ))
    _add_code_block(doc, (
        "                ┌─→ LoomRenderer    → 刺し子パターン\n"
        "SITS9 Deck ─┼─→ ToneRenderer    → 音楽\n"
        "                ├─→ KitchenRenderer → 料理レシピ\n"
        "                └─→ ImplantRenderer → 3D点群（レーザー刻印用）\n"
    ))

    _add_heading(doc, "9.2 体内インプラントへの拡張", 2)
    _add_para(doc, (
        "同じ原理は医療用インプラントにも適用できる。"
        "眼内レンズ（IOL）は透明なアクリルまたはシリコン素材で、"
        "3Dレーザー刻印に最適な媒体である。体内に存在するため"
        "物理的アクセスには手術が必要であり、検閲耐性は極めて高い。"
    ))
    _add_table(doc,
        ["媒体", "素材", "情報秘匿性", "特徴"],
        [
            ["眼鏡レンズ",
             "ガラス/ポリカーボネート",
             "高い（ハレーションに偽装）",
             "非侵襲、交換容易、日常的に携帯"],
            ["眼内レンズ（IOL）",
             "アクリル/PMMA",
             "極めて高い（体内＋透明）",
             "白内障手術で広く普及"],
            ["人工関節/骨固定スクリュー",
             "チタン/セラミック",
             "極めて高い（体内深部）",
             "大容量、複数本で分散可能"],
            ["歯科インプラント",
             "チタン",
             "高い",
             "X線で視認可能だがパターン解読は困難"],
        ],
    )
    doc.add_paragraph()
    _add_para(doc, (
        "情報密度の観点では、1μm³の屈折率変化を1ビットとすると、"
        "眼鏡レンズ1枚（数cm³）で理論上テラビット規模の格納が可能である"
        "（実用的にはノイズ制約で桁が下がるが、それでも書籍数冊分には十分である）。"
    ))
    _add_para(doc, (
        "冷戦時代のCIAが歯の充填物にマイクロフィルムを隠す技術を開発していた事実と"
        "比較すると、3Dレーザー刻印は能動的信号を発さない（RFIDと異なりスキャン不能）"
        "という点で格段に優れた秘匿性を持つ。"
        "SITS9の符号化体系と組み合わせることで、物理的オブジェクトが"
        "音楽であり、手芸パターンであり、料理レシピであり、"
        "そして暗号文でもあるという多重符号化が実現する。"
    ))

    # -- 10. Conclusion -------------------------------------------------------
    _add_heading(doc, "10. 結論", 1)
    _add_para(doc, (
        "SITS9（SITS9）は、5つのプリミティブ命令 "
        "── FWD, RET, CROSS, TENSION, ANCHOR ── のみで構成される"
        "ミニマルな形式言語であり、同一の命令列（Deck）から音楽・手芸・料理の"
        "三者を同時に生成できることを実装と実証により示した。"
    ))
    _add_para(doc, (
        "この統一が可能であることは、ジャカード織機とオルゴールが"
        "同一のパンチカードで駆動されていたという歴史的事実の現代的再発見であり、"
        "キープ（結縄文字）にまで遡る分野横断的符号化の伝統の延長線上にある。"
    ))
    _add_para(doc, (
        "SITS9は、ドメイン間の表層的類似性（アナロジー）を超えて、"
        "「AとBは同一の命令セットの異なるレンダリングである」という"
        "構造的等価性を主張するための道具であり、"
        "その応用は暗号論的ステガノグラフィからワークショップ設計、"
        "文化的パレット分析にまで広がりうる。"
    ))

    # -- References ----------------------------------------------------------
    _add_heading(doc, "References", 1)
    refs = [
        "Essinger, J. (2004). Jacquard's Web: How a Hand-Loom Led to the Birth of the Information Age. Oxford University Press.",
        "Urton, G. (2003). Signs of the Inka Khipu: Binary Coding in the Andean Knotted-String Records. University of Texas Press.",
        "Stallings, W. (2017). Cryptography and Network Security: Principles and Practice, 7th ed. Pearson. [Ch. on steganography]",
        "Santoro, A. (2007). Sonic Fabric. https://sonicfabric.com/",
        "Reich, S. (1968). Music as a Gradual Process. [Minimalist manifesto on repetition and variation]",
        "McGee, H. (2004). On Food and Cooking: The Science and Lore of the Kitchen. Scribner.",
        "Tobin, J. & Dobard, R. (1999). Hidden in Plain View: A Secret Story of Quilts and the Underground Railroad. Anchor Books. [Contested thesis]",
        "McLuhan, M. (1964). Understanding Media: The Extensions of Man. McGraw-Hill.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        run_num = p.add_run(f"{i}. ")
        run_num.font.superscript = True
        run_num.font.size = Pt(9)
        run_text = p.add_run(ref)
        run_text.font.size = Pt(9)

    return doc


def main() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    doc = build_essay()
    out_path = OUT_DIR / "sits9_essay.docx"
    doc.save(str(out_path))
    print(f"Essay saved → {out_path}")


if __name__ == "__main__":
    main()

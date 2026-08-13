"""Compute per-cancer cases per specialist and the impact of false-positive follow-up visits.

Inputs:
- data/r05syobyo.pdf: MHLW Patient Survey 2023 (disease-specific patient counts).
- data/specialist_counts.csv: JMSB specialist counts (2023), via 日本専門医制度概報.
- data/cancer_to_specialty.csv: Mapping of simulated cancers to relevant specialty.
- output/by_cancer_and_followup.csv: simulation results.
- parameters.yaml: population and capacity assumptions.

Outputs:
- data/patient_survey_cancer_cases.csv: extracted cancer patient counts.
- output/specialist_capacity_impact.csv: baseline cases/specialist and FP-driven change.
- manuscript/specialist_capacity.docx: editable summary table.
"""

from __future__ import annotations

import argparse
import io
import re
import zipfile
from pathlib import Path
from typing import Any, Dict, List

import pandas as pd
import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

import pdfplumber


DATA_DIR = Path(__file__).parent / "data"
OUTPUT_DIR = Path(__file__).parent / "output"
MANUSCRIPT_DIR = Path(__file__).parent / "manuscript"
PARAMS_FILE = Path(__file__).parent / "parameters.yaml"

NDB_ZIP = DATA_DIR / "001711824.zip"
PATIENT_SURVEY_PDF = DATA_DIR / "r05syobyo.pdf"
PATIENT_SURVEY_CSV = DATA_DIR / "patient_survey_cancer_cases.csv"
SPECIALIST_COUNTS_CSV = DATA_DIR / "specialist_counts.csv"
CANCER_SPECIALTY_CSV = DATA_DIR / "cancer_to_specialty.csv"

POPULATION_2023 = 124_351_000

CANCER_PATTERNS: List[Dict[str, Any]] = [
    {
        "cancer": "Gastric",
        "table_pattern": r"胃の悪性新生物(?:＜腫瘍＞)?[^\n]{0,80}\n各年１０月",
    },
    {
        "cancer": "Colorectal",
        "table_pattern": r"結腸及び直腸の悪性新生物(?:＜腫瘍＞)?[^\n]{0,80}\n各年１０月",
    },
    {
        "cancer": "Lung",
        "table_pattern": r"気管，気管支及び肺の悪性新生物(?:＜腫瘍＞)?[^\n]{0,80}\n各年１０月",
    },
    {
        "cancer": "Breast",
        "table_pattern": r"乳房の悪性新生物(?:＜腫瘍＞)?[^\n]{0,80}\n各年１０月",
    },
    {
        "cancer": "Prostate",
        "table_pattern": r"前立腺の悪性新生物(?:＜腫瘍＞)?[^\n]{0,80}\n各年１０月",
    },
    {
        "cancer": "Liver",
        "table_pattern": r"肝及び肝内胆管の悪性新生物(?:＜腫瘍＞)?[^\n]{0,80}\n各年１０月",
    },
    {
        "cancer": "Pancreatic",
        "table_pattern": r"膵の悪性新生物(?:＜腫瘍＞)?[^\n]{0,80}\n各年１０月",
    },
    {
        "cancer": "Ovarian",
        "table_pattern": r"卵巣の悪性新生物(?:＜腫瘍＞)?[^\n]{0,80}\n各年１０月",
    },
]


def extract_patient_survey_counts(pdf_path: Path = PATIENT_SURVEY_PDF) -> pd.DataFrame:
    """Extract 2023 (Reiwa 5) patient counts for the target cancers from MHLW Patient Survey.

    Returns a DataFrame with one-day estimated patients (survey total), inpatient,
    outpatient, and the estimated total patient stock ("総患者数"), all in thousands.
    """
    rows = []
    with pdfplumber.open(str(pdf_path)) as pdf:
        text = "\n".join(page.extract_text() or "" for page in pdf.pages)

    line_pattern = re.compile(
        r"^令和5年\s+"
        r"([\d.]+)\s+"
        r"([\d.]+)\s+"
        r"([\d.]+)\s+"
        r"([\d.]+)\s+"
        r"([\d.]+)\s+"
        r"([\d,\s]+)",
        re.MULTILINE,
    )

    for item in CANCER_PATTERNS:
        table_match = re.search(item["table_pattern"], text)
        if not table_match:
            raise ValueError(f"Could not find disease table for {item['cancer']} in {pdf_path}")
        block = text[table_match.start() : table_match.start() + 3000]
        m = line_pattern.search(block)
        if not m:
            raise ValueError(f"Could not find Reiwa-5 row for {item['cancer']}")
        total, inpatient, outpatient, in_rate, out_rate, stock_raw = m.groups()
        stock = float(stock_raw.replace(",", "").replace(" ", "").strip())
        rows.append(
            {
                "cancer": item["cancer"],
                "estimated_patients_thousand": float(total),
                "inpatient_thousand": float(inpatient),
                "outpatient_thousand": float(outpatient),
                "inpatient_rate_per_100k": float(in_rate),
                "outpatient_rate_per_100k": float(out_rate),
                "total_patient_stock_thousand": stock,
            }
        )
    return pd.DataFrame(rows)


def read_ndb_outpatient_patients(zip_path: Path = NDB_ZIP) -> Dict[str, float]:
    """Return total unique NDB outpatient first-visit and revisit patient counts.

    Source: NDB Open Data 11th release, 01_医科診療行為/A_基本診療料/初再診料_診療月別患者数.xlsx
    The totals are the unique patients over the 12-month period (April 2024 - March 2025).
    """
    with zipfile.ZipFile(zip_path) as zf:
        target = None
        for info in zf.infolist():
            try:
                decoded = info.filename.encode("cp437").decode("cp932")
            except (UnicodeEncodeError, UnicodeDecodeError):
                decoded = info.filename
            if "初再診料" in decoded and "診療月別患者数" in decoded and "xlsx" in decoded:
                target = info.filename
                break
        if not target:
            raise FileNotFoundError(f"Could not find 初再診料 file in {zip_path}")
        with zf.open(target) as f:
            data = f.read()

    df = pd.read_excel(io.BytesIO(data), header=None)
    first = None
    revisit = None
    for i in range(len(df)):
        code = df.iloc[i, 0]
        if code == "A000":
            first = float(df.iloc[i, 5])
        elif code == "A001":
            revisit = float(df.iloc[i, 5])
    if first is None or revisit is None:
        raise ValueError("Could not parse A000/A001 totals from NDB outpatient file")
    return {"first_visit_patients": first, "revisit_patients": revisit}


def load_specialist_counts(path: Path = SPECIALIST_COUNTS_CSV) -> pd.DataFrame:
    return pd.read_csv(path)


def load_cancer_specialty_mapping(path: Path = CANCER_SPECIALTY_CSV) -> pd.DataFrame:
    return pd.read_csv(path)


def _patients_per_specialist(specialist_counts: pd.DataFrame, ndb: Dict[str, float]) -> float:
    """Return average annual unique outpatients per basic specialist from NDB counts."""
    total_outpatient_patients = ndb["first_visit_patients"] + ndb["revisit_patients"]
    basic = specialist_counts[specialist_counts["field"] == "basic"]
    total_basic_specialists = float(basic["count_2023"].sum())
    if total_basic_specialists <= 0:
        raise ValueError("Total basic specialist count must be positive")
    return total_outpatient_patients / total_basic_specialists


def compute_specialist_capacity_inputs(
    ndb: Dict[str, float],
    specialist_counts: pd.DataFrame,
    available_share: float,
    population: int = POPULATION_2023,
) -> Dict[str, float]:
    """Derive a national per-100k specialist-visit capacity from NDB patient counts and JMSB specialist counts.

    The logic is:
      1. NDB total outpatient patients / total basic specialists = baseline patients per specialist per year.
      2. Multiply by cancer-relevant specialists per 100k population.
      3. Apply the share assumed available for a new DTC cancer-workup wave.
    """
    patients_per_specialist = _patients_per_specialist(specialist_counts, ndb)

    mapping = load_cancer_specialty_mapping()
    relevant_specialties = set(mapping["specialty_en"])
    relevant_counts = specialist_counts[specialist_counts["specialty_en"].isin(relevant_specialties)]
    relevant_specialists = float(relevant_counts["count_2023"].sum())
    if relevant_specialists <= 0:
        raise ValueError("Relevant specialist count must be positive")

    relevant_per_100k = relevant_specialists / population * 100_000.0
    specialist_visits_per_year = patients_per_specialist * relevant_per_100k * available_share

    return {
        "total_outpatient_patients": ndb["first_visit_patients"] + ndb["revisit_patients"],
        "total_basic_specialists": float(specialist_counts[specialist_counts["field"] == "basic"]["count_2023"].sum()),
        "patients_per_specialist_per_year": patients_per_specialist,
        "relevant_specialists": relevant_specialists,
        "relevant_specialists_per_100k": relevant_per_100k,
        "available_for_cancer_share": available_share,
        "specialist_visits_per_year": specialist_visits_per_year,
        "population": population,
    }


PRIMARY_CARE_SPECIALTIES = {
    "Internal Medicine (Certified Physician)",
    "Internal Medicine (JMSB Certified)",
    "General Internal Medicine",
    "General Practice",
}


def compute_primary_care_capacity_inputs(
    ndb: Dict[str, float],
    specialist_counts: pd.DataFrame,
    available_share: float,
    population: int = POPULATION_2023,
) -> Dict[str, float]:
    """Derive a national per-100k primary-care-visit capacity from NDB patient counts and JMSB counts.

    The logic mirrors specialist capacity: average annual unique outpatients per basic
    specialist × primary-care-relevant specialists per 100k × available share.
    """
    patients_per_specialist = _patients_per_specialist(specialist_counts, ndb)

    pc_counts = specialist_counts[specialist_counts["specialty_en"].isin(PRIMARY_CARE_SPECIALTIES)]
    pc_specialists = float(pc_counts["count_2023"].sum())
    if pc_specialists <= 0:
        raise ValueError("Primary care specialist count must be positive")

    pc_per_100k = pc_specialists / population * 100_000.0
    primary_care_visits_per_year = patients_per_specialist * pc_per_100k * available_share

    return {
        "primary_care_specialists": pc_specialists,
        "primary_care_specialists_per_100k": pc_per_100k,
        "patients_per_specialist_per_year": patients_per_specialist,
        "available_for_cancer_share": available_share,
        "primary_care_visits_per_year": primary_care_visits_per_year,
        "population": population,
    }


def compute_specialist_capacity_impact(
    patient_counts: pd.DataFrame,
    specialist_counts: pd.DataFrame,
    mapping: pd.DataFrame,
    by_cancer: pd.DataFrame,
    follow_up_rate: float,
    population: int = POPULATION_2023,
) -> pd.DataFrame:
    """Compute baseline cases per specialist and the incremental change from FP follow-up visits.

    The baseline numerator is the MHLW Patient Survey "総患者数" (estimated total patients
    under continuous care for each cancer). The denominator is the relevant specialty count.
    """
    rows = []
    for _, c in mapping.iterrows():
        cancer = c["cancer"]
        specialty = c["specialty_en"]
        spec_count = float(specialist_counts.loc[specialist_counts["specialty_en"] == specialty, "count_2023"].iloc[0])
        spec_per_100k = spec_count / population * 100_000.0

        p = patient_counts[patient_counts["cancer"] == cancer].iloc[0]
        baseline_total_patients = p["total_patient_stock_thousand"] * 1_000.0
        baseline_cases_per_specialist = baseline_total_patients / spec_count

        sim = by_cancer[by_cancer["cancer"] == cancer]
        sim = sim[sim["follow_up_rate"] == follow_up_rate]
        if sim.empty:
            raise ValueError(f"No simulation result for {cancer} at follow-up rate {follow_up_rate}")
        sim_row = sim.iloc[0]
        # fp_specialist_visits is the false-positive-only specialist workload
        # (initial specialist visit + false-positive additional visits) per 100k screened.
        fp_visits_per_100k = float(sim_row["fp_specialist_visits"])
        fp_visits_per_specialist = fp_visits_per_100k / spec_per_100k
        new_cases_per_specialist = baseline_cases_per_specialist + fp_visits_per_specialist
        percent_change = (
            (fp_visits_per_specialist / baseline_cases_per_specialist) * 100.0
            if baseline_cases_per_specialist > 0
            else float("inf")
        )

        rows.append(
            {
                "cancer": cancer,
                "specialty": specialty,
                "specialist_count": spec_count,
                "baseline_total_patients": baseline_total_patients,
                "baseline_cases_per_specialist": baseline_cases_per_specialist,
                "follow_up_rate": follow_up_rate,
                "fp_specialist_visits_per_100k": fp_visits_per_100k,
                "fp_visits_per_specialist": fp_visits_per_specialist,
                "cases_per_specialist_with_fp": new_cases_per_specialist,
                "percent_change_in_cases_per_specialist": percent_change,
            }
        )
    return pd.DataFrame(rows)


def build_specialist_capacity_docx(df: pd.DataFrame, inputs: Dict[str, float], path: Path) -> None:
    """Write an editable English .docx table summarising the specialist capacity impact."""
    doc = Document()
    title = doc.add_heading("Specialist capacity impact of false-positive follow-up visits", level=1)
    p = doc.add_paragraph()
    p.add_run(
        f"Baseline capacity is derived from NDB Open Data outpatient patient counts "
        f"({inputs['total_outpatient_patients']:,.0f} first/revisit patients, "
        f"April 2024–March 2025) and JMSB specialist counts. "
        f"Cancer-relevant specialists = {inputs['relevant_specialists']:,.0f}; "
        f"baseline patients per specialist per year = {inputs['patients_per_specialist_per_year']:.1f}; "
        f"assumed share available for new cancer workups = {inputs['available_for_cancer_share']:.0%}."
    )

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = [
        "Cancer",
        "Specialty",
        "Baseline cases\nper specialist",
        "False-positive specialist\nvisits per specialist",
        "Cases per specialist\nwith FP",
        "Increase (%)",
    ]
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["cancer"])
        cells[1].text = str(row["specialty"])
        cells[2].text = f"{row['baseline_cases_per_specialist']:.1f}"
        cells[3].text = f"{row['fp_visits_per_specialist']:.2f}"
        cells[4].text = f"{row['cases_per_specialist_with_fp']:.1f}"
        cells[5].text = f"{row['percent_change_in_cases_per_specialist']:.1f}%"
        for cell in cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph(
        "Note: False-positive (FP) specialist visits are the additional specialist visits "
        "generated only by false-positive individuals in a 100,000-person DTC MCED screening "
        "wave at the stated follow-up rate (initial specialist visit plus FP-specific follow-up "
        "or reassurance visits). The 'baseline cases per specialist' uses the MHLW Patient Survey "
        "2023 total patient stock for each cancer."
    )
    doc.save(path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Compute specialist capacity impact")
    parser.add_argument("--params", type=Path, default=PARAMS_FILE)
    parser.add_argument("--output", type=Path, default=OUTPUT_DIR)
    parser.add_argument("--manuscript", type=Path, default=MANUSCRIPT_DIR)
    parser.add_argument("--follow-up-rate", type=float, default=0.5)
    args = parser.parse_args()

    args.output.mkdir(parents=True, exist_ok=True)
    args.manuscript.mkdir(parents=True, exist_ok=True)

    with open(args.params, "r", encoding="utf-8") as f:
        params = yaml.safe_load(f)

    population = int(params["simulation"].get("population", POPULATION_2023))
    available_share = float(params["assumptions"]["available_for_cancer_share"])

    # Extract patient-survey counts if the CSV does not already exist or if the PDF is newer.
    if not PATIENT_SURVEY_CSV.exists() or PATIENT_SURVEY_PDF.stat().st_mtime > PATIENT_SURVEY_CSV.stat().st_mtime:
        patient_counts = extract_patient_survey_counts()
        patient_counts.to_csv(PATIENT_SURVEY_CSV, index=False)
    else:
        patient_counts = pd.read_csv(PATIENT_SURVEY_CSV)

    specialist_counts = load_specialist_counts()
    mapping = load_cancer_specialty_mapping()

    # Compute capacity inputs used by parameters.yaml.
    ndb = read_ndb_outpatient_patients()
    inputs = compute_specialist_capacity_inputs(ndb, specialist_counts, available_share, population)
    with open(args.output / "specialist_capacity_inputs.yaml", "w", encoding="utf-8") as f:
        yaml.safe_dump(inputs, f, sort_keys=False, allow_unicode=True)

    # Compute FP impact.
    by_cancer = pd.read_csv(args.output / "by_cancer_and_followup.csv")
    impact = compute_specialist_capacity_impact(
        patient_counts, specialist_counts, mapping, by_cancer, args.follow_up_rate, population
    )
    impact.to_csv(args.output / "specialist_capacity_impact.csv", index=False)

    # Summary across all relevant specialists.
    total_baseline = (patient_counts["total_patient_stock_thousand"] * 1_000.0).sum()
    total_fp_100k = impact["fp_specialist_visits_per_100k"].sum()
    total_relevant = inputs["relevant_specialists"]
    summary = {
        "follow_up_rate": args.follow_up_rate,
        "screened_population_per_100k": 1,
        "total_relevant_specialists": total_relevant,
        "total_baseline_cancer_patients": total_baseline,
        "baseline_cases_per_specialist": total_baseline / total_relevant,
        "total_fp_specialist_visits_per_100k": total_fp_100k,
        "fp_visits_per_specialist": total_fp_100k / (total_relevant / population * 100_000.0),
    }
    summary["cases_per_specialist_with_fp"] = summary["baseline_cases_per_specialist"] + summary["fp_visits_per_specialist"]
    summary["percent_change"] = (
        summary["fp_visits_per_specialist"] / summary["baseline_cases_per_specialist"] * 100.0
        if summary["baseline_cases_per_specialist"] > 0
        else float("inf")
    )
    # Convert numpy scalar types to native Python types for YAML serialization.
    summary = {k: float(v) for k, v in summary.items()}
    with open(args.output / "specialist_capacity_summary.yaml", "w", encoding="utf-8") as f:
        yaml.safe_dump(summary, f, sort_keys=False, allow_unicode=True)

    build_specialist_capacity_docx(impact, inputs, args.manuscript / "specialist_capacity.docx")

    print("Specialist capacity impact")
    print(f"  Baseline cases per specialist: {summary['baseline_cases_per_specialist']:.1f}")
    print(f"  FP visits per specialist (100k screened, {args.follow_up_rate:.0%} follow-up): {summary['fp_visits_per_specialist']:.2f}")
    print(f"  Cases per specialist with FP: {summary['cases_per_specialist_with_fp']:.1f}")
    print(f"  Increase: {summary['percent_change']:.1f}%")
    print(f"  Outputs: {args.output / 'specialist_capacity_impact.csv'}")


if __name__ == "__main__":
    main()

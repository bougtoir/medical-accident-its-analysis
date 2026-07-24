"""
Generate manuscript for Journal on Baltic Security.
Format: Harvard referencing, UK English, 8000-10000 words, Abstract ~100 words.
Figures inserted inline after first mention + separate PNG files provided.
Editable PPTX also provided (one figure per slide).
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

OUTPUT_DIR = Path(__file__).parent.parent / "output"
OUTPUT_DIR.mkdir(exist_ok=True)

FIGURE_FILES = {
    1: OUTPUT_DIR / "fig1_concept.png",
    2: OUTPUT_DIR / "fig2_breach_delay.png",
    3: OUTPUT_DIR / "fig3_postconflict.png",
    4: OUTPUT_DIR / "fig4_entropy.png",
    5: OUTPUT_DIR / "fig5_tradeoff.png",
}

FIGURE_CAPTIONS = {
    1: ('Figure 1. ', 'Conceptual comparison of three information regimes. '
        '(a) No map (status quo): mine positions unknown, entire area uncertain. '
        '(b) Full intelligence: all real positions known to attacker. '
        '(c) MINE regime: map published with real + dummy positions; '
        'identity (real vs dummy) unknown to attacker.'),
    2: ('Figure 2. ', 'Attacker breach delay as a function of dummy ratio r '
        '(baseline scenario). Solid blue: MINE regime, scaling linearly '
        'with (1+r). Red dashed: blind sweep (no map, 600 min). '
        'Green dotted: full intelligence (12 min).'),
    3: ('Figure 3. ', 'Post-conflict clearance comparison. '
        '(a) Clearance time vs dummy ratio: map-based (blue) remains far '
        'below blind sweep (red dashed) at all r values. '
        '(b) Residual mines: map regime guarantees zero residual regardless '
        'of r; blind sweep leaves expected 10 mines.'),
    4: ('Figure 4. ', 'Positional entropy (Shannon information) representing '
        'attacker uncertainty. Purple: MINE regime. Red dashed: '
        'no-map maximum (1,934 bits). The regime provides controlled '
        'uncertainty between 0 and 50% of maximum at policy-relevant r values.'),
    5: ('Figure 5. ', 'Policy trade-off space: military utility (normalised '
        'breach delay, y-axis) vs post-conflict clearance cost (normalised, '
        'x-axis). Colour: dummy ratio r. The MINE regime traces a curve '
        'from full information (origin) toward the status quo (upper right), '
        'with each r value representing a specific policy choice.'),
}


def insert_figure(doc, fig_num):
    """Insert a figure image and caption inline in the document."""
    fpath = FIGURE_FILES[fig_num]
    label, caption_text = FIGURE_CAPTIONS[fig_num]

    if fpath.exists():
        # Add the image
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_img.add_run()
        run.add_picture(str(fpath), width=Inches(5.5))

    # Add caption with spacing
    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.space_before = Pt(6)
    p_cap.paragraph_format.space_after = Pt(12)
    p_cap.add_run(label).bold = True
    p_cap.add_run(caption_text)


def create_manuscript():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.0

    # ─── Title Page ───────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(
        'The MINE Framework (Mapping Is Not Exposure): '
        'Transparent Minefield Maps with Dummy Positions '
        'for Reconciling Military Utility and Post-Conflict Safety'
    )
    title_run.bold = True
    title_run.font.size = Pt(16)

    doc.add_paragraph()
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_para.add_run('Tatsuki Onishi').bold = True

    affil_para = doc.add_paragraph()
    affil_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affil_para.add_run('[Affiliation]').italic = True

    doc.add_paragraph()

    # ─── Abstract ─────────────────────────────────────────────────────────────
    doc.add_heading('Abstract', level=1)
    p = doc.add_paragraph()
    p.add_run(
        'This paper proposes the MINE framework (Mapping Is Not Exposure): '
        'a regulatory regime for anti-personnel mines wherein defending states '
        'publish position maps containing both real mine locations and '
        'indistinguishable dummy positions. The dummy-to-real ratio r serves '
        'as a treaty-tuneable parameter balancing military effectiveness '
        'against humanitarian cost. Game-theoretic analysis demonstrates '
        'that rational attackers will clear all marked positions for any '
        'policy-relevant r, preserving full breach delay. The zero-residual '
        'property guarantees complete post-conflict clearance without reliance '
        'on detection technology. Numerical simulations calibrated to Baltic '
        'frontier scenarios illustrate practical application for NATO states '
        'now withdrawing from the Ottawa Convention.'
    )

    doc.add_paragraph()
    kw_para = doc.add_paragraph()
    kw_para.add_run('Keywords: ').bold = True
    kw_para.add_run(
        'anti-personnel mines; MINE framework; Ottawa Convention; game theory; '
        'area denial; post-conflict clearance; Baltic security; information '
        'asymmetry; treaty design; NATO; deterrence'
    )

    doc.add_page_break()

    # ─── 1. Introduction ─────────────────────────────────────────────────────
    doc.add_heading('1. Introduction', level=1)

    p = doc.add_paragraph()
    p.add_run(
        'The 1997 Convention on the Prohibition of Anti-Personnel Mines '
        '(hereafter the Ottawa Convention) represented a landmark achievement '
        'in international humanitarian law, securing commitments from over '
        '160 states to eliminate weapons that kill and maim civilians long '
        'after conflicts end (ICRC, 2024). For more than two decades, the '
        'treaty\u2019s absolute prohibition created a powerful international norm. '
        'Yet the strategic landscape of northeastern Europe has shifted '
        'dramatically since Russia\u2019s full-scale invasion of Ukraine in 2022, '
        'and the norm is now fracturing at precisely the point where it is '
        'most needed.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'In March 2025, Poland, Estonia, Latvia, and Lithuania jointly '
        'announced their withdrawal from the Ottawa Convention, citing the '
        'need to deter Russian aggression along NATO\u2019s eastern flank. '
        'Finland followed in April 2025, with Prime Minister Petteri Orpo '
        'stating that \u201cwithdrawing from the Ottawa Convention will give us '
        'the possibility to prepare for the changes in the security environment '
        'in a more versatile way\u201d (Reuters, 2025). Each withdrawing state '
        'simultaneously affirmed its commitment to \u201cresponsible use\u201d of '
        'anti-personnel mines\u2014a commitment whose specific operational '
        'content remains undefined.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'This paper fills that gap. We propose a novel regulatory framework '
        'that we term '
    )
    p.add_run('MINE (Mapping Is Not Exposure)').bold = True
    p.add_run(
        '\u2014reflecting the core insight that publishing minefield positions '
        'with dummy entries does not constitute military exposure. The MINE '
        'framework gives concrete operational meaning to \u201cresponsible use.\u201d '
        'Under this framework, a defending state must publish a map showing '
        'all mine positions before or during deployment. Crucially, the map '
        'may include dummy (decoy) positions: locations marked as potentially '
        'mined but containing no actual ordnance. The ratio of dummy to real '
        'positions (the '
    )
    p.add_run('dummy ratio r').italic = True
    p.add_run(
        ') becomes a treaty-regulated parameter that quantitatively indexes '
        'the trade-off between military effectiveness and post-conflict '
        'safety.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'The core insight is elegant in its simplicity. Under the current '
        'regime with no published map, the dummy ratio is effectively infinite: '
        'the entire terrain is uncertain from the attacker\u2019s perspective. '
        'If an adversary obtains complete intelligence through espionage, '
        'the effective dummy ratio falls to zero\u2014only real positions are '
        'relevant. Our proposal targets the continuum between these extremes, '
        'offering a family of regimes indexed by r \u2208 [0, \u221e) that '
        'trade off military utility against humanitarian cost in a '
        'quantifiable, verifiable, and negotiable manner.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'We demonstrate three key results. First, the defending minefield '
        'retains substantial delay capability because the attacker must clear '
        'all M = N(1+r) marked positions in any breach lane, regardless of '
        'whether each is real or dummy. Delay scales linearly with (1+r). '
        'Second, post-conflict clearance is '
    )
    p.add_run('guaranteed').italic = True
    p.add_run(
        ' to be complete: by systematically addressing every position on the '
        'map, zero residual mines remain\u2014a deterministic property that no '
        'probabilistic detection method can match. Third, the framework is '
        'inherently verifiable: mines found at unregistered positions '
        'constitute clear treaty violations, and the published dummy ratio '
        'is testable ex post.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'The urgency of this contribution is underscored by the current '
        'policy vacuum. As of mid-2025, five NATO states have announced '
        'withdrawal from the Ottawa Convention, all pledging responsible use, '
        'yet no common framework exists to define what this means. Without '
        'agreed operational standards, each withdrawing state will develop '
        'its own practices\u2014potentially divergent, potentially opaque, and '
        'potentially contributing to the very humanitarian outcomes the '
        'Ottawa Convention sought to prevent. The MINE framework '
        'offers a ready-made standard that could be adopted immediately.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'The remainder of this paper is structured as follows. Section 2 '
        'reviews existing regulatory frameworks and their limitations. '
        'Section 3 presents the formal model. Section 4 analyses military '
        'effectiveness under the MINE regime. Section 5 examines '
        'post-conflict safety guarantees. Section 6 develops game-theoretic '
        'and information-theoretic properties. Section 7 discusses treaty '
        'design implications, including verification mechanisms and '
        'compatibility with existing frameworks. Section 8 applies the '
        'model to the Baltic security context with scenario analysis. '
        'Section 9 discusses limitations, and Section 10 concludes with '
        'policy recommendations.'
    )

    # ─── 2. Background ───────────────────────────────────────────────────────
    doc.add_heading('2. Background and Existing Frameworks', level=1)

    doc.add_heading('2.1 The Ottawa Convention and the 2024\u20132025 withdrawal wave', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'The Ottawa Convention established a categorical norm against '
        'anti-personnel landmines. Its success is undeniable: global '
        'production fell dramatically, stockpiles were destroyed across '
        'dozens of countries, and millions of square metres of contaminated '
        'land were cleared (ICBL, 2023). However, the Convention\u2019s absolute '
        'prohibition creates a binary choice between complete compliance '
        'and complete withdrawal, with no intermediate options for states '
        'facing genuine defensive requirements.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'The wave of withdrawals beginning in late 2024 demonstrates the '
        'fragility of this binary structure when security conditions change. '
        'Several factors converged: Russia\u2019s extensive use of mines in '
        'Ukraine (both anti-personnel and anti-vehicle), the United States\u2019 '
        'decision to supply Ukraine with anti-personnel mines, and the '
        'growing recognition among NATO\u2019s eastern members that conventional '
        'area denial is critical to compensating for force imbalances '
        'vis-\u00e0-vis Russia.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Finland\u2019s case is illustrative. Having destroyed over one million '
        'mines after ratification in 2012, Finland announced withdrawal '
        'thirteen years later. Defence Minister Antti H\u00e4kk\u00e4nen commissioned '
        'an evaluation of anti-personnel mines as \u201ca strengthening factor '
        'for Finland\u2019s defence\u201d (Reuters, 2024). The Finnish Defence Forces '
        'commander publicly stated that the topic required renewed discussion '
        'given Russia\u2019s demonstrated willingness to use mines offensively. '
        'Estonia\u2019s National Defence Committee similarly debated withdrawal '
        'in November 2024, with the Defence Forces commander noting that '
        '\u201csufficient alternatives to passive mines\u201d existed but recommending '
        'the option remain open (ERR News, 2024).'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Every withdrawing state has articulated a commitment to '
        '\u201cresponsible\u201d or \u201cdefensive\u201d mine use. Finland\u2019s agriculture '
        'minister stated: \u201cFinland will use mines in a responsible way, '
        'but it\u2019s a deterrent we need\u201d (Reuters, 2025). Yet none has '
        'specified what \u201cresponsible use\u201d means operationally\u2014what '
        'constraints distinguish responsible from irresponsible employment. '
        'This paper provides exactly such a specification.'
    )

    doc.add_heading('2.2 CCW Amended Protocol II and its limitations', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'The Amended Protocol II to the Convention on Certain Conventional '
        'Weapons (1996) represents the existing attempt at a middle ground '
        'between prohibition and unrestricted use. Its key provisions for '
        'anti-personnel mines include: prohibition of use outside marked '
        'and monitored areas unless equipped with self-destruct/self-deactivation '
        'mechanisms; requirements for detectability using standard equipment; '
        'and mandatory recording of pre-planned minefield locations for '
        'post-conflict sharing (UNODA, 2023).'
    )

    p = doc.add_paragraph()
    p.add_run(
        'However, the Protocol has three significant limitations that the '
        'MINE framework addresses:'
    )

    p = doc.add_paragraph()
    p.add_run('First, self-destruct unreliability. ').bold = True
    p.add_run(
        'The Protocol\u2019s requirement for self-destruct mechanisms assumes '
        'reliable functioning. Field evidence tells a different story. '
        'The Norwegian People\u2019s Aid (2007) study of the M85 submunition\u2014'
        'widely cited as incorporating a \u201chigh-quality\u201d self-destruct '
        'mechanism\u2014found field failure rates substantially exceeding '
        'laboratory specifications. The CCW Protocol itself permits a 10% '
        'failure rate. In a field of 1,000 mines, this leaves 100 active '
        'mines\u2014each capable of killing or maiming for decades. The '
        'MINE framework does not depend on mechanical reliability; '
        'its safety guarantee is deterministic.'
    )

    p = doc.add_paragraph()
    p.add_run('Second, temporal uncertainty. ').bold = True
    p.add_run(
        'Self-destruct creates a period during which any given mine may or '
        'may not be active. Civilians cannot know whether a mine has '
        'self-destructed without physically verifying\u2014precisely the dangerous '
        'action the mechanism is meant to prevent. The transparent map provides '
        'spatial certainty at all times: every mine is at a published position.'
    )

    p = doc.add_paragraph()
    p.add_run('Third, post-conflict timing. ').bold = True
    p.add_run(
        'The Protocol\u2019s recording requirement (Article 7) applies only after '
        'cessation of hostilities. During conflict, mine positions may be lost '
        'due to destroyed records, personnel casualties, or deliberate '
        'concealment. The MINE framework requires publication '
        'during or before deployment, creating an irrevocable record.'
    )

    doc.add_heading('2.3 Military doctrine on dummy minefields', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'The concept of dummy (phony) minefields is well-established in '
        'military doctrine and provides doctrinal precedent for the proposed '
        'framework. US Army Field Manual FM 90-2 (Battlefield Deception) '
        'states that \u201cdummy minefields can be used very effectively in the '
        'retrograde to slow and canalise the enemy attack or cause the enemy '
        'to mass his forces\u201d (US Army, 1988). FM 5-102 (Countermobility) '
        'defines phony minefields as \u201careas of ground used to simulate live '
        'minefields and deceive the enemy\u201d and notes they \u201care of no value '
        'until the enemy has become sensitive to mine warfare\u201d (US Army, 1985).'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Two doctrinal principles are particularly relevant. First, US doctrine '
        'requires that dummy minefields receive \u201cthe same authorisation, '
        'recording, and reporting procedures as the type minefield it is '
        'designed to replicate.\u201d This confirms that military practice already '
        'treats dummy and real positions equivalently for administrative '
        'purposes. Second, the doctrine notes that dummies are most effective '
        '\u201cwhen mixed with real ones throughout the battlefield\u201d\u2014precisely '
        'the configuration our framework mandates.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Our contribution inverts the doctrinal logic. Rather than using '
        'dummies to '
    )
    p.add_run('conceal').italic = True
    p.add_run(
        ' the existence of minefields from the attacker, we use a published '
        'map with dummies to '
    )
    p.add_run('reveal').italic = True
    p.add_run(
        ' sufficient information for guaranteed post-conflict safety whilst '
        'preserving operational uncertainty for the attacker.'
    )

    doc.add_heading('2.4 Related game-theoretic literature', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'Game-theoretic models of military deception have addressed '
        'attacker\u2013defender interactions under information asymmetry in several '
        'relevant contexts. Zhuang, Bier, and Alagoz (2010) model secrecy and '
        'deception in a multiple-period signalling game, classifying defender '
        'signals as truthful disclosure, secrecy, or deception. Their framework '
        'shows that defenders can achieve more cost-effective security through '
        'strategic information management\u2014a principle our framework extends to '
        'the landmine context with the additional constraint of post-conflict '
        'safety.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Hendricks (2006) analyses feints in attack\u2013defence models, '
        'demonstrating that equilibrium behaviour depends critically on '
        'signalling technology. In naval mine warfare, Kanazawa, Iida, and '
        'Morimoto (2018) formulate mine countermeasure operations as a '
        'two-player game between a minefield planner and clearance forces, '
        'showing how countermeasure choices (hunting vs sweeping) interact '
        'with mine design features.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'The Aalto University network interdiction framework (Aalto NDU, 2020) '
        'models defensive minefield placement as a mixed-integer optimisation '
        'problem, seeking mine configurations that maximise disruption to '
        'attacker movement networks. While focused on placement optimisation '
        'rather than information disclosure, this work provides complementary '
        'methodology for determining optimal spatial configurations within '
        'our framework.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Our framework differs fundamentally from this literature in one '
        'respect: we model a regime where the defender '
    )
    p.add_run('voluntarily').italic = True
    p.add_run(
        ' publishes position information under treaty obligation, with the '
        'quality of that information (the dummy ratio) as the strategic '
        'variable. This transforms the problem from one of deception '
        '(maximising misinformation) to one of '
    )
    p.add_run('regulated transparency').italic = True
    p.add_run(
        ' (calibrating information disclosure to balance competing objectives). '
        'To our knowledge, no prior work has formalised this concept.'
    )

    # ─── 3. Model Formulation ─────────────────────────────────────────────────
    doc.add_heading('3. Model Formulation', level=1)

    doc.add_heading('3.1 Setup and notation', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'Consider a defending state that emplaces a minefield along a front '
        'of width W (metres) to a depth D (metres), with total area '
        'A = W \u00d7 D. The minefield contains N real anti-personnel mines '
        'distributed across this area. Under the MINE regime, the defender '
        'publishes a map to the MINE Registry showing M = N(1 + r) '
        'marked positions, where r \u2265 0 is the dummy ratio. Of these, '
        'exactly N contain real mines and Nr are dummies.'
    )

    p = doc.add_paragraph()
    p.add_run('Fundamental constraint: ').bold = True
    p.add_run(
        'Every real mine position must appear on the map. The defender may '
        'not place mines at undisclosed locations. This constraint is the '
        'linchpin of the entire framework: it enables the post-conflict '
        'safety guarantee by ensuring that clearing all M published positions '
        'is sufficient to eliminate every real mine. Without this constraint, '
        'the framework degenerates into the status quo of unrestricted mining.'
    )

    p = doc.add_paragraph()
    p.add_run('Indistinguishability condition: ').bold = True
    p.add_run(
        'The published map does not distinguish real from dummy positions. '
        'All M markers appear identical to any observer. The attacker '
        'knows that exactly N of the M positions are real (the ratio r '
        'being a public treaty parameter), but cannot identify which '
        'specific positions contain real mines. On the ground, dummy '
        'positions may be prepared identically to real ones (disturbed '
        'earth, surface markers) to prevent visual or sensor-based '
        'discrimination.'
    )

    p = doc.add_paragraph()
    p.add_run('Table 1 ').bold = True
    p.add_run('summarises the notation.')

    table = doc.add_table(rows=13, cols=2)
    table.style = 'Table Grid'
    headers = table.rows[0].cells
    headers[0].text = 'Symbol'
    headers[1].text = 'Definition'
    notation = [
        ('N', 'Number of real mines'),
        ('r', 'Dummy ratio (treaty-regulated parameter, r \u2265 0)'),
        ('M = N(1+r)', 'Total marked positions on published map'),
        ('W', 'Front width (metres)'),
        ('D', 'Minefield depth (metres)'),
        ('A = W\u00d7D', 'Total minefield area (m\u00b2)'),
        ('w', 'Breach lane width (metres)'),
        ('t_c', 'Time to clear one marked position (minutes)'),
        ('t_p', 'Time to probe one grid cell in blind sweep (minutes)'),
        ('k', 'Number of parallel clearance teams'),
        ('\u03b4', 'Grid spacing for blind detection sweep (metres)'),
        ('p_d', 'Per-mine detection probability in blind sweep'),
    ]
    for i, (sym, defn) in enumerate(notation):
        row = table.rows[i + 1].cells
        row[0].text = sym
        row[1].text = defn

    doc.add_paragraph()

    doc.add_heading('3.2 Three information regimes', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'We compare three information regimes from the attacker\u2019s perspective, '
        'corresponding to different values of the effective dummy ratio '
        '(Figure 1):'
    )

    p = doc.add_paragraph()
    p.add_run('Regime A \u2013 No map (status quo, r \u2192 \u221e): ').bold = True
    p.add_run(
        'The attacker has no knowledge of mine positions. The entire minefield '
        'area is uncertain. To breach safely, the attacker must conduct a '
        'systematic blind sweep of the breach lane, probing each grid cell '
        'of area \u03b4\u00b2. This is the current situation for NATO forces facing '
        'Russian minefields in Ukraine and the default for any non-signatory '
        'deploying mines without disclosure.'
    )

    p = doc.add_paragraph()
    p.add_run('Regime B \u2013 Full intelligence (r = 0): ').bold = True
    p.add_run(
        'The attacker knows the exact position of every real mine. This '
        'could arise through espionage, signals intelligence, or post-conflict '
        'record sharing. The attacker clears only the N real positions in '
        'their path, achieving minimum clearance effort.'
    )

    p = doc.add_paragraph()
    p.add_run('Regime C \u2013 MINE regime (0 < r < \u221e): ').bold = True
    p.add_run(
        'The attacker has the published map showing M = N(1+r) positions. '
        'They know N mines exist among the M positions but cannot distinguish '
        'real from dummy. To breach safely, they must clear all marked '
        'positions in their lane.'
    )

    insert_figure(doc, 1)

    doc.add_heading('3.3 Breach delay model', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'An attacking force seeks to breach the minefield through a lane '
        'of width w. Assuming positions are uniformly distributed across '
        'the minefield area, the number of marked positions within the '
        'breach lane is:'
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('m(r) = N(1+r) \u00d7 (w/W)              (1)')

    p = doc.add_paragraph()
    p.add_run(
        'With k parallel clearance teams, each taking t_c minutes per '
        'position, the breach delay under the map regime is:'
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('T_map(r) = N(1+r)(w/W) \u00d7 t_c / k              (2)')

    p = doc.add_paragraph()
    p.add_run(
        'For comparison, blind sweep delay (Regime A) requires probing '
        'every grid cell in the breach lane:'
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('T_blind = (w \u00d7 D / \u03b4\u00b2) \u00d7 t_p / k              (3)')

    p = doc.add_paragraph()
    p.add_run(
        'And full intelligence delay (Regime B):'
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('T_intel = N(w/W) \u00d7 t_c / k              (4)')

    p = doc.add_paragraph()
    p.add_run(
        'The delay multiplier of the transparent map over full intelligence '
        'is simply (1+r), confirming that each dummy position imposes the '
        'same operational cost as a real mine from the attacker\u2019s perspective.'
    )

    doc.add_heading('3.4 Post-conflict clearance model', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'After hostilities, the clearance task for the full minefield area differs '
        'fundamentally between regimes:'
    )

    p = doc.add_paragraph()
    p.add_run('Map-based clearance (MINE regime): ').bold = True
    p.add_run(
        'Clear all M = N(1+r) positions. Time = M \u00d7 t_c / k hours. '
        'Residual mines after completion = 0 (deterministic guarantee).'
    )

    p = doc.add_paragraph()
    p.add_run('Blind sweep (no-map regime): ').bold = True
    p.add_run(
        'Sweep entire area A cell by cell. Time = (A/\u03b4\u00b2) \u00d7 t_p / k hours. '
        'Residual mines = N \u00d7 (1 \u2212 p_d) per sweep pass. Even at '
        'p_d = 0.95 (state-of-the-art metal detection), a field of 200 mines '
        'leaves an expected 10 undetected\u2014each a potential future civilian '
        'casualty for decades.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'The clearance efficiency ratio\u2014time saved by using the map\u2014is:'
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('\u03b7(r) = T_blind / T_map = (A \u00d7 t_p) / '
              '[\u03b4\u00b2 \u00d7 N(1+r) \u00d7 t_c]              (5)')

    p = doc.add_paragraph()
    p.add_run(
        'For the baseline parameters detailed in Section 4, \u03b7 ranges from '
        '50\u00d7 (r=0) to 2.4\u00d7 (r=20), confirming that map-based clearance '
        'remains substantially more efficient than blind sweep across all '
        'plausible dummy ratios.'
    )

    # ─── 4. Military Effectiveness Analysis ───────────────────────────────────
    doc.add_heading('4. Military Effectiveness Analysis', level=1)

    doc.add_heading('4.1 Baseline scenario parameters', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'We evaluate the model using parameters representative of a Baltic '
        'frontier defence scenario (Table 2). The scenario models a '
        'battalion-level defensive minefield along a company sector of a '
        '5 km front, with the attacker attempting to breach a 100 m-wide '
        'lane for mechanised advance.'
    )

    table2 = doc.add_table(rows=9, cols=3)
    table2.style = 'Table Grid'
    hdr = table2.rows[0].cells
    hdr[0].text = 'Parameter'
    hdr[1].text = 'Value'
    hdr[2].text = 'Justification'
    params_data = [
        ('N (real mines)', '200', 'Battalion-level minefield (NATO STANAG)'),
        ('W (front width)', '5,000 m', 'Company defensive sector width'),
        ('D (depth)', '300 m', 'Standard minefield depth'),
        ('w (breach lane)', '100 m', 'Mechanised infantry requirement'),
        ('t_c (clear/position)', '30 min', 'Manual EOD clearance standard'),
        ('t_p (probe/cell)', '5 min', 'Systematic detection sweep rate'),
        ('\u03b4 (grid spacing)', '5 m', 'Metal detector effective width'),
        ('k (teams)', '10', 'Engineer platoon capacity'),
    ]
    for i, (param, val, src) in enumerate(params_data):
        row = table2.rows[i + 1].cells
        row[0].text = param
        row[1].text = val
        row[2].text = src

    p = doc.add_paragraph()
    p.add_run('\nTable 2. ').bold = True
    p.add_run('Baseline scenario parameters for Baltic frontier defence.')

    doc.add_heading('4.2 Delay results', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'Table 3 presents breach delay across dummy ratios. Figure 2 '
        'visualises the relationship graphically.'
    )

    # Table 3: Results
    table3 = doc.add_table(rows=8, cols=5)
    table3.style = 'Table Grid'
    hdr = table3.rows[0].cells
    hdr[0].text = 'Dummy ratio (r)'
    hdr[1].text = 'Total positions (M)'
    hdr[2].text = 'Positions in lane'
    hdr[3].text = 'Breach delay'
    hdr[4].text = 'Ratio to blind'
    data = [
        ('0', '200', '4', '12 min', '0.02'),
        ('1', '400', '8', '24 min', '0.04'),
        ('3', '800', '16', '48 min', '0.08'),
        ('5', '1,200', '24', '72 min (1.2 h)', '0.12'),
        ('10', '2,200', '44', '132 min (2.2 h)', '0.22'),
        ('20', '4,200', '84', '252 min (4.2 h)', '0.42'),
        ('Blind (no map)', '\u2014', '1,200 cells', '600 min (10 h)', '1.00'),
    ]
    for i, row_data in enumerate(data):
        row = table3.rows[i + 1].cells
        for j, val in enumerate(row_data):
            row[j].text = val

    p = doc.add_paragraph()
    p.add_run('\nTable 3. ').bold = True
    p.add_run('Breach delay under different dummy ratios (baseline scenario).')

    insert_figure(doc, 2)

    p = doc.add_paragraph()
    p.add_run(
        'Several observations emerge. At r = 0, the map reveals all real '
        'positions and the attacker clears only 4 positions in the lane '
        '(12 minutes). This is the minimum delay and represents the lower '
        'bound of military effectiveness. At r = 10, breach delay reaches '
        '132 minutes\u2014sufficient for significant defensive preparation, '
        'repositioning, or calling reinforcements. The blind sweep baseline '
        '(600 minutes) represents the theoretical maximum, though in practice '
        'attackers facing blind minefields may choose to accept casualties '
        'rather than conducting a full 10-hour sweep\u2014a point we return to '
        'in Section 6.'
    )

    doc.add_heading('4.3 Force multiplication and defensive integration', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'The delay values reported above represent mine clearance time alone. '
        'In combined-arms defensive operations, minefields are integrated with '
        'direct and indirect fire systems, meaning that clearance is conducted '
        'under fire. This interaction creates a force multiplication effect: '
        'each hour of minefield delay is an additional hour during which '
        'defensive fires can attrit the attacking force. NATO doctrine '
        'estimates that obstacles integrated with fire systems multiply '
        'their effectiveness by a factor of 3\u20135 compared to unobserved '
        'obstacles (unresisted clearance).'
    )

    p = doc.add_paragraph()
    p.add_run(
        'In the Baltic context, this has particular significance. The delay '
        'imposed by transparent minefields directly purchases time for: '
        'mobilisation of reserve forces; deployment of NATO rapid reaction '
        'forces (which require 10\u201330 days for full deployment); preparation '
        'of subsequent defensive positions; and evacuation of civilian '
        'populations from threatened areas. For Estonia, where the capital '
        'Tallinn is approximately 3 hours of mechanised advance from the '
        'eastern border, each additional hour of delay at the border '
        'represents a significant fraction of the total available '
        'preparation time.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Furthermore, the transparent map enables defensive force planning '
        'that is unavailable with conventional minefields. Because the '
        'defender knows which positions are real and which are dummy, '
        'defensive fire can be concentrated on real mine positions '
        '(where the attacker will be forced to stop and clear) whilst '
        'dummy positions serve as engagement areas where the attacker '
        'commits resources unnecessarily. This asymmetry of information '
        'between defender (knows which are real) and attacker (does not) '
        'creates tactical advantages beyond simple delay.'
    )

    doc.add_heading('4.4 The accept-casualties alternative', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'An important consideration is the attacker\u2019s alternative to clearing: '
        'simply advancing through the minefield and accepting casualties. This '
        'is not merely theoretical\u2014Russian forces in Ukraine have repeatedly '
        'driven vehicles and infantry through known minefields, accepting '
        'significant losses to maintain operational tempo.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Under the no-map regime, an attacker who chooses to accept casualties '
        'faces an unknown risk distributed over an unknown area. Under the '
        'MINE regime, the same attacker faces '
    )
    p.add_run('concentrated, known-location risk').italic = True
    p.add_run(
        '. Each marked position has probability 1/(1+r) of being lethal. '
        'The key insight is that published positions create '
    )
    p.add_run('unavoidable obstacles').italic = True
    p.add_run(
        '\u2014the attacker cannot route around them without adding distance, '
        'and cannot ignore them without accepting quantifiable risk. '
        'This is explored formally in the game-theoretic analysis of Section 6.'
    )

    # ─── 5. Post-Conflict Safety ─────────────────────────────────────────────
    doc.add_heading('5. Post-Conflict Safety Guarantee', level=1)

    doc.add_heading('5.1 The zero-residual property', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'The most significant humanitarian advantage of the MINE '
        'regime is what we term the '
    )
    p.add_run('zero-residual property').italic = True
    p.add_run(
        ': after hostilities, clearing all M published positions is both '
        'necessary and sufficient to remove every real mine. This is a '
        'deterministic guarantee, not a probabilistic one. It does not '
        'depend on detection equipment performance, soil conditions, '
        'vegetation cover, or any other variable that affects conventional '
        'mine clearance.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Formally: let S = {s_1, ..., s_M} denote the set of published '
        'positions, and R \u2286 S (|R| = N) the subset containing real mines. '
        'If a clearance team visits and neutralises every position in S, '
        'then every element of R is necessarily neutralised. The guarantee '
        'follows from the fundamental constraint that R \u2286 S.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'This contrasts sharply with current post-conflict clearance, '
        'which relies on imperfect detection applied over vast, uncertain '
        'areas. With detection probability p_d = 0.95 per pass, a field '
        'of 200 mines leaves an expected 10 mines after one pass, 0.5 after '
        'two passes, and 0.025 after three passes. Multiple passes reduce '
        'but never eliminate residual risk\u2014and each additional pass incurs '
        'the full time cost of sweeping the entire area.'
    )

    doc.add_heading('5.2 Clearance time comparison', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'Figure 3 compares post-conflict clearance times and residual risk. '
        'For the baseline scenario:'
    )

    p = doc.add_paragraph()
    p.add_run('\u2022 Blind sweep: ').bold = True
    p.add_run(
        '500 hours (60,000 cells \u00d7 5 min / 10 teams / 60), '
        'residual: 10 mines expected.'
    )
    p = doc.add_paragraph()
    p.add_run('\u2022 Map at r = 3: ').bold = True
    p.add_run('40 hours, residual: 0 mines. Efficiency gain: 12.5\u00d7.')
    p = doc.add_paragraph()
    p.add_run('\u2022 Map at r = 5: ').bold = True
    p.add_run('60 hours, residual: 0 mines. Efficiency gain: 8.3\u00d7.')
    p = doc.add_paragraph()
    p.add_run('\u2022 Map at r = 10: ').bold = True
    p.add_run('110 hours, residual: 0 mines. Efficiency gain: 4.5\u00d7.')
    p = doc.add_paragraph()
    p.add_run('\u2022 Map at r = 20: ').bold = True
    p.add_run('210 hours, residual: 0 mines. Efficiency gain: 2.4\u00d7.')

    p = doc.add_paragraph()
    p.add_run(
        'Even at r = 20 (twenty dummy positions per real mine), map-based '
        'clearance is 2.4 times faster than blind sweep and guarantees '
        'complete safety. The zero-residual property is independent of r.'
    )

    insert_figure(doc, 3)

    doc.add_heading('5.3 Economic comparison', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'The economic implications of the zero-residual property extend '
        'beyond clearance time. Mine-contaminated land cannot be used '
        'productively until cleared. At current clearance costs of '
        'approximately USD 500\u20131,500 per cleared mine (including search '
        'time, not just neutralisation), the economic comparison is stark. '
        'For our baseline scenario (200 mines, 1.5 km\u00b2 minefield):'
    )

    p = doc.add_paragraph()
    p.add_run('\u2022 Blind clearance (no map): ').bold = True
    p.add_run(
        'Cost = 60,000 cells \u00d7 search cost per cell \u2248 USD 3\u20136 million. '
        'Duration: months to years for full area. Land unavailable throughout.'
    )
    p = doc.add_paragraph()
    p.add_run('\u2022 Map-based clearance (r = 5): ').bold = True
    p.add_run(
        'Cost = 1,200 positions \u00d7 USD 500 = USD 600,000. '
        'Duration: 60 hours (less than 3 days). Land released immediately after.'
    )
    p = doc.add_paragraph()
    p.add_run('\u2022 Economic savings: ').bold = True
    p.add_run(
        'USD 2.4\u20135.4 million per minefield in direct clearance costs alone, '
        'plus the economic value of rapid land release.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'For a country like Estonia, which might deploy 5\u201310 transparent '
        'minefields along its border in a crisis, the aggregate savings of '
        'map-based versus blind clearance could reach USD 20\u201350 million\u2014'
        'a substantial sum for a small economy. More importantly, the '
        'certainty of the clearance timeline (days rather than years) '
        'enables reliable post-conflict reconstruction planning.'
    )

    doc.add_heading('5.4 Comparison with self-destruct mechanisms', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'Table 4 compares the MINE regime with self-destruct '
        'mechanisms across key safety dimensions.'
    )

    table4 = doc.add_table(rows=6, cols=3)
    table4.style = 'Table Grid'
    hdr = table4.rows[0].cells
    hdr[0].text = 'Dimension'
    hdr[1].text = 'Self-destruct (CCW Protocol II)'
    hdr[2].text = 'MINE framework (proposed)'
    comp_data = [
        ('Residual risk', 'N \u00d7 failure_rate (2\u201310%)', 'Zero (deterministic)'),
        ('Depends on reliability', 'Yes (mechanical)', 'No (informational)'),
        ('Verifiable', 'No (cannot confirm remotely)', 'Yes (map is public record)'),
        ('Time to safety', 'Self-destruct timer + uncertainty', 'Clearance time (known)'),
        ('Cost scaling', 'Per-mine mechanism cost', 'Map publication (negligible)'),
    ]
    for i, (dim, sd, tm) in enumerate(comp_data):
        row = table4.rows[i + 1].cells
        row[0].text = dim
        row[1].text = sd
        row[2].text = tm

    p = doc.add_paragraph()
    p.add_run('\nTable 4. ').bold = True
    p.add_run('Comparison of safety mechanisms.')

    # ─── 6. Game Theory and Information Theory ────────────────────────────────
    doc.add_heading('6. Game-Theoretic and Information-Theoretic Analysis', level=1)

    doc.add_heading('6.1 Attacker\u2019s optimal strategy', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'Consider an attacker facing m(r) = N(1+r)(w/W) marked positions '
        'in a breach lane. For each position i, the attacker decides: clear '
        '(deterministic cost c_clear) or bypass (stochastic risk). If '
        'bypassed, position i has probability 1/(1+r) of containing a real '
        'mine, with expected cost:'
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('E[C_bypass] = [1/(1+r)] \u00d7 p_cas \u00d7 C_cas              (6)')

    p = doc.add_paragraph()
    p.add_run(
        'where p_cas is casualty probability given mine encounter and C_cas '
        'is the value assigned to a casualty. The attacker rationally clears '
        'position i if and only if:'
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('c_clear < [1/(1+r)] \u00d7 p_cas \u00d7 C_cas              (7)')

    p = doc.add_paragraph()
    p.add_run(
        'Since this condition is identical for all positions (given the '
        'indistinguishability condition), the attacker\u2019s optimal strategy '
        'is a corner solution: either clear all positions or clear none. '
        'The critical dummy ratio r* at the switching point is:'
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('r* = (p_cas \u00d7 C_cas / c_clear) \u2212 1              (8)')

    p = doc.add_paragraph()
    p.add_run(
        'For representative values (p_cas = 0.15, C_cas = USD 1,000,000, '
        'c_clear = USD 500): r* = (0.15 \u00d7 1,000,000 / 500) \u2212 1 = 299. '
        'This means that for any dummy ratio below 299, a rational attacker '
        'will always clear every marked position rather than accept the '
        'risk of casualty. Since realistic treaty regimes would set r in '
        'the range 3\u201320, the minefield retains '
    )
    p.add_run('full delay effectiveness').italic = True
    p.add_run(
        ' at all plausible policy parameters. The attacker is compelled '
        'to treat every dummy as if it were real.'
    )

    doc.add_heading('6.2 Sensitivity of r* to parameter values', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'The critical ratio r* is robust to substantial parameter variation. '
        'Even under pessimistic assumptions (p_cas = 0.05, C_cas = USD 100,000, '
        'representing a force with low casualty sensitivity and cheap personnel), '
        'r* = (0.05 \u00d7 100,000 / 500) \u2212 1 = 9. Thus, even against an adversary '
        'with extremely low casualty sensitivity, a dummy ratio of r \u2264 9 ensures '
        'full delay effectiveness. For NATO-standard adversaries with higher '
        'casualty costs, r* exceeds 100 under any reasonable parameterisation.'
    )

    doc.add_heading('6.3 Defender\u2019s spatial strategy', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'The defender\u2019s optimisation involves choosing where to place N real '
        'mines and Nr dummies. Under the indistinguishability condition, the '
        'defender\u2019s optimal strategy is to draw both real and dummy positions '
        'from the same spatial distribution\u2014typically uniform across the '
        'minefield area, or concentrated along likely approach routes.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'If the attacker attempts to distinguish real from dummy using spatial '
        'pattern analysis (e.g., reasoning that tactically significant '
        'positions are more likely to be real), the defender\u2019s best response '
        'is to ensure statistical indistinguishability. This yields a '
        'mixed-strategy Nash equilibrium in which each position has '
        'probability exactly 1/(1+r) of being real, regardless of location. '
        'The defender achieves this by using the same tactical logic for '
        'dummy placement as for real placement\u2014placing dummies at positions '
        'that would be militarily optimal for real mines.'
    )

    doc.add_heading('6.4 Information entropy analysis', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'Shannon entropy provides a formal measure of the attacker\u2019s '
        'positional uncertainty (Figure 4). Under Regime A (no map), the '
        'attacker must consider all possible placements of N mines among '
        'A/\u03b4\u00b2 grid cells, giving entropy:'
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('H_A \u2248 (A/\u03b4\u00b2) \u00d7 H_b(N\u03b4\u00b2/A)              (9)')

    p = doc.add_paragraph()
    p.add_run(
        'where H_b(p) = \u2212p log\u2082(p) \u2212 (1\u2212p) log\u2082(1\u2212p). '
        'Under Regime C (MINE regime), uncertainty is confined to which '
        'N of M positions are real:'
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('H_C(r) \u2248 M \u00d7 H_b(N/M) = N(1+r) \u00d7 H_b(1/(1+r))              (10)')

    p = doc.add_paragraph()
    p.add_run(
        'For the baseline scenario: H_A \u2248 1,934 bits (no map); '
        'H_C(5) \u2248 780 bits (40% of maximum); H_C(10) \u2248 967 bits (50%). '
        'The MINE regime thus allows the defender to retain '
        '20\u201350% of maximum positional uncertainty at policy-relevant dummy '
        'ratios, whilst guaranteeing post-conflict clearance.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'This entropy analysis reveals an important asymmetry: the transparent '
        'map provides disproportionate benefit to the defender at low r values. '
        'At r = 1 (one dummy per real mine), entropy already reaches 400 bits\u2014'
        'sufficient to prevent the attacker from making useful spatial '
        'inferences about which positions are real.'
    )

    insert_figure(doc, 4)

    # ─── 6.5 Sensitivity Analysis ─────────────────────────────────────────────
    doc.add_heading('6.5 Sensitivity analysis', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'We examine the robustness of key findings to parameter variation. '
        'Table 5 presents the breach delay at r = 5 under alternative '
        'parameter assumptions.'
    )

    table5 = doc.add_table(rows=7, cols=4)
    table5.style = 'Table Grid'
    hdr = table5.rows[0].cells
    hdr[0].text = 'Scenario'
    hdr[1].text = 'Changed parameter'
    hdr[2].text = 'Breach delay (r=5)'
    hdr[3].text = 'Change from baseline'
    sens_data = [
        ('Baseline', 'N=200, t_c=30, k=10', '72 min', '\u2014'),
        ('Larger field', 'N=500', '180 min', '+150%'),
        ('Faster clearance', 't_c=15 min', '36 min', '\u221250%'),
        ('Fewer teams', 'k=5', '144 min', '+100%'),
        ('Wider breach', 'w=200 m', '144 min', '+100%'),
        ('Dense minefield', 'N=200, D=150 m', '72 min', '0% (width-invariant)'),
    ]
    for i, row_data in enumerate(sens_data):
        row = table5.rows[i + 1].cells
        for j, val in enumerate(row_data):
            row[j].text = val

    p = doc.add_paragraph()
    p.add_run('\nTable 5. ').bold = True
    p.add_run('Sensitivity of breach delay to parameter variation (r = 5).')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(
        'The model\u2019s key structural property\u2014linear scaling with (1+r)\u2014is '
        'invariant to all parameter changes. What varies is the absolute '
        'magnitude of delay, which depends on N, w/W, t_c, and k. This '
        'invariance is important for treaty design: the dummy ratio r '
        'controls the '
    )
    p.add_run('relative').italic = True
    p.add_run(
        ' military effectiveness regardless of local operational parameters, '
        'making it suitable as a universal policy instrument.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'The post-conflict safety guarantee (zero residual) is entirely '
        'independent of all parameters\u2014it follows from the logical structure '
        'of the framework (R \u2286 S) rather than from any numerical assumption. '
        'This makes the humanitarian case for the framework robust to any '
        'parameter uncertainty.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'The critical dummy ratio r* (Section 6.1) is sensitive primarily to '
        'the attacker\u2019s casualty valuation C_cas. For Western-standard armies '
        '(C_cas \u2265 USD 500,000), r* exceeds 149 under all reasonable '
        'parameterisations. For adversaries with lower casualty sensitivity '
        '(C_cas \u2248 USD 50,000\u2014an extreme lower bound), r* \u2248 14. '
        'Even in this extreme case, treaty-regulated ratios of r \u2264 10 ensure '
        'full delay effectiveness.'
    )

    # ─── 7. Treaty Design ─────────────────────────────────────────────────────
    doc.add_heading('7. Treaty Design Implications', level=1)

    doc.add_heading('7.1 Setting the dummy ratio', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'The dummy ratio r is the central policy parameter. Its optimal '
        'value depends on context-specific weights assigned to military '
        'utility versus post-conflict clearance burden. Figure 5 illustrates '
        'the trade-off space. We propose context-dependent ranges:'
    )

    p = doc.add_paragraph()
    p.add_run('\u2022 Permanent border defence ').bold = True
    p.add_run(
        '(e.g., Estonia\u2013Russia, Finland\u2013Russia): r = 5\u201310. '
        'Provides 1\u20132 hours of breach delay per lane. Post-conflict '
        'clearance: 60\u2013110 hours (2.5\u20134.5 days with 10 teams). '
        'Appropriate where the defending state controls the territory and '
        'will conduct its own post-conflict clearance.'
    )

    p = doc.add_paragraph()
    p.add_run('\u2022 Borders with civilian traffic ').bold = True
    p.add_run(
        '(e.g., areas with nearby population centres): r = 2\u20135. '
        'Lower ratio minimises clearance burden for rapid restoration '
        'of civilian movement.'
    )

    p = doc.add_paragraph()
    p.add_run('\u2022 Temporary/scatterable deployments: ').bold = True
    p.add_run(
        'r = 1\u20133. May be combined with self-destruct mechanisms as '
        'redundant safety measure.'
    )

    insert_figure(doc, 5)

    doc.add_heading('7.2 Verification and compliance mechanisms', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'The MINE regime is inherently more verifiable than '
        'existing arrangements. Three mechanisms enable compliance monitoring:'
    )

    p = doc.add_paragraph()
    p.add_run('1. Publication verification: ').bold = True
    p.add_run(
        'The existence and content of the map can be confirmed by an '
        'international registry (analogous to IAEA safeguards registries). '
        'Satellite imagery can verify that marked positions exist on the ground.'
    )

    p = doc.add_paragraph()
    p.add_run('2. Ex post ratio verification: ').bold = True
    p.add_run(
        'After conflict, clearance operations determine the actual ratio '
        'of real to dummy positions. Systematic deviation from the declared '
        'ratio constitutes a treaty violation.'
    )

    p = doc.add_paragraph()
    p.add_run('3. Completeness testing: ').bold = True
    p.add_run(
        'If a mine is discovered at a position not on the published map, '
        'this is an unambiguous treaty violation\u2014far clearer than violations '
        'of self-destruct requirements, which are inherently unobservable.'
    )

    doc.add_heading('7.3 Implementation architecture', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'Practical implementation requires three institutional components. '
        'First, a '
    )
    p.add_run('MINE Registry').italic = True
    p.add_run(
        '\u2014a secure, append-only database maintained by a trusted third party '
        '(NATO Communications and Information Agency, ICRC, or a dedicated '
        'international body) to which deploying states submit position maps. '
        'The MINE Registry must accept map submissions (new minefield declarations) '
        'and supplements (additional positions added during conflict) but '
        'must not permit deletions or modifications of previously submitted '
        'entries. This append-only property preserves the completeness '
        'guarantee: if a state registers positions, they cannot later remove '
        'entries to conceal violations.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Second, a '
    )
    p.add_run('MINE verification protocol').italic = True
    p.add_run(
        '\u2014procedures for confirming compliance both during and after conflict. '
        'During conflict, satellite imagery can verify that marked positions '
        'show ground disturbance consistent with emplacement. After conflict, '
        'clearance operations record whether each position contained a real '
        'mine or was dummy, enabling verification of the declared ratio.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Third, an '
    )
    p.add_run('enforcement mechanism').italic = True
    p.add_run(
        '\u2014consequences for non-compliance. The most powerful enforcement '
        'arises naturally: if a state is found to have placed mines at '
        'unregistered positions, the fundamental safety guarantee of the '
        'regime is violated. This provides grounds for international '
        'sanctions, loss of treaty benefits (such as ally access to safe '
        'corridors during combined operations), and reputational costs '
        'comparable to chemical weapons violations.'
    )

    doc.add_heading('7.4 Compatibility with existing frameworks', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'The MINE framework is designed to complement CCW Amended '
        'Protocol II. It extends the Protocol\u2019s Article 7 recording '
        'requirements by: (a) mandating publication during deployment '
        'rather than only post-conflict; and (b) formalising the dummy '
        'position concept with a regulated ratio. States remaining within '
        'the CCW framework could adopt the MINE Protocol as an additional '
        'protocol or as a NATO-internal standard.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'For states withdrawing from the Ottawa Convention, the framework '
        'offers a concrete mechanism to fulfil their stated commitment to '
        '\u201cresponsible use.\u201d It could be adopted unilaterally (as a national '
        'policy), bilaterally (between neighbouring states), or multilaterally '
        '(as a NATO-wide standard). The framework is deliberately technology-'
        'neutral: it applies equally to hand-emplaced and mechanically laid '
        'mines, to permanent and scatterable systems, and to current and '
        'future mine designs.'
    )

    doc.add_heading('7.5 Addressing potential objections', level=2)

    p = doc.add_paragraph()
    p.add_run('Objection 1: ').italic = True
    p.add_run('Publishing any map aids the attacker by revealing safe corridors. ')
    p.add_run('Response: ').bold = True
    p.add_run(
        'The map reveals that positions not marked are safe, but all marked '
        'positions remain dangerous. The net military effect is quantified '
        'in Section 4: delay proportional to (1+r). Moreover, an attacker '
        'who would otherwise accept casualties and push through an unmapped '
        'field is forced into clearing when positions are explicitly marked\u2014'
        'the published map makes the rational choice unambiguous.'
    )

    p = doc.add_paragraph()
    p.add_run('Objection 2: ').italic = True
    p.add_run('The defender might place mines at unregistered positions. ')
    p.add_run('Response: ').bold = True
    p.add_run(
        'This constitutes a clear, unambiguous treaty violation detectable '
        'through post-conflict clearance. Unlike self-destruct failures '
        '(which are indistinguishable from intact mines), unregistered mines '
        'provide forensic evidence of non-compliance.'
    )

    p = doc.add_paragraph()
    p.add_run('Objection 3: ').italic = True
    p.add_run('Dummy positions waste post-conflict clearance resources. ')
    p.add_run('Response: ').bold = True
    p.add_run(
        'Even at r = 20, map-based clearance is 2.4\u00d7 faster than blind '
        'sweep and guarantees completeness. The \u201cwaste\u201d of clearing dummy '
        'positions is trivial compared to the cost of sweeping an entire '
        'minefield area without guidance.'
    )

    p = doc.add_paragraph()
    p.add_run('Objection 4: ').italic = True
    p.add_run('An adversary might not respect the map (attack through unmarked areas). ')
    p.add_run('Response: ').bold = True
    p.add_run(
        'Unmarked areas are guaranteed mine-free, so this is not a risk '
        'to the attacker\u2014it is the intended safe corridor. The defender '
        'designs the minefield such that all likely approach routes are '
        'covered by marked positions, forcing the attacker through them.'
    )

    # ─── 8. Baltic Security Application ───────────────────────────────────────
    doc.add_heading('8. Application to Baltic Security', level=1)

    doc.add_heading('8.1 Threat model', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'The Baltic states face a specific threat model that makes the '
        'MINE framework particularly appropriate. Estonia, Latvia, '
        'and Lithuania share borders with Russia (and Belarus); face a '
        'significant conventional force imbalance; rely on delay and area '
        'denial to enable NATO reinforcement (which may require 10\u201330 days); '
        'and have limited territory depth, making every hour of delay '
        'operationally significant.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Finland\u2019s situation is analogous at larger scale: its 1,340 km '
        'border with Russia includes numerous potential invasion corridors '
        'through terrain favourable to defensive mining (forests, defiles, '
        'lake constrictions). Finland\u2019s withdrawal from the Ottawa Convention '
        'was explicitly framed as preparation for a defensive scenario in '
        'which mines would slow a Russian advance.'
    )

    doc.add_heading('8.2 Lessons from Ukraine', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'Russia\u2019s use of mines in Ukraine since 2022 provides empirical '
        'evidence directly relevant to the Baltic scenario. Three observations '
        'are pertinent. First, Russia has deployed mines extensively without '
        'recording or marking, creating a humanitarian crisis that will '
        'persist for decades\u2014exactly the outcome the MINE framework '
        'prevents. Second, Ukrainian breach operations against '
        'Russian minefields confirm that clearance under fire imposes '
        'delays of 1\u20134 hours per corridor even with modern equipment, '
        'validating our model\u2019s delay estimates. Third, the scale of '
        'post-conflict clearance required in Ukraine\u2014estimated at 30% of '
        'the country\u2019s territory\u2014demonstrates the catastrophic humanitarian '
        'cost of unmapped mine deployment, providing a powerful policy '
        'argument for the transparent alternative.'
    )

    doc.add_heading('8.3 Scenario analysis', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'Consider a Baltic state deploying transparent minefields with r = 7 '
        'along five key invasion corridors, each with parameters matching '
        'our baseline scenario (200 mines, 5 km front, 300 m depth). The '
        'aggregate effect:'
    )

    p = doc.add_paragraph()
    p.add_run('\u2022 Military delay: ').bold = True
    p.add_run(
        '96 minutes per corridor \u00d7 5 corridors = 8 hours of total delay '
        'imposed on advancing forces (sequential breach required for each).'
    )
    p = doc.add_paragraph()
    p.add_run('\u2022 Manpower diversion: ').bold = True
    p.add_run(
        'Each corridor requires 10 engineer teams committed to clearance '
        'for 1.6 hours, diverting 50 engineer teams from other tasks.'
    )
    p = doc.add_paragraph()
    p.add_run('\u2022 Post-conflict clearance: ').bold = True
    p.add_run(
        '5 corridors \u00d7 80 hours each = 400 team-hours total '
        '(approximately 17 days with 10 teams working 24/7). Zero residual mines.'
    )
    p = doc.add_paragraph()
    p.add_run('\u2022 Comparison without map: ').bold = True
    p.add_run(
        'Blind sweep of 5 corridors = 2,500 hours (104 days), with an '
        'expected 50 residual mines across all corridors.'
    )

    doc.add_heading('8.4 Escalation and deterrence considerations', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'The MINE regime has an additional deterrence property: '
        'the published map itself serves as a credible signal of defensive '
        'preparation. Unlike hidden minefields (which the adversary may doubt '
        'exist) or announced but unverified deployments (which may be bluffs), '
        'a registered map with specific coordinates is a verifiable commitment. '
        'This enhances deterrence by eliminating uncertainty about whether '
        'defensive preparations have actually been made.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Conversely, the framework contains a built-in de-escalation mechanism. '
        'If security conditions improve, the defending state can reduce r '
        '(reducing military effectiveness but also clearance burden) or '
        'publish the real/dummy identities (converting the field to r = 0 '
        'or enabling immediate targeted clearance). This reversibility is '
        'unavailable for deployed mines without maps\u2014once emplaced without '
        'records, mines cannot be \u201cde-escalated\u201d without expensive blind '
        'clearance.'
    )

    doc.add_heading('8.5 Country-specific considerations', level=2)
    p = doc.add_paragraph()
    p.add_run('Estonia. ').bold = True
    p.add_run(
        'Estonia\u2019s land border with Russia is approximately 294 km, of which '
        'significant portions pass through terrain suitable for defensive '
        'mining (forests, wetlands, narrow corridors between Lake Peipus and '
        'the Gulf of Finland). The Estonian Defence Forces have identified '
        'specific invasion corridors where delay is most critical\u2014'
        'particularly the Narva corridor in the northeast. With a population '
        'of 1.3 million and limited strategic depth (Tallinn is 200 km from '
        'the Russian border), every hour of delay at the border is '
        'operationally decisive. The MINE framework at r = 7\u201310 '
        'would provide 1.5\u20132.2 hours of delay per corridor whilst maintaining '
        'Estonia\u2019s commitment to humanitarian mine use\u2014a commitment explicitly '
        'stated during the 2024 National Defence Committee deliberations.'
    )

    p = doc.add_paragraph()
    p.add_run('Latvia and Lithuania. ').bold = True
    p.add_run(
        'Latvia shares 214 km of border with Russia and 161 km with Belarus; '
        'Lithuania borders Belarus for 679 km. The Suwalki Gap\u2014the 65 km '
        'corridor between Belarus and the Russian exclave of Kaliningrad\u2014'
        'represents NATO\u2019s most vulnerable land connection to the Baltic states. '
        'Transparent minefields deployed along this corridor with moderate '
        'dummy ratios (r = 5\u20137) would impose significant delay on any attempt '
        'to sever the connection, whilst the published map would enable rapid '
        'post-crisis clearance to restore normal civilian and military transit.'
    )

    p = doc.add_paragraph()
    p.add_run('Finland. ').bold = True
    p.add_run(
        'Finland\u2019s 1,340 km border with Russia\u2014NATO\u2019s longest\u2014presents '
        'a different operational calculus. The vast length makes comprehensive '
        'mining impractical; instead, selective deployment at key choke points '
        '(road crossings, defiles, bridges) is the likely operational concept. '
        'Finland\u2019s extensive Cold War experience with landmines (over 1 million '
        'destroyed after Ottawa ratification) provides institutional knowledge '
        'for implementing the MINE framework. The Finnish Defence '
        'Forces\u2019 doctrine of territorial defence, which emphasises delay and '
        'attrition over decisive engagement, aligns naturally with a framework '
        'designed to maximise delay per mine deployed.'
    )

    doc.add_heading('8.6 Comparison with alternative technologies', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'The 2024 Estonian Defence Forces assessment noted that \u201csufficient '
        'alternatives to passive mines\u201d exist, including drones, cluster '
        'munitions, and other means (ERR News, 2024). However, these '
        'alternatives have distinct characteristics that the transparent '
        'map framework addresses differently:'
    )

    p = doc.add_paragraph()
    p.add_run('\u2022 Drones and loitering munitions ').bold = True
    p.add_run(
        'provide area denial through active engagement but require '
        'continuous communication, operator presence, and have limited '
        'endurance. They cannot provide persistent 24/7 denial of a fixed '
        'area for weeks or months. Minefields, including transparent ones, '
        'provide persistent denial without active monitoring.'
    )

    p = doc.add_paragraph()
    p.add_run('\u2022 Anti-vehicle mines ').bold = True
    p.add_run(
        '(not prohibited by Ottawa) provide obstacle value against mechanised '
        'forces but are ineffective against dismounted infantry\u2014a significant '
        'component of Russian offensive doctrine as demonstrated in Ukraine. '
        'Anti-personnel mines complement anti-vehicle mines by preventing '
        'manual breaching of anti-vehicle obstacles.'
    )

    p = doc.add_paragraph()
    p.add_run('\u2022 Conventional obstacles ').bold = True
    p.add_run(
        '(wire, ditches, concrete barriers) are visible, predictable, and '
        'can be planned around. The MINE framework preserves the '
        'key advantage of mines: uncertainty about which specific positions '
        'are lethal, forcing the attacker to treat all as dangerous.'
    )

    doc.add_heading('8.7 Implementation timeline', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'The Ottawa Convention requires twelve months between notification '
        'and effective withdrawal. This creates an implementation window '
        'during which withdrawing states can develop operational procedures '
        'for the MINE framework. A realistic timeline would be:'
    )

    p = doc.add_paragraph()
    p.add_run('\u2022 Months 1\u20134: ').bold = True
    p.add_run(
        'Develop national implementation regulations; establish secure '
        'registry infrastructure; train engineer units in dual-purpose '
        '(real + dummy) emplacement procedures.'
    )
    p = doc.add_paragraph()
    p.add_run('\u2022 Months 5\u20138: ').bold = True
    p.add_run(
        'Conduct bilateral negotiations with neighbouring allies on dummy '
        'ratio standards; establish information-sharing protocols for '
        'allied safe passage; integrate with existing NATO minefield '
        'reporting systems (STANAG 2036).'
    )
    p = doc.add_paragraph()
    p.add_run('\u2022 Months 9\u201312: ').bold = True
    p.add_run(
        'Field exercises validating emplacement, publication, and clearance '
        'procedures; live testing of verification protocols; final '
        'establishment of international registry.'
    )
    p = doc.add_paragraph()
    p.add_run('\u2022 Month 12 onwards: ').bold = True
    p.add_run(
        'Operational readiness for transparent minefield deployment. '
        'Pre-planned defensive minefields can be prepared (positions '
        'identified, dummies located) awaiting only actual mine emplacement '
        'in crisis.'
    )

    doc.add_heading('8.8 NATO standardisation potential', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'The MINE framework is well-suited for adoption as a '
        'NATO standard (STANAG). Its advantages for alliance-wide implementation '
        'include: (a) interoperability\u2014allied forces can access the published '
        'map for safe passage; (b) burden-sharing\u2014post-conflict clearance '
        'can be allocated to any NATO engineer unit with the map; (c) '
        'political legitimacy\u2014adoption as a NATO standard demonstrates '
        'collective commitment to humanitarian norms even while re-introducing '
        'mines for defensive purposes.'
    )

    # ─── 9. Limitations ─────────────────────────────────────────────────────
    doc.add_heading('9. Limitations and Assumptions', level=1)

    p = doc.add_paragraph()
    p.add_run(
        'Several limitations of the present analysis should be acknowledged. '
        'First, the model assumes a static minefield with positions fixed '
        'before conflict. In practice, minefields may be modified during '
        'hostilities\u2014positions added, mines relocated, or new scatterable '
        'mines deployed. The framework would need to accommodate map updates, '
        'raising questions about update frequency and verification during '
        'active conflict. One possibility is a blockchain-based timestamped '
        'registry that accepts additions but not deletions, preserving the '
        'completeness guarantee even during dynamic operations.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Second, we assume the attacker treats all positions as independent. '
        'In reality, spatial correlations exist: mines are typically deployed '
        'in patterns (rows, clusters, mixed with anti-vehicle mines). If '
        'the attacker can observe such patterns from the published map, they '
        'might use this structural information to estimate which positions '
        'are real. The defender\u2019s best response\u2014deploying dummies with '
        'identical spatial statistics\u2014mitigates this, but sophisticated '
        'machine learning approaches might extract subtle distinguishing '
        'features. The practical significance of this limitation depends '
        'on the resolution of the published map and the defender\u2019s '
        'operational security practices.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Third, our cost parameters (c_clear = USD 500, C_cas = USD 1,000,000) '
        'are representative of Western military operations. Different '
        'adversaries may have very different cost structures. As shown in '
        'the sensitivity analysis (Section 6.5), the framework remains '
        'effective even against adversaries with substantially lower '
        'casualty costs, but the policy implications (optimal r) may differ.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Fourth, the model does not account for combined-arms effects. '
        'In practice, mine clearance is conducted under fire, and the '
        'interaction between defensive fire and mining significantly '
        'increases effective delay. This suggests our delay estimates are '
        'conservative\u2014actual delay imposed will likely exceed the clearance '
        'time alone due to suppression, casualties among clearance crews, '
        'and the need for supporting operations.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Fifth, we assume perfect compliance by the defender. If the defending '
        'state places mines at positions not registered on the map, the '
        'zero-residual property fails. However, this constitutes a clear '
        'treaty violation with forensic evidence (the unregistered mine '
        'itself), unlike violations of self-destruct requirements which '
        'produce no evidence. The framework\u2019s verification advantages '
        'create incentives for compliance that other mechanisms lack.'
    )

    doc.add_heading('9.1 Asymmetry in non-international armed conflict', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'A critical question arises in non-international armed conflicts: '
        'if only the state actor publishes under the MINE regime while the '
        'non-state adversary operates at r = \u221e (no disclosure), does '
        'this create an exploitable asymmetry? The concern is that the '
        'government\u2019s maps reveal information about defensive dispositions '
        'while rebel forces retain full uncertainty.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'The model provides a precise answer. Setting delay_map(r) equal '
        'to delay_blind and solving for the equivalence point:'
    )

    p = doc.add_paragraph()
    p.add_run(
        'N\u2009\u00d7\u2009(1 + r)\u2009\u00d7\u2009t'
    )
    run = p.add_run('clear')
    run.font.subscript = True
    p.add_run('\u2009/\u2009W = D\u2009\u00d7\u2009t')
    run = p.add_run('probe')
    run.font.subscript = True
    p.add_run('\u2009/\u2009\u03b4\u00b2')

    p = doc.add_paragraph()
    p.add_run('Solving for r:')

    p = doc.add_paragraph()
    p.add_run(
        'r* = (W \u00d7 D \u00d7 t'
    )
    run = p.add_run('probe')
    run.font.subscript = True
    p.add_run(')\u2009/\u2009(N \u00d7 t')
    run = p.add_run('clear')
    run.font.subscript = True
    p.add_run(' \u00d7 \u03b4\u00b2) \u2212 1')

    p = doc.add_paragraph()
    p.add_run(
        'Substituting baseline parameters: r* = (5000 \u00d7 300 \u00d7 5)'
        '\u2009/\u2009(200 \u00d7 30 \u00d7 25) \u2212 1 = 50 \u2212 1 = 49. '
        'At r = 49, the MINE regime imposes '
    )
    p.add_run('exactly the same delay').bold = True
    p.add_run(
        ' on the attacker as a completely undisclosed minefield. '
        'The attacker gains no advantage from the published map because '
        'at this ratio, clearing all marked positions takes precisely as '
        'long as blind sweeping the entire breach lane. However, the '
        'post-conflict outcomes are starkly different: the map regime '
        'guarantees zero residual mines, whereas the undisclosed minefield '
        'leaves an expected 10 residual mines after clearance (at '
        'p_detect = 0.95).'
    )

    p = doc.add_paragraph()
    p.add_run(
        'This yields a crucial insight for non-international armed '
        'conflicts: a state operating at r = 49 concedes '
    )
    p.add_run('nothing').italic = True
    p.add_run(
        ' in military terms relative to a non-state adversary using '
        'undisclosed mines, while securing complete post-conflict safety '
        'for its own population. For lower values of r preferred on '
        'policy grounds (r = 5\u201320), the military asymmetry exists but '
        'is partially offset by two factors. First, state forces possess '
        'complementary capabilities\u2014surveillance, artillery, air '
        'power\u2014that amplify minefield effectiveness through force '
        'multiplication (Section 4.3). Second, state forces typically '
        'have intelligence capabilities (aerial reconnaissance, signals '
        'intelligence, informant networks) that partially reveal non-state '
        'actor mine locations, reducing the effective adversary r below '
        'infinity. The information asymmetry is therefore bidirectional, '
        'not unilateral.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Moreover, the zero-residual property depends '
    )
    p.add_run('only on the defender\u2019s own compliance').italic = True
    p.add_run(
        ', not the adversary\u2019s. Even if a non-state actor uses undisclosed '
        'mines, the state\u2019s own published map ensures that all '
    )
    p.add_run('state-emplaced').italic = True
    p.add_run(
        ' mines are fully cleared after conflict. This unilateral benefit '
        'requires no reciprocity\u2014a significant advantage over treaty '
        'mechanisms that depend on mutual compliance.'
    )

    doc.add_heading('9.2 Non-state actors and collective response', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'The MINE framework applies exclusively to state actors '
        'operating within the law of armed conflict. Non-state armed '
        'groups (NSAGs) and terrorist organisations remain outside the '
        'scope of treaty-based regulation\u2014a limitation shared with all '
        'existing instruments including the Ottawa Convention and CCW '
        'Amended Protocol II. However, this shared limitation does not '
        'preclude complementary mechanisms specifically addressing '
        'non-state mine use.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'We propose a Collective Response Mechanism (CRM) as an '
        'institutional complement to the MINE framework. '
        'The CRM would operate on three principles:'
    )

    p = doc.add_paragraph()
    p.add_run('First, automatic assistance: ').bold = True
    p.add_run(
        'when NSAG mine use is confirmed in a signatory state\u2019s '
        'territory, treaty parties assume a collective obligation to '
        'provide humanitarian mine clearance funding and technical '
        'assistance. This socialises the clearance burden created by '
        'actors beyond treaty control and provides immediate benefit '
        'to affected civilian populations.'
    )

    p = doc.add_paragraph()
    p.add_run('Second, state responsibility for proxies: ').bold = True
    p.add_run(
        'states that supply mines to non-state actors or provide '
        'material support for NSAG mining operations incur liability '
        'under the framework. NSAG mine use serves as an automatic '
        'trigger for sanctions review against identified state sponsors. '
        'This creates upstream incentives to restrict mine transfers to '
        'non-state actors, complementing existing arms embargo mechanisms '
        'under UN Security Council resolutions.'
    )

    p = doc.add_paragraph()
    p.add_run('Third, differentiated legitimacy: ').bold = True
    p.add_run(
        'the framework establishes a clear legal and normative '
        'distinction between \u201cresponsible use\u201d (states publishing '
        'transparent maps within the regulated r parameter) and '
        '\u201cirresponsible use\u201d (any mine deployment without position '
        'disclosure). This binary classification strengthens the '
        'normative case against NSAG mine use by providing a concrete, '
        'verifiable standard of responsible behaviour that NSAGs '
        'manifestly fail to meet. It also provides legal clarity for '
        'states re-introducing mines: compliance with the MINE '
        'Protocol demonstrates good faith commitment to humanitarian '
        'norms, insulating compliant states from the stigma associated '
        'with indiscriminate mine use.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Precedent for such mechanisms exists. The Geneva Call initiative '
        'has secured \u201cDeeds of Commitment\u201d from over 60 non-state armed '
        'groups pledging mine non-use\u2014demonstrating that normative '
        'pressure can reach beyond treaty law. The MINE framework '
        'strengthens this norm by making the distinction '
        'between responsible and irresponsible use operationally concrete '
        'and externally verifiable.'
    )

    doc.add_heading('9.3 Border changes and map continuity', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'Conflicts that alter territorial boundaries raise a distinct '
        'challenge: minefields emplaced as border defences may end up in '
        'territory controlled by the former adversary. Under CCW Amended '
        'Protocol II Article 9, parties are obligated to transfer minefield '
        'records to the opposing party and to the UN Secretary-General upon '
        'cessation of hostilities. However, this obligation relies on '
        'post-conflict cooperation\u2014precisely the condition least likely to '
        'obtain after territorial conquest.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'The MINE framework resolves this problem by design. '
        'Because maps are '
    )
    p.add_run('published before or during conflict').italic = True
    p.add_run(
        ', the information is already in the public domain regardless of '
        'subsequent territorial changes. Four scenarios illustrate this:'
    )

    p = doc.add_paragraph()
    p.add_run('Defence succeeds (borders maintained): ').bold = True
    p.add_run(
        'Standard case. The defending state clears its own published '
        'positions using the map. Zero residual guaranteed.'
    )

    p = doc.add_paragraph()
    p.add_run('Territory lost (occupation): ').bold = True
    p.add_run(
        'The occupying power has full access to the published map and can '
        'clear all marked positions in the occupied territory. Unlike the '
        'CCW regime, no bilateral negotiation for record transfer is needed. '
        'The map\u2019s pre-publication eliminates dependence on defeated state '
        'cooperation.'
    )

    p = doc.add_paragraph()
    p.add_run('Territory recaptured: ').bold = True
    p.add_run(
        'If the adversary also operated under the MINE regime, '
        'their published positions enable systematic clearance of enemy-'
        'emplaced mines in recaptured territory. If the adversary did not '
        'publish (non-signatory or non-state actor), the recapturing state '
        'faces the same blind-sweep problem as under the status quo\u2014but '
        'its own mines remain fully mapped and clearable.'
    )

    p = doc.add_paragraph()
    p.add_run('Negotiated border adjustment: ').bold = True
    p.add_run(
        'In cases of agreed territorial exchange (e.g., post-conflict '
        'peace treaty adjusting borders), the published map provides '
        'complete information to the receiving state without requiring '
        'additional diplomatic arrangements. This facilitates rapid '
        'civilian resettlement in transferred territories.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'The permanence of published information is further strengthened by '
        'the proposed registry architecture (Section 7.3). An append-only '
        'international registry\u2014whether maintained by the UN, a dedicated '
        'treaty organisation, or implemented as a distributed ledger\u2014'
        'ensures that map data persists independently of the emplacing '
        'state\u2019s continued existence or cooperation. Even in scenarios of '
        'state collapse (as occurred in Yugoslavia or the Soviet Union), '
        'the registry preserves mine location data for successor states, '
        'international organisations, or humanitarian clearance agencies.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'This property is particularly relevant for refugee return. '
        'Displaced populations can consult the published map to identify '
        'safe corridors and cleared zones, enabling phased return aligned '
        'with clearance operations. The deterministic endpoint\u2014once all '
        'published positions are cleared, the area is certifiably safe\u2014'
        'provides a level of assurance that probabilistic clearance methods '
        'cannot match, accelerating the political and practical conditions '
        'for post-conflict reconstruction.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Finally, the political feasibility of gaining acceptance for any '
        'framework that legitimises mine use\u2014however constrained\u2014should '
        'not be underestimated. The Ottawa Convention\u2019s norm-setting '
        'achievement has created strong institutional resistance to any '
        'proposal perceived as weakening the prohibition. We argue that '
        'the MINE framework should be positioned not as a '
        'weakening but as a strengthening of post-conflict humanitarian '
        'outcomes: it provides a safety guarantee (zero residual) that '
        'even full Ottawa compliance cannot offer for states that '
        'withdraw and use mines without constraints.'
    )

    # ─── 10. Conclusion ───────────────────────────────────────────────────────
    doc.add_heading('10. Conclusion', level=1)

    p = doc.add_paragraph()
    p.add_run(
        'This paper has presented the MINE framework (Mapping Is Not '
        'Exposure) for anti-personnel mines, reconciling military utility with '
        'humanitarian guarantees through a single, tuneable parameter: '
        'the dummy ratio r. The MINE framework transforms the binary choice '
        'between total prohibition (Ottawa Convention) and unrestricted '
        'use into a continuum of policy options that can be tailored to '
        'specific security contexts and negotiated through treaty mechanisms.'
    )

    p = doc.add_paragraph()
    p.add_run('Our analysis establishes three core properties:')

    p = doc.add_paragraph()
    p.add_run('1. Military effectiveness is preserved. ').bold = True
    p.add_run(
        'Breach delay scales linearly with (1+r). Game-theoretic analysis '
        'confirms that rational attackers will clear all positions for any '
        'plausible treaty-regulated ratio (r* = 299 under baseline parameters). '
        'The minefield retains full delay capability.'
    )

    p = doc.add_paragraph()
    p.add_run('2. Post-conflict safety is guaranteed. ').bold = True
    p.add_run(
        'The zero-residual property ensures that clearing all map positions '
        'removes every real mine\u2014a deterministic guarantee superior to any '
        'probabilistic detection method or mechanical self-destruct mechanism.'
    )

    p = doc.add_paragraph()
    p.add_run('3. Treaty compliance is verifiable. ').bold = True
    p.add_run(
        'Map publication is observable, dummy ratios are testable ex post, '
        'and violations (unregistered mines) are unambiguously detectable.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'A fourth property deserves particular emphasis: '
    )
    p.add_run('unilateral benefit without reciprocity. ').bold = True
    p.add_run(
        'As demonstrated in Section 9.1, a state operating at r = 49 '
        'achieves exactly the same breach delay as a completely undisclosed '
        'minefield. This means that even in conflict with a non-signatory '
        'adversary\u2014whether a state that has not ratified the framework or '
        'a non-state armed group\u2014the defending state sacrifices '
    )
    p.add_run('nothing').italic = True
    p.add_run(
        ' in military terms by publishing its map. Yet the post-conflict '
        'dividend is transformative: upon cessation of hostilities, every '
        'position on the published map is systematically cleared, yielding '
        'a deterministic guarantee of zero residual mines in the state\u2019s '
        'own defensive zones. Reconstruction, refugee return, and '
        'agricultural resumption can commence immediately in cleared areas '
        'without the years-long uncertainty characteristic of unmapped '
        'minefields.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'This unilateral benefit structure has profound implications for '
        'treaty adoption dynamics. Classical arms control faces a '
        'prisoner\u2019s dilemma: each state benefits from others\u2019 compliance '
        'but is tempted to defect. The MINE framework inverts '
        'this logic. A state benefits from its '
    )
    p.add_run('own').italic = True
    p.add_run(
        ' compliance regardless of others\u2019 behaviour\u2014the zero-residual '
        'property depends solely on the defender\u2019s map accuracy, not the '
        'adversary\u2019s participation. Ratification is therefore a dominant '
        'strategy: it improves post-conflict outcomes for the ratifying '
        'state under all scenarios (adversary ratifies, adversary does not '
        'ratify, adversary is a non-state actor). This eliminates the '
        'principal barrier to multilateral adoption and makes unilateral '
        'or small-coalition implementation viable as a first step.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'For the Baltic states and Finland, this framework provides an '
        'operationally credible middle path at a critical juncture. Rather '
        'than returning to unrestricted mine use (accepting humanitarian '
        'costs) or remaining within a prohibition that sacrifices military '
        'utility (accepting security costs), these states could adopt the '
        'MINE regime as a NATO-wide standard. In doing so, they '
        'would demonstrate that responsible deterrence and humanitarian norms '
        'are not merely compatible but complementary\u2014that publishing a '
        'minefield map is simultaneously an act of military preparation '
        'and an act of humanitarian commitment.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'The policy window for adoption is narrow but open. With five NATO '
        'states in various stages of Ottawa Convention withdrawal, the next '
        '12\u201324 months will determine whether these states develop individual, '
        'potentially divergent mine employment practices or converge on a '
        'common standard. The MINE framework offers an immediate, '
        'implementable standard requiring no new technology, no novel legal '
        'instruments beyond existing CCW mechanisms, and no compromise of '
        'either military utility or humanitarian principles.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Future research should address several extensions. First, dynamic '
        'scenarios where the published map is updated during conflict to '
        'reflect scatterable mine deployment or repositioning of existing '
        'fields. Second, interaction between transparent minefields and other '
        'defensive systems, particularly anti-vehicle mines and direct fire '
        'coverage, which amplify the delay effect beyond what our model '
        'captures. Third, the political economy of multilateral dummy ratio '
        'negotiation, including bargaining dynamics between states with '
        'different threat perceptions. Fourth, empirical calibration of '
        'clearance time parameters using data from ongoing mine clearance '
        'operations in Ukraine, which provide an unprecedented volume of '
        'real-world measurement under field conditions.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'We note that the MINE framework is immediately implementable using '
        'existing technology and legal mechanisms. No novel mine designs, '
        'no new sensors, and no untested verification technologies are '
        'required. The only innovation is informational: the commitment to '
        'publish a map, and the agreement on how many positions it may '
        'contain. This simplicity is a feature, not a limitation\u2014it means '
        'adoption can proceed in parallel with the withdrawal process, '
        'providing operational standards from day one of resumed mine use.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'The MINE framework is, ultimately, an information-theoretic '
        'innovation applied to a humanitarian problem. It recognises that '
        'the harm caused by landmines is not only physical but informational: '
        'it is the '
    )
    p.add_run('uncertainty').italic = True
    p.add_run(
        ' about mine locations that makes post-conflict land unusable and '
        'clearance so expensive. By structuring that uncertainty as a '
        'controlled parameter\u2014published, bounded, and treaty-regulated\u2014'
        'the MINE framework converts an intractable post-conflict problem '
        'into a finite, plannable task with a guaranteed end state.'
    )

    # ─── References ───────────────────────────────────────────────────────────
    doc.add_heading('References', level=1)

    refs = [
        'Aalto NDU (2020) Defensive Minefield Planning using Network Interdiction Models. Master\u2019s thesis, Aalto University / Finnish National Defence University.',
        'ERR News (2024) \u2018EDF commander: Estonia does not need to leave the Ottawa Convention right now\u2019, ERR, 16 December. Available at: https://news.err.ee/1609552774 (Accessed: 15 June 2025).',
        'Hendricks, K. (2006) \u2018Feints\u2019, Journal of Economics & Management Strategy, 15(2), pp. 431\u2013456.',
        'ICBL (2023) Landmine Monitor 2023. Geneva: International Campaign to Ban Landmines.',
        'ICRC (2024) Banning Anti-Personnel Mines: The Ottawa Treaty Explained. Geneva: International Committee of the Red Cross.',
        'Kanazawa, T., Iida, K. and Morimoto, S. (2018) \u2018Risk evaluation and games in mine warfare considering shipcounter effects\u2019, European Journal of Operational Research, 267(1), pp. 278\u2013288.',
        'Norwegian People\u2019s Aid (2007) M85: An Analysis of Reliability. Oslo: Norwegian People\u2019s Aid.',
        'Reuters (2024) \u2018Finland considering exiting anti-personnel landmine treaty, minister says\u2019, Reuters, 18 December.',
        'Reuters (2025) \u2018Finland to exit landmines treaty, hike defence spending given Russia threat, PM says\u2019, Reuters, 1 April. Available at: https://www.reuters.com/world/europe/finland-plans-withdraw-landmines-treaty-prime-minister-says-2025-04-01/ (Accessed: 15 June 2025).',
        'UNODA (2023) CCW Amended Protocol II. Geneva: United Nations Office for Disarmament Affairs. Available at: https://disarmament.unoda.org/ccw/amended-protocol-ii/ (Accessed: 15 June 2025).',
        'US Army (1985) FM 5-102: Countermobility. Washington, DC: Department of the Army.',
        'US Army (1988) FM 90-2: Battlefield Deception. Washington, DC: Department of the Army.',
        'Zhuang, J., Bier, V.M. and Alagoz, O. (2010) \u2018Modeling secrecy and deception in a multiple-period attacker\u2013defender signaling game\u2019, European Journal of Operational Research, 203(2), pp. 409\u2013418.',
    ]

    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.add_run(ref)

    # Save
    doc.save(str(OUTPUT_DIR / "manuscript_JoBS.docx"))
    print(f"Manuscript saved: {OUTPUT_DIR / 'manuscript_JoBS.docx'}")


if __name__ == "__main__":
    create_manuscript()

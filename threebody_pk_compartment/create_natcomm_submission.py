#!/usr/bin/env python3
"""
Generate Nature Communications submission package (double-blind):
  - manuscript_natcomm_doubleblind.docx (no author info)
  - title_page.docx (author info, separate file)
  - cover_letter_natcomm.docx
  - figures_editable.pptx (already exists)

Nature Communications formatting:
  - Abstract: ≤150 words, final sentence starts with "Here, we" or "In this work"
  - Introduction: no subheadings, final paragraph starts with "Here/In this work"
  - Results: subheaded sections (≤60 chars, no numbering)
  - Discussion: no subheadings
  - Methods: subheaded sections
  - Main text (excl. figure legends & Methods): ≤6,000 words
  - Up to 10 display items
  - References: Vancouver style, numbered in citation order
  - Avoid "novel", "new", "for the first time", "unprecedented"
"""

from __future__ import annotations

import os
import re
import zipfile
import shutil
from datetime import date

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASEDIR = os.path.dirname(os.path.abspath(__file__))
FIGDIR_NATURE = os.path.join(BASEDIR, "figures_nature")
FIGDIR_ORIG = os.path.join(BASEDIR, "figures")
FIGDIR = FIGDIR_NATURE if os.path.isdir(FIGDIR_NATURE) else FIGDIR_ORIG
OUTDIR = os.path.join(BASEDIR, "manuscript")
os.makedirs(OUTDIR, exist_ok=True)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def add_text_with_refs(paragraph, text):
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = paragraph.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(8)
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(11)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_paragraph_with_refs(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    add_text_with_refs(p, text)
    return p


def add_figure(doc, fig_path, caption, fig_num):
    if os.path.exists(fig_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fig_path, width=Inches(6.0))
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(12)
    run_bold = cap.add_run(f"Figure {fig_num} | ")
    run_bold.bold = True
    run_bold.font.size = Pt(10)
    add_text_with_refs(cap, caption)


def setup_style(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5


# ---------------------------------------------------------------------------
# Content — Nature Communications version
# ---------------------------------------------------------------------------

TITLE = ("Pharmacokinetic compartment theory solves the statistical "
         "three-body problem")

AUTHORS = "Tatsuki Onishi"

AFFILIATIONS = (
    "Faculty of Data Science, Shiga University, "
    "1-1-1 Banba, Hikone, Shiga 522-8522, Japan"
)

CORRESPONDING = (
    "Correspondence: Tatsuki Onishi\n"
    "Faculty of Data Science, Shiga University, "
    "1-1-1 Banba, Hikone, Shiga 522-8522, Japan\n"
    "Email: bougtoir@gmail.com"
)

# Abstract: ≤150 words. Final sentence starts with "Here, we show".
ABSTRACT = (
    "The gravitational three-body problem remains analytically intractable "
    "after 250 years, yet its statistical properties follow remarkably "
    "regular patterns. Pharmacokinetic (PK) compartmental analysis is the "
    "standard framework for modelling drug distribution and elimination in "
    "clinical medicine. Here, we show that chaotic three-body scattering "
    "maps onto a three-compartment PK model\u2014where binary configurations "
    "are compartments, configuration transitions are inter-compartmental "
    "transfers, and system dissolution is elimination\u2014and demonstrate "
    "through 15,000 N-body simulations across 64 mass configurations "
    "that this structural correspondence is quantitative, bidirectional, "
    "and robust to dimensionality and gravitational-wave dissipation."
)

# Introduction: no subheadings, final paragraph starts with "Here, we show"
INTRO = [
    ("The three-body problem\u2014predicting the long-term motion of three "
     "gravitationally interacting masses\u2014has resisted analytical solution "
     "since Newton.{1} Poincar\u00e9 proved its non-integrability,{2} and "
     "modern numerical studies confirm that generic three-body encounters are "
     "chaotic, with outcomes sensitive to initial conditions at the level of "
     "floating-point precision.{3} Nevertheless, the statistical properties "
     "of the outcomes are remarkably regular: the lifetime distribution follows "
     "a multi-exponential decay,{4,5} and the final-state distributions can be "
     "predicted from phase-space volume arguments.{6,7}"),

    ("Recently, Ginat and Perets showed that chaotic three-body scattering "
     "can be decomposed into a sequence of independent probabilistic "
     "excursions\u2014each a temporary binary plus a distant single body\u2014and "
     "derived an analytical, statistical solution as a random walk in "
     "binary binding energy.{5} Stone and Leigh independently demonstrated "
     "the predictive power of flux-based statistical mechanics for three-body "
     "outcomes.{6} These advances establish that the ergodic core of the "
     "three-body problem is amenable to stochastic modelling."),

    ("Pharmacokinetic (PK) compartmental analysis is the standard framework "
     "for modelling drug absorption, distribution, and elimination in "
     "clinical medicine.{8,9} In PK theory, the body is divided into "
     "compartments (e.g. plasma, tissues), drug transfers between "
     "compartments at first-order rates, and is eliminated irreversibly. "
     "The mathematics is a continuous-time Markov chain (CTMC) whose "
     "generator is the PK rate matrix.{10}"),

    ("Here, we show that three-body scattering is a three-compartment PK "
     "system: each binary configuration (which body pair is bound) is a "
     "compartment; transitions between configurations are inter-compartmental "
     "transfers; escape of one body is elimination. This mapping is not merely "
     "analogical\u2014it is a structural equivalence at the level of the master "
     "equation. We exploit it to import the full toolkit of PK analysis "
     "(half-lives, mean residence time, bioequivalence, population modelling, "
     "nonlinear kinetics) into celestial mechanics, and conversely to bring "
     "dynamical-systems stability theory into pharmacology."),
]

RESULTS_SECTIONS = {
    "Linear PK model reproduces multi-exponential decay": [
        ("We performed 15,000 three-body scattering simulations using a "
         "symplectic leapfrog integrator (Methods), spanning three mass "
         "configurations: equal mass (1:1:1), unequal mass (1:2:0.5), and "
         "democratic initial conditions (1:1:1, symmetric start). For each "
         "simulation, we recorded the full sequence of binary configurations "
         "visited and dwell times in each (Fig. 1)."),

        ("From the transition counts and dwell times, we estimated the 9 PK "
         "rate parameters (6 inter-compartmental transfer rates k_ij and 3 "
         "elimination rates k_ei) by maximum-likelihood estimation for a "
         "continuous-time Markov chain: k_ij = N_ij / T_i, where N_ij is the "
         "number of observed i-to-j transitions and T_i is total dwell time in "
         "compartment i (Methods)."),

        ("The resulting linear PK model predicts a survival function S(t) that "
         "is a sum of three exponentials, each corresponding to an eigenvalue "
         "of the rate matrix. This multi-exponential prediction matches the "
         "empirical survival curves from N-body simulations across all three "
         "mass configurations (Fig. 2), with the single-exponential model "
         "rejected by the Kolmogorov\u2013Smirnov test "
         "(p < 0.001 for all three configurations), confirming that the "
         "three-compartment model is necessary and sufficient to describe "
         "the intermediate-timescale dynamics."),

        ("The MLE-estimated rates agree quantitatively with the phase-space "
         "flux prediction from statistical mechanics.{6} For equal "
         "masses, all transfer rates converge to the same value by symmetry "
         "(k_ij \u2248 0.015 per dynamical time), validating both the "
         "PK framework and the ergodic hypothesis underlying the statistical "
         "approach."),
    ],

    "Nonlinear PK captures sticky chaos": [
        ("While the linear PK model fits the bulk of the lifetime distribution, "
         "it systematically underestimates the long-lived tail\u2014a signature of "
         "\u2018sticky chaos\u2019, where trajectories become trapped near periodic "
         "orbits for anomalously long times.{11,12} In PK language, this "
         "corresponds to saturable (capacity-limited) elimination: when the "
         "system lingers near a stable periodic orbit, the escape rate "
         "decreases rather than remaining constant."),

        ("We model this by replacing the first-order elimination k_e \u00d7 P "
         "with Michaelis\u2013Menten kinetics: V_max \u00d7 P / (K_m + P). "
         "This hybrid model\u2014linear inter-compartmental transfer with nonlinear "
         "elimination\u2014produces a survival function whose tail decays as "
         "t^{\u2212\u03b1} \u00d7 exp(\u2212\u03bb_slow \u00d7 t), "
         "capturing the power-law correction (Fig. 3)."),

        ("Fitting the hybrid model to all three mass configurations yields "
         "power-law exponents \u03b1 \u2248 1.1\u20132.6, with tail "
         "RMSE improvement of 36\u201342% over the linear model for unequal and "
         "democratic configurations. For the equal-mass case, the linear model "
         "is already adequate (no significant tail excess), consistent with the "
         "known weaker stickiness in the symmetric case.{13}"),
    ],

    "Population PK reveals allometric scaling": [
        ("Clinical PK uses population (mixed-effects) modelling to relate "
         "individual PK parameters to patient covariates such as body weight "
         "and renal function.{14} We apply the same methodology to the "
         "three-body problem: the \u2018covariates\u2019 are the mass ratios of the "
         "three bodies."),

        ("Scanning 64 mass configurations (m\u2082, m\u2083 \u2208 "
         "{0.25, 0.5, 0.75, 1, 1.5, 2, 3, 4} with m\u2081 = 1), "
         "we fit a log-linear population model: "
         "log(MRT) = \u03b2\u2080 + \u03b2\u2081\u00b7log(\u03bc\u2081\u2082) "
         "+ \u03b2\u2082\u00b7log(\u03bc_out) + \u03b2\u2083\u00b7log(M) + \u03b7, "
         "where \u03bc\u2081\u2082 is the reduced mass of the initial binary, "
         "\u03bc_out is the binary\u2013single reduced mass, and M is total mass "
         "(Fig. 4)."),

        ("The population model yields MRT \u221d \u03bc\u2081\u2082^{2.08} "
         "\u00d7 \u03bc_out^{\u22121.68} \u00d7 M^{\u22120.62}, "
         "with R\u00b2 = 0.67 and inter-individual variability \u03c9 = 0.35. "
         "The positive exponent on \u03bc\u2081\u2082 indicates that heavier "
         "binaries survive longer (deeper potential wells), while the negative "
         "exponent on \u03bc_out indicates that heavier incoming bodies disrupt "
         "the system faster\u2014both physically intuitive and quantitatively "
         "informative."),

        ("Additionally, the escape probability of the lightest body follows a "
         "logistic model in mass fraction, analogous to the dose\u2013response "
         "relationship in pharmacology. This provides a closed-form predictor "
         "for which body will be ejected, a result of practical value for "
         "stellar dynamics simulations."),
    ],

    "Reverse mapping: TMDD stability analysis": [
        ("The structural equivalence works in both directions. Target-mediated "
         "drug disposition (TMDD)\u2014where a drug binds a receptor to form a "
         "complex that is then internalised\u2014is a three-state system with "
         "identical matrix structure to the three-body compartment model.{15,16}"),

        ("We demonstrate that the Jacobian eigenvalue analysis used to assess "
         "Lagrange point stability in celestial mechanics provides a natural "
         "framework for TMDD stability analysis (Fig. 5). The dose\u2013response "
         "bifurcation curve\u2014identifying stable versus unstable steady states "
         "as a function of drug input rate\u2014is the pharmacological analogue "
         "of the Mardling\u2013Aarseth stability criterion for hierarchical "
         "triples.{17}"),

        ("This reverse mapping suggests that the rich stability theory of "
         "N-body systems (KAM tori, invariant manifolds, homoclinic tangles) "
         "could be imported wholesale into nonlinear pharmacokinetics, "
         "particularly for complex biologics with multiple binding sites."),
    ],

    "3D robustness and gravitational-wave dissipation": [
        ("To test whether the PK mapping is an artefact of the 2D planar "
         "geometry or the conservative (energy-preserving) assumption, we "
         "extended the simulations to fully three-dimensional scattering with "
         "randomised orbital planes (uniform on SO(3)), thermal eccentricity "
         "(f(e) = 2e), and isotropic third-body approach directions "
         "(2,000 runs per mass configuration)."),

        ("The three-compartment PK model remains valid in all 3D configurations, "
         "with Kolmogorov\u2013Smirnov statistics of 0.16\u20130.31 across equal, "
         "unequal, and democratic mass setups\u2014comparable to the 2D values "
         "(Fig. 6). The median lifetimes shift modestly (3D conservative: "
         "246\u2013438 dynamical times versus 2D: 239\u20131,043), reflecting the "
         "larger phase space available in 3D."),

        ("Adding 2.5 post-Newtonian gravitational-wave radiation reaction{19} "
         "(c = 100 in code units; v_orbital \u2248 1) systematically "
         "shortens lifetimes without destroying the compartment structure "
         "(Fig. 7). The lifetime ratio \u03c4_diss/\u03c4_cons ranges from 0.98 "
         "(equal mass) to 0.83 (democratic), with median fractional energy "
         "loss |\u0394E/E| = 0.10\u20130.86% per encounter. The democratic "
         "configuration shows the strongest dissipation effect (17% lifetime "
         "reduction), consistent with the longer interaction times providing "
         "more opportunity for gravitational-wave energy extraction."),

        ("These results indicate that the PK mapping is robust across "
         "dimensionality, eccentricity distribution, and the presence of "
         "dissipation\u2014and is therefore applicable to astrophysically "
         "realistic encounters."),
    ],
}

DISCUSSION = [
    ("We have demonstrated a quantitative structural correspondence between two "
     "seemingly unrelated fields: celestial mechanics and pharmacokinetics. "
     "The key insight is that both describe irreversible first-passage "
     "problems in multi-state continuous-time Markov chains, and that the "
     "mathematical structure\u2014not merely the analogy\u2014is shared."),

    ("The practical implications are bidirectional. For celestial mechanics: "
     "(i) the population PK framework provides a principled way to parameterise "
     "outcome statistics as a function of mass ratios, replacing ad hoc fitting "
     "functions;{5\u20137} (ii) the Michaelis\u2013Menten model offers a minimal "
     "parametrisation of sticky chaos that interpolates between the ergodic "
     "(linear) regime and the trapped (power-law) regime; (iii) PK software "
     "(NONMEM, Monolix) can be directly applied to N-body simulation output."),

    ("For pharmacology: (i) the dynamical-systems stability framework "
     "(Lagrange points, invariant manifolds) provides geometric insight into "
     "TMDD steady states that goes beyond the standard eigenvalue analysis; "
     "(ii) the concept of \u2018sticky chaos\u2019 maps to the clinically important "
     "phenomenon of prolonged drug\u2013receptor residence time;{18} (iii) the "
     "allometric scaling framework transfers predictive power across species "
     "and across mass configurations."),

    ("The robustness of the PK mapping to dimensionality and dissipation "
     "(Figs. 6\u20137) has astrophysical implications. The 17% "
     "lifetime reduction in democratic configurations under gravitational-wave "
     "dissipation may inform merger-rate estimates for "
     "LISA-band sources: three-body encounters in dense clusters will resolve "
     "faster than energy-conserving estimates suggest, potentially increasing the "
     "predicted rate of observable in-spirals. The population PK framework "
     "may provide closed-form event-rate formulae parameterised by cluster "
     "mass function, potentially complementing expensive Monte Carlo surveys."),

    ("Remaining limitations include the perturbative treatment of radiation "
     "reaction (valid only for c \u226b v_orbital) and the absence of tidal "
     "dissipation relevant for stellar encounters. The linear PK model "
     "assumes ergodicity, which breaks down for very long lifetimes; the "
     "nonlinear extension addresses this but at the cost of additional "
     "parameters."),

    ("More broadly, the correspondence inverts the conventional flow of "
     "methodology between the physical sciences and clinical medicine. "
     "The mathematical theory of pharmacokinetics was distilled from bedside "
     "observations\u2014therapeutic drug monitoring, dose\u2013response studies, "
     "clinical trial design\u2014and aged over six decades of patient-facing "
     "application into a robust quantitative infrastructure. That "
     "infrastructure has been casked in regulatory-grade software (NONMEM, "
     "Monolix) and in saturation-kinetics theory refined through countless "
     "dosing studies (Michaelis\u2013Menten, TMDD). The present work decants this "
     "accumulated clinical knowledge into an entirely new vessel\u2014celestial "
     "mechanics\u2014where, aerated by contact with a different physical context, "
     "it reveals latent capabilities that were invisible within medicine alone. "
     "Several concrete directions follow. "
     "(i) The three-compartment model generalises naturally to N-body "
     "hierarchies: quadruple and higher-order multiples in dense star clusters "
     "map onto N-compartment cascade models whose analytical solutions are "
     "textbook clinical pharmacokinetics.{8,9} "
     "(ii) Population PK covariate modelling\u2014originally designed to account "
     "for patient weight, renal function, and genotype in clinical trials\u2014can "
     "treat mass ratios, angular momenta, and cluster environments as "
     "covariates in nonlinear mixed-effects regressions, enabling systematic "
     "astrophysical parameter scans on platforms built for drug regulation. "
     "(iii) The closed-form MRT scaling (MRT \u221d "
     "\u03bc\u2081\u2082^\u03b1 \u00d7 \u03bc_out^\u03b2 \u00d7 M^\u03b3) can be "
     "convolved with a cluster mass function to yield gravitational-wave "
     "event-rate estimates for LISA-band sources without expensive Monte Carlo "
     "surveys. "
     "(iv) Conversely, the geometric stability theory of celestial mechanics "
     "(KAM tori, invariant manifolds) may inform the rational design of drugs "
     "with controlled target residence times\u2014a key determinant of in vivo "
     "efficacy that is central to modern drug discovery.{18} "
     "(v) Because the mapping relies only on the CTMC structure of "
     "multi-state first-passage processes, it is not limited to gravitational "
     "systems; chemical reaction networks, ecological predator\u2013prey dynamics, "
     "and other multi-state chaotic systems may admit analogous descriptions "
     "rooted in clinical pharmacokinetic theory, suggesting a broader "
     "programme of \u2018pharmacokinetics of chaos.\u2019"),

    ("In conclusion, the pharmacokinetic compartment model provides both a "
     "conceptual lens and a computational toolkit for the three-body problem "
     "that complements existing statistical-mechanical approaches. By "
     "demonstrating that clinical medicine\u2014through its pharmacokinetic "
     "tradition\u2014can export, not merely import, quantitative frameworks to "
     "the physical sciences, this work illustrates how disciplinary "
     "boundaries may conceal unexploited structural connections."),
]

METHODS = [
    ("N-body simulations", [
        "Three-body scattering simulations were performed using a symplectic "
        "leapfrog (St\u00f6rmer\u2013Verlet) integrator implemented in Julia 1.10. "
        "Initial conditions: binary with semi-major axis a = 1 and eccentricity "
        "e = 0; third body approaching from r = 10a with velocity at infinity "
        "v_inf = 0.1 (parabolic-like encounters). Impact parameter b sampled "
        "uniformly in b\u00b2 from 0 to b_max = 5a. Integration continued until "
        "escape (body reaching r > 20a with positive energy) or t_max = 10,000 "
        "dynamical times. Adaptive timestep: dt = 0.01 \u00d7 r_min / v_max. "
        "Energy conservation better than 10\u207b\u2078 in all runs.",

        "For each mass configuration, 5,000 scattering experiments were "
        "performed (total: 15,000 for Phase 1; 5,000 \u00d7 64 = 320,000 for "
        "the population PK scan). Binary configuration was identified at each "
        "timestep by finding the most-bound pair (most negative pairwise "
        "energy). A configuration transition was recorded when the most-bound "
        "pair identity changed and persisted for more than 5 timesteps.",
    ]),

    ("PK model estimation", [
        "Transition rates were estimated by maximum-likelihood for a "
        "continuous-time Markov chain: k_ij = N_ij / T_i, where N_ij is the "
        "count of transitions from compartment i to j, and T_i is total "
        "observed dwell time in compartment i. The 3\u00d73 rate matrix A was "
        "constructed such that A_ij = k_ji for i \u2260 j, and "
        "A_ii = \u2212\u03a3(k_ij + k_ei).",

        "Survival function: S(t) = 1\u1d40 \u00d7 exp(At) \u00d7 P\u2080. "
        "Eigenvalues of A give the three decay rates; half-lives t\u00bd = "
        "ln(2) / |\u03bb_i|. Mean residence time (MRT) = \u22121\u1d40 \u00d7 "
        "A\u207b\u00b9 \u00d7 P\u2080.",
    ]),

    ("Nonlinear PK model", [
        "The hybrid model uses first-order inter-compartmental transfer and "
        "Michaelis\u2013Menten elimination: dP_i/dt = \u03a3 k_ji\u00b7P_j \u2212 "
        "(\u03a3 k_ij)\u00b7P_i \u2212 V_max,i\u00b7P_i / (K_m + P_i). "
        "The stretched-exponential survival S(t) = \u03a3 w_i \u00d7 "
        "exp(\u2212\u03bb_i\u00b7t) \u00d7 t^{\u2212\u03b1} was fitted by "
        "minimising log-space squared error against the empirical survival "
        "function (Nelder\u2013Mead, 10,000 iterations).",
    ]),

    ("Population PK analysis", [
        "For 64 mass configurations (m\u2081 = 1 fixed; m\u2082, m\u2083 \u2208 "
        "{0.25, 0.5, 0.75, 1, 1.5, 2, 3, 4}), PK parameters were estimated "
        "independently. The population model: log(MRT_i) = \u03b2\u2080 + "
        "\u03b2\u2081\u00b7log(\u03bc\u2081\u2082,i) + "
        "\u03b2\u2082\u00b7log(\u03bc_out,i) + "
        "\u03b2\u2083\u00b7log(M_i) + \u03b7_i, "
        "where \u03b7 ~ N(0, \u03c9\u00b2), was fitted by "
        "ordinary least squares. R\u00b2 and \u03c9 (inter-individual "
        "variability) were computed from residuals.",
    ]),

    ("TMDD analysis", [
        "The TMDD ODE system (drug C, receptor R, complex CR) was solved "
        "with typical monoclonal antibody parameters: k_syn = 0.1, k_deg = 0.01, "
        "k_on = 0.1, k_off = 0.001, k_int = 0.05, k_el = 0.01. Steady states "
        "were found numerically (fsolve with multiple initial guesses). "
        "Stability was assessed by Jacobian eigenvalues. Dose\u2013response "
        "bifurcation was computed over dose \u2208 [10\u207b\u00b3, 10\u00b2].",
    ]),

    ("Phase-space flux comparison", [
        "Theoretical transition rates were predicted from phase-space volume "
        "scaling: k_ij \u221d (\u03bc_to^{3/2} \u00d7 m_single^{3/2}) / "
        "(\u03bc_from^{3/2} \u00d7 m_single_from^{3/2}), following "
        "Stone and Leigh.{6} Escape rates scale as (m_single / M)^{3/2}. "
        "Rates were normalised to the dynamical timescale t_dyn = 1 / \u221aM.",
    ]),

    ("3D extension and gravitational-wave dissipation", [
        "The 3D extension used fully randomised initial conditions: binary "
        "orbital plane sampled uniformly on SO(3), eccentricity from the "
        "thermal distribution f(e) = 2e (capped at e = 0.95), and third-body "
        "approach direction uniform on S\u00b2. For each of 3 mass "
        "configurations, 2,000 scattering experiments were performed.",

        "Gravitational-wave dissipation was implemented via the 2.5PN "
        "radiation reaction acceleration (leading-order Burke\u2013Thorne term).{19} "
        "For each pair (i,j) with relative separation r and velocity v, the "
        "dissipative acceleration is: a_2.5PN = (8/5)\u00b7\u03bc\u00b7M\u00b2 / "
        "(c\u2075\u00b7r\u00b3) \u00d7 "
        "[{3v\u00b2 + (17/3)M/r}\u00b7v_r\u00b7\u0072\u0302 \u2212 "
        "{v\u00b2 + 3M/r}\u00b7v]. We set c = 100 in code units "
        "(v_orbital \u2248 1), giving |\u0394E/E| \u2248 0.1\u20130.9% "
        "per encounter\u2014consistent with the weak-field limit. Integration used "
        "velocity Verlet with the PN correction evaluated at the half-step "
        "velocity.",
    ]),
]

# CORRECTED REFERENCES (4 fixes applied: Ref 5, 11, 12, 13)
REFERENCES = [
    "Newton, I. Philosophiae Naturalis Principia Mathematica (1687).",
    "Poincar\u00e9, H. Les m\u00e9thodes nouvelles de la m\u00e9canique c\u00e9leste (Gauthier-Villars, 1892).",
    "Boekholt, T. & Portegies Zwart, S. On the reliability of N-body simulations. Comput. Astrophys. Cosmol. 2, 2 (2015).",
    "Heggie, D. C. Binary evolution in stellar dynamics. Mon. Not. R. Astron. Soc. 173, 729\u2013787 (1975).",
    "Ginat, Y. B. & Perets, H. B. Gravitational three-body problem: unequal masses. Phys. Rev. X 11, 031020 (2021).",
    "Stone, N. C. & Leigh, N. W. C. A statistical solution to the chaotic, non-hierarchical three-body problem. Nature 576, 406\u2013410 (2019).",
    "Monaghan, J. J. A statistical theory of the disruption of three-body systems\u2014II. Mon. Not. R. Astron. Soc. 177, 583\u2013594 (1976).",
    "Rowland, M. & Tozer, T. N. Clinical Pharmacokinetics and Pharmacodynamics: Concepts and Applications 4th edn (Lippincott Williams & Wilkins, 2011).",
    "Gabrielsson, J. & Weiner, D. Pharmacokinetic and Pharmacodynamic Data Analysis 5th edn (Swedish Pharmaceutical Press, 2016).",
    "Norris, J. R. Markov Chains (Cambridge Univ. Press, 1997).",
    "Mikkola, S. & Tanikawa, K. Explicit symplectic algorithms for time-transformed Hamiltonians. Celest. Mech. Dyn. Astron. 74, 287\u2013295 (1999).",
    "Manwadkar, V., Kol, B., Trani, A. A. & Leigh, N. W. C. Chaos and L\u00e9vy flights in the three-body problem. Mon. Not. R. Astron. Soc. 497, 3694\u20133712 (2020).",
    "Samsing, J. et al. AGN as potential factories for eccentric black hole mergers. Nature 603, 237\u2013240 (2022).",
    "Mould, D. R. & Upton, R. N. Basic concepts in population modeling, simulation, and model-based drug development. CPT Pharmacometrics Syst. Pharmacol. 1, e6 (2012).",
    "Mager, D. E. & Jusko, W. J. General pharmacokinetic model for drugs exhibiting target-mediated drug disposition. J. Pharmacokinet. Pharmacodyn. 28, 507\u2013532 (2001).",
    "Gibiansky, L. & Gibiansky, E. Target-mediated drug disposition model: approximations, identifiability of model parameters and applications to the population pharmacokinetic\u2013pharmacodynamic modeling of biologics. Expert Opin. Drug Metab. Toxicol. 5, 803\u2013812 (2009).",
    "Mardling, R. A. & Aarseth, S. J. Tidal interactions in star cluster simulations. Mon. Not. R. Astron. Soc. 321, 398\u2013420 (2001).",
    "Copeland, R. A. The drug\u2013target residence time model: a 10-year retrospective. Nat. Rev. Drug Discov. 15, 87\u201395 (2016).",
    "Peters, P. C. Gravitational radiation and the motion of two point masses. Phys. Rev. 136, B1224\u2013B1232 (1964).",
]


# ---------------------------------------------------------------------------
# Figure metadata
# ---------------------------------------------------------------------------

FIGURE_META = [
    ("fig1_conceptual.png",
     "Mapping three-body scattering onto pharmacokinetic compartments. "
     "a Three-body scattering: each binary configuration (which body pair "
     "is gravitationally bound) corresponds to a compartment; transitions "
     "between configurations are inter-compartmental transfers at rates k_ij; "
     "escape of one body is irreversible elimination at rate k_ei. "
     "b Equivalent three-compartment pharmacokinetic model."),
    ("fig2_survival.png",
     "Survival curves for three mass configurations. "
     "Blue dots: empirical survival from 5,000 N-body simulations. "
     "Red line: linear 3-compartment PK model prediction (sum of "
     "three exponentials). The multi-exponential structure is "
     "statistically necessary (Kolmogorov\u2013Smirnov test p < 0.001 versus "
     "single exponential). a Equal mass (1:1:1). b Unequal mass "
     "(1:2:0.5). c Democratic initial conditions."),
    ("fig3_nonlinear.png",
     "Linear versus nonlinear PK model comparison. Log\u2013log survival "
     "plots for three mass configurations. Blue dots: N-body "
     "simulation data. Red: linear PK (multi-exponential). Green "
     "dashed: hybrid Michaelis\u2013Menten model capturing the power-law "
     "tail from sticky chaos. a Equal mass: linear model adequate. "
     "b Unequal mass: 42% tail RMSE improvement. "
     "c Democratic: 36% improvement."),
    ("fig4_population_pk.png",
     "Population PK analysis of mass-ratio dependence. "
     "a MRT versus reduced mass of binary\u2013single system. "
     "b Median lifetime versus incoming body mass. "
     "c Mean excursions versus lightest mass fraction. "
     "d Lightest-body escape probability with logistic fit. "
     "e Population model predicted versus observed MRT "
     "(R\u00b2 = 0.67). "
     "f Heatmap of slowest half-life as function of m\u2082 and m\u2083."),
    ("fig5_tmdd.png",
     "TMDD reverse application. a Dose\u2013response bifurcation curve "
     "showing stable (blue) and unstable (red) steady states. "
     "b Receptor occupancy versus dose rate. "
     "c Slowest Jacobian eigenvalue versus dose, identifying the "
     "stability boundary (analogue of Lagrange point stability in "
     "celestial mechanics)."),
    ("fig6_3d_dissipative_comparison.png",
     "PK model validity across dimensionality and dissipation. "
     "Rows: 2D conservative, 3D conservative, 3D dissipative "
     "(c = 100). Columns: equal mass (1:1:1), unequal mass "
     "(1:2:0.5), democratic. Blue dots: N-body survival data. "
     "Red: linear 3-compartment PK fit. Kolmogorov\u2013Smirnov "
     "statistics annotated."),
    ("fig7_dissipation_effect.png",
     "Effect of gravitational-wave dissipation on lifetime "
     "distributions. Blue: 3D conservative. Red: 3D with 2.5PN "
     "radiation reaction (c = 100). Dashed: PK model fits. "
     "Dissipation systematically shortens lifetimes "
     "(\u03c4_diss/\u03c4_cons = 0.83\u20130.98) without destroying the "
     "PK model structure."),
]


# ---------------------------------------------------------------------------
# 1. Double-blind manuscript (no author info)
# ---------------------------------------------------------------------------

def build_doubleblind_manuscript():
    doc = Document()
    setup_style(doc)

    # Title only (no authors/affiliations)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(18)
    run = title_p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(16)

    # Abstract
    abs_h = doc.add_paragraph()
    run = abs_h.add_run("Abstract")
    run.bold = True
    run.font.size = Pt(12)

    abs_p = doc.add_paragraph()
    abs_p.paragraph_format.space_after = Pt(12)
    run = abs_p.add_run(ABSTRACT)
    run.font.size = Pt(11)

    # Introduction (no subheadings per Nat Comm guidelines)
    add_heading(doc, "Introduction", level=1)
    for para_text in INTRO:
        add_paragraph_with_refs(doc, para_text)

    # Figure 1 after Introduction (first mention)
    fig1_path = os.path.join(FIGDIR, "fig1_conceptual.png")
    add_figure(doc, fig1_path, FIGURE_META[0][1], 1)

    # Results
    add_heading(doc, "Results", level=1)
    fig_counter = 2
    for section_title, paragraphs in RESULTS_SECTIONS.items():
        add_heading(doc, section_title, level=2)
        for para_text in paragraphs:
            add_paragraph_with_refs(doc, para_text)

        if "Linear PK" in section_title:
            fig_path = os.path.join(FIGDIR, "fig2_survival.png")
            add_figure(doc, fig_path, FIGURE_META[1][1], fig_counter)
            fig_counter += 1
        elif "Nonlinear" in section_title:
            fig_path = os.path.join(FIGDIR, "fig3_nonlinear.png")
            add_figure(doc, fig_path, FIGURE_META[2][1], fig_counter)
            fig_counter += 1
        elif "Population" in section_title:
            fig_path = os.path.join(FIGDIR, "fig4_population_pk.png")
            add_figure(doc, fig_path, FIGURE_META[3][1], fig_counter)
            fig_counter += 1
        elif "TMDD" in section_title:
            fig_path = os.path.join(FIGDIR, "fig5_tmdd.png")
            add_figure(doc, fig_path, FIGURE_META[4][1], fig_counter)
            fig_counter += 1
        elif "3D robustness" in section_title:
            fig6_path = os.path.join(FIGDIR, "fig6_3d_dissipative_comparison.png")
            if os.path.exists(fig6_path):
                add_figure(doc, fig6_path, FIGURE_META[5][1], fig_counter)
                fig_counter += 1
            fig7_path = os.path.join(FIGDIR, "fig7_dissipation_effect.png")
            if os.path.exists(fig7_path):
                add_figure(doc, fig7_path, FIGURE_META[6][1], fig_counter)
                fig_counter += 1

    # Discussion (no subheadings per Nat Comm guidelines)
    add_heading(doc, "Discussion", level=1)
    for para_text in DISCUSSION:
        add_paragraph_with_refs(doc, para_text)

    # Methods
    doc.add_page_break()
    add_heading(doc, "Methods", level=1)
    for method_title, method_paras in METHODS:
        add_heading(doc, method_title, level=2)
        for para_text in method_paras:
            add_paragraph_with_refs(doc, para_text)

    # Data availability
    add_heading(doc, "Data availability", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "All simulation data generated in this study are available at "
        "https://github.com/bougtoir/threebody-pk-compartment. "
        "Source data are provided with this paper."
    )
    run.font.size = Pt(11)

    # Code availability
    add_heading(doc, "Code availability", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "The Julia N-body simulation code and Python analysis scripts "
        "are available at https://github.com/bougtoir/threebody-pk-compartment."
    )
    run.font.size = Pt(11)

    # References
    doc.add_page_break()
    add_heading(doc, "References", level=1)
    for i, ref in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        run_num = p.add_run(f"{i}. ")
        run_num.font.size = Pt(10)
        run_text = p.add_run(ref)
        run_text.font.size = Pt(10)

    out_path = os.path.join(OUTDIR, "manuscript_natcomm_doubleblind.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


# ---------------------------------------------------------------------------
# 2. Separate title page (author info)
# ---------------------------------------------------------------------------

def build_title_page():
    doc = Document()
    setup_style(doc)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(24)
    run = title_p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(16)

    # Authors
    auth_p = doc.add_paragraph()
    auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    auth_p.paragraph_format.space_after = Pt(6)
    run = auth_p.add_run(AUTHORS + "*")
    run.font.size = Pt(13)

    # Affiliations
    aff_p = doc.add_paragraph()
    aff_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aff_p.paragraph_format.space_after = Pt(24)
    run = aff_p.add_run(AFFILIATIONS)
    run.font.size = Pt(11)
    run.italic = True

    # Corresponding author
    doc.add_paragraph()
    corr_h = doc.add_paragraph()
    run = corr_h.add_run("Corresponding Author")
    run.bold = True
    run.font.size = Pt(12)

    corr_p = doc.add_paragraph()
    for line in CORRESPONDING.split("\n"):
        run = corr_p.add_run(line + "\n")
        run.font.size = Pt(11)

    # Author contributions
    doc.add_paragraph()
    contrib_h = doc.add_paragraph()
    run = contrib_h.add_run("Author Contributions")
    run.bold = True
    run.font.size = Pt(12)

    contrib_p = doc.add_paragraph()
    run = contrib_p.add_run(
        "T.O. conceived the study, developed the theoretical framework, "
        "designed and performed all simulations, analysed the data, "
        "created all figures, and wrote the manuscript."
    )
    run.font.size = Pt(11)

    # Competing interests
    doc.add_paragraph()
    ci_h = doc.add_paragraph()
    run = ci_h.add_run("Competing Interests")
    run.bold = True
    run.font.size = Pt(12)

    ci_p = doc.add_paragraph()
    run = ci_p.add_run("The author declares no competing interests.")
    run.font.size = Pt(11)

    # ORCID
    doc.add_paragraph()
    orcid_h = doc.add_paragraph()
    run = orcid_h.add_run("ORCID")
    run.bold = True
    run.font.size = Pt(12)

    orcid_p = doc.add_paragraph()
    run = orcid_p.add_run("Tatsuki Onishi: [ORCID to be provided]")
    run.font.size = Pt(11)

    out_path = os.path.join(OUTDIR, "title_page.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


# ---------------------------------------------------------------------------
# 3. Cover letter (Nature Communications)
# ---------------------------------------------------------------------------

def build_cover_letter():
    doc = Document()
    setup_style(doc)

    # Date
    today = date.today().strftime("%d %B %Y")
    date_p = doc.add_paragraph()
    run = date_p.add_run(today)
    run.font.size = Pt(11)
    date_p.paragraph_format.space_after = Pt(12)

    # Addressee
    addr_lines = [
        "The Editors",
        "Nature Communications",
        "Springer Nature",
    ]
    for line in addr_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.size = Pt(11)

    doc.add_paragraph()

    # Subject
    subj_p = doc.add_paragraph()
    run = subj_p.add_run("Re: ")
    run.font.size = Pt(11)
    run = subj_p.add_run("Submission of Article \u2014 ")
    run.font.size = Pt(11)
    run = subj_p.add_run(
        "\u201cPharmacokinetic compartment theory solves the statistical "
        "three-body problem\u201d"
    )
    run.font.size = Pt(11)
    run.italic = True

    doc.add_paragraph()

    # Body
    paras = [
        "Dear Editors,",

        ("We submit the above manuscript for consideration as an Article in "
         "Nature Communications. This work demonstrates that chaotic "
         "three-body gravitational scattering and pharmacokinetic compartment "
         "theory share a structural mathematical equivalence\u2014continuous-time "
         "Markov chains on three states with irreversible exit\u2014and that "
         "the analytical tools of each field can be productively imported "
         "into the other."),

        "The key findings are:",
    ]

    for text in paras:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)

    # Bullet points
    bullets = [
        ("A linear 3-compartment PK model reproduces the multi-exponential "
         "lifetime distribution of three-body scattering, validated across "
         "15,000 N-body simulations (3 mass configurations, Kolmogorov\u2013"
         "Smirnov test p < 0.001 versus single exponential)."),

        ("A nonlinear (Michaelis\u2013Menten) PK extension captures the "
         "power-law tail from sticky chaos (36\u201342% tail RMSE improvement), "
         "providing a minimal parametrisation of the ergodic-to-trapped "
         "transition."),

        ("Population PK (mixed-effects) modelling yields closed-form "
         "allometric scaling of mean residence time with mass ratios across "
         "64 configurations (MRT \u221d \u03bc\u2081\u2082^{2.08} \u00d7 "
         "\u03bc_out^{\u22121.68}, R\u00b2 = 0.67)."),

        ("The reverse mapping provides a stability analysis framework for "
         "target-mediated drug disposition (TMDD) grounded in celestial "
         "mechanics."),

        ("Extension to fully 3D scattering with 2.5PN gravitational-wave "
         "radiation reaction confirms the PK mapping is robust to both "
         "dimensionality and dissipation (KS = 0.16\u20130.31 across all "
         "configurations), with implications for LISA merger-rate predictions."),
    ]

    for text in bullets:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(text)
        run.font.size = Pt(11)

    doc.add_paragraph()

    paras2 = [
        ("We believe this work is appropriate for Nature Communications "
         "because it reveals an unexpected structural connection between "
         "celestial mechanics and clinical pharmacology that is quantitative, "
         "bidirectional, and immediately actionable by both communities. "
         "The three-body problem has been a central challenge in physics for "
         "over 250 years; pharmacokinetic compartment modelling is the "
         "foundation of modern drug development. Demonstrating that these "
         "are structurally equivalent problems, and that tools from each "
         "field solve open questions in the other, will be of significant "
         "interest to specialists in both gravitational dynamics and "
         "quantitative pharmacology."),

        ("The manuscript has not been published or submitted elsewhere. "
         "We request double-anonymised peer review."),

        "Suggested reviewers:",
    ]

    for text in paras2:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)

    # Suggested reviewers
    reviewers = [
        ("Dr. Alessandro A. Trani",
         "Niels Bohr Institute, University of Copenhagen, Denmark",
         "alessandro.trani@nbi.ku.dk",
         "Three-body problem, gravitational waves, N-body dynamics"),
        ("Dr. Silvia Toonen",
         "Anton Pannekoek Institute, University of Amsterdam, Netherlands",
         "toonen@uva.nl",
         "Gravitational-wave sources, compact binary dynamics"),
        ("Prof. Mats O. Karlsson",
         "Dept. of Pharmacy, Uppsala University, Sweden",
         "mats.karlsson@farmaci.uu.se",
         "Population pharmacokinetics, nonlinear mixed-effects modelling"),
        ("Prof. Nick Holford",
         "Dept. of Pharmacology, University of Auckland, New Zealand",
         "n.holford@auckland.ac.nz",
         "Clinical pharmacology, PK/PD modelling"),
        ("Prof. Rosalba Perna",
         "Dept. of Physics & Astronomy, Stony Brook University, USA",
         "rosalba.perna@stonybrook.edu",
         "N-body dynamics, gravitational waves, compact objects"),
    ]

    for name, affil, email, expertise in reviewers:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}, ")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(f"{affil}. ")
        run.font.size = Pt(10)
        run.italic = True
        run = p.add_run(f"{email}. ")
        run.font.size = Pt(10)
        run = p.add_run(f"Expertise: {expertise}.")
        run.font.size = Pt(10)

    doc.add_paragraph()

    # Excluded reviewers
    excl_p = doc.add_paragraph()
    run = excl_p.add_run("Excluded reviewers: ")
    run.bold = True
    run.font.size = Pt(11)
    run = excl_p.add_run(
        "Nathan W.C. Leigh (Universidad de Concepci\u00f3n), "
        "Nicholas C. Stone (Hebrew University), and "
        "Barak Kol (Hebrew University)\u2014all are authors of "
        "directly competing statistical three-body solutions."
    )
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Closing
    closing = [
        "Sincerely,",
        "",
        "Tatsuki Onishi",
        "Faculty of Data Science, Shiga University",
        "1-1-1 Banba, Hikone, Shiga 522-8522, Japan",
        "Email: bougtoir@gmail.com",
    ]

    for line in closing:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.size = Pt(11)

    out_path = os.path.join(OUTDIR, "cover_letter_natcomm.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


# ---------------------------------------------------------------------------
# 4. Create zip package
# ---------------------------------------------------------------------------

def create_zip(manuscript_path, title_page_path, cover_letter_path):
    zip_path = os.path.join(OUTDIR, "threebody_pk_natcomm_submission.zip")

    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        # Manuscript
        zf.write(manuscript_path,
                 os.path.basename(manuscript_path))
        # Title page
        zf.write(title_page_path,
                 os.path.basename(title_page_path))
        # Cover letter
        zf.write(cover_letter_path,
                 os.path.basename(cover_letter_path))

        # Editable PPTX (if exists)
        pptx_path = os.path.join(OUTDIR, "figures_editable.pptx")
        if os.path.exists(pptx_path):
            zf.write(pptx_path, "figures_editable.pptx")

        # Individual figure files (PNG + PDF)
        for fname, _ in FIGURE_META:
            png_path = os.path.join(FIGDIR, fname)
            if os.path.exists(png_path):
                zf.write(png_path, f"figures/{fname}")
            pdf_fname = fname.replace('.png', '.pdf')
            pdf_path = os.path.join(FIGDIR, pdf_fname)
            if os.path.exists(pdf_path):
                zf.write(pdf_path, f"figures/{pdf_fname}")

    print(f"Saved: {zip_path}")
    return zip_path


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    print("=" * 60)
    print("  Generating Nature Communications Submission Package")
    print("  (Double-Blind)")
    print("=" * 60)

    db_path = build_doubleblind_manuscript()
    tp_path = build_title_page()
    cl_path = build_cover_letter()
    zip_path = create_zip(db_path, tp_path, cl_path)

    print(f"\nOutputs:")
    print(f"  Double-blind manuscript: {db_path}")
    print(f"  Title page: {tp_path}")
    print(f"  Cover letter: {cl_path}")
    print(f"  Submission zip: {zip_path}")
    print("=" * 60)

#!/usr/bin/env python3
"""
Generate Physical Review X (PRX) submission package:
  - manuscript_prx.docx (with author info — single-blind)
  - cover_letter_prx.docx (includes justification)
  - popular_summary.txt (~250 words, nontechnical)
  - figures_editable.pptx (already exists)

PRX formatting:
  - Research Article: no strict word limit (~20 pages before surcharge)
  - Abstract: ~5% of article length, < 500 words, no citations
  - Popular Summary: ~250 words, nontechnical, no math, required
  - Justification: ~100 words (why paper meets PRX criteria), in cover letter
  - References: APS numbered style [1], [2], ...
  - Single-blind (author info in manuscript)
  - Subject areas: up to 3 from PRX list
  - Figures: separate files
"""

from __future__ import annotations

import os
import re
import zipfile
from datetime import date

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASEDIR = os.path.dirname(os.path.abspath(__file__))
FIGDIR_NATURE = os.path.join(BASEDIR, "figures_nature")
FIGDIR_ORIG = os.path.join(BASEDIR, "figures")
FIGDIR = FIGDIR_NATURE if os.path.isdir(FIGDIR_NATURE) else FIGDIR_ORIG
OUTDIR = os.path.join(BASEDIR, "manuscript")
os.makedirs(OUTDIR, exist_ok=True)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _add_run(paragraph, text, font_size=Pt(11), superscript=False,
             subscript=False, italic=False, bold=False):
    """Add a run with optional formatting."""
    if not text:
        return
    run = paragraph.add_run(text)
    run.font.size = font_size
    if superscript:
        run.font.superscript = True
    if subscript:
        run.font.subscript = True
    if italic:
        run.italic = True
    if bold:
        run.bold = True
    return run


def _render_math_segment(paragraph, segment, font_size=Pt(11)):
    """Render a segment that may contain _{...} and ^{...} math notation.

    Handles patterns like:
      k_{ij}        → k  (normal) + ij (subscript)
      V_{max}       → V  (normal) + max (subscript)
      t^{−α}        → t  (normal) + −α (superscript)
      μ_{12}^{2.08} → μ  (normal) + 12 (subscript) + 2.08 (superscript)
      \\hat{r}      → r̂
    """
    # Handle \hat{x} → x̂ (combining circumflex)
    segment = re.sub(r'\\hat\{([^}])\}', r'\1' + '\u0302', segment)
    # Handle \\hat{x} (double-escaped)
    segment = re.sub(r'\\\\hat\{([^}])\}', r'\1' + '\u0302', segment)
    # Handle v_{\\infty} or v_{\infty}
    segment = segment.replace('\\\\infty', '∞').replace('\\infty', '∞')

    # Split into tokens: normal text, _{...}, ^{...}, _X (single char)
    # Pattern matches:
    #   _{...}  → multi-char subscript
    #   ^{...}  → multi-char superscript
    #   _X      → single-char subscript (letter or digit, not followed by {)
    tokens = re.split(
        r'(_\{[^}]*\}|\^\{[^}]*\}|_(?=[a-zA-Z0-9])([a-zA-Z0-9]))', segment
    )
    # re.split with groups produces extra captures; flatten
    # Use findall approach instead
    tokens = re.findall(
        r'_\{[^}]*\}|\^\{[^}]*\}|_[a-zA-Z0-9](?!\{)|[^_^]+|[_^]',
        segment
    )

    for token in tokens:
        if not token:
            continue
        if token.startswith('_{') and token.endswith('}'):
            inner = token[2:-1]
            _add_run(paragraph, inner, font_size=font_size, subscript=True,
                     italic=True)
        elif token.startswith('^{') and token.endswith('}'):
            inner = token[2:-1]
            _add_run(paragraph, inner, font_size=font_size, superscript=True)
        elif len(token) == 2 and token[0] == '_':
            # Single-char subscript: _e, _1, _i etc.
            _add_run(paragraph, token[1], font_size=font_size, subscript=True,
                     italic=True)
        else:
            _add_run(paragraph, token, font_size=font_size)


def add_text_with_refs(paragraph, text, font_size=Pt(11)):
    """Render text with citation references {N} and math notation.

    Citations: {1}, {5–7} → superscript [1], [5–7]
    Math: k_{ij}, V_{max}, t^{−α} → proper Word sub/superscript
    """
    # Split on citation refs {N} or {N–M} or {N,M}, but NOT when preceded
    # by _ or ^ (those are math subscript/superscript markers).
    parts = re.split(r'(?<![_^])(\{[\d,\u2013\u2014\u2212 –-]+\})', text)
    for part in parts:
        if not part:
            continue
        # Check if this is a citation reference (contains only digits,
        # commas, dashes, en-dash, spaces)
        if (part.startswith('{') and part.endswith('}')
                and re.match(r'^\{[\d,\u2013\u2014\u2212 –-]+\}$', part)):
            ref_text = part[1:-1]
            run = paragraph.add_run(f" [{ref_text}]")
            run.font.size = font_size
            run.font.superscript = True
        else:
            _render_math_segment(paragraph, part, font_size=font_size)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_paragraph_with_refs(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    add_text_with_refs(p, text)
    return p


def add_figure_placeholder(doc, caption, fig_num):
    """Add a figure placeholder (no embedded image) with caption.
    PRX requires figures as separate files uploaded to the submission system."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"[FIG. {fig_num}]")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(128, 128, 128)
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(12)
    run_bold = cap.add_run(f"FIG. {fig_num}. ")
    run_bold.bold = True
    run_bold.font.size = Pt(10)
    add_text_with_refs(cap, caption, font_size=Pt(10))


def setup_style(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5


# ---------------------------------------------------------------------------
# Content — PRX version (single-blind, with author info)
# ---------------------------------------------------------------------------

TITLE = ("Pharmacokinetic compartment theory solves the statistical "
         "three-body problem")

AUTHORS = "Tatsuki Onishi"

AFFILIATIONS = (
    "Faculty of Data Science, Shiga University, "
    "1-1-1 Banba, Hikone, Shiga 522-8522, Japan"
)

CORRESPONDING = "bougtoir@gmail.com"

# Popular Summary: ~250 words, nontechnical, no math. Required by PRX.
POPULAR_SUMMARY = (
    "When three stars meet in space, their gravitational dance is chaotic "
    "and unpredictable\u2014a puzzle that has challenged physicists since "
    "Newton. Meanwhile, in hospitals and pharmacies, scientists have spent "
    "decades perfecting mathematical models that describe how drugs move "
    "through the human body, passing between blood, tissues, and organs "
    "before being eliminated.\n\n"
    "This study reveals a surprising connection: the mathematics governing "
    "these two seemingly unrelated phenomena are structurally identical. "
    "In the three-body problem, at any moment two stars form a pair while "
    "the third orbits at a distance. The pair can swap partners, and "
    "eventually one star is flung away permanently. In pharmacokinetics, "
    "a drug molecule occupies one body compartment (say, blood), can "
    "transfer to another (say, liver), and is eventually eliminated from "
    "the body. Both processes are described by the same type of random "
    "process\u2014a continuous-time Markov chain on three states with "
    "irreversible exit.\n\n"
    "By running 15,000 computer simulations of three-body encounters "
    "across 64 different mass combinations, the author shows that "
    "pharmacokinetic models quantitatively reproduce the statistical "
    "behavior of three-body scattering\u2014including how long the "
    "interaction lasts, which star gets ejected, and how the outcome "
    "depends on the masses involved. The mapping works in both "
    "directions: drug-design tools from clinical medicine can be applied "
    "to astrophysics, and stability analysis from celestial mechanics "
    "can inform pharmacology. The correspondence holds even in realistic "
    "three-dimensional encounters with gravitational-wave energy loss, "
    "suggesting broad applicability to real astrophysical systems."
)

# Abstract: < 500 words, no citations, self-contained
ABSTRACT = (
    "The gravitational three-body problem remains analytically intractable "
    "after 250 years, yet its statistical properties follow remarkably "
    "regular patterns. Here we show that chaotic three-body scattering "
    "maps onto a three-compartment pharmacokinetic (PK) model\u2014where "
    "binary configurations are compartments, configuration transitions "
    "are inter-compartmental transfers, and system dissolution is "
    "elimination\u2014and that this structural correspondence is "
    "quantitative, bidirectional, and robust. Using 15,000 N-body "
    "simulations across 64 mass configurations, we demonstrate that "
    "(i) a linear three-compartment PK model reproduces the "
    "multi-exponential lifetime distribution predicted by random-walk "
    "theory; (ii) a nonlinear (Michaelis\u2013Menten) elimination term "
    "captures the power-law tail from sticky chaos, reducing tail RMSE "
    "by 36\u201342%; (iii) population PK (mixed-effects) modeling reveals "
    "allometric scaling of mean residence time with mass ratios "
    "(MRT \u221d \u03bc_{12}^{2.08} \u00d7 "
    "\u03bc_{out}^{\u22121.68}, R\u00b2 = 0.67); and "
    "(iv) the reverse mapping yields a stability analysis framework for "
    "target-mediated drug disposition derived from celestial mechanics. "
    "Extending to fully three-dimensional scattering with 2.5 "
    "post-Newtonian gravitational-wave radiation reaction confirms that "
    "the PK mapping is robust across dimensionality and dissipation. "
    "This cross-disciplinary bridge opens computational strategies in "
    "both celestial mechanics and quantitative pharmacology."
)

INTRO = [
    ("The three-body problem\u2014predicting the long-term motion of three "
     "gravitationally interacting masses\u2014has resisted analytical solution "
     "since Newton.{1} Poincar\u00e9 proved its non-integrability,{2} and "
     "modern numerical studies confirm that generic three-body encounters are "
     "chaotic, with outcomes sensitive to initial conditions at the level of "
     "floating-point precision.{3} Nevertheless, the statistical properties "
     "of the outcomes are remarkably regular: the lifetime distribution follows "
     "a multi-exponential decay,{4,5} and the final-state distributions can be "
     "predicted from phase-space volume arguments.{6,7}"),

    ("Recently, Ginat and Perets showed that chaotic three-body scattering "
     "can be decomposed into a sequence of independent probabilistic "
     "excursions\u2014each a temporary binary plus a distant single body\u2014and "
     "derived an analytical, statistical solution as a random walk in "
     "binary binding energy.{5} Stone and Leigh independently demonstrated "
     "the predictive power of flux-based statistical mechanics for three-body "
     "outcomes.{6} These advances establish that the ergodic core of the "
     "three-body problem is amenable to stochastic modeling."),

    ("Pharmacokinetic (PK) compartmental analysis is the standard framework "
     "for modeling drug absorption, distribution, and elimination in "
     "clinical medicine.{8,9} In PK theory, the body is divided into "
     "compartments (e.g., plasma, tissues), drug transfers between "
     "compartments at first-order rates, and is eliminated irreversibly. "
     "The mathematics is a continuous-time Markov chain (CTMC) whose "
     "generator is the PK rate matrix.{10}"),

    ("Here, we show that three-body scattering is a three-compartment PK "
     "system: each binary configuration (which body pair is bound) is a "
     "compartment; transitions between configurations are inter-compartmental "
     "transfers; escape of one body is elimination. This mapping is not merely "
     "analogical\u2014it is a structural equivalence at the level of the master "
     "equation. We exploit it to import the full toolkit of PK analysis "
     "(half-lives, mean residence time, bioequivalence, population modeling, "
     "nonlinear kinetics) into celestial mechanics, and conversely to bring "
     "dynamical-systems stability theory into pharmacology."),
]

RESULTS_SECTIONS = {
    "Linear PK model reproduces multi-exponential decay": [
        ("We performed 15,000 three-body scattering simulations using a "
         "symplectic leapfrog integrator (Sec. Methods), spanning three mass "
         "configurations: equal mass (1:1:1), unequal mass (1:2:0.5), and "
         "democratic initial conditions (1:1:1, symmetric start). For each "
         "simulation, we recorded the full sequence of binary configurations "
         "visited and dwell times in each (Fig. 1)."),

        ("From the transition counts and dwell times, we estimated the 9 PK "
         "rate parameters (6 inter-compartmental transfer rates k_{ij} and 3 "
         "elimination rates k_{ei}) by maximum-likelihood estimation for a "
         "continuous-time Markov chain: k_{ij} = N_{ij} / T_i, where N_{ij} "
         "is the number of observed i\u2192j transitions and T_i is total dwell "
         "time in compartment i (Sec. Methods)."),

        ("The resulting linear PK model predicts a survival function S(t) that "
         "is a sum of three exponentials, each corresponding to an eigenvalue "
         "of the rate matrix. This multi-exponential prediction matches the "
         "empirical survival curves from N-body simulations across all three "
         "mass configurations (Fig. 2), with the single-exponential model "
         "rejected by the Kolmogorov\u2013Smirnov test "
         "(p < 0.001 for all three configurations), confirming that the "
         "three-compartment model is necessary and sufficient to describe "
         "the intermediate-timescale dynamics."),

        ("The MLE-estimated rates agree quantitatively with the phase-space "
         "flux prediction from statistical mechanics.{6} For equal "
         "masses, all transfer rates converge to the same value by symmetry "
         "(k_{ij} \u2248 0.015 per dynamical time), validating both the "
         "PK framework and the ergodic hypothesis underlying the statistical "
         "approach."),
    ],

    "Nonlinear PK captures sticky chaos": [
        ("While the linear PK model fits the bulk of the lifetime distribution, "
         "it systematically underestimates the long-lived tail\u2014a signature of "
         "\u201csticky chaos,\u201d where trajectories become trapped near periodic "
         "orbits for anomalously long times.{11,12} In PK language, this "
         "corresponds to saturable (capacity-limited) elimination: when the "
         "system lingers near a stable periodic orbit, the escape rate "
         "decreases rather than remaining constant."),

        ("We model this by replacing the first-order elimination k_e \u00d7 P "
         "with Michaelis\u2013Menten kinetics: V_{max} \u00d7 P / (K_m + P). "
         "This hybrid model\u2014linear inter-compartmental transfer with nonlinear "
         "elimination\u2014produces a survival function whose tail decays as "
         "t^{\u2212\u03b1} \u00d7 exp(\u2212\u03bb_{slow} \u00d7 t), "
         "capturing the power-law correction (Fig. 3)."),

        ("Fitting the hybrid model to all three mass configurations yields "
         "power-law exponents \u03b1 \u2248 1.1\u20132.6, with tail "
         "RMSE improvement of 36\u201342% over the linear model for unequal and "
         "democratic configurations. For the equal-mass case, the linear model "
         "is already adequate (no significant tail excess), consistent with the "
         "known weaker stickiness in the symmetric case.{13}"),
    ],

    "Population PK reveals allometric scaling": [
        ("Clinical PK uses population (mixed-effects) modeling to relate "
         "individual PK parameters to patient covariates such as body weight "
         "and renal function.{14} We apply the same methodology to the "
         "three-body problem: the \u201ccovariates\u201d are the mass ratios of the "
         "three bodies."),

        ("Scanning 64 mass configurations (m_2, m_3 \u2208 "
         "{0.25, 0.5, 0.75, 1, 1.5, 2, 3, 4} with m_1 = 1), "
         "we fit a log-linear population model: "
         "log(MRT) = \u03b2_0 + \u03b2_1 log(\u03bc_{12}) "
         "+ \u03b2_2 log(\u03bc_{out}) + \u03b2_3 log(M) + \u03b7, "
         "where \u03bc_{12} is the reduced mass of the initial binary, "
         "\u03bc_{out} is the binary\u2013single reduced mass, and M is total mass "
         "(Fig. 4)."),

        ("The population model yields MRT \u221d \u03bc_{12}^{2.08} "
         "\u00d7 \u03bc_{out}^{\u22121.68} \u00d7 M^{\u22120.62}, "
         "with R\u00b2 = 0.67 and inter-individual variability \u03c9 = 0.35. "
         "The positive exponent on \u03bc_{12} indicates that heavier "
         "binaries survive longer (deeper potential wells), while the negative "
         "exponent on \u03bc_{out} indicates that heavier incoming bodies disrupt "
         "the system faster\u2014both physically intuitive and quantitatively "
         "informative."),

        ("Additionally, the escape probability of the lightest body follows a "
         "logistic model in mass fraction, analogous to the dose\u2013response "
         "relationship in pharmacology. This provides a closed-form predictor "
         "for which body will be ejected, a result of practical value for "
         "stellar dynamics simulations."),
    ],

    "Reverse mapping: TMDD stability analysis": [
        ("The structural equivalence works in both directions. Target-mediated "
         "drug disposition (TMDD)\u2014where a drug binds a receptor to form a "
         "complex that is then internalized\u2014is a three-state system with "
         "identical matrix structure to the three-body compartment model.{15,16}"),

        ("We demonstrate that the Jacobian eigenvalue analysis used to assess "
         "Lagrange point stability in celestial mechanics provides a natural "
         "framework for TMDD stability analysis (Fig. 5). The dose\u2013response "
         "bifurcation curve\u2014identifying stable versus unstable steady states "
         "as a function of drug input rate\u2014is the pharmacological analogue "
         "of the Mardling\u2013Aarseth stability criterion for hierarchical "
         "triples.{17}"),

        ("This reverse mapping suggests that the rich stability theory of "
         "N-body systems (KAM tori, invariant manifolds, homoclinic tangles) "
         "could be imported wholesale into nonlinear pharmacokinetics, "
         "particularly for complex biologics with multiple binding sites."),
    ],

    "3D robustness and gravitational-wave dissipation": [
        ("To test whether the PK mapping is an artifact of the 2D planar "
         "geometry or the conservative (energy-preserving) assumption, we "
         "extended the simulations to fully three-dimensional scattering with "
         "randomized orbital planes (uniform on SO(3)), thermal eccentricity "
         "(f(e) = 2e), and isotropic third-body approach directions "
         "(2,000 runs per mass configuration)."),

        ("The three-compartment PK model remains valid in all 3D configurations, "
         "with Kolmogorov\u2013Smirnov statistics of 0.16\u20130.31 across equal, "
         "unequal, and democratic mass setups\u2014comparable to the 2D values "
         "(Fig. 6). The median lifetimes shift modestly (3D conservative: "
         "246\u2013438 dynamical times versus 2D: 239\u20131,043), reflecting the "
         "larger phase space available in 3D."),

        ("Adding 2.5 post-Newtonian gravitational-wave radiation reaction{19} "
         "(c = 100 in code units; v_{orbital} \u2248 1) systematically "
         "shortens lifetimes without destroying the compartment structure "
         "(Fig. 7). The lifetime ratio \u03c4_{diss}/\u03c4_{cons} ranges from "
         "0.98 (equal mass) to 0.83 (democratic), with median fractional energy "
         "loss |\u0394E/E| = 0.10\u20130.86% per encounter. The democratic "
         "configuration shows the strongest dissipation effect (17% lifetime "
         "reduction), consistent with the longer interaction times providing "
         "more opportunity for gravitational-wave energy extraction."),

        ("These results indicate that the PK mapping is robust across "
         "dimensionality, eccentricity distribution, and the presence of "
         "dissipation\u2014and is therefore applicable to astrophysically "
         "realistic encounters."),
    ],
}

DISCUSSION = [
    ("We have demonstrated a quantitative structural correspondence between two "
     "seemingly unrelated fields: celestial mechanics and pharmacokinetics. "
     "The key insight is that both describe irreversible first-passage "
     "problems in multi-state continuous-time Markov chains, and that the "
     "mathematical structure\u2014not merely the analogy\u2014is shared."),

    ("The practical implications are bidirectional. For celestial mechanics: "
     "(i) the population PK framework provides a principled way to parameterize "
     "outcome statistics as a function of mass ratios, replacing ad hoc fitting "
     "functions;{5\u20137} (ii) the Michaelis\u2013Menten model offers a minimal "
     "parametrization of sticky chaos that interpolates between the ergodic "
     "(linear) regime and the trapped (power-law) regime; (iii) PK software "
     "(NONMEM, Monolix) can be directly applied to N-body simulation output."),

    ("For pharmacology: (i) the dynamical-systems stability framework "
     "(Lagrange points, invariant manifolds) provides geometric insight into "
     "TMDD steady states that goes beyond the standard eigenvalue analysis; "
     "(ii) the concept of \u201csticky chaos\u201d maps to the clinically important "
     "phenomenon of prolonged drug\u2013receptor residence time;{18} (iii) the "
     "allometric scaling framework transfers predictive power across species "
     "and across mass configurations."),

    ("The robustness of the PK mapping to dimensionality and dissipation "
     "(Figs. 6\u20137) has astrophysical implications. The 17% "
     "lifetime reduction in democratic configurations under gravitational-wave "
     "dissipation may inform merger-rate estimates for "
     "LISA-band sources: three-body encounters in dense clusters will resolve "
     "faster than energy-conserving estimates suggest, potentially increasing the "
     "predicted rate of observable inspirals. The population PK framework "
     "may provide closed-form event-rate formulae parameterized by cluster "
     "mass function, potentially complementing expensive Monte Carlo surveys."),

    ("Remaining limitations include the perturbative treatment of radiation "
     "reaction (valid only for c \u226b v_{orbital}) and the absence of tidal "
     "dissipation relevant for stellar encounters. The linear PK model "
     "assumes ergodicity, which breaks down for very long lifetimes; the "
     "nonlinear extension addresses this but at the cost of additional "
     "parameters."),

    ("More broadly, the correspondence inverts the conventional flow of "
     "methodology between the physical sciences and clinical medicine. "
     "The mathematical theory of pharmacokinetics was distilled from bedside "
     "observations\u2014therapeutic drug monitoring, dose\u2013response studies, "
     "clinical trial design\u2014and aged over six decades of patient-facing "
     "application into a robust quantitative infrastructure. That "
     "infrastructure has been casked in regulatory-grade software (NONMEM, "
     "Monolix) and in saturation-kinetics theory refined through countless "
     "dosing studies (Michaelis\u2013Menten, TMDD). The present work decants this "
     "accumulated clinical knowledge into an entirely new vessel\u2014celestial "
     "mechanics\u2014where, aerated by contact with a different physical context, "
     "it reveals latent capabilities that were invisible within medicine alone. "
     "Several concrete directions follow. "
     "(i) The three-compartment model generalizes naturally to N-body "
     "hierarchies: quadruple and higher-order multiples in dense star clusters "
     "map onto N-compartment cascade models whose analytical solutions are "
     "textbook clinical pharmacokinetics.{8,9} "
     "(ii) Population PK covariate modeling\u2014originally designed to account "
     "for patient weight, renal function, and genotype in clinical trials\u2014can "
     "treat mass ratios, angular momenta, and cluster environments as "
     "covariates in nonlinear mixed-effects regressions, enabling systematic "
     "astrophysical parameter scans on platforms built for drug regulation. "
     "(iii) The closed-form MRT scaling (MRT \u221d "
     "\u03bc_{12}^\u03b1 \u00d7 \u03bc_{out}^\u03b2 \u00d7 M^\u03b3) can be "
     "convolved with a cluster mass function to yield gravitational-wave "
     "event-rate estimates for LISA-band sources without expensive Monte Carlo "
     "surveys. "
     "(iv) Conversely, the geometric stability theory of celestial mechanics "
     "(KAM tori, invariant manifolds) may inform the rational design of drugs "
     "with controlled target residence times\u2014a key determinant of in vivo "
     "efficacy that is central to modern drug discovery.{18} "
     "(v) Because the mapping relies only on the CTMC structure of "
     "multi-state first-passage processes, it is not limited to gravitational "
     "systems; chemical reaction networks, ecological predator\u2013prey dynamics, "
     "and other multi-state chaotic systems may admit analogous descriptions "
     "rooted in clinical pharmacokinetic theory, suggesting a broader "
     "program of \u201cpharmacokinetics of chaos.\u201d"),

    ("In conclusion, the pharmacokinetic compartment model provides both a "
     "conceptual lens and a computational toolkit for the three-body problem "
     "that complements existing statistical-mechanical approaches. By "
     "demonstrating that clinical medicine\u2014through its pharmacokinetic "
     "tradition\u2014can export, not merely import, quantitative frameworks to "
     "the physical sciences, this work illustrates how disciplinary "
     "boundaries may conceal unexploited structural connections."),
]

METHODS = [
    ("N-body simulations", [
        "Three-body scattering simulations were performed using a symplectic "
        "leapfrog (St\u00f6rmer\u2013Verlet) integrator implemented in Julia 1.10. "
        "Initial conditions: binary with semi-major axis a = 1 and eccentricity "
        "e = 0; third body approaching from r = 10a with velocity at infinity "
        "v_{\\infty} = 0.1 (parabolic-like encounters). Impact parameter b sampled "
        "uniformly in b^2 from 0 to b_{max} = 5a. Integration continued until "
        "escape (body reaching r > 20a with positive energy) or t_{max} = 10,000 "
        "dynamical times. Adaptive timestep: dt = 0.01 \u00d7 r_{min} / v_{max}. "
        "Energy conservation better than 10^{\u22128} in all runs.",

        "For each mass configuration, 5,000 scattering experiments were "
        "performed (total: 15,000 for Phase 1; 5,000 \u00d7 64 = 320,000 for "
        "the population PK scan). Binary configuration was identified at each "
        "timestep by finding the most-bound pair (most negative pairwise "
        "energy). A configuration transition was recorded when the most-bound "
        "pair identity changed and persisted for more than 5 timesteps.",
    ]),

    ("PK model estimation", [
        "Transition rates were estimated by maximum-likelihood for a "
        "continuous-time Markov chain: k_{ij} = N_{ij} / T_i, where N_{ij} is "
        "the count of transitions from compartment i to j, and T_i is total "
        "observed dwell time in compartment i. The 3\u00d73 rate matrix A was "
        "constructed such that A_{ij} = k_{ji} for i \u2260 j, and "
        "A_{ii} = \u2212\u03a3(k_{ij} + k_{ei}).",

        "Survival function: S(t) = 1^T exp(At) P_0. "
        "Eigenvalues of A give the three decay rates; half-lives t_{1/2} = "
        "ln(2) / |\u03bb_i|. Mean residence time (MRT) = \u22121^T A^{\u22121} P_0.",
    ]),

    ("Nonlinear PK model", [
        "The hybrid model uses first-order inter-compartmental transfer and "
        "Michaelis\u2013Menten elimination: dP_i/dt = \u03a3 k_{ji} P_j \u2212 "
        "(\u03a3 k_{ij}) P_i \u2212 V_{max,i} P_i / (K_m + P_i). "
        "The stretched-exponential survival S(t) = \u03a3 w_i "
        "exp(\u2212\u03bb_i t) t^{\u2212\u03b1} was fitted by "
        "minimizing log-space squared error against the empirical survival "
        "function (Nelder\u2013Mead, 10,000 iterations).",
    ]),

    ("Population PK analysis", [
        "For 64 mass configurations (m_1 = 1 fixed; m_2, m_3 \u2208 "
        "{0.25, 0.5, 0.75, 1, 1.5, 2, 3, 4}), PK parameters were estimated "
        "independently. The population model: log(MRT_i) = \u03b2_0 + "
        "\u03b2_1 log(\u03bc_{12,i}) + "
        "\u03b2_2 log(\u03bc_{out,i}) + "
        "\u03b2_3 log(M_i) + \u03b7_i, "
        "where \u03b7 ~ N(0, \u03c9^2), was fitted by "
        "ordinary least squares. R\u00b2 and \u03c9 (inter-individual "
        "variability) were computed from residuals.",
    ]),

    ("TMDD analysis", [
        "The TMDD ODE system (drug C, receptor R, complex CR) was solved "
        "with typical monoclonal antibody parameters: k_{syn} = 0.1, "
        "k_{deg} = 0.01, k_{on} = 0.1, k_{off} = 0.001, k_{int} = 0.05, "
        "k_{el} = 0.01. Steady states were found numerically (fsolve with "
        "multiple initial guesses). Stability was assessed by Jacobian "
        "eigenvalues. Dose\u2013response bifurcation was computed over "
        "dose \u2208 [10^{\u22123}, 10^2].",
    ]),

    ("Phase-space flux comparison", [
        "Theoretical transition rates were predicted from phase-space volume "
        "scaling: k_{ij} \u221d (\u03bc_{to}^{3/2} m_{single}^{3/2}) / "
        "(\u03bc_{from}^{3/2} m_{single,from}^{3/2}), following "
        "Stone and Leigh.{6} Escape rates scale as (m_{single} / M)^{3/2}. "
        "Rates were normalized to the dynamical timescale t_{dyn} = 1 / M^{1/2}.",
    ]),

    ("3D extension and gravitational-wave dissipation", [
        "The 3D extension used fully randomized initial conditions: binary "
        "orbital plane sampled uniformly on SO(3), eccentricity from the "
        "thermal distribution f(e) = 2e (capped at e = 0.95), and third-body "
        "approach direction uniform on S^2. For each of 3 mass "
        "configurations, 2,000 scattering experiments were performed.",

        "Gravitational-wave dissipation was implemented via the 2.5PN "
        "radiation reaction acceleration (leading-order Burke\u2013Thorne term).{19} "
        "For each pair (i,j) with relative separation r and velocity v, the "
        "dissipative acceleration is: a_{2.5PN} = (8/5) \u03bc M^2 / "
        "(c^5 r^3) [{3v^2 + (17/3)M/r} v_r \\hat{r} \u2212 "
        "{v^2 + 3M/r} v]. We set c = 100 in code units "
        "(v_{orbital} \u2248 1), giving |\u0394E/E| \u2248 0.1\u20130.9% "
        "per encounter\u2014consistent with the weak-field limit. Integration used "
        "velocity Verlet with the PN correction evaluated at the half-step "
        "velocity.",
    ]),
]

# CORRECTED REFERENCES (APS style: [N] Author, Title, Journal Volume, Pages (Year).)
REFERENCES = [
    "I. Newton, Philosophiae Naturalis Principia Mathematica (1687).",
    "H. Poincar\u00e9, Les M\u00e9thodes Nouvelles de la M\u00e9canique C\u00e9leste (Gauthier-Villars, Paris, 1892).",
    "T. Boekholt and S. Portegies Zwart, On the reliability of N-body simulations, Comput. Astrophys. Cosmol. 2, 2 (2015).",
    "D. C. Heggie, Binary evolution in stellar dynamics, Mon. Not. R. Astron. Soc. 173, 729 (1975).",
    "Y. B. Ginat and H. B. Perets, Gravitational three-body problem: Unequal masses, Phys. Rev. X 11, 031020 (2021).",
    "N. C. Stone and N. W. C. Leigh, A statistical solution to the chaotic, non-hierarchical three-body problem, Nature (London) 576, 406 (2019).",
    "J. J. Monaghan, A statistical theory of the disruption of three-body systems\u2014II, Mon. Not. R. Astron. Soc. 177, 583 (1976).",
    "M. Rowland and T. N. Tozer, Clinical Pharmacokinetics and Pharmacodynamics: Concepts and Applications, 4th ed. (Lippincott Williams & Wilkins, Philadelphia, 2011).",
    "J. Gabrielsson and D. Weiner, Pharmacokinetic and Pharmacodynamic Data Analysis, 5th ed. (Swedish Pharmaceutical Press, Stockholm, 2016).",
    "J. R. Norris, Markov Chains (Cambridge University Press, Cambridge, 1997).",
    "S. Mikkola and K. Tanikawa, Explicit symplectic algorithms for time-transformed Hamiltonians, Celest. Mech. Dyn. Astron. 74, 287 (1999).",
    "V. Manwadkar, B. Kol, A. A. Trani, and N. W. C. Leigh, Chaos and L\u00e9vy flights in the three-body problem, Mon. Not. R. Astron. Soc. 497, 3694 (2020).",
    "J. Samsing et al., AGN as potential factories for eccentric black hole mergers, Nature (London) 603, 237 (2022).",
    "D. R. Mould and R. N. Upton, Basic concepts in population modeling, simulation, and model-based drug development, CPT: Pharmacometrics Syst. Pharmacol. 1, e6 (2012).",
    "D. E. Mager and W. J. Jusko, General pharmacokinetic model for drugs exhibiting target-mediated drug disposition, J. Pharmacokinet. Pharmacodyn. 28, 507 (2001).",
    "L. Gibiansky and E. Gibiansky, Target-mediated drug disposition model: Approximations, identifiability of model parameters and applications to the population pharmacokinetic\u2013pharmacodynamic modeling of biologics, Expert Opin. Drug Metab. Toxicol. 5, 803 (2009).",
    "R. A. Mardling and S. J. Aarseth, Tidal interactions in star cluster simulations, Mon. Not. R. Astron. Soc. 321, 398 (2001).",
    "R. A. Copeland, The drug\u2013target residence time model: A 10-year retrospective, Nat. Rev. Drug Discovery 15, 87 (2016).",
    "P. C. Peters, Gravitational radiation and the motion of two point masses, Phys. Rev. 136, B1224 (1964).",
]


# ---------------------------------------------------------------------------
# Figure metadata (APS style: FIG. 1., panels (a), (b), etc.)
# ---------------------------------------------------------------------------

FIGURE_META = [
    ("fig1_conceptual.png",
     "Mapping three-body scattering onto pharmacokinetic compartments. "
     "(a) Three-body scattering: each binary configuration (which body pair "
     "is gravitationally bound) corresponds to a compartment; transitions "
     "between configurations are inter-compartmental transfers at rates k_{ij}; "
     "escape of one body is irreversible elimination at rate k_{ei}. "
     "(b) Equivalent three-compartment pharmacokinetic model."),
    ("fig2_survival.png",
     "Survival curves for three mass configurations. "
     "Blue dots: empirical survival from 5,000 N-body simulations. "
     "Red line: linear 3-compartment PK model prediction (sum of "
     "three exponentials). The multi-exponential structure is "
     "statistically necessary (Kolmogorov\u2013Smirnov test p < 0.001 versus "
     "single exponential). (a) Equal mass (1:1:1). (b) Unequal mass "
     "(1:2:0.5). (c) Democratic initial conditions."),
    ("fig3_nonlinear.png",
     "Linear versus nonlinear PK model comparison. Log-log survival "
     "plots for three mass configurations. Blue dots: N-body "
     "simulation data. Red: linear PK (multi-exponential). Green "
     "dashed: hybrid Michaelis\u2013Menten model capturing the power-law "
     "tail from sticky chaos. (a) Equal mass: linear model adequate. "
     "(b) Unequal mass: 42% tail RMSE improvement. "
     "(c) Democratic: 36% improvement."),
    ("fig4_population_pk.png",
     "Population PK analysis of mass-ratio dependence. "
     "(a) MRT versus reduced mass of binary\u2013single system. "
     "(b) Median lifetime versus incoming body mass. "
     "(c) Mean excursions versus lightest mass fraction. "
     "(d) Lightest-body escape probability with logistic fit. "
     "(e) Population model predicted versus observed MRT "
     "(R\u00b2 = 0.67). "
     "(f) Heatmap of slowest half-life as function of m_2 and m_3."),
    ("fig5_tmdd.png",
     "TMDD reverse application. (a) Dose\u2013response bifurcation curve "
     "showing stable (blue) and unstable (red) steady states. "
     "(b) Receptor occupancy versus dose rate. "
     "(c) Slowest Jacobian eigenvalue versus dose, identifying the "
     "stability boundary (analogue of Lagrange point stability in "
     "celestial mechanics)."),
    ("fig6_3d_dissipative_comparison.png",
     "PK model validity across dimensionality and dissipation. "
     "Rows: 2D conservative, 3D conservative, 3D dissipative "
     "(c = 100). Columns: equal mass (1:1:1), unequal mass "
     "(1:2:0.5), democratic. Blue dots: N-body survival data. "
     "Red: linear 3-compartment PK fit. Kolmogorov\u2013Smirnov "
     "statistics annotated."),
    ("fig7_dissipation_effect.png",
     "Effect of gravitational-wave dissipation on lifetime "
     "distributions. Blue: 3D conservative. Red: 3D with 2.5PN "
     "radiation reaction (c = 100). Dashed: PK model fits. "
     "Dissipation systematically shortens lifetimes "
     "(\u03c4_{diss}/\u03c4_{cons} = 0.83\u20130.98) without destroying the "
     "PK model structure."),
]


# ---------------------------------------------------------------------------
# 1. Manuscript (single-blind, with author info)
# ---------------------------------------------------------------------------

def build_manuscript():
    doc = Document()
    setup_style(doc)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(12)
    run = title_p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(16)

    # Authors (single-blind: include author info)
    auth_p = doc.add_paragraph()
    auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    auth_p.paragraph_format.space_after = Pt(6)
    run = auth_p.add_run(AUTHORS)
    run.font.size = Pt(13)

    # Affiliations
    aff_p = doc.add_paragraph()
    aff_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aff_p.paragraph_format.space_after = Pt(18)
    run = aff_p.add_run(AFFILIATIONS)
    run.font.size = Pt(11)
    run.italic = True

    # Abstract
    abs_h = doc.add_paragraph()
    run = abs_h.add_run("Abstract")
    run.bold = True
    run.font.size = Pt(12)

    abs_p = doc.add_paragraph()
    abs_p.paragraph_format.space_after = Pt(12)
    add_text_with_refs(abs_p, ABSTRACT)

    # Introduction
    add_heading(doc, "I. INTRODUCTION", level=1)
    for para_text in INTRO:
        add_paragraph_with_refs(doc, para_text)

    # Figure 1 (placeholder — figures uploaded separately)
    add_figure_placeholder(doc, FIGURE_META[0][1], 1)

    # Results
    add_heading(doc, "II. RESULTS", level=1)
    section_labels = ["A", "B", "C", "D", "E"]
    fig_counter = 2
    for idx, (section_title, paragraphs) in enumerate(RESULTS_SECTIONS.items()):
        label = section_labels[idx] if idx < len(section_labels) else ""
        add_heading(doc, f"{label}. {section_title}", level=2)
        for para_text in paragraphs:
            add_paragraph_with_refs(doc, para_text)

        if "Linear PK" in section_title:
            add_figure_placeholder(doc, FIGURE_META[1][1], fig_counter)
            fig_counter += 1
        elif "Nonlinear" in section_title:
            add_figure_placeholder(doc, FIGURE_META[2][1], fig_counter)
            fig_counter += 1
        elif "Population" in section_title:
            add_figure_placeholder(doc, FIGURE_META[3][1], fig_counter)
            fig_counter += 1
        elif "TMDD" in section_title:
            add_figure_placeholder(doc, FIGURE_META[4][1], fig_counter)
            fig_counter += 1
        elif "3D robustness" in section_title:
            add_figure_placeholder(doc, FIGURE_META[5][1], fig_counter)
            fig_counter += 1
            add_figure_placeholder(doc, FIGURE_META[6][1], fig_counter)
            fig_counter += 1

    # Discussion
    add_heading(doc, "III. DISCUSSION", level=1)
    for para_text in DISCUSSION:
        add_paragraph_with_refs(doc, para_text)

    # Methods
    doc.add_page_break()
    add_heading(doc, "IV. METHODS", level=1)
    method_labels = ["A", "B", "C", "D", "E", "F", "G"]
    for idx, (method_title, method_paras) in enumerate(METHODS):
        label = method_labels[idx] if idx < len(method_labels) else ""
        add_heading(doc, f"{label}. {method_title}", level=2)
        for para_text in method_paras:
            add_paragraph_with_refs(doc, para_text)

    # Data availability
    add_heading(doc, "DATA AVAILABILITY", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "All simulation data generated in this study are available at "
        "https://github.com/bougtoir/threebody-pk-compartment."
    )
    run.font.size = Pt(11)

    # Code availability
    add_heading(doc, "CODE AVAILABILITY", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "The Julia N-body simulation code and Python analysis scripts "
        "are available at https://github.com/bougtoir/threebody-pk-compartment."
    )
    run.font.size = Pt(11)

    # Acknowledgments
    add_heading(doc, "ACKNOWLEDGMENTS", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "The author thanks the developers of the Julia programming language "
        "and the Python scientific computing ecosystem."
    )
    run.font.size = Pt(11)

    # References (APS style: [N])
    doc.add_page_break()
    add_heading(doc, "REFERENCES", level=1)
    for i, ref in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        run_num = p.add_run(f"[{i}] ")
        run_num.font.size = Pt(10)
        run_text = p.add_run(ref)
        run_text.font.size = Pt(10)

    out_path = os.path.join(OUTDIR, "manuscript_prx.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


# ---------------------------------------------------------------------------
# 2. Cover letter (includes justification)
# ---------------------------------------------------------------------------

def build_cover_letter():
    doc = Document()
    setup_style(doc)

    # Date
    today = date.today().strftime("%d %B %Y")
    date_p = doc.add_paragraph()
    run = date_p.add_run(today)
    run.font.size = Pt(11)
    date_p.paragraph_format.space_after = Pt(12)

    # Addressee
    addr_lines = [
        "The Editors",
        "Physical Review X",
        "American Physical Society",
    ]
    for line in addr_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.size = Pt(11)

    doc.add_paragraph()

    # Subject
    subj_p = doc.add_paragraph()
    run = subj_p.add_run("Re: ")
    run.font.size = Pt(11)
    run = subj_p.add_run("Submission of Research Article \u2014 ")
    run.font.size = Pt(11)
    run = subj_p.add_run(
        "\u201cPharmacokinetic compartment theory solves the statistical "
        "three-body problem\u201d"
    )
    run.font.size = Pt(11)
    run.italic = True

    doc.add_paragraph()

    # Body
    paras = [
        "Dear Editors,",

        ("We submit the above manuscript for consideration as a Research "
         "Article in Physical Review X. This work demonstrates that chaotic "
         "three-body gravitational scattering and pharmacokinetic compartment "
         "theory share a structural mathematical equivalence at the level of "
         "the continuous-time Markov chain master equation, and that the "
         "analytical tools of each field can be productively imported into "
         "the other."),
    ]

    for text in paras:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)

    doc.add_paragraph()

    # Justification (required by PRX, ~100 words)
    just_h = doc.add_paragraph()
    run = just_h.add_run("Justification for PRX:")
    run.bold = True
    run.font.size = Pt(11)

    just_p = doc.add_paragraph()
    add_text_with_refs(just_p,
        "This paper establishes an unexpected structural equivalence between "
        "the statistical three-body problem and pharmacokinetic compartment "
        "theory\u2014two fields with no prior connection. The mapping is not "
        "merely analogical but rests on identical continuous-time Markov chain "
        "generators, validated quantitatively across 15,000+ N-body simulations. "
        "It enables bidirectional transfer of analytical tools: population PK "
        "methods yield closed-form allometric scaling of three-body lifetimes, "
        "while celestial-mechanics stability theory provides geometric insight "
        "into drug disposition. The result is robust to 3D geometry and "
        "gravitational-wave dissipation. We believe this interdisciplinary "
        "bridge meets PRX\u2019s criteria for significant, original research with "
        "broad impact across physics and allied fields."
    )

    doc.add_paragraph()

    # Key findings
    p = doc.add_paragraph()
    run = p.add_run("Key findings:")
    run.font.size = Pt(11)

    bullets = [
        ("A linear 3-compartment PK model reproduces the multi-exponential "
         "lifetime distribution of three-body scattering (15,000 simulations, "
         "3 mass configurations, KS test p < 0.001 vs. single exponential)."),

        ("A nonlinear (Michaelis\u2013Menten) extension captures the power-law "
         "tail from sticky chaos (36\u201342% tail RMSE improvement)."),

        ("Population PK (mixed-effects) modeling yields closed-form allometric "
         "scaling of mean residence time across 64 configurations "
         "(MRT \u221d \u03bc_{12}^{2.08} \u00d7 \u03bc_{out}^{\u22121.68}, "
         "R\u00b2 = 0.67)."),

        ("The reverse mapping provides a stability analysis framework for "
         "target-mediated drug disposition grounded in celestial mechanics."),

        ("Extension to 3D + 2.5PN gravitational-wave radiation reaction "
         "confirms robustness (KS = 0.16\u20130.31), with implications for "
         "LISA merger-rate predictions."),
    ]

    for text in bullets:
        p = doc.add_paragraph(style='List Bullet')
        add_text_with_refs(p, text)

    doc.add_paragraph()

    # Context
    context_p = doc.add_paragraph()
    run = context_p.add_run(
        "We note that PRX has published closely related work on the "
        "statistical three-body problem: Ginat and Perets, Phys. Rev. X 11, "
        "031020 (2021). The present study builds on that foundation by "
        "revealing that the same stochastic structure can be mapped onto "
        "pharmacokinetic compartment theory, opening a bidirectional bridge "
        "between astrophysics and clinical medicine."
    )
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Subject areas
    subj_areas_p = doc.add_paragraph()
    run = subj_areas_p.add_run("Suggested subject areas: ")
    run.bold = True
    run.font.size = Pt(11)
    run = subj_areas_p.add_run(
        "(1) Astrophysics; "
        "(2) Computational astrophysics; "
        "(3) Interdisciplinary physics"
    )
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Suggested reviewers
    p = doc.add_paragraph()
    run = p.add_run("Suggested reviewers:")
    run.font.size = Pt(11)

    reviewers = [
        ("Dr. Alessandro A. Trani",
         "Niels Bohr Institute, University of Copenhagen, Denmark",
         "alessandro.trani@nbi.ku.dk",
         "Three-body problem, gravitational waves, N-body dynamics"),
        ("Dr. Silvia Toonen",
         "Anton Pannekoek Institute, University of Amsterdam, Netherlands",
         "toonen@uva.nl",
         "Gravitational-wave sources, compact binary dynamics"),
        ("Prof. Mats O. Karlsson",
         "Dept. of Pharmacy, Uppsala University, Sweden",
         "mats.karlsson@farmaci.uu.se",
         "Population pharmacokinetics, nonlinear mixed-effects modeling"),
        ("Prof. Nick Holford",
         "Dept. of Pharmacology, University of Auckland, New Zealand",
         "n.holford@auckland.ac.nz",
         "Clinical pharmacology, PK/PD modeling"),
        ("Prof. Rosalba Perna",
         "Dept. of Physics & Astronomy, Stony Brook University, USA",
         "rosalba.perna@stonybrook.edu",
         "N-body dynamics, gravitational waves, compact objects"),
    ]

    for name, affil, email, expertise in reviewers:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}, ")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(f"{affil}. ")
        run.font.size = Pt(10)
        run.italic = True
        run = p.add_run(f"{email}. ")
        run.font.size = Pt(10)
        run = p.add_run(f"Expertise: {expertise}.")
        run.font.size = Pt(10)

    doc.add_paragraph()

    # Excluded reviewers
    excl_p = doc.add_paragraph()
    run = excl_p.add_run("Excluded reviewers: ")
    run.bold = True
    run.font.size = Pt(11)
    run = excl_p.add_run(
        "Nathan W.C. Leigh (Universidad de Concepci\u00f3n), "
        "Nicholas C. Stone (Hebrew University), and "
        "Barak Kol (Hebrew University)\u2014all are authors of "
        "directly competing statistical three-body solutions."
    )
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Closing
    closing = [
        "Sincerely,",
        "",
        "Tatsuki Onishi",
        "Faculty of Data Science, Shiga University",
        "1-1-1 Banba, Hikone, Shiga 522-8522, Japan",
        "Email: bougtoir@gmail.com",
    ]

    for line in closing:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.size = Pt(11)

    out_path = os.path.join(OUTDIR, "cover_letter_prx.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


# ---------------------------------------------------------------------------
# 3. Popular Summary (text file, ~250 words, required by PRX)
# ---------------------------------------------------------------------------

def build_popular_summary():
    out_path = os.path.join(OUTDIR, "popular_summary.txt")
    with open(out_path, 'w') as f:
        f.write("POPULAR SUMMARY\n")
        f.write("=" * 60 + "\n\n")
        f.write(POPULAR_SUMMARY + "\n")
    print(f"Saved: {out_path}")

    # Word count check
    word_count = len(POPULAR_SUMMARY.split())
    print(f"  Popular Summary word count: {word_count} (target: ~250)")
    return out_path


# ---------------------------------------------------------------------------
# 4. Create zip package
# ---------------------------------------------------------------------------

def create_zip(manuscript_path, cover_letter_path, popular_summary_path):
    zip_path = os.path.join(OUTDIR, "threebody_pk_prx_submission.zip")

    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        # Manuscript
        zf.write(manuscript_path, os.path.basename(manuscript_path))
        # Cover letter
        zf.write(cover_letter_path, os.path.basename(cover_letter_path))
        # Popular summary
        zf.write(popular_summary_path, os.path.basename(popular_summary_path))

        # Editable PPTX (if exists)
        pptx_path = os.path.join(OUTDIR, "figures_editable.pptx")
        if os.path.exists(pptx_path):
            zf.write(pptx_path, "figures_editable.pptx")

        # Individual figure files (PNG + PDF)
        for fname, _ in FIGURE_META:
            png_path = os.path.join(FIGDIR, fname)
            if os.path.exists(png_path):
                zf.write(png_path, f"figures/{fname}")
            pdf_fname = fname.replace('.png', '.pdf')
            pdf_path = os.path.join(FIGDIR, pdf_fname)
            if os.path.exists(pdf_path):
                zf.write(pdf_path, f"figures/{pdf_fname}")

    print(f"Saved: {zip_path}")
    return zip_path


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    print("=" * 60)
    print("  Generating Physical Review X (PRX) Submission Package")
    print("=" * 60)

    ms_path = build_manuscript()
    cl_path = build_cover_letter()
    ps_path = build_popular_summary()
    zip_path = create_zip(ms_path, cl_path, ps_path)

    print(f"\nOutputs:")
    print(f"  Manuscript: {ms_path}")
    print(f"  Cover letter: {cl_path}")
    print(f"  Popular summary: {ps_path}")
    print(f"  Submission zip: {zip_path}")
    print("=" * 60)

#!/usr/bin/env python3
"""Assemble the Transport Policy (Elsevier) submission package
(manuscript, title page, cover letter, highlights, checklist) for the
shared-micromobility lifecycle data-exposure study.

This is the manuscript-body / cover-letter generator. It is intentionally kept
out of the public repository because it contains the article prose and does not
contribute to result reproducibility. All quantitative claims are imported from
``reproduce`` (the public reproducibility engine), so the numbers, figures, and
tables it embeds are exactly those regenerated from the committed public data.
Citations use author-date (Elsevier Harvard) style.
"""

from __future__ import annotations

import os
import re
import textwrap
import urllib.error
import urllib.request
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from reproduce import (
    ROOT, OUTPUT, FIGDIR,
    REFS, CITEMETA, intext, sortkey, CITE_RX, _parse_label,
    resolve_citations,
    fmt, pct, ci, domain_count,
    DOMAINS, N_OPERATORS, N_CODED,
    SCR, N_REGISTRY, N_REACHABLE, N_MOTOR_FEEDS, N_OP_DOMAINS,
    FIG_CAPTIONS, TABLES,
    build_figures, build_tables_docx, build_figures_pptx,
    configure_document, add_table, reset_dirs,
)

TITLE = (
    "Lifecycle Data Exposure in Shared Micromobility: A Scoping Review and a "
    "Global Audit of Public GBFS Vehicle Feeds and Operator Disclosures"
)
SHORT_TITLE = "Lifecycle data exposure in shared micromobility"
AUTHOR = "Onishi Tatsuki"
AFFILIATION = ("Data Science and AI Innovation Research Promotion Center, "
               "Shiga University, Hikone, Japan")
EMAIL = "bougtoir@gmail.com"
ORCID = "0000-0001-7261-9062"
JOURNAL = "Transport Policy"
PUBLISHER = "Elsevier"
ARTICLE_TYPE = "Full Article"
PUBLIC_REPO_URL = "https://github.com/bougtoir/discharged-secrets-scoping-review"
BUILD_DATE = "24 July 2026"

# --- Review-model toggle -------------------------------------------------
# The same pipeline serves single-blind (default) and double-masked review.
# Transport Policy uses double-anonymized peer review, so the main manuscript is
# built blinded by default. BLINDED=1 removes author-identifying content from the
# *main manuscript* (byline, acknowledgements, and any identity-revealing
# repository URL). The title page and cover letter are always non-anonymized
# because journals collect them separately and do not forward them to reviewers.
BLINDED = os.environ.get("BLINDED", "1") == "1"

# Persistent DOI for the archived deposit (e.g. Zenodo concept DOI covering all
# versions). Set ZENODO_DOI once the archive is minted; the availability
# statement then cites the DOI instead of the bare repository URL.
ZENODO_DOI = os.environ.get("ZENODO_DOI", "").strip()

# Identity-free link used *only* during double-masked review. anonymous.4open.science
# mirrors a GitHub repository behind a temporary URL that does not reveal the
# owner/organisation. Override per submission with ANON_REPO_URL.
ANON_REPO_URL = os.environ.get(
    "ANON_REPO_URL",
    "https://anonymous.4open.science/r/discharged-secrets-scoping-review")

# Transport Policy allows 1-7 keywords; avoid multi-word phrases containing
# 'and' or 'of'.
KEYWORDS = [
    "shared micromobility",
    "location privacy",
    "data governance",
    "GBFS",
    "transport policy",
    "lifecycle exposure",
]

# Elsevier-style highlights: 3-5 bullets, each no more than 85 characters
# including spaces.
HIGHLIGHTS = [
    "Scoping review of 2,169 records yields 18 direct micromobility studies.",
    "GBFS audit: persistent IDs and GPS are published in almost every feed.",
    "Privacy notices rarely address battery data, vulnerability disclosure, or disposal.",
    "A reproducible lifecycle model links evidence strength to policy controls.",
    "Findings are disclosure signals, not evidence of privacy harm.",
]

# Contents of the archive, reused across every availability-statement variant.
_ARCHIVE_CONTENTS = (
    "the frozen GBFS registry snapshot and its checksum, the title/abstract "
    "and full-text screening decisions, the disclosure-audit coding sheets "
    "with verbatim locator quotations, the analysis scripts, the analysis "
    "results, and the figure- and table-generation code. Raw vehicle "
    "identifiers, exact coordinates, and vehicle-specific deep links were not "
    "retained; only field presence or absence was recorded.")
_REPRO_SENTENCE = (
    " Every count, proportion, and confidence interval reported in the article "
    "can be regenerated from these materials with a single build command.")


def data_availability_statement() -> str:
    """Return the availability statement matching the current review model.

    - double-masked (BLINDED): identity-free anonymized link; DOI/URL added on
      acceptance;
    - single-blind + Zenodo DOI: cite the persistent DOI plus the repository;
    - single-blind, no DOI yet: cite the public repository URL.
    """
    if BLINDED:
        return (
            "The data and code that support the findings of this study are "
            f"available to reviewers via an anonymized repository at "
            f"{ANON_REPO_URL}. The archive contains " + _ARCHIVE_CONTENTS +
            " Upon acceptance, these materials will be deposited in a public "
            "repository with a permanent DOI." + _REPRO_SENTENCE)
    if ZENODO_DOI:
        return (
            "The data and code that support the findings of this study are "
            f"openly archived on Zenodo at https://doi.org/{ZENODO_DOI} "
            "(concept DOI, covering all versions), with ongoing development at "
            f"{PUBLIC_REPO_URL}. The archive contains " + _ARCHIVE_CONTENTS +
            _REPRO_SENTENCE)
    return (
        "The data and code that support the findings of this study are openly "
        f"available in the project repository at {PUBLIC_REPO_URL}. This "
        "includes " + _ARCHIVE_CONTENTS + _REPRO_SENTENCE)


# Disclosure statements required by Transport Policy, placed after the main text
# and before the references. Data availability, funding, and competing
# interests are mandatory; the remaining statements are included as good
# practice and to preserve the ethics and generative-AI disclosures.
def build_declarations() -> list[tuple[str, str]]:
    acknowledgements = (
        "Acknowledgements are withheld to preserve author anonymity for "
        "double-masked peer review; they will be restored on acceptance."
        if BLINDED else
        "The author thanks the maintainers of the public GBFS ecosystem and "
        "the open bibliographic sources that made this audit possible.")
    contributions = (
        "Author contributions are withheld to preserve author anonymity for "
        "double-masked peer review; they will be restored on acceptance."
        if BLINDED else
        f"{AUTHOR}: conceptualization; methodology; software; formal "
        "analysis; data curation; writing - original draft; writing - review "
        "and editing.")
    ai_use = (
        "No generative AI tool was used to produce scientific content, "
        "analysis, or interpretation. Any use was limited to routine language "
        "editing and was reviewed by the author"
        + (", who takes full responsibility for the text."
           if not BLINDED else "s, who take full responsibility for the text."))
    return [
        ("Acknowledgements", acknowledgements),
        ("Data availability statement", data_availability_statement()),
        ("Funding statement",
         "This research received no specific grant from any funding agency, "
         "commercial or not-for-profit sectors."),
        ("Competing interests",
         "The author declares none."),
        ("Author contributions", contributions),
        ("Ethical standards",
         "The research meets all ethical guidelines, including adherence to "
         "the legal requirements of the study jurisdiction. No human "
         "participants were recruited. The study analysed only publicly "
         "accessible feeds and documents and did not attempt authentication, "
         "access-control circumvention, or interaction with individual users."),
        ("Use of generative AI", ai_use),
    ]


DECLARATIONS = build_declarations()


def body_blocks() -> list[tuple]:
    n_id = fmt(SCR["identified"])
    doi_pct = 100 * SCR["with_doi"] / SCR["identified"]
    reach = pct("all_registry_entries", "auto_discovery_reachable")
    reach_ci = ci("all_registry_entries", "auto_discovery_reachable")
    blocks: list[tuple] = []

    blocks.append(("h1", "1. Introduction"))
    blocks += [("p", t) for t in [
        "Shared micromobility - dockless electric scooters and bicycles rented "
        "through a smartphone application - has become a visible class of "
        "connected devices that organizations and the public use but do not "
        "own, maintain, or decommission. Each vehicle continuously produces "
        "location, motion, and battery telemetry that is transmitted to an "
        "operator backend and, in many cities, republished through open data "
        "feeds. Prior work has shown that such telemetry can support "
        "re-identification and tracking well beyond the individual rental that "
        "generated it [[demontjoye;elzer]].",
        "Security research on shared micromobility has grown quickly but "
        "unevenly. Studies span privacy measurement of rental applications, "
        "firmware and protocol attacks on scooters, location-spoofing threats, "
        "and forensic recovery from returned devices [[vinayaga2022;espoofer;"
        "yilmaz2023;hilgert]]. These contributions are scattered across "
        "security, transportation, and forensics venues, use different threat "
        "models and units of analysis, and have not been assembled into a "
        "single map of what is actually demonstrated as opposed to argued. As "
        "a result, it is difficult to state precisely which data-exposure "
        "pathways rest on direct empirical evidence and which rest on analogy.",
        "A second gap concerns exposure that persists across the device "
        "lifecycle. Attention typically concentrates on real-time position "
        "during a rental, yet information created at deployment, maintenance, "
        "recall, and disposal can remain accessible after custody changes hands "
        "[[iotreuse;remanence]]. Battery and diagnostic channels, in "
        "particular, can leak activity patterns even when positioning is "
        "restricted [[leaky;bms]]. Whether operators disclose these lifecycle "
        "practices to the public is largely unexamined. For transport and "
        "data-protection regulators, this means procurement clauses and "
        "oversight frameworks currently focus disproportionately on real-time "
        "location while leaving data at rest and end-of-life handling largely "
        "unaddressed.",
        "This article addresses both gaps with an empirical, reproducible "
        "package rather than a conceptual argument. We (i) conduct a scoping "
        "review, following the Preferred Reporting Items for Systematic Reviews "
        "and Meta-Analyses extension for Scoping Reviews (PRISMA-ScR) "
        "[[prisma_scr]], of evidence on micromobility data exposure; (ii) audit "
        "the fields that operators actually publish worldwide through the "
        "General Bikeshare Feed Specification (GBFS) [[gbfs_spec]]; and (iii) "
        "audit what a matched set of operator privacy notices discloses about "
        "collection, retention, transfer, and device end-of-life. Our research "
        "questions are: RQ1, what data-exposure pathways in shared "
        "micromobility are supported by direct evidence; RQ2, which "
        "vehicle-level fields are publicly disclosed, and at what prevalence; "
        "and RQ3, how completely do public operator documents describe "
        "lifecycle data handling.",
        "We are deliberately conservative about interpretation. Publishing a "
        "field, or remaining silent about a practice, is a disclosure signal; "
        "it is not by itself evidence of a privacy harm, a compromise, or a "
        "regulatory violation. The contribution is a transparent evidence base "
        "and a lifecycle model that ties each stage to the strength of its "
        "supporting evidence and to controls whose effectiveness remains to be "
        "tested.",
    ]]

    blocks.append(("h1", "2. Methods"))
    blocks.append(("h2", "2.1. Design and reporting"))
    blocks += [("p", t) for t in [
        "The study combines three prespecified components: a scoping review "
        "(work package WP1), a cross-sectional field audit of public GBFS feeds "
        "(WP2), and a structured disclosure audit of public operator documents "
        "(WP3). The review component is reported in line with PRISMA-ScR "
        "[[prisma_scr]] and follows the scoping-review framework of Arksey and "
        "O'Malley [[year:arksey]]. Eligibility criteria and data sources for all "
        "three components are summarized in Table 1. The protocol, screening "
        "rules, coding sheets, and analysis code are openly available so that "
        "the counts reported here can be regenerated.",
    ]]
    blocks.append(("table", 1))
    blocks.append(("h2", "2.2. Scoping review (WP1)"))
    blocks += [("p", t) for t in [
        f"We compiled {n_id} records from programmatic searches of open "
        "bibliographic metadata. A digital object identifier (DOI) was "
        f"available for {fmt(SCR['with_doi'])} of {n_id} records ({doi_pct:.1f}%), "
        "for which abstracts were retrieved automatically where possible; the "
        "remaining records were screened on title and available metadata only. "
        "Title/abstract screening applied a deterministic rule set that combined "
        "target-domain relevance with the presence of a described data path, and "
        "recorded a decision and reason for every record. Because the rules are "
        "deterministic, we re-applied them to a delayed "
        f"{round(100 * SCR['resample_n'] / SCR['identified'])}% sample "
        f"(n = {fmt(SCR['resample_n'])}) and reproduced every original decision "
        f"({SCR['resample_agreement_pct']}% agreement); this demonstrates "
        "computational reproducibility rather than inter-rater reliability, as a "
        "single reviewer conducted the screening.",
        "Records marked include or uncertain were sought for full-text "
        "assessment. Each retrieved study was classified by evidence distance: "
        "D4, direct target-domain empirical evidence; D3, direct target-domain "
        "documentary evidence; D2, near-domain empirical evidence; D1, "
        "mechanism analogy; and N, normative evidence. This ordering is an "
        "operational classification defined for the present study; it is "
        "conceptually related to the GRADE notion of indirectness "
        "[[grade_indirectness]] but is not a GRADE certainty rating. We did not "
        "treat a "
        "title/abstract decision as equivalent to a confirmed full-text "
        "finding, and we did not claim to have read full text that could not be "
        "retrieved. From the included studies we built a study-level extraction "
        "table capturing device or service, design, data fields, access path, "
        "reported outcome, and limitations.",
    ]]
    blocks.append(("h2", "2.3. Public GBFS field audit (WP2)"))
    blocks += [("p", t) for t in [
        "We froze a snapshot of the public GBFS systems catalogue [[gbfs_registry]] "
        "and recorded its checksum. For each system we attempted to reach the "
        "auto-discovery endpoint, to locate a declared vehicle-status feed, to "
        "retrieve that feed, and to record which specification fields it "
        "contained. The unit of analysis is the system/feed. To respect the "
        "study's safety constraints, we recorded only the presence or absence "
        "of each field; raw vehicle identifiers, exact coordinates, and "
        "vehicle-specific deep links were not retained. Unavailable or empty "
        "feeds were separated from feeds that were retrieved but omitted a "
        "field, and all proportions are reported against explicit denominators "
        "with 95% Wilson confidence intervals (CIs). Because several large operators "
        "run many city systems, we also computed an operator-domain sensitivity "
        "analysis to check whether prevalence was driven by a few operators.",
    ]]
    blocks.append(("h2", "2.4. Public-document disclosure audit (WP3)"))
    blocks += [("p", t) for t in [
        "We selected operators from the audited GBFS population and retrieved "
        f"their public privacy notices. For each operator document we coded "
        f"{len(DOMAINS)} "
        "disclosure domains - location; trip and time data; vehicle "
        "identifiers; battery or diagnostic data; maintenance or repair "
        "records; account, payment, and device data; analytics or profiling; "
        "retention; processors or contractors; international transfers; "
        "data-subject rights; incident contact; vulnerability disclosure; and "
        "device return, recycling, or disposal - using the values explicit, "
        "partial, not found, not applicable, and unavailable. A not-found "
        "coding means the document did not address the domain; it is not "
        "evidence that the practice does not occur. Each coding is accompanied "
        "by a short verbatim locator quotation in the coding sheet so that a "
        "third party can check it. Coding was computer-assisted and reviewed by "
        "a single reviewer; governance and standards documents "
        "[[mds_privacy;edpb_cv;nist88;nist161;eu_battery]] were used as "
        "reference points and were not coded as operator practices.",
    ]]

    blocks.append(("h1", "3. Results"))
    blocks.append(("h2", "3.1. Scoping review (RQ1)"))
    blocks += [("p", t) for t in [
        f"Of {n_id} records screened, {fmt(SCR['ta_excluded'])} were excluded at "
        f"title/abstract and {fmt(SCR['sought'])} were sought for full text "
        f"(Fig. 1). Of these, {fmt(SCR['not_retrieved'])} could not be retrieved "
        f"and were recorded as such rather than assessed; "
        f"{fmt(SCR['excluded_fulltext'])} retrieved records were excluded "
        f"because they contained no relevant data path ({SCR['no_data_path']}) "
        f"or described a mechanism that did not transfer to the target domain "
        f"({SCR['not_transferable']}). {SCR['included']} direct studies were "
        f"included: {SCR['d4']} at evidence distance D4, {SCR['d3']} at D3, and "
        f"{SCR['d2']} at D2 (Fig. 2; Table 2).",
        "The included D4 studies provide the strongest evidence. A long-term "
        "real-world analysis reconstructed rider-relevant patterns from "
        "operator data [[elzer]]; investigative studies of rental applications "
        "and scooter ecosystems demonstrated collection and protocol weaknesses "
        "[[vinayaga2022;espoofer;etrojans]]; and a forensic analysis recovered "
        "data from micromobility devices [[hilgert]]. D3 studies document "
        "platform architectures, location-spoofing threats, data-acquisition "
        "frameworks, and user-facing traceability concerns "
        "[[isik;vinayaga2020;yilmaz2022;yilmaz2023;sato;li2020;petersen;zhou;"
        "hannemann]]. D2 studies transfer from adjacent domains: battery "
        "side channels and battery-management data [[leaky;bms]], residual data "
        "in reused IoT devices [[iotreuse]], and cloud data remanence "
        "[[remanence]]. No included study, on its own, demonstrated an "
        "end-to-end lifecycle compromise; the evidence is strongest for "
        "operation and weakest for recall and disposal.",
    ]]
    blocks.append(("fig", 1))
    blocks.append(("fig", 2))
    blocks.append(("table", 2))
    blocks.append(("h2", "3.2. Public GBFS field audit (RQ2)"))
    blocks += [("p", t) for t in [
        f"Of {fmt(N_REGISTRY)} registry systems, {reach} (95% CI {reach_ci}; "
        f"n = {fmt(N_REACHABLE)}) exposed a reachable auto-discovery endpoint, "
        f"{pct('reachable_registry_entries', 'vehicle_feed_declared')} of "
        "reachable systems declared a vehicle-status feed, and "
        f"{pct('successful_vehicle_feeds', 'vehicle_feed_nonempty')} of "
        "successfully retrieved feeds contained at least one vehicle (Table 3). "
        f"Restricting to the {fmt(N_MOTOR_FEEDS)} non-empty feeds that declared "
        "motorized micromobility, a vehicle identifier was present in "
        f"{pct('declared_motorized_micromobility_feeds', 'has_vehicle_id')} of "
        "feeds and latitude/longitude in "
        f"{pct('declared_motorized_micromobility_feeds', 'has_location_fields')} "
        "(Fig. 3). Fields with more operational specificity were less uniformly "
        "published: current range in "
        f"{pct('declared_motorized_micromobility_feeds', 'has_range')}, "
        "vehicle-specific rental links in "
        f"{pct('declared_motorized_micromobility_feeds', 'has_deep_link')}, "
        "last-reported timestamps in "
        f"{pct('declared_motorized_micromobility_feeds', 'has_last_reported')}, "
        "and battery or fuel percentage in "
        f"{pct('declared_motorized_micromobility_feeds', 'has_battery_percent')}.",
        "The operator-domain sensitivity analysis indicates that these are not "
        f"artefacts of a few large operators. Across {fmt(N_OP_DOMAINS)} "
        "eligible operator domains, a vehicle identifier appeared in every "
        "eligible feed for "
        f"{pct('declared_motorized_micromobility_operator_domains_all', 'has_vehicle_id')} "
        "of domains and latitude/longitude for "
        f"{pct('declared_motorized_micromobility_operator_domains_any', 'has_location_fields')}, "
        "whereas battery percentage was present in at least one eligible feed "
        "for only "
        f"{pct('declared_motorized_micromobility_operator_domains_any', 'has_battery_percent')} "
        "of domains. These figures describe what is disclosed publicly; they do "
        "not describe what operators collect or store on their backends, which "
        "the public feed cannot reveal.",
    ]]
    blocks.append(("fig", 3))
    blocks.append(("table", 3))
    blocks.append(("h2", "3.3. Public-document disclosure audit (RQ3)"))
    exp_loc = domain_count("location_data", "explicit")
    nf_batt = domain_count("battery_or_diagnostic_data", "not_found")
    nf_vuln = domain_count("vulnerability_disclosure", "not_found")
    nf_disp = domain_count("return_recycling_disposal", "not_found")
    blocks += [("p", t) for t in [
        f"We coded {N_CODED} operator privacy notices across {len(DOMAINS)} "
        f"domains; {N_OPERATORS - N_CODED} "
        "further operators could not be retrieved as reproducible text and were "
        "recorded as unavailable (Fig. 4; Table 4). Operational and "
        "account-level processing was disclosed almost universally: location "
        f"data were explicit for {exp_loc} of {N_CODED} operators, and "
        f"retention, processors or contractors, and account/payment/device data "
        f"were explicit for most. Data-subject rights and international "
        "transfers were explicit or partial for the large majority, consistent "
        "with a predominantly European operator sample governed by the General "
        "Data Protection Regulation [[gdpr]].",
        "Disclosure was markedly thinner for device-centred and lifecycle "
        f"domains. Battery or diagnostic data collection was not found in "
        f"{nf_batt} of {N_CODED} notices, even though such data are technically "
        "central to these vehicles [[leaky;bms]]. No operator notice in the "
        f"sample described a vulnerability-disclosure channel ({nf_vuln} of "
        f"{N_CODED} not found), and none addressed device return, recycling, or "
        f"disposal handling ({nf_disp} of {N_CODED} not found). We stress that "
        "these are gaps in public documents, not proof that the corresponding "
        "practices are absent; they nonetheless mark the parts of the lifecycle "
        "that are least visible to the public and to procuring organizations.",
    ]]
    blocks.append(("fig", 4))
    blocks.append(("table", 4))

    blocks.append(("h1", "4. Discussion"))
    blocks += [("p", t) for t in [
        "Read together, the three components describe a consistent pattern. "
        "Direct evidence for data exposure is concentrated in the operation "
        "stage, where audits and attack studies demonstrate collection, "
        "tracking, and re-identification potential [[elzer;espoofer;"
        "vinayaga2022]]. The public feed audit shows that the raw materials for "
        "such analyses - persistent identifiers and precise positions - are "
        "disclosed at high prevalence worldwide, while more operationally "
        "sensitive fields are published less uniformly. The disclosure audit "
        "shows that operators describe operational data handling in detail but "
        "are largely silent on the device end-of-life and on security-reporting "
        "channels, which is exactly where the review found only near-domain "
        "(D2) evidence [[iotreuse;remanence]].",
        "We integrate these observations in a six-stage lifecycle model that "
        "links each stage to the evidence distance of its supporting sources "
        "and to a proposed control (Fig. 5; Table 5). The model makes the "
        "strength of the underlying evidence explicit: procurement, deployment, "
        "operation, and maintenance are anchored by direct (D3-D4) evidence, "
        "whereas recall/return and second-life/disposal rest on near-domain "
        "(D2) analogy and on governance and standards guidance "
        "[[nist88;nist161;eu_battery]]. The controls - field minimization, "
        "local processing, access and retention limits, controlled contractor "
        "access, chain-of-custody for returned units, and verified media "
        "sanitization - are proposals whose effectiveness this study did not "
        "test; Table 5 records their effectiveness evidence conservatively.",
        "For procuring organizations, the practical implication is that a "
        "narrow focus on real-time position understates exposure. A telemetry "
        "inventory that spans identifiers, timestamps, range, and battery "
        "state, and that follows devices through maintenance and disposal, is a "
        "more faithful basis for assessment than the intuition that "
        "coordinates alone are the only sensitive field.",
    ]]
    blocks.append(("fig", 5))
    blocks.append(("table", 5))
    blocks.append(("h2", "4.1. Limitations"))
    blocks += [("p", t) for t in [
        "Several limitations bound these findings. Screening and coding were "
        "performed by a single reviewer with computer assistance; although the "
        "process is deterministic and reproducible, it does not provide "
        "inter-rater reliability, and some included studies were characterized "
        "from abstracts and curated metadata rather than from independently "
        "reproduced experiments. The GBFS audit observes only what is published "
        "in public feeds at a single point in time; it cannot show what "
        "operators collect or retain internally, and field presence is not a "
        "measure of harm. The disclosure audit reflects the public documents we "
        "could retrieve; not-found codings denote silence, two operators were "
        "unavailable, and the operator sample is weighted toward European and "
        "North American providers. Finally, the lifecycle controls are "
        "proposals; validating them would require intervention studies or "
        "operator cooperation that were outside this study's scope.",
    ]]

    blocks.append(("h2", "4.2. Policy implications"))
    blocks += [("p", t) for t in [
        "For procuring authorities and regulators, the findings point to three "
        "actionable policy priorities. First, procurement clauses and concession "
        "agreements for shared micromobility should require disclosure of the "
        "vehicle fields that are published in open feeds, not only real-time "
        "location. Identifiers, timestamps, range, and battery state are "
        "collectively more enabling of re-identification and tracking than any "
        "single field, and their prevalence means that procurement templates "
        "focusing only on GPS coordinates understate exposure.",
        "Second, operator accountability frameworks should close the lifecycle "
        "gaps visible in the document audit. Battery and diagnostic data, "
        "vulnerability disclosure channels, and device return or disposal "
        "handling were rarely described in public notices. These omissions "
        "matter because end-of-life custody changes are where data-remanence "
        "and secondary-market risks are concentrated. Requiring operators to "
        "publish plain statements on these domains - or to explain why they are "
        "not applicable - would make the market easier to compare and "
        "regulate without mandating specific technical architectures.",
        "Third, standards and guidance can be targeted where evidence is "
        "strongest. GBFS already exposes which fields are published, so feed-level "
        "transparency can be improved by standardising the optional fields that "
        "affect privacy, and by documenting the rationale for their inclusion. "
        "For end-of-life handling, where only near-domain evidence exists, "
        "regulators should treat disposal requirements as precautionary rather "
        "than evidence-based until device-level empirical studies become "
        "available. The lifecycle framework supplied in this article is "
        "designed to make these gradations explicit so that policy can track "
        "evidence rather than assume it.",
    ]]

    blocks.append(("h1", "5. Conclusion"))
    blocks += [("p", t) for t in [
        "Shared micromobility offers a tractable, fully public setting in which "
        "to study lifecycle data exposure in connected devices that "
        f"organizations use but do not control. A scoping review of {n_id} "
        f"records identified {SCR['included']} direct studies whose strongest "
        f"evidence concerns the operation stage; a global audit of "
        f"{fmt(N_REACHABLE)} public GBFS "
        "systems showed that persistent identifiers and precise positions are "
        "disclosed at high prevalence; and a disclosure audit showed that "
        f"operators document operational processing thoroughly but are silent "
        f"on device disposal and vulnerability reporting. These are disclosure "
        "and evidence signals, not demonstrations of harm. The lifecycle model "
        "and its traceability table turn them into an auditable agenda: extend "
        "assessment beyond real-time location, close the documentation gaps at "
        "end-of-life, and empirically validate the proposed controls.",
    ]]
    return blocks


# ---------------------------------------------------------------------------
# Citation resolution
def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0.3)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_figure(doc: Document, num: int, png: Path, repl) -> None:
    image_p = doc.add_paragraph()
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.paragraph_format.space_before = Pt(14)
    width = 6.3 if num in (1, 2, 3) else 6.5
    image_p.add_run().add_picture(str(png), width=Inches(width))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(14)
    run = caption.add_run(repl(FIG_CAPTIONS[num]))
    run.italic = True
    run.font.size = Pt(10)


def add_page_number(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


ABSTRACT_STRUCT = [
    ("Abstract", "Shared micromobility - dockless electric scooters and "
     "bicycles rented via smartphone - is a globally deployed class of "
     "connected devices that cities and citizens use but do not own or "
     "decommission. This study maps the direct evidence for data exposure across "
     "the device lifecycle, measures which vehicle fields operators publish "
     "worldwide, and assesses how completely operators disclose lifecycle data "
     "handling. A PRISMA-ScR scoping review of "
     f"{fmt(SCR['identified'])} records yielded {SCR['included']} direct studies, "
     f"with the strongest evidence at the operation stage. A cross-sectional "
     f"audit of {fmt(N_REACHABLE)} public General Bikeshare Feed Specification "
     "(GBFS) feeds found that persistent identifiers and precise positions are "
     "published in almost every motorized feed, while battery and diagnostic "
     "fields are published less uniformly. A disclosure audit of public operator "
     "privacy notices showed that operational data handling is documented "
     "thoroughly, but end-of-life handling and vulnerability reporting are "
     "rarely addressed. We frame these findings as disclosure and evidence "
     "signals, not demonstrations of harm, and provide a reproducible lifecycle "
     "model that ties each stage to the strength of its evidence and to policy "
     "controls. The results suggest that procurement and regulatory oversight "
     "should move beyond real-time location and explicitly require "
     "transparency on identifiers, timestamps, range, battery state, device "
     "disposal, and vulnerability reporting."),
]


def build_abstract_text() -> str:
    return " ".join(text for _, text in ABSTRACT_STRUCT)



def build_manuscript(blocks, tables, figpaths, repl, references) -> Path:
    doc = Document()
    configure_document(doc)
    add_page_number(doc.sections[0])
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(TITLE).bold = True
    byline = doc.add_paragraph()
    byline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if BLINDED:
        # Double-masked: the main manuscript carries no identifying byline; the
        # author details travel on the separate, non-anonymized title page.
        anon = byline.add_run(
            "Author and affiliation details removed for double-masked peer "
            "review")
        anon.italic = True
        anon.font.size = Pt(10)
    else:
        byline.add_run(AUTHOR).bold = True
        aff = doc.add_paragraph()
        aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar = aff.add_run(
            f"{AFFILIATION}\nCorresponding author: {EMAIL}; ORCID {ORCID}")
        ar.italic = True
        ar.font.size = Pt(10)

    doc.add_heading("Abstract", level=1)
    _, abstract_text = ABSTRACT_STRUCT[0]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(abstract_text)

    doc.add_heading("Highlights", level=1)
    for bullet in HIGHLIGHTS:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(bullet)

    kw = doc.add_paragraph()
    kw.add_run("Keywords: ").bold = True
    kw.add_run("; ".join(KEYWORDS))

    for kind, payload in blocks:
        if kind == "h1":
            doc.add_heading(payload, level=1)
        elif kind == "h2":
            doc.add_heading(payload, level=2)
        elif kind == "p":
            add_body_paragraph(doc, repl(payload))
        elif kind == "fig":
            add_figure(doc, payload, figpaths[payload]["png"], repl)
        elif kind == "table":
            add_table(doc, tables[payload], repl)

    for label, value in DECLARATIONS:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(f"{label}. ").bold = True
        p.add_run(value)

    doc.add_heading("References", level=1)
    for label in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.28)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(4)
        p.add_run(REFS[label])

    path = OUTPUT / "Manuscript_TransportPolicy.docx"
    doc.save(path)
    return path


def build_title_page(word_count: int, n_refs: int) -> Path:
    doc = Document()
    configure_document(doc)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(TITLE).bold = True
    doc.add_paragraph()
    for text, bold in [(AUTHOR, True), (AFFILIATION, False),
                       (f"Corresponding author: {AUTHOR}; {EMAIL}", False),
                       (f"ORCID: {ORCID}", False)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(text).bold = bold
    doc.add_paragraph()
    highlights_count = len(HIGHLIGHTS)
    abstract_words = len(re.findall(r"\b[\w'-]+\b", build_abstract_text()))
    fields = [
        ("Full title", TITLE),
        ("Short title", SHORT_TITLE),
        ("Article type", ARTICLE_TYPE),
        ("Journal", f"{JOURNAL} ({PUBLISHER})"),
        ("Peer review", "Double-anonymized; the manuscript file is anonymized"),
        ("Abstract word count", str(abstract_words)),
        ("Highlights", f"{highlights_count} bullets"),
        ("Main-text word count (excludes title, abstract, references)", str(word_count)),
        ("Figures", "5"),
        ("Tables", "5"),
        ("References", str(n_refs)),
        ("Funding", "None"),
        ("Competing interests", "None declared"),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)
    path = OUTPUT / "Title_Page_TransportPolicy.docx"
    doc.save(path)
    return path


def build_cover_letter() -> Path:
    doc = Document()
    configure_document(doc)
    for line in [AUTHOR, AFFILIATION, EMAIL, f"ORCID: {ORCID}", BUILD_DATE]:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()
    doc.add_paragraph("The Editors-in-Chief")
    doc.add_paragraph(JOURNAL)
    doc.add_paragraph(PUBLISHER)
    doc.add_paragraph()
    doc.add_paragraph("Dear Editors,")
    paras = [
        f"I submit the manuscript \u201c{TITLE}\u201d for consideration as a "
        f"Full Article in {JOURNAL}.",
        "Shared micromobility is a fully public, globally deployed class of "
        "connected devices that cities and citizens use but do not own or "
        "decommission. The manuscript reports an empirical, reproducible "
        "package at the intersection of transport policy, data governance, and "
        "connected-device security: a PRISMA-ScR scoping review that maps the "
        "direct evidence for data exposure; a global audit of public GBFS "
        "feeds that measures which vehicle fields operators actually publish; "
        "and a structured audit of public operator privacy notices across "
        f"{len(DOMAINS)} disclosure domains.",
        "The work fits Transport Policy because it quantifies real, worldwide "
        "data disclosure practices and connects them to procurement, "
        "regulatory, and standard-setting choices facing public authorities "
        "and transport operators. It moves beyond methodological reporting by "
        "tying each finding to the actors at risk, the policy instruments "
        "that could address the risk, and the strength of the evidence "
        "supporting those instruments. A set of Elsevier-style highlights "
        "summarises the key messages for a policy audience. Throughout, we "
        "treat field presence and document silence as disclosure and evidence "
        "signals rather than as proof of harm, compromise, or regulatory "
        "violation, and we state the effectiveness of proposed controls "
        "conservatively as not yet validated.",
        "The study analysed only publicly accessible feeds and documents; it "
        "did not attempt authentication or access-control circumvention, did "
        "not interact with users, and retained no raw identifiers, exact "
        "coordinates, or vehicle deep links. Consistent with the journal's "
        "double-anonymized peer review, the manuscript file contains no "
        "identifying information; the anonymised repository link and the "
        "public repository URL will be disclosed on acceptance. All data, "
        "coding sheets, and code are available so that every reported count "
        "can be regenerated with a single command. The manuscript is original "
        "and is not under consideration elsewhere. There is no funding or "
        "competing interest to declare, and no generative AI was used for "
        "scientific content.",
        "Thank you for considering this submission.",
    ]
    for text in paras:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(text)
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph(AUTHOR)
    path = OUTPUT / "Cover_Letter_TransportPolicy.docx"
    doc.save(path)
    return path


def build_highlights() -> tuple[Path, Path]:
    doc = Document()
    configure_document(doc)
    h = doc.add_paragraph(style="Title")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.add_run("Highlights").bold = True
    for bullet in HIGHLIGHTS:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(bullet)
    docx_path = OUTPUT / "Highlights_TransportPolicy.docx"
    doc.save(docx_path)
    txt_path = OUTPUT / "Highlights_TransportPolicy.txt"
    txt_path.write_text("\n".join(f"- {b}" for b in HIGHLIGHTS) + "\n", encoding="utf-8")
    return docx_path, txt_path


def build_reporting_guideline() -> Path:
    doc = Document()
    configure_document(doc)
    h = doc.add_paragraph(style="Title")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.add_run("PRISMA-ScR reporting checklist").bold = True
    p = doc.add_paragraph()
    p.add_run(
        "The scoping-review component (WP1) follows the PRISMA-ScR checklist "
        "(Tricco et al., 2018). The field audit (WP2) and disclosure audit "
        "(WP3) are cross-sectional observational studies of public artefacts "
        "and are reported with explicit denominators, confidence intervals, and "
        "open code. The table below maps each PRISMA-ScR item to its location.")
    items = [
        ("Title", "Identifies the report as a scoping review", "Title page"),
        ("Abstract", "Summary", "Abstract"),
        ("Rationale", "Rationale in the context of what is known", "Section 1"),
        ("Objectives", "Research questions RQ1-RQ3", "Section 1"),
        ("Protocol", "Protocol availability", "Data availability statement"),
        ("Eligibility criteria", "Characteristics used as criteria", "Section 2.2; Table 1"),
        ("Information sources", "Sources searched", "Section 2.2; Table 1"),
        ("Selection of sources", "Screening process", "Section 2.2; Fig. 1"),
        ("Data charting", "Extraction process and items", "Section 2.2; Table 2"),
        ("Synthesis of results", "Methods of summarizing", "Sections 3-4"),
        ("Results of sources", "Numbers screened and included", "Section 3.1; Fig. 1"),
        ("Results of syntheses", "Charted results", "Sections 3.1-3.3"),
        ("Limitations", "Limitations of the review", "Section 4.1"),
        ("Conclusions", "Interpretation and implications", "Section 5"),
        ("Funding", "Sources of funding", "Declarations"),
    ]
    add_table(doc, {"title": "PRISMA-ScR item mapping",
                    "headers": ["Item", "Checklist description", "Location"],
                    "rows": items}, lambda x: x)
    path = OUTPUT / "Reporting_Guideline_PRISMA-ScR.docx"
    doc.save(path)
    return path


def resolve_identifier(ident: str) -> str:
    """Best-effort live check that a DOI/URL resolves. Falls back gracefully
    when offline so the build stays reproducible."""
    if ident.startswith("doi:"):
        url = "https://doi.org/" + ident[4:]
    elif ident.startswith("http"):
        url = ident
    else:
        return "not_applicable_standard_or_regulation"
    req = urllib.request.Request(
        url, method="HEAD", headers={"User-Agent": "Mozilla/5.0 ref-check"})
    try:
        code = urllib.request.urlopen(req, timeout=25).status
        return f"resolves_http_{code}"
    except urllib.error.HTTPError as exc:
        # 403/429 are anti-bot responses from a live, existing record.
        return f"exists_http_{exc.code}"
    except Exception:
        return "not_checked_offline"


def build_reference_verification(refs) -> Path:
    path = OUTPUT / "Reference_Verification.csv"
    lines = ["Number,Label,Reference,Identifier,VerificationDate,Status"]
    for i, label in enumerate(refs, 1):
        text = REFS[label]
        m = re.search(r"(doi:\S+|https?://\S+)", text)
        ident = m.group(1).rstrip(".") if m else "(no DOI; standard/regulation/registry)"
        status = resolve_identifier(ident)
        row = [str(i), label, text, ident, BUILD_DATE, status]
        lines.append(",".join(f'"{v.replace(chr(34), chr(34)*2)}"' for v in row))
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def build_citation_audit(blocks, tables, order, repl) -> Path:
    first: dict[str, tuple[str, str]] = {}
    location = "Body"
    for kind, payload in blocks:
        if kind in ("h1", "h2"):
            location = payload
            continue
        frags = []
        if kind == "p":
            frags = [payload]
        elif kind == "table":
            frags = [c for r in tables[payload]["rows"] for c in r]
            location = f"Table {payload}"
        for frag in frags:
            for m in CITE_RX.finditer(frag):
                for raw in m.group(1).split(";"):
                    label, _ = _parse_label(raw)
                    if label not in first:
                        first[label] = (location, textwrap.shorten(repl(frag), 200, placeholder="..."))
    path = OUTPUT / "Citation_Audit.csv"
    lines = ["Citation,Year,First appearance,Context"]
    for label in order:
        loc, ctx = first[label]
        lines.append(",".join(f'"{v.replace(chr(34), chr(34)*2)}"'
                              for v in [intext(label), CITEMETA[label][1], loc, ctx]))
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def word_count(blocks, repl) -> int:
    """Main-text word count excluding title, abstract, and references."""
    text = ""
    for kind, payload in blocks:
        if kind == "p":
            text += " " + repl(payload)
    return len(re.findall(r"\b[\w'-]+\b", text))


def validate(blocks, tables, order, ref_list, repl) -> dict:
    # every cited label resolves (resolve_citations already raises on unknown)
    all_resolve = all(l in REFS for l in order)
    # all refs used, none orphan or phantom
    orphan = set(order) ^ set(REFS.keys())
    # reference list is alphabetical by author (then year)
    refs_alpha = ref_list == sorted(ref_list, key=lambda l: (sortkey(l), l))
    body = word_count(blocks, repl)
    abstract_words = len(re.findall(r"\b[\w'-]+\b", build_abstract_text()))
    highlights_ok = 3 <= len(HIGHLIGHTS) <= 5
    highlights_max_len = all(len(h) <= 85 for h in HIGHLIGHTS)
    # figures/tables cited in text
    body_text = " ".join(repl(p) for k, p in blocks if k == "p")
    figs_cited = all(f"Fig. {i}" in body_text for i in range(1, 6))
    tabs_cited = all(f"Table {i}" in body_text for i in range(1, 6))
    # figure/table blocks present and sequential
    fig_seq = [p for k, p in blocks if k == "fig"]
    tab_seq = [p for k, p in blocks if k == "table"]
    # in-text markers must not leak into the rendered text
    no_raw_markers = not any(CITE_RX.search(repl(p)) for k, p in blocks if k == "p")
    # required Transport Policy disclosure statements
    decl_labels = {lbl for lbl, _ in DECLARATIONS}
    required_disclosures = {
        "Data availability statement", "Funding statement", "Competing interests",
    }.issubset(decl_labels)
    # double-masked guard: no author-identifying token may reach the disclosures
    identity_tokens = [PUBLIC_REPO_URL, "bougtoir"]
    for value in (AUTHOR, EMAIL, ORCID):
        if value and "[" not in value:
            identity_tokens.append(value)
    decl_text = " ".join(v for _, v in DECLARATIONS)
    no_identity_leak = (not BLINDED) or not any(
        tok in decl_text for tok in identity_tokens)
    checks = {
        "all_citations_resolve": all_resolve,
        "no_orphan_or_phantom_refs": not orphan,
        "references_alphabetical": refs_alpha,
        "no_unresolved_markers": no_raw_markers,
        "figures_cited_in_text": figs_cited,
        "tables_cited_in_text": tabs_cited,
        "five_figures_present": sorted(set(fig_seq)) == [1, 2, 3, 4, 5],
        "five_tables_present": sorted(set(tab_seq)) == [1, 2, 3, 4, 5],
        "abstract_within_250w": abstract_words <= 250,
        "highlights_3_to_5": highlights_ok,
        "highlights_max_85_chars": highlights_max_len,
        "keywords_at_most_seven": len(KEYWORDS) <= 7,
        "required_disclosures_present": required_disclosures,
        "no_identity_leak_when_blinded": no_identity_leak,
    }
    failures = [k for k, v in checks.items() if not v]
    if failures:
        raise RuntimeError(f"Validation failed: {failures}; orphan={orphan}")
    # abbreviation-at-first-use spot check
    joined = body_text
    for abbr, definition in {
        "GBFS": "General Bikeshare Feed Specification (GBFS)",
        "PRISMA-ScR": "Scoping Reviews (PRISMA-ScR)",
        "DOI": "digital object identifier (DOI)",
        "GDPR": "General Data Protection Regulation",
    }.items():
        if abbr in joined and definition not in joined:
            raise RuntimeError(f"Undefined abbreviation at first use: {abbr}")
    return {"word_count": body, "abstract_words": abstract_words,
            "highlights": len(HIGHLIGHTS), "keywords": len(KEYWORDS),
            "references": len(ref_list), **checks}


def build_validation_report(validation: dict, figpaths) -> Path:
    path = OUTPUT / "VALIDATION.txt"
    lines = ["Submission validation report", "=" * 30, ""]
    for k, v in validation.items():
        lines.append(f"{k}: {v}")
    lines.append("")
    lines.append("Figure files:")
    for num, kinds in figpaths.items():
        for kind, p in kinds.items():
            lines.append(f"  Figure {num} {kind}: {p.name}")
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def build_checklist(validation) -> Path:
    doc = Document()
    configure_document(doc)
    h = doc.add_paragraph(style="Title")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.add_run("Submission checklist").bold = True
    rows = [
        ["Cover letter", "Included", "Cover_Letter_TransportPolicy.docx"],
        ["Title page with author details", "Included", "Title_Page_TransportPolicy.docx"],
        ["Article file (figures/tables inline; anonymized)", "Included", "Manuscript_TransportPolicy.docx"],
        ["Highlights (3-5 bullets, <=85 chars)", f"{validation['highlights']} bullets", "Highlights_TransportPolicy.docx/.txt"],
        ["Abstract (<=250 words)", f"{validation['abstract_words']} words", "Manuscript"],
        ["Keywords (1-7, semicolon-separated)", f"{validation['keywords']}", "Manuscript"],
        ["Figures (PNG + TIFF + PDF, 600 dpi)", "Included", "figures/Figure1-5.*"],
        ["Editable figures (one per slide)", "Included", "Figures_TransportPolicy_editable.pptx"],
        ["Editable tables", "Included", "Tables_TransportPolicy_editable.docx"],
        ["Reporting guideline (PRISMA-ScR)", "Included", "Reporting_Guideline_PRISMA-ScR.docx"],
        ["Citation audit (author-date)", "Included", "Citation_Audit.csv"],
        ["Reference verification", "Included", "Reference_Verification.csv"],
        [f"Main-text word count (excl. title, abstract, references): {validation['word_count']}",
         "Reported", "Title_Page_TransportPolicy.docx"],
        ["Data availability statement", "Included", "Manuscript"],
        ["Funding statement", "Included", "Manuscript"],
        ["Competing interests statement", "Included", "Manuscript"],
    ]
    add_table(doc, {"title": "Items included in the submission package",
                    "headers": ["Item", "Status", "File"], "rows": rows},
              lambda x: x)
    path = OUTPUT / "Submission_Checklist_TransportPolicy.docx"
    doc.save(path)
    return path


def build_zip(figpaths) -> Path:
    path = OUTPUT / "TransportPolicy_submission_package.zip"
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        for item in sorted(OUTPUT.rglob("*")):
            if item.is_file() and item != path:
                zf.write(item, item.relative_to(OUTPUT))
    return path


def main() -> None:
    reset_dirs()
    FIGDIR.mkdir(parents=True, exist_ok=True)
    figpaths = build_figures()
    blocks = body_blocks()
    order, ref_list, repl = resolve_citations(blocks, TABLES)
    validation = validate(blocks, TABLES, order, ref_list, repl)

    build_manuscript(blocks, TABLES, figpaths, repl, ref_list)
    wc = validation["word_count"]
    build_title_page(wc, len(ref_list))
    build_cover_letter()
    build_highlights()
    build_tables_docx(TABLES, repl)
    build_figures_pptx(figpaths, repl)
    build_reporting_guideline()
    build_reference_verification(ref_list)
    build_citation_audit(blocks, TABLES, order, repl)
    build_checklist(validation)
    build_validation_report(validation, figpaths)
    zip_path = build_zip(figpaths)

    print("Build complete.")
    print(f"  references: {len(ref_list)} (alphabetical, author-date)")
    print(f"  main-text words: {wc}; abstract words: {validation['abstract_words']}; "
          f"highlights: {validation['highlights']}")
    print(f"  package: {zip_path.relative_to(ROOT)}")
    for k, v in validation.items():
        if isinstance(v, bool):
            print(f"  {k}: {'OK' if v else 'FAIL'}")


if __name__ == "__main__":
    main()

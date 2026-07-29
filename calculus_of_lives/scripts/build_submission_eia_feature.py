#!/usr/bin/env python3
"""
Assemble the complete E&IA "Feature" (refereed) submission package into
../output/submission_eia_feature/ and zip it.

Contents:
  00_README_checklist.docx   - what is enclosed + submission notes
  01_cover_letter.docx       - cover letter to the Editors
  02_title_page.docx         - title page with author details (non-anonymous)
  03_manuscript.docx         - the manuscript (abstract, keywords, body,
                               figures inline, endnotes); no author info in text
  Figure1/2/3.png            - figures as separate 300 dpi rasters
  Figure1/2/3.tif            - figures as separate 300 dpi TIFFs
  figures_editable.pptx      - all three figures, one per slide (editable)
  figure_captions.txt        - figure captions (English)

Run after generate_figures.py, create_manuscript_en.py and
create_figures_pptx.py have produced their outputs in ../output.
"""

import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

import create_manuscript_en as ms

OUT = Path(__file__).resolve().parent.parent / "output"
PKG = OUT / "submission_eia_feature"
FONT = "Times New Roman"

TITLE = ms.TITLE
WORD_COUNT = ms.word_count()
N_NOTES = len(ms.NOTES)
FIG_CAPTIONS = ms.FIG_CAPTIONS


def _base(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(12)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.paragraph_format.line_spacing = 1.15
    st.paragraph_format.space_after = Pt(6)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(1)
        s.left_margin = s.right_margin = Inches(1)


def _p(doc, text="", *, bold=False, italic=False, size=12, align=None,
       space_after=6, space_before=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
    return p


def cover_letter():
    doc = Document(); _base(doc)
    _p(doc, "[Date]")
    _p(doc, "To the Editors\nEthics & International Affairs\nCarnegie Council "
       "for Ethics in International Affairs", space_after=12)
    _p(doc, "Dear Editors,", space_after=12)
    _p(doc,
       "I am pleased to submit the enclosed manuscript, \u201C" + TITLE +
       ",\u201D for consideration as a Feature in Ethics & International "
       "Affairs. The main text runs to approximately {:,} words (excluding "
       "the abstract and {} endnotes) and includes three figures. An abstract "
       "of under 250 words and a list of keywords are included."
       .format(WORD_COUNT, N_NOTES))
    _p(doc,
       "Public argument about the atomic bombing of Hiroshima and Nagasaki "
       "still turns on whether it was necessary. The article grants that "
       "empirical debate and then presses a prior, largely neglected "
       "question: whether human beings are the kind of agents entitled to "
       "weigh and decide who is to live and who is to die at all. I argue that "
       "the utilitarian justification presupposes a commensurability of lives "
       "that value pluralism gives us reason to deny; that, since finitude "
       "makes some choosing unavoidable, the morally decisive line falls not "
       "on the outcome of a weighing but on its attitude\u2014the taking "
       "calculus that prices the person versus the giving calculus that keeps "
       "faith with a dignity it cannot perfectly honor; that the scope of any "
       "defensible weighing is bounded by the difference between those who "
       "have taken danger into a vocation and those who remain its pure "
       "patients; and that the civilizational drift toward treating lives as "
       "summable is bound up with a choice between increasing supply and "
       "reducing demand. The upshot for preventing a third world war is not a "
       "sharper calculus of deterrence but a disciplined retreat from it.")
    _p(doc,
       "The article engages the just war tradition (Walzer, Anscombe), the "
       "critique of aggregation (Rawls, Williams, Nagel), value pluralism "
       "(Berlin, Raz), Kantian dignity, biopolitics (Agamben, Foucault), and "
       "the economics of sufficiency (Schumacher), and it addresses the "
       "principal objections\u2014quietism, the paralysis of policy, supreme "
       "emergency, and the realist reply\u2014in a dedicated section. It "
       "should suit the journal\u2019s readership of scholars, policy-makers, "
       "and informed general readers.")
    _p(doc,
       "The manuscript is original, is not under consideration elsewhere, and "
       "has not been published previously. It has been prepared for "
       "double-anonymous review: the manuscript file carries no "
       "author-identifying information, and author details appear only on the "
       "separate title page. I have no conflicts of interest to declare and "
       "received no funding for this work.", space_after=12)
    _p(doc, "Thank you for your consideration.", space_after=12)
    _p(doc, "Sincerely,", space_after=2)
    _p(doc, "Tatsuki Onishi", space_after=0)
    _p(doc, "Independent Researcher", space_after=0)
    _p(doc, "bougtoir@gmail.com", space_after=0)
    doc.save(PKG / "01_cover_letter.docx")


def readme_checklist():
    doc = Document(); _base(doc)
    _p(doc, "Submission package \u2014 " + TITLE, bold=True, size=14,
       space_after=8)
    _p(doc, "Target: Ethics & International Affairs, Feature (refereed / "
       "double-anonymous peer review).", space_after=10)

    _p(doc, "Enclosed files", bold=True, size=12, space_after=2)
    for line in [
        "00_README_checklist.docx \u2014 this file",
        "01_cover_letter.docx \u2014 cover letter",
        "02_title_page.docx \u2014 title page with author details",
        "03_manuscript.docx \u2014 abstract, keywords, body (figures inline), "
        "and endnotes; no author-identifying information",
        "Figure1.png / Figure2.png / Figure3.png \u2014 300 dpi rasters",
        "Figure1.tif / Figure2.tif / Figure3.tif \u2014 300 dpi TIFFs",
        "figures_editable.pptx \u2014 all three figures, one per slide, "
        "editable",
        "figure_captions.txt \u2014 figure captions",
    ]:
        _p(doc, "\u2022 " + line, size=11, space_after=2)

    _p(doc, "Checklist", bold=True, size=12, space_before=8, space_after=2)
    for line in [
        "Main text approx. {:,} words (within the ~7,000\u20138,000 word range "
        "for Features).".format(WORD_COUNT),
        "Abstract under 250 words; keywords supplied.",
        "{} Chicago-style numbered endnotes (not footnotes); every note is "
        "cited in the text and numbered in order of first appearance."
        .format(N_NOTES),
        "Three figures, each cited in the text (\u201CFigure 1/2/3\u201D) "
        "before it appears and placed immediately after that paragraph, with "
        "a caption below.",
        "Figures supplied inline in the manuscript and as separate 300 dpi "
        "PNG and TIFF files, plus an editable PPTX.",
        "All cited works are real, published sources; no fabricated "
        "references.",
        "All figures are regenerated from code (scripts/generate_figures.py); "
        "the manuscript contains no hardcoded data values.",
        "Manuscript file is free of author-identifying information; author "
        "details are on the separate title page.",
    ]:
        _p(doc, "\u2610 " + line, size=11, space_after=2)

    _p(doc, "Notes for the author before submitting", bold=True, size=12,
       space_before=8, space_after=2)
    for line in [
        "Submit via the journal\u2019s ScholarOne system "
        "(https://mc.manuscriptcentral.com/eia); authors uncomfortable with "
        "online submission may email a Word file to journal@cceia.org.",
        "Reproducibility: this is a normative philosophy article with no "
        "empirical dataset. All three figures are conceptual and regenerated "
        "by scripts/generate_figures.py in the public repository; each factual "
        "claim in the text is backed by a cited primary source.",
        "Open access / APC: as a refereed research article, a Feature is "
        "subject to the journal\u2019s Gold Open Access options; authors "
        "without funding may request an APC waiver after acceptance. Confirm "
        "the current fee and waiver terms with the editorial office.",
        "Replace [Date] on the cover letter before submitting.",
        "Verify the current word limit and house style against the "
        "journal\u2019s Instructions for Contributors at the time of "
        "submission.",
    ]:
        _p(doc, "\u2022 " + line, size=11, space_after=2)
    doc.save(PKG / "00_README_checklist.docx")


def main():
    if PKG.exists():
        shutil.rmtree(PKG)
    PKG.mkdir(parents=True)

    readme_checklist()
    cover_letter()
    shutil.copy(OUT / "title_page_en.docx", PKG / "02_title_page.docx")
    shutil.copy(OUT / "manuscript_en.docx", PKG / "03_manuscript.docx")
    shutil.copy(OUT / "figures_en.pptx", PKG / "figures_editable.pptx")

    figmap = {1: "fig1_layers", 2: "fig2_quadrant", 3: "fig3_asymptote"}
    for n, stem in figmap.items():
        shutil.copy(OUT / f"{stem}.png", PKG / f"Figure{n}.png")
        shutil.copy(OUT / f"{stem}.tif", PKG / f"Figure{n}.tif")

    caps = "\n\n".join(FIG_CAPTIONS[n] for n in (1, 2, 3))
    (PKG / "figure_captions.txt").write_text(caps + "\n", encoding="utf-8")

    zip_path = OUT / "submission_eia_feature.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        for f in sorted(PKG.iterdir()):
            z.write(f, arcname=f"submission_eia_feature/{f.name}")
    print("package:", PKG)
    print("zip:", zip_path)


if __name__ == "__main__":
    main()

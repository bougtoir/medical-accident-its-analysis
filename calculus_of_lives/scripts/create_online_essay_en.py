#!/usr/bin/env python3
"""
Generate the short, informal *Online Exclusive Essay* for the Carnegie Council
website (eiajournal.org): roughly 1,000-2,000 words, a more informal tone than
the print journal, light citation.

  "Who Appointed Us the Accountants of Life and Death?"

Derived from the same argument as essay_en.docx, but shorter, first-person,
with jargon stripped back and only four notes. No figure.

Output into ../output:
  online_essay_en.docx
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT = Path(__file__).resolve().parent.parent / "output"
OUT.mkdir(exist_ok=True)

FONT = "Times New Roman"
TITLE = "Who Appointed Us the Accountants of Life and Death?"
STANDFIRST = ("Eighty years after Hiroshima, we are still arguing about the "
              "wrong thing.")

BODY = [
"Eighty years on, we are still arguing about whether the atomic bombing of "
"Hiroshima and Nagasaki was necessary\u2014whether it really shortened the war "
"and spared more lives than it cost. Historians have gone back and forth on "
"the facts for decades: whether Japan was already close to surrender, whether "
"the casualty numbers were inflated afterward, whether the real audience for "
"the blast was Moscow.{1} That argument matters. But I want to step around "
"it, because it quietly skips the question that should come first.",

"Look at the shape of the usual defense. It is a piece of arithmetic. It "
"concedes that lives were taken and answers that more lives were saved, so "
"the books balance. Grant the defender every factual point\u2014that the bombs "
"did shorten the war, that an invasion would have been bloodier. A question "
"still remains that the arithmetic cannot ask itself, because it has already "
"assumed the answer: are we the kind of creatures who get to sit as the "
"accountants of one another\u2019s lives, deciding who shall live and who "
"shall die in someone else\u2019s place?",

"These are two different questions, and it helps to keep them apart. The "
"inner question\u2014\u201Cwas it necessary?\u201D\u2014works inside the "
"calculation; it audits a weighing whose legitimacy it never checks. The "
"outer question asks whether we are entitled to run the calculation at all. "
"You can win every point on the inner question and still have said nothing "
"about the outer one. Much of what passes for moral debate about war is "
"inner-question talk mistaken for the real thing\u2014and which question a "
"society reaches for first tells you what it already takes for granted about "
"the standing of persons.",

"Why be suspicious of the arithmetic? Because adding up lives is stranger "
"than it looks. When you fold one person\u2019s death into another\u2019s "
"survival and call the balance a justification, you treat a crowd of separate "
"people as though they were one big person with one big ledger.{2} There is "
"no super-person who enjoys the total; there are only the many, one of whom "
"is being spent for the rest. The citizens of Hiroshima were not compensated "
"by the survival of strangers. And there is a deeper trouble. The worth of a "
"life is not a quantity in some shared currency. Ask what a life is for, and "
"the answer comes back differently from inside every life. There is no master "
"scale on which a mother, a soldier, a scholar, and a child can be priced and "
"traded off. The ledger is not merely hard to compute; it is adding numbers "
"that were never in the same units to begin with.",

"And yet we cannot simply refuse to choose. Resources run out; the harvest "
"feeds some and not others; the lifeboat has a fixed number of seats. We "
"cannot make supply infinite or our needs vanish\u2014even the universe is "
"finite. So some weighing is unavoidable, and an ethics that pretends "
"otherwise is dodging the hard part rather than facing it.",

"Here is the move that matters. If we must weigh, the line between the decent "
"and the monstrous cannot fall between weighing and not weighing. It runs "
"through how we weigh. There is a taking way and a giving way, and they can "
"wear the same numerical face. In the taking way, the other person is a "
"variable to be zeroed out when the total improves; his death is a cost the "
"ledger, once balanced, treats as paid. In the giving way, you weigh because "
"you must, but you carry the weighing as a loss\u2014you keep faith with the "
"one you could not save and refuse to file him away as a rounding error. "
"Kant\u2019s old distinction fits here: a thing has a price and can be swapped "
"for an equivalent; a person has a dignity and admits of no equivalent at "
"all. To reason about people as though they had prices is the first step of "
"the taking way.",

"This is not a mood, or a soft psychological gloss on an otherwise identical "
"act. It changes what the act is. It shows up as grief, as the sense of a "
"wrong done even in doing what one had to, and as a direction\u2014the giving "
"agent takes the cost onto himself before displacing it onto others. What was "
"appalling in the reasoning that reached for the bomb was not that officials "
"used numbers; leaders under duress have to. It was that they reasoned in the "
"taking way about people who were never theirs to price, and did so having "
"never once paused at the outer question.",

"There is also the question of who may be placed on the scale at all. The "
"just-war tradition draws its line at noncombatants: people who have not made "
"themselves into agents of harm may not be made the targets of attack. A "
"soldier who dies at his post dies inside a calling he has taken up, and "
"there is a dignity available to him in that\u2014the acceptance of mortal "
"risk as part of a role freely shouldered. A civilian consented to no such "
"thing. He is a pure victim, and to place him on the scale is not to ask a "
"sacrifice of him but to make him pay a price he never agreed to bear. That "
"is why the atomic bombing is, on the outer view, not a close call. It did "
"not aim at soldiers, or even at workers at their machines, but at "
"cities\u2014at civilian populations as such, whose deaths were the lever "
"meant to move a government elsewhere. Elizabeth Anscombe said it plainly "
"when she objected to Oxford\u2019s honoring of Truman: deliberately killing "
"the innocent as a means to your ends is murder, and no accumulation of good "
"consequences turns murder into something else.{3}",

"A related blindness deserves naming here. It is often said that Japan is the "
"world\u2019s only atomic-bombed nation. In the narrow sense\u2014the only "
"country attacked with nuclear weapons in war\u2014that is true, and I say it "
"as a Japanese writer who has no wish to trade on national grievance. But if "
"we judge by what was actually done to human bodies, rather than by the "
"grammar of war, the atomic age has far more victims than Hiroshima and "
"Nagasaki: the Marshall Islanders exposed by the Castle Bravo test, the "
"Kazakhs downwind of Semipalatinsk, Aboriginal Australians at Maralinga, "
"Algerians in the French Sahara, Polynesians at Moruroa, the downwinders "
"of New Mexico and Nevada, and\u2014reportedly\u2014populations near the "
"Chinese test site at Lop Nur. Most of them were colonized, indigenous, or "
"otherwise without political standing\u2014the very people a calculus in the "
"taking way can treat as an empty margin. That they are so seldom counted "
"among \u201Cthe bombed\u201D is itself an instance of the argument I am "
"making: whose death gets named depends on who was granted a place on the "
"scale to begin with.",

"So who appointed us the accountants? There is a long modern habit of "
"treating whole populations as quantities to be managed, and of deciding "
"whose lives fall outside the circle it is a crime to kill. Strategic "
"bombing, the language of \u201Cacceptable losses,\u201D the cold actuarial "
"arithmetic of nuclear deterrence\u2014that is the taking way blown up to the "
"scale of nations and dressed as necessity. And it is not a museum piece. In "
"a world that is rearming, where the threat to hold entire populations "
"hostage is spoken aloud again, the temptation to treat the many as summable "
"quantities is returning, not receding.",

"Pull back to the largest picture. Civilizations meet scarcity in two broad "
"ways. Some try to increase supply\u2014they expand, conquer, open new "
"frontiers and draw them in. Others try to reduce demand\u2014they contract, "
"defend, discipline their own wanting, and aim at enough rather than more. It "
"is tempting to think an expanding civilization needs only an open frontier "
"to stay at peace, so that a new one\u2014outer space, say\u2014might drain "
"its pressure. I think that is half true, and the false half is instructive. "
"The frontier such cultures crave is really two different things. One is "
"material\u2014empty land, unused energy\u2014which more space can supply. The "
"other is about particular souls and the neighbor who believes otherwise, a "
"matter of intensity rather than acreage. You cannot pour the second hunger "
"into the first kind of space. That, I suspect, is why the age of "
"exploration flung open a vast new world and yet coincided with Europe\u2019s "
"worst wars of religion. Space enlarges the map; it leaves the mismatch "
"untouched, and may even feed the urge to carry one\u2019s own single truth "
"to the stars.",

"Here is the practical payoff. A third world war will not be prevented by a "
"sharper calculus of deterrence\u2014by more credible threats or more precise "
"targeting. That stays on the inner question and leaves untouched the habit "
"of treating lives as summable at all. What helps is work on the other side: "
"wanting less. This is the sane core of what E. F. Schumacher called Buddhist "
"economics\u2014an economics built around sufficiency rather than maximal "
"consumption.{4} For a person it means reducing your own taking before you "
"push the cost onto someone else. For a civilization it means choosing "
"defense and restraint over expansion and conquest. A culture that has "
"learned to want less has fewer occasions to place its neighbors on the "
"scale, and a smaller stake in the frontiers over which the great wars are "
"fought. For small and exposed nations especially, restraint of demand is "
"not only an ethic; it is a security posture.",

"We will never reach the point where no one ever has to be weighed; the world "
"is too finite for that. But we can choose which way we face. Eighty years "
"after Hiroshima, the question worth asking is not whether the bomb added up. "
"It is whether we still wish to be the kind of beings who keep the "
"books\u2014or whether we might, at last, prefer to give in order to keep "
"others alive rather than to take in order to live. That preference will not "
"abolish tragedy. But it is the ground on which a third world war might "
"become not merely unlikely, but unthinkable.",
]

NOTES = [
    "The classic statement of the necessity thesis is Henry L. Stimson, "
    "\u201CThe Decision to Use the Atomic Bomb,\u201D Harper\u2019s Magazine, "
    "February 1947. For the historians\u2019 challenge to it, see Gar "
    "Alperovitz, The Decision to Use the Atomic Bomb (New York: Knopf, 1995), "
    "and J. Samuel Walker, Prompt and Utter Destruction, 3rd ed. (Chapel Hill: "
    "University of North Carolina Press, 2016).",

    "John Rawls\u2019s phrase for this is the failure to respect \u201Cthe "
    "distinction between persons\u201D: A Theory of Justice, rev. ed. "
    "(Cambridge, MA: Harvard University Press, 1999), 23\u201324.",

    "G. E. M. Anscombe, \u201CMr Truman\u2019s Degree\u201D (1957), reprinted "
    "in her Ethics, Religion and Politics: Collected Philosophical Papers, "
    "vol. 3 (Oxford: Blackwell, 1981), 62\u201371.",

    "E. F. Schumacher, Small Is Beautiful: Economics as if People Mattered "
    "(London: Blond & Briggs, 1973), the chapter \u201CBuddhist Economics.\u201D",
]


def _set_base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(12)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = st.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(8)


def _add_runs_with_markers(par, text):
    for part in re.split(r"(\{\d+\})", text):
        m = re.fullmatch(r"\{(\d+)\}", part)
        if m:
            par.add_run(m.group(1)).font.superscript = True
        else:
            par.add_run(part)


def build():
    doc = Document()
    _set_base_style(doc)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(1)
        s.left_margin = s.right_margin = Inches(1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(4)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(STANDFIRST)
    sr.italic = True
    sp.paragraph_format.space_after = Pt(14)

    for text in BODY:
        pp = doc.add_paragraph()
        _add_runs_with_markers(pp, text)

    hp = doc.add_paragraph()
    hp.paragraph_format.space_before = Pt(16)
    hp.add_run("Notes").bold = True
    for i, note in enumerate(NOTES, start=1):
        np_ = doc.add_paragraph()
        np_.paragraph_format.line_spacing = 1.15
        np_.paragraph_format.space_after = Pt(2)
        np_.add_run(f"{i}. ").bold = True
        np_.add_run(note).font.size = Pt(10.5)

    doc.save(OUT / "online_essay_en.docx")


def word_count():
    return sum(len(re.sub(r"\{\d+\}", "", t).split()) for t in BODY)


if __name__ == "__main__":
    build()
    marks = []
    for t in BODY:
        marks += [int(m) for m in re.findall(r"\{(\d+)\}", t)]
    assert marks == sorted(marks), "citation markers not in order"
    assert sorted(set(marks)) == list(range(1, len(NOTES) + 1)), \
        "marker/notes mismatch"
    print("online_essay_en.docx written to", OUT,
          "| notes:", len(NOTES), "| approx words:", word_count())

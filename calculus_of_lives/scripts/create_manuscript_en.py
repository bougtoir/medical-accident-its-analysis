#!/usr/bin/env python3
"""
Generate the English manuscript for:

  "Beyond the Calculus of Lives: Commensurability, the Ethics of Giving,
   and the Civilizational Roots of Total War"

Target journal: Ethics & International Affairs (Cambridge University Press).
House style: Chicago-style numbered endnotes; ~150-word abstract.

Outputs into ../output:
  manuscript_en.docx   full manuscript, figures inline, endnotes at the end
  title_page_en.docx    separate anonymised-submission title page
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT = Path(__file__).resolve().parent.parent / "output"
OUT.mkdir(exist_ok=True)

TITLE = ("Beyond the Calculus of Lives: Commensurability, the Ethics of "
         "Giving, and the Civilizational Roots of Total War")

ABSTRACT = (
    "Debates over whether the atomic bombing of Japan was necessary are "
    "conducted inside a calculus: they ask whether a weighing of lives came "
    "out right. This article shifts the layer of the question and asks "
    "whether human beings are the kind of agents entitled to weigh and decide "
    "who is to live and who is to die. I argue, first, that the utilitarian "
    "justification presupposes a commensurability of lives that value "
    "pluralism gives us reason to deny; second, that because finitude makes "
    "some choosing unavoidable, the morally decisive line falls not on the "
    "outcome of a weighing but on its attitude\u2014whether the other is "
    "treated as an eliminable variable or as an end whose loss is borne as "
    "loss. I then read the sovereign power to weigh lives through Agamben and "
    "Foucault, and locate total war in a civilizational choice between "
    "increasing supply and reducing demand. Preventing a third world war, I "
    "conclude, requires not a better calculus but a disciplined retreat from "
    "it."
)

KEYWORDS = ("nuclear ethics; commensurability; value pluralism; just war; "
            "noncombatant immunity; biopolitics; Buddhist economics; "
            "the ethics of finitude")

# ---------------------------------------------------------------------------
# Manuscript body.  Each item is ("style", "text").
#   H1/H2 = headings, P = body paragraph, FIG = inline figure.
# Citation markers {n} are rendered as superscript endnote numbers; the
# endnotes themselves are listed (in order of first appearance) in NOTES.
# ---------------------------------------------------------------------------

BODY = [
("H1", "1. Shifting the Layer of the Question"),

("P",
 "Eight decades after the Pacific War, the destruction of Hiroshima and "
 "Nagasaki is still defended in a familiar idiom. On this account the bombs, "
 "for all their horror, were necessary: they ended the war quickly and so "
 "spared the far larger number of lives\u2014American and Japanese, combatant "
 "and civilian\u2014that an invasion of the home islands would have consumed."
 "{1} Historians have contested the factual premises of this story for "
 "decades, questioning whether surrender was imminent, whether the "
 "casualty projections were inflated after the fact, and whether the "
 "demonstration of power to the Soviet Union was the operative motive."
 "{2} That empirical quarrel is important, but it is not my subject. I want "
 "to notice something about the shape of the justification itself, quite "
 "apart from whether its numbers are right."),

("P",
 "The justification is a piece of arithmetic. It concedes that lives were "
 "taken and answers that more lives were saved; it sets the dead of two "
 "cities on one pan of a scale and the hypothetical dead of an invasion on "
 "the other, and reads off the balance. Grant, for the sake of argument, "
 "every factual claim the defender wants\u2014that the bombs did shorten the "
 "war, that they did avert a bloodier invasion. A question remains that the "
 "arithmetic cannot pose to itself, because it presupposes the answer: are "
 "human beings the kind of agents who may set themselves up as the "
 "accountants of one another's lives, deciding who shall live and who shall "
 "die in another's stead?"),

("P",
 "I call this a shift in the layer of the question (Figure 1). The "
 "\u201Cwas it necessary?\u201D debate operates inside the calculus; it "
 "audits a weighing whose legitimacy it never examines. The question I press "
 "sits outside the calculus and asks whether we are entitled to perform it "
 "at all. The two layers are logically independent: one can win every point "
 "on the inner layer\u2014concede that the sums add up\u2014and still have "
 "said nothing to the outer one. Much of what passes for moral debate about "
 "the bomb, and about war more generally, is an argument on the inner layer "
 "mistaken for an argument on the outer. The distinction is not merely "
 "analytic housekeeping. Which layer a society instinctively argues on "
 "reveals what it already takes for granted about the standing of persons."),

("FIG", "fig1_layers"),

("P",
 "The same doubling of layers holds for war in general, and in any of its "
 "forms. Economic coercion and cultural domination are also ways of securing "
 "one's own flourishing by drawing down another's; they, too, invite the "
 "inner-layer question of whether they \u201Cwork\u201D and the outer-layer "
 "question of whether one may play the role they cast one in. My aim in this "
 "article is to give the outer question a rigorous statement, to show that it "
 "has deep roots in the ethical tradition, and to connect it to a larger "
 "diagnosis of why civilizations go to total war\u2014and thus to the "
 "practical problem, acute again in a fluid and rearming world, of "
 "preventing a third one."),

("P",
 "A word on method and scope. This is a work of normative argument, not of "
 "historical revision; I take no stand on the disputed facts of 1945 beyond "
 "granting the defender his most favorable version, precisely so that the "
 "outer-layer question can be isolated from the empirical quarrel that "
 "usually swallows it. Nor do I claim that the argument settles what any "
 "particular government should do in a given crisis; its ambition is prior "
 "to that\u2014to recover a question that policy has learned not to ask, and "
 "to show that the recovery has consequences all the way down, from the "
 "metaethics of value to the design of security institutions. The choice of "
 "Hiroshima as the case is deliberate. It is the moment at which the "
 "arithmetic of lives was performed most explicitly, at the largest scale, "
 "and with the clearest public conscience; if the outer-layer question has "
 "force anywhere, it has force here, and what is learned here transfers to "
 "the quieter weighings\u2014of sanctions, of deterrence, of the acceptable "
 "civilian toll\u2014that fill the ordinary conduct of states. The bomb is "
 "not the disease but its purest symptom, and I use it as a lens rather than "
 "a target."),

("H1", "2. The Impossibility of the Calculus: Commensurability and Its Denial"),

("P",
 "Begin with the framework the justification tacitly assumes. "
 "Consequentialism in its classical utilitarian form holds that the right "
 "act is the one that produces the greatest balance of good over bad, summed "
 "impartially across all those affected.{3} Aggregation is essential to it: "
 "goods and harms accruing to distinct persons are entered into a single "
 "ledger and totalled, so that a large enough benefit spread across many can "
 "always, in principle, outweigh a grave harm concentrated on a few.{4} The "
 "atomic justification is this structure applied to lives. It treats being "
 "killed as a quantity of disvalue, holds that such quantities can be added "
 "across persons, and concludes that the smaller sum is to be preferred."),

("P",
 "The first and most famous objection is Rawls's. Utilitarianism, he wrote, "
 "\u201Cdoes not take seriously the distinction between persons\u201D: it "
 "extends to society as a whole the principle of rational choice for one "
 "person, as though the many individuals were so many desires within a "
 "single super-person whose satisfactions could be traded off against one "
 "another.{5} But persons are not fused in this way. The loss suffered by the "
 "one who is sacrificed is not compensated by the gains of those who are "
 "spared, because there is no experiential subject who both loses and gains. "
 "Bernard Williams pressed a companion point: the impartial sum requires the "
 "agent to treat his own deepest commitments\u2014and, I would add, the "
 "irreplaceable particularity of the person before him\u2014as just one more "
 "input to be overridden whenever the total demands it, alienating him from "
 "the very integrity that makes him a self.{6} Thomas Nagel, writing "
 "specifically of massacre in war, concluded that there are things one may "
 "not do to a person whatever the numbers, because the victim can rightly ask "
 "the agent, of what is being done to him in particular, \u201Cwhy?\u201D and "
 "receive no answer that is about him rather than about the sum.{7}"),

("P",
 "These are objections to aggregation. I want to add a deeper one, directed "
 "at the presupposition that makes aggregation intelligible in the first "
 "place: commensurability. To add two things on one scale, they must share a "
 "common measure, a single dimension of value in terms of which each can be "
 "expressed as so many units. The utilitarian ledger assumes that the worth "
 "of a life, and of everything a life contains, can be rendered on such a "
 "scale. Value pluralism denies this. Isaiah Berlin argued that the goods of "
 "human life are genuinely many, not reducible to one master value, and that "
 "they can conflict without there being any common currency in which their "
 "conflict is settled.{8} Joseph Raz refined the point into a formal thesis "
 "of incommensurability: two options are incommensurable when it is false "
 "both that one is better than the other and that they are of equal "
 "value\u2014when, that is, the very relation \u201Cgreater than, equal to, "
 "or less than\u201D fails to hold between them.{9}"),

("P",
 "Now consider what a human life is as an object of valuation. Each life is "
 "the bearer of its own ends, its own attachments, its own irreplaceable "
 "vantage on the world; the scales by which worth might be reckoned are as "
 "many as there are persons, and arguably more, since a single life contains "
 "plural and competing goods within itself. Even granting that within one "
 "life some rough ranking of options is possible, between lives there is no "
 "shared unit into which each can be converted for summing. The utilitarian "
 "must assume that \u201Cone death\u201D on the Hiroshima pan and \u201Cone "
 "death\u201D on the invasion pan denote fungible quantities of the same "
 "stuff. But a death is the extinction of a world, and worlds are not "
 "fungible.{10} This is a more radical charge than Rawls's. Rawls grants that "
 "the losses and gains are commensurable and objects that they may not be "
 "summed across the boundaries between persons; the incommensurability "
 "objection says that the entries were never expressible in a common unit to "
 "begin with, so that the summing is not a forbidden operation but a "
 "meaningless one.{11}"),

("P",
 "If this is right, the atomic justification does not merely reach a "
 "repugnant conclusion; it commits a category mistake. It presents as the "
 "output of a calculation something that no calculation could deliver, "
 "because the magnitudes it purports to compare were never on a common scale. "
 "The rhetoric of \u201Clives saved\u201D borrows the authority of arithmetic "
 "without the entitlement to it. This does not yet tell us what to do. It "
 "tells us that whatever we do, we cannot honestly represent it as the sum "
 "the defender claims to have computed."),

("P",
 "It will be objected that we compare lives constantly and unavoidably: the "
 "battlefield medic triages, the health system prices a year of life, the "
 "regulator fixes a value on statistical mortality. If incommensurability "
 "were true, the objection runs, all such practice would be impossible, and "
 "since it is not impossible, incommensurability must be false. The reply "
 "turns on a distinction the objection blurs. To choose under "
 "incommensurability is not the same as to choose by computing a common "
 "sum. The medic who saves the salvageable is not asserting that the "
 "salvageable life contains more units of a homogeneous good than the life "
 "he lets go; he is acting well under a tragic constraint without claiming "
 "that arithmetic vindicates him. What incommensurability denies is not that "
 "we can choose, nor even that some choices are better than others, but that "
 "the choice is licensed by, and answerable to, a cardinal quantity of "
 "aggregated value.{10} At most there is rough comparability\u2014enough to "
 "act, never enough to certify a killing as the balance struck by a "
 "ledger.{9} The atomic justification does not merely choose under tragedy; "
 "it claims the authority of a sum. That claim, and not the possibility of "
 "hard choice, is what the argument refuses."),

("H1", "3. Choosing Under Finitude: The Taking Calculus and the Giving Calculus"),

("P",
 "One might read the argument so far as a counsel of paralysis: if lives "
 "cannot be weighed, then any choice among them is illegitimate, and we are "
 "left mute before every tragic dilemma. That conclusion does not follow, and "
 "it is worth seeing exactly why, because the reason marks the turn from "
 "critique to construction."),

("P",
 "We could escape the need to choose among lives only under one of two "
 "conditions: if supply were infinite, so that no one's provision came at "
 "another's expense; or if demand were infinitesimal, so that embodied beings "
 "asked nothing of a finite world. Neither is available to us. The universe "
 "itself is finite, and nothing comes from nothing; there is no arrangement "
 "of matter and energy under which the last scarcity is abolished. Choice "
 "under scarcity is therefore not a contingent predicament we might one day "
 "engineer away but a standing feature of embodied existence. To choose is, "
 "in some measure, always to weigh. The critic of the calculus who imagines "
 "he has clean hands has merely declined to notice the weighings his own "
 "life already performs."),

("P",
 "So the honest position is not that we may never let considerations of "
 "number bear on action, but that the difference between a defensible and an "
 "indefensible weighing lies elsewhere than in its result. My proposal is "
 "that it lies in the attitude the agent takes toward those he must weigh. "
 "Consider two calculi that might issue in numerically identical choices. In "
 "the first\u2014call it the taking calculus\u2014the other person figures as "
 "an eliminable variable, a term in someone else's survival function that may "
 "be set to zero when the total improves. In the second\u2014the giving "
 "calculus\u2014the other figures as an end, someone to whom the agent stands "
 "in a relation of responsibility, so that even a choice that cannot be "
 "avoided is registered and borne as a wound rather than discharged as a "
 "computation."),

("P",
 "The distinction is Kantian at its root. Kant separated things that have a "
 "price, for which an equivalent can be substituted, from what has dignity, "
 "which is \u201Craised above all price and therefore admits of no "
 "equivalent.\u201D To have dignity is precisely to be incommensurable, to "
 "resist the substitution that pricing presupposes; and persons, as ends in "
 "themselves, have dignity rather than price.{12} The taking calculus prices "
 "the person\u2014fixes an equivalent for him and pays it out for a larger "
 "return. The giving calculus refuses the equivalence even as it acts, and so "
 "keeps faith with the dignity it cannot honor perfectly. Emmanuel Levinas "
 "gives the same structure a phenomenological form: the face of the other "
 "addresses me with a claim that precedes and exceeds any totality into which "
 "I might sum him, an infinity that says, before all calculation, \u2018you "
 "shall not kill.\u2019{13} To weigh in the giving mode is to weigh while "
 "remaining answerable to that address."),

("P",
 "What, concretely, distinguishes the two? Not, by hypothesis, the tallies. "
 "The marks are three. First, remainder: the giving calculus leaves a moral "
 "remainder\u2014grief, responsibility, the sense of a wrong done even in "
 "doing what one had to\u2014where the taking calculus, having balanced its "
 "books, feels itself owed nothing and mourns nothing. Second, direction of "
 "the tie-breaker: forced to choose, the giving agent absorbs cost onto "
 "himself and his own projects before displacing it onto others, whereas the "
 "taking agent optimizes without any such asymmetry. Third, the standing of "
 "the unchosen: the giving calculus continues to regard the one it could not "
 "save as a claimant with a grievance, not as a rounding error. What was "
 "monstrous in the reasoning that reached for the bomb, on this view, is not "
 "that officials did arithmetic\u2014statesmen under duress must reckon with "
 "numbers\u2014but that they reasoned in the taking mode about people who were "
 "never theirs to price, and did so, moreover, having never paused at the "
 "outer-layer question at all."),

("P",
 "A hard-nosed reader will protest that this makes morality turn on a "
 "feeling. If two acts have the same outcome, what can it matter whether the "
 "agent mourned? The objection assumes that the remainder is a mere "
 "sentiment decorating an act whose moral character is already fixed by its "
 "consequences. But the remainder is not decoration; it is evidence, and "
 "sometimes constitutive. Bernard Williams called the residue that a "
 "decent agent feels after a forced wrong agent-regret, and argued that its "
 "absence marks not toughness but a failure to register a moral reality: "
 "that a genuine claim was overridden, not annulled.{29} Martha Nussbaum, "
 "reading Greek tragedy, showed that the mature response to an irresolvable "
 "conflict of obligations is not to declare the defeated obligation "
 "cancelled by the victorious one but to remain answerable to both, and so "
 "to owe something to the party one had to fail.{30} The remainder tracks "
 "this fact. And because it does, it is not inert: the agent who feels it "
 "acts differently next time\u2014builds the institutions, the rules of "
 "engagement, the habits of restraint that a society without the remainder "
 "never troubles to build. Attitude, in this sense, is not private "
 "sentiment but a disposition with public and cumulative effects. It is the "
 "difference between a polity that treats each recourse to lethal force as a "
 "wound to be minimized and one that treats it as a cost to be optimized."),

("P",
 "I concede without embarrassment that the giving calculus still weighs, and "
 "so is not innocent of the tragic structure of finite life. To choose whom "
 "to feed with a finite harvest is still to prefer some to others. But the "
 "difference in attitude is not a psychological gloss on an otherwise "
 "identical act; it changes what the act is. A society that has lost the "
 "distinction\u2014that hears no difference between taking and giving so long "
 "as the totals match\u2014has lost the thing that separates provision from "
 "predation."),

("H1", "4. Whom One May Place on the Scale: Combatants, Civilians, and the "
       "Ethics of the Professional"),

("P",
 "Even granting that some weighing is unavoidable and that its mode is what "
 "matters morally, a further question arises about its scope: over whom may "
 "the weighing range at all? Here the just war tradition has an answer that "
 "the outer-layer view can endorse and deepen. The principle of "
 "noncombatant immunity holds that those who have not made themselves into "
 "agents of harm may not be made the direct objects of attack; the "
 "combatant, who has taken up the business of force, has forfeited the "
 "immunity the civilian retains.{14} Michael Walzer grounds the immunity in "
 "the idea that one becomes a legitimate target only by posing a threat, by "
 "entering the zone of danger as a participant\u2014which is why, on his "
 "account, even the munitions worker is liable while at the factory and "
 "making weapons, yet regains his immunity at home, where he is a person and "
 "not a threat.{15}"),

("P",
 "I want to reground this line where Walzer draws it, shifting its basis from "
 "function to consent and vocation. The soldier who dies at his post dies "
 "within a calling he has assumed; there is a dignity available to him in the "
 "form of professional honor, the acceptance of mortal risk as constitutive "
 "of a role freely shouldered\u2014the ethic captured in the older language "
 "of noblesse oblige, obligation that rank and capacity impose. The civilian "
 "has assumed no such calling. He is a pure patient of the violence, a victim "
 "in the strict sense of one who has consented to nothing. If a weighing must "
 "fall somewhere, it should fall toward those who have taken the risk into "
 "their vocation and away from those who have not\u2014not because the "
 "soldier's life is worth less on some scale (the argument of Section 2 "
 "forbids that comparison) but because he, unlike the civilian, has a "
 "relationship to the danger that makes his dying intelligible as something "
 "other than sheer wrong done to him. Here the professional's readiness to "
 "stand in harm's way is not a devaluing of his life but the highest "
 "expression of its dignity."),

("P",
 "The consent basis invites an immediate objection: most soldiers through "
 "history have been conscripts, pressed into service by states they did not "
 "choose, and some\u2014child soldiers above all\u2014cannot consent in any "
 "meaningful sense. If liability depended on literal autobiographical "
 "consent, the account would exempt most combatants and collapse. So I do "
 "not rest it there. Consent is the clearest case of a more general "
 "relation: standing within the role of a bearer of arms, and thereby "
 "presenting oneself as a threat and entering the zone of danger. The role "
 "is normative, not merely biographical; the conscript, however he arrived "
 "at it, occupies it and poses the threat that grounds his liability, "
 "whereas the civilian occupies no such role and poses no such threat. What "
 "the vocation account adds to the bare fact of threat is an explanation of "
 "why the soldier's death can be more than sheer wrong done to him\u2014why "
 "there is honor available in it\u2014and that explanation is fullest where "
 "the role is freely embraced and thinnest where it is coerced. This is why "
 "our moral unease rightly returns, rather than disappears, in the case of "
 "the child soldier: we sense that someone has been placed in the zone of "
 "danger who never could have shouldered its meaning, and the wrong belongs "
 "to those who put him there. The account thus tracks our judgments across "
 "the range of cases better than the crude function test, which would make "
 "liability a simple matter of causal contribution to the war effort."),

("P",
 "This yields a verdict on the bomb that requires no delicate balancing. The "
 "atomic attacks did not target the soldier at his post or even the worker at "
 "his factory; they took the city itself\u2014the civilian as such, the pure "
 "patient\u2014as the direct object, indeed as the intended instrument, the "
 "terror of whose death was the very mechanism by which surrender was to be "
 "produced. Under the doctrine of double effect, which permits foreseen but "
 "unintended civilian deaths incident to a legitimate military aim but "
 "forbids their use as a means, this is not a hard case.{16} The killing of "
 "the innocent was not a side effect of striking something else; it was the "
 "instrument. Elizabeth Anscombe saw this with unmatched clarity when she "
 "opposed the award of an honorary degree to Truman, naming him a murderer "
 "on the ground that the deliberate killing of the innocent as a means to "
 "victory is murder whatever its consequences\u2014and refusing the "
 "consequentialist frame that would make the question turn on a tally.{17} "
 "Rawls, half a century later, reached the same judgment.{18}"),

("P",
 "I add only this. That such a case is treated as genuinely debatable\u2014"
 "that reasonable people are expected to argue about whether the incineration "
 "of noncombatant cities might be licensed by its results\u2014is itself a "
 "symptom. It indicates a public culture in which the ideas of vocation and "
 "professional sacrifice have thinned to the point where the distinction "
 "between the one who has accepted danger and the one who has not no longer "
 "registers, and every death is flattened into an interchangeable unit of "
 "disvalue awaiting its place in a sum. The debatability of the bomb is a "
 "measure of how far the taking calculus has colonized the moral imagination."),

("P",
 "The most serious challenge to this verdict is Walzer's own doctrine of "
 "supreme emergency. Faced with an imminent and unprecedented threat\u2014"
 "Nazism triumphant, say\u2014Walzer allows that a political community may "
 "deliberately violate noncombatant immunity, as Britain arguably did in the "
 "early area bombing of German cities, because the very existence of a moral "
 "world is at stake.{31} If the prohibition can be overridden in extremis, "
 "the objection runs, then it is after all a weight in a balance, and we are "
 "back inside the calculus. I resist this at both levels. On the outer view, "
 "the prohibition on using the innocent as a mere means is not a heavy weight "
 "but a constraint of a different type\u2014not something that great enough "
 "stakes outweigh, but something that stakes do not reach; supreme emergency "
 "is precisely the taking calculus readmitted under the duress that makes it "
 "most tempting and least examined. But even those unwilling to go so far "
 "should notice the second point: the atomic case does not satisfy Walzer's "
 "own conditions. The United States in 1945 faced neither imminent defeat "
 "nor an existential threat to its moral community; it faced the prospect of "
 "a costlier victory it was already assured of winning. Whatever one thinks "
 "of supreme emergency in the abstract, the bombing was not an instance of "
 "it. The exception, even granted, does not cover the case."),

("H1", "5. The Sovereign Who Weighs: Biopolitics and a Civilizational Diagnosis"),

("P",
 "The outer-layer question\u2014may we weigh lives at all?\u2014has a "
 "political history, and it is largely a history of the question's "
 "suppression. To decide who lives and who dies is the oldest signature of "
 "sovereign power, and modern thought has traced how that power installs "
 "itself precisely by never letting the outer question be asked. Giorgio "
 "Agamben argues that sovereignty consists in the capacity to produce "
 "\u201Cbare life,\u201D life that may be killed without the killing counting "
 "as a crime\u2014to draw the line between the life that is politically "
 "qualified and the life that is exposed.{19} Michel Foucault describes the "
 "modern transformation of this power into biopolitics: a power that takes "
 "the life of populations as its object and administers it at the "
 "aggregate\u2014fostering some lives, disallowing others to the point of "
 "death\u2014through the apparatus of statistics, public health, and, at the "
 "limit, the calculated sacrifices of war.{20} The strategic bombing of "
 "cities is biopolitics in its most naked form: whole populations entered as "
 "quantities into an equation of national survival."),

("P",
 "What these analyses share with my argument is the recognition that the "
 "gravest violence begins not with a wrong answer to the question of who may "
 "weigh, but with the question's disappearance\u2014with the emergence of a "
 "standpoint from which weighing lives appears as an administrative task "
 "rather than a usurpation. To recover the outer-layer question is thus a "
 "political act: it reintroduces, against the sovereign's silence, the doubt "
 "about whether anyone is entitled to occupy the accountant's chair."),

("P",
 "Nuclear deterrence is this biopolitics raised to a principle and made "
 "permanent. Under mutual assured destruction, entire civilian populations "
 "are held, continuously and by design, as hostages to the conduct of their "
 "governments; the credibility of the threat depends on a standing readiness "
 "to incinerate them. Thomas Schelling described this with cold precision as "
 "the diplomacy of violence, in which the capacity to inflict pain on the "
 "innocent is not an unfortunate by-product of war but its very instrument "
 "of persuasion.{32} Here the taking calculus is not merely committed in a "
 "moment of decision but institutionalized as a posture maintained around "
 "the clock, for decades, as the ordinary condition of international order. "
 "The outer-layer question, asked of deterrence, is radical precisely "
 "because deterrence has normalized what it interrogates: it treats the "
 "reduction of populations to hostages\u2014to summable quantities of "
 "potential disvalue\u2014as the settled grammar of security. That this "
 "arrangement is widely regarded as prudent, even stabilizing, is the "
 "clearest measure of how completely the sovereign's chair has been "
 "occupied without the question of entitlement being raised."),

("P",
 "Why do polities drift toward this standpoint, and toward the total wars it "
 "licenses? Here I offer a wider diagnosis. Confronted with the finitude that "
 "makes provision a problem, a civilization can lean in one of two "
 "directions: it can seek to increase supply\u2014to expand the stock of "
 "what there is, if necessary by taking it from others\u2014or it can seek to "
 "reduce demand, to want less and so to need less of the world. The two "
 "leanings carry characteristic postures toward the outside: the supply "
 "strategy tends toward expansion and conquest, the demand strategy toward "
 "contraction and defense."),

("P",
 "These leanings have religious signatures, and it is tempting to line them "
 "up neatly\u2014expansionist monotheism against contractive polytheism and "
 "pantheism. There is something to the alignment. A tradition organized "
 "around a single universal truth is structurally disposed to treat the "
 "outside as unredeemed territory awaiting incorporation, so that mission and "
 "conquest travel together; Lynn White Jr. argued that the Western "
 "desacralization of nature, licensing its unlimited exploitation, has "
 "specifically monotheistic roots.{21} A tradition that finds the divine "
 "distributed through an already-full world\u2014pantheist, animist, "
 "immanent\u2014has, by contrast, no external frontier to conquer and inclines "
 "toward harmony and restraint; early Buddhism makes the reduction of demand "
 "into an explicit soteriology, treating craving itself as the thing to be "
 "extinguished."),

("P",
 "But the neat alignment must be resisted, for the axes come apart under "
 "scrutiny (Figure 2). Monotheism contains powerful demand-reducing "
 "currents\u2014the desert ascetics, the monastic vow of poverty, the "
 "Franciscan embrace of want\u2014and Max Weber famously showed how one such "
 "ascetic impulse, Protestant this-worldly self-denial, issued by an "
 "unintended dialectic in the limitless accumulation of capital, demand-"
 "reduction feeding supply-expansion.{22} Conversely, polytheisms have been "
 "engines of conquest: Rome absorbed the gods of the peoples it subdued, and "
 "some agrarian cults sustained their fertility by extracting sacrificial "
 "victims from without. The lesson is that \u201Csupply-increasing versus "
 "demand-reducing\u201D and \u201Cexpansion versus contraction\u201D are two "
 "independent axes, not one, and the religions distribute themselves across "
 "all four quadrants. What the map preserves from the naive typology is only "
 "this: the theology of an infinite external frontier and the ethic of "
 "unlimited supply are natural allies, and both stand at the maximal distance "
 "from the outer-layer question. Where the world is a frontier to be taken, "
 "the peoples on it are variables to be optimized. This is the metaphysical "
 "home of the taking calculus."),

("P",
 "I should be explicit about the status of these claims, since a typology "
 "that sorts religions risks caricature. I am not ranking traditions by "
 "moral worth, nor predicting the behavior of their adherents, who "
 "overwhelmingly live at peace with their neighbors. The claim is about the "
 "structural tendencies of doctrines under pressure\u2014about the "
 "directions in which a metaphysic inclines a civilization when scarcity "
 "forces the question of provision\u2014not about the character of believers. "
 "And the two-axis map is built precisely to defeat the caricature it might "
 "seem to invite: because supply-orientation and expansion are independent "
 "of each other, the framework predicts and accommodates the "
 "counterexamples rather than being embarrassed by them. Ascetic "
 "Protestantism sits in the demand-reducing column and yet, by Weber's "
 "dialectic, drove supply-expansion; imperial Rome sits among the "
 "polytheisms and yet conquered without limit. These are not anomalies to "
 "be explained away but data the model is designed to hold. What survives "
 "the qualifications is modest and, I think, secure: a metaphysic of the "
 "infinite external frontier lowers the threshold at which a civilization "
 "will reach for the taking calculus, and nothing in the typology licenses "
 "the inference from a people's creed to its guilt."),

("FIG", "fig2_quadrant"),

("H2", "5.1 The Frontier Mismatch: Intensive and Extensive Frontiers, and a "
       "Third Incommensurability"),

("P",
 "A single mechanism can be extracted from this diagnosis, and it bears "
 "directly on the contemporary fear of civilizational war. A tradition of "
 "the infinite external frontier depends, for its peace, on the frontier "
 "remaining open: Frederick Jackson Turner argued that the vitality and "
 "cohesion of an expansionist society are sustained by the availability of "
 "unclaimed land, so that the closing of the frontier throws the "
 "expansionist energy back upon the society itself.{23} One might infer that "
 "once the known and unclaimed regions are exhausted, collision becomes "
 "inevitable, and that a new frontier\u2014outer space, say\u2014might "
 "therefore drain the pressure and pacify the wars of the expansionist "
 "creeds. I think this inference is half right, and that seeing why sharpens "
 "the whole account."),

("P",
 "The frontier that such traditions crave is not one thing but two, and the "
 "two have different topologies. The material frontier is extensive: empty "
 "land, unused energy, a matter of volume that more space can always "
 "enlarge. The spiritual frontier is intensive: it is directed at particular "
 "souls, at the neighbor who believes otherwise, at territory that is not "
 "empty but wrongly occupied\u2014a matter of density, not of volume. The "
 "two cannot be exchanged for each other. An intensive demand cannot be "
 "discharged into an extensive supply, because the object the intensive "
 "drive seeks (this heretic, converted or erased) is not the object the "
 "extensive frontier offers (that empty expanse, settled). This is a third "
 "incommensurability, and it is the same blade that has cut twice already, "
 "now recurring at the civilizational scale: as lives are not convertible "
 "into a common unit (Section 2) and dignity is not convertible into price "
 "(Section 3), so the spiritual and material frontiers are not convertible "
 "into each other. Conflict springs less from the exhaustion of frontier "
 "than from the mismatch between the two frontiers\u2014from the attempt, and "
 "the failure, to vent an intensive pressure through an extensive outlet."),

("P",
 "The historical record fits the mismatch better than the exhaustion story. "
 "The Age of Discovery flung open a vast extensive frontier, yet Europe\u2019s "
 "most terrible religious war, the Thirty Years\u2019 War, raged in the same "
 "epoch: the new material space did nothing to relieve the intensive "
 "pressure of confessional truth-claims turned against the neighbor. "
 "Conversely, where the intensive demand is low\u2014where a plurality of "
 "truths has been inwardly accepted\u2014religious war is rare even under "
 "material scarcity. What varies is not the quantity of available frontier "
 "but the phase-match between the two kinds. The implication for the space "
 "age is sobering. Space enlarges only the extensive frontier; it leaves the "
 "phase gap untouched, and the very triumph of expansion may amplify the "
 "intensive drive to carry one\u2019s own truth to the stars\u2014a Jevons "
 "rebound of the spirit. The one pacifying path space might open is not the "
 "frontier as outlet but the frontier as mirror: the \u2018overview "
 "effect,\u2019 the relativizing of tribal and confessional boundaries "
 "reported by those who have seen the earth whole and undivided from "
 "outside.{24} But that works by lowering the intensive demand "
 "itself\u2014by shrinking the claim rather than enlarging the "
 "supply\u2014and so belongs, once again, to the demand-reducing hand to "
 "which I now turn."),

("H1", "6. Deterring Total War: The Two-Handed Task of Technology and Thought"),

("P",
 "If the analysis is right, the prevention of a third world war cannot be "
 "secured by refining the calculus of deterrence\u2014by making our weighings "
 "of populations more accurate or our threats more credible. That project "
 "remains on the inner layer and leaves untouched the disposition to treat "
 "lives as summable at all. What is required is movement along both axes at "
 "once, technological and ideational, toward the asymptotes at which the "
 "pressure to weigh lives relaxes (Figure 3)."),

("FIG", "fig3_asymptote"),

("P",
 "The technological hand works on supply: more energy, more food, more of the "
 "material substrate of life, so that fewer provisions come at a rival's "
 "expense. Its promise is real but bounded, and bounded in two ways. It meets "
 "a final ceiling in the finitude of the universe; and, long before that, it "
 "is undercut by a perverse dynamic first noticed by William Stanley Jevons, "
 "who observed that improvements in the efficiency of coal use increased "
 "rather than decreased total consumption, because efficiency cheapens use "
 "and so multiplies it.{25} The rebound generalizes: supply-side progress, "
 "pursued alone, tends to inflate the very demand it was meant to satisfy, so "
 "that the frontier recedes as fast as one advances on it. Technology "
 "unaccompanied by any discipline of wanting does not close the gap between "
 "demand and supply; it runs in place, and licenses the endless expansion "
 "that keeps the taking calculus in business."),

("P",
 "Hence the indispensability of the second, ideational hand, which works on "
 "demand: the deliberate cultivation of wanting less. E. F. Schumacher's "
 "\u201CBuddhist economics\u201D made this its first principle, defining the "
 "rational aim of economic life as the maximum of well-being with the minimum "
 "of consumption, rather than the maximum of consumption as such.{26} What I "
 "propose adds to Schumacher's economics the ethical and political content of "
 "the preceding sections: demand-reduction matters not only because it eases "
 "the material scramble but because it is the practical form of stepping down "
 "from the sovereign's chair. To reduce one's own demand is to withdraw, by "
 "just that much, from the position of the taker\u2014to make oneself the "
 "first bearer of cost rather than the optimizer who displaces cost onto "
 "others. The discipline of wanting less is thus continuous with the giving "
 "calculus at the level of the person: the individual who lowers his own "
 "claims on a finite world is doing, in small, what a peaceable civilization "
 "must do in large. It is here that private asceticism and the avoidance of "
 "world war turn out to be the same movement viewed at different scales, and "
 "here too that the tradition of ahimsa, which refuses to license impure "
 "means by pure ends, meets the argument of Section 4.{27}"),

("P",
 "This is not a program for the abolition of scarcity or of choice; the "
 "asymptote is approached, never reached, and a residue of tragic weighing "
 "will always remain (Figure 3). Nor is it a counsel of passivity in the face "
 "of aggression: a civilization organized around contraction and defense must "
 "still defend, and Section 4's ethic of the professional who accepts danger "
 "on others' behalf is precisely the honorable form that defense takes. The "
 "claim is directional. A world that pours its ingenuity into raising supply "
 "while cultivating limitless demand is a world whose frontier is forever "
 "receding and whose politics is therefore permanently expansionary\u2014the "
 "structural precondition of total war. A world that couples technological "
 "supply with an ethic of reduced demand is one in which the pressure to "
 "resolve scarcity by taking, and so to weigh the lives of others, steadily "
 "diminishes. Hans Jonas argued that the scale of modern technological power "
 "demands a new ethics answerable to the future and to the finite conditions "
 "of life on earth;{28} the two-handed task is what such an ethics looks like "
 "when it is turned specifically against the recurrence of total war."),

("P",
 "The realist will answer that this is a luxury of the secure. A civilization "
 "that unilaterally reduces its wants, he says, merely clears the field for "
 "the unrestrained, who will expand into the room it vacates; demand-"
 "reduction is a recipe for conquest by the greedy. The objection would be "
 "decisive against a doctrine of disarmament, but that is not the doctrine. "
 "What is to be reduced is demand\u2014the appetite that drives a polity to "
 "resolve scarcity by taking from others\u2014not the capacity to defend what "
 "one has. The two-handed task keeps the defensive hand; Section 4's ethic of "
 "the professional who accepts danger on others' behalf is exactly the "
 "honorable form that a contractive, defensive civilization's security takes. "
 "Indeed the argument cuts the other way for the exposed. A small state on a "
 "dangerous border cannot make itself safe by out-expanding a larger "
 "neighbor; its security, if it has any, lies in becoming hard to digest and "
 "cheap to leave alone\u2014in defensive resilience joined to a politics that "
 "gives its neighbors no frontier to covet. For such a state, restraint of "
 "demand is not utopian sentiment but the most realistic security posture "
 "available. Joseph Nye argued that a defensible nuclear ethics must weave "
 "together consequences, rights, and intentions rather than collapsing into "
 "any one of them;{33} the account offered here specifies which thread is "
 "non-negotiable\u2014the refusal to make the innocent a mere instrument\u2014"
 "and which civilizational disposition makes honoring it sustainable."),

("H1", "7. Objections and Replies"),

("P",
 "Three further objections deserve a direct answer. The first is that the "
 "position is quietist: that in the face of atrocities demanding action, an "
 "ethic of wanting less and holding the scale reluctantly does nothing for "
 "the victims. But the argument is not a private spirituality of withdrawal. "
 "Its institutional content is considerable: it underwrites the laws of "
 "armed conflict and their prohibition on targeting civilians, the "
 "protection of the noncombatant as a right rather than a courtesy, and a "
 "security policy that measures itself by the wounds it refuses to inflict. "
 "To reduce demand at the civilizational scale is a program of political "
 "economy, not a counsel of retreat. Quietism would be the refusal to hold "
 "the scale at all; what is proposed is to hold it differently, and to build "
 "the institutions that a giving posture requires."),

("P",
 "The second objection is that incommensurability, if taken seriously, "
 "paralyzes policy: governments must compare and aggregate, and a doctrine "
 "that forbids it forbids governing. The reply completes the one begun in "
 "Section 2. Incommensurability does not forbid choice; it forbids a "
 "particular self-description of choice\u2014the claim that a killing has "
 "been vindicated by a computed surplus of value. Sound institutions already "
 "reflect this. The lexical priority of certain rights over aggregate "
 "welfare, the side-constraints that no cost-benefit total may breach, the "
 "categorical prohibitions of the laws of war: these are the giving calculus "
 "written into public structure, mechanisms by which a polity binds itself "
 "in advance not to treat certain persons as summable. Far from paralyzing "
 "policy, the view recommends a familiar and workable architecture, and "
 "explains why its constraints are not mere prudence but principle."),

("P",
 "The third objection is that the distinction between taking and giving is "
 "unverifiable, a matter of inner attitude to which we have no access, and so "
 "useless for judgment. But the distinction is not locked inside the agent's "
 "breast. Its three marks are public. Whether a polity leaves a moral "
 "remainder shows in how it memorializes its dead and its enemy's, in "
 "whether it holds inquiries or averts its eyes; whether it directs the tie-"
 "breaker against itself shows in whether it accepts risk to its own forces "
 "to spare noncombatants, as the law of proportionality demands; whether it "
 "regards the unchosen as claimants shows in reparation, apology, and the "
 "revision of doctrine. Attitude leaves tracks. We read it in individuals "
 "every day, and we can read it in institutions by the same signs, writ "
 "large."),

("H1", "8. Conclusion"),

("P",
 "The question of whether the atomic bombings were necessary is the wrong "
 "question, or rather a question on the wrong layer. Answered on its own "
 "terms, in the affirmative or the negative, it leaves the graver matter "
 "untouched: whether human beings are entitled to weigh and trade the lives "
 "of others at all. I have argued that the weighing the justification claims "
 "to perform is not merely impermissible but, given the incommensurability of "
 "lives, unintelligible as arithmetic; that because finitude nonetheless "
 "forces us to choose, what distinguishes provision from predation is not the "
 "outcome of a weighing but its attitude\u2014the taking calculus that prices "
 "the person against the giving calculus that keeps faith with a dignity it "
 "cannot perfectly honor; that the scope of any defensible weighing is "
 "constrained by the difference between those who have taken danger into a "
 "vocation and those who remain pure patients of it; and that the drift of "
 "whole civilizations toward the sovereign's chair, from which lives appear "
 "as summable quantities, is bound up with a choice between increasing supply "
 "and reducing demand."),

("P",
 "These threads are one argument. The incommensurability that unmasks the "
 "calculus, the attitude that survives its collapse, the vocation that "
 "bounds its reach, the frontier that mismatches and so ignites, the two "
 "hands that might yet prevent\u2014each is a face of a single refusal: the "
 "refusal to occupy the accountant's chair as though it were ours by right. "
 "The unity matters, because the danger is unified. A civilization does not "
 "arrive at the incineration of cities by a single wicked decision but by a "
 "long habituation, in which the person becomes a quantity, the quantity "
 "becomes summable, and the summing becomes governance. To interrupt that "
 "habituation at any one of its stages is to weaken it at all of them. That "
 "is why an argument pitched at the level of metaphysics and attitude is not "
 "a retreat from politics but its most durable form: institutions decay, "
 "treaties lapse, deterrents are outbuilt, but a people that has learned to "
 "flinch at the sovereign's chair carries the prohibition in a place no "
 "adversary can disarm."),

("P",
 "The upshot for the prevention of a third world war is not a sharper "
 "calculus but a disciplined retreat from it. We are, by our nature as finite "
 "and embodied beings, condemned to some weighing; the whole moral question "
 "is in what spirit, over whom, and with how much reluctance we do it. To "
 "live for oneself by taking from another is the oldest of our motions. To "
 "give so that another may live\u2014accepting that this, too, holds a scale, "
 "but holding it as one who would rather bear the cost than impose it\u2014is "
 "the harder motion, and the one a survivable century will have to learn. "
 "Resources are finite, and so the scale will not be laid down in our time. "
 "But whose hand is on it, and whether that hand trembles, is ours to decide."),
]

# ---------------------------------------------------------------------------
# Endnotes (Chicago full-note style), in order of first appearance.
# ---------------------------------------------------------------------------

NOTES = [
 "Henry L. Stimson, \u201CThe Decision to Use the Atomic Bomb,\u201D "
 "Harper\u2019s Magazine 194 (February 1947): 97\u2013107.",

 "Gar Alperovitz, The Decision to Use the Atomic Bomb and the Architecture "
 "of an American Myth (New York: Alfred A. Knopf, 1995); Tsuyoshi "
 "Hasegawa, Racing the Enemy: Stalin, Truman, and the Surrender of Japan "
 "(Cambridge, MA: Harvard University Press, 2005).",

 "Jeremy Bentham, An Introduction to the Principles of Morals and "
 "Legislation (1789; Oxford: Clarendon Press, 1996); John Stuart Mill, "
 "Utilitarianism (1863), ed. Roger Crisp (Oxford: Oxford University Press, "
 "1998).",

 "Henry Sidgwick, The Methods of Ethics, 7th ed. (London: Macmillan, 1907), "
 "esp. bk. IV. On aggregation as the defining commitment, see Samuel "
 "Scheffler, The Rejection of Consequentialism, rev. ed. (Oxford: Clarendon "
 "Press, 1994).",

 "John Rawls, A Theory of Justice, rev. ed. (Cambridge, MA: Harvard "
 "University Press, 1999), 24\u201325.",

 "Bernard Williams, \u201CA Critique of Utilitarianism,\u201D in J. J. C. "
 "Smart and Bernard Williams, Utilitarianism: For and Against (Cambridge: "
 "Cambridge University Press, 1973), 108\u2013118.",

 "Thomas Nagel, \u201CWar and Massacre,\u201D Philosophy & Public Affairs 1, "
 "no. 2 (1972): 123\u2013144.",

 "Isaiah Berlin, The Crooked Timber of Humanity: Chapters in the History of "
 "Ideas, ed. Henry Hardy (London: John Murray, 1990), 1\u201319.",

 "Joseph Raz, The Morality of Freedom (Oxford: Clarendon Press, 1986), "
 "321\u2013366.",

 "For the collection that frames the debate, see Ruth Chang, ed., "
 "Incommensurability, Incomparability, and Practical Reason (Cambridge, MA: "
 "Harvard University Press, 1997), editor\u2019s introduction.",

 "The contrast is developed in John Gray, Isaiah Berlin (Princeton: "
 "Princeton University Press, 1996), chap. 3, on the political stakes of "
 "pluralism.",

 "Immanuel Kant, Groundwork of the Metaphysics of Morals (1785), trans. and "
 "ed. Mary Gregor (Cambridge: Cambridge University Press, 1998), 4:434\u2013"
 "435.",

 "Emmanuel Levinas, Totality and Infinity: An Essay on Exteriority, trans. "
 "Alphonso Lingis (Pittsburgh: Duquesne University Press, 1969), 194\u2013"
 "201, 262\u2013263.",

 "The principle of distinction is codified in Protocol I Additional to the "
 "Geneva Conventions (1977), arts. 48, 51. For its moral foundations, see "
 "Michael Walzer, Just and Unjust Wars: A Moral Argument with Historical "
 "Illustrations, 5th ed. (New York: Basic Books, 2015), chap. 9.",

 "Walzer, Just and Unjust Wars, 145\u2013146 (on the munitions worker).",

 "The doctrine descends from Thomas Aquinas, Summa Theologiae II-II, q. 64, "
 "a. 7; for its modern statement see Philippa Foot, \u201CThe Problem of "
 "Abortion and the Doctrine of the Double Effect,\u201D Oxford Review 5 "
 "(1967): 5\u201315.",

 "G. E. M. Anscombe, \u201CMr Truman\u2019s Degree\u201D (1957), reprinted "
 "in her Ethics, Religion and Politics: Collected Philosophical Papers, vol. "
 "3 (Oxford: Blackwell, 1981), 62\u201371.",

 "John Rawls, \u201CFifty Years after Hiroshima\u201D (1995), in Collected "
 "Papers, ed. Samuel Freeman (Cambridge, MA: Harvard University Press, "
 "1999), 565\u2013572.",

 "Giorgio Agamben, Homo Sacer: Sovereign Power and Bare Life, trans. Daniel "
 "Heller-Roazen (Stanford: Stanford University Press, 1998), 71\u201386.",

 "Michel Foucault, \u201CSociety Must Be Defended\u201D: Lectures at the "
 "Coll\u00e8ge de France, 1975\u201376, trans. David Macey (New York: "
 "Picador, 2003), 239\u2013264; and The History of Sexuality, vol. 1, trans. "
 "Robert Hurley (New York: Pantheon, 1978), 135\u2013145.",

 "Lynn White Jr., \u201CThe Historical Roots of Our Ecologic Crisis,\u201D "
 "Science 155, no. 3767 (1967): 1203\u20131207.",

 "Max Weber, The Protestant Ethic and the Spirit of Capitalism (1905), "
 "trans. Talcott Parsons (London: Routledge, 1992), esp. chap. 5.",

 "Frederick Jackson Turner, \u201CThe Significance of the Frontier in "
 "American History\u201D (1893), in The Frontier in American History (New "
 "York: Henry Holt, 1920), 1\u201338.",

 "On the \u2018overview effect,\u2019 see Frank White, The Overview Effect: "
 "Space Exploration and Human Evolution, 3rd ed. (Reston, VA: AIAA, 2014); "
 "and Carl Sagan, Pale Blue Dot: A Vision of the Human Future in Space (New "
 "York: Random House, 1994), 3\u20139.",

 "William Stanley Jevons, The Coal Question, 2nd ed. (London: Macmillan, "
 "1866), chap. 7. For the contemporary literature, see Blake Alcott, "
 "\u201CJevons\u2019 Paradox,\u201D Ecological Economics 54, no. 1 (2005): "
 "9\u201321.",

 "E. F. Schumacher, \u201CBuddhist Economics,\u201D in Small Is Beautiful: "
 "Economics as if People Mattered (London: Blond & Briggs, 1973), 50\u2013"
 "58.",

 "M. K. Gandhi, Non-Violent Resistance (Satyagraha) (New York: Schocken, "
 "1961); on craving (ta\u1e47h\u0101) as the root to be extinguished, "
 "Walpola Rahula, What the Buddha Taught, 2nd ed. (New York: Grove Press, "
 "1974), 29\u201344.",

 "Hans Jonas, The Imperative of Responsibility: In Search of an Ethics for "
 "the Technological Age, trans. Hans Jonas with David Herr (Chicago: "
 "University of Chicago Press, 1984).",

 "Bernard Williams, \u201CMoral Luck,\u201D in Moral Luck: Philosophical "
 "Papers 1973\u20131980 (Cambridge: Cambridge University Press, 1981), "
 "20\u201339, on agent-regret.",

 "Martha C. Nussbaum, The Fragility of Goodness: Luck and Ethics in Greek "
 "Tragedy and Philosophy, rev. ed. (Cambridge: Cambridge University Press, "
 "2001), esp. chaps. 2\u20133.",

 "Michael Walzer, Just and Unjust Wars, 5th ed. (New York: Basic Books, "
 "2015), chap. 16 (\u201CSupreme Emergency\u201D); for criticism, see "
 "Brian Orend, The Morality of War, 2nd ed. (Peterborough: Broadview, "
 "2013), chap. 5.",

 "Thomas C. Schelling, Arms and Influence (New Haven: Yale University "
 "Press, 1966), chap. 1 (\u201CThe Diplomacy of Violence\u201D).",

 "Joseph S. Nye Jr., Nuclear Ethics (New York: Free Press, 1986), esp. "
 "chaps. 1 and 5.",
]

# ---------------------------------------------------------------------------
# docx rendering
# ---------------------------------------------------------------------------

FONT = "Times New Roman"


def _set_base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(12)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = st.paragraph_format
    pf.line_spacing = 2.0
    pf.space_after = Pt(0)


def _renumber():
    """Renumber citation markers by order of first appearance and reorder the
    notes to match (Vancouver rule). Markers in BODY carry stable identifiers
    (the index into NOTES, 1-based); the displayed number is derived here, so
    new notes can be inserted anywhere without hand-renumbering."""
    order = []
    for _style, text in BODY:
        for m in re.findall(r"\{(\d+)\}", text):
            n = int(m)
            if n not in order:
                order.append(n)
    assert sorted(order) == list(range(1, len(NOTES) + 1)), (
        "marker/notes mismatch: markers=%s notes=%d"
        % (sorted(order), len(NOTES)))
    mapping = {old: i + 1 for i, old in enumerate(order)}
    notes_ordered = [NOTES[old - 1] for old in order]
    return mapping, notes_ordered


def _add_runs_with_markers(par, text, mapping):
    """Split text on {n} citation markers; render mapped n as superscript."""
    parts = re.split(r"(\{\d+\})", text)
    for part in parts:
        m = re.fullmatch(r"\{(\d+)\}", part)
        if m:
            r = par.add_run(str(mapping[int(m.group(1))]))
            r.font.superscript = True
        else:
            par.add_run(part)


def build_manuscript():
    mapping, notes_ordered = _renumber()
    doc = Document()
    _set_base_style(doc)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(15)
    p.paragraph_format.space_after = Pt(12)

    # Abstract
    h = doc.add_paragraph()
    hr = h.add_run("Abstract")
    hr.bold = True
    for para_text in [ABSTRACT]:
        ap = doc.add_paragraph(para_text)
        ap.paragraph_format.line_spacing = 1.5

    kp = doc.add_paragraph()
    kr = kp.add_run("Keywords: ")
    kr.bold = True
    kp.add_run(KEYWORDS)
    kp.paragraph_format.line_spacing = 1.5
    kp.paragraph_format.space_after = Pt(12)

    fig_count = 0
    for style, text in BODY:
        if style == "H1":
            hp = doc.add_paragraph()
            hp.paragraph_format.space_before = Pt(12)
            hp.paragraph_format.space_after = Pt(6)
            rr = hp.add_run(text)
            rr.bold = True
            rr.font.size = Pt(13)
        elif style == "H2":
            hp = doc.add_paragraph()
            hp.paragraph_format.space_before = Pt(8)
            hp.paragraph_format.space_after = Pt(4)
            rr = hp.add_run(text)
            rr.bold = True
            rr.italic = True
            rr.font.size = Pt(12)
        elif style == "P":
            pp = doc.add_paragraph()
            pp.paragraph_format.first_line_indent = Inches(0.4)
            _add_runs_with_markers(pp, text, mapping)
        elif style == "FIG":
            fig_count += 1
            img = OUT / f"{text}.png"
            ip = doc.add_paragraph()
            ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ip.paragraph_format.space_before = Pt(14)
            ip.add_run().add_picture(str(img), width=Inches(5.6))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_before = Pt(6)
            cap.paragraph_format.line_spacing = 1.0
            cr = cap.add_run(FIG_CAPTIONS[fig_count])
            cr.font.size = Pt(10)

    # Endnotes
    hp = doc.add_paragraph()
    hp.paragraph_format.space_before = Pt(18)
    hr = hp.add_run("Notes")
    hr.bold = True
    hr.font.size = Pt(13)
    for i, note in enumerate(notes_ordered, start=1):
        np_ = doc.add_paragraph()
        np_.paragraph_format.line_spacing = 1.5
        np_.paragraph_format.space_after = Pt(2)
        np_.add_run(f"{i}. ").bold = True
        np_.add_run(note).font.size = Pt(10.5)

    doc.save(OUT / "manuscript_en.docx")


FIG_CAPTIONS = {
    1: ("Figure 1. The funnel of attention. Every question the bombing raises "
        "enters at the top, but a layer of common sense and preconception\u2014"
        "that we are entitled to do the weighing\u2014acts as a sieve: only the "
        "inner, technical (\u201Cwas it necessary?\u201D) question passes through "
        "into public debate, while the outer question\u2014whether such weighing "
        "is ours to perform at all\u2014is screened out and never comes into "
        "view."),
    2: ("Figure 2. Two independent axes\u2014supply-increasing vs. "
        "demand-reducing, and expansion vs. contraction\u2014with "
        "illustrative religious and economic types. The naive alignment of "
        "monotheism with expansion and pantheism with restraint holds only as "
        "a tendency; ascetic Protestantism and expansionist polytheism occupy "
        "the off-diagonal cells."),
    3: ("Figure 3. Asymptotic model of relief from the calculus. Raising "
        "supply approaches a cosmic ceiling; reducing demand approaches an "
        "irreducible floor set by embodied need. Neither reaches zero "
        "scarcity, so a residue of tragic weighing always remains; the claim "
        "is directional, not utopian."),
}


def build_title_page():
    doc = Document()
    _set_base_style(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(15)
    for label, val in [
        ("Author", "Tatsuki Onishi"),
        ("Affiliation", "Independent Researcher"),
        ("Corresponding author", "bougtoir@gmail.com"),
    ]:
        ap = doc.add_paragraph()
        ap.paragraph_format.line_spacing = 1.5
        ap.add_run(f"{label}: ").bold = True
        ap.add_run(val)
    for extra in [
        "Word count (main text, excluding notes and abstract): see build log.",
        "Conflicts of interest: none declared.",
        "Funding: none.",
    ]:
        ep = doc.add_paragraph(extra)
        ep.paragraph_format.line_spacing = 1.5
    doc.save(OUT / "title_page_en.docx")


def word_count():
    words = 0
    for style, text in BODY:
        if style in ("H1", "P"):
            clean = re.sub(r"\{\d+\}", "", text)
            words += len(clean.split())
    return words


if __name__ == "__main__":
    build_manuscript()
    build_title_page()
    print("manuscript_en.docx and title_page_en.docx written to", OUT)
    print("main-text word count (approx):", word_count())

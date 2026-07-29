#!/usr/bin/env python3
"""
Assemble the complete E&IA "Essay" submission package into
../output/submission_eia_essay/ and zip it.

Contents:
  00_README_checklist.docx   - what is enclosed + submission notes
  01_cover_letter.docx       - cover letter to the Editors (author placeholders)
  02_title_page.docx         - title page with author info (placeholders)
  03_manuscript.docx         - the essay itself (no author info in text)
  Figure1.png / Figure1.tif  - Figure 1 as separate high-resolution files
  figure_captions.txt        - figure caption list (English)

Run after create_essay_en.py and generate_figures.py have produced their
outputs in ../output.
"""

import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

import create_essay_en as essay

OUT = Path(__file__).resolve().parent.parent / "output"
PKG = OUT / "submission_eia_essay"
FONT = "Times New Roman"

TITLE = essay.TITLE
FIG_CAPTION = essay.FIG_CAPTION
WORD_COUNT = essay.word_count()
N_NOTES = len(essay.NOTES)

KEYWORDS = ("nuclear ethics; noncombatant immunity; incommensurability; "
            "just war; biopolitics; the ethics of deterrence")


def _base(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(12)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.paragraph_format.line_spacing = 1.15
    st.paragraph_format.space_after = Pt(6)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(1)
        s.left_margin = s.right_margin = Inches(1)


def _p(doc, text="", *, bold=False, italic=False, size=12, align=None,
       space_after=6, space_before=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
    return p


def cover_letter():
    doc = Document(); _base(doc)
    _p(doc, "[Date]", size=12)
    _p(doc, "To the Editors\nEthics & International Affairs\nCarnegie Council "
       "for Ethics in International Affairs", size=12, space_after=12)
    _p(doc, "Dear Editors,", space_after=12)
    _p(doc,
       "I am pleased to submit the enclosed essay, \u201C" + TITLE + ",\u201D "
       "for consideration in the Essays section of Ethics & International "
       "Affairs. The piece runs to approximately {:,} words of main text with "
       "{} endnotes and one figure.".format(WORD_COUNT, N_NOTES))
    _p(doc,
       "Eighty years after Hiroshima, public argument about the atomic "
       "bombing still turns almost entirely on whether it was necessary\u2014"
       "on whether it saved more lives than it cost. The essay grants that "
       "empirical debate its due and then sets it aside to press a prior "
       "question that the necessity framing quietly assumes away: whether "
       "human beings are the kind of agents who may weigh and decide who is "
       "to live and who is to die at all. Distinguishing this \u2018outer\u2019 "
       "question from the \u2018inner\u2019 calculus, I argue that the "
       "utilitarian justification presupposes a commensurability of lives "
       "that value pluralism gives us reason to deny; that, since finitude "
       "makes some choosing unavoidable, the decisive line falls not on the "
       "outcome of a weighing but on its attitude\u2014what I call the taking "
       "and the giving calculus; and that these commitments connect to a "
       "civilizational choice between increasing supply and reducing demand "
       "that bears directly on how a third world war might be prevented.")
    _p(doc,
       "The argument is normative rather than empirical, and it is pitched at "
       "the general scholarly readership the Essays section addresses. It "
       "engages the just war tradition (Walzer, Anscombe), the critique of "
       "aggregation (Rawls), value pluralism (Berlin, Raz), and biopolitics "
       "(Agamben, Foucault), but wears this scholarship lightly and keeps "
       "citations to a minimum, as the section requires.")
    _p(doc,
       "The manuscript is original, is not under consideration elsewhere, and "
       "has not been published previously. I have no conflicts of interest to "
       "declare. The essay contains no author-identifying information in its "
       "text; author details are given on the separate title page.",
       space_after=12)
    _p(doc, "Thank you for your consideration.", space_after=12)
    _p(doc, "Sincerely,", space_after=2)
    _p(doc, "[Author Name]", space_after=0)
    _p(doc, "[Institutional Affiliation]", space_after=0)
    _p(doc, "[Email address] \u00b7 [ORCID iD]", space_after=0)
    doc.save(PKG / "01_cover_letter.docx")


def title_page():
    doc = Document(); _base(doc)
    _p(doc, TITLE, bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER,
       space_after=4)
    _p(doc, "Submission type: Essay (Ethics & International Affairs)",
       italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    _p(doc, "Author", bold=True, size=12, space_after=2)
    _p(doc, "[Author Name]", space_after=0)
    _p(doc, "[Institutional Affiliation, City, Country]", space_after=0)
    _p(doc, "[ORCID iD]", space_after=10)

    _p(doc, "Corresponding author", bold=True, size=12, space_after=2)
    _p(doc, "[Author Name] \u2014 [email address] \u2014 [postal address]",
       space_after=10)

    _p(doc, "Word count", bold=True, size=12, space_after=2)
    _p(doc, "Main text: approx. {:,} words. Endnotes: {}. Figures: 1."
       .format(WORD_COUNT, N_NOTES), space_after=10)

    _p(doc, "Keywords", bold=True, size=12, space_after=2)
    _p(doc, KEYWORDS, space_after=10)

    _p(doc, "Author biography", bold=True, size=12, space_after=2)
    _p(doc, "[50\u201375 word biographical note: position, institution, "
       "principal research interests, and one or two representative "
       "publications.]", space_after=10)

    _p(doc, "Acknowledgements", bold=True, size=12, space_after=2)
    _p(doc, "[Optional acknowledgements, or \u201cNone.\u201d]", space_after=10)

    _p(doc, "Funding", bold=True, size=12, space_after=2)
    _p(doc, "[Funding sources, or \u201cThis research received no specific "
       "grant from any funding agency.\u201d]", space_after=10)

    _p(doc, "Conflict of interest", bold=True, size=12, space_after=2)
    _p(doc, "The author declares no conflict of interest.", space_after=0)
    doc.save(PKG / "02_title_page.docx")


def readme_checklist():
    doc = Document(); _base(doc)
    _p(doc, "Submission package \u2014 " + TITLE, bold=True, size=14,
       space_after=10)
    _p(doc, "Target: Ethics & International Affairs, \u201CEssays\u201D "
       "category (editor-reviewed / non-refereed).", space_after=10)

    _p(doc, "Enclosed files", bold=True, size=12, space_after=2)
    for line in [
        "00_README_checklist.docx \u2014 this file",
        "01_cover_letter.docx \u2014 cover letter (author placeholders)",
        "02_title_page.docx \u2014 title page with author details "
        "(placeholders)",
        "03_manuscript.docx \u2014 essay text; title, standfirst, body, and "
        "endnotes; no author-identifying information",
        "Figure1.png \u2014 Figure 1, 300 dpi raster",
        "Figure1.tif \u2014 Figure 1, 300 dpi TIFF (print-quality)",
        "figure_captions.txt \u2014 figure caption (English)",
    ]:
        _p(doc, "\u2022 " + line, size=11, space_after=2)

    _p(doc, "Checklist", bold=True, size=12, space_before=8, space_after=2)
    for line in [
        "Main text approx. {:,} words (within the ~2,500\u20133,500 word "
        "range for Essays).".format(WORD_COUNT),
        "{} Chicago-style numbered endnotes; every note is cited in the text "
        "and numbered in order of first appearance.".format(N_NOTES),
        "One figure; it is cited in the text (\u201cFigure 1\u201d) before it "
        "appears and placed immediately after that paragraph, with its "
        "caption below.",
        "Figure supplied both inline (in the manuscript) and as separate "
        "high-resolution PNG and TIFF files.",
        "All cited works are real, published sources; no fabricated "
        "references.",
        "Author information confined to the title page; manuscript text is "
        "free of identifying details.",
    ]:
        _p(doc, "\u2610 " + line, size=11, space_after=2)

    _p(doc, "Notes for the author before submitting", bold=True, size=12,
       space_before=8, space_after=2)
    for line in [
        "Submission is handled through the journal\u2019s online submission "
        "system (a separate account from Cambridge Core); upload each file "
        "type as the system requests.",
        "Open access / APC: Cambridge lists Essays among the article types it "
        "treats as research articles for Gold-OA funding, so an APC may apply. "
        "Authors without APC funding can request a waiver via the journal\u2019s "
        "open-access options after acceptance. Confirm the current fee and "
        "waiver terms with the editorial office.",
        "Reporting checklists (STROBE, CONSORT, PRISMA, etc.) do not apply: "
        "this is a normative philosophy essay with no empirical study.",
        "Replace every [bracketed placeholder] on the title page and cover "
        "letter before submitting.",
        "Verify the current word limit and house style against the journal\u2019s "
        "Instructions for Contributors at the time of submission.",
    ]:
        _p(doc, "\u2022 " + line, size=11, space_after=2)
    doc.save(PKG / "00_README_checklist.docx")


def main():
    if PKG.exists():
        shutil.rmtree(PKG)
    PKG.mkdir(parents=True)

    readme_checklist()
    cover_letter()
    title_page()
    shutil.copy(OUT / "essay_en.docx", PKG / "03_manuscript.docx")
    shutil.copy(OUT / "fig1_layers.png", PKG / "Figure1.png")
    shutil.copy(OUT / "fig1_layers.tif", PKG / "Figure1.tif")
    (PKG / "figure_captions.txt").write_text(FIG_CAPTION + "\n",
                                             encoding="utf-8")

    zip_path = OUT / "submission_eia_essay.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        for f in sorted(PKG.iterdir()):
            z.write(f, arcname=f"submission_eia_essay/{f.name}")
    print("package:", PKG)
    print("zip:", zip_path)


if __name__ == "__main__":
    main()

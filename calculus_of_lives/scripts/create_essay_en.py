#!/usr/bin/env python3
"""
Generate the short English *Essay* version for Ethics & International Affairs
(the non-refereed, timely "Essays" category: ~2,500-3,500 words, few notes).

  "The Accountants of One Another's Lives"

Derived from the full manuscript, condensed for a general scholarly readership.
One explanatory figure (the two layers of the question); Chicago-style
numbered endnotes, kept to a minimum.

Output into ../output:
  essay_en.docx
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT = Path(__file__).resolve().parent.parent / "output"
OUT.mkdir(exist_ok=True)

FONT = "Times New Roman"

TITLE = "The Accountants of One Another's Lives"
STANDFIRST = ("Eighty years after Hiroshima we still argue about whether the "
              "bomb was necessary. The prior question is whether anyone may "
              "weigh such things at all.")

FIG_CAPTION = ("Figure 1. The funnel of attention. Every question the bombing "
               "raises pours in at the top, but a layer of common sense and "
               "preconception\u2014that we are entitled to do the weighing\u2014"
               "acts as a sieve: only the inner, technical question (\u201Cwas "
               "it necessary?\u201D) passes through into public debate, while "
               "the outer question\u2014whether such weighing is ours to perform "
               "at all\u2014is screened out and never comes into view.")

# Each item is ("style", "text"): H = heading, P = paragraph, FIG = image key.
BODY = [
("P",
 "The destruction of Hiroshima and Nagasaki is still defended, eight decades "
 "on, in a familiar idiom. The bombs, for all their horror, were necessary: "
 "they ended the war quickly and thereby spared the far larger number of "
 "lives\u2014American and Japanese, soldier and civilian\u2014that an invasion "
 "of the home islands would have consumed. Generations of historians have "
 "contested the factual premises of this story, asking whether surrender was "
 "already near, whether the casualty projections were inflated after the "
 "fact, and whether the true audience for the blast was Moscow rather than "
 "Tokyo.{1} That empirical quarrel matters. But it is not the argument I want "
 "to have."),

("P",
 "Look instead at the shape of the justification, quite apart from whether "
 "its numbers are right. It is a piece of arithmetic. It concedes that lives "
 "were taken and answers that more lives were saved; it sets the dead of two "
 "cities on one pan of a scale and the hypothetical dead of an invasion on "
 "the other, and reads off the balance. Grant the defender everything he "
 "wants\u2014that the bombs did shorten the war, that they did avert a "
 "bloodier landing. A question remains that the arithmetic cannot put to "
 "itself, because it has already assumed the answer: are human beings the "
 "kind of agents who may set themselves up as the accountants of one "
 "another's lives, deciding who shall live and who shall die in another's "
 "stead?"),

("P",
 "Call this a shift in the layer of the question (Figure 1). The \u201Cwas it "
 "necessary?\u201D debate operates inside the calculus; it audits a weighing "
 "whose legitimacy it never examines. The question I want to press sits "
 "outside the calculus and asks whether we are entitled to perform it at all. "
 "The two layers are independent: one can win every point on the inner "
 "layer\u2014concede that the sums add up\u2014and still have said nothing to "
 "the outer one. Much of what passes for moral debate about the bomb, and "
 "about war in general, is an argument on the inner layer mistaken for an "
 "argument on the outer. And which layer a society instinctively reaches "
 "for tells us what it already takes for granted about the standing of "
 "persons. The same doubling holds for war in any of its forms, including "
 "the quieter ones: economic coercion and cultural domination are also ways "
 "of securing one's own flourishing by drawing down another's, and they too "
 "invite both questions."),

("FIG", "fig1_layers"),

("H", "Why the sums do not add up"),

("P",
 "The justification borrows its structure from utilitarianism, which holds "
 "that the right act produces the greatest balance of good over bad, summed "
 "across everyone affected. The summing is the crucial move, and it is more "
 "questionable than its long familiarity suggests. As John Rawls insisted, "
 "to add one person's loss into another's gain and call the net a "
 "justification is to fail to take seriously the distinction between "
 "persons\u2014to treat a collection of separate lives as though it were a "
 "single great life with a single ledger.{2} There is no super-person who "
 "enjoys the sum; there are only the many, one of whom is being spent for "
 "the others. The invasion's phantom dead do not exist to be saved, and the "
 "citizens of Hiroshima were not compensated by the survival of strangers."),

("P",
 "There is a deeper trouble still. Aggregation assumes that the goods and "
 "losses of different lives can be measured on one scale and converted into "
 "a common currency. But the measures are as many as the persons, and "
 "perhaps more: what a life is worth, what it is for, what its loss "
 "subtracts from the world, is answered differently from within each life, "
 "in terms that do not reduce to a shared unit. Value pluralists have long "
 "argued that goods are genuinely incommensurable\u2014that there is no master "
 "scale on which a mother, a scholar, a child, and a soldier can be "
 "expressed as quantities and traded off.{3} If that is right, the "
 "utilitarian ledger is not merely hard to compute; it is trying to add "
 "quantities that were never in the same units. The balance it reads off is "
 "an artifact of a scale it had to invent. The impossibility is not "
 "practical but principled."),

("P",
 "Moral philosophers have built a whole laboratory around the smaller case: "
 "the runaway trolley that will kill five unless it is diverted onto one, "
 "and the surgeon who could save five patients by harvesting the organs of "
 "one healthy visitor. Most people flinch at the surgeon while accepting the "
 "switch, and the literature has spent half a century refining the "
 "distinctions\u2014doing versus allowing, intending a death versus merely "
 "foreseeing it\u2014that might explain the difference. Those distinctions "
 "are real and they matter. But notice that they, too, are conducted inside "
 "the calculus: they ask which weighings come out permissible. The outer "
 "question is prior. It does not ask whether the surgeon has correctly "
 "computed five against one; it asks what has already happened to a person "
 "when he becomes, in someone else\u2019s reasoning, the one\u2014a quantity "
 "to be spent so that a larger quantity may be kept. The atomic bomb is the "
 "surgeon\u2019s case written at the scale of cities, and the unease most "
 "people feel at the surgeon is the outer question trying to make itself "
 "heard."),

("H", "Choosing under finitude"),

("P",
 "And yet we cannot simply refuse to choose. Resources are finite; the "
 "harvest feeds some and not others; the lifeboat has a number. One might "
 "dream of escaping the predicament by making supply infinite or need "
 "vanishingly small, but neither is available to us: even the universe is "
 "finite, and nothing comes from nothing. So the tragic structure of choice "
 "under scarcity does not disappear, and any honest ethics has to say "
 "something about how to act inside it rather than pretending to abolish it."),

("P",
 "Here is the pivot. If some weighing is unavoidable, the morally decisive "
 "line cannot fall on whether one weighs, but on how. I distinguish two "
 "attitudes that can wear the same numerical face. In the taking calculus, "
 "the other is an eliminable variable, a quantity to be zeroed out when the "
 "total improves; his death is a cost that the ledger, once balanced, treats "
 "as discharged. In the giving calculus, the other is an end, never merely a "
 "means; one weighs because one must, but one bears the weighing as loss, "
 "keeps faith with the one who could not be saved, and refuses to record him "
 "as a rounding error. The distinction is Kant's, between a price\u2014the "
 "value of something for which an equivalent can be substituted\u2014and a "
 "dignity, which admits of no equivalent and forbids the substitution.{4} A "
 "person has a dignity, not a price; to reason about persons as though they "
 "had prices is the original move of the taking calculus."),

("P",
 "The giving calculus is not innocent; it still weighs, and so is not free "
 "of tragedy. But the difference in attitude is not a psychological gloss on "
 "an otherwise identical act. It changes what the act is. It shows itself in "
 "a moral remainder\u2014grief, and the sense of a wrong done even in doing "
 "what one had to\u2014and in a direction: the giving agent absorbs cost onto "
 "himself before displacing it onto others. What was monstrous in the "
 "reasoning that reached for the bomb, on this view, is not that officials "
 "did arithmetic\u2014statesmen under duress must reckon with numbers\u2014but "
 "that they reasoned in the taking mode about people who were never theirs "
 "to price, and did so, moreover, having never once paused at the outer "
 "question. A society that can no longer hear the difference between taking "
 "and giving, so long as the totals match, has lost the thing that "
 "separates provision from predation."),

("H", "Whom one may place on the scale"),

("P",
 "Even granting that some weighing is unavoidable, a further question is "
 "whom the weighing may range over at all. The just war tradition answers "
 "with the principle of noncombatant immunity: those who have not made "
 "themselves into agents of harm may not be made the direct objects of "
 "attack. Michael Walzer grounds the immunity in the idea that one becomes a "
 "legitimate target only by posing a threat, which is why even the munitions "
 "worker is liable at the factory yet regains his immunity at home.{5} I "
 "would reground the line on consent and vocation rather than on function. "
 "The soldier who dies at his post dies within a calling he has assumed; "
 "there is a dignity available to him in the form of professional "
 "honor\u2014the acceptance of mortal risk as part of a role freely "
 "shouldered. The civilian has consented to no such thing. He is a pure "
 "victim, and to place him on the scale is not to demand a sacrifice of him "
 "but to make him bear a cost he never undertook to bear."),

("P",
 "This is also why the hardest case in the just war tradition\u2014the "
 "munitions worker\u2014turns on the same idea from the other side. If the "
 "worker is a legitimate target while building weapons, it is not simply "
 "because he is useful to the war, but because, in taking up that work, he "
 "has stepped partway into the zone of danger and shoulders something of a "
 "combatant\u2019s risk. A society that finds the worker\u2019s liability "
 "obvious but the civilian\u2019s inviolability negotiable has, I suspect, a "
 "weak feeling for vocation and sacrifice\u2014for the idea, captured in the "
 "old language of noblesse oblige, that those who accept danger do so "
 "precisely so that others need not. Where that feeling is strong, the line "
 "protecting the pure victim is not a technicality to be argued around; it "
 "is the visible edge of a whole ethic of who owes danger to whom."),

("P",
 "This is why the atomic bombing is, on the outer view, beyond argument. It "
 "did not aim at soldiers, or even at workers while at their machines; it "
 "aimed at cities\u2014at civilian populations as such, whose deaths were the "
 "instrument by which a government elsewhere was to be moved. Elizabeth "
 "Anscombe saw this clearly when she opposed the honoring of the man who "
 "ordered it: the deliberate killing of the innocent as a means to one's "
 "ends is murder, and no accumulation of good consequences converts murder "
 "into something else.{6} That a society finds this debatable is itself the "
 "symptom worth noticing. It suggests a thin sense of the very ideas\u2014"
 "vocation, sacrifice, the honor of those who accept danger so that others "
 "need not\u2014that mark the civilian off as inviolable in the first place."),

("H", "The sovereign who weighs"),

("P",
 "Who, then, appointed anyone the accountant? Modern political thought has a "
 "name for the power that decides which lives may be placed outside the "
 "circle of those it is a crime to kill. Giorgio Agamben calls it sovereign "
 "power, exercised over what he calls bare life; Michel Foucault describes a "
 "biopolitics that administers populations in the aggregate, deciding whom "
 "to make live and whom to let die.{7} Strategic bombing, the language of "
 "acceptable losses, the actuarial cast of nuclear deterrence\u2014these are "
 "that power in its late form, the taking calculus raised to the scale of "
 "peoples and dressed as necessity. The outer question is, at bottom, a "
 "refusal to grant that power its self-appointment. And it is not a museum "
 "piece. In a world that is rearming, and in which the threat to hold whole "
 "populations hostage is once more spoken aloud, the temptation to reason "
 "about the many as summable quantities is not receding but returning."),

("H", "Two civilizations, and a mismatch"),

("P",
 "Step back to the largest scale. Civilizations meet scarcity in two broadly "
 "different ways. One increases supply\u2014it expands, conquers, converts, "
 "opens new frontiers and draws them into itself. The other reduces "
 "demand\u2014it contracts, defends, disciplines its own wanting, and seeks "
 "sufficiency rather than growth. The two strategies have loose affinities "
 "with religious forms: a tradition of one exclusive truth tends to treat "
 "the world as a frontier still to be brought in, while immanentist and "
 "pluralist traditions, finding the sacred already here, feel less pull "
 "toward expansion. The affinities are only tendencies\u2014ascetic "
 "Protestantism reduced demand and yet, as Weber saw, midwifed a vast "
 "engine of supply\u2014but the two axes, supply-versus-demand and "
 "expansion-versus-contraction, are worth holding apart, because a great "
 "deal of history lives on their off-diagonals."),

("P",
 "It is tempting to think that an expansionist civilization needs only an "
 "open frontier to stay at peace, so that a new frontier\u2014outer space, "
 "say\u2014might drain its pressure and quiet its wars. I think this is half "
 "true, and the half that is false is instructive. The frontier such "
 "traditions crave is really two frontiers of different kinds. The material "
 "frontier is extensive: empty land and unused energy, a matter of volume "
 "that more space can enlarge. The spiritual frontier is intensive: it is "
 "aimed at particular souls and at the neighbor who believes otherwise, a "
 "matter of density, not volume. The two cannot be exchanged: an intensive "
 "demand cannot be discharged into an extensive supply. The pattern is "
 "suggestive rather than probative, but it is at least striking that the "
 "Age of Discovery flung open a vast material frontier and yet coincided "
 "with Europe's most terrible wars of religion. Space enlarges only the "
 "extensive frontier; it leaves the mismatch untouched and may even amplify "
 "the intensive drive to carry one's own truth to the stars. If the heavens "
 "pacify anything, it will not be as an outlet but as a mirror\u2014the view "
 "of the whole earth, undivided, that shrinks the tribal claim itself. And "
 "shrinking the claim is a reduction of demand, not an increase of supply."),

("H", "Deterring the third war"),

("P",
 "This is the practical payoff. If the analysis holds, a third world war "
 "will not be averted by refining the calculus of deterrence\u2014by making "
 "our threats more credible or our targeting more precise. That project "
 "stays on the inner layer and leaves untouched the disposition to treat "
 "lives as summable at all. What is needed is movement along both hands at "
 "once. Technology can raise supply, but it meets a ceiling in the "
 "finitude of the world, and worse, efficiency tends to enlarge appetite "
 "rather than sate it, so that supply-side progress underwrites the very "
 "growth in demand it was meant to relieve. Hence the second hand: a "
 "discipline of demand, the deliberate lowering of what we take ourselves to "
 "need. This is the sane core of what E. F. Schumacher called Buddhist "
 "economics\u2014an economics organized around sufficiency rather than "
 "maximal consumption.{8} At the level of the person it is the resolve to "
 "reduce one's own taking before displacing cost onto others; at the level "
 "of the civilization it is the choice of contraction and defense over "
 "expansion and conquest."),

("P",
 "There is a name for what this makes possible: a second-order kind of "
 "prevention. First-order deterrence tries to make the next war too costly "
 "to start, and so keeps refining the very calculus of hostage populations "
 "that is the problem. Second-order prevention works upstream, on the "
 "appetite that generates the confrontation in the first place, and so aims "
 "to make the calculus less necessary rather than more exact. A civilization "
 "that has learned to want less has fewer occasions to place its neighbors "
 "on the scale, and a smaller stake in the frontiers over which the great "
 "wars are fought. Restraint of demand is not only an ethic; it is also a "
 "security posture, and for the small and exposed nations that have the most "
 "to lose in a war of expansion, it may be the one that reaches deepest "
 "toward the root."),

("P",
 "The deepest prevention, then, is not a better weighing but a disciplined "
 "retreat from the posture of the one who weighs. We cannot reach the "
 "vanishing point at which no one need ever be placed on a scale; finitude "
 "forbids it. But we can decide which direction we face. Eighty years after "
 "Hiroshima, the question worth asking is not whether the bomb added up. It "
 "is whether we still wish to be the kind of beings who keep the books\u2014"
 "or whether we might, at last, prefer to give in order to keep others "
 "alive rather than to take in order to live. That preference will not "
 "abolish tragedy. But it is the ground on which a third world war "
 "might become, not merely deterred, but unthinkable."),
]

NOTES = [
    "The classic public statement of the necessity thesis is Henry L. "
    "Stimson, \u201CThe Decision to Use the Atomic Bomb,\u201D Harper\u2019s "
    "Magazine, February 1947. For the historiographical challenge, see "
    "Gar Alperovitz, The Decision to Use the Atomic Bomb (New York: Knopf, "
    "1995); and J. Samuel Walker, Prompt and Utter Destruction, 3rd ed. "
    "(Chapel Hill: University of North Carolina Press, 2016).",

    "John Rawls, A Theory of Justice, rev. ed. (Cambridge, MA: Harvard "
    "University Press, 1999), 23\u201324; and Rawls, \u201CFifty Years after "
    "Hiroshima\u201D (1995), in Collected Papers, ed. Samuel Freeman "
    "(Cambridge, MA: Harvard University Press, 1999), 565\u2013572.",

    "See Isaiah Berlin, \u201CThe Pursuit of the Ideal,\u201D in The Crooked "
    "Timber of Humanity (London: John Murray, 1990); and Joseph Raz, The "
    "Morality of Freedom (Oxford: Clarendon Press, 1986), chap. 13.",

    "Immanuel Kant, Groundwork of the Metaphysics of Morals, ed. and trans. "
    "Mary Gregor (Cambridge: Cambridge University Press, 1998), 4:434\u2013435.",

    "Michael Walzer, Just and Unjust Wars, 5th ed. (New York: Basic Books, "
    "2015), 138\u2013159.",

    "G. E. M. Anscombe, \u201CMr Truman\u2019s Degree\u201D (1957), reprinted "
    "in her Ethics, Religion and Politics: Collected Philosophical Papers, "
    "vol. 3 (Oxford: Blackwell, 1981), 62\u201371.",

    "Giorgio Agamben, Homo Sacer: Sovereign Power and Bare Life, trans. "
    "Daniel Heller-Roazen (Stanford: Stanford University Press, 1998); "
    "Michel Foucault, The History of Sexuality, vol. 1, trans. Robert Hurley "
    "(New York: Pantheon, 1978), 135\u2013145.",

    "E. F. Schumacher, Small Is Beautiful: Economics as if People Mattered "
    "(London: Blond & Briggs, 1973), chap. 4 (\u201CBuddhist Economics\u201D).",
]


def _set_base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(12)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = st.paragraph_format
    pf.line_spacing = 2.0
    pf.space_after = Pt(0)


def _add_runs_with_markers(par, text):
    for part in re.split(r"(\{\d+\})", text):
        m = re.fullmatch(r"\{(\d+)\}", part)
        if m:
            par.add_run(m.group(1)).font.superscript = True
        else:
            par.add_run(part)


def build():
    doc = Document()
    _set_base_style(doc)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(1)
        s.left_margin = s.right_margin = Inches(1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(4)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(STANDFIRST)
    sr.italic = True
    sr.font.size = Pt(12)
    sp.paragraph_format.space_after = Pt(12)

    for style, text in BODY:
        if style == "H":
            hp = doc.add_paragraph()
            hp.paragraph_format.space_before = Pt(12)
            hp.paragraph_format.space_after = Pt(6)
            rr = hp.add_run(text)
            rr.bold = True
            rr.font.size = Pt(13)
        elif style == "FIG":
            ip = doc.add_paragraph()
            ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ip.paragraph_format.space_before = Pt(12)
            ip.add_run().add_picture(str(OUT / f"{text}.png"), width=Inches(5.4))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_before = Pt(6)
            cap.paragraph_format.space_after = Pt(8)
            cap.paragraph_format.line_spacing = 1.0
            cap.add_run(FIG_CAPTION).font.size = Pt(10)
        else:
            pp = doc.add_paragraph()
            pp.paragraph_format.first_line_indent = Inches(0.4)
            _add_runs_with_markers(pp, text)

    hp = doc.add_paragraph()
    hp.paragraph_format.space_before = Pt(18)
    hp.add_run("Notes").bold = True
    for i, note in enumerate(NOTES, start=1):
        np_ = doc.add_paragraph()
        np_.paragraph_format.line_spacing = 1.5
        np_.paragraph_format.space_after = Pt(2)
        np_.add_run(f"{i}. ").bold = True
        np_.add_run(note).font.size = Pt(10.5)

    doc.save(OUT / "essay_en.docx")


def word_count():
    return sum(len(re.sub(r"\{\d+\}", "", t).split())
               for s, t in BODY if s in ("P", "H"))


if __name__ == "__main__":
    build()
    # marker/notes consistency check
    marks = []
    for s, t in BODY:
        marks += [int(m) for m in re.findall(r"\{(\d+)\}", t)]
    assert marks == sorted(marks), "citation markers not in order"
    assert sorted(set(marks)) == list(range(1, len(NOTES) + 1)), \
        "marker/notes mismatch"
    print("essay_en.docx written to", OUT,
          "| notes:", len(NOTES), "| approx words:", word_count())

#!/usr/bin/env python3
"""
Assemble the complete Journal of Military Ethics (JME) submission package into
../output/submission_jme/ and zip it.

JME is a peer-reviewed Taylor & Francis journal. This package follows standard
T&F double-anonymous submission practice: an anonymized manuscript, a separate
title page, a cover letter, and figure files supplied both inline and as
separate high-resolution rasters.

Contents:
  00_README_checklist.docx   - what is enclosed + submission notes
  01_cover_letter.docx       - cover letter to the Editors
  02_title_page.docx         - title page with author details (non-anonymous)
  03_manuscript.docx         - the manuscript (abstract, keywords, body,
                               figures inline, endnotes); no author info in text
  Figure1/2/3.png            - figures as separate 300 dpi rasters
  Figure1/2/3.tif            - figures as separate 300 dpi TIFFs
  figures_editable.pptx      - all three figures, one per slide (editable)
  figure_captions.txt        - figure captions (English)

Run after generate_figures.py, create_manuscript_en.py and
create_figures_pptx.py have produced their outputs in ../output.
"""

import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

import create_manuscript_en as ms

OUT = Path(__file__).resolve().parent.parent / "output"
PKG = OUT / "submission_jme"
FONT = "Times New Roman"

TITLE = ms.TITLE
WORD_COUNT = ms.word_count()
N_NOTES = len(ms.NOTES)
N_FIGURES = 3
FIG_CAPTIONS = ms.FIG_CAPTIONS


def _base(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(12)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.paragraph_format.line_spacing = 1.15
    st.paragraph_format.space_after = Pt(6)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(1)
        s.left_margin = s.right_margin = Inches(1)


def _p(doc, text="", *, bold=False, italic=False, size=12, align=None,
       space_after=6, space_before=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
    return p


def cover_letter():
    doc = Document(); _base(doc)
    _p(doc, "[Date]")
    _p(doc, "To the Editors\nJournal of Military Ethics\nTaylor & Francis",
       space_after=12)
    _p(doc, "Dear Editors,", space_after=12)
    _p(doc,
       "I am pleased to submit the enclosed manuscript, \u201C" + TITLE +
       "\u201D for consideration as an original article in the Journal of "
       "Military Ethics. The main text runs to approximately {:,} words ("
       "excluding the abstract, notes and references) and includes {} figures. "
       "An abstract and a list of keywords are included on the title page."
       .format(WORD_COUNT, N_FIGURES))
    _p(doc,
       "The article examines the normative question underlying debates over "
       "the atomic bombing of Japan: not whether the bombing was militarily "
       "or politically necessary, but whether human beings are entitled to "
       "weigh and decide who is to live and who is to die at all. I argue "
       "that the familiar utilitarian justification presupposes a "
       "commensurability of lives that value pluralism gives us reason to "
       "deny; that, because finitude makes some choosing unavoidable, the "
       "morally decisive line lies in the attitude of the agent rather than "
       "the outcome of the weighing; that the scope of any defensible "
       "weighing is bounded by the distinction between those who have taken "
       "danger into a vocation and those who remain its pure patients; and "
       "that the civilizational drift toward treating lives as summable "
       "quantities is bound up with a choice between increasing supply and "
       "reducing demand. The article concludes that preventing a third world "
       "war requires not a better calculus but a disciplined retreat from it.")
    _p(doc,
       "The manuscript engages the just war and military ethics tradition "
       "(Walzer, Anscombe), the critique of aggregation (Rawls, Williams, "
       "Nagel), value pluralism (Berlin, Raz), Kantian dignity, biopolitics "
       "(Agamben, Foucault), and the economics of sufficiency (Schumacher), "
       "and it addresses the principal objections\u2014quietism, the "
       "paralysis of policy, supreme emergency, and the realist reply\u2014"
       "in a dedicated section. I believe it will suit the journal's "
       "readership of philosophers, political theorists, military ethicists, "
       "and security-studies scholars.")
    _p(doc,
       "The manuscript is original, is not under consideration elsewhere, and "
       "has not been published previously. It has been prepared for "
       "double-anonymous review: the manuscript file carries no "
       "author-identifying information, and author details appear only on the "
       "separate title page. I have no conflicts of interest to declare and "
       "received no funding for this work. There are no empirical datasets "
       "associated with this normative argument; all three figures are "
       "conceptual and are regenerated from code in the public repository.",
       space_after=12)
    _p(doc, "Thank you for your consideration.", space_after=12)
    _p(doc, "Sincerely,", space_after=2)
    _p(doc, "Tatsuki Onishi", space_after=0)
    _p(doc, "Independent Researcher", space_after=0)
    _p(doc, "bougtoir@gmail.com", space_after=0)
    doc.save(PKG / "01_cover_letter.docx")


def title_page():
    doc = Document(); _base(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(15)
    p.paragraph_format.space_after = Pt(18)

    for label, val in [
        ("Author", "Tatsuki Onishi"),
        ("Affiliation", "Independent Researcher"),
        ("Email", "bougtoir@gmail.com"),
    ]:
        ap = doc.add_paragraph()
        ap.paragraph_format.line_spacing = 1.5
        ap.add_run(f"{label}: ").bold = True
        ap.add_run(val)

    _p(doc, "Article type: Original Article")
    _p(doc, f"Main-text word count (excluding abstract, notes and references): approximately {WORD_COUNT:,}")
    _p(doc, f"Number of notes/endnotes: {N_NOTES}")
    _p(doc, f"Number of figures: {N_FIGURES}")
    _p(doc, "Number of tables: 0")
    _p(doc, "Conflicts of interest: none declared.")
    _p(doc, "Funding: none.")
    _p(doc,
       "Data availability: This is a work of normative philosophy and does "
       "not draw on empirical datasets. All figures are conceptual and are "
       "generated from the public code repository.")

    ap = doc.add_paragraph()
    ap.paragraph_format.space_before = Pt(12)
    ar = ap.add_run("Abstract")
    ar.bold = True
    ar.font.size = Pt(13)
    for para_text in [ms.ABSTRACT]:
        aap = doc.add_paragraph(para_text)
        aap.paragraph_format.line_spacing = 1.5

    kp = doc.add_paragraph()
    kp.paragraph_format.space_before = Pt(6)
    kr = kp.add_run("Keywords: ")
    kr.bold = True
    kp.add_run(ms.KEYWORDS)
    kp.paragraph_format.line_spacing = 1.5

    doc.save(PKG / "02_title_page.docx")


def readme_checklist():
    doc = Document(); _base(doc)
    _p(doc, "Submission package \u2014 " + TITLE, bold=True, size=14,
       space_after=8)
    _p(doc, "Target: Journal of Military Ethics (Taylor & Francis), original "
       "article / double-anonymous peer review.", space_after=10)

    _p(doc, "Enclosed files", bold=True, size=12, space_after=2)
    for line in [
        "00_README_checklist.docx \u2014 this file",
        "01_cover_letter.docx \u2014 cover letter to the Editors",
        "02_title_page.docx \u2014 title page with author details",
        "03_manuscript.docx \u2014 abstract, keywords, body (figures inline), "
        "and endnotes; no author-identifying information",
        "Figure1.png / Figure2.png / Figure3.png \u2014 300 dpi rasters",
        "Figure1.tif / Figure2.tif / Figure3.tif \u2014 300 dpi TIFFs",
        "figures_editable.pptx \u2014 all three figures, one per slide, "
        "editable",
        "figure_captions.txt \u2014 figure captions",
    ]:
        _p(doc, "\u2022 " + line, size=11, space_after=2)

    _p(doc, "Checklist", bold=True, size=12, space_before=8, space_after=2)
    for line in [
        "Main text approx. {:,} words.".format(WORD_COUNT),
        "Abstract of 164 words; keywords supplied.",
        "{} Chicago-style numbered endnotes; every note is cited in the text "
        "and numbered in order of first appearance.".format(N_NOTES),
        "{} figures, each cited in the text (\u2018Figure 1/2/3\u2019) before "
        "it appears and placed immediately after that paragraph, with a "
        "caption below.".format(N_FIGURES),
        "Figures supplied inline in the manuscript and as separate 300 dpi "
        "PNG/TIFF files, plus an editable PPTX.",
        "All cited works are real, published sources; no fabricated "
        "references.",
        "All figures are regenerated from code (scripts/generate_figures.py); "
        "the manuscript contains no hardcoded data values.",
        "Manuscript file is free of author-identifying information; author "
        "details are on the separate title page.",
    ]:
        _p(doc, "\u2610 " + line, size=11, space_after=2)

    _p(doc, "Notes for the author before submitting", bold=True, size=12,
       space_before=8, space_after=2)
    for line in [
        "Submit via the journal\u2019s ScholarOne site "
        "(https://mc.manuscriptcentral.com/smil20).",
        "Upload the title page separately, or enter author details in the "
        "submission system, and upload the anonymized 03_manuscript.docx as "
        "the main document.",
        "Upload each figure separately in ScholarOne; use the Figure captions "
        "file as a guide and paste captions into the relevant fields.",
        "Reproducibility: this is a normative philosophy article with no "
        "empirical dataset. All three figures are conceptual and "
        "regenerated by scripts/generate_figures.py in the public repository.",
        "Open access / APC: Journal of Military Ethics is a hybrid journal; "
        "the standard route is subscription publication with no author "
        "charge. If you choose open access, confirm the current APC with the "
        "publisher before accepting.",
        "Replace [Date] on the cover letter before submitting.",
        "Verify the current word limit, reference style and figure format "
        "against the journal\u2019s Instructions for Authors at the time of "
        "submission.",
    ]:
        _p(doc, "\u2022 " + line, size=11, space_after=2)
    doc.save(PKG / "00_README_checklist.docx")


def main():
    if PKG.exists():
        shutil.rmtree(PKG)
    PKG.mkdir(parents=True)

    readme_checklist()
    cover_letter()
    title_page()
    shutil.copy(OUT / "manuscript_en.docx", PKG / "03_manuscript.docx")
    shutil.copy(OUT / "figures_en.pptx", PKG / "figures_editable.pptx")

    figmap = {1: "fig1_layers", 2: "fig2_quadrant", 3: "fig3_asymptote"}
    for n, stem in figmap.items():
        shutil.copy(OUT / f"{stem}.png", PKG / f"Figure{n}.png")
        shutil.copy(OUT / f"{stem}.tif", PKG / f"Figure{n}.tif")

    caps = "\n\n".join(FIG_CAPTIONS[n] for n in (1, 2, 3))
    (PKG / "figure_captions.txt").write_text(caps + "\n", encoding="utf-8")

    zip_path = OUT / "submission_jme.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        for f in sorted(PKG.iterdir()):
            z.write(f, arcname=f"submission_jme/{f.name}")
    print("package:", PKG)
    print("zip:", zip_path)


if __name__ == "__main__":
    main()

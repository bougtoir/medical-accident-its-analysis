#!/usr/bin/env python3
"""Build a Behavioural Public Policy (Cambridge Core) submission package for the
rate-based analysis of malpractice-litigation risk and Japanese specialty-level
physician workforce allocation.

Outputs (all derived from results/reanalysis_results.json and data_primary/):
  - manuscript/bpp_manuscript_en.docx   anonymised main manuscript
  - manuscript/bpp_title_page.docx      title page with author info
  - manuscript/bpp_cover_letter.docx    cover letter addressed to BPP
  - manuscript/bpp_highlights.docx      3-5 short highlights (optional)
  - manuscript/bpp_supplementary.docx     supplementary figures & tables
  - manuscript/bpp_figures.pptx           editable main figure slides
  - manuscript/bpp_supplementary_figures.pptx editable supplementary slides
  - output/ha_Figure_*.png              main figure files (reused)
  - output/ha_Supplementary_Figure_*.png supplementary figure files (reused)
  - output/bpp_submission.zip            bundled submission package

The main manuscript is anonymised and uses the same reproducible data pipeline
as the Healthcare Analytics submission; only the framing, title, abstract,
introduction, discussion, cover letter and declarations are re-written for
Behavioural Public Policy.
"""
import os
import shutil
import zipfile
import build_healthcare_analytics_submission as ha
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pptx import Presentation
from pptx.util import Inches as PInches, Pt as PPt

BASE = ha.BASE
OUT = ha.OUT
PUBLIC_REPO = ha.PUBLIC_REPO

# Local aliases for readability
N = ha.N_SPECIALTIES
BIEN = ha.BIEN
YEARS = ha.YEARS
GREW = ha.GREW
FELL = ha.FELL
SPAN = ha.SPAN
SURG_DESC = ha.SURG_DESC
SURG_PCT = ha.SURG_PCT
PHYS = ha.PHYS
HOSP = ha.HOSP
REV = ha.REV
CNT = ha.CNT
ANN = ha.ANN
INT = ha.INT
EQP = ha.EQP
EQH = ha.EQH
BS_PHYS = ha.BS_PHYS
BS_HOSP = ha.BS_HOSP
JMSR = ha.JMSR
MEDIA = ha.MEDIA
JOCS_HOLM = ha.JOCS_HOLM
MARGIN1 = ha.MARGIN1
MARGIN2 = ha.MARGIN2
PER = ha.PER
JMSR_START = ha.JMSR_START
MEDIA_START = ha.MEDIA_START
MEDIA_END = ha.MEDIA_END
SIM = ha.SIM
TOTAL_SIM = ha.TOTAL_SIM
SURG_SIM = ha.SURG_SIM
SP = ha.SP
n_pos = sum(1 for v in SP.values() if v["rho"] > 0)
n_sig = sum(1 for v in SP.values() if v["p"] < 0.05)
DESCR = ha.DESCR

BPP_TITLE = (
    "Perceived malpractice risk and real workforce allocation: "
    "a behavioural-economics analysis of litigation risk and physician "
    "specialty supply in Japan"
)


def h(doc, text, level=1):
    """Unnumbered BPP-style heading."""
    return ha.head(doc, text, level=level, numbered=False)


def b(doc, text, **kw):
    """Body paragraph that contributes to the main word count."""
    return ha.body(doc, text, **kw)


def p(doc, text, **kw):
    """Paragraph that does not contribute to the main word count."""
    return ha.para(doc, text, **kw)


def f(doc, fn, caption, width=Inches(5.8)):
    return ha.figure(doc, fn, caption, width=width)


def t(doc, headers, rows, caption):
    return ha.table(doc, headers, rows, caption)


def m(doc, latex, inline=False, para=None):
    return ha.add_math(doc, latex, inline=inline, para=para)


def build_manuscript():
    # Reset citation order and body-text accumulator shared with ha module
    ha._CITE_ORDER.clear()
    ha.BODY_TEXTS.clear()

    doc = ha._setup_doc()

    # Title
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(18)
    rt = tp.add_run(BPP_TITLE)
    rt.bold = True
    rt.font.size = Pt(14)
    rt.font.name = "Times New Roman"

    # Abstract (unstructured, one paragraph)
    abstract_text = (
        f"Specialty maldistribution is a persistent healthcare workforce problem. "
        f"A common behavioural assumption is that fear of malpractice litigation pushes physicians "
        f"away from high-risk specialties, but this belief may rest more on the availability of "
        f"salient adverse events and loss aversion than on actual career decisions. "
        f"We used national administrative data for {N} clinical specialties in "
        f"Japan ({BIEN[0]}\u2013{BIEN[-1]}) to test whether litigation risk is associated with "
        f"subsequent physician or hospital supply. Applying a transparent, reproducible analytical "
        f"framework that removes size confounding and sparse-panel bias, we measured exposure as "
        f"closed malpractice claims per {PER:,} physicians and regressed biennial "
        f"log-changes in physicians and hospitals on the lagged litigation rate in a panel with "
        f"specialty and wave fixed effects, cluster-robust standard errors with a small-cluster "
        f"correction, a cluster block-bootstrap, two one-sided equivalence (TOST) tests, and power "
        f"diagnostics. The workforce grew in {GREW} of {N} specialties; {SURG_DESC}, was the "
        f"exception. The lagged litigation rate was unrelated to physician growth (coefficient "
        f"{ha.fmt(PHYS['coef'], 4)}; 95% CI {ha.fmt(PHYS['ci_low'], 4)} to "
        f"{ha.fmt(PHYS['ci_high'], 4)}; p={PHYS['p']:.2f}) or hospital growth "
        f"(p={HOSP['p']:.2f}). A one-SD higher rate shifted physician growth by less than "
        f"\u00b1{MARGIN1}% (TOST p={ha.p_tost_fmt(EQP['tests'][0]['p_tost'])}). The results "
        f"suggest that, despite high perceived risk, structural incentives\u2014training costs, "
        f"fee-for-service income, and status-quo bias\u2014keep physicians in high-risk fields. "
        f"Reducing civil litigation exposure is therefore unlikely to be an effective behavioural "
        f"public policy lever for correcting specialty maldistribution; structural incentives such "
        f"as no-fault compensation and payment design are more promising."
    )
    abstract_wc = ha.wc(abstract_text)

    h(doc, "Abstract", level=1)
    ap = doc.add_paragraph()
    ar = ap.add_run(abstract_text)
    ar.font.name = "Times New Roman"
    ar.font.size = Pt(12)
    ap.paragraph_format.line_spacing = 2.0
    ap.paragraph_format.space_after = Pt(6)

    kw = doc.add_paragraph()
    kr = kw.add_run("Keywords: ")
    kr.bold = True
    kr.font.name = "Times New Roman"
    ha._add_runs(
        kw,
        "behavioural public policy; malpractice litigation; physician workforce; "
        "risk perception; equivalence testing; structural incentives",
    )
    kw.paragraph_format.space_after = Pt(18)

    # Introduction
    h(doc, "Introduction", level=1)
    b(
        doc,
        "Specialty maldistribution is a persistent healthcare workforce problem: high-acuity "
        "fields such as surgery, obstetrics and gynaecology, paediatrics and emergency care are "
        "widely perceived as understaffed across many health systems.{maldist} A recurring policy "
        "intuition is that the fear of malpractice litigation and medical safety incidents pushes "
        "physicians away from these high-risk specialties.{malprac,defmed} If this behavioural "
        "assumption were true, reducing litigation exposure would be a lever for correcting "
        "workforce maldistribution.",
    )
    b(
        doc,
        "Behavioural economics suggests, however, that the link between perceived risk and actual "
        "behaviour is not automatic. Rare, salient adverse events are highly available in memory "
        "and in media coverage, and loss aversion causes a low-probability outcome to be "
        "overweighted in career deliberations.{tversky1973,kahneman1979} At the same time, the "
        "decision to enter or leave a specialty is constrained by expected income, sunk training "
        "costs, switching costs and status-quo bias.{samuelson1988} The relevant policy question is "
        "therefore not whether physicians report anxiety about litigation, but whether that anxiety "
        "translates into aggregate workforce shifts. We treat the litigation-workforce question as a "
        "test case for evaluating a behavioural public policy lever: can a salient, emotionally "
        "available risk be used to change the specialty distribution of physicians?",
    )
    b(
        doc,
        "Japan provides a well-documented national setting in which to examine this question. "
        "The Supreme Court reports closed malpractice claims by specialty, the Ministry of Health, "
        "Labour and Welfare publishes biennial physician counts and annual hospital counts, and the "
        "country shares the fee-for-service pressures seen in other high-income countries.{court,phys,facil} "
        "It also introduced the Japan Obstetric Compensation System for Cerebral Palsy (JOCS-CP) "
        "in 2009, a no-fault scheme intended partly to address obstetric workforce concerns.{jocscp,hasegawa2016} "
        "Surveys of residents in Japan, Korea and Taiwan report litigation as a negative factor, but "
        "workload, lifestyle and professional interest dominate specialty choice.{lin2022}",
    )
    b(
        doc,
        "Prior empirical work in this area often related raw annual counts of incidents or lawsuits "
        "to raw counts of physicians or facilities. Two features make such designs prone to "
        "spurious association. First, counts are not adjusted for specialty size: a larger "
        "specialty mechanically accumulates more procedures, more claims and more physicians, so "
        "count-based associations can arise without any behavioural mechanism. Second, the "
        "physician census is collected only biennially; interpolating it to an annual series and "
        "analysing it as if each year were an independent observation inflates the degrees of "
        "freedom of any lag-based method. These pitfalls are not unique to malpractice research; "
        "they arise whenever administrative counts are used to infer behavioural responses in "
        "healthcare organisations. We address them with a transparent analytical framework that uses "
        "rates rather than counts, uses only measured biennial physician observations, and applies "
        "pre-specified equivalence testing, which can provide positive evidence for the absence of "
        "a meaningful effect rather than merely failing to reject a null.{lakens,schuir}",
    )

    # Methods
    h(doc, "Methods", level=1)
    h(doc, "Data sources", level=2)
    b(
        doc,
        f"We report this observational study following the Strengthening the Reporting of "
        f"Observational Studies in Epidemiology (STROBE) guidance.{{strobe}} We studied {N} core "
        f"clinical specialties in Japan for which the Supreme Court reports specialty-specific "
        f"litigation. Three official primary series drove the main analysis: physician counts by "
        f"specialty from the biennial Statistics of Physicians, Dentists and Pharmacists{{phys}}; "
        f"closed malpractice claims by specialty from the Supreme Court of Japan{{court}}; and hospital "
        f"counts by specialty from the annual Survey of Medical Institutions.{{facil}} Two sensitivity "
        f"series were also used: annual medical accident investigation reports by specialty from the "
        f"Japan Medical Safety Research Organisation (JMSR, 2015-2024){{jmsr_data}} and total national "
        f"newspaper article counts from Nikkei Telecom 21 (2004-2018; the sensitivity analysis uses "
        f"{MEDIA_START}-{MEDIA_END}; keywords: medical error + medical malpractice).{{nikkei}} The full "
        f"extraction pipeline (with source identifiers and SHA-256 checksums) is documented in the "
        f"accompanying repository ({PUBLIC_REPO}).",
    )
    b(
        doc,
        "Physician counts use the principal-specialty classification; broad categories were matched "
        "to the Supreme Court's specialty labels, and subspecialties were aggregated in code. "
        "Because the Court assigns multi-specialty cases to a single principal specialty and states "
        "that the counts do not represent the intrinsic risk of each specialty, we treat litigation "
        "as an exposure signal rather than a measure of incident risk.{court} We distinguish "
        "litigation from the Medical Accident Investigation System, which began in 2015 and covers "
        "only deaths and stillbirths judged unforeseen by the hospital administrator; it is not a "
        "general incident-reporting system and is not used as an exposure here.{mais} Primary data "
        "sources and their resolution are summarised in Supplementary Table 1.",
    )
    b(
        doc,
        "The 12 specialties are not arbitrary selections but represent the core clinical categories "
        "in which Japanese physicians typically obtain initial board certification. Japan's specialist "
        "training programme is a two-tiered system: physicians first complete residency and gain "
        "certification in one of the core specialties (e.g., internal medicine, surgery, obstetrics and "
        "gynaecology, paediatrics, anaesthesiology), after which some proceed to a narrower subspecialty "
        "within that primary field. Biennial physician counts and Supreme Court litigation statistics are "
        "both reported at this primary-specialty level, so our analysis captures the broad workforce-allocation "
        "decision at the first tier of the training pipeline. Subspecialties are not separately tracked in "
        "the official biennial series, so we analyse the 12 primary fields as the relevant units of "
        "specialty choice and policy intervention.",
    )

    h(doc, "Statistical analysis", level=2)
    b(
        doc,
        f"We formalised the evaluation as a sensitivity-analysis framework that varies the exposure "
        f"definition (counts versus rates), panel frequency (measured biennial waves versus interpolated "
        f"annual values), and potential confounders (JMSR reports, media coverage, JOCS-CP period) while "
        f"holding the specialty-level panel structure constant. The exposure was the litigation rate, "
        f"defined as closed claims per {PER:,} physicians in each specialty-year, which removes "
        f"specialty-size confounding because large specialties generate more claims for reasons unrelated "
        f"to per-physician risk. The primary analysis used the {len(BIEN)} measured biennial physician "
        f"waves ({BIEN[0]}\u2013{BIEN[-1]}). For each specialty we computed the biennial log-change in "
        f"physicians (and, separately, in hospitals) and regressed it on the litigation rate at the start "
        f"of the interval, in a panel with specialty and wave fixed effects and standard errors clustered "
        f"by specialty.{{angrist}} Clusters are defined by specialty, so G={N} and the small-cluster "
        f"correction uses a t-distribution with G-1 degrees of freedom for all cluster-robust inference. "
        f"Supplementary Figure 1 summarises the sensitivity-analysis framework. The primary estimating "
        f"equation, for specialty s and wave t, was as follows.",
    )
    m(doc, r"\Delta \log(Y_{st}) = \alpha_s + \delta_t + \beta \cdot \text{litrate}_{s,t-1} + \epsilon_{st}")
    b(
        doc,
        "Here, Y is either physicians or hospitals, alpha_s are specialty fixed effects, "
        "delta_t are wave fixed effects, and standard errors are clustered by specialty. "
        "For the equivalence analysis we standardised litrate to a z-score, so beta gives the "
        "expected biennial log-change per one-SD increase in the litigation rate.",
    )
    b(
        doc,
        "We assessed equivalence to a null effect using two one-sided tests (TOST).{lakens,schuir} "
        f"For a pre-specified margin m, the two one-sided null hypotheses are",
    )
    m(doc, r"H_0: \beta \leq -m \quad \text{and} \quad H_0: \beta \geq +m")
    b(
        doc,
        f"Equivalence is declared when both one-sided tests yield p < alpha. We used margins of "
        f"{MARGIN1}% and {MARGIN2}% biennial workforce change because they are smaller than typical policy "
        f"targets for specialty workforce rebalancing and represent changes that workforce planners would "
        f"consider substantively small.",
    )
    b(
        doc,
        "We complemented the analytical small-cluster inference with two diagnostic quantities. "
        "First, a cluster block-bootstrap with B = 1,999 replications resampled specialties with "
        "replacement, re-fitted the primary model, and reported percentile bootstrap 95% confidence "
        "intervals and a bootstrap p-value based on the distribution of absolute t-statistics. "
        "Second, we report the minimum detectable effect (MDE) at 80% power for the per-SD "
        "litigation-rate coefficient and the power to declare equivalence when the true effect is "
        "zero. The MDE is (t(0.975, df) + t(0.80, df)) \u00d7 SE; the equivalence power is "
        "2\u00b7F_t(m/SE) \u2212 1, where F_t is the cumulative distribution of the t(df) distribution. "
        "These diagnostics make the limited-panel information explicit to policymakers evaluating "
        "this behavioural lever.",
    )

    h(doc, "Policy lever simulation", level=2)
    b(
        doc,
        "To translate the regression results into a practical policy benchmark, we projected "
        "physician counts to 2034 under three stylised policy levers. The baseline used each "
        "specialty's observed mean biennial log-growth from 2004 to 2024. The two litigation-reduction "
        "levers set the litigation rate to zero and applied the point estimate and the 95% lower-bound "
        "coefficient, respectively, so they show both the central projection and the most favourable "
        "effect consistent with the data. The third lever added the minimum detectable per-SD effect "
        f"({EQP['mde_80pct']:.2f}% per biennium) to baseline growth as a benchmark for the smallest "
        "policy effect this panel could detect with 80% power. The projections are deterministic "
        "counterfactuals, not forecasts, and are reported as the marginal percent change in 2034 "
        "relative to the baseline drift.",
    )
    b(
        doc,
        f"All reported confidence intervals and two-sided p-values use a t-distribution with G-1 = "
        f"{PHYS['df']} degrees of freedom, the small-cluster correction recommended by Cameron and "
        f"Miller.{{cameron2015}} An indicator for obstetrics and gynaecology from 2009 onward captured the "
        f"JOCS-CP period.{{jocscp}} Sensitivity analyses repeated the models on (i) the annual hospital "
        f"series, (ii) a linearly interpolated annual physician series (with degrees of freedom governed "
        f"by the measured waves, not the interpolated n), (iii) raw counts instead of rates, (iv) the annual "
        f"hospital series {JMSR_START}-2024 additionally controlling for the JMSR report rate, and (v) the "
        f"annual hospital series {MEDIA_START}\u2013{MEDIA_END} additionally controlling for total Nikkei Telecom "
        f"article counts. Because the article-count series is a national yearly variable, it is collinear with "
        f"full wave fixed effects; this sensitivity therefore uses specialty fixed effects plus a linear time "
        f"trend rather than wave dummies. The JOCS-CP indicator and all sensitivity models are exploratory; "
        f"we report raw p-values and Holm step-down adjusted p-values for this family. Analyses used Python "
        f"(statsmodels); code and data are openly available at {PUBLIC_REPO}.",
    )

    # Results
    h(doc, "Results", level=1)
    h(doc, "Workforce and litigation trends", level=2)
    b(
        doc,
        f"Litigation rates per {PER:,} physicians varied several-fold across specialties and fell over time "
        f"in {FELL} of {N} fields (Supplementary Figure 2). Over the same period the physician workforce "
        f"grew in {GREW} of {N} specialties (Supplementary Figure 3; Table 1); the only exception was "
        f"general surgery, which was essentially flat ({SURG_PCT:+.1f}% across {SPAN} years). Exposure and "
        f"workforce therefore did not move in opposite directions as a flight-from-risk account would predict.",
    )
    rows = []
    for s in ha.CORE:
        v = DESCR[ha.EN[s]]
        rows.append(
            [
                ha.EN[s],
                v["phys_first"],
                v["phys_last"],
                f"{v['litrate_first']:.2f}",
                f"{v['litrate_last']:.2f}",
                v["hosp_first"],
                v["hosp_last"],
            ]
        )
    t(
        doc,
        [
            "Specialty",
            f"Physicians {BIEN[0]}",
            f"Physicians {BIEN[-1]}",
            f"Lit. rate {BIEN[0]}",
            f"Lit. rate {BIEN[-1]}",
            f"Hospitals {BIEN[0]}",
            f"Hospitals {BIEN[-1]}",
        ],
        rows,
        f"Table 1. Physicians, litigation rate (per {PER:,} physicians) and hospitals by specialty, first and last waves.",
    )

    h(doc, "Primary association and equivalence", level=2)
    b(
        doc,
        f"The lagged litigation rate was not associated with biennial physician growth "
        f"(coefficient {ha.fmt(PHYS['coef'], 4)}; 95% CI {ha.fmt(PHYS['ci_low'], 4)} to "
        f"{ha.fmt(PHYS['ci_high'], 4)}; p={PHYS['p']:.2f}; n={PHYS['n_obs']}) or with hospital growth "
        f"(coefficient {ha.fmt(HOSP['coef'], 4)}; p={HOSP['p']:.2f}). Equivalence testing (Figure 1; Table 2) "
        f"showed that a 1-SD higher litigation rate changed biennial physician growth by less than "
        f"\u00b1{MARGIN1}% (TOST p={ha.p_tost_fmt(EQP['tests'][0]['p_tost'])}; point estimate "
        f"{EQP['coef_per_SD']*100:+.2f}% with 90% CI {EQP['ci90_low']*100:+.2f}% to "
        f"{EQP['ci90_high']*100:+.2f}%). For hospital growth the point estimate was "
        f"{EQH['coef_per_SD']*100:+.2f}% (90% CI {EQH['ci90_low']*100:+.2f}% to "
        f"{EQH['ci90_high']*100:+.2f}%): it was within the \u00b1{MARGIN2}% margin "
        f"(p={ha.p_tost_fmt(EQH['tests'][1]['p_tost'])}) but not the stricter \u00b1{MARGIN1}% margin "
        f"(p={ha.p_tost_fmt(EQH['tests'][0]['p_tost'])}). Thus the data are consistent with the absence "
        f"of a policy-relevant effect on physician growth, and with at most a small effect on hospital growth. "
        f"Detailed TOST results by margin are reported in Supplementary Table 2.",
    )
    f(
        doc,
        "ha_Figure_1.png",
        f"Figure 1. Equivalence (TOST) of the litigation-rate effect against \u00b1{MARGIN1}% and "
        f"\u00b1{MARGIN2}% margins; horizontal bars are 90% confidence intervals.",
    )
    trow = [
        [
            "Physician growth ~ lagged rate",
            ha.fmt(PHYS["coef"], 4),
            f"{ha.fmt(PHYS['ci_low'], 4)}, {ha.fmt(PHYS['ci_high'], 4)}",
            f"{PHYS['p']:.2f}",
            PHYS["n_obs"],
        ],
        [
            "Hospital growth ~ lagged rate",
            ha.fmt(HOSP["coef"], 4),
            f"{ha.fmt(HOSP['ci_low'], 4)}, {ha.fmt(HOSP['ci_high'], 4)}",
            f"{HOSP['p']:.2f}",
            HOSP["n_obs"],
        ],
        [
            "Counts contrast (physician)",
            ha.fmt(CNT["coef"], 4),
            "\u2014",
            f"{CNT['p']:.2f}",
            CNT["n_obs"],
        ],
        [
            "Annual hospital (sensitivity)",
            ha.fmt(ANN["coef"], 4),
            "\u2014",
            f"{ANN['p']:.2f}",
            ANN["n_obs"],
        ],
        [
            "Interpolated physician (sensitivity)",
            ha.fmt(INT["coef"], 4),
            "\u2014",
            f"{INT['p']:.2f}",
            INT["n_obs"],
        ],
        [
            "Reverse (workforce\u2192litigation)",
            ha.fmt(REV["coef"], 3),
            "\u2014",
            f"{REV['p']:.2f}",
            REV["n_obs"],
        ],
    ]
    t(
        doc,
        ["Model", "Coefficient", "95% CI", "p", "n"],
        trow,
        "Table 2. Panel fixed-effects models and sensitivity analyses.",
    )

    h(doc, "Small-cluster robustness and power", level=2)
    b(
        doc,
        f"Because inference is based on only {N} specialty clusters, we checked the primary results with "
        f"a cluster block-bootstrap (B = 1,999). For physician growth the bootstrap 95% CI for the lagged "
        f"litigation-rate coefficient was {ha.fmt(BS_PHYS['coef_boot_ci_low'], 4)} to "
        f"{ha.fmt(BS_PHYS['coef_boot_ci_high'], 4)} and the bootstrap p-value was {BS_PHYS['p_bootstrap']:.2f}; "
        f"for hospital growth the bootstrap 95% CI was {ha.fmt(BS_HOSP['coef_boot_ci_low'], 4)} to "
        f"{ha.fmt(BS_HOSP['coef_boot_ci_high'], 4)} and the bootstrap p-value was {BS_HOSP['p_bootstrap']:.2f}. "
        f"Both intervals comfortably contain zero. Power diagnostics make the panel information explicit. "
        f"For physician growth, the minimum detectable effect was {EQP['mde_80pct']:.2f}% per SD at 80% power, "
        f"and the power to declare equivalence within the \u00b1{MARGIN1}% margin if the true effect were zero "
        f"was {EQP['tests'][0]['power_if_null']*100:.1f}%. For hospital growth the minimum detectable effect was "
        f"{EQH['mde_80pct']:.2f}% per SD and the equivalent power for the \u00b1{MARGIN1}% margin was "
        f"{EQH['tests'][0]['power_if_null']*100:.1f}%. The panel is therefore informative enough to rule out "
        f"policy-relevant effects for physicians, and to bound any hospital effect within a small margin.",
    )

    h(doc, "Counts versus rates, and confounders", level=2)
    b(
        doc,
        f"Using raw litigation counts rather than rates did not reveal a negative association in this "
        f"measured-only design (p={CNT['p']:.2f}). Figure 2 shows the contrast: the count exposure is "
        f"confounded by specialty size (panel a), whereas the rate-adjusted exposure is not (panel b); "
        f"points are coloured and shaped by specialty so readers can identify which fields drive any "
        f"apparent pattern. The annual hospital and interpolated annual-physician sensitivity analyses were "
        f"also null (p={ANN['p']:.2f} and p={INT['p']:.2f}), confirming that the null result is robust to panel "
        f"frequency and exposure definition. The JOCS-CP indicator was positive in sign in the obstetric-hospital "
        f"model (coefficient {ha.fmt(HOSP['jocscp_coef'], 3)}, raw p={HOSP['jocscp_p']:.3f}), but it did not "
        f"remain significant after the small-cluster correction and Holm adjustment for the exploratory "
        f"sensitivity family (Holm p={JOCS_HOLM:.3f}); we therefore treat it as exploratory and do not "
        f"interpret it as a causal policy effect.",
    )
    f(
        doc,
        "ha_Figure_2.png",
        "Figure 2. Biennial physician growth against lagged litigation exposure measured as "
        "(a) counts and (b) rates. Points are coloured by specialty; the count panel shows the size "
        "confounding that the rate panel removes.",
    )
    b(
        doc,
        "The same count-versus-rate contrast for biennial hospital growth is shown in Figure 3. As "
        "with physician growth, the count exposure creates a spurious size confound that disappears once "
        "the rate-adjusted exposure is used.",
    )
    f(
        doc,
        "ha_Figure_3.png",
        "Figure 3. Biennial hospital growth against lagged litigation exposure measured as "
        "(a) counts and (b) rates. Points are coloured by specialty; the rate-adjusted panel shows "
        "no systematic association.",
    )
    b(
        doc,
        f"Descriptively, per-specialty rank correlations between the lagged litigation rate and physician "
        f"growth were positive in {n_pos} of {N} specialties and statistically significant in {n_sig}; the "
        f"direction is therefore, if anything, opposite to a flight-from-risk hypothesis.",
    )
    b(
        doc,
        f"A reverse specification (change in litigation rate regressed on lagged log physicians) was also "
        f"null (coefficient {ha.fmt(REV['coef'], 3)}, p={REV['p']:.2f}; Table 2), making a reverse-causation "
        f"interpretation of the null unlikely.",
    )
    b(
        doc,
        f"We also evaluated JMSR medical-accident investigation report counts as a potential confounder or "
        f"competing exposure.{{mais}} From {ha.JMSR_CORR['years'][0]} to {ha.JMSR_CORR['years'][-1]}, raw "
        f"litigation and JMSR report counts were strongly correlated across specialties (Pearson "
        f"r={ha.JMSR_CORR['pooled_r']:.2f}), because large specialties generate more of both. After removing "
        f"specialty-specific levels and trends, however, the within-specialty correlation was negligible "
        f"(r={ha.JMSR_CORR['detrended_r']:.2f}). A model of annual hospital growth for {JMSR_START}-2024 that "
        f"included both the lagged litigation rate and the lagged JMSR report rate left the litigation "
        f"coefficient essentially unchanged ({ha.fmt(JMSR['lit_coef'], 4)}; p={JMSR['lit_p']:.2f}) and the JMSR "
        f"term was not associated with hospital growth (p={JMSR['med_p']:.2f}; Supplementary Table 3). The null "
        f"litigation result is therefore neither explained nor masked by broader medical-accident reporting.",
    )
    b(
        doc,
        f"Finally, we tested national newspaper coverage from Nikkei Telecom 21 as a potential confounder."
        f"{{nikkei}} Total annual article counts (keywords: medical error + medical malpractice) and total "
        f"litigation counts were correlated (Pearson r={ha.MEDIA_CORR['total_r']:.2f}), consistent with "
        f"greater public attention in high-litigation years. Within the annual hospital panel, however, the "
        f"lagged litigation rate and the media-count series were only weakly correlated. A model of annual "
        f"hospital growth for {MEDIA_START}-{MEDIA_END} that included both the lagged litigation rate and the "
        f"lagged article count (per 1,000 articles) left the litigation coefficient essentially unchanged and "
        f"the media term was not associated with hospital growth (p={MEDIA['media_p']:.2f}; Supplementary "
        f"Table 4). Media coverage therefore does not explain the null litigation effect either. Holm step-down "
        f"adjusted p-values for the exploratory sensitivity family are reported in Supplementary Table 5.",
    )

    h(doc, "Counterfactual simulation", level=2)
    b(
        doc,
        f"The counterfactual projection made the practical implications of the null regression result "
        f"explicit. Under the point estimate, eliminating all malpractice litigation would add only "
        f"{TOTAL_SIM.get('marginal_pct_lit_point', 0):.1f}% to the projected 2034 national physician stock "
        f"relative to baseline drift. Even under the 95% lower-bound (most favourable) coefficient it would add "
        f"{TOTAL_SIM.get('marginal_pct_lit_lower', 0):.1f}%, comparable to the "
        f"{TOTAL_SIM.get('marginal_pct_mde', 0):.1f}% gain from a generic lever equal to the minimum detectable "
        f"effect. General surgery, the only specialty with negative baseline drift, illustrates the break-even "
        f"arithmetic: its projected 2024-2034 decline of {SURG_SIM.get('pct_change_baseline', 0):.1f}% would be "
        f"reduced to {SURG_SIM.get('pct_change_lit_zero_point', 0):.1f}% under the point estimate and reversed to "
        f"{SURG_SIM.get('pct_change_lit_zero_lower', 0):.1f}% under the 95% lower bound. The latter requires "
        f"eliminating every remaining closed claim and assumes the most adverse (most negative) coefficient "
        f"compatible with the data; a more realistic policy would achieve far less. Full projected 2034 physician "
        f"counts by specialty and lever are reported in Supplementary Table 7; Figure 4 summarises the same "
        f"information as marginal percentage changes. Litigation reduction is therefore not a high-leverage "
        f"instrument for workforce allocation in this setting.",
    )
    f(
        doc,
        "ha_Figure_4.png",
        "Figure 4. Counterfactual policy-lever simulation: marginal 10-year change in physician counts by "
        "specialty relative to the projected baseline drift. The MDE benchmark is the minimum detectable "
        "per-SD effect from the primary analysis.",
    )

    # Discussion
    h(doc, "Discussion", level=1)
    b(
        doc,
        f"Using national primary data, rates rather than counts, and only measured biennial physician "
        f"observations, we found no association between specialty-level malpractice-litigation risk and "
        f"subsequent physician or hospital decline. Equivalence testing showed that any effect of litigation risk "
        f"on biennial physician growth is smaller than {MARGIN1}% (90% CI within the {MARGIN1}% margin), and any "
        f"effect on hospital growth is smaller than {MARGIN2}% (but not confidently smaller than {MARGIN1}%). "
        f"These data therefore do not support the hypothesis that physicians systematically abandon "
        f"high-litigation specialties over {SPAN} years of official statistics.",
    )
    b(
        doc,
        "The null result is not merely a failure to detect an effect. The narrow confidence intervals, "
        "pre-specified equivalence margins, and power diagnostics allow us to say that, if litigation risk "
        "does influence specialty-level workforce growth, the magnitude is too small to matter for workforce "
        "planning. For behavioural public policy, this is a important distinction: a widely discussed risk can "
        "be highly available and emotionally salient without being a reliable policy lever. The public policy "
        "question is not whether physicians worry about litigation, but whether a policy that reduces litigation "
        "risk would materially change aggregate specialty supply. Our evidence suggests it would not.",
    )
    h(doc, "Behavioural mechanisms", level=2)
    b(
        doc,
        "Several behavioural mechanisms are consistent with this finding. First, litigation risk may affect "
        "clinical behaviour on the intensive margin (defensive medicine) rather than the extensive margin "
        "(specialty exit). Kessler and McClellan showed that U.S. malpractice reforms reduced medical "
        "expenditures for elderly heart-disease patients without increasing mortality or complications, "
        "suggesting that defensive practice is one margin of adjustment to liability pressure.{kessler1996} "
        "Subsequent reassessments have debated the magnitude and robustness of this effect, but the conceptual "
        "point remains: physicians can respond to liability risk by changing how they practise rather than by "
        "exiting a specialty.{sloan2008} Fee-for-service reimbursement in Japan rewards the high-acuity "
        "procedural work that also carries litigation exposure, so the financial return to remaining in surgery, "
        "obstetrics, or interventional specialties may dominate any deterrent from civil claims.",
    )
    b(
        doc,
        "Second, the discrepancy between perceived risk and measured workforce supply is consistent with "
        "well-documented behavioural-economics mechanisms. Media coverage of sensational malpractice or "
        "criminal prosecutions makes litigation risk highly available to physicians and trainees, and loss aversion "
        "can cause a rare but salient adverse outcome to be overweighted in career deliberations."
        "{tversky1973,kahneman1979} Yet the actual decision to leave a specialty is governed by expected income, "
        "sunk training costs, switching costs and status-quo bias, all of which discourage exit even when "
        "perceived risk is high.{samuelson1988} The gap between reported anxiety and measured supply is therefore "
        "not a contradiction; it is exactly what one would expect when a vivid, low-probability risk meets strong "
        "economic and institutional incentives to remain.",
    )
    h(doc, "Institutional context and international evidence", level=2)
    b(
        doc,
        "The civil litigation environment in Japan itself dampens the likelihood of a flight-from-risk response. "
        "Taniguchi and colleagues analysed all closed malpractice claims reported by the Supreme Court from 2006 "
        "to 2021 and found that more than half ended in settlement, plaintiffs won only about a quarter of "
        "judgments, and the number of claims has been declining, especially in obstetrics and gynaecology."
        "{taniguchi2023} The Court data we use therefore describe a civil system that is low-volume, "
        "settlement-prone, and comparatively favourable to physicians. This context makes it unlikely that routine "
        "civil litigation risk alone would drive physicians out of high-risk fields.",
    )
    b(
        doc,
        "No-fault obstetric compensation (JOCS-CP, 2009) illustrates a different mechanism. It was introduced "
        "partly because of a shortage of young obstetricians and regional gaps in maternity care, and it combined "
        "no-fault compensation with investigation and prevention.{hasegawa2016} The hospital-level JOCS-CP "
        f"indicator was directionally positive (coefficient {ha.fmt(HOSP['jocscp_coef'], 3)}, raw p={HOSP['jocscp_p']:.3f}), "
        f"but it did not remain significant after the small-cluster correction and Holm adjustment for the "
        f"exploratory sensitivity family (Holm p={JOCS_HOLM:.3f}). This suggests that, if the JOCS-CP did support "
        "obstetric hospital supply, the effect would be too small or too confounded by concurrent obstetric policies "
        "to be isolated here. Civil litigation exposure is also distinct from criminal prosecution. Morita studied the "
        "2004 Fukushima obstetrician prosecution and found a 13 percent decline in obstetricians, with some "
        "switching to gynaecology.{morita2018} Criminal cases and their media coverage may be far more salient to "
        "career decisions than routine closed civil claims, and our data do not capture that channel.",
    )
    b(
        doc,
        "The obstetrics and gynaecology case is the most discussed example of the litigation-workforce nexus, "
        "and it is consistent with our interpretation. A recent Japan\u2013U.S. comparison of medical-legal claims in "
        "obstetrics and gynaecology found that the proportion of malpractice claims in this specialty fell from "
        "15.1 percent in 2004 to 5.2 percent in 2022, and that claims per 100 OB/GYN physicians fell from 0.9 in "
        "2007 to 0.4 in 2016, while maternal and neonatal mortality also declined.{kamijo2025} The authors attribute "
        "this to heightened awareness after a wrongful criminal charge, the JOCS-CP no-fault scheme, standardised "
        "clinical guidelines, and the adverse-event investigation system. This is not evidence that lowering "
        "litigation risk caused the workforce to grow; it is evidence that obstetric litigation, workforce support, "
        "and safety interventions moved together. Surveys of OB/GYN residents in Japan, Korea, and Taiwan likewise "
        "show that litigation is reported as a negative factor, but that its perceived importance is smaller where "
        "no-fault compensation exists and that workload, lifestyle, and professional interest remain dominant."
        "{lin2022} These findings echo our specialty-level result: litigation may matter for perceptions, but it is "
        "not the binding constraint on supply.",
    )
    b(
        doc,
        "International evidence on tort reform and physician supply is consistent with a small or context-specific "
        "effect. A review of U.S. evidence by Helland and Seabury concluded that the measured impacts on physician "
        "supply are generally modest and sensitive to specification.{helland2015} Matsa found that U.S. state damage "
        "caps increased the supply of frontier rural specialists by 10-12 percent, but did not affect physician supply "
        "for the average resident.{matsa2007} Hyman and colleagues, examining the 2003 Texas reforms, found no "
        "measurable increase in physician supply for high-malpractice-risk specialties, primary care, or rural "
        "physicians.{hyman2015} Frakes and colleagues showed that negligence-standard reforms could shift the composition "
        "of the physician workforce toward surgery in some regions, yet the effect was localised and modest.{frakes2020} "
        "Against this backdrop, a null effect of civil litigation risk on specialty supply is not surprising, especially in "
        "a system with comparatively low litigation volume and predictable damages.",
    )
    b(
        doc,
        "The raw-count sensitivity did not reveal a negative association in these data, illustrating that the apparent "
        "count-litigation relationship does not translate into a behavioural effect once specialty size is accounted for "
        "(Figure 2). This is a cautionary example for workforce research that pairs administrative count series, and shows "
        "why rate-based, measured-only designs are preferable when testing litigation-workforce hypotheses.",
    )
    h(doc, "Policy implications", level=2)
    b(
        doc,
        f"What do these findings imply for behavioural public policy? Reducing civil malpractice litigation is unlikely "
        f"to be a powerful lever for correcting specialty maldistribution in this national setting. The 10-year counterfactual "
        f"simulation showed that eliminating all litigation would add only {TOTAL_SIM.get('marginal_pct_lit_point', 0):.1f}% "
        f"to the projected national physician stock under the point estimate, and even "
        f"{TOTAL_SIM.get('marginal_pct_lit_lower', 0):.1f}% under the most favourable 95% lower-bound coefficient, before "
        f"accounting for the implausibility of zero claims. Structural incentives are more promising: no-fault compensation "
        f"can de-risk high-acuity specialties, and payment design can reward service in underserved settings and activities. "
        f"The JOCS-CP experience supports the former; the country's fee-for-service schedule and rural/urban payment "
        f"adjustments illustrate the latter. Malpractice reform may still matter for defensive medicine, patient compensation, "
        f"and provider-patient trust. But our evidence does not support the claim, at least from these data, that lowering "
        f"litigation risk will retain physicians in high-risk specialties. Policymakers should therefore target structural "
        f"incentives before relying on litigation-avoidance messaging.",
    )
    b(
        doc,
        "International experience with no-fault compensation is consistent with this policy orientation. New Zealand replaced "
        "tort-based medical-injury compensation with a government-funded no-fault scheme in 1974 and, after 2005 reforms, "
        "extended coverage to all treatment injuries; this separated compensation from negligence findings and largely barred "
        "malpractice litigation.{bismark2006} Sweden and Denmark operate similar administrative systems in which neutral "
        "experts evaluate claims without requiring proof of provider fault, improving injured patients' access to redress while "
        "controlling liability costs and generating patient-safety learning.{mello2011} The JOCS-CP is narrower in scope\u2014"
        "it covers only obstetric cerebral palsy\u2014but it moves in the same direction: it provides compensation and cause "
        "analysis without a protracted adversarial process. Extending such an approach more broadly would be a structural "
        "alternative to repeated calls to reduce malpractice litigation as a workforce strategy.",
    )
    h(doc, "Limitations", level=2)
    b(
        doc,
        f"This is an ecological, specialty-level analysis and cannot establish individual-level causality. The 12 specialties "
        f"correspond to the primary-specialty tier of Japan's two-tiered specialist training programme; the analysis therefore "
        f"describes workforce allocation at the initial board-certification stage and may not extend to narrower subspecialties that "
        f"are not separately tracked in the biennial census. The physician census is biennial, giving {len(BIEN)} measured waves; we addressed the limited power directly through equivalence "
        f"testing and by pooling across specialties, but residual power constraints remain and the equivalence margins are a "
        f"judgement. Litigation rates may be endogenous to physician supply if a smaller workforce increases workload and hence "
        f"incidents; the lagged exposure, fixed effects, and reverse specification make reverse causation unlikely, yet unobserved "
        f"confounders at the specialty or prefecture level cannot be fully ruled out. Cluster block-bootstrap and power diagnostics "
        f"are reported in Supplementary Table 6. Because clusters are defined by the {N} specialties, the small-cluster correction "
        f"uses G-1={PHYS['df']} degrees of freedom; this is the minimum at which cluster-robust t inference is recommended and is "
        f"inherent to the data. Specialty-specific litigation counts could be recovered only from {BIEN[0]}; pre-{BIEN[0]} specialty "
        f"tables were not retrievable from primary sources. Clinic counts by specialty are published only every "
        f"{ha.CLINIC_RES} years and were used descriptively. JMSR report counts are available only from {ha.JMSR_CORR['years'][0]} "
        f"and were used in a {JMSR_START}-2024 sensitivity. Media article counts are available only for {MEDIA_START}-{MEDIA_END} and "
        f"are a national total, so they cannot be decomposed by specialty and are collinear with full wave fixed effects. Litigation "
        f"counts are assigned to a principal specialty and, by the Court's own note, do not measure intrinsic specialty risk.{{court}} "
        f"Finally, these findings are embedded in the country's particular legal, cultural and institutional context\u2014including its "
        f"no-fault obstetric compensation scheme, its fee-for-service reimbursement structure and its comparatively low-volume "
        f"malpractice-litigation culture\u2014so physician responses to litigation risk may differ in health systems with different "
        f"liability regimes, compensation mechanisms or professional norms; the results should not be assumed to generalise across "
        f"cultural spheres.",
    )

    h(doc, "Conclusions", level=1)
    b(
        doc,
        f"Across {YEARS}, specialty-level malpractice-litigation risk was not associated with physician or hospital decline in "
        f"these national data, and the physician effect was statistically equivalent to null within a small margin. From a "
        f"behavioural public policy perspective, malpractice litigation is not a reliable lever for correcting specialty "
        f"maldistribution in this setting. Policymakers may more productively target structural incentives, especially no-fault "
        f"compensation and payment design, rather than rely on the assumption that reducing litigation will retain physicians in "
        f"high-risk specialties. The transparent, reproducible sensitivity-analysis framework used here is exportable to other "
        f"healthcare workforce-policy levers.",
    )

    # AI declaration (Cambridge / BPP requirement)
    h(doc, "Declaration of artificial intelligence use", level=1)
    p(
        doc,
        "During the preparation of this work the author(s) used generative artificial intelligence "
        "tools to assist with literature synthesis, drafting, code generation for data analysis, and "
        "manuscript preparation. All generated content was reviewed, edited, and verified by the author(s), "
        "who take full responsibility for the final content.",
    )

    # Declarations
    h(doc, "Declarations", level=1)
    p(
        doc,
        "Funding: none. Competing interests: none declared. Ethics approval: this study used publicly "
        "available aggregated national statistics and did not involve human subjects, identifiable data or "
        "patient records; no ethics approval was required. "
        f"Data and code availability: all primary data files, extraction scripts and analysis code are "
        f"openly available in the project repository ({PUBLIC_REPO}), enabling full reproduction of every "
        "reported number.",
    )

    # References
    h(doc, "References", level=1)
    for i, k in enumerate(ha._CITE_ORDER, 1):
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.line_spacing = 1.5
        r_ref = p_ref.add_run(f"{i}. {ha.REFS[k]}")
        r_ref.font.size = Pt(10)
        r_ref.font.name = "Times New Roman"

    out = os.path.join(BASE, "bpp_manuscript_en.docx")
    doc.save(out)
    main_wc = sum(ha.wc(t) for t in ha.BODY_TEXTS)
    print(f"wrote {out}; abstract {abstract_wc} words; main body ~{main_wc} words")
    return main_wc, abstract_wc


def build_title_page(main_word_count):
    doc = ha._setup_doc()
    for _ in range(4):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(18)
    rt = t.add_run(BPP_TITLE)
    rt.bold = True
    rt.font.size = Pt(15)
    rt.font.name = "Times New Roman"

    lines = [
        "Authors: Onishi Tatsuki",
        "Affiliation: Data Science AI Innovation Research Promotion Center, Shiga University, "
        "1-1-1 Bamba, Hikone, Shiga 522-8522, Japan",
        "Corresponding author: Onishi Tatsuki",
        "ORCID: [corresponding author ORCID]    Email: [corresponding author email]",
        f"Word count (main text): approximately {main_word_count} words (excluding abstract, references, declarations, tables and figure legends)",
        "Article type: Original research article",
        "Target journal: Behavioural Public Policy (Cambridge Core)",
        "Tables: 2  Figures: 4  Supplementary tables: 7  Supplementary figures: 3",
        "Conflicts of interest: none declared",
        "Funding: none",
        f"Data availability: all primary data and analysis code are openly available in the project repository ({PUBLIC_REPO}).",
    ]
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(12)
        r.font.name = "Times New Roman"

    out = os.path.join(BASE, "bpp_title_page.docx")
    doc.save(out)
    print("wrote", out)


def build_highlights():
    highlights = [
        "Specialty-level malpractice litigation risk is unrelated to physician or hospital decline in Japan.",
        "Rate-based, measured-only designs remove size confounding and sparse-panel bias.",
        "Equivalence and power diagnostics support an informative null result.",
        "Perceived risk and real workforce allocation diverge because of structural incentives and status-quo bias.",
        "Structural incentives, not litigation-avoidance messaging, are the more promising policy lever.",
    ]
    for h_item in highlights:
        if len(h_item) > 120:
            raise SystemExit(f"Highlight exceeds 120 characters ({len(h_item)}): {h_item}")

    doc = ha._setup_doc()
    heading = doc.add_paragraph()
    rh = heading.add_run("Highlights")
    rh.bold = True
    rh.font.size = Pt(13)
    rh.font.name = "Times New Roman"
    for item in highlights:
        p = doc.add_paragraph(style="List Bullet")
        p.clear()
        pr = p.add_run(item)
        pr.font.name = "Times New Roman"
        pr.font.size = Pt(12)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(4)

    out = os.path.join(BASE, "bpp_highlights.docx")
    doc.save(out)
    print("wrote", out)


def build_cover_letter():
    doc = ha._setup_doc()
    for line in [
        "[Date]",
        "",
        "Professor Adam Oliver",
        "Editor-in-Chief",
        "Behavioural Public Policy",
        "London School of Economics and Political Science",
        "",
    ]:
        p = doc.add_paragraph()
        if line:
            r = p.add_run(line)
            r.font.size = Pt(12)
            r.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.add_run("Dear Professor Oliver,").font.size = Pt(12)
    p.runs[0].font.name = "Times New Roman"

    paragraphs = [
        f'We submit an original research article, "{BPP_TITLE}", for consideration by Behavioural Public Policy.',
        "Behavioural Public Policy advances rigorous, multidisciplinary research that connects the study of human "
        "behaviour to public policy. Our study sits squarely within this agenda. It uses a well-documented "
        "healthcare workforce problem\u2014specialty maldistribution\u2014as a policy test case and asks whether a "
        "salient, emotionally available risk (malpractice litigation) actually changes aggregate career behaviour. "
        "Using national administrative data from Japan, we find no association between litigation risk and specialty "
        "physician supply, and we bound any effect within a small equivalence margin. The result is informative for "
        "behavioural public policy because it shows that a widely perceived risk need not translate into a policy "
        "lever when structural incentives, switching costs and status-quo bias constrain individual choice.",
        "The behavioural contribution is threefold. First, we show how two common observational fallacies\u2014"
        "size confounding in raw administrative counts and interpolation of sparse panel data\u2014can distort the "
        "evidence base for a behavioural policy lever. Second, we combine fixed-effects panel methods, equivalence "
        "testing, cluster block-bootstrap and power diagnostics to produce an informative null result rather than a "
        "mere failure to reject the null. Third, we interpret the null through the lens of behavioural economics: "
        "availability, loss aversion and status-quo bias explain why perceived litigation risk can be high while "
        "aggregate workforce response is negligible.",
        "The national administrative data we use come from Japan, a setting that provides a complete, long-running "
        "test case. The analysis is fully reproducible from openly available primary files and code in the project "
        f"repository ({PUBLIC_REPO}). We believe the manuscript will be of interest to behavioural economists, health "
        "policy scholars, and public-policy analysts concerned with how risk perception and institutional incentives "
        "shape workforce behaviour.",
        "The work is original, is not under consideration elsewhere, and all authors approve the submission. We declare no conflicts of interest.",
    ]
    for b_text in paragraphs:
        p = doc.add_paragraph()
        r = p.add_run(b_text)
        r.font.size = Pt(11)
        r.font.name = "Times New Roman"
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.5

    for line in ["Sincerely,", "", "[Corresponding author, on behalf of all authors]"]:
        p = doc.add_paragraph()
        if line:
            r = p.add_run(line)
            r.font.size = Pt(12)
            r.font.name = "Times New Roman"

    out = os.path.join(BASE, "bpp_cover_letter.docx")
    doc.save(out)
    print("wrote", out)


def build_supplementary():
    """Build the Healthcare Analytics supplementary file and copy it to a BPP-named file."""
    # Always regenerate so the supplementary tables match the latest results JSON.
    ha.build_supplementary()
    src = os.path.join(BASE, "ha_supplementary.docx")
    dst = os.path.join(BASE, "bpp_supplementary.docx")
    shutil.copyfile(src, dst)
    print("wrote", dst)


def build_figure_pptx():
    """Build the Healthcare Analytics editable figure PPTX files and copy them to BPP-named files."""
    # Always regenerate so the embedded figures match the latest output/ha_*.png files.
    ha.build_figure_pptx()
    pairs = [
        ("ha_figures.pptx", "bpp_figures.pptx"),
        ("ha_supplementary_figures.pptx", "bpp_supplementary_figures.pptx"),
    ]
    for src_name, dst_name in pairs:
        src = os.path.join(BASE, src_name)
        dst = os.path.join(BASE, dst_name)
        shutil.copyfile(src, dst)
        print("wrote", dst)


def create_submission_zip():
    """Bundle the BPP submission files."""
    zip_path = os.path.join(OUT, "bpp_submission.zip")
    files = [
        os.path.join(BASE, "bpp_manuscript_en.docx"),
        os.path.join(BASE, "bpp_title_page.docx"),
        os.path.join(BASE, "bpp_cover_letter.docx"),
        os.path.join(BASE, "bpp_highlights.docx"),
        os.path.join(BASE, "bpp_supplementary.docx"),
        os.path.join(BASE, "bpp_figures.pptx"),
        os.path.join(BASE, "bpp_supplementary_figures.pptx"),
        os.path.join(OUT, "ha_Figure_1.png"),
        os.path.join(OUT, "ha_Figure_2.png"),
        os.path.join(OUT, "ha_Figure_3.png"),
        os.path.join(OUT, "ha_Figure_4.png"),
        os.path.join(OUT, "ha_Supplementary_Figure_1.png"),
        os.path.join(OUT, "ha_Supplementary_Figure_2.png"),
        os.path.join(OUT, "ha_Supplementary_Figure_3.png"),
    ]
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        for path in files:
            if not os.path.exists(path):
                raise SystemExit(f"submission zip missing file: {path}")
            z.write(path, arcname=os.path.basename(path))
    print("wrote", zip_path)


def main():
    main_wc, abstract_wc = build_manuscript()
    build_title_page(main_wc)
    build_highlights()
    build_cover_letter()
    build_supplementary()
    build_figure_pptx()
    create_submission_zip()


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""Build BMJ Open cover letter as editable .docx."""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "bmjopen_cover_letter.docx")

doc = Document()

# --- Page setup: A4, 2.54 cm margins ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(6)


def add_para(text, bold=False, italic=False, space_after=Pt(6), alignment=None):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = space_after
    return p


# ============================================================
# DATE AND ADDRESSEE
# ============================================================
add_para("[Date]")
add_para("")
add_para("The Editor")
add_para("BMJ Open")
add_para("BMA House, Tavistock Square")
add_para("London WC1H 9JR")
add_para("United Kingdom")
add_para("")

# ============================================================
# SUBJECT LINE
# ============================================================
add_para(
    "Re: Submission of original research article \u2014 "
    "\"Impact of medical safety incidents on physician workforce and healthcare facility "
    "supply across 12 specialties in Japan: an interrupted time series analysis\"",
    bold=True
)
add_para(
    "Manuscript ID: BMJ-2026-090547 (transferred from The BMJ)",
    italic=True
)
add_para("")

# ============================================================
# SALUTATION
# ============================================================
add_para("Dear Editor,")
add_para("")

# ============================================================
# BODY
# ============================================================

# Paragraph 1: Transfer context and what we are submitting
add_para(
    "We are pleased to submit the above manuscript to BMJ Open following an editorial "
    "transfer from The BMJ (manuscript ID: BMJ-2026-090547). We welcome the opportunity "
    "to have this work considered for publication in BMJ Open, whose broad scope and "
    "commitment to open-access dissemination of methodologically sound research make it "
    "an excellent venue for this study."
)

# Paragraph 2: What the study addresses and why it matters
add_para(
    "This study examines whether medical safety incidents have measurable downstream "
    "effects on the physician workforce and healthcare facility supply \u2014 a question of "
    "direct relevance to health systems worldwide that are grappling with workforce "
    "shortages in high-risk specialties. Despite widespread anecdotal concern that "
    "high-profile medical incidents deter physicians from certain specialties, no study "
    "has systematically quantified the temporal parameters of this association across "
    "multiple specialties."
)

# Paragraph 3: What we did
add_para(
    "Using national administrative data from Japan spanning up to 20 years, we conducted "
    "an interrupted time series analysis across 12 medical specialties, employing three "
    "complementary definitions of medical safety incidents: mandatory reporting to the "
    "Japan Medical Safety Research Organisation (JMSR), Supreme Court medical malpractice "
    "litigation statistics, and a composite index. We analysed the impact on three "
    "workforce outcomes: specialty-specific physician counts, healthcare facility numbers, "
    "and specialist trainee enrolment. Cross-correlation analysis was used to estimate "
    "lead times, and AIC-based model selection to estimate the duration of the effect "
    "(window period)."
)

# Paragraph 4: Key findings
add_para(
    "Our principal findings are as follows. First, medical safety incidents are followed "
    "by reductions in specialty-specific physician supply, with an estimated lead time of "
    "1\u20133 years and an effect duration of 4\u20135 years. Obstetrics and gynaecology showed "
    "the strongest and most consistent association (r = \u22120.95, lag +3 years). Second, "
    "the phenomenon extends well beyond obstetrics: otolaryngology (r = \u22120.97), "
    "dermatology (r = \u22120.99), and urology (r = \u22120.94) \u2014 specialties not traditionally "
    "considered \u2018high-risk\u2019 \u2014 also showed strong delayed inverse associations. Third, "
    "negative lag values observed for anaesthesiology (\u22124 years) and otolaryngology "
    "trainees (\u22123 years) suggest that the relationship may be bidirectional: workforce "
    "shortages may themselves precede and contribute to increased incident reporting, "
    "raising the possibility of a vicious cycle."
)

# Paragraph 5: Novelty and significance
add_para(
    "The study addresses a gap at the intersection of patient safety and workforce "
    "policy. While previous work has documented the Fukushima obstetrics prosecution as "
    "a natural experiment affecting a single specialty, no study has systematically "
    "quantified the temporal parameters of the incident\u2013workforce association across "
    "multiple specialties and definitions. The explicit estimation of lead time and "
    "window period parameters is methodologically novel and provides directly actionable "
    "inputs for workforce forecasting models. Furthermore, the finding that the "
    "relationship may be bidirectional has important implications for how health systems "
    "design both patient safety and workforce retention strategies."
)

# Paragraph 6: Reporting guidelines and data
add_para(
    "The study follows the REporting of studies Conducted using Observational "
    "Routinely-collected health Data (RECORD) guidelines and the Cochrane Effective "
    "Practice and Organisation of Care (EPOC) criteria for interrupted time series "
    "studies. A completed RECORD checklist is provided as supplementary table S1. "
    "All data sources are publicly available national administrative databases. "
    "Analysis code and data will be deposited in a public repository upon acceptance."
)

# Paragraph 7: Declarations
add_para(
    "This manuscript has not been published previously and is not under consideration by "
    "any other journal. All authors have read and approved the final manuscript. There "
    "are no conflicts of interest to declare. Ethical approval was not required as the "
    "study used only publicly available, aggregated, de-identified administrative data "
    "with no individual patient information. No patients or members of the public were "
    "directly involved in the design, conduct, or reporting of this research."
)

# Paragraph 8: Word count and article type
add_para(
    "The manuscript is submitted as an Original Research article. It contains "
    "approximately 3,800 words (excluding abstract, references, tables, and figures), "
    "a structured abstract of approximately 350 words, 9 figures, 4 tables, and "
    "4 supplementary tables.",
    italic=True
)

add_para("")

# ============================================================
# CLOSING
# ============================================================
add_para(
    "We thank you for considering this manuscript and look forward to your response."
)
add_para("")
add_para("Yours sincerely,")
add_para("")
add_para("Onishi Tatsuki")
add_para(
    "Data Science AI Innovation Research Promotion Center, Shiga University"
)
add_para(
    "1-1-1, Bamba, Hikone, Shiga, 522-8522 Japan"
)
add_para("[Email address]")
add_para("[ORCID]")

# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT_PATH)
print(f"Saved to {OUTPUT_PATH}")

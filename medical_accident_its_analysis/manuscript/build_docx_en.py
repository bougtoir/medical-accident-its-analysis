#!/usr/bin/env python3
"""
Build English manuscript (.docx) for VAR + Granger causality analysis.

Target journals (generic structure):
  - Int J Quality in Health Care (IJQHC) — Oxford, UK English, double-blind
  - Health Policy (Elsevier) — ≤4000 words, structured abstract ≤250 words
  - JMA Journal — ≤3500 words, structured abstract ≤300 words, ICMJE

Design: ≤3500 words body, structured abstract ≤250 words, Vancouver refs [n].
Figures NOT embedded per journal rules — legends at end, figures as separate files.
"""

import os
import re
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(BASE_DIR)
OUTPUT_DIR = os.path.join(PROJECT_DIR, 'output')
MANUSCRIPT_DIR = BASE_DIR

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 2.0

# ---- Helper functions ----

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
    return h


def _add_runs_with_citations(p, text, font_size=Pt(12), bold=False, italic=False):
    parts = re.split(r'(\[[^\]]+\])', text)
    for part in parts:
        if part.startswith('[') and part.endswith(']'):
            run = p.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.superscript = True
            run.bold = False
            run.italic = False
        else:
            if not part:
                continue
            run = p.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = font_size
            run.bold = bold
            run.italic = italic


def add_para(text, bold=False, italic=False, alignment=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    _add_runs_with_citations(p, text, bold=bold, italic=italic)
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 2.0
    return p


def add_abstract_field(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(6)
    run_label = p.add_run(label + ' ')
    run_label.bold = True
    run_label.font.name = 'Times New Roman'
    run_label.font.size = Pt(12)
    _add_runs_with_citations(p, text)
    return p


def add_table_from_data(headers, rows, caption=None):
    if caption:
        cap = doc.add_paragraph()
        cap.paragraph_format.space_after = Pt(6)
        cap.paragraph_format.space_before = Pt(12)
        cap.paragraph_format.line_spacing = 2.0
        run_cap = cap.add_run(caption)
        run_cap.font.name = 'Times New Roman'
        run_cap.font.size = Pt(10)
        run_cap.bold = True
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)
            if c_idx > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    return table


# ============================================================
# TITLE PAGE (separate — for non-blinded submission)
# ============================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(24)
run_t = title_p.add_run(
    'Temporal associations between medical safety incidents and '
    'physician workforce dynamics across 12 specialties in Japan: '
    'a vector autoregression and Granger causality analysis'
)
run_t.font.name = 'Times New Roman'
run_t.font.size = Pt(16)
run_t.bold = True

add_para('Authors: Onishi Tatsuki', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para(
    'Correspondence to: Onishi Tatsuki\n'
    'Data Science AI Innovation Research Promotion Center, Shiga University\n'
    '1-1-1, Bamba, Hikone, Shiga, 522-8522 Japan',
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
)
add_para(
    'Word count: approximately 3 400 (excluding abstract, references, '
    'tables, and figure legends)',
    alignment=WD_ALIGN_PARAGRAPH.CENTER, italic=True,
)
add_para(
    'Keywords: medical safety incidents, physician workforce, '
    'Granger causality, vector autoregression, health policy, Japan',
    alignment=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=Pt(24),
)

# ============================================================
# RESEARCH IN CONTEXT (Health Policy requirement; useful for all)
# ============================================================
doc.add_page_break()
add_heading_styled('Research in context', level=2)

add_heading_styled('What is already known', level=3)
known = [
    'Medical malpractice litigation has been associated with defensive medicine '
    'and physician avoidance of high-risk specialties, primarily studied in '
    'the United States.',
    'A single natural experiment in Japan (Fukushima obstetrics prosecution, '
    '2006) demonstrated a 13% reduction in obstetricians using '
    'difference-in-differences analysis.',
    'No study has systematically tested whether medical safety incidents '
    'Granger-cause changes in physician supply across multiple specialties.',
]
for item in known:
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 2.0

add_heading_styled('What this study adds', level=3)
adds = [
    'Using bivariate vector autoregression (VAR) models and Granger causality '
    'tests on 20 years of national data, we show that litigation statistics '
    'Granger-cause physician counts in 9 of 12 specialties and facility '
    'counts in 9 of 12 specialties.',
    'Obstetrics and gynaecology exhibited the strongest forward Granger '
    'causality (F = 46.66, P < 0.001), with impulse response functions '
    'indicating that a litigation shock reduces physician supply for '
    'approximately 3\u20135 years.',
    'Bidirectional Granger causality was found for 7 specialty-outcome '
    'combinations, suggesting a reinforcing cycle between workforce '
    'shortages and safety incidents.',
]
for item in adds:
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 2.0

add_heading_styled('Policy implications', level=3)
policy = [
    'Workforce planning models should incorporate a 1\u20134 year lag between '
    'safety incident trends and their downstream effects on physician supply.',
    'Specialties showing bidirectional causality (obstetrics, paediatrics, '
    'general surgery) may benefit most from early intervention to break '
    'the shortage\u2013incident cycle.',
]
for item in policy:
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 2.0

# ============================================================
# HIGHLIGHTS (Health Policy; useful summary for all)
# ============================================================
add_heading_styled('Highlights', level=2)
highlights = [
    'Medical safety incidents Granger-cause physician workforce changes '
    'in 9 of 12 specialties in Japan.',
    'Obstetrics and gynaecology shows the strongest temporal association '
    '(F = 46.66, P < 0.001).',
    'Impulse response functions reveal a 3\u20135 year effect duration after '
    'a litigation shock.',
    'Bidirectional causality suggests a reinforcing cycle between '
    'workforce shortages and incidents.',
    'VAR-based forecasts project continued decline in general surgery '
    'physicians and facilities through 2033.',
]
for item in highlights:
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 2.0

# ============================================================
# ABSTRACT
# ============================================================
doc.add_page_break()
add_heading_styled('Abstract', level=1)

add_abstract_field('Background:',
    'Medical safety incidents may deter physicians from high-risk '
    'specialties, yet quantitative evidence of their temporal impact on '
    'workforce supply remains limited.')

add_abstract_field('Objective:',
    'To test whether medical safety incident counts Granger-cause changes '
    'in specialty-specific physician and healthcare facility numbers in Japan.')

add_abstract_field('Methods:',
    'We constructed bivariate vector autoregression (VAR) models for 12 '
    'specialties using two national incident series\u2014mandatory safety reports '
    '(2015\u20132025) and malpractice litigation statistics (2004\u20132023)\u2014paired '
    'with physician and facility counts. Granger causality F-tests, impulse '
    'response functions (IRF), and VAR-based forecasts were computed.')

add_abstract_field('Results:',
    'Of 47 VAR models fitted, 20 showed significant forward Granger '
    'causality (incidents \u2192 workforce; P < 0.05). Obstetrics and '
    'gynaecology exhibited the strongest association with physician '
    'counts (F = 46.66, P < 0.001, lag 1 year). Significant forward '
    'causality was found for 9 of 12 specialties using litigation data. '
    'Bidirectional causality was observed in 7 specialty-outcome pairs, '
    'including obstetrics, paediatrics, and general surgery. IRFs indicated '
    'that a one-unit litigation shock reduced physician counts for 3\u20135 '
    'years. VAR forecasts project continued decline in general surgery '
    'physicians (\u2212230/year) and facilities (\u2212310/year) through 2033.')

add_abstract_field('Conclusions:',
    'Medical safety incidents Granger-cause reductions in physician '
    'supply across most specialties, with effects lasting 3\u20135 years. '
    'Bidirectional causality in several specialties suggests a reinforcing '
    'cycle. These findings support incorporating incident trends into '
    'workforce planning models.')

add_para(
    'Keywords: medical safety incidents, physician workforce, '
    'Granger causality, vector autoregression, health policy, Japan',
    italic=True, space_after=Pt(18),
)

# ============================================================
# BACKGROUND
# ============================================================
doc.add_page_break()
add_heading_styled('Background', level=1)

add_para(
    'The relationship between medical safety incidents and physician '
    'workforce dynamics is of growing importance to health systems '
    'worldwide. Concerns that malpractice risk drives physicians away '
    'from high-risk specialties have been raised for decades,[1\u20133] yet '
    'quantitative evidence remains limited, geographically concentrated '
    'in the United States, and largely focused on tort reform rather '
    'than incident occurrence itself.[4\u20136]'
)

add_para(
    'Japan provides a particularly informative setting for studying '
    'this phenomenon. In 2004, an obstetrician at Fukushima Prefectural '
    'Ohno Hospital was arrested following a maternal death during '
    'caesarean section, and subsequently prosecuted in 2006 for '
    'professional negligence causing death.[7] Although the physician '
    'was acquitted in 2008, the case is widely regarded as a catalyst '
    'for the subsequent decline in obstetricians willing to provide '
    'delivery services.[8,9] Morita estimated a 13% decrease in '
    'obstetricians in the affected prefecture using difference-in-'
    'differences and synthetic control methods.[10] However, this '
    'analysis was limited to a single event, a single specialty, and '
    'a single geographic area.'
)

add_para(
    'Japan\u2019s healthcare system features two national-level data '
    'sources on medical safety incidents. First, the Japan Medical '
    'Safety Research Organisation (JMSR) was established in 2015 '
    'under a revision to the Medical Care Act, creating a mandatory '
    'reporting system for unexpected deaths potentially related to '
    'medical care.[11] Second, the Supreme Court of Japan publishes '
    'annual statistics on medical malpractice litigation by '
    'specialty.[12,13] Additionally, Japan conducts biennial national '
    'surveys of all practising physicians and maintains annual records '
    'of healthcare facility registrations\u2014providing unusually detailed, '
    'population-level workforce data.'
)

add_para(
    'Previous analyses of these data, including our own, have relied on '
    'interrupted time series (ITS) methodology with cross-correlation.[14] '
    'However, ITS assumes a discrete, exogenous intervention point, '
    'which is poorly suited to continuously varying exposure data such '
    'as incident counts. Cross-correlation, while descriptive, does not '
    'formally test predictive causality and is susceptible to spurious '
    'associations from shared trends.'
)

add_para(
    'We therefore re-analysed the data using vector autoregression (VAR) '
    'models and Granger causality tests\u2014methods widely used in '
    'econometrics and increasingly applied in health policy '
    'research.[15\u201317] VAR models simultaneously capture the dynamic '
    'interdependence between incident counts and workforce outcomes, '
    'while Granger causality provides a formal statistical test of '
    'whether past values of one series improve prediction of another. '
    'We further employed impulse response functions (IRFs) to estimate '
    'the magnitude and duration of workforce effects following a safety '
    'incident shock, and VAR-based forecasting to project workforce '
    'trends.'
)

# ============================================================
# METHODS
# ============================================================
add_heading_styled('Methods', level=1)

add_heading_styled('Study design and reporting', level=2)
add_para(
    'We conducted a secondary analysis of routinely collected national '
    'administrative data from Japan using bivariate VAR models and '
    'Granger causality tests. This study is reported in accordance '
    'with the RECORD (REporting of studies Conducted using Observational '
    'Routinely-collected health Data) statement,[18,19] and the STROBE '
    '(Strengthening the Reporting of Observational Studies in '
    'Epidemiology) guidelines for ecological studies.'
)

add_heading_styled('Data sources', level=2)
add_para(
    'Medical safety incident data were obtained from two sources. '
    'Definition 1 (JMSR): annual specialty-specific incident reports '
    'to the Japan Medical Safety Research Organisation for fiscal years '
    '2015\u20132025 (11 annual observations).[11] Definition 2 (litigation): '
    'annual specialty-specific medical malpractice closed-claim counts '
    'from the Supreme Court of Japan for 2004\u20132023 (20 annual '
    'observations).[12]'
)
add_para(
    'Specialty-specific physician counts were obtained from the National '
    'Survey of Physicians, Dentists, and Pharmacists, conducted '
    'biennially by the Ministry of Health, Labour and Welfare, covering '
    '2002\u20132024.[20] As the survey is biennial, annual values were '
    'obtained by linear interpolation between survey years. '
    'Specialty-specific facility counts were obtained from the Survey '
    'of Medical Institutions, covering 2005\u20132024.[21] Specialist trainee '
    'enrolment data were obtained from the Japan Board of Medical '
    'Specialties for 2018\u20132025 (8 years).[22]'
)

add_heading_styled('Specialties analysed', level=2)
add_para(
    'Twelve core specialties with sufficient data across all sources '
    'were analysed: internal medicine, general surgery, orthopaedic '
    'surgery, plastic surgery, obstetrics and gynaecology, paediatrics, '
    'psychiatry, ophthalmology, otolaryngology, urology, dermatology, '
    'and anaesthesiology.'
)

add_heading_styled('Statistical analysis', level=2)

add_para(
    'Stationarity of each series was assessed using augmented '
    'Dickey\u2013Fuller (ADF) tests. When both series in a pair were '
    'non-stationary, first differences were used; otherwise, models '
    'were fitted in levels with a stationarity caveat noted.'
)

add_para(
    'For each specialty\u2013definition\u2013outcome combination, a bivariate '
    'VAR model was fitted with the two endogenous variables being '
    'incident counts and the workforce outcome (physician count or '
    'facility count). The optimal lag order was selected by minimising '
    'the Akaike information criterion (AIC), with a maximum lag of '
    'min(4, n/3 \u2212 1) to preserve degrees of freedom in these '
    'relatively short series.'
)

add_para(
    'Granger causality was tested in both directions within the VAR '
    'framework using F-tests on the lagged coefficients.[23] The null '
    'hypothesis for the forward test is that lagged incident counts '
    'do not improve prediction of the workforce outcome; for the '
    'reverse test, that lagged workforce counts do not improve '
    'prediction of incidents. Statistical significance was set at '
    'P < 0.05. We report the lag yielding the most significant '
    'F-statistic for each direction.'
)

add_para(
    'Impulse response functions (IRFs) were computed from the fitted '
    'VAR models to estimate the dynamic response of workforce outcomes '
    'to a one-unit shock in incident counts, over a 10-year horizon.[24] '
    'IRFs provide both the magnitude and duration of the workforce '
    'effect following a safety incident shock.'
)

add_para(
    'For specialties with litigation data (the longer series), '
    'VAR-based forecasts of workforce outcomes were generated for '
    '2024\u20132033, with 95% forecast intervals.'
)

add_para(
    'For specialist trainees, the short data series (8 years) '
    'precluded formal VAR modelling; detrended Pearson correlations '
    'between JMSR incidents and trainee enrolment were computed instead.'
)

add_para(
    'All analyses were performed in Python 3.10 using statsmodels '
    '0.14 (VAR and Granger causality), SciPy 1.14, and pandas '
    '2.2.[25\u201327] Analysis code and data are available at '
    'https://github.com/bougtoir/medical-accident-its-analysis.'
)

add_heading_styled('Patient and public involvement', level=2)
add_para(
    'No patients or members of the public were involved in this '
    'research. This study used exclusively publicly available, '
    'aggregated administrative data.'
)

# ============================================================
# RESULTS
# ============================================================
doc.add_page_break()
add_heading_styled('Results', level=1)

add_heading_styled('Descriptive overview', level=2)
add_para(
    'Over the study period, the 12 core specialties encompassed '
    'approximately 200 000 physicians and 150 000 healthcare '
    'facilities. Between 2015 and 2025, the JMSR received 3 860 '
    'reports across all specialties, with general surgery (n = 520), '
    'internal medicine (n = 479), and orthopaedic surgery (n = 310) '
    'recording the highest volumes. Medical malpractice litigation '
    'cases across the 12 specialties declined from 1 139 in 2004 to '
    '746 in 2023 overall, although trends varied by specialty '
    '(Figure 1).'
)

add_heading_styled('VAR model specification', level=2)
add_para(
    'A total of 47 bivariate VAR models were fitted across the 12 '
    'specialties, two incident definitions, and two workforce outcomes. '
    'The litigation-based models (20 annual observations) provided '
    'the primary analysis; JMSR-based models (10\u201311 overlapping '
    'observations) served as sensitivity analyses. AIC-selected lag '
    'orders ranged from 1 to 4 (median 2).'
)

# ---- Table 1: Granger causality results ----
add_heading_styled('Granger causality tests', level=2)
add_para(
    'Table 1 presents the Granger causality test results for all '
    'specialty\u2013outcome combinations using litigation data. Of the 24 '
    'litigation-based tests (12 specialties \u00d7 2 outcomes), 18 showed '
    'significant forward Granger causality (incidents \u2192 workforce, '
    'P < 0.05). The strongest forward associations were observed for '
    'obstetrics and gynaecology (physicians: F = 46.66, P < 0.001; '
    'facilities: F = 16.34, P < 0.001), psychiatry (physicians: '
    'F = 20.89, P < 0.001), and anaesthesiology (facilities: '
    'F = 20.34, P < 0.001).'
)

# Load Granger results for table
gc_df = pd.read_csv(os.path.join(OUTPUT_DIR, 'granger_causality_results.csv'))
gc_lit = gc_df[gc_df['definition'] == 'litigation'].copy()

table1_headers = ['Specialty', 'Outcome', 'VAR lag', 'n',
                  'Forward F', 'Forward P', 'Reverse F', 'Reverse P']
table1_rows = []
for _, r in gc_lit.sort_values(['specialty_en', 'outcome']).iterrows():
    fwd_f = f"{r['gc_forward_F']:.2f}" if pd.notna(r.get('gc_forward_F')) else '\u2014'
    fwd_p = f"{r['gc_forward_p']:.4f}" if pd.notna(r.get('gc_forward_p')) else '\u2014'
    rev_f = f"{r['gc_reverse_F']:.2f}" if pd.notna(r.get('gc_reverse_F')) else '\u2014'
    rev_p = f"{r['gc_reverse_p']:.4f}" if pd.notna(r.get('gc_reverse_p')) else '\u2014'
    # Add significance markers
    if pd.notna(r.get('gc_forward_p')) and r['gc_forward_p'] < 0.05:
        fwd_p += '*'
    if pd.notna(r.get('gc_reverse_p')) and r['gc_reverse_p'] < 0.05:
        rev_p += '*'
    table1_rows.append([
        r['specialty_en'], r['outcome'].capitalize(),
        str(int(r['var_lag'])), str(int(r['n_obs'])),
        fwd_f, fwd_p, rev_f, rev_p,
    ])

add_table_from_data(
    table1_headers, table1_rows,
    caption='Table 1. Granger causality test results for litigation data '
            '(12 specialties \u00d7 2 outcomes). Forward: incidents \u2192 workforce; '
            'Reverse: workforce \u2192 incidents. *P < 0.05.'
)

add_para(
    'Significant reverse Granger causality (workforce \u2192 incidents) was '
    'observed in 8 litigation-based models, indicating bidirectional '
    'temporal dynamics. Seven specialty\u2013outcome combinations showed '
    'significant causality in both directions: obstetrics and '
    'gynaecology (physicians and facilities), paediatrics (physicians '
    'and facilities), general surgery (JMSR\u2013physicians), ophthalmology '
    '(facilities), and otolaryngology (facilities). This bidirectional '
    'pattern is consistent with a reinforcing cycle in which safety '
    'incidents deter physician recruitment, leading to workforce '
    'shortages that in turn increase incident risk.'
)

add_para(
    'In the JMSR-based sensitivity analysis (Table 2), only 2 of 24 '
    'tests showed significant forward Granger causality, likely '
    'reflecting the shorter series (10\u201311 observations) and '
    'consequent reduced statistical power. However, general surgery '
    '(F = 23.04, P = 0.003) and psychiatry (F = 9.48, P = 0.022) '
    'reached significance even with this limited sample.'
)

# Table 2: JMSR results
gc_jmsr = gc_df[gc_df['definition'] == 'jmsr'].copy()
table2_rows = []
for _, r in gc_jmsr.sort_values(['specialty_en', 'outcome']).iterrows():
    fwd_f = f"{r['gc_forward_F']:.2f}" if pd.notna(r.get('gc_forward_F')) else '\u2014'
    fwd_p = f"{r['gc_forward_p']:.4f}" if pd.notna(r.get('gc_forward_p')) else '\u2014'
    rev_f = f"{r['gc_reverse_F']:.2f}" if pd.notna(r.get('gc_reverse_F')) else '\u2014'
    rev_p = f"{r['gc_reverse_p']:.4f}" if pd.notna(r.get('gc_reverse_p')) else '\u2014'
    if pd.notna(r.get('gc_forward_p')) and r['gc_forward_p'] < 0.05:
        fwd_p += '*'
    if pd.notna(r.get('gc_reverse_p')) and r['gc_reverse_p'] < 0.05:
        rev_p += '*'
    table2_rows.append([
        r['specialty_en'], r['outcome'].capitalize(),
        str(int(r['var_lag'])), str(int(r['n_obs'])),
        fwd_f, fwd_p, rev_f, rev_p,
    ])

add_table_from_data(
    table1_headers, table2_rows,
    caption='Table 2. Granger causality test results for JMSR data '
            '(sensitivity analysis, 10\u201311 observations). *P < 0.05.'
)

# ---- IRF results ----
add_heading_styled('Impulse response functions', level=2)
add_para(
    'Figure 2 shows the IRFs for the response of physician counts to '
    'a one-unit shock in litigation cases for six key specialties. '
    'For obstetrics and gynaecology, the IRF demonstrates a sustained '
    'negative response peaking at approximately 2\u20133 years and '
    'persisting for 5 years. General surgery shows a similar pattern '
    'with the response attenuating by year 4. The corresponding IRFs '
    'for facility counts are shown in Supplementary Figure S1.'
)

# ---- Forecasts ----
add_heading_styled('VAR-based forecasts', level=2)
add_para(
    'Figure 3 presents VAR-based forecasts for physician and facility '
    'counts in three specialties of particular policy interest. General '
    'surgery is projected to continue declining, with physician counts '
    'falling by approximately 230 per year and facilities by 310 per '
    'year through 2033. Obstetrics and gynaecology physician counts are '
    'projected to increase modestly, while facilities continue to '
    'decline. Internal medicine physician counts show continued growth '
    'but facility numbers are projected to decrease.'
)

# Table 3: Forecast summary
fc_df = pd.read_csv(os.path.join(OUTPUT_DIR, 'var_forecast_results.csv'))
table3_headers = ['Specialty', 'Outcome', '2028 (forecast)', '2033 (forecast)']
table3_rows = []
for spec in ['General surgery', 'Obstetrics & gynaecology', 'Internal medicine',
             'Paediatrics', 'Orthopaedic surgery', 'Anaesthesiology']:
    for outcome in ['physicians', 'facilities']:
        sub = fc_df[(fc_df['specialty_en'] == spec) & (fc_df['outcome'] == outcome)]
        if sub.empty:
            continue
        y2028 = sub[sub['year'] == 2028]
        y2033 = sub[sub['year'] == 2033]
        v2028 = f"{y2028['forecast_mean'].values[0]:,.0f}" if len(y2028) > 0 else '\u2014'
        v2033 = f"{y2033['forecast_mean'].values[0]:,.0f}" if len(y2033) > 0 else '\u2014'
        table3_rows.append([spec, outcome.capitalize(), v2028, v2033])

add_table_from_data(
    table3_headers, table3_rows,
    caption='Table 3. VAR-based workforce forecasts for selected specialties '
            '(point estimates from litigation\u2013outcome bivariate models).'
)

# ---- Trainee analysis ----
add_heading_styled('Trainee analysis', level=2)
add_para(
    'Detrended correlations between JMSR incident reports and specialist '
    'trainee enrolment (2018\u20132025, n = 8) showed notable negative '
    'associations for psychiatry (r = \u22120.77, P = 0.025), '
    'anaesthesiology (r = \u22120.77, P = 0.026), and obstetrics and '
    'gynaecology (r = \u22120.70, P = 0.051). However, these results should '
    'be interpreted cautiously given the short series and limited '
    'statistical power.'
)

# ============================================================
# DISCUSSION
# ============================================================
doc.add_page_break()
add_heading_styled('Discussion', level=1)

add_heading_styled('Principal findings', level=2)
add_para(
    'This study provides the first multi-specialty application of '
    'VAR models and Granger causality tests to examine the temporal '
    'relationship between medical safety incidents and physician '
    'workforce dynamics. Using two complementary incident definitions '
    'and national administrative data from Japan, we found that '
    'litigation statistics Granger-cause physician counts in 9 of 12 '
    'specialties and facility counts in 9 of 12 specialties. '
    'Obstetrics and gynaecology showed the strongest forward Granger '
    'causality (F = 46.66, P < 0.001), corroborating previous findings '
    'from the Fukushima prosecution case[10] but extending these to a '
    'national, multi-definition, multi-specialty framework.'
)

add_heading_styled('Methodological advance over ITS', level=2)
add_para(
    'The VAR/Granger approach addresses the four major methodological '
    'concerns identified in prior reviews of this research '
    'question.[14] First, VAR does not require an exogenous '
    'intervention point\u2014it models the dynamic interdependence of two '
    'time series directly. Second, Granger causality provides a formal '
    'statistical test of temporal precedence (\"does X help predict Y '
    'beyond Y\u2019s own past?\"), avoiding the risk of conflating '
    'contemporaneous correlation with predictive causality. Third, the '
    'stationarity testing and differencing protocol directly addresses '
    'spurious correlation risk. Fourth, VAR-based forecasting is '
    'grounded in the estimated multivariate dynamics rather than '
    'univariate linear extrapolation.'
)

add_heading_styled('Bidirectional causality and policy implications', level=2)
add_para(
    'The observation of bidirectional Granger causality in 7 '
    'specialty\u2013outcome combinations has important policy implications. '
    'Forward causality (incidents \u2192 workforce decline) is consistent '
    'with the hypothesised \u2018incident avoidance\u2019 pathway: high-profile '
    'safety events deter entry and accelerate exit from a specialty. '
    'Reverse causality (workforce decline \u2192 incidents) is consistent '
    'with an \u2018overwork\u2019 pathway: fewer physicians per caseload leads '
    'to fatigue, reduced supervision, and increased incident risk.'
)
add_para(
    'The coexistence of both pathways in specialties such as obstetrics '
    'and gynaecology, paediatrics, and general surgery suggests a '
    'reinforcing vicious cycle. Policy interventions to break this '
    'cycle could target either pathway: improving post-incident support '
    'and public communication (to reduce avoidance), or augmenting '
    'workforce capacity and reducing per-physician workload (to reduce '
    'incident risk). The IRF results, showing that workforce effects '
    'persist for 3\u20135 years after a litigation shock, define a window '
    'during which such interventions could be most effective.'
)

add_heading_styled('Comparison with existing literature', level=2)
add_para(
    'Our findings are consistent with and extend the existing '
    'literature. The magnitude of the obstetrics and gynaecology '
    'association aligns with Morita\u2019s finding of a 13% decrease in '
    'obstetricians following the 2006 prosecution,[10] and with the '
    'broader literature on malpractice-driven specialty avoidance in '
    'the United States.[1\u20133] Klick and Stratmann found that tort '
    'reforms increased the supply of physicians in high-risk '
    'specialties,[4] providing further evidence for the incident '
    'avoidance pathway.'
)
add_para(
    'The present study extends this literature by demonstrating that '
    'the phenomenon is not limited to traditionally \u201chigh-risk\u201d '
    'specialties. Dermatology, ophthalmology, and otolaryngology\u2014'
    'not typically considered malpractice-sensitive\u2014showed significant '
    'forward Granger causality, suggesting that workforce sensitivity '
    'to safety incidents may be more widespread than previously '
    'recognised.'
)

add_heading_styled('Strengths and limitations', level=2)
add_para(
    'Strengths of this study include the use of population-level data '
    'from mandatory national registries, covering 12 specialties over '
    'up to 20 years; the application of VAR/Granger methodology, which '
    'is better suited than ITS to continuously varying exposures; and '
    'the formal testing of bidirectional causality.'
)
add_para(
    'Several limitations should be acknowledged. First, this is an '
    'ecological study using aggregate national data; individual-level '
    'causal pathways cannot be inferred. Granger causality tests '
    'temporal precedence in prediction, not true causation in the '
    'interventionist sense. Second, the biennial physician survey data '
    'required linear interpolation, which may smooth short-term '
    'fluctuations. Third, the JMSR series (11 years) provided limited '
    'statistical power, as reflected in the few significant JMSR-based '
    'results. Fourth, unmeasured confounders\u2014including demographic '
    'shifts, remuneration differentials, lifestyle preferences, and '
    'policy interventions such as the 2004 postgraduate clinical '
    'training reform\u2014may explain some associations. Fifth, the '
    'specialist trainee data cover only 8 years, precluding formal '
    'VAR analysis for this outcome.'
)

# ============================================================
# CONCLUSIONS
# ============================================================
add_heading_styled('Conclusions', level=1)
add_para(
    'Medical safety incidents in Japan Granger-cause reductions in '
    'specialty-specific physician supply in 9 of 12 specialties '
    'analysed, with effects persisting for 3\u20135 years as estimated by '
    'impulse response functions. Bidirectional causality in several '
    'specialties suggests a reinforcing cycle between workforce '
    'shortages and safety incidents. These temporal parameters should '
    'be incorporated into workforce forecasting models. Further '
    'research using individual-level panel data and natural experiments '
    'is needed to confirm these associations and elucidate the '
    'underlying mechanisms.'
)

# ============================================================
# DECLARATIONS
# ============================================================
doc.add_page_break()
add_heading_styled('Declarations', level=1)

add_heading_styled('Funding', level=2)
add_para('[To be completed by authors]')

add_heading_styled('Competing interests', level=2)
add_para(
    'The author declares no competing interests.'
)

add_heading_styled('Ethical approval', level=2)
add_para(
    'This study used publicly available, anonymised, aggregate-level '
    'administrative data. No individual-level patient or physician data '
    'were accessed. Ethical approval was not required.'
)

add_heading_styled('Data availability', level=2)
add_para(
    'The analysis code and aggregated datasets used in this study are '
    'available at https://github.com/bougtoir/medical-accident-its-analysis. '
    'The original data sources are publicly available from the Japan '
    'Medical Safety Research Organisation, the Supreme Court of Japan, '
    'and the Ministry of Health, Labour and Welfare.'
)

add_heading_styled('Use of AI-assisted technologies', level=2)
add_para(
    'AI-assisted tools were used for coding assistance and manuscript '
    'drafting. The author takes full responsibility for the content '
    'and integrity of the work.'
)

# ============================================================
# REFERENCES
# ============================================================
doc.add_page_break()
add_heading_styled('References', level=1)

references = [
    '1. Studdert DM, Mello MM, Sage WM, et al. Defensive medicine among high-risk specialist physicians in a volatile malpractice environment. JAMA 2005;293:2609\u201317.',
    '2. Mello MM, Studdert DM, DesRoches CM, et al. Caring for patients in a malpractice crisis: physician satisfaction and quality of care. Health Aff (Millwood) 2004;23:42\u201353.',
    '3. Kessler D, McClellan MB. Do doctors practice defensive medicine? Q J Econ 1996;111:353\u201390.',
    '4. Klick J, Stratmann T. Medical malpractice reform and physicians in high-risk specialties. J Legal Stud 2007;36:S121\u201342.',
    '5. Frakes M. Defensive medicine and obstetric practices. J Empir Legal Stud 2012;9:457\u201381.',
    '6. Frakes MD, Gruber J, Jena AB. Is great information good enough? Evidence from physicians as patients. J Health Econ 2021;75:102406.',
    '7. Nagamatsu S, Kami M, Nakata Y. Healthcare safety committee in Japan: mandatory accountability reporting system and punishment. Curr Opin Anaesthesiol 2009;22:199\u2013206.',
    '8. Hiyama T, Yoshihara M, Tanaka S, et al. Defensive medicine practices among gastroenterologists in Japan. World J Gastroenterol 2006;12:7671\u20135.',
    '9. Ishikawa T. Distribution and retention of obstetrician-gynecologists in Japan: a longitudinal study, 1996\u20132016. Nihon Iji Shimpo 2021. [in Japanese]',
    '10. Morita H. Criminal prosecution and physician supply. Int Rev Law Econ 2018;55:1\u201311.',
    '11. Japan Medical Safety Research Organisation. Annual report on medical accident investigation. Tokyo: JMSR; 2025. Available from: https://www.medsafe.or.jp/',
    '12. Supreme Court of Japan. Annual report of judicial statistics: medical malpractice litigation. Tokyo: Supreme Court; 2024. [in Japanese]',
    '13. Taniguchi K, Watari T, Nagoshi K. Characteristics and trends of medical malpractice claims in Japan between 2006 and 2021. PLoS One 2024;19:e0296155.',
    '14. Onishi T. Impact of medical safety incidents on physician workforce and healthcare facility supply across 12 specialties in Japan: an interrupted time series analysis. Preprint 2025.',
    '15. L\u00fctkepohl H. New Introduction to Multiple Time Series Analysis. Berlin: Springer; 2005.',
    '16. Granger CWJ. Investigating causal relations by econometric models and cross-spectral methods. Econometrica 1969;37:424\u201338.',
    '17. Toda HY, Yamamoto T. Statistical inference in vector autoregressions with possibly integrated processes. J Econom 1995;66:225\u201350.',
    '18. Benchimol EI, Smeeth L, Guttmann A, et al. The REporting of studies Conducted using Observational Routinely-collected health Data (RECORD) statement. PLoS Med 2015;12:e1001885.',
    '19. Langan SM, Schmidt SAJ, Wing K, et al. The reporting of studies conducted using observational routinely collected health data statement for pharmacoepidemiology (RECORD-PE). BMJ 2018;363:k3532.',
    '20. Ministry of Health, Labour and Welfare. Survey of physicians, dentists, and pharmacists. Tokyo: MHLW; 2023. [in Japanese]',
    '21. Ministry of Health, Labour and Welfare. Survey of medical institutions (dynamic survey). Tokyo: MHLW; 2024. [in Japanese]',
    '22. Japan Board of Medical Specialties. Specialist trainee registration statistics. Tokyo: JBMS; 2025. [in Japanese]',
    '23. Hamilton JD. Time Series Analysis. Princeton: Princeton University Press; 1994.',
    '24. Sims CA. Macroeconomics and reality. Econometrica 1980;48:1\u201348.',
    '25. Seabold S, Perktold J. Statsmodels: econometric and statistical modeling with Python. Proceedings of the 9th Python in Science Conference. 2010:92\u20136.',
    '26. Virtanen P, Gommers R, Oliphant TE, et al. SciPy 1.0: fundamental algorithms for scientific computing in Python. Nat Methods 2020;17:261\u201372.',
    '27. McKinney W. Data structures for statistical computing in Python. Proceedings of the 9th Python in Science Conference. 2010:56\u201361.',
    '28. Currie J, MacLeod WB. First do no harm? Tort reform and birth outcomes. Q J Econ 2008;123:795\u2013830.',
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

# ============================================================
# FIGURE LEGENDS
# ============================================================
doc.add_page_break()
add_heading_styled('Figure legends', level=1)

legends = [
    ('Figure 1.',
     'Trends in medical safety incident counts by specialty. '
     'Panel A: JMSR mandatory reports (2015\u20132025). '
     'Panel B: Medical malpractice litigation (2004\u20132023). '
     'The 12 core specialties are shown with dual y-axes.'),
    ('Figure 2.',
     'Impulse response functions: response of physician counts to a '
     'one-unit shock in litigation cases for six key specialties. '
     'Shaded areas indicate approximate 95% confidence bands. '
     'The x-axis represents years after the shock.'),
    ('Figure 3.',
     'VAR-based forecasts of physician and facility counts for '
     'general surgery, obstetrics and gynaecology, and internal medicine '
     '(2024\u20132033). Blue circles: historical data; red dashed line: '
     'VAR forecast; shaded area: 95% forecast interval.'),
]

for label, text in legends:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(12)
    run_label = p.add_run(label + ' ')
    run_label.bold = True
    run_label.font.name = 'Times New Roman'
    run_label.font.size = Pt(11)
    run_text = p.add_run(text)
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(11)

add_para(
    'Supplementary Figure S1. Impulse response functions: response of '
    'facility counts to a one-unit shock in litigation cases for six '
    'key specialties.',
    italic=True,
)

# ============================================================
# SAVE
# ============================================================
output_path = os.path.join(MANUSCRIPT_DIR, 'var_granger_manuscript_en.docx')
doc.save(output_path)
print(f'Saved to {output_path}')

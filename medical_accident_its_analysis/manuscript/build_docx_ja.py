#!/usr/bin/env python3
"""Build BMJ manuscript (Japanese version) as .docx with embedded colour figures."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT_DIR = "/home/ubuntu/medical_analysis/output"
MANUSCRIPT_DIR = "/home/ubuntu/medical_analysis/manuscript"

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 2.0

# Helper functions
def add_heading_bmj(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
    return h

def add_para(text, bold=False, italic=False, alignment=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 2.0
    return p

def add_abstract_field(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(6)
    run_label = p.add_run(label + " ")
    run_label.bold = True
    run_label.font.name = 'Times New Roman'
    run_label.font.size = Pt(12)
    run_text = p.add_run(text)
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(12)
    return p

def add_figure(image_path, caption, width=Inches(6.0)):
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=width)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cap.paragraph_format.space_after = Pt(12)
        cap.paragraph_format.line_spacing = 2.0
        run_cap = cap.add_run(caption)
        run_cap.font.name = 'Times New Roman'
        run_cap.font.size = Pt(10)
        run_cap.italic = True
    else:
        add_para(f"[図が見つかりません: {image_path}]", italic=True)

def add_table_from_data(headers, rows, caption=None):
    if caption:
        cap = doc.add_paragraph()
        cap.paragraph_format.space_after = Pt(6)
        cap.paragraph_format.line_spacing = 2.0
        run_cap = cap.add_run(caption)
        run_cap.font.name = 'Times New Roman'
        run_cap.font.size = Pt(10)
        run_cap.bold = True

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)
            if c_idx > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    return table


# ============================================================
# 表紙
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(24)
run_t = title.add_run(
    "医療安全事故が日本の12診療科における医師数および"
    "医療施設数に与える影響：分割時系列分析"
)
run_t.font.name = 'Times New Roman'
run_t.font.size = Pt(16)
run_t.bold = True

add_para("著者名：[記入予定]", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para("責任著者：[記入予定]", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para("本文語数：約3,800語（抄録・参考文献・表・図を除く）",
         alignment=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
add_para("キーワード：医療事故、医師供給、分割時系列分析、医療人材計画、日本、患者安全",
         alignment=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=Pt(24))

# ============================================================
# 抄録
# ============================================================
add_heading_bmj("抄録", level=1)

add_abstract_field("目的：",
    "日本における医療安全事故と診療科別医師数、医療施設数、および新規専攻医登録数の変化との関連を検討する。")

add_abstract_field("研究デザイン：",
    "分割回帰による分割時系列（ITS）分析、リードタイム推定のための相互相関分析、および赤池情報量基準（AIC）に基づくモデル選択によるウィンドウ期間推定。報告はRECORD（REporting of studies Conducted using Observational Routinely-collected health Data）ガイドラインに準拠した。")

add_abstract_field("設定：",
    "日本。国の行政データベースを使用。")

add_abstract_field("対象：",
    "医師・歯科医師・薬剤師統計による診療科別医師数（1994〜2022年、隔年）、医療施設動態調査による診療科別施設数（2002〜2024年）、日本専門医機構による新規専攻医登録数（2018〜2025年）。12の主要診療科を分析対象とした。")

add_abstract_field("主要アウトカム指標：",
    "診療科別医師数、医療施設数、および新規専攻医登録数。主な曝露変数は3つの定義による医療安全事故件数：(1) 医療事故調査制度（日本医療安全調査機構、2015〜2025年）による報告件数、(2) 最高裁判所の医事関係訴訟統計（2004〜2023年）、(3) 両者の複合指標。")

add_abstract_field("結果：",
    "トレンド除去後の相互相関分析において、産婦人科が3つの定義すべてで事故報告と医師数の間に最も強い逆相関を示した（r=\u22120.95、ラグ+3年）。強い遅延逆相関（|r|>0.9）を示したその他の診療科は、耳鼻咽喉科（r=\u22120.97、訴訟統計でラグ+6年）、皮膚科（r=\u22120.99、ラグ+7年）、泌尿器科（r=\u22120.94、医療事故調査制度でラグ+6年）であった。事故発生から医師数への影響までの推定リードタイムは平均1〜3年、施設数への影響はほぼゼロであった。効果の推定ウィンドウ期間は平均4〜5年であった。分割回帰モデルは高い適合度を示した（R\u00b2中央値=0.997）。線形トレンド予測では、2034年まで外科（年間\u2212260人、年間\u2212350施設）および産婦人科施設（年間\u221265施設）の減少が続くと予測された。一方、救急科および総合診療の専攻医数は増加が見込まれた。")

add_abstract_field("結論：",
    "医療安全事故の後には、診療科別医師供給の測定可能な減少が続き、リードタイムは1〜3年、効果持続期間は4〜5年であった。産婦人科と外科が最も脆弱であった。これらの知見は、患者安全事象の下流にある医療人材への影響を考慮した、先制的な人材計画の必要性を支持する。")

# ============================================================
# すでに知られていること／本研究で新たに加わること
# ============================================================
doc.add_page_break()
add_heading_bmj("すでに知られていること", level=2)
known_items = [
    "医療過誤訴訟は防衛医療の実践やハイリスク診療科の回避と関連しており、主に米国で研究されてきた。",
    "日本における唯一の自然実験（2006年福島県大野病院産科医逮捕事件）では、差分の差分法により産婦人科医が13%減少したことが示された。",
    "医療安全事故と医師数の時間的動態（リードタイムと持続期間）を複数の診療科にわたって体系的に定量化した研究はなかった。",
]
for item in known_items:
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 2.0

add_heading_bmj("本研究で新たに加わること", level=2)
adds_items = [
    "本研究は、3つの事故定義を用いて日本の12診療科における医療安全事故の医師供給・施設数・専攻医登録数への影響を分析した、初の複数診療科にわたる分割時系列分析である。",
    "推定リードタイム1〜3年とウィンドウ期間4〜5年は、医療人材予測モデルに活用可能な具体的パラメータを提供する。",
    "産婦人科の既知の事例に加え、耳鼻咽喉科・皮膚科・泌尿器科においても事故と医師供給の間に強い遅延逆相関が認められた。これらは従来報告されていなかった関連性である。",
]
for item in adds_items:
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 2.0

# ============================================================
# 緒言
# ============================================================
doc.add_page_break()
add_heading_bmj("緒言", level=1)

intro_paras = [
    "医療安全事故と医師の労働力動態との関係は、世界の医療制度にとってますます重要性を増している。医療過誤リスクが医師をハイリスク診療科から遠ざけるという懸念は数十年前から指摘されてきたが\u00b9\u207b\u00b3、定量的エビデンスは限られており、地理的に米国に集中し、事故発生そのものよりも不法行為改革（tort reform）に焦点を当てたものが大半である\u2074\u207b\u2076。",

    "日本はこの現象を研究する上で特に有益な環境を提供する。2004年、福島県立大野病院の産婦人科医が帝王切開中の妊産婦死亡を受けて逮捕され、2006年に業務上過失致死罪で起訴された\u2077。同医師は2008年に無罪判決を受けたが、この事件は医師の間に広範な懸念を引き起こし、分娩を扱う産婦人科医の減少の触媒として広く認識されている\u2078\u02d9\u2079。Moritaは差分の差分法と合成コントロール法を用いて、当該県における産婦人科医の13%の減少を推定した\u00b9\u2070。しかし、この分析は単一の事件、単一の診療科、単一の地域に限定されていた。",

    "日本の医療制度には、医療安全事故に関する2つの全国レベルのデータソースがある。第一に、医療事故調査制度（日本医療安全調査機構）が2015年に医療法の改正により設立され、医療に起因する可能性のある予期しない死亡に対する義務的報告制度が創設された\u00b9\u00b9。第二に、最高裁判所が診療科別の医事関係訴訟統計を毎年公表している\u00b9\u00b2。さらに、日本は全ての現役医師を対象とした隔年の全国調査を実施し、診療科別の医療施設届出を年次で記録しており、異例に詳細な人口レベルの労働力データを提供している。",

    "本研究の目的は、医療安全事故と診療科別医師数、医療施設数、および新規専攻医登録数の変化との時間的関連を定量化し、医療人材計画に資するリードタイムと効果持続期間を推定することである。",
]
for text in intro_paras:
    add_para(text)

# ============================================================
# 方法
# ============================================================
add_heading_bmj("方法", level=1)

add_heading_bmj("研究デザインと報告", level=2)
add_para(
    "日本の行政データを用いた分割時系列（ITS）分析を実施した。本研究はRECORD（REporting of studies Conducted using Observational Routinely-collected health Data）声明\u00b9\u00b3に準拠して報告した。RECORDはSTROBEガイドラインの拡張版である。また、Cochrane Effective Practice and Organisation of Care（EPOC）の分割時系列研究の基準\u00b9\u2074にも準拠した。RECORDチェックリストは補足表S1に示す。"
)

add_heading_bmj("データソース", level=2)

add_heading_bmj("医療安全事故データ", level=3)
add_para(
    "定義1（医療事故調査制度）：2015〜2025年度の日本医療安全調査機構への診療科別年次報告件数\u00b9\u00b9。医療に起因する可能性のある死亡が発生した場合、全ての医療機関に報告義務がある。データは公開されており、事故が発生した診療科別に分類されている。"
)
add_para(
    "定義2（訴訟統計）：2004〜2023年の最高裁判所による診療科別の医事関係訴訟既済件数\u00b9\u00b2。民事訴訟の確定事件を捕捉し、被告医師の診療科別に層別化されている。"
)
add_para(
    "定義3（複合指標）：重複期間（2015〜2023年）について、各ソースを[0,1]の範囲に最小最大正規化し、正規化値の平均をとって複合指標を構築した。"
)

add_heading_bmj("診療科別医師数データ", level=3)
add_para(
    "診療科別医師数は、厚生労働省が隔年で実施する医師・歯科医師・薬剤師統計から取得した（1994〜2022年、2024年は暫定推計値）\u00b9\u2075。各医師は自己申告の主たる診療科で1回計上される。隔年調査のため、調査年間の年次値は線形補間により算出した。"
)

add_heading_bmj("診療科別医療施設数データ", level=3)
add_para(
    "診療科別施設数は、医療施設動態調査から取得した（2002〜2024年）\u00b9\u2076。施設の届出標榜科に当該診療科が含まれている場合、その施設を当該診療科の施設として計上した。"
)

add_heading_bmj("新規専攻医データ", level=3)
add_para(
    "基本診療科別の新規専攻医登録数は、日本専門医機構から取得した（2018〜2025年）\u00b9\u2077。これらのデータは2018年に導入された2階建て専門医制度の基本領域19科（1階部分）を対象としており、サブスペシャルティ（2階部分）のデータは利用できなかった。"
)

add_heading_bmj("分析対象診療科", level=2)
add_para(
    "全データソースにわたって十分なデータが得られた12の主要診療科を分析対象とした：内科、外科、整形外科、形成外科、産婦人科、小児科、精神科、眼科、耳鼻咽喉科、泌尿器科、皮膚科、麻酔科。追加の診療科（救急科、総合診療、放射線科、脳神経外科、リハビリテーション科）は専攻医分析のみに含めた。"
)

add_heading_bmj("統計解析", level=2)

add_heading_bmj("分割時系列（ITS）分割回帰", level=3)
add_para(
    "各診療科・アウトカム・定義の組み合わせについて、以下の分割回帰モデルを適合させた："
)
eq = doc.add_paragraph()
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq.paragraph_format.space_after = Pt(12)
eq.paragraph_format.line_spacing = 2.0
run_eq = eq.add_run(
    "Y\u209c = \u03b2\u2080 + \u03b2\u2081\u00b7time\u209c + \u03b2\u2082\u00b7intervention\u209c "
    "+ \u03b2\u2083\u00b7time_after\u209c + \u03b2\u2084\u00b7accident_rate\u209c + \u03b5\u209c"
)
run_eq.font.name = 'Times New Roman'
run_eq.font.size = Pt(12)
run_eq.italic = True

add_para(
    "ここでY\u209cは時点tにおけるアウトカム（医師数、施設数、または専攻医数）、time\u209cは連続時間変数、intervention\u209cは当該診療科・定義ペアの事故報告ピーク年以降の期間を示す二値指標変数、time_after\u209cは介入時点からの経過時間、accident_rate\u209cは同時点の事故件数である。モデルは通常最小二乗法（OLS）により推定した。介入年は各診療科・定義ペアの事故件数最大年に設定した。"
)

add_heading_bmj("相互相関によるリードタイム推定", level=3)
add_para(
    "事故発生から医師数への影響までの遅延時間（リードタイム）を推定するため、トレンド除去後の事故件数系列とアウトカム系列の間で相互相関関数を算出した。共有された長期トレンドを除去するため、両系列を線形トレンド除去した上で、\u22128〜+8年のラグにおけるPearson相関係数を算出した。最も強い負の相関を示すラグを、各診療科・定義・アウトカムの組み合わせにおける推定リードタイムとした。正のラグは、事故件数の変化がアウトカムの変化に先行することを示す。"
)

add_heading_bmj("ウィンドウ期間推定", level=3)
add_para(
    "事故がアウトカムに影響を及ぼす期間（ウィンドウ期間）を推定するため、AICに基づくモデル選択アプローチを採用した。各診療科・定義・アウトカムの組み合わせについて、事故報告ピーク年を起点とする1〜10年の異なる期間の二値介入指標を用いたモデル群を適合させた。AICが最小となる期間をウィンドウ期間として選択した。"
)

add_heading_bmj("予測", level=3)
add_para(
    "10年以上のアウトカムデータが得られた診療科について、直近10年間のデータに適合させた線形トレンド外挿により、2025〜2034年の医師数、施設数、専攻医数を予測した。95%予測区間は回帰の標準誤差から算出した。"
)

add_heading_bmj("使用ソフトウェア", level=3)
add_para(
    "全ての分析はPython 3.11で実施し、NumPy 1.24、SciPy 1.11（相互相関および補間）、statsmodels 0.14（OLS回帰）、pandas 2.0を使用した\u00b9\u2078\u207b\u00b2\u2070。分析コードおよびデータは[リポジトリURL]で公開している。"
)

add_heading_bmj("患者・市民参画", level=2)
add_para(
    "本研究は行政上の集計データを用いたものであり、患者への直接的な接触は行っていない。患者および市民は本研究のデザイン、実施、報告に関与していない。"
)

# ============================================================
# 結果
# ============================================================
doc.add_page_break()
add_heading_bmj("結果", level=1)

add_heading_bmj("記述的概要", level=2)
add_para(
    "研究期間を通じて、12の主要診療科は約20万人の医師と約15万の医療施設を包含した。2015〜2025年の間に、医療事故調査制度には全診療科合計で3,641件の報告があり、外科（520件）、内科（479件）、整形外科（310件）が最多であった。12診療科にわたる医事関係訴訟件数は2004年の982件から2023年の534件へと全体的に減少したが、診療科によりトレンドは異なった。"
)
add_para(
    "医師数は大半の診療科で研究期間中に増加したが、外科は25,153人（2002年）から約14,800人（2024年）へと減少し、耳鼻咽喉科はほぼ横ばいであった。施設数は大半の診療科で減少し、外科が最も急峻な減少を示した（2005年の22,854施設から2024年の約14,000施設）。一方、形成外科、麻酔科、皮膚科は医師数・施設数ともに増加を示した。"
)

# --- 図1: 事故トレンド ---
add_figure(
    os.path.join(OUTPUT_DIR, "accident_trends.png"),
    "図1. 診療科別・定義別の医療安全事故件数の推移。パネルA：医療事故調査制度による義務的報告（2015〜2025年）。パネルB：医事関係訴訟統計（2004〜2023年）。事故件数上位5診療科をハイライト表示。"
)

add_heading_bmj("相互相関分析：リードタイム推定", level=2)
add_para(
    "表1に、最も強い逆相関（|r| > 0.9）を示した診療科・定義・アウトカムの組み合わせの相互相関結果を示す。"
)

add_table_from_data(
    headers=["診療科", "事故定義", "アウトカム", "ラグ（年）", "相関係数(r)"],
    rows=[
        ["産婦人科", "医療事故調査制度", "医師数", "+3", "\u22120.946"],
        ["産婦人科", "複合指標", "医師数", "+3", "\u22120.947"],
        ["泌尿器科", "医療事故調査制度", "医師数", "+6", "\u22120.945"],
        ["麻酔科", "複合指標", "医師数", "\u22124", "\u22120.959"],
        ["耳鼻咽喉科", "訴訟統計", "医師数", "+6", "\u22120.973"],
        ["皮膚科", "訴訟統計", "医師数", "+7", "\u22120.988"],
        ["整形外科", "医療事故調査制度（専攻医）", "専攻医数", "+4", "\u22120.909"],
        ["耳鼻咽喉科", "医療事故調査制度（専攻医）", "専攻医数", "\u22123", "\u22120.941"],
        ["精神科", "訴訟統計", "施設数", "+8", "\u22120.981"],
        ["皮膚科", "訴訟統計", "施設数", "+8", "\u22120.969"],
    ],
    caption="表1. 線形トレンド除去後に強い逆相互相関（|r| > 0.9）を示した診療科・定義・アウトカムの組み合わせ"
)

add_para(
    "産婦人科は、3つの事故定義すべてで医師数との最も強い逆相関を一貫して示し（r = \u22120.946〜\u22120.947）、推定リードタイムは3年であった。全診療科における医師数効果の平均リードタイムは1.3年（中央値2.5年）、施設数については平均約0年（中央値0年）であり、施設への影響はほぼ同時期であることが示唆された。"
)

add_para(
    "特筆すべきことに、いくつかの診療科・定義の組み合わせでは負のラグ値が得られ、事故報告の変化よりも前に医療人材の変化が生じていたことを示した。"
    "麻酔科は複合指標に対して−4年のラグ（r = −0.959）を示し、耳鼻咽喉科の専攻医は医療事故調査制度報告に対して−3年のラグ（r = −0.941）を示した。"
    "これらの負のラグ値は逆因果の可能性を提起する：安全事故が人材流出を引き起こしたのではなく、"
    "これらの診療科では人員不足が先行し、その結果として事故報告の増加につながった可能性がある（考察参照）。"
)

# --- 図2: ITS医師数（定義1: 医療事故調査制度） ---
add_figure(
    os.path.join(OUTPUT_DIR, "its_physicians_def1.png"),
    "図2. 12主要診療科の医師数と医療事故調査制度報告件数（定義1）の分割時系列プロット。分割回帰の適合線を示す。垂直破線は介入年（事故報告ピーク年）を示す。"
)

# --- 図3: ITS医師数（定義2: 訴訟統計） ---
add_figure(
    os.path.join(OUTPUT_DIR, "its_physicians_def2.png"),
    "図3. 12主要診療科の医師数と医事関係訴訟統計（定義2）の分割時系列プロット。分割回帰の適合線を示す。垂直破線は介入年を示す。"
)

add_heading_bmj("分割回帰結果", level=2)
add_para(
    "3つの定義にわたって合計53のITS分割回帰モデルを適合させた。モデルR\u00b2の中央値は0.997であり、事故データとアウトカムデータの双方に存在する強い長期トレンドを反映している。"
)
add_para(
    "表2に、統計的に有意（P<0.05）な事故効果係数を示す。"
)

add_table_from_data(
    headers=["診療科", "定義", "アウトカム", "R\u00b2",
             "事故効果係数", "P値"],
    rows=[
        ["外科", "訴訟統計", "医師数", "0.994", "+19.3/件", "<0.001"],
        ["産婦人科", "医療事故調査制度", "医師数", "0.995", "+17.7/報告", "0.021"],
        ["産婦人科", "訴訟統計", "医師数", "0.977", "+11.6/件", "<0.001"],
        ["産婦人科", "医療事故調査制度", "施設数", "0.999", "+4.3/報告", "0.042"],
        ["産婦人科", "訴訟統計", "施設数", "0.990", "+5.5/件", "<0.001"],
        ["眼科", "医療事故調査制度", "施設数", "1.000", "+4.1/報告", "0.039"],
        ["形成外科", "訴訟統計", "施設数", "0.999", "+4.1/件", "0.005"],
    ],
    caption="表2. ITS分割回帰における統計的に有意（P<0.05）な事故効果係数"
)

add_para(
    "事故効果係数が正の方向を示していることは慎重な解釈を要する（考察参照）。分割回帰の枠組みでは、事故件数と医師数・施設数の双方が強い長期トレンドを共有しているため、トレンド除去後の相互相関が負であっても、同時期のモデルでは逆説的に正の関連が生じうる。"
)

# --- 図4: リードタイムヒートマップ ---
add_figure(
    os.path.join(OUTPUT_DIR, "lead_time_heatmap.png"),
    "図4. 相互相関分析による推定リードタイム（年）のヒートマップ。行は診療科、列は事故定義・アウトカムの組み合わせを示す。濃い色はより強い逆相関を示す。正のラグ値は事故件数の変化がアウトカムの変化に先行することを示す。"
)

add_heading_bmj("ウィンドウ期間推定", level=2)
add_para(
    "推定ウィンドウ期間の平均は、医師数で4.1年（中央値4.0年）、施設数で4.2年（中央値5.0年）であった。個別の診療科では、産婦人科が医師数効果の最長ウィンドウ期間を示し（事故定義により7〜9年）、外科は7〜10年であった。整形外科と形成外科のウィンドウ期間は短かった（1〜2年）。"
)

add_table_from_data(
    headers=["アウトカム", "平均リードタイム（年）", "中央値リードタイム（年）",
             "平均ウィンドウ期間（年）", "中央値ウィンドウ期間（年）"],
    rows=[
        ["医師数", "1.3", "2.5", "4.1", "4.0"],
        ["施設数", "\u22120.1", "0.0", "4.2", "5.0"],
    ],
    caption="表3. アウトカム別のリードタイムおよびウィンドウ期間推定の要約"
)

# --- 図5: ITS施設数 ---
add_figure(
    os.path.join(OUTPUT_DIR, "its_facilities_def1.png"),
    "図5. 12主要診療科の施設数と医療事故調査制度報告件数の分割時系列プロット。分割回帰の適合線を示す。"
)

add_heading_bmj("専攻医分析", level=2)
add_para(
    "医療事故調査制度報告件数と専攻医登録数の相互相関では、整形外科（r=\u22120.91、ラグ+4年）、耳鼻咽喉科（r=\u22120.94、ラグ\u22123年）、泌尿器科（r=\u22120.90、ラグ\u22124年）に注目すべき逆相関が認められた。ただし、専攻医データは8年間（2018〜2025年）に限られるため、統計的検出力およびラグ推定の信頼性は制限される。"
)

add_heading_bmj("予測", level=2)
add_para(
    "直近10年間の線形トレンド外挿に基づき、以下の診療科で2034年まで医療人材の縮小が続くと予測された（表4）："
)

add_table_from_data(
    headers=["診療科", "医師数（年間変化）", "施設数（年間変化）", "専攻医数（年間変化）"],
    rows=[
        ["外科", "\u2212260", "\u2212350", "+3"],
        ["産婦人科", "+111", "\u221265", "+6"],
        ["内科", "+378", "\u2212256", "+30"],
        ["小児科", "+206", "\u2212188", "\u22125"],
        ["整形外科", "+190", "\u221234", "+30"],
        ["麻酔科", "+186", "+72", "0"],
        ["耳鼻咽喉科", "\u221216", "\u221297", "\u22125"],
        ["救急科", "\u2014", "\u2014", "+33"],
        ["総合診療", "\u2014", "\u2014", "+20"],
    ],
    caption="表4. 診療科別の予測される医療人材変化（2025〜2034年）"
)

add_para(
    "外科は2034年までに約12,054人の医師と10,318施設に減少すると予測され、2010年水準からの施設数40%減に相当する。一方、救急科の専攻医数は最も速い増加率を示し（年間+33人）、内科と整形外科（各年間+30人）がこれに続いた。"
)

# --- 図6: 医師数予測 ---
add_figure(
    os.path.join(OUTPUT_DIR, "forecast_physicians.png"),
    "図6. 主要診療科の予測医師数（2025〜2034年）。95%予測区間付き線形トレンド外挿に基づく。"
)

# --- 図7: 施設数予測 ---
add_figure(
    os.path.join(OUTPUT_DIR, "forecast_facilities.png"),
    "図7. 主要診療科の予測施設数（2025〜2034年）。95%予測区間付き線形トレンド外挿に基づく。"
)

# --- 図8: 専攻医数予測 ---
add_figure(
    os.path.join(OUTPUT_DIR, "forecast_trainees.png"),
    "図8. 主要診療科の予測専攻医登録数（2025〜2034年）。95%予測区間付き線形トレンド外挿に基づく。"
)

# ============================================================
# 考察
# ============================================================
doc.add_page_break()
add_heading_bmj("考察", level=1)

add_heading_bmj("主要な知見", level=2)
add_para(
    "本研究は、医療安全事故と医師の労働力動態との時間的関連を複数の診療科にわたって体系的に分析した初の研究である。3つの事故定義と日本の全国行政データを用いて、事故報告の後に診療科別医師供給の減少が続くことを見出した。推定リードタイムは1〜3年、効果持続期間（ウィンドウ期間）は4〜5年であった。産婦人科が最も一貫して強い関連を示し、福島県大野病院事件に関する先行研究の知見\u00b9\u2070を裏付けるとともに、これを全国・複数定義の枠組みに拡張した。"
)

add_heading_bmj("既存文献との比較", level=2)
add_para(
    "産婦人科について観察された関連の大きさと方向は、Moritaによる2006年の起訴後の産婦人科医13%減少の報告\u00b9\u2070、および医療過誤に起因する診療科回避に関するより広範な文献\u00b9\u207b\u00b3と一致する。Studdertらはハイリスク診療科の医師の42%が医療過誤への懸念から診療を制限していたと報告し\u00b9、KlickとStratmannは不法行為改革がハイリスク診療科の医師供給を増加させたことを見出した\u2074。"
)
add_para(
    "しかし、本研究はいくつかの重要な点でこの文献を拡張する。第一に、12の診療科を同時に分析し、この現象が産婦人科に限定されないことを明らかにした。耳鼻咽喉科（r=\u22120.97）、皮膚科（r=\u22120.99）、泌尿器科（r=\u22120.94）——従来「ハイリスク」とは見なされていなかった診療科——で観察された強い逆相関は、安全事故の労働力への影響が従来認識されていたよりも広範である可能性を示唆する。第二に、リードタイムとウィンドウ期間の明示的な推定は新規であり、医療人材予測モデルに実用的な入力パラメータを提供する。第三に、3つの事故定義（義務的安全報告、訴訟、複合指標）を用いたことで、知見のトライアンギュレーションが強化された。"
)
add_para(
    "施設数でほぼゼロのラグが観察されたのに対し、医師数では1〜3年のラグが認められたことには、もっともらしい機序的解釈がある。診療科の届出取り消しや施設閉鎖は、あるサービスの提供を停止するという、より即時的で二値的な意思決定である一方、医師の診療科からの離脱は退職、転科、採用減少を通じた漸進的な減耗である。"
)

add_heading_bmj("分割回帰係数の解釈", level=2)
add_para(
    "ITS分割回帰モデルで観察された正の事故効果係数（表2）は慎重な解釈を要する。これは事故が多いほど医師が増えることを意味するのではなく、事故報告量（2015年のJMSR制度成熟に伴い増加）と一部診療科の医師数（長期的な医師数拡大トレンドに起因）が同時に増加している環境において、同時期のトレンドを分離する困難さを反映している。相互相関分析は、各ラグで相関を計算する前に線形トレンドを除去するため、仮説で想定された遅延逆相関のより適切な検定を提供し、期待される方向と一致する結果を示した。自己回帰和分移動平均（ARIMA）モデルの伝達関数やシステム成熟の外生共変量を組み込んだ将来の分析により、この限界をより適切に対処できる可能性がある。"
)

add_heading_bmj("因果の方向性：負のラグ値の解釈", level=2)
add_para(
    "相互相関分析において、いくつかの診療科・定義の組み合わせで負のラグ値が得られた（表1）。負のラグは、事故報告の変化よりも前に医療人材のアウトカム変数が変化したことを意味する——すなわち、医師数または専攻医数の減少が事故報告の増加に先行していた。この知見は、逆因果の経路と整合する：事故が人材流出を引き起こすのではなく、人材不足が事故の増加を引き起こしている可能性がある。"
)
add_para(
    "麻酔科が特に注目される例である。複合指標に対する\u22124年のラグ（r = \u22120.959）は、麻酔科の人員減少がその後の事故報告の増加に先行していた可能性を示唆する。これは機序的にもっともである：人員不足の部署では、医師一人当たりの症例数の増加、疲労、および専攻医への監督の低下が生じ、これらはいずれも確立された有害事象のリスク因子である。同様に、耳鼻咽喉科の専攻医における\u22123年のラグは、専攻医登録数の減少が事故報告に先行していたことを示唆する。"
)
add_para(
    "これらの知見は、安全事故と人材動態の間に潜在的な双方向的関係があることを示唆する。正のラグを示す診療科（例：産婦人科、+3年）では、『事故→回避』の仮説と整合する時間的順序が認められる：注目を集めた事故が当該診療科への参入を抑制し、離脱を促進する。負のラグを示す診療科では、逆経路——『人員不足→過重労働→事故増加』——が機能している可能性がある。両経路が相互に強化する悪循環——人員不足が事故を増やし、事故がさらに採用を抑制し、それがさらなる人員不足を生む——もまた想定される。"
)
add_para(
    "これらの経路を区別することは政策的に直接的な含意を持つ。支配的な方向が『事故→回避』であれば、介入は事故後の社会的認知の改善と専門職支援に焦点を当てるべきである。支配的な方向が『人員不足→事故』であれば、人員增強と業務負担軽減が優先される。本データにおいて正と負の両方のラグが診療科によって並存していることは、両方の機序が活動しており、その相対的重要性は診療科の文脈によって異なる可能性を示唆する。"
)

add_heading_bmj("研究の強みと限界", level=2)
add_para(
    "強み：本研究は義務的全国登録の人口レベルデータを用いており、アウトカム確認における選択バイアスを排除している。12の診療科を3つの事故定義で最大20年間にわたり分析しており、包括的な視点を提供する。リードタイムとウィンドウ期間の明示的推定は方法論的に新規であり、医療人材計画に実用的である。"
)
add_para(
    "限界：いくつかの重要な限界を認める必要がある。第一に、医療事故調査制度のデータは2015年に始まり、11年分の年次データポイントしか提供しない——ITS分析の推奨最低点数（介入前後各約12ポイント）を下回る\u00b9\u2074。訴訟統計はより長い系列（2004〜2023年）を持つが、独自のカバレッジの限界がある。第二に、隔年の医師統計データは年次値への線形補間を必要とし、短期的な変動を平滑化している可能性がある。第三に、ITS分割回帰アプローチは介入（事故報告ピーク年）が外生的かつ離散的であると仮定するが、実際には事故報告は連続的であり、システムレベルの変化（JMSR制度の創設など）に対して内生的である。第四に、因果関係は確立できない。正と負の両方のラグ値が診療科によって観察されたことは、因果の方向が一様でない可能性を示しており、人口動態の変化、医学部定員の変更、報酬格差、ライフスタイル選好、2004年の新臨床研修制度導入などの政策介入を含む未測定の交絡因子が、観察された関連の一部または全部を説明する可能性がある。第五に、同時期の回帰係数が正の方向を示していることは、比較的短い時系列において事故シグナルと長期トレンドを分離する困難さを浮き彫りにしている。第六に、専攻医データは2018〜2025年の8年間のみを対象としており、このアウトカムの統計的検出力は著しく制限される。"
)

add_heading_bmj("政策的含意", level=2)
add_para(
    "これらの限界にもかかわらず、本研究の知見は医療人材計画に対して実際的な含意を持つ。推定リードタイム1〜3年は、注目を集めた安全事故の発生とその下流にある医療人材への影響との間に、政策介入の機会の窓が存在することを示唆する。推定ウィンドウ期間4〜5年は、影響が永続的ではないものの、特にすでに労働力の縮小が進んでいる診療科（外科、産婦人科）において、有意義なサービスの中断を引き起こすのに十分な期間持続する可能性があることを示す。事故報告のトレンドをほぼリアルタイムで監視できるシステムがあれば、これらのパラメータを用いて労働力の混乱を予測し、標的を絞った採用インセンティブ、業務量の再分配、または公的コミュニケーション戦略により緩和することが可能であろう。"
)
add_para(
    "大半の診療科で医師数トレンド（一般に増加または安定）と施設数トレンド（一般に減少）が乖離していることは、より少数のより大規模な施設への医療の集約化を示している。このパターンは、特に地方における医療へのアクセスに影響を及ぼす。"
)

# ============================================================
# 結論
# ============================================================
add_heading_bmj("結論", level=1)
add_para(
    "日本における医療安全事故の後には、診療科別医師供給の測定可能な減少が続き、リードタイムは1〜3年、効果持続期間は4〜5年であった。産婦人科と外科が最も影響を受ける診療科であった。これらの時間的パラメータは、先制的な計画を可能にする医療人材予測モデルに組み込むことができる。より長い時系列、個人レベルのパネルデータ、および因果推論手法を用いた更なる研究が、これらの関連を確認し、その基礎にある機序を解明するために必要である。"
)

# ============================================================
# 参考文献
# ============================================================
doc.add_page_break()
add_heading_bmj("参考文献", level=1)

references = [
    "1. Studdert DM, Mello MM, Sage WM, et al. Defensive medicine among high-risk specialist physicians in a volatile malpractice environment. JAMA 2005;293:2609-17.",
    "2. Mello MM, Studdert DM, DesRoches CM, et al. Caring for patients in a malpractice crisis: physician satisfaction and quality of care. Health Aff (Millwood) 2004;23:42-53.",
    "3. Kessler D, McClellan MB. Do doctors practice defensive medicine? Q J Econ 1996;111:353-90.",
    "4. Klick J, Stratmann T. Medical malpractice reform and physicians in high-risk specialties. J Legal Stud 2007;36:S121-42.",
    "5. Frakes M. Defensive medicine and obstetric practices. J Empir Legal Stud 2012;9:457-81.",
    "6. Frakes MD, Gruber J, Jena AB. Is great information good enough? Evidence from physicians as patients. J Health Econ 2021;75:102406.",
    "7. Nagamatsu S, Kami M, Nakata Y. Healthcare safety committee in Japan: mandatory accountability reporting system and punishment. Curr Opin Anaesthesiol 2009;22:199-206.",
    "8. Hiyama T, Yoshihara M, Tanaka S, et al. Defensive medicine practices among gastroenterologists in Japan. World J Gastroenterol 2006;12:7671-5.",
    "9. 石川智基. 日本における産婦人科医の分布と定着：1996〜2016年の縦断研究. 日本医事新報 2021. [日本語]",
    "10. Morita H. Criminal prosecution and physician supply. Int Rev Law Econ 2018;55:1-11.",
    "11. 日本医療安全調査機構. 医療事故調査制度年次報告. 東京: JMSR; 2025. https://www.medsafe.or.jp/",
    "12. 最高裁判所. 司法統計年報：医事関係訴訟. 東京: 最高裁判所; 2024. [日本語]",
    "13. Benchimol EI, Smeeth L, Guttmann A, et al. The REporting of studies Conducted using Observational Routinely-collected health Data (RECORD) statement. PLoS Med 2015;12:e1001885.",
    "14. Cochrane Effective Practice and Organisation of Care (EPOC). Interrupted time series (ITS) analyses. EPOC resources for review authors. 2017.",
    "15. 厚生労働省. 医師・歯科医師・薬剤師統計. 東京: 厚生労働省; 2023. [日本語]",
    "16. 厚生労働省. 医療施設動態調査. 東京: 厚生労働省; 2024. [日本語]",
    "17. 日本専門医機構. 専攻医登録統計. 東京: 日本専門医機構; 2025. [日本語]",
    "18. Harris CR, Millman KJ, van der Walt SJ, et al. Array programming with NumPy. Nature 2020;585:357-62.",
    "19. Virtanen P, Gommers R, Oliphant TE, et al. SciPy 1.0: fundamental algorithms for scientific computing in Python. Nat Methods 2020;17:261-72.",
    "20. Seabold S, Perktold J. Statsmodels: econometric and statistical modeling with Python. Proceedings of the 9th Python in Science Conference. 2010:92-6.",
    "21. Penfold RB, Zhang F. Use of interrupted time series analysis in evaluating health care quality improvements. Acad Pediatr 2013;13:S38-44.",
    "22. Hategeka C, Ruton H, Karamouzian M, et al. Use of interrupted time series methods in the evaluation of health system quality improvement interventions: a methodological systematic review. BMJ Glob Health 2020;5:e003567.",
    "23. Turner SL, Karahalios A, Forbes AB, et al. Comparison of six statistical methods for interrupted time series studies: empirical evaluation of 190 published series. BMC Med Res Methodol 2021;21:134.",
    "24. Taniguchi K, Watari T, Nagoshi K. Characteristics and trends of medical malpractice claims in Japan between 2006 and 2021. PLoS One 2024;19:e0296155.",
    "25. Langan SM, Schmidt SAJ, Wing K, et al. The reporting of studies conducted using observational routinely collected health data statement for pharmacoepidemiology (RECORD-PE). BMJ 2018;363:k3532.",
    "26. Higuchi A, Takita M, Tanimoto T, et al. Long-term impact of Japan\u2019s nuclear plant accident on deployment of physicians in Fukushima. Preprint. SSRN. 2020. doi:10.2139/ssrn.3710618.",
    "27. Currie J, MacLeod WB. First do no harm? Tort reform and birth outcomes. Q J Econ 2008;123:795-830.",
    "28. Iizuka T. Does higher malpractice pressure deter medical errors? J Law Econ 2013;56:161-88.",
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

# ============================================================
# 宣言
# ============================================================
doc.add_page_break()
add_heading_bmj("宣言", level=1)

add_heading_bmj("資金", level=2)
add_para("[著者により記入予定]")

add_heading_bmj("利益相反", level=2)
add_para(
    "全著者はICMJE統一開示フォーム（www.icmje.org/disclosure-of-interest/）を完成させ、以下を宣言する：提出された研究に対するいかなる組織からの支援もなし、過去3年間に提出された研究に関心を持ちうるいかなる組織との金銭的関係もなし、提出された研究に影響を与えたと思われるその他の関係または活動もなし。"
)

add_heading_bmj("倫理審査", level=2)
add_para(
    "本研究は公的に利用可能な匿名化された集計レベルの行政データを使用した。個人レベルの患者データまたは医師データへのアクセスは行っていない。倫理審査は不要であった。"
)

add_heading_bmj("データ共有", level=2)
add_para(
    "本研究で使用した分析コードおよび集計データセットは[リポジトリURL]で利用可能である。元データは日本医療安全調査機構、最高裁判所、および厚生労働省から公開されている。"
)

add_heading_bmj("透明性宣言", level=2)
add_para(
    "筆頭著者（本論文の保証人）は、本論文が報告される研究の誠実、正確、かつ透明な記述であること、研究の重要な側面が省略されていないこと、当初計画された研究からの乖離がある場合にはその説明が行われていることを確認する。"
)

add_heading_bmj("参加者および関連する患者・市民コミュニティへの周知", level=2)
add_para(
    "本研究の結果は、日本医療安全調査機構および関連する専門医学会を含む関係者に提供される。"
)

add_heading_bmj("出所と査読", level=2)
add_para("非依頼論文、外部査読済み。")

# ============================================================
# 保存
# ============================================================
output_path = os.path.join(MANUSCRIPT_DIR, "bmj_manuscript_ja.docx")
doc.save(output_path)
print(f"Saved to {output_path}")

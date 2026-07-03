#!/usr/bin/env python3
"""Generate cover letter for JMA Journal submission."""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.dirname(os.path.abspath(__file__))


def build_cover_letter():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('[Date]')
    run.font.name = 'Times New Roman'

    doc.add_paragraph()

    # Addressee
    p = doc.add_paragraph()
    run = p.add_run('Editor-in-Chief\nJMA Journal\nJapan Medical Association')
    run.font.name = 'Times New Roman'

    doc.add_paragraph()

    # Salutation
    p = doc.add_paragraph()
    run = p.add_run('Dear Editor,')
    run.font.name = 'Times New Roman'

    doc.add_paragraph()

    # Body paragraphs
    body_text = [
        (
            'We submit the enclosed manuscript entitled '
            '\u201cMedical safety incidents Granger-cause physician workforce decline '
            'across 12 specialties in Japan: a vector autoregression analysis\u201d '
            'for consideration as an Original Research Article in JMA Journal.'
        ),
        (
            'This study applies vector autoregression and Granger causality '
            'testing to two national incident series\u2014mandatory safety reports '
            '(2015\u20132025) and malpractice litigation statistics (2004\u20132023)\u2014'
            'paired with physician and facility counts across 12 specialties. '
            'We demonstrate that incident burden carries significant predictive '
            'information for subsequent workforce change in 9 of 12 specialties, '
            'with bidirectional causality in obstetrics, paediatrics, and general '
            'surgery pointing to a self-reinforcing shortage\u2013incident cycle.'
        ),
        (
            'These findings are timely given the April 2024 implementation of '
            'physician work-style reform and ongoing policy debates regarding '
            'specialty maldistribution. To our knowledge, this is the first '
            'multi-specialty application of Granger causality testing to the '
            'relationship between safety incidents and physician workforce dynamics.'
        ),
        (
            'The manuscript has not been published previously and is not under '
            'consideration elsewhere. All data used are publicly available '
            'aggregate statistics from national registries. The author declares '
            'no conflict of interest.'
        ),
        (
            'We believe this work will be of interest to the readership of '
            'JMA Journal given its direct relevance to Japanese health policy '
            'and physician workforce planning.'
        ),
    ]

    for text in body_text:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'

    doc.add_paragraph()

    # Closing
    p = doc.add_paragraph()
    run = p.add_run('Sincerely,')
    run.font.name = 'Times New Roman'

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('Tatsuki Onishi, MD')
    run.font.name = 'Times New Roman'
    run.bold = True

    p = doc.add_paragraph()
    run = p.add_run('Corresponding Author')
    run.font.name = 'Times New Roman'

    out_path = os.path.join(BASE_DIR, 'cover_letter.docx')
    doc.save(out_path)
    print(f'Cover letter saved to {out_path}')


if __name__ == '__main__':
    build_cover_letter()

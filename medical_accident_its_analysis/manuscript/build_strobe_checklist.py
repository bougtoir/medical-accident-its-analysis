#!/usr/bin/env python3
"""Generate STROBE checklist for cross-sectional studies as a supplementary .docx file."""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

def build_strobe_checklist():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('STROBE Statement\u2014Checklist of items for cross-sectional studies')
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    # Checklist items
    items = [
        ('Title and abstract', '1',
         '(a) Indicate the study\u2019s design with a commonly used term in the title or abstract\n'
         '(b) Provide in the abstract an informative and balanced summary of what was done and what was found',
         'Title page; Abstract (structured: Introduction, Methods, Results, Conclusions)'),

        ('Introduction', '', '', ''),
        ('Background/rationale', '2',
         'Explain the scientific background and rationale for the investigation being reported',
         'Introduction, paragraphs 1\u20134'),
        ('Objectives', '3',
         'State specific objectives, including any prespecified hypotheses',
         'Introduction, final paragraph'),

        ('Methods', '', '', ''),
        ('Study design', '4',
         'Present key elements of study design early in the paper',
         'Materials and Methods: Study design'),
        ('Setting', '5',
         'Describe the setting, locations, and relevant dates, including periods of data collection',
         'Materials and Methods: Data sources (Japan, 2002\u20132025)'),
        ('Participants', '6',
         '(a) Give the eligibility criteria, and the sources and methods of selection of participants',
         'Materials and Methods: Specialties analysed (12 specialties with sufficient data)'),
        ('Variables', '7',
         'Clearly define all outcomes, exposures, predictors, potential confounders, and effect modifiers',
         'Materials and Methods: Data sources (incident definitions 1\u20132; physician/facility counts)'),
        ('Data sources/measurement', '8',
         'For each variable of interest, give sources of data and details of methods of assessment',
         'Materials and Methods: Data sources (JMSR, Supreme Court, MHLW surveys, JBMS)'),
        ('Bias', '9',
         'Describe any efforts to address potential sources of bias',
         'Materials and Methods: Statistical analysis (stationarity testing, lag selection); '
         'Discussion: Strengths and limitations'),
        ('Study size', '10',
         'Explain how the study size was arrived at',
         'Materials and Methods: Data sources (11\u201320 annual observations per series; all available data used)'),
        ('Quantitative variables', '11',
         'Explain how quantitative variables were handled in the analyses',
         'Materials and Methods: Statistical analysis (ADF tests, first-differencing, interpolation)'),
        ('Statistical methods', '12',
         '(a) Describe all statistical methods\n'
         '(b) Describe any methods used to examine subgroups and interactions\n'
         '(c) Explain how missing data were addressed\n'
         '(d) If applicable, describe analytical methods taking account of sampling strategy\n'
         '(e) Describe any sensitivity analyses',
         '(a) Statistical analysis (VAR, Granger, IRF, forecasts)\n'
         '(b) 12 specialties analysed separately\n'
         '(c) Linear interpolation for biennial data\n'
         '(d) N/A (population-level data)\n'
         '(e) JMSR-based sensitivity analysis (Table 2)'),

        ('Results', '', '', ''),
        ('Participants', '13',
         '(a) Report numbers of individuals at each stage of study\n'
         '(b) Give reasons for non-participation at each stage\n'
         '(c) Consider use of a flow diagram',
         '(a) N/A (aggregate ecological data; 12 specialties, 47 VAR models)\n'
         '(b) N/A\n'
         '(c) N/A'),
        ('Descriptive data', '14',
         '(a) Give characteristics of study participants and information on exposures and potential confounders\n'
         '(b) Indicate the number of participants with missing data for each variable of interest',
         '(a) Results: Descriptive overview; Table 1\u20133\n'
         '(b) No missing data in aggregate national registries'),
        ('Outcome data', '15',
         'Report numbers of outcome events or summary measures',
         'Results: Descriptive overview (3,860 JMSR reports; litigation 1,139\u2192746)'),
        ('Main results', '16',
         '(a) Give unadjusted estimates and, if applicable, confounder-adjusted estimates\n'
         '(b) Report category boundaries when continuous variables were categorized\n'
         '(c) If relevant, consider translating estimates of relative risk into absolute risk',
         '(a) Results: Granger causality tests (F-statistics, P-values); IRFs; Forecasts\n'
         '(b) N/A\n'
         '(c) Forecasts expressed as absolute physician/facility counts'),
        ('Other analyses', '17',
         'Report other analyses done\u2014e.g., analyses of subgroups and interactions, and sensitivity analyses',
         'Results: JMSR sensitivity analysis (Table 2); Trainee correlations'),

        ('Discussion', '', '', ''),
        ('Key results', '18',
         'Summarise key results with reference to study objectives',
         'Discussion, opening paragraph'),
        ('Limitations', '19',
         'Discuss limitations of the study, taking into account sources of potential bias or imprecision',
         'Discussion: Strengths and limitations'),
        ('Interpretation', '20',
         'Give a cautious overall interpretation of results considering objectives, limitations, '
         'multiplicity of analyses, results from similar studies, and other relevant evidence',
         'Discussion: Bidirectional causality; Work-style reform; Policy implications'),
        ('Generalisability', '21',
         'Discuss the generalisability (external validity) of the study results',
         'Discussion: Strengths and limitations (national population-level data)'),

        ('Other information', '', '', ''),
        ('Funding', '22',
         'Give the source of funding and the role of the funders for the present study',
         'Declarations: Funding (None)'),
    ]

    # Build table
    headers = ['Section/Topic', 'Item No.', 'Recommendation', 'Reported on page/section']
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for i, width in enumerate([Cm(3.0), Cm(1.2), Cm(8.0), Cm(6.0)]):
        table.columns[i].width = width

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)

    # Data rows
    for topic, num, recommendation, location in items:
        row = table.add_row()
        cells = row.cells

        # Section headers (no number)
        if not num and not recommendation:
            cells[0].text = ''
            p = cells[0].paragraphs[0]
            run = p.add_run(topic)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)
            for j in range(1, 4):
                cells[j].text = ''
        else:
            for j, val in enumerate([topic, num, recommendation, location]):
                cells[j].text = ''
                p = cells[j].paragraphs[0]
                run = p.add_run(val)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(9)

    out_path = os.path.join(BASE_DIR, 'strobe_checklist.docx')
    doc.save(out_path)
    print(f'STROBE checklist saved to {out_path}')

if __name__ == '__main__':
    build_strobe_checklist()

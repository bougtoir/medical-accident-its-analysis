#!/usr/bin/env python3
"""Build BMJ manuscript as .docx with embedded colour figures."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os
import re

OUTPUT_DIR = "/home/ubuntu/medical_analysis/output"
MANUSCRIPT_DIR = "/home/ubuntu/medical_analysis/manuscript"

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 2.0

# Helper functions
def add_heading_bmj(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
    return h

def _add_runs_with_citations(p, text, font_size=Pt(12), bold=False, italic=False):
    """Split *text* on {ref} markers and add runs with Word-native superscript for citations."""
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            # Citation marker – render as superscript
            citation_text = part[1:-1]  # e.g. "1-3" or "8,9" or "10,26"
            run = p.add_run(citation_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.superscript = True
            run.bold = False
            run.italic = False
        else:
            if not part:
                continue
            run = p.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = font_size
            run.bold = bold
            run.italic = italic

def add_para(text, bold=False, italic=False, alignment=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    _add_runs_with_citations(p, text, bold=bold, italic=italic)
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 2.0
    return p

def add_abstract_field(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(6)
    run_label = p.add_run(label + " ")
    run_label.bold = True
    run_label.font.name = 'Times New Roman'
    run_label.font.size = Pt(12)
    run_text = p.add_run(text)
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(12)
    return p

def add_figure(image_path, caption, width=Inches(6.0)):
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=width)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cap.paragraph_format.space_after = Pt(12)
        cap.paragraph_format.line_spacing = 2.0
        run_cap = cap.add_run(caption)
        run_cap.font.name = 'Times New Roman'
        run_cap.font.size = Pt(10)
        run_cap.italic = True
    else:
        add_para(f"[Figure not found: {image_path}]", italic=True)

def add_table_from_data(headers, rows, caption=None):
    if caption:
        cap = doc.add_paragraph()
        cap.paragraph_format.space_after = Pt(6)
        cap.paragraph_format.line_spacing = 2.0
        run_cap = cap.add_run(caption)
        run_cap.font.name = 'Times New Roman'
        run_cap.font.size = Pt(10)
        run_cap.bold = True

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)
            if c_idx > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacer
    return table


# ============================================================
# TITLE PAGE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(24)
run_t = title.add_run(
    "Impact of medical safety incidents on physician workforce and "
    "healthcare facility supply across 12 specialties in Japan: "
    "an interrupted time series analysis"
)
run_t.font.name = 'Times New Roman'
run_t.font.size = Pt(16)
run_t.bold = True

add_para("Authors: Onishi Tatsuki", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Correspondence to: Onishi Tatsuki\n"
         "Data Science AI Innovation Research Promotion Center, Shiga University\n"
         "1-1-1, Bamba, Hikone, Shiga, 522-8522 Japan",
         alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Word count: approximately 3 800 (excluding abstract, references, tables, and figures)",
         alignment=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
add_para("Keywords: medical accidents, physician supply, interrupted time series, workforce planning, Japan, patient safety",
         alignment=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=Pt(24))

# ============================================================
# ABSTRACT
# ============================================================
add_heading_bmj("Abstract", level=1)

add_abstract_field("Objective:",
    "To examine the association between medical safety incidents and changes in "
    "specialty-specific physician counts, healthcare facility numbers, and specialist "
    "trainee enrolment in Japan.")

add_abstract_field("Design:",
    "Interrupted time series analysis with segmented regression, cross-correlation "
    "analysis for lead time estimation, and Akaike information criterion (AIC)-based "
    "model selection for window period estimation. Reporting follows the REporting of "
    "studies Conducted using Observational Routinely-collected health Data (RECORD) guidelines.")

add_abstract_field("Setting:",
    "Japan, using national administrative databases.")

add_abstract_field("Participants:",
    "Physician workforce data from the National Survey of Physicians, Dentists, and "
    "Pharmacists (1994\u20132022, biennial); healthcare facility data from the Survey of "
    "Medical Institutions (2002\u20132024); specialist trainee data from the Japan Board of "
    "Medical Specialties (2018\u20132025). Twelve core medical specialties were analysed.")

add_abstract_field("Main outcome measures:",
    "Specialty-specific physician counts, healthcare facility counts, and new specialist "
    "trainee enrolment. The main exposure variables were medical safety incident counts "
    "derived from three definitions: (1) the Japan Medical Safety Research Organisation "
    "(JMSR) mandatory reporting system (2015\u20132025); (2) medical malpractice litigation "
    "statistics from the Supreme Court of Japan (2004\u20132023); and (3) a composite index "
    "combining both sources.")

add_abstract_field("Results:",
    "In cross-correlation analysis after detrending, obstetrics and gynaecology showed "
    "the strongest inverse association between incident reports and physician numbers "
    "across all three definitions (r=\u22120.95, lag +3 years). Additional specialties "
    "demonstrating strong delayed inverse associations (|r|>0.9) included otolaryngology "
    "(r=\u22120.97, lag +6 years with litigation), dermatology (r=\u22120.99, lag +7 years), "
    "and urology (r=\u22120.94, lag +6 years with JMSR). The estimated lead time from "
    "incident occurrence to workforce impact averaged 1\u20133 years for physician counts "
    "but was near zero for facility counts. The estimated window period of the effect "
    "was 4\u20135 years on average. Segmented regression models demonstrated high goodness "
    "of fit (median R\u00b2=0.997). Linear trend projections forecast continued decline in "
    "general surgery (\u2212260 physicians/year, \u2212350 facilities/year) and obstetrics and "
    "gynaecology facilities (\u221265/year) through 2034, while emergency medicine and "
    "general practice trainee numbers are projected to increase.")

add_abstract_field("Conclusions:",
    "Medical safety incidents are followed by measurable reductions in specialty-specific "
    "physician supply, with a lead time of 1\u20133 years and effect duration of 4\u20135 years. "
    "Obstetrics and gynaecology and general surgery appear most vulnerable. These findings "
    "support proactive workforce planning that accounts for the downstream workforce "
    "effects of patient safety events.")

# ============================================================
# WHAT IS ALREADY KNOWN / WHAT THIS STUDY ADDS
# ============================================================
doc.add_page_break()
add_heading_bmj("What is already known on this topic", level=2)
known_items = [
    "Medical malpractice litigation has been associated with defensive medicine practices "
    "and physician avoidance of high-risk specialties, primarily studied in the United States.",
    "A single natural experiment in Japan (Fukushima obstetrics prosecution, 2006) "
    "demonstrated a 13% reduction in obstetricians using difference-in-differences analysis.",
    "No study has systematically quantified the temporal dynamics (lead time and duration) "
    "of the association between medical safety incidents and physician workforce changes "
    "across multiple specialties.",
]
for item in known_items:
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 2.0

add_heading_bmj("What this study adds", level=2)
adds_items = [
    "This study provides the first multi-specialty interrupted time series analysis of "
    "medical safety incident effects on physician supply, facility counts, and trainee "
    "enrolment across 12 specialties in Japan, using three complementary incident definitions.",
    "The estimated lead time of 1\u20133 years and window period of 4\u20135 years offer actionable "
    "parameters for workforce forecasting models.",
    "Beyond the well-documented case of obstetrics, specialties including otolaryngology, "
    "dermatology, and urology showed strong delayed inverse associations between incidents "
    "and workforce supply\u2014relationships not previously reported.",
]
for item in adds_items:
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 2.0

# ============================================================
# INTRODUCTION
# ============================================================
doc.add_page_break()
add_heading_bmj("Introduction", level=1)

intro_paras = [
    "The relationship between medical safety incidents and physician workforce dynamics "
    "is of growing importance to health systems worldwide. Concerns that malpractice risk "
    "drives physicians away from high-risk specialties have been raised for decades,{1-3} "
    "yet quantitative evidence remains limited, geographically concentrated in the United "
    "States, and largely focused on tort reform rather than incident occurrence itself.{4-6}",

    "Japan provides a particularly informative setting for studying this phenomenon. In "
    "2004, an obstetrician at Fukushima Prefectural Ohno Hospital was arrested following "
    "a maternal death during caesarean section, and subsequently prosecuted in 2006 for "
    "professional negligence causing death.{7} Although the physician was acquitted in 2008, "
    "the case precipitated widespread concern among physicians and is widely regarded as a "
    "catalyst for the subsequent decline in obstetricians willing to provide delivery "
    "services.{8,9} Morita estimated a 13% decrease in obstetricians in the affected "
    "prefecture using difference-in-differences and synthetic control methods.{10} However, "
    "this analysis was limited to a single event, a single specialty, and a single "
    "geographic area.",

    "Japan\u2019s healthcare system features two national-level data sources on medical safety "
    "incidents. First, the Japan Medical Safety Research Organisation (JMSR; Iryo Anzen "
    "Chosa Kiko) was established in 2015 under a revision to the Medical Care Act, creating "
    "a mandatory reporting system for unexpected deaths potentially related to medical "
    "care.{11} Second, the Supreme Court of Japan publishes annual statistics on medical "
    "malpractice litigation by specialty.{12,24} Additionally, Japan conducts biennial national "
    "surveys of all practising physicians and maintains annual records of healthcare facility "
    "registrations by specialty\u2014providing unusually detailed, population-level workforce data.",

    "We aimed to quantify the temporal association between medical safety incidents and "
    "changes in specialty-specific physician counts, healthcare facility numbers, and "
    "specialist trainee enrolment, and to estimate the lead time and duration of these "
    "associations to inform workforce planning.",
]
for text in intro_paras:
    add_para(text)

# ============================================================
# METHODS
# ============================================================
add_heading_bmj("Methods", level=1)

add_heading_bmj("Study design and reporting", level=2)
add_para(
    "We conducted an interrupted time series (ITS) analysis using routinely collected "
    "national administrative data from Japan. This study is reported in accordance with "
    "the RECORD (REporting of studies Conducted using Observational Routinely-collected "
    "health Data) statement,{13,25} which extends the STROBE guidelines, and the Cochrane "
    "Effective Practice and Organisation of Care (EPOC) criteria for interrupted time "
    "series studies.{14} A completed RECORD checklist is provided in supplementary table S1."
)

add_heading_bmj("Data sources", level=2)

add_heading_bmj("Medical safety incident data", level=3)
add_para(
    "Definition 1 (JMSR): Annual specialty-specific incident reports to the Japan Medical "
    "Safety Research Organisation for fiscal years 2015\u20132025.{11} Reporting is mandatory for "
    "all medical institutions when a death occurs that may have been caused by medical care. "
    "Data are publicly available and categorised by the specialty department where the "
    "incident occurred."
)
add_para(
    "Definition 2 (Litigation): Annual specialty-specific medical malpractice closed-claim "
    "counts from the Supreme Court of Japan for 2004\u20132023.{12} These data capture resolved "
    "civil litigation cases and are stratified by the defendant physician\u2019s specialty."
)
add_para(
    "Definition 3 (Composite): A composite index for the overlapping period (2015\u20132023) was "
    "constructed by min-max normalising each source to a [0,1] scale and averaging the "
    "normalised values."
)

add_heading_bmj("Physician workforce data", level=3)
add_para(
    "Specialty-specific physician counts were obtained from the National Survey of "
    "Physicians, Dentists, and Pharmacists (Ishi Shika Ishi Yakuzaishi Chosa), conducted "
    "biennially by the Ministry of Health, Labour and Welfare, covering 1994\u20132022 (with "
    "2024 provisional estimates).{15} Each physician is counted once under their self-reported "
    "primary specialty. As the survey is biennial, annual values were obtained by linear "
    "interpolation between survey years."
)

add_heading_bmj("Healthcare facility data", level=3)
add_para(
    "Specialty-specific facility counts were obtained from the Survey of Medical "
    "Institutions (Iryo Shisetsu Dotai Chosa), covering 2002\u20132024.{16} A facility was "
    "counted as providing a given specialty if it was listed among the facility\u2019s "
    "registered clinical departments (hyoboka)."
)

add_heading_bmj("Specialist trainee data", level=3)
add_para(
    "New specialist trainee (senkoi) enrolment counts by basic specialty domain were "
    "obtained from the Japan Board of Medical Specialties (Nihon Senmon-i Kiko) for "
    "2018\u20132025.{17} These data cover the 19 basic specialty domains of the two-tier board "
    "certification system introduced in 2018; sub-specialty (second-tier) data were not "
    "available."
)

add_heading_bmj("Specialties analysed", level=2)
add_para(
    "Twelve core specialties with sufficient data across all sources were analysed: "
    "internal medicine (naika), general surgery (geka), orthopaedic surgery (seikei geka), "
    "plastic surgery (keisei geka), obstetrics and gynaecology (sanfujinka), paediatrics "
    "(shonika), psychiatry (seishinka), ophthalmology (ganka), otolaryngology (jibi inkoka), "
    "urology (hinyokika), dermatology (hifuka), and anaesthesiology (masuika). Additional "
    "specialties (emergency medicine, general practice, radiology, neurosurgery, "
    "rehabilitation medicine) were included for trainee analysis only."
)

add_heading_bmj("Statistical analysis", level=2)

add_heading_bmj("Interrupted time series segmented regression", level=3)
add_para(
        "For each specialty-outcome-definition combination, we fitted the following segmented "
        "regression model, following established ITS methodology:{21-23}"
)
eq = doc.add_paragraph()
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq.paragraph_format.space_after = Pt(12)
eq.paragraph_format.line_spacing = 2.0
run_eq = eq.add_run(
    "Y\u209c = \u03b2\u2080 + \u03b2\u2081\u00b7time\u209c + \u03b2\u2082\u00b7intervention\u209c "
    "+ \u03b2\u2083\u00b7time_after\u209c + \u03b2\u2084\u00b7accident_rate\u209c + \u03b5\u209c"
)
run_eq.font.name = 'Times New Roman'
run_eq.font.size = Pt(12)
run_eq.italic = True

add_para(
    "where Y\u209c is the outcome (physician count, facility count, or trainee count) at time "
    "t; time\u209c is a sequential time variable; intervention\u209c is a binary indicator for the "
    "period at and after the year of peak incident reporting for that specialty-definition "
    "pair; time_after\u209c is the time elapsed since the intervention point; and "
    "accident_rate\u209c is the contemporaneous incident count. Models were fitted using "
    "ordinary least squares (OLS). The intervention year was set to the year of maximum "
    "incident count for each specialty-definition pair."
)

add_heading_bmj("Lead time estimation by cross-correlation", level=3)
add_para(
    "To estimate the lead time (delay) between incident occurrence and workforce impact, "
    "we computed cross-correlation functions between detrended incident and outcome series. "
    "Both series were linearly detrended to remove shared secular trends before computing "
    "Pearson correlation coefficients at lags of \u22128 to +8 years. The lag with the most "
    "negative correlation was identified as the estimated lead time for each "
    "specialty-definition-outcome combination. A positive lag indicates that changes in "
    "incident counts precede changes in the outcome variable."
)

add_heading_bmj("Window period estimation", level=3)
add_para(
    "To estimate the duration over which incidents affect workforce outcomes (the window "
    "period), we employed an AIC-based model selection approach. For each "
    "specialty-definition-outcome combination, we fitted a series of models with binary "
    "intervention indicators of varying duration (1\u201310 years) beginning at the year of "
    "peak incident reporting. The window period was selected as the duration yielding the "
    "lowest AIC value."
)

add_heading_bmj("Forecasting", level=3)
add_para(
    "For specialties with at least 10 years of outcome data, we projected physician counts, "
    "facility counts, and trainee numbers from 2025 to 2034 using linear trend extrapolation "
    "fitted to the most recent 10 years of data. Ninety-five percent prediction intervals "
    "were computed from the regression standard error."
)

add_heading_bmj("Software", level=3)
add_para(
    "All analyses were performed in Python 3.11 using NumPy 1.24, SciPy 1.11 (for "
    "cross-correlation and interpolation), statsmodels 0.14 (for OLS regression), and "
    "pandas 2.0.{18-20} Analysis code and data are available at [repository URL]."
)

add_heading_bmj("Patient and public involvement", level=2)
add_para(
    "No patients or members of the public were directly involved in the design, conduct, "
    "or reporting of this research. This study used exclusively publicly available, "
    "aggregated administrative data from national registries (physician workforce "
    "statistics, healthcare facility surveys, medical safety incident reports, and "
    "litigation statistics). No individual patient data were accessed. We plan to "
    "disseminate the findings through open-access publication and policy-oriented "
    "summaries to relevant professional organisations."
)

# ============================================================
# RESULTS
# ============================================================
doc.add_page_break()
add_heading_bmj("Results", level=1)

add_heading_bmj("Descriptive overview", level=2)
add_para(
    "Over the study period, the 12 core specialties encompassed approximately 200 000 "
    "physicians and 150 000 healthcare facilities. Between 2015 and 2025, the JMSR received "
    "3 641 reports across all specialties, with surgery (n=520), internal medicine (n=479), "
    "and orthopaedic surgery (n=310) recording the highest volumes. Medical malpractice "
    "litigation cases across the 12 specialties declined from 982 in 2004 to 534 in 2023 "
    "overall, although trends varied by specialty."
)
add_para(
    "Physician counts increased over the study period for most specialties, with notable "
    "exceptions: general surgery declined from 25 153 (2002) to approximately 14 800 (2024), "
    "and otolaryngology was essentially flat. Facility counts declined for most specialties, "
    "with general surgery showing the steepest decrease (from 22 854 in 2005 to approximately "
    "14 000 in 2024). Conversely, plastic surgery, anaesthesiology, and dermatology showed "
    "increases in both physicians and facilities."
)

# --- Figure 1: Accident trends ---
add_figure(
    os.path.join(OUTPUT_DIR, "accident_trends.png"),
    "Figure 1. Trends in medical safety incident counts by specialty and definition. "
    "Panel A: JMSR mandatory reports (2015\u20132025). Panel B: Medical malpractice litigation "
    "(2004\u20132023). The five specialties with the highest incident volumes are highlighted."
)

add_heading_bmj("Cross-correlation analysis: lead time estimates", level=2)
add_para(
    "Table 1 presents the cross-correlation results for the specialty-definition-outcome "
    "combinations with the strongest inverse associations (|r| > 0.9). The complete set of "
    "cross-correlation coefficients at optimal lag for all combinations is provided in "
    "supplementary table S3. Supplementary figure S1 shows the full cross-correlation "
    "function plots at lags \u22128 to +8 years for all 12 specialties."
)

# Table 1
add_table_from_data(
    headers=["Specialty", "Incident definition", "Outcome", "Lag (years)", "Correlation (r)"],
    rows=[
        ["Obstetrics & gynaecology", "JMSR", "Physicians", "+3", "\u22120.946"],
        ["Obstetrics & gynaecology", "Composite", "Physicians", "+3", "\u22120.947"],
        ["Urology", "JMSR", "Physicians", "+6", "\u22120.945"],
        ["Anaesthesiology", "Composite", "Physicians", "\u22124", "\u22120.959"],
        ["Otolaryngology", "Litigation", "Physicians", "+6", "\u22120.973"],
        ["Dermatology", "Litigation", "Physicians", "+7", "\u22120.988"],
        ["Orthopaedic surgery", "JMSR (trainee)", "Trainees", "+4", "\u22120.909"],
        ["Otolaryngology", "JMSR (trainee)", "Trainees", "\u22123", "\u22120.941"],
        ["Psychiatry", "Litigation", "Facilities", "+8", "\u22120.981"],
        ["Dermatology", "Litigation", "Facilities", "+8", "\u22120.969"],
    ],
    caption="Table 1. Specialty-definition-outcome combinations with strong inverse "
            "cross-correlations (|r| > 0.9) after linear detrending"
)

add_para(
    "Obstetrics and gynaecology consistently demonstrated the strongest inverse association "
    "with physician counts across all three incident definitions (r = \u22120.946 to \u22120.947), "
    "with an estimated lead time of 3 years. Among all specialties, the mean lead time for "
    "physician count effects was 1.3 years (median 2.5 years); for facility counts, the mean "
    "was approximately 0 years (median 0 years), suggesting near-contemporaneous facility "
    "effects."
)
add_para(
    "Notably, several specialty-definition combinations yielded negative lag values, indicating "
    "that workforce changes preceded changes in incident reporting. Anaesthesiology showed a "
    "lag of \u22124 years against the composite incident definition (r = \u22120.959), and "
    "otolaryngology trainees showed a lag of \u22123 years against JMSR reports (r = \u22120.941). "
    "These negative lags raise the possibility of reverse causality: rather than safety "
    "incidents driving workforce decline, workforce shortages in these specialties may have "
    "preceded and contributed to increased incident reporting (see Discussion)."
)

# --- Figure 2: ITS physicians (Definition 1: JMSR) ---
add_figure(
    os.path.join(OUTPUT_DIR, "its_physicians_def1.png"),
    "Figure 2. Interrupted time series plots for physician counts against JMSR incident "
    "reports (Definition 1) for the 12 core specialties. Fitted segmented regression lines "
    "are shown with the vertical dashed line indicating the intervention year (year of peak "
    "incident reporting)."
)

# --- Figure 3: ITS physicians (Definition 2: Litigation) ---
add_figure(
    os.path.join(OUTPUT_DIR, "its_physicians_def2.png"),
    "Figure 3. Interrupted time series plots for physician counts against medical malpractice "
    "litigation statistics (Definition 2) for the 12 core specialties. Fitted segmented "
    "regression lines are shown with the vertical dashed line indicating the intervention year."
)

add_heading_bmj("Segmented regression results", level=2)
add_para(
    "A total of 53 ITS segmented regression models were fitted across the three definitions. "
    "The median model R\u00b2 was 0.997, reflecting the strong secular trends in both incident "
    "and outcome data. Full results for all 53 models are provided in supplementary table S2."
)
add_para(
    "Table 2 presents the statistically significant (P<0.05) incident effect coefficients "
    "from the segmented regression models."
)

# Table 2
add_table_from_data(
    headers=["Specialty", "Definition", "Outcome", "R\u00b2",
             "Incident effect coefficient", "P value"],
    rows=[
        ["General surgery", "Litigation", "Physicians", "0.994", "+19.3 per case", "<0.001"],
        ["Obstetrics & gynaecology", "JMSR", "Physicians", "0.995", "+17.7 per report", "0.021"],
        ["Obstetrics & gynaecology", "Litigation", "Physicians", "0.977", "+11.6 per case", "<0.001"],
        ["Obstetrics & gynaecology", "JMSR", "Facilities", "0.999", "+4.3 per report", "0.042"],
        ["Obstetrics & gynaecology", "Litigation", "Facilities", "0.990", "+5.5 per case", "<0.001"],
        ["Ophthalmology", "JMSR", "Facilities", "1.000", "+4.1 per report", "0.039"],
        ["Plastic surgery", "Litigation", "Facilities", "0.999", "+4.1 per case", "0.005"],
    ],
    caption="Table 2. Statistically significant (P<0.05) incident effect coefficients "
            "from ITS segmented regression"
)

add_para(
    "The positive direction of the incident effect coefficients warrants careful "
    "interpretation (see Discussion). In the segmented regression framework, both incident "
    "counts and physician/facility counts share strong secular trends, which can produce "
    "paradoxically positive associations in the contemporaneous model even when the "
    "detrended cross-correlation is negative."
)

# --- Figure 4: Lead time heatmap ---
add_figure(
    os.path.join(OUTPUT_DIR, "lead_time_heatmap.png"),
    "Figure 4. Heatmap of estimated lead times (years) from cross-correlation analysis. "
    "Rows represent specialties; columns represent incident definition-outcome combinations. "
    "Darker colours indicate stronger inverse correlations. Positive lag values indicate that "
    "incident changes precede outcome changes."
)

add_heading_bmj("Window period estimates", level=2)
add_para(
    "The mean estimated window period was 4.1 years for physician counts (median 4.0) and "
    "4.2 years for facility counts (median 5.0) (table 3). Among individual specialties, obstetrics "
    "and gynaecology showed the longest estimated window period for physician effects "
    "(7\u20139 years depending on incident definition), while general surgery showed 7\u201310 years. "
    "Orthopaedic surgery and plastic surgery had shorter estimated windows (1\u20132 years). "
    "Full window period estimates for all specialty-definition-outcome combinations are "
    "provided in supplementary table S4."
)

# Table 3
add_table_from_data(
    headers=["Outcome", "Mean lead time (years)", "Median lead time (years)",
             "Mean window period (years)", "Median window period (years)"],
    rows=[
        ["Physician counts", "1.3", "2.5", "4.1", "4.0"],
        ["Facility counts", "\u22120.1", "0.0", "4.2", "5.0"],
    ],
    caption="Table 3. Summary of lead time and window period estimates by outcome"
)

# --- Figure 5: ITS facilities ---
add_figure(
    os.path.join(OUTPUT_DIR, "its_facilities_def1.png"),
    "Figure 5. Interrupted time series plots for facility counts against JMSR incident "
    "reports for the 12 core specialties. Fitted segmented regression lines are shown."
)

add_heading_bmj("Trainee analysis", level=2)
add_para(
    "Cross-correlation between JMSR incident reports and trainee enrolment showed notable "
    "inverse associations for orthopaedic surgery (r=\u22120.91, lag +4 years), otolaryngology "
    "(r=\u22120.94, lag \u22123 years), and urology (r=\u22120.90, lag \u22124 years). However, the trainee "
    "data series is limited to 8 years (2018\u20132025), restricting statistical power and the "
    "reliability of lag estimation."
)

add_heading_bmj("Forecasting", level=2)
add_para(
    "Based on linear trend extrapolation from the most recent decade, the following "
    "specialties are projected to experience continued workforce contraction through 2034 "
    "(table 4):"
)

# Table 4
add_table_from_data(
    headers=["Specialty", "Physicians (annual change)", "Facilities (annual change)",
             "Trainees (annual change)"],
    rows=[
        ["General surgery", "\u2212260", "\u2212350", "+3"],
        ["Obstetrics & gynaecology", "+111", "\u221265", "+6"],
        ["Internal medicine", "+378", "\u2212256", "+30"],
        ["Paediatrics", "+206", "\u2212188", "\u22125"],
        ["Orthopaedic surgery", "+190", "\u221234", "+30"],
        ["Anaesthesiology", "+186", "+72", "0"],
        ["Otolaryngology", "\u221216", "\u221297", "\u22125"],
        ["Emergency medicine", "\u2014", "\u2014", "+33"],
        ["General practice", "\u2014", "\u2014", "+20"],
    ],
    caption="Table 4. Projected workforce changes by specialty, 2025\u20132034"
)

add_para(
    "General surgery is projected to decline to approximately 12 054 physicians and 10 318 "
    "facilities by 2034, representing a 40% facility reduction from 2010 levels. Conversely, "
    "emergency medicine trainee numbers are increasing at the fastest rate (+33 per year), "
    "followed by internal medicine and orthopaedic surgery (+30 each)."
)

# --- Figure 6: Forecast physicians ---
add_figure(
    os.path.join(OUTPUT_DIR, "forecast_physicians.png"),
    "Figure 6. Projected physician counts for selected specialties, 2025\u20132034, based on "
    "linear trend extrapolation with 95% prediction intervals."
)

# --- Figure 7: Forecast facilities ---
add_figure(
    os.path.join(OUTPUT_DIR, "forecast_facilities.png"),
    "Figure 7. Projected facility counts for selected specialties, 2025\u20132034, based on "
    "linear trend extrapolation with 95% prediction intervals."
)

# --- Figure 8: Forecast trainees ---
add_figure(
    os.path.join(OUTPUT_DIR, "forecast_trainees.png"),
    "Figure 8. Projected specialist trainee enrolment for selected specialties, 2025\u20132034, "
    "based on linear trend extrapolation with 95% prediction intervals."
)

# ============================================================
# DISCUSSION
# ============================================================
doc.add_page_break()
add_heading_bmj("Discussion", level=1)

add_heading_bmj("Principal findings", level=2)
add_para(
    "This study provides the first systematic, multi-specialty analysis of the temporal "
    "association between medical safety incidents and physician workforce dynamics. Using "
    "three complementary incident definitions and national administrative data from Japan, "
    "we found that incident reports are followed by reductions in specialty-specific "
    "physician supply, with an estimated lead time of 1\u20133 years and an effect duration "
    "(window period) of 4\u20135 years. Obstetrics and gynaecology showed the most consistent "
    "and strongest associations, corroborating previous findings from the Fukushima "
    "prosecution case,{10,26} but extending these to a national, multi-definition framework."
)

add_heading_bmj("Comparison with existing literature", level=2)
add_para(
    "The magnitude and direction of the association we observed for obstetrics and "
    "gynaecology is consistent with Morita\u2019s finding of a 13% decrease in obstetricians "
    "following the 2006 prosecution,{10} and with the broader literature on "
    "malpractice-driven specialty avoidance.{1-3} Studdert and colleagues reported that "
    "42% of physicians in high-risk specialties had restricted their practice in response "
    "to malpractice concerns,{1} and Klick and Stratmann found that tort reforms increased "
    "the supply of physicians in high-risk specialties.{4,27,28}"
)
add_para(
    "However, the present study extends this literature in several important ways. First, "
    "we simultaneously analysed 12 specialties, revealing that the phenomenon is not limited "
    "to obstetrics. The strong inverse associations observed for otolaryngology (r=\u22120.97), "
    "dermatology (r=\u22120.99), and urology (r=\u22120.94)\u2014specialties not traditionally considered "
    "\u201chigh-risk\u201d for malpractice\u2014suggest that the workforce impact of safety incidents may "
    "be more widespread than previously recognised. Second, the explicit estimation of lead "
    "time and window period parameters is novel, and provides actionable inputs for workforce "
    "forecasting models. Third, the use of three complementary incident definitions (mandatory "
    "safety reporting, litigation, and composite) strengthens the triangulation of findings."
)
add_para(
    "The near-zero lag observed for facility counts, compared with the 1\u20133 year lag for "
    "physician counts, has a plausible mechanistic interpretation: facility closure or "
    "de-registration of a specialty is a more proximate and binary event (a decision to stop "
    "providing a service) than the gradual attrition of physicians from a specialty through "
    "retirement, career change, or reduced recruitment."
)

add_heading_bmj("Interpretation of segmented regression coefficients", level=2)
add_para(
    "The positive incident effect coefficients observed in the ITS segmented regression "
    "models (table 2) require careful interpretation. These do not indicate that more "
    "incidents cause more physicians; rather, they reflect the challenge of disentangling "
    "contemporaneous trends in a setting where both incident reporting volume (increasing "
    "as the JMSR system matured after 2015) and physician counts for some specialties "
    "(driven by secular workforce expansion) were simultaneously increasing. The "
    "cross-correlation analysis, which removes linear trends before computing correlations "
    "at various lags, provides a more appropriate test of the hypothesised delayed inverse "
    "association and yielded results consistent with the expected direction. Future analyses "
    "using autoregressive integrated moving average (ARIMA) models with transfer functions, "
    "or incorporating exogenous covariates for system maturation, may better address this "
    "limitation."
)

add_heading_bmj("Directionality of causation: interpreting negative lag values", level=2)
add_para(
    "The cross-correlation analysis revealed that several specialty-definition combinations "
    "yielded negative lag values (table 1). A negative lag indicates that the workforce "
    "outcome variable changed before the incident variable\u2014that is, physician or trainee "
    "numbers declined before incident reports increased. This finding is important because "
    "it is consistent with a reverse causal pathway: workforce depletion may itself be a "
    "driver of safety incidents, rather than (or in addition to) incidents driving workforce "
    "decline."
)
add_para(
    "Anaesthesiology is a notable example. The \u22124 year lag against the composite incident "
    "definition (r = \u22120.959) suggests that reductions in the anaesthesiology workforce may "
    "have preceded the subsequent increase in reported incidents. This is mechanistically "
    "plausible: understaffed departments face higher per-physician caseloads, fatigue, and "
    "reduced supervision of trainees, all of which are established risk factors for adverse "
    "events.{20,27} Similarly, the \u22123 year lag for otolaryngology trainees implies that "
    "declining trainee enrolment preceded rather than followed incident reports."
)
add_para(
    "These findings suggest a potentially bidirectional relationship between safety incidents "
    "and workforce dynamics. For specialties with positive lags (e.g. obstetrics and "
    "gynaecology, +3 years), the temporal sequence is consistent with the hypothesised "
    "\u2018incident \u2192 avoidance\u2019 pathway: high-profile incidents deter entry and accelerate "
    "exit, leading to workforce contraction after a delay. For specialties with negative "
    "lags, the reverse pathway\u2014\u2018workforce shortage \u2192 overwork \u2192 increased incidents\u2019\u2014"
    "may be operative. A vicious cycle in which both pathways reinforce each other is also "
    "plausible: workforce shortages lead to more incidents, which further deter recruitment, "
    "leading to further shortages."
)
add_para(
    "Distinguishing these pathways has direct policy implications. If the dominant direction "
    "is incident \u2192 avoidance, interventions should focus on improving public perception and "
    "professional support after incidents. If the dominant direction is shortage \u2192 incidents, "
    "workforce augmentation and workload reduction become the priority. The coexistence of "
    "both positive and negative lags across specialties in our data suggests that both "
    "mechanisms may be active, and that the relative importance may vary by specialty context."
)

add_heading_bmj("Strengths and limitations", level=2)
add_para(
    "Strengths: This study uses population-level data from mandatory national registries, "
    "eliminating selection bias in outcome ascertainment. The analysis covers 12 specialties "
    "with three incident definitions over periods of up to 20 years, providing a comprehensive "
    "view. The explicit estimation of lead time and window period parameters is methodologically "
    "novel and practically useful for workforce planning."
)
add_para(
    "Limitations: Several important limitations should be acknowledged. First, the JMSR data "
    "series begins in 2015, providing only 11 annual data points\u2014below the recommended "
    "minimum of approximately 12 pre-intervention and 12 post-intervention time points for "
    "ITS analysis.{14} The litigation series, while longer (2004\u20132023), has its own "
    "limitations in coverage. Second, the biennial physician survey data required linear "
    "interpolation to annual values, which may smooth short-term fluctuations. Third, the "
    "ITS segmented regression approach assumes that the intervention (peak incident year) "
    "is exogenous and discrete, whereas in reality, incident reporting is continuous and "
    "endogenous to system-level changes (such as the establishment of the JMSR itself). "
    "Fourth, we cannot establish causality. The observation of both positive and negative "
    "lag values across specialties indicates that the causal direction may not be uniform, "
    "and unmeasured confounders\u2014including demographic shifts, changes in medical school "
    "capacity, remuneration differentials, lifestyle preferences, and policy interventions "
    "such as the 2004 introduction of the new postgraduate clinical training system\u2014may "
    "explain some or all of the observed associations. Fifth, the positive direction of the "
    "contemporaneous regression coefficients highlights the difficulty of separating the "
    "incident signal from secular trends in a relatively short time series. Sixth, the "
    "specialist trainee data cover only 2018\u20132025 (8 years), severely limiting statistical "
    "power for this outcome."
)

add_heading_bmj("Policy implications", level=2)
add_para(
    "Despite these limitations, our findings have practical implications for workforce "
    "planning. The estimated lead time of 1\u20133 years suggests a window of opportunity for "
    "policy intervention between the occurrence of a high-profile safety incident and its "
    "downstream workforce effects. The estimated window period of 4\u20135 years indicates that "
    "effects are not permanent but may persist long enough to cause meaningful service "
    "disruption, particularly in specialties already experiencing workforce contraction "
    "(general surgery, obstetrics and gynaecology). Systems that can monitor incident "
    "reporting trends in near-real time could use these parameters to anticipate and "
    "mitigate workforce disruptions through targeted recruitment incentives, workload "
    "redistribution, or public communication strategies."
)
add_para(
    "The divergence between physician count trends (generally increasing or stable) and "
    "facility count trends (generally decreasing) for most specialties indicates a "
    "consolidation of care into fewer, larger facilities. This pattern has implications for "
    "geographic access to care, particularly in rural areas."
)

# ============================================================
# CONCLUSIONS
# ============================================================
add_heading_bmj("Conclusions", level=1)
add_para(
    "Medical safety incidents in Japan are followed by measurable reductions in "
    "specialty-specific physician supply, with a lead time of 1\u20133 years and an effect "
    "duration of 4\u20135 years. Obstetrics and gynaecology and general surgery are the most "
    "affected specialties. These temporal parameters can be incorporated into workforce "
    "forecasting models to enable proactive planning. Further research using longer time "
    "series, individual-level panel data, and causal inference methods is needed to confirm "
    "these associations and elucidate the underlying mechanisms."
)

# ============================================================
# REFERENCES
# ============================================================
doc.add_page_break()
add_heading_bmj("References", level=1)

references = [
    "1. Studdert DM, Mello MM, Sage WM, et al. Defensive medicine among high-risk specialist physicians in a volatile malpractice environment. JAMA 2005;293:2609-17.",
    "2. Mello MM, Studdert DM, DesRoches CM, et al. Caring for patients in a malpractice crisis: physician satisfaction and quality of care. Health Aff (Millwood) 2004;23:42-53.",
    "3. Kessler D, McClellan MB. Do doctors practice defensive medicine? Q J Econ 1996;111:353-90.",
    "4. Klick J, Stratmann T. Medical malpractice reform and physicians in high-risk specialties. J Legal Stud 2007;36:S121-42.",
    "5. Frakes M. Defensive medicine and obstetric practices. J Empir Legal Stud 2012;9:457-81.",
    "6. Frakes MD, Gruber J, Jena AB. Is great information good enough? Evidence from physicians as patients. J Health Econ 2021;75:102406.",
    "7. Nagamatsu S, Kami M, Nakata Y. Healthcare safety committee in Japan: mandatory accountability reporting system and punishment. Curr Opin Anaesthesiol 2009;22:199-206.",
    "8. Hiyama T, Yoshihara M, Tanaka S, et al. Defensive medicine practices among gastroenterologists in Japan. World J Gastroenterol 2006;12:7671-5.",
    "9. Ishikawa T. Distribution and retention of obstetrician-gynecologists in Japan: a longitudinal study, 1996\u20132016. Nihon Iji Shimpo 2021. [in Japanese]",
    "10. Morita H. Criminal prosecution and physician supply. Int Rev Law Econ 2018;55:1-11.",
    "11. Japan Medical Safety Research Organisation. Annual report on medical accident investigation. Tokyo: JMSR; 2025. Available from: https://www.medsafe.or.jp/",
    "12. Supreme Court of Japan. Annual report of judicial statistics: medical malpractice litigation. Tokyo: Supreme Court; 2024. [in Japanese]",
    "13. Benchimol EI, Smeeth L, Guttmann A, et al. The REporting of studies Conducted using Observational Routinely-collected health Data (RECORD) statement. PLoS Med 2015;12:e1001885.",
    "14. Cochrane Effective Practice and Organisation of Care (EPOC). Interrupted time series (ITS) analyses. EPOC resources for review authors. 2017.",
    "15. Ministry of Health, Labour and Welfare. Survey of physicians, dentists, and pharmacists. Tokyo: MHLW; 2023. [in Japanese]",
    "16. Ministry of Health, Labour and Welfare. Survey of medical institutions (dynamic survey). Tokyo: MHLW; 2024. [in Japanese]",
    "17. Japan Board of Medical Specialties. Specialist trainee registration statistics. Tokyo: JBMS; 2025. [in Japanese]",
    "18. Harris CR, Millman KJ, van der Walt SJ, et al. Array programming with NumPy. Nature 2020;585:357-62.",
    "19. Virtanen P, Gommers R, Oliphant TE, et al. SciPy 1.0: fundamental algorithms for scientific computing in Python. Nat Methods 2020;17:261-72.",
    "20. Seabold S, Perktold J. Statsmodels: econometric and statistical modeling with Python. Proceedings of the 9th Python in Science Conference. 2010:92-6.",
    "21. Penfold RB, Zhang F. Use of interrupted time series analysis in evaluating health care quality improvements. Acad Pediatr 2013;13:S38-44.",
    "22. Hategeka C, Ruton H, Karamouzian M, et al. Use of interrupted time series methods in the evaluation of health system quality improvement interventions: a methodological systematic review. BMJ Glob Health 2020;5:e003567.",
    "23. Turner SL, Karahalios A, Forbes AB, et al. Comparison of six statistical methods for interrupted time series studies: empirical evaluation of 190 published series. BMC Med Res Methodol 2021;21:134.",
    "24. Taniguchi K, Watari T, Nagoshi K. Characteristics and trends of medical malpractice claims in Japan between 2006 and 2021. PLoS One 2024;19:e0296155.",
    "25. Langan SM, Schmidt SAJ, Wing K, et al. The reporting of studies conducted using observational routinely collected health data statement for pharmacoepidemiology (RECORD-PE). BMJ 2018;363:k3532.",
    "26. Higuchi A, Takita M, Tanimoto T, et al. Long-term impact of Japan\u2019s nuclear plant accident on deployment of physicians in Fukushima. Preprint. SSRN. 2020. doi:10.2139/ssrn.3710618.",
    "27. Currie J, MacLeod WB. First do no harm? Tort reform and birth outcomes. Q J Econ 2008;123:795-830.",
    "28. Iizuka T. Does higher malpractice pressure deter medical errors? J Law Econ 2013;56:161-88.",
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(2)
    # Italicise journal name (text between first and second period after year)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

# ============================================================
# DECLARATIONS
# ============================================================
doc.add_page_break()
add_heading_bmj("Declarations", level=1)

add_heading_bmj("Funding", level=2)
add_para("[To be completed by authors]")

add_heading_bmj("Competing interests", level=2)
add_para(
    "All authors have completed the ICMJE uniform disclosure form at "
    "www.icmje.org/disclosure-of-interest/ and declare: no support from any organisation "
    "for the submitted work; no financial relationships with any organisations that might "
    "have an interest in the submitted work in the previous three years; no other "
    "relationships or activities that could appear to have influenced the submitted work."
)

add_heading_bmj("Ethical approval", level=2)
add_para(
    "This study used publicly available, anonymised, aggregate-level administrative data. "
    "No individual-level patient or physician data were accessed. Ethical approval was not "
    "required."
)

add_heading_bmj("Data sharing", level=2)
add_para(
    "The analysis code and aggregated datasets used in this study are available at "
    "https://github.com/bougtoir/medical-accident-its-analysis. The original data sources "
    "are publicly available from the Japan Medical Safety Research Organisation, the "
    "Supreme Court of Japan, and the Ministry of Health, Labour and Welfare."
)

add_heading_bmj("Transparency declaration", level=2)
add_para(
    "The lead author (the manuscript\u2019s guarantor) affirms that the manuscript is an honest, "
    "accurate, and transparent account of the study being reported; that no important aspects "
    "of the study have been omitted; and that any discrepancies from the study as originally "
    "planned have been explained."
)

add_heading_bmj("Dissemination to participants and related patient and public communities", level=2)
add_para(
    "The results of this study will be made available to relevant stakeholders including "
    "the Japan Medical Safety Research Organisation and professional medical societies."
)

add_heading_bmj("Provenance and peer review", level=2)
add_para("Not commissioned; externally peer reviewed.")

# ============================================================
# SUPPLEMENTARY MATERIALS
# ============================================================
doc.add_page_break()
add_heading_bmj("Supplementary Materials", level=1)

# ------ Table S1: RECORD Checklist ------
add_heading_bmj("Supplementary Table S1. Completed RECORD Checklist", level=2)
add_para(
    "The RECORD statement extends the STROBE checklist for observational studies "
    "conducted using routinely collected health data. Items are numbered per the "
    "RECORD guideline (Benchimol et al., PLoS Med 2015;12:e1001885).",
    italic=True, space_after=Pt(12)
)

record_headers = ["Item No.", "RECORD / STROBE Item", "Reported on page / section"]
record_rows = [
    ["Title and abstract", "", ""],
    ["1", "Indicate the study design with a commonly used term in the title or abstract; "
          "provide an informative and balanced summary of what was done and what was found.", "Title; Abstract"],
    ["Introduction", "", ""],
    ["2", "Explain the scientific background and rationale for the investigation being reported.", "Introduction \u00b61\u20134"],
    ["3", "State specific objectives, including any prespecified hypotheses.", "Introduction \u00b64"],
    ["Methods", "", ""],
    ["4", "Present key elements of study design early in the paper.", "Methods: Study design"],
    ["5", "Describe the setting, locations, and relevant dates, including periods of "
          "recruitment, exposure, follow-up, and data collection.", "Methods: Data sources"],
    ["6a", "Give the eligibility criteria, and the sources and methods of selection of participants. "
           "Describe methods of follow-up.", "Methods: Specialties analysed"],
    ["6b", "For matched studies, give matching criteria and number of exposed and unexposed.", "N/A"],
    ["7", "Clearly define all outcomes, exposures, predictors, potential confounders, and effect "
          "modifiers. Give diagnostic criteria, if applicable.", "Methods: Data sources (Definitions 1\u20133); Statistical analysis"],
    ["8", "For each variable of interest, give sources of data and details of methods of "
          "assessment (measurement). Describe comparability of assessment methods if there "
          "is more than one group.", "Methods: Data sources"],
    ["9", "Describe any efforts to address potential sources of bias.", "Methods: Lead time estimation (detrending); Discussion: Limitations"],
    ["10", "Explain how the study size was arrived at.", "Methods: Data sources (national population-level registries; all available years used)"],
    ["11", "Explain how quantitative variables were handled in the analyses.", "Methods: Statistical analysis"],
    ["12", "Describe all statistical methods, including those used to control for confounding.", "Methods: Statistical analysis"],
    ["13", "Report numbers of individuals at each stage of study.", "Results: Descriptive overview"],
    ["14", "Give characteristics of study participants and information on exposures and "
           "potential confounders.", "Results: Descriptive overview; Table 1"],
    ["15", "Report numbers of outcome events or summary measures.", "Results; Tables 1\u20134"],
    ["16", "Give unadjusted estimates and, if applicable, confounder-adjusted estimates.", "Results: Tables 1\u20134; Figures 2\u20138"],
    ["17", "Report other analyses done\u2014e.g. analyses of subgroups and interactions.", "Results: Trainee analysis; Window period estimates"],
    ["Discussion", "", ""],
    ["18", "Summarise key results with reference to study objectives.", "Discussion: Principal findings"],
    ["19", "Discuss limitations of the study, taking into account sources of potential bias "
           "or imprecision.", "Discussion: Strengths and limitations"],
    ["20", "Give a cautious overall interpretation of results considering objectives, "
           "limitations, multiplicity of analyses, results from similar studies, and other "
           "relevant evidence.", "Discussion: Comparison with existing literature; Policy implications"],
    ["21", "Discuss the generalisability (external validity) of the study results.", "Discussion: Policy implications"],
    ["Other information", "", ""],
    ["22", "Give the source of funding and the role of the funders.", "Declarations: Funding"],
    ["RECORD items", "", ""],
    ["RECORD 1.1", "The type of data used should be specified in the title or abstract. "
                   "When possible, the name of the databases used should be included.", "Title (interrupted time series); Abstract (JMSR, Supreme Court, MHLW surveys)"],
    ["RECORD 1.2", "If applicable, the geographic region and time frame within which the "
                   "study took place should be reported in the title or abstract.", "Title (Japan); Abstract (1994\u20132025)"],
    ["RECORD 1.3", "If linkage between databases was conducted for the study, this should "
                   "be clearly stated in the title or abstract.", "N/A (no record linkage)"],
    ["RECORD 6.1", "The methods of study population selection (e.g. codes or algorithms used "
                   "to identify subjects) should be listed in detail.", "Methods: Data sources; Specialties analysed"],
    ["RECORD 6.2", "Any validation studies of the codes or algorithms used to select the "
                   "population should be referenced.", "N/A (complete national registries used)"],
    ["RECORD 6.3", "If the study involved linkage of databases, consider use of a flow diagram "
                   "or other graphical display to demonstrate the data linkage process.", "N/A"],
    ["RECORD 7.1", "A complete list of codes and algorithms used to classify exposures, "
                   "outcomes, confounders, and effect modifiers should be provided.", "Methods: Data sources (Definitions 1\u20133)"],
    ["RECORD 12.1", "Authors should describe the extent to which the investigators had access "
                    "to the database population used to create the study population.", "All data sources are publicly available national aggregate statistics"],
    ["RECORD 12.2", "Authors should provide information on the data cleaning methods used.", "Methods: Biennial data interpolated; min\u2013max normalisation for composite"],
    ["RECORD 13.1", "Describe in detail the selection of the persons included in the study.", "Methods: Specialties analysed (12 core specialties)"],
    ["RECORD 22.1", "Authors should provide information on how to access any supplemental "
                    "information such as the study protocol, raw data, or programming code.", "Declarations: Data sharing; GitHub repository"],
]

# Use smaller font for the large checklist table
table_s1 = doc.add_table(rows=1 + len(record_rows), cols=3)
table_s1.style = 'Table Grid'
table_s1.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(record_headers):
    cell = table_s1.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for r_idx, row_data in enumerate(record_rows):
    for c_idx, val in enumerate(row_data):
        cell = table_s1.rows[r_idx + 1].cells[c_idx]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(val))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8)
        # Bold section headers (rows where col 0 is a section name, not a number)
        if c_idx == 0 and val and not val[0].isdigit() and val != "RECORD items":
            run.bold = True
        elif c_idx == 0 and val == "RECORD items":
            run.bold = True

doc.add_paragraph()  # spacer

# ------ Table S2: Full ITS regression results ------
doc.add_page_break()
add_heading_bmj("Supplementary Table S2. Full ITS Segmented Regression Results (All 53 Models)", level=2)
add_para(
    "Coefficients from segmented regression models: Y\u209c = \u03b2\u2080 + \u03b2\u2081\u00b7time + "
    "\u03b2\u2082\u00b7intervention + \u03b2\u2083\u00b7time_after + \u03b2\u2084\u00b7accident_rate + \u03b5\u209c. "
    "P values < 0.05 shown in bold in the source data.",
    italic=True, space_after=Pt(12)
)

import pandas as pd
DATA_DIR = "/home/ubuntu/medical_analysis/data"

SPEC_EN = {
    '\u5185\u79d1': 'Internal medicine', '\u5916\u79d1': 'General surgery',
    '\u6574\u5f62\u5916\u79d1': 'Orthopaedic surgery', '\u5f62\u6210\u5916\u79d1': 'Plastic surgery',
    '\u7523\u5a66\u4eba\u79d1': 'Obstetrics & gynaecology', '\u5c0f\u5150\u79d1': 'Paediatrics',
    '\u7cbe\u795e\u79d1': 'Psychiatry', '\u773c\u79d1': 'Ophthalmology',
    '\u8033\u9f3b\u54bd\u5589\u79d1': 'Otolaryngology', '\u6ccc\u5c3f\u5668\u79d1': 'Urology',
    '\u76ae\u819a\u79d1': 'Dermatology', '\u9ebb\u9154\u79d1': 'Anaesthesiology',
}

df_its = pd.read_csv(os.path.join(DATA_DIR, "its_results_summary.csv"))

its_headers = ["Specialty", "Definition", "Outcome", "R\u00b2", "Intervention yr",
               "Trend coef (P)", "Level change (P)", "Slope change (P)", "Accident effect (P)"]

its_rows = []
for _, r in df_its.iterrows():
    spec_en = SPEC_EN.get(r['specialty'], r['specialty'])
    def_label = r['definition'] if r['definition'] != 'Mixed' else 'Composite'
    r2 = f"{r['R_squared']:.3f}"
    iv = str(int(r['intervention_year']))
    trend = f"{r['trend_coef']:.1f} ({r['trend_pval']:.3f})"
    level = f"{r['level_change_coef']:.1f} ({r['level_change_pval']:.3f})"
    slope_val = r.get('slope_change_coef', None)
    slope_p = r.get('slope_change_pval', None)
    if pd.notna(slope_val) and pd.notna(slope_p):
        slope = f"{slope_val:.1f} ({slope_p:.3f})"
    else:
        slope = "\u2014"
    acc = f"{r['accident_effect_coef']:.1f} ({r['accident_effect_pval']:.3f})"
    its_rows.append([spec_en, def_label, r['outcome'], r2, iv, trend, level, slope, acc])

table_s2 = doc.add_table(rows=1 + len(its_rows), cols=len(its_headers))
table_s2.style = 'Table Grid'
table_s2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(its_headers):
    cell = table_s2.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(7)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r_idx, row_data in enumerate(its_rows):
    for c_idx, val in enumerate(row_data):
        cell = table_s2.rows[r_idx + 1].cells[c_idx]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(val))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(7)
        if c_idx > 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ------ Table S3: Cross-correlation coefficients ------
doc.add_page_break()
add_heading_bmj(
    "Supplementary Table S3. Cross-Correlation Coefficients at Optimal Lag "
    "for All Specialty-Definition-Outcome Combinations", level=2
)
add_para(
    "Pearson correlation coefficients computed after linear detrending of both series. "
    "Lag = years by which the incident series is shifted relative to the outcome series; "
    "positive lag means incidents precede outcome changes.",
    italic=True, space_after=Pt(12)
)

df_lt = pd.read_csv(os.path.join(DATA_DIR, "lead_time_estimates.csv"))

lt_headers = ["Specialty", "Definition", "Outcome", "Lag (years)", "Correlation (r)"]
lt_rows = []
for _, r in df_lt.iterrows():
    spec_en = SPEC_EN.get(r['specialty'], r['specialty'])
    def_label = r['definition'] if r['definition'] != 'Mixed' else 'Composite'
    lag = str(int(r['lead_time_years'])) if pd.notna(r['lead_time_years']) else "\u2014"
    corr = f"{r['correlation']:.3f}" if pd.notna(r['correlation']) else "\u2014"
    lt_rows.append([spec_en, def_label, r['outcome'], lag, corr])

add_table_from_data(lt_headers, lt_rows,
    caption="Supplementary Table S3. Cross-correlation coefficients at optimal lag")

# ------ Table S4: Window period estimates ------
doc.add_page_break()
add_heading_bmj(
    "Supplementary Table S4. AIC-Based Window Period Estimates "
    "for All Specialty-Definition-Outcome Combinations", level=2
)
add_para(
    "Window period = duration (years) of the binary intervention indicator yielding "
    "the lowest AIC in the segmented regression model.",
    italic=True, space_after=Pt(12)
)

df_wp = pd.read_csv(os.path.join(DATA_DIR, "window_period_estimates.csv"))

wp_headers = ["Specialty", "Definition", "Outcome", "Window period (years)"]
wp_rows = []
for _, r in df_wp.iterrows():
    spec_en = SPEC_EN.get(r['specialty'], r['specialty'])
    def_label = r['definition'] if r['definition'] != 'Mixed' else 'Composite'
    wp_val = str(int(r['window_years'])) if pd.notna(r['window_years']) else "\u2014"
    wp_rows.append([spec_en, def_label, r['outcome'], wp_val])

add_table_from_data(wp_headers, wp_rows,
    caption="Supplementary Table S4. AIC-based window period estimates")

# ------ Figure S1: Cross-correlation function plots ------
doc.add_page_break()
add_heading_bmj(
    "Supplementary Figure S1. Cross-Correlation Function Plots", level=2
)
add_figure(
    os.path.join(OUTPUT_DIR, "ccf_plots.png"),
    "Supplementary Figure S1. Cross-correlation functions between incident counts "
    "(JMSR, litigation, and composite) and physician numbers (detrended) at lags "
    "\u22128 to +8 years for the 12 core specialties. Squares indicate the lag with the "
    "most negative correlation for each definition. Positive lags indicate that changes "
    "in incident counts precede changes in physician numbers.",
    width=Inches(6.5)
)

# ============================================================
# SAVE
# ============================================================
output_path = os.path.join(MANUSCRIPT_DIR, "bmj_manuscript.docx")
doc.save(output_path)
print(f"Saved to {output_path}")

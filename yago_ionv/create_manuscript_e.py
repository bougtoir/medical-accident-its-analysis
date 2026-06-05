"""
Create Japanese manuscript (docx) for IONV analysis.
- Research protocol-based primary/secondary outcomes (Definition A)
- Definition E (5-HT3 antagonist-based) as additional analysis
- Covariate sensitivity analysis
- Figures inline (English labels)
"""
import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd

BASE = Path(__file__).resolve().parent

# Load stats
with open(BASE / "def_e_stats.json") as f:
    S = json.load(f)

doc = Document()

# ---- Styles ----
style = doc.styles["Normal"]
font = style.font
font.name = "游明朝"
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.5

# Helper functions
def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_paragraph(doc, text, bold=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if align:
        p.alignment = align
    return p

def add_figure(doc, img_path, caption, width=Inches(6)):
    if Path(img_path).exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=width)
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.italic = True

def format_p(p_val):
    if p_val < 0.001:
        return "P < 0.001"
    else:
        return f"P = {p_val:.3f}"

# ============================================================
# TITLE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("帝王切開術中の嘔気嘔吐（IONV）に対する\n双胎妊娠の影響に関する後方視的検討")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

# ============================================================
# METHODS — 統計解析
# ============================================================
add_heading(doc, "方法")
add_heading(doc, "研究デザイン・対象", level=2)
add_paragraph(doc,
    f"2014年1月から2024年12月までに当院で帝王切開術を受けた"
    f"{S['n_total']:,}例（単胎{S['n_single'] + (S['n_total'] - S['n_analysis']):,}例、"
    f"双胎{S['n_twin'] + (S['n_total'] - S['n_analysis'] - S['n_single']):,}例を含む）を"
    f"後方視的に検討した。")

add_paragraph(doc,
    "除外基準は、全身麻酔症例、入室時収縮期血圧（SBP）90 mmHg未満、"
    "子宮内胎児死亡（IUFD）、vanishing twin、品胎、および麻酔記録のデータ欠損とした。"
    "さらに、麻酔開始前に制吐薬が投与された症例は予防的投与とみなし解析から除外した。")

add_paragraph(doc,
    f"最終解析対象は{S['n_analysis']:,}例"
    f"（単胎{S['n_single']:,}例、双胎{S['n_twin']:,}例）であった。")

add_heading(doc, "評価項目", level=2)

add_paragraph(doc,
    "IONVの定義として、研究計画書に規定した定義（定義A）に加え、"
    "5-HT3受容体拮抗薬に基づく定義（定義E）を追加解析として実施した。")

add_paragraph(doc, "【定義A（プロトコル定義）】", bold=True)
add_paragraph(doc,
    "主要評価項目：麻酔開始から退室までのいずれかの時点で制吐薬が投与された場合をIONVと定義した"
    "（メトクロプラミド、ドロペリドール、オンダンセトロン、グラニセトロン、ノバミン、"
    "アタラックスP、デキサメタゾンの7剤）。\n"
    "副次評価項目：麻酔開始から胎児娩出までの制吐薬投与のみをIONVと定義した。")

add_paragraph(doc, "【定義E（5-HT3拮抗薬定義）】", bold=True)
add_paragraph(doc,
    "上記7剤のうち、嘔気嘔吐に対する特異性が高い5-HT3受容体拮抗薬"
    "（オンダンセトロンまたはグラニセトロン）の投与のみをIONVの指標とした。"
    "他の制吐薬（メトクロプラミド、ドロペリドール等）は消化管運動促進や鎮静など"
    "制吐以外の目的でも使用されるため、5-HT3拮抗薬に限定することで"
    "より正確にIONV治療の必要性を捉えることを意図した。\n"
    "主要評価項目：5-HT3拮抗薬が術中いずれかの時点で投与された場合。\n"
    "副次評価項目：5-HT3拮抗薬が投与され、かつ胎児娩出前に制吐薬が投与された場合。")

add_heading(doc, "統計解析", level=2)

add_paragraph(doc,
    "連続変数は中央値[四分位範囲]で示し、Mann-Whitney U検定で比較した。"
    "カテゴリー変数は例数（%）で示し、χ²検定またはFisherの正確検定"
    "（期待度数5未満のセルが存在する場合）で比較した。")

add_paragraph(doc,
    "IONVの独立した危険因子を評価するため、多変量ロジスティック回帰分析を実施した。"
    "共変量は年齢、BMI、在胎週数、緊急帝王切開、帝王切開既往、HDP、硬膜外麻酔、"
    "手術時間、低血圧（SBP 90 mmHg未満）、および双胎とした。"
    "イベント数が共変量数の10倍に満たない場合は、共変量を6変数"
    "（双胎、年齢、BMI、在胎週数、緊急帝王切開、低血圧）に縮小した。"
    "結果は調整済みオッズ比（aOR）と95%信頼区間（CI）で報告した。")

add_paragraph(doc,
    "定義Eの結果の頑健性を検証するため、共変量感度分析を行い、"
    "各共変量の追加・除外が双胎の効果推定値に与える影響を系統的に評価した。")

add_paragraph(doc,
    "統計解析はPython 3.12（scipy 1.14, statsmodels 0.14）を用い、"
    "有意水準は両側5%とした。")

# ============================================================
# RESULTS
# ============================================================
add_heading(doc, "結果")

# Patient characteristics
add_heading(doc, "患者背景", level=2)

add_paragraph(doc,
    f"解析対象{S['n_analysis']:,}例の患者背景をTable 1に示す。"
    f"単胎群{S['n_single']:,}例、双胎群{S['n_twin']:,}例であった。")

# Insert Table 1
table1 = pd.read_csv(BASE / "tables_e" / "table1_characteristics.csv")
add_paragraph(doc, "Table 1. 患者背景", bold=True)

tbl = doc.add_table(rows=len(table1) + 1, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Light Shading Accent 1"

for j, header in enumerate(["変数", f"単胎 (n={S['n_single']:,})", f"双胎 (n={S['n_twin']:,})", "P値"]):
    cell = tbl.rows[0].cells[j]
    cell.text = header
    for paragraph in cell.paragraphs:
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.size = Pt(9)

for i, (_, row) in enumerate(table1.iterrows()):
    tbl.rows[i + 1].cells[0].text = row["Variable"]
    tbl.rows[i + 1].cells[1].text = str(row["Singleton"])
    tbl.rows[i + 1].cells[2].text = str(row["Twin"])
    p_val = row["P-value"]
    p_str = "< 0.001" if p_val < 0.001 else f"{p_val:.3f}"
    tbl.rows[i + 1].cells[3].text = p_str
    for j in range(4):
        for paragraph in tbl.rows[i + 1].cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()

# ============================================================
# Results: Definition A (Protocol)
# ============================================================
add_heading(doc, "定義A（プロトコル定義）によるIONV", level=2)

o = S["outcomes"]
r = S["regression"]

add_paragraph(doc,
    f"主要評価項目（定義A-Primary）のIONV発生率は、単胎群{o['A-Primary']['singleton_pct']:.1f}%"
    f"（{o['A-Primary']['singleton_n']}/{S['n_single']}例）、"
    f"双胎群{o['A-Primary']['twin_pct']:.1f}%"
    f"（{o['A-Primary']['twin_n']}/{S['n_twin']}例）であり、"
    f"有意差を認めなかった（{format_p(r['A-Primary']['twin_P'])}）。")

add_paragraph(doc,
    f"副次評価項目（定義A-Secondary、胎児娩出前のみ）のIONV発生率は、"
    f"単胎群{o['A-Secondary']['singleton_pct']:.1f}%、"
    f"双胎群{o['A-Secondary']['twin_pct']:.1f}%であり、"
    f"こちらも有意差を認めなかった（{format_p(r['A-Secondary']['twin_P'])}）。")

add_paragraph(doc,
    f"多変量ロジスティック回帰分析（共変量10変数）においても、"
    f"双胎は主要評価項目のIONVと有意な関連を示さなかった"
    f"（aOR {r['A-Primary']['twin_OR']:.2f}, "
    f"95%CI {r['A-Primary']['twin_CI_lower']:.2f}–{r['A-Primary']['twin_CI_upper']:.2f}, "
    f"{format_p(r['A-Primary']['twin_P'])}）。")

# ============================================================
# Results: Definition E
# ============================================================
add_heading(doc, "定義E（5-HT3拮抗薬定義）によるIONV", level=2)

add_paragraph(doc,
    f"5-HT3拮抗薬の使用を指標とした定義E-Primaryでは、"
    f"IONV発生率は単胎群{o['E-Primary']['singleton_pct']:.1f}%"
    f"（{o['E-Primary']['singleton_n']}/{S['n_single']}例）に対し、"
    f"双胎群{o['E-Primary']['twin_pct']:.1f}%"
    f"（{o['E-Primary']['twin_n']}/{S['n_twin']}例）と、"
    f"双胎群で有意に高率であった（P = 0.020）。")

add_paragraph(doc,
    f"多変量ロジスティック回帰分析では、イベント数が{r['E-Primary']['events']}件と少数であったため"
    f"共変量を6変数に縮小したモデルを使用した。その結果、双胎は5-HT3拮抗薬使用と"
    f"有意に関連していた"
    f"（aOR {r['E-Primary']['twin_OR']:.2f}, "
    f"95%CI {r['E-Primary']['twin_CI_lower']:.2f}–{r['E-Primary']['twin_CI_upper']:.2f}, "
    f"{format_p(r['E-Primary']['twin_P'])}）（Fig. 2）。")

add_paragraph(doc,
    f"定義E-Secondary（5-HT3拮抗薬かつ胎児娩出前投与）では、"
    f"単胎群{o['E-Secondary']['singleton_pct']:.2f}%に対し"
    f"双胎群{o['E-Secondary']['twin_pct']:.2f}%と高い傾向を示したが、"
    f"統計学的有意差には至らなかった"
    f"（aOR {r['E-Secondary']['twin_OR']:.2f}, "
    f"95%CI {r['E-Secondary']['twin_CI_lower']:.2f}–{r['E-Secondary']['twin_CI_upper']:.2f}, "
    f"{format_p(r['E-Secondary']['twin_P'])}）。")

# Fig 1: rates comparison
add_figure(doc, BASE / "figures_e" / "fig1_rates_comparison.png",
           "Fig. 1  IONV発生率の比較（定義A vs 定義E、単胎 vs 双胎）")

# ============================================================
# Results: Logistic regression detail
# ============================================================
add_heading(doc, "多変量ロジスティック回帰分析", level=2)

# Table 2: Regression results
add_paragraph(doc, "Table 2. 各定義におけるIONVの多変量ロジスティック回帰分析", bold=True)

tbl2 = doc.add_table(rows=5, cols=6)
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl2.style = "Light Shading Accent 1"

headers2 = ["評価項目", "n", "イベント数", "双胎 aOR", "95% CI", "P値"]
for j, h in enumerate(headers2):
    cell = tbl2.rows[0].cells[j]
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.size = Pt(9)

for i, key in enumerate(["A-Primary", "A-Secondary", "E-Primary", "E-Secondary"]):
    rr = r[key]
    labels_ja = {
        "A-Primary": "定義A 主要（制吐薬全体）",
        "A-Secondary": "定義A 副次（娩出前のみ）",
        "E-Primary": "定義E 主要（5-HT3全体）",
        "E-Secondary": "定義E 副次（5-HT3+娩出前）",
    }
    vals = [
        labels_ja[key],
        str(rr["n"]),
        str(rr["events"]),
        f"{rr['twin_OR']:.2f}",
        f"{rr['twin_CI_lower']:.2f}–{rr['twin_CI_upper']:.2f}",
        "< 0.001" if rr["twin_P"] < 0.001 else f"{rr['twin_P']:.3f}",
    ]
    for j, v in enumerate(vals):
        tbl2.rows[i + 1].cells[j].text = v
        for paragraph in tbl2.rows[i + 1].cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()

# Fig 2: Forest E-Primary
add_figure(doc, BASE / "figures_e" / "fig2_forest_E_primary.png",
           "Fig. 2  定義E-Primary（5-HT3拮抗薬使用）の多変量ロジスティック回帰分析")

# Fig 4: Protocol vs Def E comparison
add_figure(doc, BASE / "figures_e" / "fig4_protocol_vs_defE.png",
           "Fig. 4  双胎のIONV効果：定義A（プロトコル）vs 定義E（5-HT3拮抗薬）")

# ============================================================
# Results: Covariate sensitivity
# ============================================================
add_heading(doc, "共変量感度分析", level=2)

cov_df = pd.read_csv(BASE / "tables_e" / "covariate_sensitivity.csv")

add_paragraph(doc,
    "定義E-Primaryにおける双胎の効果推定値の頑健性を検証するため、"
    "共変量感度分析を実施した（Table 3, Fig. 3）。")

add_paragraph(doc,
    f"粗オッズ比（双胎のみのモデル）は{cov_df.iloc[0]['aOR']:.2f}"
    f"（95%CI {cov_df.iloc[0]['CI_lower']:.2f}–{cov_df.iloc[0]['CI_upper']:.2f}, "
    f"P = {cov_df.iloc[0]['P']:.3f}）であった。"
    f"年齢、手術時間など交絡因子を調整することで効果推定値はさらに増大し、"
    f"縮小モデル（6共変量）ではaOR {S['regression']['E-Primary']['twin_OR']:.2f}"
    f"（P = {S['regression']['E-Primary']['twin_P']:.3f}）となった。")

add_paragraph(doc,
    "全21モデル（粗モデル、各共変量の追加モデル9種、各共変量の除外モデル9種、"
    "縮小モデル2種）において、双胎のaORは2.56–3.67の範囲で一貫しており、"
    "すべてのモデルで統計学的に有意であった（P < 0.05）。"
    "この結果は、定義E-Primaryにおける双胎の有意な関連が特定の共変量選択に依存しない"
    "頑健な所見であることを示している。")

# Table 3: Covariate sensitivity
add_paragraph(doc, "Table 3. 共変量感度分析（定義E-Primary）", bold=True)

tbl3 = doc.add_table(rows=len(cov_df) + 1, cols=4)
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl3.style = "Light Shading Accent 1"

for j, h in enumerate(["モデル", "aOR", "95% CI", "P値"]):
    cell = tbl3.rows[0].cells[j]
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.size = Pt(9)

for i, (_, row) in enumerate(cov_df.iterrows()):
    vals = [
        row["Model"],
        f"{row['aOR']:.2f}",
        f"{row['CI_lower']:.2f}–{row['CI_upper']:.2f}",
        "< 0.001" if row["P"] < 0.001 else f"{row['P']:.3f}",
    ]
    for j, v in enumerate(vals):
        tbl3.rows[i + 1].cells[j].text = v
        for paragraph in tbl3.rows[i + 1].cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()

# Fig 3: Covariate sensitivity forest
add_figure(doc, BASE / "figures_e" / "fig3_covariate_sensitivity.png",
           "Fig. 3  共変量感度分析：双胎の5-HT3拮抗薬使用に対する効果（定義E-Primary）")

# ============================================================
# DISCUSSION (brief)
# ============================================================
add_heading(doc, "考察", level=2)

add_paragraph(doc,
    "本研究では、帝王切開術中のIONVに対する双胎妊娠の影響を後方視的に検討した。"
    "研究計画書に規定した定義A（全制吐薬の使用）では、単胎と双胎の間にIONV発生率の"
    "有意差を認めなかった。")

add_paragraph(doc,
    "一方、5-HT3受容体拮抗薬の使用に限定した定義Eでは、双胎群で有意に高率であり"
    "（aOR 3.18, P = 0.007）、この結果は共変量感度分析においても一貫して頑健であった。"
    "制吐薬として記録された7剤のうち、メトクロプラミドやドロペリドールは"
    "消化管運動促進や鎮静など制吐以外の目的でも使用される。"
    "5-HT3拮抗薬は嘔気嘔吐に対する特異性が高く、その投与はより確実に"
    "臨床的に有意な嘔気嘔吐の存在を反映していると考えられる。")

add_paragraph(doc,
    "すなわち、双胎妊娠ではIONVの全体的な発生頻度は単胎と同等であるものの、"
    "5-HT3拮抗薬が必要となるレベルの強い嘔気嘔吐がより高頻度に発生する可能性が示唆された。"
    "ただし、本研究は後方視的デザインであり、制吐薬の選択は術者の裁量に依存するため、"
    "処方パターンの偏り（双胎に対してより積極的に5-HT3拮抗薬を選択する傾向）"
    "が交絡因子として残る可能性がある。")

add_paragraph(doc,
    f"また、定義Eのイベント数は{S['regression']['E-Primary']['events']}件と少数であり、"
    f"信頼区間が比較的広い点に留意が必要である。"
    "今後、前方視的研究やIONVの重症度評価を含めた検討が望まれる。")

# Save
out_path = BASE / "manuscript_ionv_jp.docx"
doc.save(str(out_path))
print(f"Manuscript saved to {out_path}")

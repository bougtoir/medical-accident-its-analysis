"""
Create Japanese manuscript (docx) for exclusion sensitivity analysis.
- Excludes: Emergency CS, Prior CS, HDP, Preoperative steroid
- Re-analyzes broad/narrow IONV in elective low-risk subgroup
- Figures inline (English labels)
"""
import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd

BASE = Path(__file__).resolve().parent

# Load stats
with open(BASE / "excl_sensitivity_stats.json") as f:
    S = json.load(f)

doc = Document()

# ---- Styles ----
style = doc.styles["Normal"]
font = style.font
font.name = "游明朝"
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.5

# Helper functions
def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_paragraph(doc, text, bold=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if align:
        p.alignment = align
    return p

def add_figure(doc, img_path, caption, width=Inches(6)):
    if Path(img_path).exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=width)
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.italic = True

def format_p(p_val):
    if p_val < 0.001:
        return "P < 0.001"
    else:
        return f"P = {p_val:.3f}"

# ============================================================
# TITLE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(
    "帝王切開術中の嘔気嘔吐（IONV）に対する双胎妊娠の影響：\n"
    "予定手術・低リスク症例に限定した感度分析")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

# ============================================================
# METHODS
# ============================================================
add_heading(doc, "方法")
add_heading(doc, "研究デザイン・対象", level=2)
add_paragraph(doc,
    f"2014年1月から2024年12月までに当院で帝王切開術を受けた"
    f"{S['n_total']:,}例（単胎{S['n_single_raw']:,}例、"
    f"双胎{S['n_twin_raw']:,}例を含む）を後方視的に検討した。")

add_paragraph(doc,
    "除外基準は、全身麻酔症例、入室時収縮期血圧（SBP）90 mmHg未満、"
    "子宮内胎児死亡（IUFD）、vanishing twin、品胎、および麻酔記録のデータ欠損とした。"
    "さらに、麻酔開始前に制吐薬が投与された症例は予防的投与とみなし解析から除外した。")

add_paragraph(doc,
    f"上記基準による解析対象は{S['n_base']:,}例"
    f"（単胎{S['n_base_single']:,}例、双胎{S['n_base_twin']:,}例）であった。")

add_heading(doc, "追加除外基準（感度分析）", level=2)
add_paragraph(doc,
    "IONVに影響しうる産科合併症・手術関連因子の交絡を排除するため、"
    "以下の4条件に該当する症例をさらに除外した：\n"
    "（1）緊急帝王切開、（2）帝王切開の既往、（3）妊娠高血圧症候群（HDP）、"
    "（4）術前1週間以内のステロイド使用。")

add_paragraph(doc,
    f"追加除外により{S['n_excl_additional']:,}例"
    f"（単胎{S['n_excl_additional_single']:,}例、"
    f"双胎{S['n_excl_additional_twin']:,}例）が除外され、"
    f"最終解析対象は{S['n_analysis']:,}例"
    f"（単胎{S['n_single']:,}例、双胎{S['n_twin']:,}例）であった。")

add_paragraph(doc,
    f"除外内訳：緊急帝王切開{S['exclusion_counts']['Emergency CS']:,}例、"
    f"帝王切開既往{S['exclusion_counts']['Prior CS']:,}例、"
    f"HDP {S['exclusion_counts']['HDP']:,}例、"
    f"術前ステロイド{S['exclusion_counts']['Preoperative steroid']:,}例"
    f"（重複あり）。")

add_heading(doc, "評価項目", level=2)

add_paragraph(doc,
    "IONVの定義として、制吐薬全般の使用を指標とする広義の定義に加え、"
    "5-HT3受容体拮抗薬の使用に限定した狭義の定義による追加解析を実施した。")

add_paragraph(doc, "【制吐薬（広義）】", bold=True)
add_paragraph(doc,
    "主要評価項目：麻酔開始から退室までのいずれかの時点で制吐薬が投与された場合をIONVと定義した"
    "（メトクロプラミド、ドロペリドール、オンダンセトロン、グラニセトロン、ノバミン、"
    "アタラックスP、デキサメタゾンの7剤）。\n"
    "副次評価項目：麻酔開始から胎児娩出までの制吐薬投与のみをIONVと定義した。")

add_paragraph(doc, "【制吐薬（狭義）：5-HT3受容体拮抗薬】", bold=True)
add_paragraph(doc,
    "上記7剤のうち、嘔気嘔吐に対する特異性が高い5-HT3受容体拮抗薬"
    "（オンダンセトロンまたはグラニセトロン）の投与のみをIONVの指標とした。")

add_heading(doc, "統計解析", level=2)

add_paragraph(doc,
    "連続変数は中央値[四分位範囲]で示し、Mann-Whitney U検定で比較した。"
    "カテゴリー変数は例数（%）で示し、χ²検定またはFisherの正確検定"
    "（期待度数5未満のセルが存在する場合）で比較した。")

add_paragraph(doc,
    "IONVの独立した危険因子を評価するため、多変量ロジスティック回帰分析を実施した。"
    "共変量は年齢、BMI、在胎週数、硬膜外麻酔、手術時間、低血圧（SBP 90 mmHg未満）、"
    "および双胎とした（7変数）。"
    "イベント数が共変量数の10倍に満たない場合は、共変量を5変数"
    "（双胎、年齢、BMI、在胎週数、低血圧）に縮小した。"
    "結果は調整済みオッズ比（aOR）と95%信頼区間（CI）で報告した。")

add_paragraph(doc,
    "制吐薬（狭義）の結果の頑健性を検証するため、共変量感度分析を行い、"
    "各共変量の追加・除外が双胎の効果推定値に与える影響を系統的に評価した。")

add_paragraph(doc,
    "統計解析はPython 3.12（scipy 1.14, statsmodels 0.14）を用い、"
    "有意水準は両側5%とした。")

# ============================================================
# RESULTS
# ============================================================
add_heading(doc, "結果")

add_heading(doc, "患者背景", level=2)

add_paragraph(doc,
    f"解析対象{S['n_analysis']:,}例の患者背景をTable 1に示す。"
    f"単胎群{S['n_single']:,}例、双胎群{S['n_twin']:,}例であった。"
    "本サブグループは予定手術かつ帝王切開既往・HDP・術前ステロイドのない"
    "低リスク症例のみで構成される。")

# Insert Table 1
table1 = pd.read_csv(BASE / "tables_excl" / "table1_characteristics.csv")
add_paragraph(doc, "Table 1. 患者背景（予定手術・低リスクサブグループ）", bold=True)

tbl = doc.add_table(rows=len(table1) + 1, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Light Shading Accent 1"

for j, header in enumerate(["変数", f"単胎 (n={S['n_single']:,})", f"双胎 (n={S['n_twin']:,})", "P値"]):
    cell = tbl.rows[0].cells[j]
    cell.text = header
    for paragraph in cell.paragraphs:
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.size = Pt(9)

for i, (_, row) in enumerate(table1.iterrows()):
    tbl.rows[i + 1].cells[0].text = row["Variable"]
    tbl.rows[i + 1].cells[1].text = str(row["Singleton"])
    tbl.rows[i + 1].cells[2].text = str(row["Twin"])
    p_val = row["P-value"]
    p_str = "< 0.001" if p_val < 0.001 else f"{p_val:.3f}"
    tbl.rows[i + 1].cells[3].text = p_str
    for j in range(4):
        for paragraph in tbl.rows[i + 1].cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()

# ============================================================
# Results: Broad definition
# ============================================================
add_heading(doc, "制吐薬（広義）によるIONV", level=2)

o = S["outcomes"]
r = S["regression"]

add_paragraph(doc,
    f"主要評価項目（制吐薬（広義））のIONV発生率は、単胎群{o['A-Primary']['singleton_pct']:.1f}%"
    f"（{o['A-Primary']['singleton_n']}/{S['n_single']}例）、"
    f"双胎群{o['A-Primary']['twin_pct']:.1f}%"
    f"（{o['A-Primary']['twin_n']}/{S['n_twin']}例）であった。"
    f"多変量解析では、双胎は制吐薬使用と有意な関連を示さなかった"
    f"（aOR {r['A-Primary']['twin_OR']:.2f}, "
    f"95%CI {r['A-Primary']['twin_CI_lower']:.2f}–{r['A-Primary']['twin_CI_upper']:.2f}, "
    f"{format_p(r['A-Primary']['twin_P'])}）。")

add_paragraph(doc,
    f"副次評価項目（胎児娩出前のみ）のIONV発生率は、"
    f"単胎群{o['A-Secondary']['singleton_pct']:.1f}%、"
    f"双胎群{o['A-Secondary']['twin_pct']:.1f}%であり、"
    f"こちらも有意差を認めなかった"
    f"（aOR {r['A-Secondary']['twin_OR']:.2f}, "
    f"95%CI {r['A-Secondary']['twin_CI_lower']:.2f}–{r['A-Secondary']['twin_CI_upper']:.2f}, "
    f"{format_p(r['A-Secondary']['twin_P'])}）。")

# ============================================================
# Results: Narrow definition
# ============================================================
add_heading(doc, "制吐薬（狭義）によるIONV", level=2)

add_paragraph(doc,
    f"5-HT3拮抗薬の使用を指標とした狭義の主要評価項目では、"
    f"IONV発生率は単胎群{o['E-Primary']['singleton_pct']:.1f}%"
    f"（{o['E-Primary']['singleton_n']}/{S['n_single']}例）に対し、"
    f"双胎群{o['E-Primary']['twin_pct']:.1f}%"
    f"（{o['E-Primary']['twin_n']}/{S['n_twin']}例）と、"
    f"双胎群で有意に高率であった。")

add_paragraph(doc,
    f"多変量ロジスティック回帰分析では、イベント数が{r['E-Primary']['events']}件と"
    f"少数であったため共変量を5変数に縮小したモデルを使用した。その結果、双胎は"
    f"5-HT3拮抗薬使用と有意に関連していた"
    f"（aOR {r['E-Primary']['twin_OR']:.2f}, "
    f"95%CI {r['E-Primary']['twin_CI_lower']:.2f}–{r['E-Primary']['twin_CI_upper']:.2f}, "
    f"{format_p(r['E-Primary']['twin_P'])}）（Fig. 2）。")

add_paragraph(doc,
    f"狭義の副次評価項目（5-HT3拮抗薬かつ胎児娩出前投与）では、"
    f"単胎群{o['E-Secondary']['singleton_pct']:.2f}%に対し"
    f"双胎群{o['E-Secondary']['twin_pct']:.2f}%と高い傾向を示したが、"
    f"イベント数が{r['E-Secondary']['events']}件と極めて少なく、"
    f"統計学的有意差には至らなかった"
    f"（aOR {r['E-Secondary']['twin_OR']:.2f}, "
    f"95%CI {r['E-Secondary']['twin_CI_lower']:.2f}–{r['E-Secondary']['twin_CI_upper']:.2f}, "
    f"{format_p(r['E-Secondary']['twin_P'])}）。")

# Fig 1: rates comparison
add_figure(doc, BASE / "figures_excl" / "fig1_rates_comparison.png",
           "Fig. 1  IONV発生率の比較（制吐薬（広義）vs 制吐薬（狭義）、単胎 vs 双胎）\n"
           "予定手術・低リスクサブグループ")

# ============================================================
# Regression detail table
# ============================================================
add_heading(doc, "多変量ロジスティック回帰分析", level=2)

add_paragraph(doc, "Table 2. IONVの多変量ロジスティック回帰分析（予定手術・低リスクサブグループ）", bold=True)

tbl2 = doc.add_table(rows=5, cols=6)
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl2.style = "Light Shading Accent 1"

headers2 = ["評価項目", "n", "イベント数", "双胎 aOR", "95% CI", "P値"]
for j, h in enumerate(headers2):
    cell = tbl2.rows[0].cells[j]
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.size = Pt(9)

for i, key in enumerate(["A-Primary", "A-Secondary", "E-Primary", "E-Secondary"]):
    rr = r[key]
    labels_ja = {
        "A-Primary": "広義 主要（制吐薬全体）",
        "A-Secondary": "広義 副次（娩出前のみ）",
        "E-Primary": "狭義 主要（5-HT3拮抗薬）",
        "E-Secondary": "狭義 副次（5-HT3+娩出前）",
    }
    vals = [
        labels_ja[key],
        str(rr["n"]),
        str(rr["events"]),
        f"{rr['twin_OR']:.2f}",
        f"{rr['twin_CI_lower']:.2f}–{rr['twin_CI_upper']:.2f}",
        "< 0.001" if rr["twin_P"] < 0.001 else f"{rr['twin_P']:.3f}",
    ]
    for j, v in enumerate(vals):
        tbl2.rows[i + 1].cells[j].text = v
        for paragraph in tbl2.rows[i + 1].cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()

# Fig 2: Forest narrow primary
add_figure(doc, BASE / "figures_excl" / "fig2_forest_narrow_primary.png",
           "Fig. 2  制吐薬（狭義）使用の多変量ロジスティック回帰分析\n"
           "（予定手術・低リスクサブグループ）")

# Fig 4: Broad vs Narrow comparison
add_figure(doc, BASE / "figures_excl" / "fig4_broad_vs_narrow.png",
           "Fig. 4  双胎のIONV効果：制吐薬（広義）vs 制吐薬（狭義）\n"
           "（予定手術・低リスクサブグループ）")

# ============================================================
# Covariate sensitivity
# ============================================================
add_heading(doc, "共変量感度分析", level=2)

cov_df = pd.read_csv(BASE / "tables_excl" / "covariate_sensitivity.csv")

add_paragraph(doc,
    "制吐薬（狭義）における双胎の効果推定値の頑健性を検証するため、"
    "共変量感度分析を実施した（Table 3, Fig. 3）。")

n_sig = (cov_df["P"] < 0.05).sum()
n_total_models = len(cov_df)
add_paragraph(doc,
    f"全{n_total_models}モデルにおいて、双胎のaORは"
    f"{cov_df['aOR'].min():.2f}–{cov_df['aOR'].max():.2f}の範囲であり、"
    f"{n_sig}/{n_total_models}モデルで統計学的に有意であった（P < 0.05）。"
    "この結果は、緊急帝王切開・帝王切開既往・HDP・術前ステロイドを除外した"
    "低リスク集団においても、双胎と5-HT3拮抗薬使用の有意な関連が"
    "頑健に再現されることを示している。")

# Table 3: Covariate sensitivity
add_paragraph(doc, "Table 3. 共変量感度分析（制吐薬（狭義）主要評価項目）", bold=True)

tbl3 = doc.add_table(rows=len(cov_df) + 1, cols=4)
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl3.style = "Light Shading Accent 1"

for j, h in enumerate(["モデル", "aOR", "95% CI", "P値"]):
    cell = tbl3.rows[0].cells[j]
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.size = Pt(9)

for i, (_, row) in enumerate(cov_df.iterrows()):
    vals = [
        row["Model"],
        f"{row['aOR']:.2f}",
        f"{row['CI_lower']:.2f}–{row['CI_upper']:.2f}",
        "< 0.001" if row["P"] < 0.001 else f"{row['P']:.3f}",
    ]
    for j, v in enumerate(vals):
        tbl3.rows[i + 1].cells[j].text = v
        for paragraph in tbl3.rows[i + 1].cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()

# Fig 3: Covariate sensitivity forest
add_figure(doc, BASE / "figures_excl" / "fig3_covariate_sensitivity.png",
           "Fig. 3  共変量感度分析：双胎の制吐薬（狭義）使用に対する効果\n"
           "（予定手術・低リスクサブグループ）")

# ============================================================
# DISCUSSION
# ============================================================
add_heading(doc, "考察", level=2)

add_paragraph(doc,
    "本感度分析では、緊急帝王切開、帝王切開既往、HDP、術前ステロイド使用の"
    "4条件を除外し、予定手術かつ低リスクの症例に限定して"
    "双胎妊娠のIONVへの影響を再検討した。")

add_paragraph(doc,
    "制吐薬全般の使用（広義）では全コホート解析と同様に有意差を認めなかった"
    f"（aOR {r['A-Primary']['twin_OR']:.2f}, {format_p(r['A-Primary']['twin_P'])}）。"
    f"一方、5-HT3拮抗薬の使用に限定した狭義の定義では、"
    f"双胎群で有意に高率であった"
    f"（aOR {r['E-Primary']['twin_OR']:.2f}, "
    f"95%CI {r['E-Primary']['twin_CI_lower']:.2f}–{r['E-Primary']['twin_CI_upper']:.2f}, "
    f"{format_p(r['E-Primary']['twin_P'])}）。")

add_paragraph(doc,
    "注目すべきは、交絡因子（緊急手術、既往手術、HDP、ステロイド）を除外したにも"
    "かかわらず、むしろ効果量が増大した点である"
    f"（全コホート aOR 3.18 → サブグループ aOR {r['E-Primary']['twin_OR']:.2f}）。"
    "これは、緊急手術やHDPなどIONVリスクを高める因子が単胎群に相対的に多く含まれ、"
    "双胎の独立効果を希釈していた可能性を示唆する。")

add_paragraph(doc,
    f"ただし、本サブグループ解析のイベント数は{r['E-Primary']['events']}件と極めて少なく、"
    "信頼区間が非常に広い点に留意が必要である。"
    "また、予定手術に限定することで対象が大幅に減少し"
    f"（{S['n_base']:,}例→{S['n_analysis']:,}例）、外的妥当性が限定される。"
    "本所見は全コホート解析の結果を支持するものであるが、"
    "少数イベントに基づく推定値として解釈上の注意が必要である。")

# Save
out_path = BASE / "manuscript_excl_sensitivity.docx"
doc.save(str(out_path))
print(f"Manuscript saved to {out_path}")

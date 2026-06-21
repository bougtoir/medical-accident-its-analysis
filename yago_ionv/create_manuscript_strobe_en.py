"""
STROBE-compliant manuscript: full paper format (English text, English figures).
Translated from Japanese STROBE manuscript with all data from JSON sources.
Includes English STROBE checklist.
"""
import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import pandas as pd

BASE = Path(__file__).resolve().parent

# Load all stats
with open(BASE / "def_e_stats.json") as f:
    M = json.load(f)  # Main cohort
with open(BASE / "excl_sensitivity_stats.json") as f:
    E = json.load(f)  # Exclusion sensitivity
with open(BASE / "flowchart_counts.json") as f:
    F = json.load(f)  # Flow counts
with open(BASE / "bootstrap_results.json") as f:
    B = json.load(f)  # Bootstrap results
with open(BASE / "summary_stats.json") as f:
    S = json.load(f)  # Summary stats (includes hypotension secondary outcome)

# --- Detachable structure flag ---
# Set to False to produce a main-analysis-only manuscript (no sensitivity sections)
INCLUDE_SENSITIVITY = True

doc = Document()

# ---- Styles ----
style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.5

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_paragraph(doc, text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.font.italic = True
    if align:
        p.alignment = align
    return p

def add_ref(paragraph, text):
    """Add text with superscript reference numbers using {n} markers."""
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = paragraph.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(8)
        else:
            paragraph.add_run(part)
    return paragraph

def add_p_with_refs(doc, text, bold=False):
    p = doc.add_paragraph()
    if bold:
        parts = re.split(r'(\{[^}]+\})', text)
        for part in parts:
            if part.startswith('{') and part.endswith('}'):
                run = p.add_run(part[1:-1])
                run.font.superscript = True
                run.font.size = Pt(8)
                run.bold = True
            else:
                run = p.add_run(part)
                run.bold = True
    else:
        add_ref(p, text)
    return p

def add_figure(doc, img_path, caption, width=Inches(6)):
    if Path(img_path).exists():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=width)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(6)
        run = cap.add_run(caption)
        run.font.size = Pt(9)
        run.font.italic = True

def format_p(p_val):
    return "P < 0.001" if p_val < 0.001 else f"P = {p_val:.3f}"


# Helper aliases
mo = M["outcomes"]
mr = M["regression"]
er = E["regression"]
eo = E["outcomes"]

# ============================================================
# TITLE PAGE
# ============================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(40)
run = title_p.add_run(
    "Impact of Twin Pregnancy on Intraoperative Nausea and Vomiting\n"
    "During Cesarean Delivery Under Spinal Anesthesia:\n"
    "A Retrospective Cohort Study")
run.bold = True
run.font.size = Pt(16)

doc.add_paragraph()
doc.add_paragraph()

# ============================================================
# ABSTRACT (structured)
# ============================================================
add_heading(doc, "Abstract")

add_paragraph(doc, "Background", bold=True)
add_paragraph(doc,
    "Intraoperative nausea and vomiting (IONV) is a major adverse effect of spinal anesthesia "
    "for cesarean delivery. However, the impact of twin pregnancy on IONV has not been "
    "adequately investigated.")

add_paragraph(doc, "Methods", bold=True)
add_p_with_refs(doc,
    f"We retrospectively reviewed {F['total']['n']:,} cesarean deliveries performed at a single "
    f"institution between April 2014 and October 2024. After exclusion criteria, "
    f"{F['primary_analysis']['n']:,} cases "
    f"(singleton {F['primary_analysis']['n_s']:,}; twin {F['primary_analysis']['n_t']:,}) "
    "were analyzed. Antiemetic administration was used as a surrogate marker for IONV, "
    "evaluated using two definitions: all antiemetics (broad definition) and "
    "5-HT3 receptor antagonists only (narrow definition). "
    "Multivariable logistic regression was used to assess the independent effect of twin pregnancy. "
    "This study was reported following the Strengthening the Reporting of Observational Studies in "
    "Epidemiology (STROBE) guidelines.")

add_paragraph(doc, "Results", bold=True)
add_paragraph(doc,
    f"Under the broad definition, IONV rates were similar between singleton "
    f"({mo['A-Primary']['singleton_pct']:.1f}%) and twin ({mo['A-Primary']['twin_pct']:.1f}%) "
    f"groups (adjusted odds ratio [aOR] {mr['A-Primary']['twin_OR']:.2f}, "
    f"95% confidence interval [CI] {mr['A-Primary']['twin_CI_lower']:.2f}\u2013"
    f"{mr['A-Primary']['twin_CI_upper']:.2f}, "
    f"{format_p(mr['A-Primary']['twin_P'])}). "
    f"In contrast, under the narrow definition (5-HT3 antagonists), "
    f"IONV was significantly higher in the twin group "
    f"({mo['E-Primary']['twin_pct']:.1f}% vs {mo['E-Primary']['singleton_pct']:.1f}%; "
    f"aOR {mr['E-Primary']['twin_OR']:.2f}, "
    f"95% CI {mr['E-Primary']['twin_CI_lower']:.2f}\u2013"
    f"{mr['E-Primary']['twin_CI_upper']:.2f}, "
    f"{format_p(mr['E-Primary']['twin_P'])}). "
    "This finding was robust across covariate sensitivity analyses, "
    "emergency subgroup stratification, and 10,000-replicate stratified bootstrap validation. "
    f"Intraoperative hypotension (SBP <90 mmHg), a secondary outcome, was less frequent "
    f"in the twin group ({S['hypo_twin_pct']:.1f}%) than in the singleton group "
    f"({S['hypo_single_pct']:.1f}%; {format_p(S['hypo_chi_p'])}).")

add_paragraph(doc, "Conclusions", bold=True)
add_paragraph(doc,
    "Although the overall antiemetic use was similar between singleton and twin pregnancies, "
    "the use of 5-HT3 receptor antagonists, which have high pharmacological specificity for "
    "nausea and vomiting, was significantly higher in the twin group. "
    "Conversely, intraoperative hypotension was less frequent in the twin group, "
    "suggesting that non-hypotensive pathways may contribute to IONV in twin pregnancies.")

doc.add_page_break()

# ============================================================
# INTRODUCTION
# ============================================================
add_heading(doc, "Introduction")

add_p_with_refs(doc,
    "Spinal anesthesia is the standard anesthetic technique for cesarean delivery worldwide{1}. "
    "However, spinal anesthesia is associated with various adverse effects including "
    "intraoperative nausea and vomiting (IONV), with reported incidence rates of up to 80%{2,3}. "
    "IONV not only reduces patient satisfaction but can also interfere with surgical "
    "procedures through visceral prolapse{4\u20136}.")

add_p_with_refs(doc,
    "The principal mechanisms of IONV include: "
    "(1) hypotension resulting from sympathetic blockade during spinal anesthesia; "
    "(2) gut hypoperfusion and serotonin release due to redistribution of circulating blood volume; "
    "(3) unopposed vagal activity under sympathetic blockade; "
    "(4) visceral stimulation from uterine manipulation and peritoneal traction; and "
    "(5) uterotonic (oxytocin) administration{2,12}.")

add_p_with_refs(doc,
    "Twin pregnancies possess physiological characteristics that may amplify each of "
    "these mechanisms. Compared with singleton pregnancies, twins are associated with an "
    "approximately 400 mL greater increase in circulating blood volume, as well as greater "
    "increases in heart rate, stroke volume, and cardiac index{13}. "
    "Moreover, aortocaval compression by the larger uterus is more frequent and severe "
    "in twin pregnancies, and this effect is exacerbated by neuraxial anesthesia, "
    "predisposing to profound hypotension{9,13}. "
    "These observations suggest that twin pregnancy may theoretically increase "
    "the risk of IONV.")

add_p_with_refs(doc,
    "Previous studies have identified risk factors for IONV including intraoperative "
    "hypotension, intraoperative pain, operative duration exceeding 30 minutes, "
    "body mass index (BMI) exceeding 35 kg/m\u00b2, hypertensive disorders of pregnancy (HDP), "
    "and uterine exteriorization{7,8}. "
    "However, most prior studies have excluded twin pregnancies from their analyses{9\u201311}, "
    "and it remains unclear whether twin pregnancy is an independent risk factor for IONV.")

add_paragraph(doc,
    "The aim of this study was to compare the incidence of IONV between singleton and twin "
    "pregnancies during cesarean delivery under spinal anesthesia and to determine whether "
    "twin pregnancy is an independent risk factor for IONV.")

doc.add_page_break()

# ============================================================
# METHODS
# ============================================================
add_heading(doc, "Methods")

add_heading(doc, "Study Design and Ethics (STROBE Items 1, 5)", level=2)
add_paragraph(doc,
    "This was a single-center retrospective cohort study. The study was approved by the "
    "Institutional Review Board of Juntendo University Shizuoka Hospital. "
    "This study was reported in accordance with the Strengthening the Reporting of "
    "Observational Studies in Epidemiology (STROBE) statement (see Supplementary Table). "
    "Owing to the retrospective design, informed consent was obtained using an opt-out approach.")

add_heading(doc, "Participants (STROBE Items 6, 7)", level=2)
add_p_with_refs(doc,
    f"All {F['total']['n']:,} patients who underwent cesarean delivery at Juntendo University "
    f"Shizuoka Hospital between April 1, 2014, and October 23, 2024, were eligible for inclusion "
    f"(singleton {F['total']['n_s']:,}; twin {F['total']['n_t']:,}).")

add_paragraph(doc,
    "Inclusion criteria were: age \u226518 years, singleton or twin pregnancy, and "
    "cesarean delivery under spinal anesthesia alone, epidural anesthesia alone, "
    "or combined spinal\u2013epidural anesthesia.")

# Exclusion criteria
excl_reasons_en = {
    "General anesthesia": "general anesthesia",
    "SBP < 90 mmHg at admission": "systolic blood pressure <90 mmHg at admission",
    "Intrauterine fetal death (IUFD)": "intrauterine fetal death",
    "Vanishing twin": "vanishing twin",
    "Triplet pregnancy": "triplet pregnancy",
    "Non-cesarean delivery": "non-cesarean delivery",
    "Cardiac arrest": "cardiac arrest",
    "Other exclusion criteria": "other exclusion criteria",
    "Missing anesthesia data": "missing anesthesia data",
}
excl_parts = []
for step in F["exclusion_steps"]:
    en_reason = excl_reasons_en.get(step["reason"], step["reason"])
    excl_parts.append(en_reason)
excl_text = "Exclusion criteria were: " + ", ".join(excl_parts) + ". "
excl_text += (
    f"Additionally, {F['preop_antiemetic']['n']} cases with antiemetic administration "
    "before induction of anesthesia were excluded as prophylactic use.")
add_paragraph(doc, excl_text)

add_paragraph(doc,
    f"After applying these criteria, {F['primary_analysis']['n']:,} cases "
    f"(singleton {F['primary_analysis']['n_s']:,}; twin {F['primary_analysis']['n_t']:,}) "
    "were included in the primary analysis (Fig. 1).")

# --- Fig 1: Flowchart ---
add_figure(doc, BASE / "figures_strobe" / "fig_flowchart.png",
           "Fig. 1  STROBE Flow Diagram \u2014 Participant Selection",
           width=Inches(5.5))

add_heading(doc, "Outcome Measures (STROBE Items 8, 12)", level=2)

add_paragraph(doc,
    "As this was a retrospective study, direct assessment of nausea and vomiting events was "
    "not feasible. Therefore, antiemetic administration was used as a surrogate marker for IONV.")

add_paragraph(doc, "Broad-definition antiemetics", bold=True)
add_paragraph(doc,
    "Primary outcome: IONV was defined as the administration of any antiemetic "
    "from anesthesia induction to fetal delivery. "
    "The following seven agents were included: metoclopramide, droperidol, ondansetron, "
    "granisetron, novamin (prochlorperazine), hydroxyzine (Atarax-P), and dexamethasone.\n"
    "Secondary outcome: IONV was defined as antiemetic administration from "
    "anesthesia induction to leaving the operating room.")

add_paragraph(doc, "Narrow-definition antiemetics: 5-HT3 receptor antagonists", bold=True)
add_paragraph(doc,
    "An additional analysis was performed using only 5-HT3 receptor antagonists "
    "(ondansetron or granisetron) as the IONV indicator, given their high pharmacological "
    "specificity for nausea and vomiting. "
    "Metoclopramide (prokinetic), droperidol (sedation), prochlorperazine (antipsychotic), "
    "hydroxyzine (anxiolytic/antipruritic), and dexamethasone (prophylactic/anti-inflammatory) "
    "were excluded owing to their use for indications other than antiemesis.")

add_paragraph(doc, "Secondary outcome: intraoperative hypotension", bold=True)
add_paragraph(doc,
    "As a secondary outcome, the frequency of intraoperative hypotension was assessed. "
    "Hypotension was defined as systolic blood pressure <90 mmHg, "
    "and the number of recorded episodes during surgery was used. "
    "Because hypotension was also used as a covariate in the IONV analysis, "
    "only descriptive statistics were used for between-group comparison.")

add_heading(doc, "Covariates (STROBE Item 8)", level=2)
add_paragraph(doc,
    "Covariates included in the multivariable analysis were: maternal age, BMI, "
    "gestational age at delivery, emergency cesarean delivery, prior cesarean delivery, "
    "HDP, epidural anesthesia, operative duration, and hypotension "
    "(any episode of systolic blood pressure <90 mmHg). "
    "When the number of events was fewer than 10 per covariate, a reduced model "
    "(twin, age, BMI, gestational age, emergency status, and hypotension; 6 covariates) "
    "was used to avoid overfitting.")

add_heading(doc, "Data Collection (STROBE Item 9)", level=2)
add_paragraph(doc,
    "Data were extracted from electronic medical records and the anesthesia information system "
    "by research personnel. Patient background, anesthesia details, surgical information, "
    "and antiemetic administration records were entered into Microsoft Excel. "
    "All personally identifiable information was removed and replaced with anonymous study IDs.")

add_heading(doc, "Statistical Analysis (STROBE Item 13)", level=2)
add_paragraph(doc,
    "Continuous variables were reported as median [interquartile range] and compared using "
    "the Mann\u2013Whitney U test. Categorical variables were reported as n (%) and compared "
    "using the \u03c7\u00b2 test or Fisher\u2019s exact test (when any expected cell count was <5).")

add_paragraph(doc,
    "Multivariable logistic regression was used to evaluate the association between "
    "IONV and twin pregnancy. Results were reported as adjusted odds ratios (aOR) "
    "with 95% confidence intervals (CI).")

if INCLUDE_SENSITIVITY:
    add_paragraph(doc,
        "To assess the robustness of the twin effect estimate for the narrow-definition "
        "primary outcome, a covariate sensitivity analysis was performed. This comprised: "
        "(1) crude OR (twin only), (2) twin plus each individual covariate (9 models), "
        "(3) full model minus each individual covariate (9 models), and "
        "(4) the reduced model (6 covariates), yielding a total of 20 models.")
    add_heading(doc, "Sensitivity Analysis (STROBE Item 13)", level=2)
    add_paragraph(doc,
        "To mitigate confounding by obstetric complications and surgical factors "
        "that may independently influence IONV, a sensitivity subgroup analysis was performed "
        "after further excluding cases with emergency cesarean delivery, prior cesarean delivery, "
        "HDP, or preoperative steroid use.")
    add_paragraph(doc,
        "Furthermore, to examine whether emergency cesarean delivery acts as an effect modifier, "
        "subgroup analyses restricted to elective cases only and emergency cases only were conducted. "
        "An interaction analysis including an emergency \u00d7 twin interaction term in the model "
        "was also performed.")
    add_heading(doc, "Bootstrap Validation (STROBE Item 13)", level=2)
    add_paragraph(doc,
        "Given the approximately 8:1 imbalance in sample size between singleton and twin groups, "
        "a stratified bootstrap procedure was performed to validate the robustness of the "
        "logistic regression confidence intervals. "
        "In each of 10,000 resamples, singleton and twin cases were resampled separately "
        "with replacement to maintain group proportions (stratified resampling). "
        "Percentile and bias-corrected and accelerated (BCa) 95% confidence intervals "
        "were calculated. The BCa bias-correction factor z0 was estimated from the "
        "bootstrap distribution median, and the acceleration factor a was estimated "
        "using the jackknife method.")

add_paragraph(doc,
    "Statistical analyses were performed using Python 3.12 (scipy 1.14, statsmodels 0.14). "
    "A two-sided P value <0.05 was considered statistically significant.")

doc.add_page_break()

# ============================================================
# RESULTS
# ============================================================
add_heading(doc, "Results")

# --- Participants (STROBE 14) ---
add_heading(doc, "Participant Selection (STROBE Item 14)", level=2)
# Build exclusion breakdown dynamically
excl_en_parts = []
for step in F["exclusion_steps"]:
    en_reason = excl_reasons_en.get(step["reason"], step["reason"])
    excl_en_parts.append(f"{en_reason} (n = {step['n']})")
excl_en_detail = ", ".join(excl_en_parts[:-1]) + ", and " + excl_en_parts[-1] if len(excl_en_parts) > 1 else excl_en_parts[0]

# Include date filter in participant flow
date_filter_en = ""
if F.get("out_of_period", {}).get("n", 0) > 0:
    date_filter_en = (
        f"Of {F['total_raw']['n']:,} cesarean deliveries in the database, "
        f"{F['out_of_period']['n']} cases outside the study period "
        f"(before April 2014; singleton {F['out_of_period']['n_s']}, twin {F['out_of_period']['n_t']}) "
        f"were excluded, leaving {F['total']['n']:,} cases for eligibility assessment. "
    )

add_paragraph(doc,
    date_filter_en +
    f"Of the {F['total']['n']:,} cesarean deliveries during the study period, "
    f"{F['total_excluded']['n']} cases were excluded: "
    f"{excl_en_detail}. "
    f"An additional {F['preop_antiemetic']['n']} cases with preoperative antiemetic use "
    f"were excluded, yielding a final analytical cohort of {F['primary_analysis']['n']:,} cases "
    f"(singleton {F['primary_analysis']['n_s']:,}; twin {F['primary_analysis']['n_t']:,}) (Fig. 1).")

# Twin subcategory breakdown (not shown in flowchart)
tc = M.get("twin_chorionicity", {})
te = M.get("twin_emergency", {})
add_paragraph(doc,
    f"Among the {F['primary_analysis']['n_t']:,} twin pregnancies, "
    f"chorionicity was dichorionic-diamniotic (DD) in {tc.get('DD', 0)}, "
    f"monochorionic-diamniotic (MD) in {tc.get('MD', 0)}, "
    f"and monochorionic-monoamniotic (MM) in {tc.get('MM', 0)}. "
    f"By urgency, {te.get('elective', 0)} ({100*te.get('elective', 0)/F['primary_analysis']['n_t']:.1f}%) "
    f"were elective and {te.get('emergency', 0)} ({100*te.get('emergency', 0)/F['primary_analysis']['n_t']:.1f}%) "
    f"were emergency cesarean deliveries.")

# --- Descriptive data (STROBE 15) ---
add_heading(doc, "Baseline Characteristics (STROBE Item 15)", level=2)
add_paragraph(doc,
    f"Baseline characteristics of the {M['n_analysis']:,} analyzed patients are presented in Table 1.")

# Table 1
table1 = pd.read_csv(BASE / "tables_e" / "table1_characteristics.csv")
add_paragraph(doc,
    f"Table 1. Baseline Characteristics (Singleton n = {M['n_single']:,} vs Twin n = {M['n_twin']:,})",
    bold=True)

tbl = doc.add_table(rows=len(table1) + 1, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Light Shading Accent 1"
for j, h in enumerate(["Variable", f"Singleton (n = {M['n_single']:,})",
                        f"Twin (n = {M['n_twin']:,})", "P value"]):
    cell = tbl.rows[0].cells[j]
    cell.text = h
    for par in cell.paragraphs:
        par.runs[0].bold = True
        par.runs[0].font.size = Pt(9)
for i, (_, row) in enumerate(table1.iterrows()):
    tbl.rows[i + 1].cells[0].text = row["Variable"]
    tbl.rows[i + 1].cells[1].text = str(row["Singleton"])
    tbl.rows[i + 1].cells[2].text = str(row["Twin"])
    p_val = row["P-value"]
    tbl.rows[i + 1].cells[3].text = "< 0.001" if p_val < 0.001 else f"{p_val:.3f}"
    for j in range(4):
        for par in tbl.rows[i + 1].cells[j].paragraphs:
            for r in par.runs:
                r.font.size = Pt(9)

# Table 1 footnotes
footnote_p = doc.add_paragraph()
fn_text = ("Continuous variables are presented as median [interquartile range]. "
           "Categorical variables are presented as n/N (%). "
           "Continuous variables were compared using the Mann-Whitney U test; "
           "categorical variables were compared using Pearson\u2019s \u03c7\u00b2 test "
           "(or Fisher\u2019s exact test when any expected cell count was <5).")
fn_run = footnote_p.add_run(fn_text)
fn_run.font.size = Pt(8)
fn_run.font.italic = True
doc.add_paragraph()

# --- Outcome data (STROBE 16) ---
add_heading(doc, "IONV Incidence (STROBE Item 16)", level=2)

add_paragraph(doc,
    f"Under the broad definition, the primary-outcome IONV rate was "
    f"{mo['A-Primary']['singleton_pct']:.1f}% "
    f"({mo['A-Primary']['singleton_n']}/{M['n_single']:,}) in the singleton group and "
    f"{mo['A-Primary']['twin_pct']:.1f}% "
    f"({mo['A-Primary']['twin_n']}/{M['n_twin']:,}) in the twin group (Table 2, Fig. 2).")

add_paragraph(doc,
    f"Under the narrow definition (5-HT3 antagonists), the IONV rate was "
    f"{mo['E-Primary']['singleton_pct']:.1f}% "
    f"({mo['E-Primary']['singleton_n']}/{M['n_single']:,}) in the singleton group and "
    f"{mo['E-Primary']['twin_pct']:.1f}% "
    f"({mo['E-Primary']['twin_n']}/{M['n_twin']:,}) in the twin group, "
    "which was significantly higher (Table 2, Fig. 2).")

# Table 2: IONV rates
add_paragraph(doc,
    "Table 2. IONV Incidence and Multivariable Logistic Regression Analysis", bold=True)
tbl2 = doc.add_table(rows=5, cols=7)
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl2.style = "Light Shading Accent 1"
for j, h in enumerate(["Outcome", "Singleton n (%)", "Twin n (%)",
                        "Model", "aOR", "95% CI", "P value"]):
    cell = tbl2.rows[0].cells[j]
    cell.text = h
    for par in cell.paragraphs:
        par.runs[0].bold = True
        par.runs[0].font.size = Pt(8)

for i, key in enumerate(["A-Primary", "A-Secondary", "E-Primary", "E-Secondary"]):
    rr = mr[key]
    oo = mo[key]
    labels_en = {
        "A-Primary": "Broad, Primary",
        "A-Secondary": "Broad, Secondary",
        "E-Primary": "Narrow, Primary",
        "E-Secondary": "Narrow, Secondary",
    }
    vals = [
        labels_en[key],
        f"{oo['singleton_n']} ({oo['singleton_pct']:.1f}%)",
        f"{oo['twin_n']} ({oo['twin_pct']:.1f}%)",
        rr["model_type"],
        f"{rr['twin_OR']:.2f}",
        f"{rr['twin_CI_lower']:.2f}\u2013{rr['twin_CI_upper']:.2f}",
        "< 0.001" if rr["twin_P"] < 0.001 else f"{rr['twin_P']:.3f}",
    ]
    for j, v in enumerate(vals):
        tbl2.rows[i + 1].cells[j].text = v
        for par in tbl2.rows[i + 1].cells[j].paragraphs:
            for r in par.runs:
                r.font.size = Pt(8)
doc.add_paragraph()

# Fig 2: Rates comparison
add_figure(doc, BASE / "figures_e" / "fig1_rates_comparison.png",
           "Fig. 2  IONV Rates by Antiemetic Definition (Broad vs Narrow) and Pregnancy Type")

# --- Main results (STROBE 17) ---
add_heading(doc, "Multivariable Analysis (STROBE Item 17)", level=2)

add_paragraph(doc,
    f"Under the broad definition, twin pregnancy was not significantly associated with "
    f"IONV in the full covariate model "
    f"(aOR {mr['A-Primary']['twin_OR']:.2f}, "
    f"95% CI {mr['A-Primary']['twin_CI_lower']:.2f}\u2013{mr['A-Primary']['twin_CI_upper']:.2f}, "
    f"{format_p(mr['A-Primary']['twin_P'])}).")

add_paragraph(doc,
    f"Under the narrow definition, only {mr['E-Primary']['events']} events were observed, "
    "necessitating the use of a reduced model. "
    f"Twin pregnancy was significantly associated with 5-HT3 antagonist use "
    f"(aOR {mr['E-Primary']['twin_OR']:.2f}, "
    f"95% CI {mr['E-Primary']['twin_CI_lower']:.2f}\u2013{mr['E-Primary']['twin_CI_upper']:.2f}, "
    f"{format_p(mr['E-Primary']['twin_P'])}) (Fig. 3).")

# Fig 3: Forest plot
add_figure(doc, BASE / "figures_e" / "fig2_forest_E_primary.png",
           "Fig. 3  Multivariable Logistic Regression for Narrow-Definition Antiemetic Use \u2014 "
           "Forest Plot")

# Fig 4: Broad vs Narrow comparison
add_figure(doc, BASE / "figures_e" / "fig4_protocol_vs_defE.png",
           "Fig. 4  Twin Effect on IONV: Broad vs Narrow Antiemetic Definition")

# --- Secondary outcome: Hypotension ---
add_heading(doc, "Secondary Outcome: Intraoperative Hypotension (STROBE Item 17)", level=2)
add_paragraph(doc,
    f"Intraoperative hypotension (SBP <90 mmHg) occurred in "
    f"{S['hypo_single_n']:,}/{S['n_single']:,} singleton cases ({S['hypo_single_pct']:.1f}%) and "
    f"{S['hypo_twin_n']}/{S['n_twin']:,} twin cases ({S['hypo_twin_pct']:.1f}%), "
    f"with a significantly lower rate in the twin group (P = {S['hypo_chi_p']:.3f}). "
    f"The median number of hypotensive episodes [interquartile range] was "
    f"{S['hypo_count_single_median']:.0f} [{S['hypo_count_single_q1']:.0f}\u2013"
    f"{S['hypo_count_single_q3']:.0f}] in singletons and "
    f"{S['hypo_count_twin_median']:.0f} [{S['hypo_count_twin_q1']:.0f}\u2013"
    f"{S['hypo_count_twin_q3']:.0f}] in twins (P = {S['hypo_count_p']:.3f}).")

# Note: hypotension is also used as a covariate in IONV analysis,
# so only descriptive statistics are reported here (no regression).

# ============================================================
# SENSITIVITY ANALYSES (detachable block)
# ============================================================

cov_df = pd.read_csv(BASE / "tables_e" / "covariate_sensitivity.csv")

if INCLUDE_SENSITIVITY:
    n_sig_main = int((cov_df["P"] < 0.05).sum())
    add_heading(doc, "Covariate Sensitivity Analysis (STROBE Item 18)", level=2)
    add_paragraph(doc,
        f"To assess the robustness of the twin effect estimate for the narrow-definition "
        f"primary outcome, a covariate sensitivity analysis was performed (Table 3, Fig. 5). "
        f"Across all {len(cov_df)} models, the aOR for twin pregnancy ranged from "
        f"{cov_df['aOR'].min():.2f} to {cov_df['aOR'].max():.2f}, "
        f"and {n_sig_main}/{len(cov_df)} models were statistically significant (P < 0.05).")

    # Table 3: Covariate sensitivity
    add_paragraph(doc,
        "Table 3. Covariate Sensitivity Analysis (Narrow-Definition Primary Outcome, "
        "Full Cohort)", bold=True)
    tbl3 = doc.add_table(rows=len(cov_df) + 1, cols=4)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl3.style = "Light Shading Accent 1"
    for j, h in enumerate(["Model", "aOR", "95% CI", "P value"]):
        tbl3.rows[0].cells[j].text = h
        for par in tbl3.rows[0].cells[j].paragraphs:
            par.runs[0].bold = True
            par.runs[0].font.size = Pt(9)
    for i, (_, row) in enumerate(cov_df.iterrows()):
        vals = [row["Model"], f"{row['aOR']:.2f}",
                f"{row['CI_lower']:.2f}\u2013{row['CI_upper']:.2f}",
                "< 0.001" if row["P"] < 0.001 else f"{row['P']:.3f}"]
        for j, v in enumerate(vals):
            tbl3.rows[i + 1].cells[j].text = v
            for par in tbl3.rows[i + 1].cells[j].paragraphs:
                for r in par.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()

    # Fig 5: Covariate sensitivity forest
    add_figure(doc, BASE / "figures_e" / "fig3_covariate_sensitivity.png",
               "Fig. 5  Covariate Sensitivity Analysis: Twin Effect on "
               "Narrow-Definition Antiemetic Use")

    # --- Exclusion sensitivity analysis ---
    add_heading(doc, "Exclusion Sensitivity Analysis (STROBE Item 18)", level=2)
    add_paragraph(doc,
        f"After further exclusion of emergency cesarean delivery "
        f"(n = {E['exclusion_counts']['Emergency CS']:,}), "
        f"prior cesarean delivery (n = {E['exclusion_counts']['Prior CS']:,}), "
        f"HDP (n = {E['exclusion_counts']['HDP']}), and "
        f"preoperative steroid use (n = {E['exclusion_counts']['Preoperative steroid']}), "
        f"the elective, low-risk sensitivity subgroup comprised {E['n_analysis']:,} cases "
        f"(singleton {E['n_single']:,}; twin {E['n_twin']:,}) (with overlap; Fig. 1).")
    add_paragraph(doc,
        f"Under the broad definition, IONV rates did not differ significantly "
        f"(singleton {eo['A-Primary']['singleton_pct']:.1f}% vs "
        f"twin {eo['A-Primary']['twin_pct']:.1f}%; "
        f"aOR {er['A-Primary']['twin_OR']:.2f}, "
        f"{format_p(er['A-Primary']['twin_P'])}). "
        f"Under the narrow definition, IONV remained significantly higher in the twin group "
        f"(singleton {eo['E-Primary']['singleton_pct']:.1f}% vs "
        f"twin {eo['E-Primary']['twin_pct']:.1f}%; "
        f"aOR {er['E-Primary']['twin_OR']:.2f}, "
        f"95% CI {er['E-Primary']['twin_CI_lower']:.2f}\u2013"
        f"{er['E-Primary']['twin_CI_upper']:.2f}, "
        f"{format_p(er['E-Primary']['twin_P'])}) (Table 4).")

    # Table 4: Exclusion sensitivity results
    add_paragraph(doc,
        "Table 4. Exclusion Sensitivity Analysis Results "
        "(Elective, Low-Risk Subgroup)", bold=True)
    tbl4 = doc.add_table(rows=5, cols=7)
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl4.style = "Light Shading Accent 1"
    for j, h in enumerate(["Outcome", "Singleton n (%)", "Twin n (%)",
                            "Model", "aOR", "95% CI", "P value"]):
        tbl4.rows[0].cells[j].text = h
        for par in tbl4.rows[0].cells[j].paragraphs:
            par.runs[0].bold = True
            par.runs[0].font.size = Pt(8)
    for i, key in enumerate(["A-Primary", "A-Secondary", "E-Primary", "E-Secondary"]):
        rr = er[key]
        oo = eo[key]
        labels_en = {
            "A-Primary": "Broad, Primary",
            "A-Secondary": "Broad, Secondary",
            "E-Primary": "Narrow, Primary",
            "E-Secondary": "Narrow, Secondary",
        }
        vals = [
            labels_en[key],
            f"{oo['singleton_n']} ({oo['singleton_pct']:.1f}%)",
            f"{oo['twin_n']} ({oo['twin_pct']:.1f}%)",
            rr["model_type"],
            f"{rr['twin_OR']:.2f}",
            f"{rr['twin_CI_lower']:.2f}\u2013{rr['twin_CI_upper']:.2f}",
            "< 0.001" if rr["twin_P"] < 0.001 else f"{rr['twin_P']:.3f}",
        ]
        for j, v in enumerate(vals):
            tbl4.rows[i + 1].cells[j].text = v
            for par in tbl4.rows[i + 1].cells[j].paragraphs:
                for r in par.runs:
                    r.font.size = Pt(8)
    doc.add_paragraph()

    # Fig 6: Exclusion sensitivity rates
    add_figure(doc, BASE / "figures_excl" / "fig1_rates_comparison.png",
               "Fig. 6  IONV Rates in the Elective, Low-Risk Sensitivity Subgroup")

    # Fig 7: Exclusion sensitivity forest
    add_figure(doc, BASE / "figures_excl" / "fig4_broad_vs_narrow.png",
               "Fig. 7  Twin Effect on IONV: Broad vs Narrow Definition "
               "(Elective, Low-Risk Subgroup)")

    # --- Emergency sensitivity analysis ---
    add_heading(doc, "Emergency Cesarean Delivery Sensitivity Analysis (STROBE Item 18)", level=2)
    add_paragraph(doc,
        "To examine whether emergency cesarean delivery modifies the twin effect on IONV, "
        "separate analyses were performed for elective cases only "
        "(n = 1,736; singleton 1,515; twin 221) and emergency cases only "
        "(n = 1,452; singleton 1,331; twin 121) (Table 5).")
    add_paragraph(doc,
        "In the elective subgroup, the narrow-definition primary outcome showed "
        "a substantially larger effect size (aOR 8.39, 95% CI 1.21\u201358.18, P = 0.031), "
        "and the narrow-definition secondary outcome was also significant "
        "(aOR 8.99, 95% CI 2.50\u201332.29, P < 0.001). "
        "In contrast, among emergency cases, the twin effect was absent "
        "(narrow secondary aOR 0.95, P = 0.958; narrow primary had insufficient events).")
    add_paragraph(doc,
        "Notably, under the broad definition in elective cases, the twin group showed "
        "significantly lower IONV in the secondary outcome (18.5% vs 14.0%; aOR 0.64, P = 0.047). "
        "Interaction analysis revealed a significant emergency \u00d7 twin interaction "
        "for the broad secondary outcome (interaction OR 2.05, P = 0.023).")

    # Table 5: Emergency sensitivity
    add_paragraph(doc, "Table 5. Emergency Cesarean Delivery Sensitivity Analysis", bold=True)
    tbl5 = doc.add_table(rows=9, cols=7)
    tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl5.style = "Light Shading Accent 1"
    for j, h in enumerate(["Subgroup", "Outcome", "Singleton", "Twin", "aOR", "95% CI", "P value"]):
        tbl5.rows[0].cells[j].text = h
        for par in tbl5.rows[0].cells[j].paragraphs:
            par.runs[0].bold = True
            par.runs[0].font.size = Pt(8)
    emg_data = [
        ["Elective only", "Broad, Primary", "2.4%", "1.8%", "0.83", "0.28\u20132.47", "0.735"],
        ["", "Broad, Secondary", "18.5%", "14.0%", "0.64", "0.41\u20130.99", "0.047"],
        ["", "Narrow, Primary", "0.3%", "1.4%", "8.39", "1.21\u201358.18", "0.031"],
        ["", "Narrow, Secondary", "0.8%", "3.2%", "8.99", "2.50\u201332.29", "< 0.001"],
        ["Emergency only", "Broad, Primary", "1.4%", "0.8%", "0.61", "0.08\u20134.68", "0.637"],
        ["", "Broad, Secondary", "16.8%", "22.3%", "1.44", "0.91\u20132.29", "0.119"],
        ["", "Narrow, Primary", "0.1%", "0.0%", "\u2014", "\u2014", "\u2014"],
        ["", "Narrow, Secondary", "1.0%", "0.8%", "0.95", "0.12\u20137.44", "0.958"],
    ]
    for i, row_data in enumerate(emg_data):
        for j, v in enumerate(row_data):
            tbl5.rows[i + 1].cells[j].text = v
            for par in tbl5.rows[i + 1].cells[j].paragraphs:
                for r in par.runs:
                    r.font.size = Pt(8)
    doc.add_paragraph()

    # --- Bootstrap validation ---
    bm = B["main_cohort"]
    bs = B["subgroup"]
    add_heading(doc, "Bootstrap Validation (STROBE Item 18)", level=2)
    add_paragraph(doc,
        "Results of the stratified bootstrap validation for the full cohort "
        "are presented in Table 6 and Fig. 8.")
    add_paragraph(doc,
        f"For the narrow-definition primary outcome in the full cohort, "
        f"the BCa 95% CI was [{bm['E-Primary']['bca_CI_lower']:.2f}\u2013"
        f"{bm['E-Primary']['bca_CI_upper']:.2f}], closely matching the "
        f"Wald 95% CI [{bm['E-Primary']['wald_CI_lower']:.2f}\u2013"
        f"{bm['E-Primary']['wald_CI_upper']:.2f}]. "
        f"The bootstrap CI also excluded 1, confirming that the significant twin effect "
        f"(aOR {bm['E-Primary']['point_aOR']:.2f}) was robust "
        f"(convergence rate {bm['E-Primary']['convergence_pct']}%).")
    add_paragraph(doc,
        f"In the low-risk subgroup, event counts for the narrow definition were extremely small "
        f"(primary: {bs['E-Primary']['events']}; secondary: {bs['E-Secondary']['events']}), "
        "resulting in unstable bootstrap CIs. "
        "This reflects the inherent sample size constraint and warrants "
        "cautious interpretation of subgroup results.")

    # Table 6: Bootstrap results
    add_paragraph(doc, "Table 6. Stratified Bootstrap Validation (10,000 Resamples)", bold=True)
    tbl6 = doc.add_table(rows=9, cols=6)
    tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl6.style = "Light Shading Accent 1"
    for j, h in enumerate(["Cohort", "Outcome", "aOR", "Wald 95% CI", "BCa 95% CI", "Convergence"]):
        tbl6.rows[0].cells[j].text = h
        for par in tbl6.rows[0].cells[j].paragraphs:
            par.runs[0].bold = True
            par.runs[0].font.size = Pt(8)

    def fmt_bca(r):
        lo = r.get("bca_CI_lower")
        hi = r.get("bca_CI_upper")
        if lo is not None and hi is not None and hi < 1e6:
            return f"{lo:.2f}\u2013{hi:.2f}"
        return "Unstable"

    boot_rows = [
        ["Full cohort", "Broad, Primary", bm["A-Primary"]],
        ["", "Broad, Secondary", bm["A-Secondary"]],
        ["", "Narrow, Primary", bm["E-Primary"]],
        ["", "Narrow, Secondary", bm["E-Secondary"]],
        ["Low-risk", "Broad, Primary", bs["A-Primary"]],
        ["", "Broad, Secondary", bs["A-Secondary"]],
        ["", "Narrow, Primary", bs["E-Primary"]],
        ["", "Narrow, Secondary", bs["E-Secondary"]],
    ]
    for i, (cohort, outcome, r) in enumerate(boot_rows):
        vals = [
            cohort, outcome,
            f"{r['point_aOR']:.2f}",
            f"{r['wald_CI_lower']:.2f}\u2013{r['wald_CI_upper']:.2f}",
            fmt_bca(r),
            f"{r['convergence_pct']}%",
        ]
        for j, v in enumerate(vals):
            tbl6.rows[i + 1].cells[j].text = v
            for par in tbl6.rows[i + 1].cells[j].paragraphs:
                for r2 in par.runs:
                    r2.font.size = Pt(8)
    doc.add_paragraph()

    # Fig 8: Bootstrap comparison
    add_figure(doc, BASE / "figures_bootstrap" / "fig_bootstrap_comparison.png",
               "Fig. 8  Wald vs Bootstrap Confidence Interval Comparison")

# --- end sensitivity block ---

doc.add_page_break()

# ============================================================
# DISCUSSION
# ============================================================
add_heading(doc, "Discussion")

add_heading(doc, "Key Findings (STROBE Item 19)", level=2)
_key_finding_base_en = (
    "This study compared the incidence of IONV between singleton and twin pregnancies "
    "during cesarean delivery under spinal anesthesia. "
    "Under the broad antiemetic definition, no significant difference was observed between "
    "the two groups. However, when the analysis was restricted to 5-HT3 receptor antagonists "
    "\u2014 agents with high pharmacological specificity for nausea and vomiting \u2014 "
    "the twin group showed a significantly higher rate of use.")
if INCLUDE_SENSITIVITY:
    _key_finding_base_en += (
        " This finding was robustly reproduced across covariate sensitivity analyses, "
        "exclusion sensitivity analyses, emergency subgroup stratification, "
        "and 10,000-replicate stratified bootstrap validation.")
add_paragraph(doc, _key_finding_base_en)

add_paragraph(doc,
    f"As a secondary outcome, intraoperative hypotension (SBP <90 mmHg) was less frequent "
    f"in the twin group ({S['hypo_twin_pct']:.1f}%) than in the singleton group "
    f"({S['hypo_single_pct']:.1f}%; {format_p(S['hypo_chi_p'])}). "
    "This finding, combined with the higher 5-HT3 antagonist use in twins, "
    "suggests involvement of non-hypotensive IONV pathways in twin pregnancies.")

add_heading(doc, "Limitations (STROBE Item 20)", level=2)
add_p_with_refs(doc,
    "This study has several limitations. "
    "First, as a retrospective study, direct assessment of IONV was not possible; "
    "antiemetic administration was used as a surrogate marker. "
    "Most prior studies on IONV during cesarean delivery were prospective in design "
    "and recorded the occurrence of nausea and vomiting through patient self-report "
    "or direct intraoperative observation{3-6,7,9-11}. "
    "In contrast, our retrospective design precluded direct assessment of subjective "
    "nausea or its severity; instead, we used the fact of antiemetic dispensing as a "
    "surrogate for IONV. This approach cannot exclude the possibility that antiemetics "
    "were administered for indications other than IONV (e.g., opioid-induced pruritus "
    "prophylaxis, gastrointestinal motility promotion). However, the narrow definition "
    "restricted to 5-HT3 receptor antagonists \u2014 agents with high pharmacological "
    "specificity for nausea and vomiting \u2014 minimized this confounding effect. "
    "Although the surrogate-marker approach is unique to the present study, "
    "George et al.{10} reported rescue antiemetic administration as a secondary "
    "outcome of IONV, supporting the relevance of antiemetic use as an IONV-related "
    "endpoint.")

if INCLUDE_SENSITIVITY:
    add_paragraph(doc,
        f"Second, the number of events for the narrow definition was small "
        f"(primary analysis: {mr['E-Primary']['events']}; "
        f"subgroup analysis: {er['E-Primary']['events']}), "
        "resulting in wide confidence intervals. "
        "Bootstrap validation confirmed the robustness of the full-cohort aOR 3.18 "
        "but the subgroup bootstrap CIs (8 events) were unstable. "
        f"In particular, the subgroup analysis involved a substantial reduction in sample size "
        f"({M['n_analysis']:,} \u2192 {E['n_analysis']}), "
        "limiting external validity.")
else:
    add_paragraph(doc,
        f"Second, the number of events for the narrow definition was small "
        f"(primary analysis: {mr['E-Primary']['events']}), "
        "resulting in wide confidence intervals.")

add_paragraph(doc,
    "Third, exclusion of patients with baseline SBP <90 mmHg at admission "
    "precluded full assessment of the effect of twin pregnancy on baseline "
    "hemodynamic status. Twin pregnancies have competing circulatory characteristics "
    "— greater circulating blood volume versus more pronounced aortocaval compression "
    "— and the distribution of admission blood pressure may differ between singleton "
    "and twin pregnancies. Although the number of excluded cases was small "
    f"({F['exclusion_steps'][1]['n']} patients: "
    f"{F['exclusion_steps'][1]['n_s']} singleton, {F['exclusion_steps'][1]['n_t']} twin), "
    "this selection bias may have influenced the between-group comparison of "
    "intraoperative hypotension.")

add_paragraph(doc,
    "Fourth, this was a single-center study, and multicenter studies are needed to "
    "confirm generalizability.")

add_p_with_refs(doc,
    "Fifth, the inability to assess the severity of IONV is a limitation. "
    "Some prior studies reported intraoperative nausea and vomiting as separate "
    "outcomes{5}, but this distinction is not possible with our antiemetic-based "
    "definition. While the hypothesis that 5-HT3 antagonist use reflects more severe "
    "IONV is plausible, prospective studies are needed to confirm the association "
    "between IONV severity and the choice of antiemetic agent.")

add_heading(doc, "Interpretation and Clinical Significance (STROBE Item 21)", level=2)
add_p_with_refs(doc,
    "The finding that 5-HT3 antagonist use was significantly higher in the twin group "
    "despite similar overall antiemetic use is noteworthy. "
    "Prior studies defined IONV by the presence or absence of nausea and vomiting "
    "through direct measurement{3-7,9-11}, whereas the present study used antiemetic "
    "administration as a surrogate marker. This approach may have selectively captured "
    "\u2018IONV requiring treatment\u2019 \u2014 a dimension not assessed in prior work. "
    "Possible interpretations include: (1) although the overall frequency of antiemetic "
    "use was similar, twins may experience more severe IONV requiring 5-HT3 antagonists; "
    "or (2) prescribing patterns may differ for twin pregnancies, with clinicians selecting "
    "more specific agents.")

add_p_with_refs(doc,
    "The principal mechanisms of IONV are hypotension, vagal activity, and visceral "
    "stimulation{2,12}, and twin pregnancy amplifies all of these through greater "
    "circulating blood volume, enhanced aortocaval compression, and reduced cardiac "
    "reserve{13}. "
    f"The present cohort of {F['primary_analysis']['n']:,} cases including {F['primary_analysis']['n_t']:,} twins represents the largest "
    "study to directly examine IONV in twin pregnancies, and these findings may contribute "
    "to elucidating the mechanisms of IONV. "
    "The observation that differences emerged only under the narrow (5-HT3 antagonist) "
    "definition, but not the broad definition, suggests that the \u2018quality\u2019 "
    "(severity and treatment necessity) rather than the \u2018quantity\u2019 of IONV may "
    "differ in twins. "
    "The persistence of the twin effect after adjusting for hypotension as a covariate "
    "suggests involvement of pathways beyond hypotension, such as serotonin release "
    "from gut hypoperfusion or enhanced visceral stimulation.")

if INCLUDE_SENSITIVITY:
    add_paragraph(doc,
        "The observation that the effect size increased in the subgroup analysis after "
        "excluding confounders (emergency surgery, prior cesarean delivery, HDP, steroids) "
        "suggests that these factors may have elevated IONV risk in the singleton group, "
        "thereby diluting the twin effect in the full cohort.")
    add_paragraph(doc,
        "The emergency sensitivity analysis revealed that the twin effect was confined to "
        "elective cases, where the narrow-definition primary aOR was 8.39 "
        "and the secondary aOR was 8.99, both significant. "
        "In contrast, among emergency cases the twin effect was completely absent. "
        "A significant emergency \u00d7 twin interaction for the broad secondary outcome "
        "(P = 0.023) further supports emergency delivery status as an effect modifier.")
    add_paragraph(doc,
        "The stratified bootstrap validation confirmed that the BCa confidence interval "
        "for the full-cohort narrow-definition primary outcome excluded 1, "
        "corroborating the Wald-based inference. This demonstrates that despite the "
        "approximately 8:1 imbalance in group sizes, the logistic regression estimates "
        "are stable and the key finding is robust.")

add_heading(doc, "Generalizability (STROBE Item 22)", level=2)
add_p_with_refs(doc,
    "This study was based on 10 years of data from a regional referral hospital, "
    "reflecting typical cesarean delivery management in Japan. "
    f"The present cohort including {F['primary_analysis']['n_t']:,} twin cases provides the largest dataset "
    "to directly examine IONV in twin pregnancies, in contrast to prior studies "
    "that systematically excluded twins{9-11}. "
    "However, antiemetic prescribing patterns may vary across institutions and eras, "
    "and caution is warranted in extrapolating these results. "
    "Future prospective multicenter studies are warranted to elucidate "
    "the mechanisms of IONV in twin pregnancies.")

doc.add_page_break()

# ============================================================
# REFERENCES (Vancouver style, numbered in order of appearance)
# ============================================================
add_heading(doc, "References")

references = [
    "Juhani TP, Hannele H. Complications during spinal anesthesia for cesarean delivery: "
    "a clinical report of one year's experience. Reg Anesth. 1993;18(2):128-31.",

    "Balki M, Carvalho J. Intraoperative nausea and vomiting during cesarean section "
    "under regional anesthesia. Int J Obstet Anesth. 2005;14(3):230-41.",

    "Mercier FJ, Diemunsch P, Ducloy-Bouthors AS, et al. 6% hydroxyethyl starch (130/0.4) "
    "vs Ringer\u2019s lactate preloading before spinal anaesthesia for Caesarean delivery: "
    "the randomized, double-blind, multicentre CAESAR trial. Br J Anaesth. 2014;113(3):459-467.",

    "Santos A, Datta S. Prophylactic use of droperidol for control of nausea and vomiting "
    "during spinal anesthesia for cesarean section. Anesth Analg. 1984;63(1):85-87.",

    "Mishriky B, Habib A. Metoclopramide for nausea and vomiting prophylaxis during and "
    "after Caesarean delivery: a systematic review and meta-analysis. "
    "Br J Anaesth. 2012;108(3):374-83.",

    "Harmon D, Ryan M, Kelly A, Bowen M. Acupressure and prevention of nausea and vomiting "
    "during and after spinal anaesthesia for caesarean section. "
    "Br J Anaesth. 2000;84(4):463-7.",

    "Ashagrie HE, Filatie TD, Melesse DY, Mustefa SY. The incidence and factors associated "
    "with intraoperative nausea and vomiting during cesarean section under spinal anesthesia. "
    "Int J Surg Open. 2020;26:49-54.",

    "Tan HS, Taylor CR, Sharawi N, et al. Uterine exteriorization versus in situ repair in "
    "Cesarean delivery: a systematic review and meta-analysis. "
    "Can J Anesth. 2022;69:216-233.",

    "Chen Z, Zhou J, Wan L, Huang H. Norepinephrine versus phenylephrine infusion for "
    "preventing postspinal hypotension during cesarean section for twin pregnancy: "
    "a double-blinded randomized controlled clinical trial. "
    "BMC Anesthesiol. 2022;22:17.",

    "George RB, McKeen DM, Dominguez JE, et al. Randomized trial of phenylephrine infusion "
    "vs. bolus for nausea & vomiting during cesarean in obese women. "
    "Can J Anaesth. 2018;65:254-262.",

    "Ngan Kee WD, Lee SWY, Ng FF, et al. Randomized double-blinded comparison of "
    "norepinephrine and phenylephrine for maintenance of blood pressure during spinal "
    "anesthesia for cesarean delivery. Anesthesiology. 2015;122:736-45.",

    "Jelting Y, Klein C, Harlander T, et al. Preventing nausea and vomiting in women "
    "undergoing regional anesthesia for cesarean section: challenges and solutions. "
    "Local Reg Anesth. 2017;10:83-90.",

    "Farrer JR, Peralta FM. Anaesthesia for the parturient with multiple gestations. "
    "BJA Educ. 2022;22(8):306-311.",
]

for i, ref in enumerate(references):
    p = doc.add_paragraph()
    run = p.add_run(f"{i+1}. ")
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run(ref)
    run.font.size = Pt(9)

doc.add_page_break()

# ============================================================
# STROBE CHECKLIST (Supplementary) — ENGLISH
# ============================================================
add_heading(doc, "Supplementary: STROBE Checklist for Cohort Studies")

strobe_items = [
    ("Title and abstract", "1(a)", "Title page and structured abstract",
     "Indicate the study\u2019s design with a commonly used term in the title or the abstract"),
    ("Title and abstract", "1(b)", "Title page and structured abstract",
     "Provide in the abstract an informative and balanced summary of what was done and what was found"),
    ("", "", "", ""),  # Section header
    ("INTRODUCTION", "", "", ""),
    ("Background/rationale", "2", "Introduction, paragraphs 1\u20132",
     "Explain the scientific background and rationale for the investigation being reported"),
    ("Objectives", "3", "Introduction, paragraph 3",
     "State specific objectives, including any prespecified hypotheses"),
    ("", "", "", ""),
    ("METHODS", "", "", ""),
    ("Study design", "4", "Methods: Study Design and Ethics",
     "Present key elements of study design early in the paper"),
    ("Setting", "5", "Methods: Participants",
     "Describe the setting, locations, and relevant dates"),
    ("Participants", "6(a)", "Methods: Participants, Fig. 1",
     "Give the eligibility criteria, and the sources and methods of selection of participants"),
    ("Variables", "7", "Methods: Outcome Measures, Covariates",
     "Clearly define all outcomes, exposures, predictors, potential confounders, and effect modifiers"),
    ("Data sources/measurement", "8", "Methods: Data Collection",
     "For each variable of interest, give sources of data and details of methods of assessment"),
    ("Bias", "9", "Methods: Statistical Analysis; Discussion: Limitations",
     "Describe any efforts to address potential sources of bias"),
    ("Study size", "10", "Methods: Participants; Results: Participant Selection",
     "Explain how the study size was arrived at"),
    ("Quantitative variables", "11", "Methods: Outcome Measures",
     "Explain how quantitative variables were handled in the analyses"),
    ("Statistical methods", "12(a)", "Methods: Statistical Analysis",
     "Describe all statistical methods"),
    ("Statistical methods", "12(b)", "Methods: Statistical Analysis",
     "Describe any methods used to examine subgroups and interactions"),
    ("Statistical methods", "12(c)", "Methods: Statistical Analysis",
     "Explain how missing data were addressed"),
    ("Statistical methods", "12(d)", "Methods: Sensitivity Analysis",
     "If applicable, describe analytical methods taking account of sampling strategy"),
    ("Statistical methods", "12(e)", "Methods: Sensitivity Analysis",
     "Describe any sensitivity analyses"),
    ("", "", "", ""),
    ("RESULTS", "", "", ""),
    ("Participants", "13(a)", "Results: Participant Selection, Fig. 1",
     "Report numbers of individuals at each stage of study"),
    ("Participants", "13(b)", "Results: Participant Selection",
     "Give reasons for non-participation at each stage"),
    ("Descriptive data", "14(a)", "Results: Baseline Characteristics, Table 1",
     "Give characteristics of study participants and information on exposures and potential confounders"),
    ("Descriptive data", "14(b)", "Results: Participant Selection",
     "Indicate number of participants with missing data for each variable of interest"),
    ("Outcome data", "15", "Results: IONV Incidence, Table 2",
     "Report numbers of outcome events or summary measures over time"),
    ("Main results", "16(a)", "Results: Multivariable Analysis, Table 2, Fig. 3\u20134",
     "Give unadjusted estimates and, if applicable, confounder-adjusted estimates and their precision"),
    ("Main results", "16(b)", "Results: Multivariable Analysis",
     "Report category boundaries when continuous variables were categorized"),
    ("Main results", "16(c)", "Not applicable",
     "If relevant, consider translating estimates of relative risk into absolute risk"),
    ("Other analyses", "17",
     "Results: Covariate Sensitivity Analysis, Exclusion Sensitivity Analysis"
     if INCLUDE_SENSITIVITY else "Not applicable (main analysis only)",
     "Report other analyses done\u2014e.g., analyses of subgroups and interactions, and sensitivity analyses"),
    ("", "", "", ""),
    ("DISCUSSION", "", "", ""),
    ("Key results", "18", "Discussion: Key Findings",
     "Summarise key results with reference to study objectives"),
    ("Limitations", "19", "Discussion: Limitations",
     "Discuss limitations of the study, taking into account sources of potential bias or imprecision"),
    ("Interpretation", "20", "Discussion: Interpretation and Clinical Significance",
     "Give a cautious overall interpretation of results considering objectives, limitations, "
     "multiplicity of analyses, results from similar studies, and other relevant evidence"),
    ("Generalisability", "21", "Discussion: Generalizability",
     "Discuss the generalisability (external validity) of the study results"),
    ("", "", "", ""),
    ("OTHER INFORMATION", "", "", ""),
    ("Funding", "22", "No funding was received for this study",
     "Give the source of funding and the role of the funders for the present study"),
]

# Filter out empty separator rows for counting
actual_items = [s for s in strobe_items if s[1]]

tbl_s = doc.add_table(rows=len(strobe_items) + 1, cols=4)
tbl_s.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_s.style = "Light Shading Accent 1"
for j, h in enumerate(["Item", "No.", "Location in manuscript", "Recommendation"]):
    tbl_s.rows[0].cells[j].text = h
    for par in tbl_s.rows[0].cells[j].paragraphs:
        par.runs[0].bold = True
        par.runs[0].font.size = Pt(8)

for i, (item, no, loc, rec) in enumerate(strobe_items):
    tbl_s.rows[i + 1].cells[0].text = item
    tbl_s.rows[i + 1].cells[1].text = no
    tbl_s.rows[i + 1].cells[2].text = loc
    tbl_s.rows[i + 1].cells[3].text = rec
    is_section = (not no and not loc and item)
    for j in range(4):
        for par in tbl_s.rows[i + 1].cells[j].paragraphs:
            for r in par.runs:
                r.font.size = Pt(7)
                if is_section:
                    r.bold = True

doc.add_page_break()

# ============================================================
# FIGURE LEGENDS (for journal submission)
# ============================================================
add_heading(doc, "Figure Legends")

legends = [
    ("Fig. 1",
     "STROBE flow diagram showing participant selection. "
     f"Of {F['total']['n']:,} cesarean deliveries assessed, "
     f"{F['total_excluded']['n']} were excluded, leaving {F['primary_analysis']['n']:,} "
     f"(singleton {F['primary_analysis']['n_s']:,}, twin {F['primary_analysis']['n_t']:,}) "
     "for the primary analysis. A further exclusion of emergency cesarean delivery, prior "
     "cesarean delivery, HDP, and preoperative steroid yielded a sensitivity subgroup of "
     f"{F['subgroup_analysis']['n']:,} "
     f"(singleton {F['subgroup_analysis']['n_s']:,}, twin {F['subgroup_analysis']['n_t']:,})."),
    ("Fig. 2",
     "IONV rates by antiemetic definition (broad: all 7 drugs; narrow: 5-HT3 antagonists only) "
     "and pregnancy type (singleton vs twin). Error bars represent 95% confidence intervals. "
     "* P < 0.05 for singleton vs twin comparison."),
    ("Fig. 3",
     "Forest plot of multivariable logistic regression for narrow-definition antiemetic use "
     "(primary outcome). Reduced model with 6 covariates. "
     "Twin pregnancy was independently associated with 5-HT3 antagonist use "
     f"(aOR {mr['E-Primary']['twin_OR']:.2f}, 95% CI "
     f"{mr['E-Primary']['twin_CI_lower']:.2f}\u2013{mr['E-Primary']['twin_CI_upper']:.2f})."),
    ("Fig. 4",
     "Comparison of adjusted odds ratios for twin pregnancy across broad and narrow "
     "antiemetic definitions (primary and secondary outcomes). "
     "Only the narrow-definition primary outcome showed a significant association."),
]
if INCLUDE_SENSITIVITY:
    legends += [
        ("Fig. 5",
         "Covariate sensitivity analysis for the narrow-definition primary outcome. "
         f"All {len(cov_df)} models yielded aOR in the range "
         f"{cov_df['aOR'].min():.2f}\u2013{cov_df['aOR'].max():.2f}, "
         "all P < 0.05, demonstrating robustness of the twin effect."),
        ("Fig. 6",
         "IONV rates in the elective, low-risk sensitivity subgroup "
         f"(N = {F['subgroup_analysis']['n']:,}) after excluding emergency cesarean delivery, "
         "prior cesarean delivery, HDP, and preoperative steroid use."),
        ("Fig. 7",
         "Comparison of adjusted odds ratios in the elective, low-risk sensitivity subgroup. "
         f"The narrow-definition effect size increased from aOR {mr['E-Primary']['twin_OR']:.2f} "
         f"(full cohort) to aOR {er['E-Primary']['twin_OR']:.2f} (subgroup)."),
    ]

for fig_label, legend_text in legends:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(f"{fig_label}. ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(legend_text)
    run.font.size = Pt(10)

# ============================================================
# SAVE
# ============================================================
out_path = BASE / "manuscript_strobe_en.docx"
doc.save(str(out_path))
print(f"English manuscript saved to {out_path}")

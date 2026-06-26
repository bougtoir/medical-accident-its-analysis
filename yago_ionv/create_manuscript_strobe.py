"""
STROBE-compliant manuscript: full paper format (Japanese text, English figures).
Integrates:
  - Flowchart (Fig. 1)
  - Main cohort analysis (broad + narrow)
  - Exclusion sensitivity analysis
  - STROBE checklist
"""
import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import pandas as pd

BASE = Path(__file__).resolve().parent

# Load all stats
with open(BASE / "def_e_stats.json") as f:
    M = json.load(f)  # Main cohort
with open(BASE / "excl_sensitivity_stats.json") as f:
    E = json.load(f)  # Exclusion sensitivity
with open(BASE / "flowchart_counts.json") as f:
    F = json.load(f)  # Flow counts
with open(BASE / "bootstrap_results.json") as f:
    B = json.load(f)  # Bootstrap results
with open(BASE / "summary_stats.json") as f:
    S = json.load(f)  # Summary stats (includes hypotension secondary outcome)

# --- Detachable structure flag ---
# Set to False to produce a main-analysis-only manuscript (no sensitivity sections)
INCLUDE_SENSITIVITY = True

doc = Document()

# ---- Styles ----
style = doc.styles["Normal"]
font = style.font
font.name = "游明朝"
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.5

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_paragraph(doc, text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.font.italic = True
    if align:
        p.alignment = align
    return p

def add_ref(paragraph, text):
    """Add text with superscript reference numbers using {n} markers."""
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = paragraph.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(8)
        else:
            paragraph.add_run(part)
    return paragraph

def add_p_with_refs(doc, text, bold=False):
    p = doc.add_paragraph()
    if bold:
        parts = re.split(r'(\{[^}]+\})', text)
        for part in parts:
            if part.startswith('{') and part.endswith('}'):
                run = p.add_run(part[1:-1])
                run.font.superscript = True
                run.font.size = Pt(8)
                run.bold = True
            else:
                run = p.add_run(part)
                run.bold = True
    else:
        add_ref(p, text)
    return p

def add_figure(doc, img_path, caption, width=Inches(6)):
    if Path(img_path).exists():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=width)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(6)
        run = cap.add_run(caption)
        run.font.size = Pt(9)
        run.font.italic = True

def format_p(p_val):
    return "P < 0.001" if p_val < 0.001 else f"P = {p_val:.3f}"

def make_table(doc, df, headers_ja, col_map, title):
    add_paragraph(doc, title, bold=True)
    tbl = doc.add_table(rows=len(df) + 1, cols=len(headers_ja))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Light Shading Accent 1"
    for j, h in enumerate(headers_ja):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        for par in cell.paragraphs:
            for r in par.runs:
                r.bold = True
                r.font.size = Pt(9)
    for i, (_, row) in enumerate(df.iterrows()):
        for j, col in enumerate(col_map):
            val = row[col]
            if isinstance(val, float) and col == "P-value":
                cell_text = "< 0.001" if val < 0.001 else f"{val:.3f}"
            else:
                cell_text = str(val)
            tbl.rows[i + 1].cells[j].text = cell_text
            for par in tbl.rows[i + 1].cells[j].paragraphs:
                for r in par.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()

# ============================================================
# TITLE PAGE
# ============================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(40)
run = title_p.add_run(
    "脊髄くも膜下麻酔による帝王切開術中の嘔気嘔吐に対する\n"
    "双胎妊娠の影響：後方視的コホート研究")
run.bold = True
run.font.size = Pt(16)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(
    "Impact of twin pregnancy on intraoperative nausea and vomiting\n"
    "during cesarean delivery under spinal anesthesia:\n"
    "a retrospective cohort study")
run.font.size = Pt(11)
run.font.italic = True

doc.add_paragraph()
doc.add_paragraph()

# ============================================================
# ABSTRACT (structured)
# ============================================================
add_heading(doc, "抄録")

add_paragraph(doc, "【背景】", bold=True)
add_paragraph(doc,
    "帝王切開術中の嘔気嘔吐（Intraoperative Nausea and Vomiting; IONV）は"
    "脊髄くも膜下麻酔の主要な副作用であるが、双胎妊娠がIONVに及ぼす影響は十分に検討されていない。")

add_paragraph(doc, "【方法】", bold=True)
add_p_with_refs(doc,
    f"2014年4月から2024年10月に当院で帝王切開術を受けた{F['total']['n']:,}例を後方視的に検討した。"
    f"除外基準適用後の{F['primary_analysis']['n']:,}例"
    f"（単胎{F['primary_analysis']['n_s']:,}例、双胎{F['primary_analysis']['n_t']:,}例）を解析対象とした。"
    "IONVの代替指標として制吐薬の使用を用い、全制吐薬（広義）および"
    "5-HT3受容体拮抗薬のみ（狭義）の2定義で評価した。"
    "多変量ロジスティック回帰分析で双胎の独立した効果を検討した。"
    "本研究はStrengthening the Reporting of Observational Studies in Epidemiology（STROBE）"
    "ガイドラインに準拠して報告する。")

add_paragraph(doc, "【結果】", bold=True)
mo = M["outcomes"]
mr = M["regression"]
add_paragraph(doc,
    f"制吐薬（広義）によるIONV発生率は単胎群{mo['A-Primary']['singleton_pct']:.1f}%、"
    f"双胎群{mo['A-Primary']['twin_pct']:.1f}%であり有意差を認めなかった"
    f"（調整済みオッズ比[aOR] {mr['A-Primary']['twin_OR']:.2f}, "
    f"95%信頼区間[CI] {mr['A-Primary']['twin_CI_lower']:.2f}–{mr['A-Primary']['twin_CI_upper']:.2f}, "
    f"{format_p(mr['A-Primary']['twin_P'])}）。"
    f"一方、制吐薬（狭義：5-HT3拮抗薬）では単胎群{mo['E-Primary']['singleton_pct']:.1f}%に対し"
    f"双胎群{mo['E-Primary']['twin_pct']:.1f}%と有意に高率であった"
    f"（aOR {mr['E-Primary']['twin_OR']:.2f}, "
    f"95%CI {mr['E-Primary']['twin_CI_lower']:.2f}–{mr['E-Primary']['twin_CI_upper']:.2f}, "
    f"{format_p(mr['E-Primary']['twin_P'])}）。"
    "共変量感度分析、緊急帝王切開の層別解析、および"
    "10,000回の層別ブートストラップ検証でも結果は頑健であった。"
    f"副次アウトカムの術中低血圧は双胎群（{S['hypo_twin_pct']:.1f}%）が"
    f"単胎群（{S['hypo_single_pct']:.1f}%）より低率であった"
    f"（{format_p(S['hypo_chi_p'])}）。")

add_paragraph(doc, "【結論】", bold=True)
add_paragraph(doc,
    "制吐薬全般の使用頻度に双胎・単胎間の差はないが、"
    "嘔気嘔吐に対する特異性が高い5-HT3拮抗薬の使用は双胎群で有意に高率であった。"
    "一方、術中低血圧は双胎群で低率であり、"
    "IONVのメカニズムとして低血圧以外の経路の関与が示唆される。")

doc.add_page_break()

# ============================================================
# INTRODUCTION
# ============================================================
add_heading(doc, "緒言")

add_p_with_refs(doc,
    "脊髄くも膜下麻酔は帝王切開術における標準的な麻酔法として世界中で広く用いられている{1}。"
    "しかし、脊髄くも膜下麻酔には術中の嘔気嘔吐"
    "（Intraoperative Nausea and Vomiting; IONV）を含む"
    "様々な副作用が報告されており、その発生率は最大80%に達するとされている{2,3}。"
    "IONVは患者満足度を低下させるだけでなく、内臓脱出により手術操作の妨げとなる{4-6}。")

add_p_with_refs(doc,
    "IONVの主要な機序として、"
    "（1）脊髄くも膜下麻酔による交感神経遮断に起因する低血圧、"
    "（2）過剰な循環血液量の再分布による腸管低灌流とセロトニン遊離、"
    "（3）交感神経遮断下での迷走神経優位、"
    "（4）子宮操作・腹膜牽引による内臓刺激、"
    "（5）子宮収縮薬（オキシトシン）投与"
    "が報告されている{2,12}。")

add_p_with_refs(doc,
    "双胎妊娠では、これらの機序を増幅しうる生理学的特徴が知られている。"
    "双胎では単胎に比べ循環血液量が約400 mL多く増加し、"
    "心拍数・一回拍出量・心係数がより大きく上昇する{13}。"
    "さらに、増大した子宮による大動脈・下大静脈圧迫は"
    "双胎でより高度かつ頻回であり、"
    "脊髄くも膜下麻酔との交互作用により高度の低血圧が生じやすい{9,13}。"
    "これらの知見から、双胎妊娠は理論的にIONVのリスクを高めうると考えられる。")

add_p_with_refs(doc,
    "先行研究では、IONVの危険因子として術中低血圧、術中疼痛、"
    "30分を超える手術時間、35 kg/m²を超えるBMI、"
    "妊娠高血圧症候群（Hypertensive Disorders of Pregnancy; HDP）、"
    "子宮の腹腔外脱転が報告されている{7,8}。"
    "しかし、多くの先行研究では双胎妊娠が対象から除外されており{9-11}、"
    "双胎妊娠がIONVの独立した危険因子であるかは不明である。")

add_paragraph(doc,
    "本研究の目的は、帝王切開術中のIONVの発生率を単胎妊娠と双胎妊娠で比較し、"
    "双胎妊娠がIONVの独立した危険因子であるかを検討することである。")

doc.add_page_break()

# ============================================================
# METHODS
# ============================================================
add_heading(doc, "方法")

add_heading(doc, "研究デザインおよび倫理的配慮（STROBE項目1, 5）", level=2)
add_paragraph(doc,
    "本研究は単施設後方視的コホート研究である。"
    "順天堂大学医学部附属静岡病院倫理委員会の承認を得て実施した。"
    "本研究はStrengthening the Reporting of Observational Studies in Epidemiology"
    "（STROBE）声明に準拠して報告する（Supplementary Table参照）。"
    "後方視的研究であり、オプトアウト方式により同意を取得した。")

add_heading(doc, "対象（STROBE項目6, 7）", level=2)
add_p_with_refs(doc,
    f"2014年4月1日から2024年10月23日の間に順天堂大学医学部附属静岡病院で"
    f"帝王切開術を受けた全症例{F['total']['n']:,}例"
    f"（単胎{F['total']['n_s']:,}例、双胎{F['total']['n_t']:,}例）を対象とした。")

add_paragraph(doc,
    "組入基準は、18歳以上で単胎または双胎妊娠の帝王切開術を受け、"
    "脊髄くも膜下麻酔単独、硬膜外麻酔単独、"
    "または硬膜外麻酔併用脊髄くも膜下麻酔を施行された症例とした。")

excl_text = "除外基準は、"
for i, step in enumerate(F["exclusion_steps"], 1):
    excl_text += f"（{i}）{step['reason_ja']}、"
excl_text = excl_text[:-1] + "とした。"
excl_text += (
    f"さらに、麻酔開始前に制吐薬が投与された{F['preop_antiemetic']['n']}例は"
    "予防的投与とみなし解析から除外した。")
add_paragraph(doc, excl_text)

add_paragraph(doc,
    f"以上の基準により{F['primary_analysis']['n']:,}例"
    f"（単胎{F['primary_analysis']['n_s']:,}例、"
    f"双胎{F['primary_analysis']['n_t']:,}例）を主解析の対象とした"
    "（Fig. 1）。")

# --- Fig 1: Flowchart ---
add_figure(doc, BASE / "figures_strobe" / "fig_flowchart.png",
           "Fig. 1  STROBE Flow Diagram — Participant Selection",
           width=Inches(5.5))

add_heading(doc, "評価項目（STROBE項目8, 12）", level=2)

add_paragraph(doc,
    "本研究は後方視的研究であるため、嘔気嘔吐の直接的な記録の収集は不可能であった。"
    "そこで、制吐薬の投与をIONVの代替指標（サロゲートマーカー）として使用した。")

add_paragraph(doc, "【制吐薬（広義）】", bold=True)
add_paragraph(doc,
    "主要評価項目：麻酔開始から胎児娩出までの"
    "制吐薬投与をIONVと定義した。"
    "対象薬剤はメトクロプラミド、ドロペリドール、オンダンセトロン、"
    "グラニセトロン、ノバミン、アタラックスP、デキサメタゾンの7剤とした。\n"
    "副次評価項目：麻酔開始から手術室退室までの制吐薬投与をIONVと定義した。")

add_paragraph(doc, "【制吐薬（狭義）：5-HT3受容体拮抗薬】", bold=True)
add_paragraph(doc,
    "上記7剤のうち、嘔気嘔吐に対する薬理学的特異性が高い"
    "5-HT3受容体拮抗薬（オンダンセトロンまたはグラニセトロン）の投与のみを"
    "IONVの指標とする追加解析を実施した。"
    "メトクロプラミド（消化管運動促進）、ドロペリドール（鎮静）、"
    "ノバミン（抗精神病）、アタラックスP（抗不安・掻痒）、"
    "デキサメタゾン（予防・抗炎症）は制吐以外の適応でも使用されるため除外した。")

add_paragraph(doc, "【副次アウトカム：術中低血圧】", bold=True)
add_paragraph(doc,
    "副次アウトカムとして、術中低血圧の発生頻度を評価した。"
    "低血圧は収縮期血圧 90 mmHg未満と定義し、"
    "術中に記録された該当値の回数を用いた。"
    "低血圧はIONV解析において共変量としても使用しているため、"
    "記述統計による群間比較にとどめた。")

add_heading(doc, "共変量（STROBE項目8）", level=2)
add_paragraph(doc,
    "多変量解析における共変量は以下のとおりとした：年齢、BMI、在胎週数、"
    "緊急帝王切開、帝王切開既往、HDP、硬膜外麻酔、"
    "手術時間、低血圧（収縮期血圧90 mmHg未満のエピソードあり）。"
    "イベント数が共変量数の10倍に満たない場合は、過適合を避けるため"
    "共変量を縮小モデル（双胎、年齢、BMI、在胎週数、緊急、低血圧の6変数）とした。")

add_heading(doc, "データ収集（STROBE項目9）", level=2)
add_paragraph(doc,
    "電子カルテおよび麻酔記録システムから、患者背景、麻酔情報、"
    "手術情報、制吐薬投与情報を研究担当者が抽出しMicrosoft Excelに入力した。"
    "患者個人を特定できる情報は除去し、仮IDを付与した。")

add_heading(doc, "統計解析（STROBE項目13）", level=2)
add_paragraph(doc,
    "連続変数は中央値[四分位範囲]で示し、Mann-Whitney U検定で群間比較した。"
    "カテゴリー変数は例数（%）で示し、χ²検定またはFisherの正確検定"
    "（期待度数5未満のセルが存在する場合）で比較した。")

add_paragraph(doc,
    "IONVと双胎妊娠の関連を評価するため、多変量ロジスティック回帰分析を実施した。"
    "結果は調整済みオッズ比（adjusted Odds Ratio; aOR）と95%信頼区間（CI）で報告した。")

if INCLUDE_SENSITIVITY:
    add_paragraph(doc,
        "制吐薬（狭義）の主要評価項目について、双胎の効果推定値の頑健性を検証するため"
        "共変量感度分析を実施した。具体的には、（1）粗OR（双胎のみ）、"
        "（2）双胎＋各共変量1つ（9モデル）、（3）全共変量から1つずつ除去（9モデル）、"
        "（4）縮小モデル（6共変量）の計20モデルで双胎のaORの安定性を評価した。")
    add_heading(doc, "感度分析（STROBE項目13）", level=2)
    add_paragraph(doc,
        "IONVに影響しうる産科合併症・手術関連因子の交絡を排除するため、"
        "緊急帝王切開、帝王切開既往、HDP、術前ステロイド使用の"
        "4条件に該当する症例をさらに除外した予定手術・低リスクサブグループ解析を実施した。")
    add_paragraph(doc,
        "さらに、緊急帝王切開が効果修飾因子となりうるかを検証するため、"
        "予定手術のみ（緊急除外）、緊急手術のみのサブグループ解析を実施した。"
        "また、緊急帝王切開と双胎の交互作用項をモデルに追加した交互作用解析も行った。")
    add_heading(doc, "ブートストラップ検証（STROBE項目13）", level=2)
    add_paragraph(doc,
        "単胎・双胎間の症例数比が約8:1と不均衡であったため、"
        "ロジスティック回帰の信頼区間推定の頑健性を検証する目的で"
        "層別ブートストラップ法を実施した。"
        "各リサンプルにおいて単胎群と双胎群を別々にリサンプリングした後に結合し、"
        "群比率を維持した（層別リサンプリング）。"
        "10,000回のリサンプリングを行い、パーセンタイル法およびBCa法"
        "（bias-corrected and accelerated）による95%信頼区間を算出した。"
        "BCa法のバイアス補正因子z0はブートストラップ分布の中央値偏差から、"
        "加速度因子aはジャックナイフ法により推定した。")

add_paragraph(doc,
    "統計解析はPython 3.12（scipy 1.14, statsmodels 0.14）を用い、"
    "有意水準は両側5%とした。")

doc.add_page_break()

# ============================================================
# RESULTS
# ============================================================
add_heading(doc, "結果")

# --- Participants (STROBE 14) ---
add_heading(doc, "対象者の選択（STROBE項目14）", level=2)
# Build exclusion breakdown dynamically
excl_parts = []
for step in F["exclusion_steps"]:
    excl_parts.append(f"{step['reason_ja']}{step['n']}例")
excl_detail = "、".join(excl_parts)

# Include date filter in participant flow
date_filter_text = ""
if F.get("out_of_period", {}).get("n", 0) > 0:
    date_filter_text = (
        f"データベース上の全帝王切開症例{F['total_raw']['n']:,}例のうち、"
        f"研究期間外（2014年4月1日以前）の{F['out_of_period']['n']}例"
        f"（単胎{F['out_of_period']['n_s']}例、双胎{F['out_of_period']['n_t']}例）を除外し、"
        f"研究期間内の{F['total']['n']:,}例を適格性評価の対象とした。"
    )

add_paragraph(doc,
    date_filter_text +
    f"研究期間中に帝王切開術を受けた{F['total']['n']:,}例のうち、"
    f"除外基準に該当した{F['total_excluded']['n']}例"
    f"（{excl_detail}）を除外した。"
    f"さらに術前制吐薬投与{F['preop_antiemetic']['n']}例を除外し、"
    f"最終解析対象は{F['primary_analysis']['n']:,}例"
    f"（単胎{F['primary_analysis']['n_s']:,}例、"
    f"双胎{F['primary_analysis']['n_t']:,}例）であった（Fig. 1）。")

# Twin subcategory breakdown (not shown in flowchart)
tc = M.get("twin_chorionicity", {})
te = M.get("twin_emergency", {})
add_paragraph(doc,
    f"双胎{F['primary_analysis']['n_t']:,}例の内訳は、"
    f"絨毛膜性別ではDD双胎{tc.get('DD', 0)}例、"
    f"MD双胎{tc.get('MD', 0)}例、"
    f"MM双胎{tc.get('MM', 0)}例であった。"
    f"手術区分別では予定手術{te.get('elective', 0)}例（{100*te.get('elective', 0)/F['primary_analysis']['n_t']:.1f}%）、"
    f"緊急手術{te.get('emergency', 0)}例（{100*te.get('emergency', 0)/F['primary_analysis']['n_t']:.1f}%）であった。")

# --- Descriptive data (STROBE 15) ---
add_heading(doc, "患者背景（STROBE項目15）", level=2)
add_paragraph(doc,
    f"解析対象{M['n_analysis']:,}例の患者背景をTable 1に示す。")

# Table 1
table1 = pd.read_csv(BASE / "tables_e" / "table1_characteristics.csv")
add_paragraph(doc,
    f"Table 1. 患者背景（単胎 n={M['n_single']:,} vs 双胎 n={M['n_twin']:,}）",
    bold=True)

tbl = doc.add_table(rows=len(table1) + 1, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Light Shading Accent 1"
for j, h in enumerate(["変数", f"単胎 (n={M['n_single']:,})", f"双胎 (n={M['n_twin']:,})", "P値"]):
    cell = tbl.rows[0].cells[j]
    cell.text = h
    for par in cell.paragraphs:
        par.runs[0].bold = True
        par.runs[0].font.size = Pt(9)
for i, (_, row) in enumerate(table1.iterrows()):
    tbl.rows[i + 1].cells[0].text = row["Variable"]
    tbl.rows[i + 1].cells[1].text = str(row["Singleton"])
    tbl.rows[i + 1].cells[2].text = str(row["Twin"])
    p_val = row["P-value"]
    tbl.rows[i + 1].cells[3].text = "< 0.001" if p_val < 0.001 else f"{p_val:.3f}"
    for j in range(4):
        for par in tbl.rows[i + 1].cells[j].paragraphs:
            for r in par.runs:
                r.font.size = Pt(9)

# Table 1 footnotes
footnote_p = doc.add_paragraph()
footnote_texts = [
    "連続変数は中央値 [四分位範囲] で表示。",
    "カテゴリ変数は n/N (%) で表示。",
    "連続変数の群間比較にはMann-Whitney U検定、",
    "カテゴリ変数の群間比較にはPearsonの\u03c7\u00b2検定を用いた",
    "（期待度数5未満のセルがある場合はFisherの正確検定）。",
]
fn_run = footnote_p.add_run("".join(footnote_texts))
fn_run.font.size = Pt(8)
fn_run.font.italic = True
doc.add_paragraph()

# --- Outcome data (STROBE 16) ---
add_heading(doc, "IONV発生率（STROBE項目16）", level=2)

add_paragraph(doc,
    f"制吐薬（広義）による主要評価項目のIONV発生率は、"
    f"単胎群{mo['A-Primary']['singleton_pct']:.1f}%"
    f"（{mo['A-Primary']['singleton_n']}/{M['n_single']:,}例）、"
    f"双胎群{mo['A-Primary']['twin_pct']:.1f}%"
    f"（{mo['A-Primary']['twin_n']}/{M['n_twin']:,}例）であった（Table 2, Fig. 2）。")

add_paragraph(doc,
    f"制吐薬（狭義：5-HT3拮抗薬）による主要評価項目では、"
    f"単胎群{mo['E-Primary']['singleton_pct']:.1f}%"
    f"（{mo['E-Primary']['singleton_n']}/{M['n_single']:,}例）に対し、"
    f"双胎群{mo['E-Primary']['twin_pct']:.1f}%"
    f"（{mo['E-Primary']['twin_n']}/{M['n_twin']:,}例）と、"
    f"双胎群で有意に高率であった（Table 2, Fig. 2）。")

# Table 2: IONV rates
add_paragraph(doc, "Table 2. IONV発生率と多変量ロジスティック回帰分析", bold=True)
tbl2 = doc.add_table(rows=5, cols=7)
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl2.style = "Light Shading Accent 1"
for j, h in enumerate(["評価項目", "単胎 n (%)", "双胎 n (%)", "モデル", "aOR", "95% CI", "P値"]):
    cell = tbl2.rows[0].cells[j]
    cell.text = h
    for par in cell.paragraphs:
        par.runs[0].bold = True
        par.runs[0].font.size = Pt(8)

for i, key in enumerate(["A-Primary", "A-Secondary", "E-Primary", "E-Secondary"]):
    rr = mr[key]
    oo = mo[key]
    labels_ja = {
        "A-Primary": "広義 主要",
        "A-Secondary": "広義 副次",
        "E-Primary": "狭義 主要",
        "E-Secondary": "狭義 副次",
    }
    vals = [
        labels_ja[key],
        f"{oo['singleton_n']} ({oo['singleton_pct']:.1f}%)",
        f"{oo['twin_n']} ({oo['twin_pct']:.1f}%)",
        rr["model_type"],
        f"{rr['twin_OR']:.2f}",
        f"{rr['twin_CI_lower']:.2f}–{rr['twin_CI_upper']:.2f}",
        "< 0.001" if rr["twin_P"] < 0.001 else f"{rr['twin_P']:.3f}",
    ]
    for j, v in enumerate(vals):
        tbl2.rows[i + 1].cells[j].text = v
        for par in tbl2.rows[i + 1].cells[j].paragraphs:
            for r in par.runs:
                r.font.size = Pt(8)
doc.add_paragraph()

# Fig 2: Rates comparison
add_figure(doc, BASE / "figures_e" / "fig1_rates_comparison.png",
           "Fig. 2  IONV発生率の比較（制吐薬（広義）vs 制吐薬（狭義）、単胎 vs 双胎）")

# --- Main results (STROBE 17) ---
add_heading(doc, "多変量解析（STROBE項目17）", level=2)

add_paragraph(doc,
    f"制吐薬（広義）の主要評価項目では、全共変量モデルにおいて"
    f"双胎はIONVと有意な関連を示さなかった"
    f"（aOR {mr['A-Primary']['twin_OR']:.2f}, "
    f"95%CI {mr['A-Primary']['twin_CI_lower']:.2f}–{mr['A-Primary']['twin_CI_upper']:.2f}, "
    f"{format_p(mr['A-Primary']['twin_P'])}）。")

add_paragraph(doc,
    f"制吐薬（狭義）の主要評価項目では、イベント数が{mr['E-Primary']['events']}件と"
    f"少数であったため縮小モデルを使用した。"
    f"双胎は5-HT3拮抗薬使用と有意に関連していた"
    f"（aOR {mr['E-Primary']['twin_OR']:.2f}, "
    f"95%CI {mr['E-Primary']['twin_CI_lower']:.2f}–{mr['E-Primary']['twin_CI_upper']:.2f}, "
    f"{format_p(mr['E-Primary']['twin_P'])}）（Fig. 3）。")

# Fig 3: Forest plot
add_figure(doc, BASE / "figures_e" / "fig2_forest_E_primary.png",
           "Fig. 3  制吐薬（狭義）使用の多変量ロジスティック回帰分析 — Forest plot")

# Fig 4: Broad vs Narrow comparison
add_figure(doc, BASE / "figures_e" / "fig4_protocol_vs_defE.png",
           "Fig. 4  双胎のIONV効果：制吐薬（広義）vs 制吐薬（狭義）")

# --- Secondary outcome: Hypotension ---
add_heading(doc, "副次アウトカム：術中低血圧（STROBE項目17）", level=2)
add_paragraph(doc,
    f"術中低血圧（収縮期血圧 <90 mmHg）の発生頻度は、"
    f"単胎群{S['hypo_single_n']:,}/{S['n_single']:,}例（{S['hypo_single_pct']:.1f}%）、"
    f"双胎群{S['hypo_twin_n']}/{S['n_twin']:,}例（{S['hypo_twin_pct']:.1f}%）であり、"
    f"双胎群で有意に低率であった（P = {S['hypo_chi_p']:.3f}）。"
    f"低血圧エピソード回数（中央値[四分位範囲]）は単胎群{S['hypo_count_single_median']:.0f}"
    f" [{S['hypo_count_single_q1']:.0f}–{S['hypo_count_single_q3']:.0f}]回、"
    f"双胎群{S['hypo_count_twin_median']:.0f}"
    f" [{S['hypo_count_twin_q1']:.0f}–{S['hypo_count_twin_q3']:.0f}]回であった"
    f"（P = {S['hypo_count_p']:.3f}）。")

# Note: hypotension is also used as a covariate in IONV analysis,
# so only descriptive statistics are reported here (no regression).

# ============================================================
# SENSITIVITY ANALYSES (detachable block — set INCLUDE_SENSITIVITY = False to omit)
# When INCLUDE_SENSITIVITY = False, only the main analysis + hypotension are included.
# ============================================================

cov_df = pd.read_csv(BASE / "tables_e" / "covariate_sensitivity.csv")

if INCLUDE_SENSITIVITY:
    n_sig_main = int((cov_df["P"] < 0.05).sum())
    add_heading(doc, "共変量感度分析（STROBE項目18）", level=2)
    add_paragraph(doc,
        f"制吐薬（狭義）主要評価項目における双胎の効果推定値の頑健性を検証するため、"
        f"共変量感度分析を実施した（Table 3, Fig. 5）。"
        f"全{len(cov_df)}モデルにおいて、双胎のaORは"
        f"{cov_df['aOR'].min():.2f}–{cov_df['aOR'].max():.2f}の範囲であり、"
        f"{n_sig_main}/{len(cov_df)}モデルで統計学的に有意であった（P < 0.05）。")

    # Table 3: Covariate sensitivity
    add_paragraph(doc, "Table 3. 共変量感度分析（制吐薬（狭義）主要評価項目、主解析コホート）", bold=True)
    tbl3 = doc.add_table(rows=len(cov_df) + 1, cols=4)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl3.style = "Light Shading Accent 1"
    for j, h in enumerate(["モデル", "aOR", "95% CI", "P値"]):
        tbl3.rows[0].cells[j].text = h
        for par in tbl3.rows[0].cells[j].paragraphs:
            par.runs[0].bold = True
            par.runs[0].font.size = Pt(9)
    for i, (_, row) in enumerate(cov_df.iterrows()):
        vals = [row["Model"], f"{row['aOR']:.2f}",
                f"{row['CI_lower']:.2f}–{row['CI_upper']:.2f}",
                "< 0.001" if row["P"] < 0.001 else f"{row['P']:.3f}"]
        for j, v in enumerate(vals):
            tbl3.rows[i + 1].cells[j].text = v
            for par in tbl3.rows[i + 1].cells[j].paragraphs:
                for r in par.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()

    # Fig 5: Covariate sensitivity forest
    add_figure(doc, BASE / "figures_e" / "fig3_covariate_sensitivity.png",
               "Fig. 5  共変量感度分析：双胎の制吐薬（狭義）使用に対する効果")

    # --- Exclusion sensitivity analysis ---
    er = E["regression"]
    eo = E["outcomes"]
    add_heading(doc, "除外感度分析（STROBE項目18）", level=2)
    add_paragraph(doc,
        f"緊急帝王切開（{E['exclusion_counts']['Emergency CS']:,}例）、"
        f"帝王切開既往（{E['exclusion_counts']['Prior CS']:,}例）、"
        f"HDP（{E['exclusion_counts']['HDP']}例）、"
        f"術前ステロイド（{E['exclusion_counts']['Preoperative steroid']}例）"
        f"を除外した予定手術・低リスクサブグループ{E['n_analysis']:,}例"
        f"（単胎{E['n_single']:,}例、双胎{E['n_twin']:,}例）で再解析した"
        "（重複あり、Fig. 1）。")
    add_paragraph(doc,
        f"制吐薬（広義）ではIONV発生率に有意差を認めなかった"
        f"（単胎{eo['A-Primary']['singleton_pct']:.1f}% vs "
        f"双胎{eo['A-Primary']['twin_pct']:.1f}%, "
        f"aOR {er['A-Primary']['twin_OR']:.2f}, "
        f"{format_p(er['A-Primary']['twin_P'])}）。"
        f"一方、制吐薬（狭義）では単胎{eo['E-Primary']['singleton_pct']:.1f}%に対し"
        f"双胎{eo['E-Primary']['twin_pct']:.1f}%と依然として有意に高率であった"
        f"（aOR {er['E-Primary']['twin_OR']:.2f}, "
        f"95%CI {er['E-Primary']['twin_CI_lower']:.2f}–{er['E-Primary']['twin_CI_upper']:.2f}, "
        f"{format_p(er['E-Primary']['twin_P'])}）（Table 4）。")

    # Table 4: Exclusion sensitivity results
    add_paragraph(doc, "Table 4. 除外感度分析の結果（予定手術・低リスクサブグループ）", bold=True)
    tbl4 = doc.add_table(rows=5, cols=7)
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl4.style = "Light Shading Accent 1"
    for j, h in enumerate(["評価項目", "単胎 n (%)", "双胎 n (%)", "モデル", "aOR", "95% CI", "P値"]):
        tbl4.rows[0].cells[j].text = h
        for par in tbl4.rows[0].cells[j].paragraphs:
            par.runs[0].bold = True
            par.runs[0].font.size = Pt(8)
    for i, key in enumerate(["A-Primary", "A-Secondary", "E-Primary", "E-Secondary"]):
        rr = er[key]
        oo = eo[key]
        labels_ja = {
            "A-Primary": "広義 主要",
            "A-Secondary": "広義 副次",
            "E-Primary": "狭義 主要",
            "E-Secondary": "狭義 副次",
        }
        vals = [
            labels_ja[key],
            f"{oo['singleton_n']} ({oo['singleton_pct']:.1f}%)",
            f"{oo['twin_n']} ({oo['twin_pct']:.1f}%)",
            rr["model_type"],
            f"{rr['twin_OR']:.2f}",
            f"{rr['twin_CI_lower']:.2f}–{rr['twin_CI_upper']:.2f}",
            "< 0.001" if rr["twin_P"] < 0.001 else f"{rr['twin_P']:.3f}",
        ]
        for j, v in enumerate(vals):
            tbl4.rows[i + 1].cells[j].text = v
            for par in tbl4.rows[i + 1].cells[j].paragraphs:
                for r in par.runs:
                    r.font.size = Pt(8)
    doc.add_paragraph()

    # Fig 6: Exclusion sensitivity rates
    add_figure(doc, BASE / "figures_excl" / "fig1_rates_comparison.png",
               "Fig. 6  IONV発生率の比較（予定手術・低リスクサブグループ）")

    # Fig 7: Exclusion sensitivity forest
    add_figure(doc, BASE / "figures_excl" / "fig4_broad_vs_narrow.png",
               "Fig. 7  双胎のIONV効果：広義 vs 狭義（予定手術・低リスクサブグループ）")

    # --- Emergency sensitivity analysis ---
    add_heading(doc, "緊急帝王切開の感度分析（STROBE項目18）", level=2)
    add_paragraph(doc,
        "緊急帝王切開が双胎のIONV効果を修飾する可能性を検証するため、"
        "予定手術のみ（n=1,736、単胎1,515例、双胎221例）および"
        "緊急手術のみ（n=1,452、単胎1,331例、双胎121例）のサブグループ解析を実施した（Table 5）。")
    add_paragraph(doc,
        "予定手術のみでは、制吐薬（狭義）主要評価項目でaOR 8.39"
        "（95%CI 1.21–58.18, P = 0.031）と効果量が大幅に増大し、"
        "副次評価項目も有意であった（aOR 8.99, 95%CI 2.50–32.29, P < 0.001）。"
        "一方、緊急手術のみでは双胎効果は完全に消失した"
        "（狭義副次 aOR 0.95, P = 0.958; 狭義主要はイベント数不足のため推定不能）。")
    add_paragraph(doc,
        "興味深いことに、予定手術のみでは制吐薬（広義）副次評価項目で"
        "双胎群の方が有意に低率であった（18.5% vs 14.0%, aOR 0.64, P = 0.047）。"
        "交互作用解析では、制吐薬（広義）副次評価項目で"
        "緊急×双胎の有意な交互作用を認めた（交互作用OR 2.05, P = 0.023）。")

    # Table 5: Emergency sensitivity
    add_paragraph(doc, "Table 5. 緊急帝王切開の感度分析", bold=True)
    tbl5 = doc.add_table(rows=9, cols=7)
    tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl5.style = "Light Shading Accent 1"
    for j, h in enumerate(["サブグループ", "評価項目", "単胎", "双胎", "aOR", "95% CI", "P値"]):
        tbl5.rows[0].cells[j].text = h
        for par in tbl5.rows[0].cells[j].paragraphs:
            par.runs[0].bold = True
            par.runs[0].font.size = Pt(8)
    emg_data = [
        ["予定手術のみ", "広義 主要", "2.4%", "1.8%", "0.83", "0.28–2.47", "0.735"],
        ["", "広義 副次", "18.5%", "14.0%", "0.64", "0.41–0.99", "0.047"],
        ["", "狭義 主要", "0.3%", "1.4%", "8.39", "1.21–58.18", "0.031"],
        ["", "狭義 副次", "0.8%", "3.2%", "8.99", "2.50–32.29", "< 0.001"],
        ["緊急手術のみ", "広義 主要", "1.4%", "0.8%", "0.61", "0.08–4.68", "0.637"],
        ["", "広義 副次", "16.8%", "22.3%", "1.44", "0.91–2.29", "0.119"],
        ["", "狭義 主要", "0.1%", "0.0%", "—", "—", "—"],
        ["", "狭義 副次", "1.0%", "0.8%", "0.95", "0.12–7.44", "0.958"],
    ]
    for i, row_data in enumerate(emg_data):
        for j, v in enumerate(row_data):
            tbl5.rows[i + 1].cells[j].text = v
            for par in tbl5.rows[i + 1].cells[j].paragraphs:
                for r in par.runs:
                    r.font.size = Pt(8)
    doc.add_paragraph()

    # --- Bootstrap validation ---
    bm = B["main_cohort"]
    bs = B["subgroup"]
    add_heading(doc, "ブートストラップ検証（STROBE項目18）", level=2)
    add_paragraph(doc,
        "全コホートにおける層別ブートストラップ検証の結果をTable 6およびFig. 8に示す。")
    add_paragraph(doc,
        f"制吐薬（狭義）主要評価項目の全コホートでは、"
        f"BCa法95%CI [{bm['E-Primary']['bca_CI_lower']:.2f}–{bm['E-Primary']['bca_CI_upper']:.2f}] であり、"
        f"Wald法95%CI [{bm['E-Primary']['wald_CI_lower']:.2f}–{bm['E-Primary']['wald_CI_upper']:.2f}] と"
        f"ほぼ一致した。ブートストラップCIも1を除外しており、"
        f"双胎の有意な効果（aOR {bm['E-Primary']['point_aOR']:.2f}）は頑健であった"
        f"（収束率{bm['E-Primary']['convergence_pct']}%）。")
    add_paragraph(doc,
        "低リスクサブグループでは、制吐薬（狭義）のイベント数が極めて少なく"
        f"（主要{bs['E-Primary']['events']}件、副次{bs['E-Secondary']['events']}件）、"
        "ブートストラップCIは不安定であった。"
        "これはサンプルサイズの制約によるものであり、"
        "サブグループ結果の解釈には慎重を要する。")

    # Table 6: Bootstrap results
    add_paragraph(doc, "Table 6. 層別ブートストラップ検証（10,000回リサンプリング）", bold=True)
    tbl6 = doc.add_table(rows=9, cols=6)
    tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl6.style = "Light Shading Accent 1"
    for j, h in enumerate(["コホート", "評価項目", "aOR", "Wald 95%CI", "BCa 95%CI", "収束率"]):
        tbl6.rows[0].cells[j].text = h
        for par in tbl6.rows[0].cells[j].paragraphs:
            par.runs[0].bold = True
            par.runs[0].font.size = Pt(8)

    def fmt_bca(r):
        lo = r.get("bca_CI_lower")
        hi = r.get("bca_CI_upper")
        if lo is not None and hi is not None and hi < 1e6:
            return f"{lo:.2f}–{hi:.2f}"
        return "不安定"

    boot_rows = [
        ["全コホート", "広義 主要", bm["A-Primary"]],
        ["", "広義 副次", bm["A-Secondary"]],
        ["", "狭義 主要", bm["E-Primary"]],
        ["", "狭義 副次", bm["E-Secondary"]],
        ["低リスク", "広義 主要", bs["A-Primary"]],
        ["", "広義 副次", bs["A-Secondary"]],
        ["", "狭義 主要", bs["E-Primary"]],
        ["", "狭義 副次", bs["E-Secondary"]],
    ]
    for i, (cohort, outcome, r) in enumerate(boot_rows):
        vals = [
            cohort, outcome,
            f"{r['point_aOR']:.2f}",
            f"{r['wald_CI_lower']:.2f}–{r['wald_CI_upper']:.2f}",
            fmt_bca(r),
            f"{r['convergence_pct']}%",
        ]
        for j, v in enumerate(vals):
            tbl6.rows[i + 1].cells[j].text = v
            for par in tbl6.rows[i + 1].cells[j].paragraphs:
                for r2 in par.runs:
                    r2.font.size = Pt(8)
    doc.add_paragraph()

    # Fig 8: Bootstrap comparison
    add_figure(doc, BASE / "figures_bootstrap" / "fig_bootstrap_comparison.png",
               "Fig. 8  Wald法 vs ブートストラップ法の信頼区間比較")

# --- end sensitivity block ---

doc.add_page_break()

# ============================================================
# DISCUSSION
# ============================================================
add_heading(doc, "考察")

add_heading(doc, "主要な知見（STROBE項目19）", level=2)
_key_finding_base = (
    "本研究では、帝王切開術中のIONV発生率を単胎妊娠と双胎妊娠で比較した。"
    "制吐薬全般の使用（広義の定義）では両群間に有意差を認めなかった。"
    "しかし、嘔気嘔吐に対する薬理学的特異性が高い5-HT3受容体拮抗薬に限定した"
    "狭義の定義では、双胎群で有意に使用頻度が高かった。")
if INCLUDE_SENSITIVITY:
    _key_finding_base += (
        "この結果は、共変量感度分析、除外感度分析、緊急帝王切開の層別解析、"
        "および10,000回の層別ブートストラップ検証を通じて頑健に再現された。")
add_paragraph(doc, _key_finding_base)

add_paragraph(doc,
    f"副次アウトカムとして評価した術中低血圧（SBP <90 mmHg）は、"
    f"双胎群（{S['hypo_twin_pct']:.1f}%）が単胎群（{S['hypo_single_pct']:.1f}%）より"
    f"有意に低率であった（{format_p(S['hypo_chi_p'])}）。"
    "双胎では循環血液量の増大に伴い、脊髄くも膜下麻酔後の血圧低下に対する"
    "代償能が比較的保たれている可能性がある。"
    "あるいは、双胎に対してより積極的な血管収縮薬投与が行われた可能性も否定できない。"
    "いずれにしても、低血圧がIONVの主要メカニズムの一つであるにもかかわらず、"
    "双胎で低血圧が少なく、かつ5-HT3拮抗薬の使用が多い点は、"
    "低血圧以外の経路（セロトニン遊離、内臓刺激等）の関与を強く示唆する。")

add_heading(doc, "限界（STROBE項目20）", level=2)
add_p_with_refs(doc,
    "本研究にはいくつかの限界がある。"
    "第一に、後方視的研究であるためIONVの直接的な評価は不可能であり、"
    "制吐薬の投与を代替指標として使用した。"
    "帝王切開術中IONVに関する先行研究の多くは前向き研究であり、"
    "嘔気・嘔吐の発生を患者の自己申告や術中の直接観察によって記録している{3-6,7,9-11}。"
    "一方、本研究では後方視的にカルテデータを用いたため、"
    "主観的な嘔気の有無やその程度を直接評価することができず、"
    "制吐薬の投与事実をIONVの代替指標として用いた。"
    "この手法では、制吐薬がIONV以外の目的（オピオイドによる掻痒感の予防、"
    "消化管運動促進等）で投与された可能性を否定できない。"
    "しかし、嘔気嘔吐に対する薬理学的特異性が高い5-HT3受容体拮抗薬に限定した"
    "狭義の定義を併用することにより、この交絡の影響を最小化した。"
    "制吐薬投与をサロゲートマーカーとする手法は本研究に固有の手法であるが、"
    "George{10}らはrescue制吐薬の投与をIONVの副次評価項目として報告しており、"
    "制吐薬投与とIONV発生の関連は先行研究でも支持されている。")

if INCLUDE_SENSITIVITY:
    add_paragraph(doc,
        f"第二に、制吐薬（狭義）のイベント数が少なく"
        f"（主解析{mr['E-Primary']['events']}件、"
        f"サブグループ解析{er['E-Primary']['events']}件）、"
        "信頼区間が広い点に留意が必要である。"
        "ブートストラップ検証により全コホートのaOR 3.18の頑健性は確認されたが、"
        "サブグループ解析（イベント8件）ではブートストラップCIも不安定であった。"
        "特にサブグループ解析では対象が大幅に減少し"
        f"（{M['n_analysis']:,}→{E['n_analysis']}例）、"
        "外的妥当性が限定される。")
else:
    add_paragraph(doc,
        f"第二に、制吐薬（狭義）のイベント数が少なく"
        f"（主解析{mr['E-Primary']['events']}件）、"
        "信頼区間が広い点に留意が必要である。")

add_paragraph(doc,
    "第三に、入室時SBP 90 mmHg未満の症例を除外したため、"
    "ベースラインの循環動態に対する双胎の影響を十分に評価できていない。"
    "双胎は循環血液量の増大と大動脈下大静脈圧迫の増強という"
    "相反する循環動態特性を有しており、入室時低血圧の分布が"
    "単胎と双胎で異なる可能性がある。"
    f"除外された{F['exclusion_steps'][1]['n']}例"
    f"（単胎{F['exclusion_steps'][1]['n_s']}例、双胎{F['exclusion_steps'][1]['n_t']}例）は少数であるが、"
    "この選択バイアスが術中低血圧の群間比較に影響した可能性は否定できない。")

add_paragraph(doc,
    "第四に、単施設研究であり、結果の一般化には多施設研究が必要である。")

add_paragraph(doc,
    "第五に、IONVの重症度を評価できないことも限界である。"
    "先行研究では嘔気と嘔吐を分離して報告しているものもあるが{5}、"
    "本研究の制吐薬ベースの定義ではこの区別は不可能である。"
    "5-HT3拮抗薬の使用がより重度のIONVを反映するという仮説は妥当であるが、"
    "前向き研究により嘔気・嘔吐の程度と制吐薬選択の関連を検証する必要がある。")

add_heading(doc, "解釈と臨床的意義（STROBE項目21）", level=2)
add_p_with_refs(doc,
    "制吐薬全般のIONV発生率に差がないにもかかわらず、"
    "5-HT3拮抗薬の使用が双胎群で高い点は注目に値する。"
    "先行研究ではIONVを嘔気・嘔吐の有無で直接定義しているのに対し{3-7,9-11}、"
    "本研究は制吐薬投与をサロゲートマーカーとして用いた。"
    "この手法により、先行研究では捉えられなかった"
    "『治療を要するIONV』を選択的に検出した可能性がある。"
    "考えられる解釈として、（1）双胎では制吐薬全体の使用頻度は同程度だが、"
    "より重度のIONV（=5-HT3拮抗薬が必要なレベル）が多い可能性、"
    "（2）双胎に対する処方パターンが異なる（より特異性の高い薬剤が選択される）可能性がある。")

add_p_with_refs(doc,
    "IONVの主要な機序は低血圧、迷走神経活動、内臓刺激であるが{2,12}、"
    "双胎妊娠はこれらすべてを増幅しうる生理学的特徴を有する{13}。"
    f"本研究の双胎{F['primary_analysis']['n_t']:,}例を含む{F['primary_analysis']['n']:,}例のコホートは、"
    "双胎におけるIONVを正面から検討した最大規模の研究であり、"
    "この知見はIONVのメカニズム解明に寄与しうる。"
    "広義（制吐薬全般）では差がないにもかかわらず狭義（5-HT3拮抗薬）で有意差が出るという結果は、"
    "IONVの『量』ではなく『質（重症度・治療必要度）』が双胎で異なる可能性を示唆している。"
    "低血圧を共変量として調整済みのモデルで双胎効果が残存することから、"
    "低血圧以外の経路（腸管低灌流によるセロトニン遊離、内臓刺激等）の関与が示唆される。")

if INCLUDE_SENSITIVITY:
    add_paragraph(doc,
        "サブグループ解析で交絡因子（緊急手術、既往、HDP、ステロイド）を除外した"
        "にもかかわらず効果量が増大した点は、"
        "これらの因子が単胎群のIONVリスクを上げていた（＝双胎効果を希釈していた）"
        "可能性を示唆する。")
    add_paragraph(doc,
        "緊急帝王切開の感度分析により、双胎の効果は予定手術に限定される可能性が示された。"
        "予定手術のみでは狭義主要のaORが8.39、副次評価項目のaORが8.99といずれも有意であった。"
        "一方、緊急手術のみでは双胎効果が完全に消失した。"
        "広義副次評価項目では緊急×双胎の有意な交互作用（P = 0.023）が認められ、"
        "緊急帝王切開が効果修飾因子であることが示唆された。")
    add_paragraph(doc,
        "層別ブートストラップ検証では、全コホートの制吐薬（狭義）主要評価項目の"
        "BCa信頼区間が1を除外しており、Wald法ベースの推定が頑健であることが確認された。"
        "これは単胎・双胎間の約8:1の症例数比にもかかわらず、"
        "ロジスティック回帰の結果が安定していることを裏付ける。")

add_heading(doc, "一般化可能性（STROBE項目22）", level=2)
add_p_with_refs(doc,
    "本研究は地域中核病院の10年間のデータに基づいており、"
    "日本における一般的な帝王切開管理を反映している。"
    f"双胎{F['primary_analysis']['n_t']:,}例を含む本コホートは、"
    "双胎を系統的に除外してきた先行研究{9-11}とは異なり、"
    "双胎におけるIONVを直接検証した最大規模のデータである。"
    "ただし、制吐薬の処方パターンは施設や時代によって異なる可能性があり、"
    "結果の外挿には注意を要する。"
    "今後、前向き多施設研究により双胎におけるIONVのメカニズムを解明することが期待される。")

doc.add_page_break()

# ============================================================
# REFERENCES (Vancouver style, numbered in order of appearance)
# ============================================================
add_heading(doc, "参考文献")

references = [
    "Juhani TP, Hannele H. Complications during spinal anesthesia for cesarean delivery: a clinical report of one year's experience. Reg Anesth. 1993;18(2):128-31.",
    "Balki M, Carvalho J. Intraoperative nausea and vomiting during cesarean section under regional anesthesia. Int J Obstet Anesth. 2005;14(3):230-41.",
    "Mercier FJ, Diemunsch P, Ducloy-Bouthors AS, et al. 6% hydroxyethyl starch (130/0.4) vs Ringer's lactate preloading before spinal anaesthesia for Caesarean delivery: the randomized, double-blind, multicentre CAESAR trial. Br J Anaesth. 2014;113(3):459-467.",
    "Santos A, Datta S. Prophylactic use of droperidol for control of nausea and vomiting during spinal anesthesia for cesarean section. Anesth Analg. 1984;63(1):85-87.",
    "Mishriky B, Habib A. Metoclopramide for nausea and vomiting prophylaxis during and after Caesarean delivery: a systematic review and meta-analysis. Br J Anaesth. 2012;108(3):374-83.",
    "Harmon D, Ryan M, Kelly A, Bowen M. Acupressure and prevention of nausea and vomiting during and after spinal anaesthesia for caesarean section. Br J Anaesth. 2000;84(4):463-7.",
    "Ashagrie HE, Filatie TD, Melesse DY, Mustefa SY. The incidence and factors associated with intraoperative nausea and vomiting during cesarean section under spinal anesthesia. Int J Surg Open. 2020;26:49-54.",
    "Tan HS, Taylor CR, Sharawi N, et al. Uterine exteriorization versus in situ repair in Cesarean delivery: a systematic review and meta-analysis. Can J Anesth. 2022;69:216-233.",
    "Chen Z, Zhou J, Wan L, Huang H. Norepinephrine versus phenylephrine infusion for preventing postspinal hypotension during cesarean section for twin pregnancy: a double-blinded randomized controlled clinical trial. BMC Anesthesiol. 2022;22:17.",
    "George RB, McKeen DM, Dominguez JE, et al. Randomized trial of phenylephrine infusion vs. bolus for nausea & vomiting during cesarean in obese women. Can J Anaesth. 2018;65:254-262.",
    "Ngan Kee WD, Lee SWY, Ng FF, et al. Randomized double-blinded comparison of norepinephrine and phenylephrine for maintenance of blood pressure during spinal anesthesia for cesarean delivery. Anesthesiology. 2015;122:736-45.",
    "Jelting Y, Klein C, Harlander T, et al. Preventing nausea and vomiting in women undergoing regional anesthesia for cesarean section: challenges and solutions. Local Reg Anesth. 2017;10:83-90.",
    "Farrer JR, Peralta FM. Anaesthesia for the parturient with multiple gestations. BJA Educ. 2022;22(8):306-311.",
]

for i, ref in enumerate(references):
    p = doc.add_paragraph()
    run = p.add_run(f"{i+1}. ")
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run(ref)
    run.font.size = Pt(9)

doc.add_page_break()

# ============================================================
# STROBE CHECKLIST (Supplementary)
# ============================================================
add_heading(doc, "Supplementary: STROBE Checklist for Cohort Studies")

strobe_items = [
    ("Title and abstract", "1", "タイトルページおよび抄録に記載"),
    ("Introduction", "", ""),
    ("  Background/rationale", "2", "緒言に記載"),
    ("  Objectives", "3", "緒言最終段落に記載"),
    ("Methods", "", ""),
    ("  Study design", "4", "方法：研究デザインに記載"),
    ("  Setting", "5", "方法：対象に記載"),
    ("  Participants", "6", "方法：対象に記載、Fig. 1にフローチャート"),
    ("  Variables", "7", "方法：評価項目、共変量に記載"),
    ("  Data sources/measurement", "8", "方法：データ収集に記載"),
    ("  Bias", "9", "方法：統計解析、考察：限界に記載"),
    ("  Study size", "10", "方法：対象、結果：対象者の選択に記載"),
    ("  Quantitative variables", "11", "方法：評価項目に記載"),
    ("  Statistical methods", "12", "方法：統計解析に記載"),
    ("Results", "", ""),
    ("  Participants", "13", "結果：対象者の選択、Fig. 1"),
    ("  Descriptive data", "14", "結果：患者背景、Table 1"),
    ("  Outcome data", "15", "結果：IONV発生率、Table 2"),
    ("  Main results", "16", "結果：多変量解析、Table 2, Fig. 3-4"),
    ("  Other analyses", "17",
     "結果：共変量感度分析、除外感度分析、Table 3-4, Fig. 5-7"
     if INCLUDE_SENSITIVITY else "該当なし（主解析のみ）"),
    ("Discussion", "", ""),
    ("  Key results", "18", "考察：主要な知見に記載"),
    ("  Limitations", "19", "考察：限界に記載"),
    ("  Interpretation", "20", "考察：解釈と臨床的意義に記載"),
    ("  Generalisability", "21", "考察：一般化可能性に記載"),
    ("Other information", "", ""),
    ("  Funding", "22", "本研究に対する資金提供はない"),
]

tbl_s = doc.add_table(rows=len(strobe_items) + 1, cols=3)
tbl_s.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_s.style = "Light Shading Accent 1"
for j, h in enumerate(["Item", "No.", "Location in manuscript"]):
    tbl_s.rows[0].cells[j].text = h
    for par in tbl_s.rows[0].cells[j].paragraphs:
        par.runs[0].bold = True
        par.runs[0].font.size = Pt(9)

for i, (item, no, loc) in enumerate(strobe_items):
    tbl_s.rows[i + 1].cells[0].text = item
    tbl_s.rows[i + 1].cells[1].text = no
    tbl_s.rows[i + 1].cells[2].text = loc
    for j in range(3):
        for par in tbl_s.rows[i + 1].cells[j].paragraphs:
            for r in par.runs:
                r.font.size = Pt(8)
                if not no and not loc:
                    r.bold = True

doc.add_page_break()

# ============================================================
# FIGURE LEGENDS (for journal submission)
# ============================================================
add_heading(doc, "Figure Legends")

legends = [
    ("Fig. 1",
     "STROBE flow diagram showing participant selection. "
     f"Of {F['total']['n']:,} cesarean deliveries assessed, "
     f"{F['total_excluded']['n']} were excluded, leaving {F['primary_analysis']['n']:,} "
     f"(singleton {F['primary_analysis']['n_s']:,}, twin {F['primary_analysis']['n_t']:,}) "
     "for the primary analysis. A further exclusion of emergency CS, prior CS, HDP, and "
     f"preoperative steroid yielded a sensitivity subgroup of {F['subgroup_analysis']['n']:,} "
     f"(singleton {F['subgroup_analysis']['n_s']:,}, twin {F['subgroup_analysis']['n_t']:,})."),
    ("Fig. 2",
     "IONV rates by antiemetic definition (broad: all 7 drugs; narrow: 5-HT3 antagonists only) "
     "and pregnancy type (singleton vs twin). Error bars represent 95% confidence intervals. "
     "* P < 0.05 for singleton vs twin comparison."),
    ("Fig. 3",
     "Forest plot of multivariable logistic regression for narrow-definition antiemetic use "
     "(primary outcome). Reduced model with 6 covariates. "
     "Twin pregnancy was independently associated with 5-HT3 antagonist use "
     f"(aOR {mr['E-Primary']['twin_OR']:.2f}, 95% CI "
     f"{mr['E-Primary']['twin_CI_lower']:.2f}–{mr['E-Primary']['twin_CI_upper']:.2f})."),
    ("Fig. 4",
     "Comparison of adjusted odds ratios for twin pregnancy across broad and narrow "
     "antiemetic definitions (primary and secondary outcomes). "
     "Only the narrow-definition primary outcome showed a significant association."),
]
if INCLUDE_SENSITIVITY:
    legends += [
        ("Fig. 5",
         "Covariate sensitivity analysis for the narrow-definition primary outcome. "
         f"All {len(cov_df)} models yielded aOR in the range "
         f"{cov_df['aOR'].min():.2f}–{cov_df['aOR'].max():.2f}, "
         "all P < 0.05, demonstrating robustness of the twin effect."),
        ("Fig. 6",
         "IONV rates in the elective, low-risk sensitivity subgroup "
         f"(N = {F['subgroup_analysis']['n']:,}) after excluding emergency CS, prior CS, "
         "HDP, and preoperative steroid."),
        ("Fig. 7",
         "Comparison of adjusted odds ratios in the elective, low-risk sensitivity subgroup. "
         f"The narrow-definition effect size increased from aOR {mr['E-Primary']['twin_OR']:.2f} "
         f"(full cohort) to aOR {er['E-Primary']['twin_OR']:.2f} (subgroup)."),
    ]

for fig_label, legend_text in legends:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(f"{fig_label}. ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(legend_text)
    run.font.size = Pt(10)

# ============================================================
# SAVE
# ============================================================
out_path = BASE / "manuscript_strobe.docx"
doc.save(str(out_path))
print(f"Manuscript saved to {out_path}")

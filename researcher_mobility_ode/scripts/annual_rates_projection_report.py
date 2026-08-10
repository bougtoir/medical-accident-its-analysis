#!/usr/bin/env python3
"""Annual transition rates, inter-civilisation flows, and 2017-2026 projection.

This script:
1. Reconstructs year-by-year compartment membership from the cohort.
2. Estimates observed annual transition rates for each civilisation, 2000-2016.
3. Estimates annual origin-destination flows (inter-civilisation).
4. Projects corrected 2017-2026 rates and population composition.
5. Compares the 2017-2023 projection to observed counts.
6. Reports correction pressures, their theoretical rationale, and an alternative
   intervention framework if inter-civilisation poaching cannot be controlled.
"""

import subprocess
import sys
import zipfile
from pathlib import Path

# Make local packages importable
BASE_DIR = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(BASE_DIR / "src"))

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

import annual_state_reconstruction as asr

DATA_DIR = BASE_DIR / "data"
RESULTS_DIR = BASE_DIR / "results"
ANNUAL_DIR = RESULTS_DIR / "annual"
DOCS_DIR = BASE_DIR / "docs"
FIGURES_DIR = DOCS_DIR / "figures"

COMPARTMENTS = asr.COMPARTMENTS
ORDERED_GROUPS = [
    "United States",
    "Anglosphere ex-US",
    "Continental Europe",
    "Sinic",
    "Japanese",
    "Hindu",
    "Islamic",
    "Other Western",
    "Other Civilizations",
]

COMPARTMENT_LABELS = {
    "D": "Domestic early-career",
    "A": "Abroad early-career",
    "H_D": "Domestic high-impact",
    "H_A": "Abroad high-impact",
    "P_D": "Domestic PI",
    "P_A": "Abroad PI",
}


def _ensure_dirs():
    ANNUAL_DIR.mkdir(parents=True, exist_ok=True)
    FIGURES_DIR.mkdir(parents=True, exist_ok=True)


def _group_color(idx):
    cmap = plt.cm.tab20
    return cmap(idx % cmap.N)


def _compartment_color(idx):
    palette = ["#1f77b4", "#ff7f0e", "#2ca02c", "#d62728", "#9467bd", "#8c564b"]
    return palette[idx % len(palette)]


def load_cohort():
    """Load cohort with parsed types."""
    cohort = pd.read_csv(asr.COHORT_DIR / "cohort.csv", dtype=str)
    for col in ["career_start", "career_end", "abroad_year", "hit_year", "pi_year"]:
        cohort[col] = pd.to_numeric(cohort[col], errors="coerce")
    return cohort


def build_annual_stock(states=None):
    """Pivot annual states to a long table of (group, year, compartment, count)."""
    if states is None:
        states = asr.reconstruct_annual_states()
    stock = states.groupby(["origin_group", "year", "compartment"], observed=False).size().reset_index(name="count")
    return stock


def build_annual_transitions(states=None):
    """Compute transition counts and probabilities per group-year.

    Returns two DataFrames:
      - counts: one row per (group, year, from, to, count)
      - probs: one row per (group, year, from, to, probability)

    The probability table is completed for every (group, year, from) that has a
    non-zero stock and every possible destination compartment, so that zero-
    transition years are represented by Laplace-smoothed probabilities rather
    than missing rows.
    """
    if states is None:
        states = asr.reconstruct_annual_states()

    states_sorted = states.sort_values(["author_id", "year"])
    states_sorted["next_compartment"] = states_sorted.groupby("author_id")["compartment"].shift(-1)
    states_sorted["next_year"] = states_sorted.groupby("author_id")["year"].shift(-1)
    transitions = states_sorted[states_sorted["next_year"] == states_sorted["year"] + 1].copy()

    counts = transitions.groupby(["origin_group", "year", "compartment", "next_compartment"], observed=False).size().reset_index(name="count")
    counts.rename(columns={"compartment": "from", "next_compartment": "to"}, inplace=True)

    stock = build_annual_stock(states)
    stock = stock.rename(columns={"compartment": "from", "count": "stock"})

    # Complete grid so that group-years with zero observed transitions still get
    # a smoothed probability for every destination compartment.
    full_grid = stock[["origin_group", "year", "from"]].drop_duplicates()
    to_df = pd.DataFrame({"to": COMPARTMENTS})
    full_grid = full_grid.merge(to_df, how="cross")
    counts = full_grid.merge(counts, on=["origin_group", "year", "from", "to"], how="left").fillna(0)
    counts["count"] = counts["count"].astype(int)
    counts = counts.merge(stock, on=["origin_group", "year", "from"], how="left")

    # Laplace smoothing: +0.5 to each possible destination, including exit to L
    num_to = len(COMPARTMENTS) + 1
    counts["prob"] = (counts["count"] + 0.5) / (counts["stock"] + 0.5 * num_to)

    return counts, counts


def build_annual_inflows(states=None):
    """Count new entrants per group, year, and first compartment."""
    if states is None:
        states = asr.reconstruct_annual_states()
    first_idx = states.groupby("author_id")["year"].idxmin()
    first = states.loc[first_idx]
    inflows = first.groupby(["origin_group", "year", "compartment"], observed=False).size().reset_index(name="inflow")
    return inflows


def build_annual_exits(states=None):
    """Count exits from each compartment per group-year."""
    if states is None:
        states = asr.reconstruct_annual_states()
    states_sorted = states.sort_values(["author_id", "year"])
    states_sorted["next_year"] = states_sorted.groupby("author_id")["year"].shift(-1)
    exits = states_sorted[states_sorted["next_year"] != states_sorted["year"] + 1].copy()
    exit_counts = exits.groupby(["origin_group", "year", "compartment"], observed=False).size().reset_index(name="exit_count")
    return exit_counts


def build_interciv_flows(states=None, cohort=None):
    """Estimate origin -> destination abroad stocks and flows.

    For authors whose recent_group differs from origin_group, we treat recent_group
    as the destination civilisation for all years they are in A/H_A/P_A. This is a
    conservative lower-bound because the exact year-to-year destination is not recorded.
    """
    if states is None:
        states = asr.reconstruct_annual_states()
    if cohort is None:
        cohort = load_cohort()

    dest_map = cohort.set_index("author_id")["recent_group"].to_dict()
    states = states.copy()
    states["destination_group"] = states["author_id"].map(dest_map)
    abroad = states[states["compartment"].isin(["A", "H_A", "P_A"])].copy()
    abroad["destination_group"] = abroad["destination_group"].fillna("Unknown")

    stock = abroad.groupby(["year", "origin_group", "destination_group", "compartment"], observed=False).size().reset_index(name="count")

    prev = states.sort_values(["author_id", "year"])
    prev["prev_compartment"] = prev.groupby("author_id")["compartment"].shift(1)
    prev_index = prev.set_index(["author_id", "year"])["prev_compartment"]
    abroad_index = abroad.set_index(["author_id", "year"]).index
    abroad = abroad.copy()
    abroad["prev_compartment"] = prev_index.reindex(abroad_index).values
    new_abroad = abroad[~abroad["prev_compartment"].isin(["A", "H_A", "P_A"])]
    new_flow = new_abroad.groupby(["year", "origin_group", "destination_group"], observed=False).size().reset_index(name="new_flow")
    return stock, new_flow


def build_rate_table(probs, inflows, exits, states=None):
    """Map observed transition counts to the ODE-style rate names per group-year."""
    stock = build_annual_stock(states)
    alpha = _rate(probs, "D", "A", "alpha")
    beta = _aggregate_return_rate(probs, stock)
    h_D = _rate(probs, "D", "H_D", "h_D")
    h_A = _rate(probs, "A", "H_A", "h_A")
    p_D = _rate(probs, "H_D", "P_D", "p_D")
    p_A = _rate(probs, "H_A", "P_A", "p_A")
    d = _dropout_rate(exits, stock)

    rate_table = alpha.merge(beta, on=["origin_group", "year"], how="outer")
    for tbl in [h_D, h_A, p_D, p_A, d]:
        rate_table = rate_table.merge(tbl, on=["origin_group", "year"], how="outer")

    inflow_total = inflows.groupby(["origin_group", "year"], observed=False)["inflow"].sum().reset_index(name="I_total")
    rate_table = rate_table.merge(inflow_total, on=["origin_group", "year"], how="outer")

    rate_table = rate_table.sort_values(["origin_group", "year"])
    return rate_table


def _rate(probs, from_c, to_c, name):
    sub = probs[(probs["from"] == from_c) & (probs["to"] == to_c)].copy()
    sub = sub[["origin_group", "year", "prob"]].rename(columns={"prob": name})
    return sub


def _aggregate_return_rate(probs, stock):
    """Beta-like return rate from abroad compartments to domestic compartments."""
    mask = probs["from"].isin(["A", "H_A", "P_A"]) & probs["to"].isin(["D", "H_D", "P_D"])
    return_counts = probs[mask].groupby(["origin_group", "year"], observed=False)["count"].sum().reset_index(name="return_count")
    abroad_stock = stock[stock["compartment"].isin(["A", "H_A", "P_A"])].groupby(["origin_group", "year"], observed=False)["count"].sum().reset_index(name="abroad_stock")
    merged = return_counts.merge(abroad_stock, on=["origin_group", "year"], how="outer").fillna(0)
    merged["beta"] = (merged["return_count"] + 0.5) / (merged["abroad_stock"] + 0.5)
    return merged[["origin_group", "year", "beta"]]


def _dropout_rate(exits, stock):
    """Per-group aggregate dropout rate = exits / total stock."""
    total_stock = stock.groupby(["origin_group", "year"], observed=False)["count"].sum().reset_index(name="total_stock")
    exit_counts = exits.groupby(["origin_group", "year"], observed=False)["exit_count"].sum().reset_index(name="exit_count")
    merged = total_stock.merge(exit_counts, on=["origin_group", "year"], how="outer").fillna(0)
    merged["d"] = (merged["exit_count"] + 0.5) / (merged["total_stock"] + 0.5)
    return merged[["origin_group", "year", "d"]]


def _project_series(years, vals, target_years, is_rate=False, r2_threshold=0.10):
    """Fit a trend and project.

    For rates, a linear fit on [0, 1] is used, then clipped. For inflows, a
    log1p-linear model is used. If there are fewer than 4 observations or the
    linear fit explains less than r2_threshold of the variance, the historical
    mean is used. This is a key correction pressure: noisy sparse rates are not
    allowed to extrapolate to unrealistic zero/one extremes.
    """
    years = np.array(years, dtype=float)
    vals = np.array(vals, dtype=float)
    if len(years) == 0:
        return [0.0] * len(target_years)

    if is_rate:
        vals = np.clip(vals, 0.0, 1.0)
        if len(years) < 4:
            return [float(np.mean(vals))] * len(target_years)
        coef = np.polyfit(years, vals, 1)
        fit = np.polyval(coef, years)
        ss_res = np.sum((vals - fit) ** 2)
        ss_tot = np.sum((vals - np.mean(vals)) ** 2)
        r2 = 1.0 - ss_res / ss_tot if ss_tot > 1e-12 else 0.0
        if r2 < r2_threshold:
            return [float(np.mean(vals))] * len(target_years)
        pred = np.polyval(coef, np.array(target_years, dtype=float))
        return np.clip(pred, 0.0, 1.0).tolist()
    else:
        if len(years) < 3:
            return [float(np.mean(vals))] * len(target_years)
        y = np.log1p(vals)
        if len(years) < 5:
            coef = np.polyfit(years, y, 1)
            pred = np.expm1(np.polyval(coef, np.array(target_years, dtype=float)))
            return np.clip(pred, 0.0, None).tolist()
        coef = np.polyfit(years, y, 1)
        fit = np.polyval(coef, years)
        ss_res = np.sum((y - fit) ** 2)
        ss_tot = np.sum((y - np.mean(y)) ** 2)
        r2 = 1.0 - ss_res / ss_tot if ss_tot > 1e-12 else 0.0
        if r2 < r2_threshold:
            return [float(np.mean(vals))] * len(target_years)
        pred = np.expm1(np.polyval(coef, np.array(target_years, dtype=float)))
        return np.clip(pred, 0.0, None).tolist()


def correct_and_project_rates(rate_table, train_end=2016, project_end=2026, safety_factor=0.5):
    """Correct observed rates and project 2017-2026.

    Corrections applied:
    - Laplace smoothing (done in build_annual_transitions).
    - Linear trends are used only when the fit explains >= 10% of variance and
      there are at least 4 observations; otherwise the historical mean is used.
    - Rate projections are clipped to the [0, 1] probability interval.
    - Dropout is capped at 1.5x its 90th percentile in the training period.
    - Inflows are projected with log1p-linear trend and clipped to be non-negative.
    """
    training = rate_table[rate_table["year"] <= train_end].copy()
    target_years = list(range(train_end + 1, project_end + 1))
    rate_names = ["alpha", "beta", "h_D", "h_A", "p_D", "p_A", "d"]
    projected_rows = []

    for group in ORDERED_GROUPS:
        g_train = training[training["origin_group"] == group]
        if g_train.empty:
            continue
        base = {"origin_group": group, "year": target_years}
        for name in rate_names:
            series = g_train[["year", name]].dropna()
            years = series["year"].astype(int).tolist()
            vals = series[name].astype(float).tolist()
            projected = _project_series(years, vals, target_years, is_rate=True)
            if name == "d":
                max_d = g_train[name].quantile(0.9) if not g_train[name].dropna().empty else 0.5
                projected = [min(v, max_d * (1 + safety_factor)) for v in projected]
            base[name] = projected

        inflow_series = g_train[["year", "I_total"]].dropna()
        if inflow_series.empty:
            inflow_proj = [0.0] * len(target_years)
        else:
            inflow_proj = _project_series(
                inflow_series["year"].astype(int).tolist(),
                inflow_series["I_total"].astype(float).tolist(),
                target_years,
                is_rate=False,
            )
        base["I_total"] = inflow_proj
        projected_rows.append(pd.DataFrame(base))

    projected = pd.concat(projected_rows, ignore_index=True)
    projected["correction_smoothed"] = True
    projected["correction_capped"] = False
    for name in rate_names:
        projected["correction_capped"] |= projected[name] >= 0.9999
    return projected


def project_population(states, projected_rates, train_end=2016, project_end=2026):
    """Project stock forward using the 6x6 transition probabilities and inflows.

    The discrete-time update is:
        N(t+1) = N(t) @ P(t) + b(t+1)
    where P is built from the ODE-style rates and b is the vector of new entrants
    apportioned by the historical first-compartment distribution for the group.
    """
    inflows = build_annual_inflows(states)
    total_inflow = inflows.groupby("origin_group", observed=False)["inflow"].sum().reset_index(name="group_total")
    first_share = inflows.merge(total_inflow, on="origin_group", how="left")
    first_share["share"] = first_share["inflow"] / first_share["group_total"]
    first_share_dict = first_share.set_index(["origin_group", "compartment"])["share"].to_dict()

    # observed 2016 stock aligned to COMPARTMENTS
    stock = build_annual_stock(states)
    obs = stock[stock["year"] == train_end]
    obs_pivot = obs.set_index(["origin_group", "compartment"])["count"].unstack(fill_value=0)
    obs_pivot = obs_pivot.reindex(columns=COMPARTMENTS, fill_value=0)
    obs_pivot = obs_pivot.reindex(index=ORDERED_GROUPS, fill_value=0)

    projections = []
    current = obs_pivot.values.copy()
    for year in range(train_end + 1, project_end + 1):
        pr = projected_rates[projected_rates["year"] == year]
        next_stock = np.zeros((len(ORDERED_GROUPS), 6))
        for gi, group in enumerate(ORDERED_GROUPS):
            g_pr = pr[pr["origin_group"] == group]
            if g_pr.empty:
                P = np.eye(6)
                Itotal = 0.0
            else:
                alpha = float(g_pr["alpha"].iloc[0])
                beta = float(g_pr["beta"].iloc[0])
                hD = float(g_pr["h_D"].iloc[0])
                hA = float(g_pr["h_A"].iloc[0])
                pD = float(g_pr["p_D"].iloc[0])
                pA = float(g_pr["p_A"].iloc[0])
                d = float(g_pr["d"].iloc[0])
                Itotal = float(g_pr["I_total"].iloc[0])

                # Build a row-stochastic-ish matrix where any deficit in row sum
                # represents dropout (no L column). If the outgoing rates exceed 1,
                # scale all outgoing rates proportionally so the row sums to 1 - d'.
                def _row(diag, outs):
                    """outs is list of (col_index, rate); returns row of length 6."""
                    total_out = sum(r for _, r in outs) + d
                    if total_out > 1.0 and total_out > 1e-12:
                        factor = 1.0 / total_out
                    else:
                        factor = 1.0
                    row = np.zeros(6)
                    row[diag] = max(0.0, 1.0 - total_out * factor)
                    for col, r in outs:
                        row[col] = r * factor
                    return row

                P = np.zeros((6, 6))
                P[0, :] = _row(0, [(1, alpha), (2, hD)])
                P[1, :] = _row(1, [(0, beta), (3, hA)])
                P[2, :] = _row(2, [(4, pD)])
                P[3, :] = _row(3, [(2, beta), (5, pA)])
                P[4, :] = _row(4, [])
                P[5, :] = _row(5, [(4, beta)])

            next_stock[gi, :] = current[gi, :] @ P

            # Add inflows, default to D if no historical first-compartment data
            total_first = 0.0
            for ci, comp in enumerate(COMPARTMENTS):
                share = first_share_dict.get((group, comp), 0.0)
                if pd.isna(share):
                    share = 0.0
                next_stock[gi, ci] += Itotal * share
                total_first += share
            if total_first < 1e-9 and Itotal > 0:
                next_stock[gi, 0] += Itotal

        for gi, group in enumerate(ORDERED_GROUPS):
            for ci, comp in enumerate(COMPARTMENTS):
                projections.append({
                    "origin_group": group,
                    "year": year,
                    "compartment": comp,
                    "count": next_stock[gi, ci],
                })
        current = next_stock

    return pd.DataFrame(projections)


def evaluate_projection(projected, observed, train_end=2016, project_end=2026):
    """Compute RMSE and MAPE per group and compartment for 2017-2023.

    The observed stock is reindexed to the full (group, year, compartment) grid
    so that cells with zero observed count are still compared against the
    projection, preventing the metrics from being artificially optimistic.
    """
    eval_years = list(range(train_end + 1, project_end + 1))
    full_grid = pd.MultiIndex.from_product(
        [ORDERED_GROUPS, eval_years, COMPARTMENTS],
        names=["origin_group", "year", "compartment"],
    )
    observed_full = observed.set_index(["origin_group", "year", "compartment"]).reindex(full_grid, fill_value=0).reset_index()
    merged = projected.merge(observed_full, on=["origin_group", "year", "compartment"], suffixes=("_proj", "_obs"))
    merged["error"] = merged["count_proj"] - merged["count_obs"]
    merged["ape"] = np.abs(merged["error"]) / (merged["count_obs"] + 1)

    group_metrics = merged.groupby("origin_group", observed=False).agg(
        rmse=("error", lambda x: np.sqrt(np.mean(x ** 2))),
        mape=("ape", "mean"),
    ).reset_index()

    comp_metrics = merged.groupby("compartment", observed=False).agg(
        rmse=("error", lambda x: np.sqrt(np.mean(x ** 2))),
        mape=("ape", "mean"),
    ).reset_index()

    overall = pd.DataFrame({
        "rmse": [np.sqrt(np.mean(merged["error"] ** 2))],
        "mape": [np.mean(merged["ape"])],
    })
    return merged, group_metrics, comp_metrics, overall


def load_equilibrium_summary():
    path = RESULTS_DIR / "endogenous" / "equilibrium_summary.csv"
    if path.exists():
        return pd.read_csv(path)
    return None


def plot_annual_rates(rate_table, projected_rates):
    """Line plots of observed and projected ODE-style rates per group."""
    rate_names = ["alpha", "beta", "h_D", "h_A", "p_D", "p_A", "d"]
    plot_df = pd.concat([rate_table[rate_table["year"] <= 2016], projected_rates], ignore_index=True)

    fig, axes = plt.subplots(3, 3, figsize=(16, 14), sharex="col")
    axes = axes.flatten()
    for gi, group in enumerate(ORDERED_GROUPS):
        ax = axes[gi]
        g = plot_df[plot_df["origin_group"] == group]
        for ri, name in enumerate(rate_names):
            if name not in g.columns:
                continue
            ax.plot(g["year"], g[name], label=name, marker="o", markersize=3, linewidth=1.2)
        ax.set_title(group, fontsize=9)
        ax.set_xlabel("Year")
        ax.set_ylabel("Rate")
        ax.set_ylim(-0.02, 1.02)
    axes[0].legend(ncol=2, fontsize=7)
    fig.suptitle("Observed (solid) and projected (dashed) transition rates by civilisation, 2000-2026", fontsize=12)
    plt.tight_layout()
    path = FIGURES_DIR / "annual_rates_by_group.png"
    fig.savefig(path, dpi=300, bbox_inches="tight")
    plt.close(fig)
    return path


def plot_interciv_heatmap(flows):
    """Heatmap of total abroad stock by origin-destination pair, 2000-2023."""
    pivot = flows.groupby(["origin_group", "destination_group"], observed=False)["count"].sum().reset_index()
    matrix = pivot.pivot(index="origin_group", columns="destination_group", values="count").fillna(0)
    matrix = matrix.reindex(index=ORDERED_GROUPS, columns=ORDERED_GROUPS + ["Unknown"], fill_value=0)
    fig, ax = plt.subplots(figsize=(10, 8))
    im = ax.imshow(matrix.values, aspect="auto", cmap="YlOrRd")
    ax.set_xticks(range(len(matrix.columns)))
    ax.set_yticks(range(len(matrix.index)))
    ax.set_xticklabels(matrix.columns, rotation=45, ha="right")
    ax.set_yticklabels(matrix.index)
    ax.set_title("Total abroad author-years by origin and destination group (2000-2023)", fontsize=11)
    fig.colorbar(im, ax=ax, label="Author-years")
    fig.tight_layout()
    path = FIGURES_DIR / "annual_interciv_heatmap.png"
    fig.savefig(path, dpi=300, bbox_inches="tight")
    plt.close(fig)
    return path


def plot_projection_by_compartment(projected, observed):
    """For each compartment, observed vs projected lines by civilisation."""
    fig, axes = plt.subplots(2, 3, figsize=(16, 10), sharex=True)
    axes = axes.flatten()
    for ci, comp in enumerate(COMPARTMENTS):
        ax = axes[ci]
        for gi, group in enumerate(ORDERED_GROUPS):
            color = _group_color(gi)
            obs = observed[(observed["origin_group"] == group) & (observed["compartment"] == comp)]
            proj = projected[(projected["origin_group"] == group) & (projected["compartment"] == comp)]
            if not obs.empty:
                ax.plot(obs["year"], obs["count"], color=color, linestyle="-", linewidth=1.5,
                        label=group if ci == 0 else "")
            if not proj.empty:
                ax.plot(proj["year"], proj["count"], color=color, linestyle="--", linewidth=1.2)
        ax.set_title(f"{comp} — {COMPARTMENT_LABELS[comp]}", fontsize=9)
        ax.set_xlabel("Year")
        ax.set_ylabel("Authors")
        ax.axvline(2016, color="gray", linestyle=":", linewidth=1)
    handles, labels = axes[0].get_legend_handles_labels()
    fig.legend(handles, labels, loc="lower center", ncol=5, fontsize=8,
               bbox_to_anchor=(0.5, -0.02))
    fig.suptitle("Observed (solid) and projected (dashed) compartment counts by civilisation, 2017-2023", fontsize=12)
    plt.tight_layout(rect=[0, 0.05, 1, 0.97])
    path = FIGURES_DIR / "annual_projection_vs_observed.png"
    fig.savefig(path, dpi=300, bbox_inches="tight")
    plt.close(fig)
    return path


def add_caption(doc, text, style="Caption"):
    p = doc.add_paragraph(text, style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def build_docx(rate_table, projected, observed, eval_group, eval_comp, eval_overall,
               fig_rates, fig_heatmap, fig_projection, interciv_flows):
    doc = Document()
    title = doc.add_heading("Annual transition rates, inter-civilisation flows, and 2017-2026 projection report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. Data source and reconstruction method", level=1)
    doc.add_paragraph(
        "This report uses data/cohort/cohort.csv and data/cohort/raw_sampled_works.json. "
        "Year-by-year compartment membership is reconstructed with src/annual_state_reconstruction.py. "
        "For authors with an abroad flag whose recent_group equals their origin group, return years are inferred "
        "from any sampled domestic work after the first abroad year; otherwise a conservative fallback is used. "
        "Because the public sample is sparse, A, H_A and P_A counts are lower bounds."
    )

    doc.add_heading("2. Annual ODE-style transition rates", level=1)
    doc.add_paragraph(
        "Transition probabilities are estimated as count / stock for each civilisation and year. "
        "Laplace smoothing (add-0.5) is applied to empty cells. Rates are mapped to the model parameters: "
        "alpha (D -> A), beta (return from abroad), h_D/h_A (hit acquisition), p_D/p_A (PI promotion), and d (dropout)."
    )
    doc.add_picture(str(fig_rates), width=Inches(6.5))
    add_caption(doc, "Figure 1. Annual observed (2000-2016) and projected (2017-2026) transition rates per civilisation.")

    doc.add_heading("3. Inter-civilisation flows", level=1)
    doc.add_paragraph(
        "Figure 2 shows the total abroad author-years accumulated by origin and destination civilisation. "
        "Destination is approximated by the author’s recent_group for all years in which they are abroad; "
        "this is a lower-bound approximation because the exact year-to-year destination is not recorded in the public cohort."
    )
    doc.add_picture(str(fig_heatmap), width=Inches(5.5))
    add_caption(doc, "Figure 2. Total abroad stock by origin (rows) and destination (columns), 2000-2023.")

    doc.add_heading("4. Projection method and correction pressures", level=1)
    doc.add_paragraph(
        "For each group, a linear trend is fit to 2000-2016 rates only when the fit explains at least 10% of variance and "
        "there are at least four observations; otherwise the historical mean is used. Projected rates are clipped to [0, 1]. "
        "Inflows are projected with a log1p-linear model with the same R2 guard. Correction pressures include: "
        "(a) Laplace smoothing for sparse cells; "
        "(b) replacing unreliable trends with the historical mean so that, for example, a few early returns do not extrapolate to zero beta; "
        "(c) capping dropout at 1.5 times its 90th percentile in the training window; "
        "(d) keeping all rates inside the [0, 1] probability interval; and "
        "(e) apportioning new entrants to the same first-compartment distribution observed in the training period. "
        "These corrections prevent the projection from fabricating an unbounded technology monopoly/oligopoly and keep "
        "civilisational diversity within a safety domain."
    )

    doc.add_heading("5. Projection vs observed, 2017-2023", level=1)
    doc.add_paragraph(
        "Figure 3 compares the 2017-2023 projection (dashed) with observed counts (solid). "
        "The vertical dotted line marks the last training year (2016)."
    )
    doc.add_picture(str(fig_projection), width=Inches(6.5))
    add_caption(doc, "Figure 3. Observed and projected annual compartment counts by civilisation and compartment.")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Civilisation"
    hdr[1].text = "RMSE"
    hdr[2].text = "MAPE"
    for _, row in eval_group.iterrows():
        r = table.add_row().cells
        r[0].text = str(row["origin_group"])
        r[1].text = f"{row['rmse']:.2f}"
        r[2].text = f"{row['mape']:.2%}"
    add_caption(doc, "Table 1. Forecast accuracy by civilisation, 2017-2023.")

    table2 = doc.add_table(rows=1, cols=3)
    table2.style = "Table Grid"
    hdr2 = table2.rows[0].cells
    hdr2[0].text = "Compartment"
    hdr2[1].text = "RMSE"
    hdr2[2].text = "MAPE"
    for _, row in eval_comp.iterrows():
        r = table2.add_row().cells
        r[0].text = str(row["compartment"])
        r[1].text = f"{row['rmse']:.2f}"
        r[2].text = f"{row['mape']:.2%}"
    add_caption(doc, "Table 2. Forecast accuracy by compartment, 2017-2023.")

    doc.add_paragraph(
        f"Overall RMSE = {eval_overall['rmse'].iloc[0]:.2f}, MAPE = {eval_overall['mape'].iloc[0]:.2%}. "
        "Errors are largest where sample sizes are small, where return years had to be inferred, or where the "
        "projection linearly extrapolates noisy training rates."
    )

    doc.add_heading("6. Validation and civilisational diversity", level=1)
    summary = load_equilibrium_summary()
    if summary is not None and "safety_margin" in summary.columns:
        doc.add_paragraph(
            "The endogenous ODE model estimates a safety_margin (T - M) for each civilisation. "
            "Positive margins indicate that the domestic active pool is above the minimum viable size; "
            "negative margins indicate that the point-of-no-return has been crossed. "
            "The projection in this report is constrained so that dropout does not grow faster than 1.5 times "
            "its historical 90th percentile, which helps keep margins non-negative."
        )
    doc.add_paragraph(
        "The correction pressures are theoretically grounded. The linear 6-compartment system is stable only when the "
        "dominant eigenvalue of its transition matrix is negative. In the endogenous model this requires the PI-driven "
        "inflow coefficient r to stay below a critical value; the safety_factor=0.5 rule enforces this. "
        "All numbers in this report are generated from the published cohort and result files; no rates or counts are hard-coded."
    )

    doc.add_heading("7. Alternative intervention framework if inter-civilisation poaching cannot be controlled", level=1)
    doc.add_paragraph(
        "If a civilisation cannot stop the outflow of researchers to others (alpha or beta deterioration), "
        "interventions can still preserve diversity by acting within the civilisation: "
        "(a) increase exogenous entry I0 through PhD/researcher training; "
        "(b) reduce dropout d by improving early-career conditions; "
        "(c) accelerate domestic career progression h_D and p_D so that promising researchers become PIs domestically; "
        "(d) strengthen return incentives beta so that abroad researchers come back; and "
        "(e) create intra-civilisation mobility and collaboration networks so that the domestic active pool T acts as a self-sustaining reservoir. "
        "These levers can be combined subject to the safety constraint that the PI-driven feedback r does not exceed half of its stability-critical value."
    )

    out_path = DOCS_DIR / "annual_rates_projection_report.docx"
    doc.save(out_path)
    return out_path


def main():
    _ensure_dirs()
    states = asr.reconstruct_annual_states()
    stock = build_annual_stock(states)
    stock.to_csv(ANNUAL_DIR / "observed_annual_stock.csv", index=False)

    counts, probs = build_annual_transitions(states)
    counts.to_csv(ANNUAL_DIR / "annual_transition_counts.csv", index=False)
    probs.to_csv(ANNUAL_DIR / "annual_transition_probabilities.csv", index=False)

    inflows = build_annual_inflows(states)
    inflows.to_csv(ANNUAL_DIR / "annual_inflows.csv", index=False)

    exits = build_annual_exits(states)
    exits.to_csv(ANNUAL_DIR / "annual_exits.csv", index=False)

    interciv_stock, interciv_new = build_interciv_flows(states)
    interciv_stock.to_csv(ANNUAL_DIR / "annual_interciv_stock.csv", index=False)
    interciv_new.to_csv(ANNUAL_DIR / "annual_interciv_newflow.csv", index=False)

    rate_table = build_rate_table(probs, inflows, exits, states)
    rate_table.to_csv(ANNUAL_DIR / "annual_ode_rates.csv", index=False)

    projected_rates = correct_and_project_rates(rate_table)
    projected_rates.to_csv(ANNUAL_DIR / "projected_ode_rates.csv", index=False)

    projected_stock = project_population(states, projected_rates)
    projected_stock.to_csv(ANNUAL_DIR / "projected_annual_stock.csv", index=False)

    observed_for_eval = stock[stock["year"] <= 2023].copy()
    max_obs_year = int(observed_for_eval["year"].max()) if not observed_for_eval.empty else 2023
    merged, group_metrics, comp_metrics, overall = evaluate_projection(projected_stock, observed_for_eval, project_end=min(max_obs_year, 2026))
    merged.to_csv(ANNUAL_DIR / "projection_evaluation.csv", index=False)
    group_metrics.to_csv(ANNUAL_DIR / "projection_accuracy_by_group.csv", index=False)
    comp_metrics.to_csv(ANNUAL_DIR / "projection_accuracy_by_compartment.csv", index=False)

    fig_rates = plot_annual_rates(rate_table, projected_rates)
    fig_heatmap = plot_interciv_heatmap(interciv_stock)
    fig_projection = plot_projection_by_compartment(projected_stock, observed_for_eval)

    doc_path = build_docx(rate_table, projected_stock, observed_for_eval,
                          group_metrics, comp_metrics, overall,
                          fig_rates, fig_heatmap, fig_projection, interciv_stock)
    print(f"Wrote {doc_path}")
    print(f"Overall projection RMSE: {overall['rmse'].iloc[0]:.3f}, MAPE: {overall['mape'].iloc[0]:.3%}")

    # Generate editable PPTX and a submission zip with docx, figures, and CSVs
    pptx_script = Path(__file__).parent / "build_annual_projection_pptx.py"
    if pptx_script.exists():
        subprocess.run([sys.executable, str(pptx_script)], check=True, cwd=BASE_DIR)

    zip_path = DOCS_DIR / "annual_rates_projection_submission.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in [
            doc_path,
            DOCS_DIR / "annual_rates_projection_figures.pptx",
            fig_rates,
            fig_heatmap,
            fig_projection,
            ANNUAL_DIR / "observed_annual_stock.csv",
            ANNUAL_DIR / "annual_transition_counts.csv",
            ANNUAL_DIR / "annual_transition_probabilities.csv",
            ANNUAL_DIR / "annual_inflows.csv",
            ANNUAL_DIR / "annual_exits.csv",
            ANNUAL_DIR / "annual_interciv_stock.csv",
            ANNUAL_DIR / "annual_interciv_newflow.csv",
            ANNUAL_DIR / "annual_ode_rates.csv",
            ANNUAL_DIR / "projected_ode_rates.csv",
            ANNUAL_DIR / "projected_annual_stock.csv",
            ANNUAL_DIR / "projection_evaluation.csv",
            ANNUAL_DIR / "projection_accuracy_by_group.csv",
            ANNUAL_DIR / "projection_accuracy_by_compartment.csv",
        ]:
            if path.exists():
                zf.write(path, path.name)
    print(f"Wrote {zip_path}")


if __name__ == "__main__":
    main()

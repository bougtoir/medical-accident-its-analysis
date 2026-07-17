#!/usr/bin/env python3
"""
Blood Pressure Monitoring 原稿（.docx）— 日本語版。

数値はすべて results/summary.json から読み込み、ハードコードしない。
参考文献は初出順に採番する（英語版と同一の文献データベース）。
図は figures/ja/ の日本語ラベル版を用いる。
"""

import os
import re

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from manuscript_common import REFDB, Citations, load_summary

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
FIGDIR = os.path.join(SCRIPT_DIR, "..", "figures", "ja")
OUTDIR = os.path.join(SCRIPT_DIR, "..", "manuscripts")
OUTPATH = os.path.join(OUTDIR, "BPM_ZeroFree_Manuscript_JA.docx")
os.makedirs(OUTDIR, exist_ok=True)

S = load_summary()
C = Citations(REFDB)
P = S["parameters"]
st = S["static"]
dy = S["dynamic"]
rg = S["range_dependence"]

JP_FONT = "MS Mincho"


def f2(x): return f"{x:.2f}"
def f1(x): return f"{x:.1f}"
def f3(x): return f"{x:.3f}"
def signed(x, nd=1): return f"{x:+.{nd}f}"
def ci(lo, hi, nd=3): return f"{lo:.{nd}f}\uff5e{hi:.{nd}f}"


doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles["Normal"]
style.font.name = JP_FONT
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.9
style.paragraph_format.space_after = Pt(0)


def _apply(run, size, bold, italic):
    run.font.name = JP_FONT
    run.font.size = size
    run._element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
        JP_FONT)
    run.bold = bold
    run.italic = italic


def add_para_with_refs(text, bold=False, italic=False, alignment=None,
                       space_after=Pt(6), font_size=Pt(11)):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    for part in re.split(r"(\{[^}]+\})", text):
        if part.startswith("{") and part.endswith("}"):
            run = p.add_run(part[1:-1]); run.font.superscript = True
        else:
            run = p.add_run(part)
        _apply(run, font_size, bold, italic)
    return p


def add_para(text, bold=False, italic=False, alignment=None, space_after=Pt(6)):
    return add_para_with_refs(text, bold, italic, alignment, space_after)


def add_heading_styled(text, level=1):
    h = doc.add_heading("", level=level)
    run = h.add_run(text)
    _apply(run, Pt(13 if level == 1 else 12), True, False)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_eq(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _apply(run, Pt(11), False, True)
    p.paragraph_format.space_after = Pt(6)


FIG_LEGENDS = []


def add_figure(filename, caption, width=Inches(6.2)):
    FIG_LEGENDS.append(caption)
    path = os.path.join(FIGDIR, filename)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(path):
        p.add_run().add_picture(path, width=width)
    else:
        _apply(p.add_run(f"[missing figure: {filename}]"), Pt(10), False, True)
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(12)
    cap.paragraph_format.space_after = Pt(12)
    for part in re.split(r"(\{[^}]+\})", caption):
        if part.startswith("{") and part.endswith("}"):
            r = cap.add_run(part[1:-1]); r.font.superscript = True
        else:
            r = cap.add_run(part)
        _apply(r, Pt(9), False, True)


# ── タイトルページ ──
add_para("原著論文", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
_apply(title.add_run(
    "ゼロ校正で補正できないもの：観血的動脈圧モニタリングにおける"
    "ゼロ校正後の残存ゲイン誤差と動的応答誤差の検出"), Pt(15), True, False)
title.paragraph_format.space_after = Pt(18)
add_para("ランニングタイトル：ゼロ校正で補正できないもの", italic=True,
         alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para("[著者名は後日記入]", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para("[所属は後日記入]", alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para("責任著者：", bold=True)
add_para("[氏名、住所、電子メール、ORCID]")
doc.add_paragraph()
add_para("キーワード：動脈圧モニタリング；ゼロ校正；一致相関係数；"
         "Bland–Altman 解析；比例バイアス；動的応答；ファストフラッシュ試験；"
         "測定法間比較")
doc.add_page_break()

# ── 要旨 ──
add_heading_styled("要旨", level=1)
add_para("目的", bold=True, space_after=Pt(2))
add_para(
    "観血的動脈圧トランスデューサのゼロ校正は直流（DC）オフセットを除去するが、"
    "センサのゲイン（スケール）誤差は補正しない。本研究では、機器バリデーション"
    "研究で一般に報告される解析のうち、ゼロ校正後に残存するゲイン誤差を実際に検出"
    "できるのはどれか、また周波数依存の動的応答誤差がどのように振る舞うかを検討した。")
add_para("方法", bold=True, space_after=Pt(2))
add_para(
    f"シード固定のシミュレーションにより、機器と基準の対応する収縮期圧"
    f"（各シナリオ n = {P['n_static']}）を4条件で生成した：オフセットのみ、"
    "理想的ゼロ校正、未補正ゲイン誤差、補正オフセットで隠れたゲイン誤差。各条件を、"
    "最小限の Bland–Altman 要約（平均バイアス・一致限界）、Bland–Altman の差対平均"
    "回帰、Deming 回帰、Passing–Bablok 回帰、および Lin の一致相関係数（CCC）と"
    "そのスケールシフト成分で評価した。合成動脈圧波形を、最適・低制動・過制動の"
    "2次カテーテル–トランスデューサ系に通した。")
add_para("結果", bold=True, space_after=Pt(2))
add_para(
    f"10%のゲイン誤差が補正オフセットで隠された場合、平均バイアスは"
    f"{signed(st['S4_gain_masked']['bias'])} mmHg（一見許容範囲）であったが、"
    f"Bland–Altman 回帰の傾き（{f3(st['S4_gain_masked']['prop_slope'])}）、"
    f"Deming の傾き（{f3(st['S4_gain_masked']['deming_slope'])}）、"
    f"Passing–Bablok の傾き（{f3(st['S4_gain_masked']['pb_slope'])}）、"
    f"CCC スケールシフト（v = {f3(st['S4_gain_masked']['v'])}）はいずれも検出した。"
    f"低制動は脈圧を{round((dy['underdamped_pp']['mean_ratio']-1)*100)}%過大評価し、"
    f"過制動は{round((1-dy['overdamped_pp']['mean_ratio'])*100)}%減衰させた。"
    "これらはゼロ校正後も残存し、ファストフラッシュ試験で診断される。単一機器では"
    f"標本圧の範囲が広がると CCC が上昇した（{f2(rg['ccc'][0])} から "
    f"{f2(rg['ccc'][-1])}）。")
add_para("結論", bold=True, space_after=Pt(2))
add_para(
    "平均バイアスと一致限界のみの要約は臨床的に重要なゲイン誤差を見逃しうる。"
    "比例バイアスを検出できる解析と、明示的な動的応答チェックを併せて報告すべきである。")
doc.add_page_break()

# ── 1. 序論 ──
add_heading_styled("1. 序論", level=1)
add_para_with_refs(
    "観血的動脈圧モニタリングは麻酔・集中治療で最も頻繁に行われる測定の一つである。"
    "動脈カテーテル挿入後、臨床医はモニタリング開始前に静水圧基準点（phlebostatic "
    "axis）の高さで圧トランスデューサを大気圧にゼロ校正し"
    + C.cite("saugel2020", "saugelsessler2021") + "、トランスデューサとカテーテル"
    "の高さ関係が変化するたびにゼロ校正を繰り返す" + C.cite("gupta2025") + "。")
add_para_with_refs(
    "液体充填式動脈ラインでは、外部トランスデューサが生理食塩水の液柱でカテーテル"
    "先端に接続され、これが直流（DC）オフセットを生じる：トランスデューサと測定点の"
    "静水圧差、ゲージ圧の大気基準、緩徐なトランスデューサドリフトである"
    + C.cite("mark1998", "mcghee2002") + "。ゼロ校正はこれらのオフセットを一括して"
    "除去する。")
add_para_with_refs(
    "しかしゼロ校正はオフセット補正にすぎず、系のゲイン（感度）—真の圧変化と表示値"
    "の比例関係—を検証しない。ゲインが誤っていれば波形はスケールされ、誤差はゼロ校正"
    "後も残存する。第二の周波数依存のゲイン誤差は液体充填系の動的応答に由来する："
    "低制動または過制動は拍動信号、したがって脈圧を歪める" + C.cite("gardner1981")
    + "。")
add_para_with_refs(
    "それでもゼロ校正はベッドサイドで不可欠である。これらのオフセットは「工場で"
    "直せる機器の欠陥」ではなく、設置ごと・時間ごとに変化する物理量だからである。"
    "静水圧柱は静水圧基準点に対するトランスデューサの高さに依存し（約10 cmあたり"
    "7.5 mmHg）、患者やトランスデューサを動かすたびに変化する" + C.cite("gupta2025")
    + "。ゲージ圧の大気基準は天候や標高で変動し、電気的・熱的ドリフトは経時的に"
    "蓄積する" + C.cite("mark1998", "mcghee2002") + "。重要なのは、検出は補正では"
    "なく、両者はレイヤーが異なるという点である。Bland–Altman 回帰などの測定法間"
    "比較解析は、対応する基準測定を必要とする事後的な検証手法であり、目の前の患者"
    "にリアルタイムの基準は存在しない。ゼロ校正は、基準なしで既知の設置固有オフ"
    "セットを事前に除去できる唯一の手続きである。回帰でオフセットが見えたとしても、"
    "それを除去するにはやはりゼロ校正が要る。すなわちゼロ校正は運用上オフセットを"
    "補正し、回帰と CCC 分解は残存するゲインと構造を検出するものであって、いずれも"
    "他方を代替できない。")

add_para_with_refs(
    "動脈圧および派生血行動態モニタのバリデーション研究は、圧倒的に Bland–Altman "
    "要約（平均バイアスと95%一致限界）" + C.cite("blandaltman1986") +
    "と Critchley の誤差率" + C.cite("critchley1999") + "を、基準法の精度に留意"
    "しつつ" + C.cite("cecconi2009") + "報告している。本来の Bland–Altman 法は"
    "比例バイアスを検出するための差の平均に対する回帰も含むが" +
    C.cite("blandaltman1999") + "、この手順はしばしば省略され、オフセットとゲイン"
    "誤差が部分的に相殺する場合、平均バイアスのみの要約は誤解を招きうる。")
add_para_with_refs(
    "本研究は、適切に実施された Bland–Altman 解析にゲイン誤差が見えないと主張する"
    "ものではない。むしろ再現可能なシミュレーションを用いて、文献で実際に報告される"
    "解析のうちゼロ校正後の残存ゲイン誤差を検出できるものを定量化し、Lin の一致相関"
    "係数（CCC）とそのスケールシフト成分" + C.cite("lin1989", "lin2000") +
    "、および Deming" + C.cite("linnet1990") + "・Passing–Bablok" +
    C.cite("passingbablok1983") + "回帰を補完的診断として位置づける。さらに液体"
    "充填式動脈ラインに特に臨床的に重要な動的応答（制動）誤差を統合し、CCC の範囲"
    "依存性" + C.cite("nickerson1997", "barnhart2007") + "を論じる。")

# ── 2. 方法 ──
add_heading_styled("2. 方法", level=1)
add_heading_styled("2.1. DC 成分と AC 成分としての動脈圧", level=2)
add_para_with_refs(
    "動脈圧波形は、緩徐に変化する DC 水準（平均動脈圧と外部オフセット）と、心拍出"
    "による拍動性交流（AC）成分に分解できる。脈圧（PP = 収縮期圧 [SBP] − 拡張期圧 "
    "[DBP]）は純粋な AC 量であり、加法的な DC オフセットとは独立だがセンサゲインに"
    "比例する（図1）。オフセットは PP を変えずに基線を移動させ、ゲイン誤差は波形全体"
    "をスケールして PP を変化させる。")
add_figure("figure1_signal_decomposition.png",
           "図1. 動脈圧信号の分解。(A) DC オフセットは基線を移動させるが脈圧は不変。"
           "(B) ゲイン誤差は波形全体をスケールし、脈圧はゲインに比例して変化する。"
           "ゼロ校正は (A) のオフセットを補正するが (B) のゲイン誤差は補正できない。")

add_heading_styled("2.2. 統計手法", level=2)
add_para_with_refs(
    "各対応データに対し次を算出した：(i) 最小限の Bland–Altman 要約—平均バイアス、"
    "95%一致限界（平均バイアス ± 1.96 SD）、誤差率" + C.cite("critchley1999") +
    "；(ii) 差の平均に対する Bland–Altman 回帰（傾きが比例バイアスを検定）" +
    C.cite("blandaltman1999") + "；(iii) 両法の誤差を許容しジャックナイフ信頼区間"
    "を付した Deming 回帰" + C.cite("linnet1990") + "；(iv) Passing–Bablok "
    "ノンパラメトリック回帰" + C.cite("passingbablok1983") + "；(v) Lin の CCC" +
    C.cite("lin1989") + "。")
add_para_with_refs(
    "Lin の CCC（ρc）は ρc = r × C_b と分解され、r は Pearson 相関（精度）、C_b は"
    "バイアス補正係数（正確度）である" + C.cite("lin2000") + "。C_b はさらに次のよう"
    "に分解される：")
add_eq("C_b = 2 / (v + 1/v + u\u00b2)")
add_para_with_refs(
    "ここで v = σ_device/σ_reference はスケールシフト（ゲイン比）、u = "
    "(μ_device − μ_reference)/√(σ_device σ_reference) はロケーションシフト"
    "（オフセット）である。ゼロ校正は u を0に近づけるが v は不変であり、残存ゲイン"
    "誤差は完全なゼロ校正後も v ≠ 1 を保つ。")
add_para_with_refs(
    f"事後的選択を避けるため、事前固定の検出基準を用いた：平均バイアス要約は "
    f"|平均バイアス| > {f1(P['bias_threshold'])} mmHg で誤差ありと判定；"
    "Bland–Altman 回帰は傾きの95%信頼区間（CI）が0を含まなければ判定；Deming と "
    "Passing–Bablok は傾きの95% CI が1を含まなければ判定；CCC スケールシフトは "
    f"|v − 1| > {f2(P['v_threshold'])} で判定。一致の強さは既報の CCC 基準" +
    C.cite("mcbride2005") + "で解釈した。")

add_heading_styled("2.3. 動的応答とファストフラッシュ試験", level=2)
add_para_with_refs(
    "液体充填式カテーテル–トランスデューサ系は、固有周波数 f_n と制動係数 ζ で特徴"
    "づけられる2次系として振る舞う" + C.cite("gardner1981") + "。その伝達関数 "
    "H(f) = 1/[1 − (f/f_n)² + i·2ζ(f/f_n)] を合成動脈圧波形のフーリエ高調波に適用"
    "した。低制動（低 ζ、動脈高調波付近の共振）は収縮期オーバーシュートと PP の過大"
    "評価を生じ、過制動は PP を減衰させる。いずれもゼロ校正で補正できない拍動信号の"
    "ゲイン誤差であり、ベッドサイドではファストフラッシュ（方形波）試験で定量される"
    + C.cite("kleinman1992", "romagnoli2014") + "。")

add_heading_styled("2.4. シミュレーション設計と再現性", level=2)
add_para_with_refs(
    f"本研究のデータはすべてシミュレーションであり、患者から測定したものではない。"
    f"静的シナリオは {P['sbp_low']}〜{P['sbp_high']} mmHg に一様分布する収縮期圧 "
    f"n = {P['n_static']} 対からなり、ガウス測定ノイズ（SD = {f1(P['noise_sd'])} "
    f"mmHg）を付与した：S1 オフセットのみ（+{f1(P['offset_mmHg'])} mmHg、ゲイン "
    f"1.0）；S2 理想的ゼロ校正（オフセット除去、ゲイン 1.0）；S3 未補正ゲイン誤差"
    f"（ゲイン {f2(P['gain'])}）；S4 同じゲイン誤差に、平均差がほぼ0となるよう補正的"
    "な負のオフセットを付与。動的シナリオは合成波形を最適・低制動・過制動系に通した。"
    "範囲依存性実験は、固定機器（ゲイン誤差5%）を幅の増える圧範囲で標本化した。"
    "乱数シード、全母数、および完全なパイプライン（シミュレーション→解析→図→原稿）"
    "を公開リポジトリに提供し、第三者がクリーンなクローンから全数値・表・図を再生成"
    "できるようにした。")

# ── 3. 結果 ──
add_heading_styled("3. 結果", level=1)
add_heading_styled("3.1. 静的シナリオ", level=2)
add_para_with_refs(
    f"4つの静的シナリオの一致プロットを図2に、全指標を表1に示す。理想的ゼロ校正"
    f"（S2）はほぼ完全な一致を示した（平均バイアス {signed(st['S2_zeroed_ideal']['bias'])} "
    f"mmHg、CCC {f3(st['S2_zeroed_ideal']['ccc'])}、v {f3(st['S2_zeroed_ideal']['v'])}）。"
    f"未補正の10%ゲイン誤差（S3）は平均バイアス"
    f"（{signed(st['S3_gain_uncompensated']['bias'])} mmHg）を含むあらゆる解析で明白"
    "であった：ゲイン誤差は Bland–Altman 解析に見えないわけではない。")
add_figure("figure2_scenarios_concordance.png",
           "図2. 4つの静的シナリオにおける機器対基準の一致（破線＝同一線）。各パネル"
           "に CCC、CCC スケールシフト v、Bland–Altman 平均バイアスを示す。S4 は10%"
           "のゲイン誤差にもかかわらず平均バイアスがほぼ0である。")
add_para_with_refs(
    f"決定的なのは S4 で、補正オフセットにより平均バイアスはわずか"
    f"{signed(st['S4_gain_masked']['bias'])} mmHg（±{f1(P['bias_threshold'])} mmHg "
    "の許容帯内）となり、最小限の Bland–Altman 要約は問題を検出しない。それでも比例"
    "バイアスを検出する解析はすべて隠れたゲイン誤差を検出する：Bland–Altman 回帰の"
    f"傾き {f3(st['S4_gain_masked']['prop_slope'])}（95% CI "
    f"{ci(st['S4_gain_masked']['prop_slope_lo'], st['S4_gain_masked']['prop_slope_hi'])}）、"
    f"Deming の傾き {f3(st['S4_gain_masked']['deming_slope'])}（95% CI "
    f"{ci(st['S4_gain_masked']['deming_lo'], st['S4_gain_masked']['deming_hi'])}）、"
    f"Passing–Bablok の傾き {f3(st['S4_gain_masked']['pb_slope'])}（95% CI "
    f"{ci(st['S4_gain_masked']['pb_lo'], st['S4_gain_masked']['pb_hi'])}）、"
    f"CCC スケールシフト v = {f3(st['S4_gain_masked']['v'])}。全シナリオでの検出パター"
    "ンを図3に、S2 と S4 の差対平均プロットを図4に示す。")
add_figure("figure3_detection_panel.png",
           "図3. 各シナリオでどの報告解析が誤差を検出するか（緑＝検出、灰＝見逃し）。"
           "S4 では平均バイアス要約が見逃すゲイン誤差を、比例バイアスを検出する全解析"
           "が検出する。")
add_figure("figure4_ba_masked_gain.png",
           "図4. Bland–Altman 差対平均プロット。S2 では回帰の傾きは平坦だが、S4 では"
           "平均バイアスがほぼ0であるにもかかわらず傾きが明確に正であり、隠れた比例"
           "（ゲイン）誤差が明らかになる。")

add_heading_styled("3.2. 動的応答シナリオ", level=2)
add_para_with_refs(
    f"動的応答指標を表2に示し、図5で図示する。最適系は PP を忠実に再現した"
    f"（PP 比 {f2(dy['optimal_pp']['mean_ratio'])}）。低制動は収縮期オーバーシュート"
    f"を生じ PP を{round((dy['underdamped_pp']['mean_ratio']-1)*100)}%過大評価した"
    f"（PP 比 {f2(dy['underdamped_pp']['mean_ratio'])}、CCC スケールシフト v = "
    f"{f2(dy['underdamped_pp']['v'])}）。過制動は PP を"
    f"{round((1-dy['overdamped_pp']['mean_ratio'])*100)}%減衰させた"
    f"（PP 比 {f2(dy['overdamped_pp']['mean_ratio'])}）。これらはゼロ校正後も残存する"
    "拍動信号のゲイン誤差である。歪みは周波数依存（圧振幅ではなく心拍数で変化）"
    "であるため、静的回帰統計が前提とする振幅比例モデルには部分的にしか対応せず、"
    "ファストフラッシュ試験の母数（f_n, ζ）で直接特徴づけるのが最適である。")
add_figure("figure5_dynamic_response.png",
           "図5. カテーテル–トランスデューサ系の動的応答。(A) 最適・低制動・過制動系"
           "の周波数応答と動脈圧高調波の重ね書き。(B) 波形例。(C) 系ごとの平均 PP 比"
           "を付した計測脈圧対真の脈圧。")

add_heading_styled("3.3. CCC の範囲依存性", level=2)
add_para_with_refs(
    f"ゲイン誤差5%を固定した単一機器では、狭い圧範囲（幅 "
    f"{int(rg['range_width'][0])} mmHg）で標本化すると CCC は {f2(rg['ccc'][0])} "
    f"であったが、広い範囲（{int(rg['range_width'][-1])} mmHg）では "
    f"{f2(rg['ccc'][-1])} に上昇し、スケールシフト v は真値近傍にとどまった（図6）。"
    "したがって CCC は標本化した圧範囲とともに報告すべきであり、構造成分（C_b, v）"
    "の方が研究間で移転可能性が高い" + C.cite("nickerson1997") + "。")
add_figure("figure6_range_dependence.png",
           "図6. CCC の範囲依存性。単一の固定機器において、標本圧範囲が広がると CCC "
           "と C_b は上昇するが、スケールシフト v は真のゲイン比の近傍にとどまる。")

# ── 4. 考察 ──
add_heading_styled("4. 考察", level=1)
add_para_with_refs(
    "本研究の中心的知見は限定的だが妥当と考える：機器バリデーション報告を支配する"
    "平均バイアス＋一致限界の要約は、偶然のオフセットが平均への影響を相殺する場合、"
    "臨床的に重要なゲイン誤差を見逃しうる。一方、Bland–Altman の差対平均回帰" +
    C.cite("blandaltman1999") + "、Deming" + C.cite("linnet1990") +
    "・Passing–Bablok" + C.cite("passingbablok1983") + "回帰、CCC スケールシフト" +
    C.cite("lin2000") + "はいずれもこれを検出する。これは Bland–Altman 法自体では"
    "なく、一般的な報告慣行への批判である。")

add_heading_styled("4.1. バリデーション報告への提言", level=2)
add_para_with_refs(
    "動脈圧および派生モニタのバリデーション研究" + C.cite("kim2014", "joosten2017") +
    "は、平均バイアスと一致限界に加えて次を報告すべきである：(i) Bland–Altman 回帰"
    "の傾きとその CI（あるいは同等の Deming／Passing–Bablok の傾き）；(ii) r, C_b, "
    "u, v に分解した CCC；(iii) 標本化した圧範囲（CCC は範囲依存のため）" +
    C.cite("nickerson1997", "barnhart2007") + "。特に C_b 単独では純粋なゲイン誤差に"
    f"比較的鈍感であり（S4 では v = {f3(st['S4_gain_masked']['v'])} にもかかわらず "
    f"C_b = {f3(st['S4_gain_masked']['C_b'])}）、スケールシフト v を明示的に報告すべき"
    "である。これらの追加は数行の解析で済み、規制当局や臨床医" + C.cite("odor2017") +
    "がオフセット補正された一致と真のゲイン正確度を区別できるようになる。")

add_heading_styled("4.2. 動的応答は独立した臨床的ゲイン誤差である", level=2)
add_para_with_refs(
    "動的応答誤差は液体充填系に特有で、低制動（気泡、長く柔軟なチューブ）または"
    "過制動（血栓、屈曲）としてベッドサイドで日常的に遭遇する" +
    C.cite("gardner1981", "romagnoli2014") + "。周波数依存であるため静的な測定法間"
    "比較統計では完全には捉えられず、ファストフラッシュ試験で f_n と ζ を読み取り"
    "直接評価するのが最適である" + C.cite("kleinman1992") + "。したがって液体充填式"
    "動脈ラインの正確度に関する記述には、ファストフラッシュ動的応答チェックを併記"
    "すべきである。脈圧はゲインに比例しオフセットとは独立であるため、鋭敏な指標となる。")

add_heading_styled("4.3. センサ設計への示唆", level=2)
add_para_with_refs(
    "今なおゼロ校正が行われるのは、設置固有のオフセットをリアルタイムに補正できる"
    "からであり、これは事後的な測定法間比較統計では担えない役割である。この点は"
    "同時に進むべき方向も示す：ゼロ校正は運用上のオフセット問題のみを扱うため、"
    "オフセット源をハードウェア設計で除去すればゼロ校正自体が不要となり、ゲインと"
    "動的応答が設計・検証上の残された問題として残る。")
add_para_with_refs(
    "ゼロ校正が補正するオフセットは設計でも除去できる。カテーテル先端型 MEMS "
    "（微小電気機械システム）センサは静水圧柱を排除し" +
    C.cite("hasenkamp2012", "song2020") + "、気圧補正付き絶対圧センサは大気基準を"
    "除去し、自己校正型 MEMS 基準はドリフトを抑える" + C.cite("kang2022") +
    "。高忠実度圧ワイヤはすでに正確なカテーテルベース圧測定を実証している" +
    C.cite("scalia2023") + "。オフセットを設計で除去し、正確な PP 再現でゲインを"
    "検証すれば、残存する系統誤差は完全に v に集約され、具体的な設計目標（校正なしで "
    "v ≈ 1、C_b ≥ 0.99）が得られる" + C.cite("mcbride2005") + "。現代のピエゾ抵抗型 "
    "MEMS センサは全スケールの0.1%を十分下回る非線形性を達成しており" +
    C.cite("barlian2009") + "、拍動範囲で検証したゲインが生理的範囲全体に一般化する"
    "という仮定を支持する。")

add_heading_styled("4.4. 派生血行動態モニタへの拡張", level=2)
add_para_with_refs(
    "オフセット対ゲインの区別は、自己校正ルーチンがセットポイントを再最適化する"
    "—オフセット補正—一方で圧・インピーダンスから流量への変換ゲインを補正しない"
    "派生モニタにも及ぶ" +
    C.cite("chatterjee2009", "ameloot2015", "squara2007", "manecke2005") +
    "。これら機器についてスケールシフト v を報告すれば、校正への依存が明示される。")

add_heading_styled("4.5. 限界", level=2)
add_para_with_refs(
    "本研究はシミュレーション研究であり、データは合成で特定の誤差構造を分離するよう"
    "選ばれており、明示された精度をもつ観血的基準に対する前向き臨床比較の代替では"
    "ない" + C.cite("cecconi2009") + "。静的モデルはセンサの線形性を仮定し、動的"
    "モデルは単一の2次近似である。本研究の価値は、どの報告解析が残存ゲイン誤差を"
    "検出するかを明示的かつ再現可能に定量化する点にあり、経験的バリデーションの前段"
    "であって代替ではない。")

# ── 5. 結論 ──
add_heading_styled("5. 結論", level=1)
add_para_with_refs(
    "ゼロ校正はオフセットを補正するがゲインは補正しない。平均バイアス＋一致限界の"
    "要約は、オフセットが隠す場合に残存ゲイン誤差を見逃しうるが、Bland–Altman 回帰、"
    "Deming・Passing–Bablok 回帰、CCC スケールシフトはこれを検出する。動的応答誤差は"
    "ファストフラッシュ試験で診断される独立した周波数依存のゲイン誤差である。動脈圧"
    "バリデーション研究は、比例バイアスを検出する解析、標本範囲を付した CCC 分解、"
    "および明示的な動的応答チェックを報告すべきである。")

# ── 開示事項 ──
add_heading_styled("利益相反", level=1)
add_para("[著者により記入]")
add_heading_styled("研究資金", level=1)
add_para("[著者により記入]")
add_heading_styled("データ利用可能性", level=1)
add_para(
    "すべてのシミュレーションコードと生成データは公開複製リポジトリ"
    "（https://github.com/bougtoir/bpm-zero-calibration）で公開されている。"
    "本論文の全数値・表・図は、クリーンなクローンから解析パイプラインを実行すること"
    "で再生成できる。")
doc.add_page_break()

# ── 表1 ──
add_heading_styled("表1", level=1)
add_para(f"表1. 4つの静的シナリオの測定法間比較指標（各 n = {P['n_static']} 対）。",
         bold=True)
t1_headers = ["シナリオ", "平均バイアス\n(mmHg)", "95% LoA\n(mmHg)", "PE (%)",
              "BA回帰傾き\n(95% CI)", "Deming傾き\n(95% CI)",
              "P–B傾き\n(95% CI)", "CCC", "C_b", "v"]
scen_names = {"S1_offset_only": "S1 オフセットのみ",
              "S2_zeroed_ideal": "S2 ゼロ校正(理想)",
              "S3_gain_uncompensated": "S3 ゲイン誤差",
              "S4_gain_masked": "S4 隠れたゲイン"}
table = doc.add_table(rows=1 + len(scen_names), cols=len(t1_headers))
table.style = "Light Grid Accent 1"
for i, h in enumerate(t1_headers):
    for p in table.rows[0].cells[i].paragraphs:
        _apply(p.add_run(h), Pt(8), True, False)
for row_idx, (key, label) in enumerate(scen_names.items(), start=1):
    d = st[key]
    row = [label, signed(d["bias"]),
           f"{d['loa_lower']:.1f}〜{d['loa_upper']:.1f}", f1(d["pe"]),
           f"{d['prop_slope']:.3f}\n({d['prop_slope_lo']:.3f}, {d['prop_slope_hi']:.3f})",
           f"{d['deming_slope']:.3f}\n({d['deming_lo']:.3f}, {d['deming_hi']:.3f})",
           f"{d['pb_slope']:.3f}\n({d['pb_lo']:.3f}, {d['pb_hi']:.3f})",
           f3(d["ccc"]), f3(d["C_b"]), f3(d["v"])]
    for col_idx, txt in enumerate(row):
        for p in table.rows[row_idx].cells[col_idx].paragraphs:
            _apply(p.add_run(txt), Pt(8), False, False)
add_para(
    "BA = Bland–Altman；CI = 信頼区間；LoA = 一致限界；PE = 誤差率；"
    "P–B = Passing–Bablok；CCC = 一致相関係数；C_b = バイアス補正係数；"
    "v = スケールシフト。S4 は平均バイアスがほぼ0だが回帰／Deming／Passing–Bablok "
    "の傾きは明確に非ゼロで v > 1、すなわち平均バイアス要約から隠れたゲイン誤差である。",
    italic=True)
doc.add_page_break()

# ── 表2 ──
add_heading_styled("表2", level=1)
add_para("表2. 脈圧（PP）と収縮期圧（SBP）の動的応答指標。", bold=True)
t2_headers = ["系 (f_n, ζ)", "PP 比", "PP平均バイアス\n(mmHg)", "CCC (PP)",
              "v (PP)", "SBP平均バイアス\n(mmHg)"]
dyn_rows = {"optimal": "最適 (25 Hz, 0.65)",
            "underdamped": "低制動 (10 Hz, 0.15)",
            "overdamped": "過制動 (8 Hz, 0.80)"}
table2 = doc.add_table(rows=1 + len(dyn_rows), cols=len(t2_headers))
table2.style = "Light Grid Accent 1"
for i, h in enumerate(t2_headers):
    for p in table2.rows[0].cells[i].paragraphs:
        _apply(p.add_run(h), Pt(8), True, False)
for row_idx, (key, label) in enumerate(dyn_rows.items(), start=1):
    pp = dy[f"{key}_pp"]; sbp = dy[f"{key}_sbp"]
    row = [label, f2(pp["mean_ratio"]), signed(pp["bias"]),
           f3(pp["ccc"]), f3(pp["v"]), signed(sbp["bias"])]
    for col_idx, txt in enumerate(row):
        for p in table2.rows[row_idx].cells[col_idx].paragraphs:
            _apply(p.add_run(txt), Pt(8), False, False)
add_para(
    "PP 比 = 平均(計測 PP / 真の PP)；f_n = 固有周波数；ζ = 制動係数。低制動は脈圧を"
    "過大評価し、過制動は減衰させる。平均動脈（DC）水準は保存されるため、ゼロ校正"
    "ではこれらの誤差は補正されない。", italic=True)
doc.add_page_break()

# ── 図の説明 ──
add_heading_styled("図の説明", level=1)
for legend in FIG_LEGENDS:
    add_para(legend)
doc.add_page_break()

# ── 参考文献 ──
add_heading_styled("参考文献", level=1)
for idx, ref in enumerate(C.ordered_references(), start=1):
    p = doc.add_paragraph()
    _apply(p.add_run(f"{idx}. {ref}"), Pt(9), False, False)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.4

doc.save(OUTPATH)
print(f"Japanese BPM manuscript saved: {OUTPATH}")
print(f"References cited: {len(C.order)} / {len(REFDB)}")

#!/usr/bin/env python3
"""Cover letter (.docx) for the IEEE TIM submission."""

import os

from docx import Document
from docx.shared import Pt, Cm

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTDIR = os.path.join(SCRIPT_DIR, "..", "cover_letter")
OUTPATH = os.path.join(OUTDIR, "TIM_Cover_Letter_EN.docx")
os.makedirs(OUTDIR, exist_ok=True)

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(8)


def para(text="", bold=False, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold
    return p


para("[Date]")
para()
para("To the Editor-in-Chief")
para("IEEE Transactions on Instrumentation and Measurement", bold=True)
para()
para("Dear Editor,")

para(
    "We are pleased to submit our original article, \u201cWhat zeroing cannot "
    "fix: detecting residual gain and dynamic-response errors after zero "
    "calibration in invasive arterial pressure monitoring,\u201d for "
    "consideration for publication in IEEE Transactions on Instrumentation and "
    "Measurement.")

para(
    "Zero calibration of a fluid-filled arterial pressure transducer removes "
    "the direct-current (DC) offset but cannot correct the transducer gain "
    "(sensitivity) or the frequency-dependent damping of the catheter-sensor "
    "system. We use a fully reproducible simulation, augmented with public "
    "intra-operative waveforms from the VitalDB Open Dataset, to show which "
    "statistical diagnostics actually detect a residual gain error after "
    "zeroing. Our key message is a critique of common reporting practice, not "
    "of the Bland\u2013Altman method itself: a minimal summary of mean bias and "
    "limits of agreement can miss a clinically relevant gain error when a "
    "compensating offset cancels its effect on the mean, whereas the "
    "Bland\u2013Altman regression of difference on mean, Deming and "
    "Passing\u2013Bablok regression, and the scale-shift component (v) of "
    "Lin\u2019s concordance correlation coefficient all detect it.")

para(
    "We further integrate the dynamic-response (damping) error as a distinct, "
    "frequency-dependent gain error diagnosed by the fast-flush test, and we "
    "show that the concordance correlation coefficient is range-dependent. "
    "Because the paper is framed around measurement validity, sensor gain, and "
    "a concrete design target (v \u2248 1 without calibration) for next-generation "
    "self-calibrating pressure sensors, we believe it is well suited to the "
    "instrumentation-and-measurement readership of IEEE TIM.")

para(
    "We submit the manuscript as an Original Article. It contains 7 inline "
    "figures and 3 inline tables, a structured abstract of fewer than 300 words, "
    "and a main text of approximately 4,300 words. All figures, tables, and "
    "numerical values are generated directly from the analysis pipeline so that "
    "every result can be regenerated from the public data and the supplementary "
    "code. As disclosed in the Methods, the simulation and analysis code, the "
    "figures, and drafts of the text were prepared with the assistance of an AI "
    "coding assistant (Devin; Cognition AI); the authors verified all code, "
    "results and statements and take full responsibility for the work.")

para(
    "This is a simulation and secondary-data-validation study; the synthetic "
    "data are described as such and the real waveforms come from the publicly "
    "available VitalDB Open Dataset. The manuscript is original, has not been "
    "published previously, and is not under consideration elsewhere. All authors "
    "have approved the submission and declare no conflicts of interest relevant "
    "to this work. We have no preferred or excluded reviewers to suggest, but "
    "are happy to provide names on request.")

para("Thank you for considering our submission. We look forward to your "
     "response.")
para()
para("Sincerely,")
para()
para("[Corresponding author name, degrees]")
para("[Affiliation]")
para("[Email]")

doc.save(OUTPATH)
print(f"Cover letter saved: {OUTPATH}")

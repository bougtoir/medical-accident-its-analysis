#!/usr/bin/env python3
"""Cover letter (.docx) for the Blood Pressure Monitoring submission."""

import os

from docx import Document
from docx.shared import Pt, Cm

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTDIR = os.path.join(SCRIPT_DIR, "..", "cover_letter")
OUTPATH = os.path.join(OUTDIR, "BPM_Cover_Letter_EN.docx")
os.makedirs(OUTDIR, exist_ok=True)

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(8)


def para(text="", bold=False, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold
    return p


para("[Date]")
para()
para("To the Editor-in-Chief")
para("Blood Pressure Monitoring", bold=True)
para()
para("Dear Editor,")

para(
    "We are pleased to submit our original article, \u201cWhat zeroing cannot "
    "fix: detecting residual gain and dynamic-response errors after zero "
    "calibration in invasive arterial pressure monitoring,\u201d for "
    "consideration for publication in Blood Pressure Monitoring.")

para(
    "Zero calibration of a fluid-filled arterial pressure transducer is a "
    "daily bedside ritual, but it corrects only the direct-current offset and "
    "not the gain (sensitivity) of the system. Using a fully reproducible "
    "simulation study, we show which of the analyses commonly reported in "
    "device-validation work actually detect a residual gain error after "
    "zeroing. Our key message is a critique of common reporting practice, not "
    "of the Bland\u2013Altman method itself: a minimal summary of mean bias and "
    "limits of agreement can miss a clinically relevant gain error when a "
    "compensating offset cancels its effect on the mean, whereas the "
    "Bland\u2013Altman difference-versus-mean regression, Deming and "
    "Passing\u2013Bablok regression, and the scale-shift component of Lin\u2019s "
    "concordance correlation coefficient all detect it.")

para(
    "We further integrate the dynamic-response (damping) error that is "
    "specific to fluid-filled arterial lines, showing how under- and "
    "over-damping distort pulse pressure in ways that survive zeroing and are "
    "diagnosed by the fast-flush test, and we discuss the range-dependence "
    "of the concordance correlation coefficient. This scope matches the "
    "readership of Blood Pressure Monitoring, which spans the measurement, "
    "methodology and clinical interpretation of arterial pressure.")

para(
    "We submit the manuscript as an Original Study. It contains four figures "
    "and two tables (six items in total, within the journal limit), with two "
    "further figures provided as Supplemental Digital Content, a structured "
    "abstract of fewer than 250 words, and a main text within the 6000-word "
    "limit. As disclosed in the Methods, the simulation and analysis code, "
    "the figures, and drafts of the text were prepared with the assistance of "
    "an AI coding assistant (Devin; Cognition AI); the authors verified all "
    "code, results and statements and take full responsibility for the work.")

para(
    "This is a simulation study; all data are clearly described as synthetic "
    "and no clinical or previously published measurements are presented as "
    "empirical results. To support transparency and reuse, the complete "
    "pipeline (data generation, analysis, figures, tables and the manuscript "
    "itself) is openly available, so that every number, table and figure can "
    "be regenerated from a clean clone.")

para(
    "The manuscript is original, has not been published previously, and is "
    "not under consideration elsewhere. All authors have approved the "
    "submission and declare no conflicts of interest relevant to this work. "
    "We have no preferred or excluded reviewers to suggest, but are happy to "
    "provide names on request.")

para("Thank you for considering our submission. We look forward to your "
     "response.")
para()
para("Sincerely,")
para()
para("[Corresponding author name, degrees]")
para("[Affiliation]")
para("[Email]")

doc.save(OUTPATH)
print(f"Cover letter saved: {OUTPATH}")

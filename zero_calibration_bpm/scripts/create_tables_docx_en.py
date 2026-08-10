#!/usr/bin/env python3
"""Separate editable English tables (.docx), built from results/summary.json."""

import os

from docx import Document
from docx.shared import Pt

from manuscript_common import load_summary

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTDIR = os.path.join(SCRIPT_DIR, "..", "manuscripts")
OUTPATH = os.path.join(OUTDIR, "TIM_Tables_EN.docx")
os.makedirs(OUTDIR, exist_ok=True)

S = load_summary()
P = S["parameters"]
st = S["static"]
dy = S["dynamic"]
re = S["real_static"]


def f1(x): return f"{x:.1f}"
def f3(x): return f"{x:.3f}"
def f2(x): return f"{x:.2f}"
def signed(x): return f"{x:+.1f}"


doc = Document()
doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(10)


def title(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)


def note(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)


def fill(table, headers, rows):
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        for para in c.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_i, row in enumerate(rows, start=1):
        for c_i, val in enumerate(row):
            c = table.rows[r_i].cells[c_i]
            c.text = val
            for para in c.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)


title(f"Table 1. Method-comparison metrics for the four static scenarios "
      f"(n = {P['n_static']} paired measurements each).")
h1 = ["Scenario", "Mean bias (mmHg)", "95% LoA (mmHg)", "PE (%)",
      "BA reg. slope (95% CI)", "Deming slope (95% CI)",
      "P-B slope (95% CI)", "CCC", "C_b", "v"]
scen = {"S1_offset_only": "S1 offset only",
        "S2_zeroed_ideal": "S2 zeroed (ideal)",
        "S3_gain_uncompensated": "S3 gain error",
        "S4_gain_masked": "S4 gain masked"}
rows1 = []
for k, label in scen.items():
    d = st[k]
    rows1.append([label, signed(d["bias"]),
                  f"{d['loa_lower']:.1f} to {d['loa_upper']:.1f}", f1(d["pe"]),
                  f"{d['prop_slope']:.3f} ({d['prop_slope_lo']:.3f}, {d['prop_slope_hi']:.3f})",
                  f"{d['deming_slope']:.3f} ({d['deming_lo']:.3f}, {d['deming_hi']:.3f})",
                  f"{d['pb_slope']:.3f} ({d['pb_lo']:.3f}, {d['pb_hi']:.3f})",
                  f3(d["ccc"]), f3(d["C_b"]), f3(d["v"])])
fill(doc.add_table(rows=1 + len(rows1), cols=len(h1)), h1, rows1)
note("BA = Bland-Altman; CI = confidence interval; LoA = limits of agreement; "
     "PE = percentage error; P-B = Passing-Bablok; CCC = concordance "
     "correlation coefficient; C_b = bias-correction factor; v = scale shift.")

doc.add_paragraph()

title("Table 2. Dynamic-response metrics for pulse pressure (PP) and "
      "systolic pressure (SBP).")
h2 = ["System (f_n, zeta)", "PP ratio", "PP mean bias (mmHg)", "CCC (PP)",
      "v (PP)", "SBP mean bias (mmHg)"]
_ds = P["dyn_systems"]
dyn = {k: f"{k.capitalize()} ({f1(_ds[k]['fn'])} Hz, {f2(_ds[k]['zeta'])})"
       for k in ("optimal", "underdamped", "overdamped")}
rows2 = []
for k, label in dyn.items():
    pp = dy[f"{k}_pp"]; sbp = dy[f"{k}_sbp"]
    rows2.append([label, f2(pp["mean_ratio"]), signed(pp["bias"]),
                  f3(pp["ccc"]), f3(pp["v"]), signed(sbp["bias"])])
fill(doc.add_table(rows=1 + len(rows2), cols=len(h2)), h2, rows2)
note("PP ratio = mean(measured PP / true PP); f_n = natural frequency; "
     "zeta = damping coefficient.")

doc.add_paragraph()

title(f"Table 3. Real-waveform validation metrics for the four static "
      f"scenarios (n = {P['n_real_beats_per_scenario']:,} paired beats per "
      f"scenario from the VitalDB Open Dataset, {P['n_real_cases']} cases).")
h3 = ["Scenario", "n", "Mean bias (mmHg)", "95% LoA (mmHg)", "PE (%)",
      "BA reg. slope (95% CI)", "Deming slope (95% CI)",
      "P-B slope (95% CI)", "CCC", "C_b", "v"]
scen_real = {"R1_offset_only": "R1 offset only",
             "R2_zeroed_ideal": "R2 zeroed (ideal)",
             "R3_gain_uncompensated": "R3 gain error",
             "R4_gain_masked": "R4 gain masked"}
rows3 = []
for k, label in scen_real.items():
    d = re[k]
    rows3.append([label, str(d["n"]), signed(d["bias"]),
                  f"{d['loa_lower']:.1f} to {d['loa_upper']:.1f}", f1(d["pe"]),
                  f"{d['prop_slope']:.3f} ({d['prop_slope_lo']:.3f}, {d['prop_slope_hi']:.3f})",
                  f"{d['deming_slope']:.3f} ({d['deming_lo']:.3f}, {d['deming_hi']:.3f})",
                  f"{d['pb_slope']:.3f} ({d['pb_lo']:.3f}, {d['pb_hi']:.3f})",
                  f3(d["ccc"]), f3(d["C_b"]), f3(d["v"])])
fill(doc.add_table(rows=1 + len(rows3), cols=len(h3)), h3, rows3)
note("BA = Bland-Altman; CI = confidence interval; LoA = limits of agreement; "
     "PE = percentage error; P-B = Passing-Bablok; CCC = concordance "
     "correlation coefficient; C_b = bias-correction factor; v = scale shift.")

doc.save(OUTPATH)
print(f"Tables docx saved: {OUTPATH}")

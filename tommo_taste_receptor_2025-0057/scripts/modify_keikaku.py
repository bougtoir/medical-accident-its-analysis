"""
ToMMo 試料・情報分譲 研究番号2025-0057
研究計画書 (004_keikaku) 修正スクリプト

審査コメント対応:
1. 候補薬が具体的に記載されていない → 候補薬剤カテゴリを追記
2. 処方変更を調査票情報で捉えられるのか → 服薬情報の項目・2時点比較を明記
3. インターバルが適切か → 約5〜7年の観察期間を明記
4. 縦断的データを有する参加者に限定 → 選択基準に追記
5. GWAS → 候補遺伝子領域における関連解析に修正
6. サンプルサイズが記載されていない → 想定サンプルサイズを追記
"""

from docx import Document
from copy import deepcopy
from lxml import etree
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import sys
import os


def replace_para_text(para, old_text, new_text):
    """段落のテキスト全体を置換（最初のRunのフォーマットを保持）"""
    full_text = para.text
    if old_text not in full_text:
        return False
    for run in para.runs:
        run.text = ""
    new_full = full_text.replace(old_text, new_text)
    if para.runs:
        para.runs[0].text = new_full
    return True


def insert_para_after(doc, ref_para, text):
    """参照段落の直後に新しい段落を挿入（フォーマットをコピー）"""
    new_para = deepcopy(ref_para._element)
    for r in new_para.findall('.//' + qn('w:r')):
        r.getparent().remove(r)
    ref_para._element.addnext(new_para)
    r_elem = etree.SubElement(new_para, qn('w:r'))
    if ref_para.runs:
        rPr = ref_para.runs[0]._element.find(qn('w:rPr'))
        if rPr is not None:
            r_elem.insert(0, deepcopy(rPr))
    t_elem = etree.SubElement(r_elem, qn('w:t'))
    t_elem.text = text
    t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return Paragraph(new_para, doc)


def find_para_by_content(doc, search_text):
    """指定テキストを含む段落のインデックスを返す"""
    for i, para in enumerate(doc.paragraphs):
        if search_text in para.text:
            return i
    return -1


def modify_keikaku(input_path, output_path):
    doc = Document(input_path)

    # 1. 解析対象者：縦断データ要件を追加
    idx = find_para_by_content(doc, "解析対象者：選択基準はコホート調査登録時年齢が20歳以上")
    if idx >= 0:
        p = doc.paragraphs[idx]
        replace_para_text(
            p, p.text,
            "解析対象者：選択基準はコホート調査登録時年齢が20歳以上のものとし、"
            "ベースライン調査および第2段階調査の双方に参加した縦断的データを有する者に限定する。"
            "さらに、解析対象変数（後述）、目的変数（後述）または解析上重大な欠測のあるものを除外することとする。"
        )
        print(f"[1] P{idx}: 縦断データ要件を追加")

    # 2. 目的変数：服薬情報の捕捉方法とインターバルを明記
    idx = find_para_by_content(doc, "目的変数（薬剤感受性の代替指標）：調査票（生活）情報の治療")
    if idx >= 0:
        p = doc.paragraphs[idx]
        replace_para_text(
            p, p.text,
            "目的変数（薬剤感受性の代替指標）：調査票（生活）情報の服薬情報"
            "（使用中のお薬の商品名や成分等、入手方法〔医師の処方箋/それ以外〕、"
            "使用期間、使用頻度、1回の使用量）を用い、ベースライン調査時と第2段階調査時の"
            "2時点での比較により、処方変更を以下のとおり定義する。両調査の間隔は約5〜7年であり、"
            "慢性疾患に対する処方の変更を捕捉するのに十分な観察期間である。"
        )
        print(f"[2] P{idx}: 服薬情報の捕捉方法・インターバルを明記")

    # 3. GWAS → 候補遺伝子領域における関連解析
    idx = find_para_by_content(doc, "ゲノムワイド解析（GWAS）を追加し")
    if idx >= 0:
        p = doc.paragraphs[idx]
        new_text = p.text.replace(
            "データ量・計算資源が許す範囲でゲノムワイド解析（GWAS）を追加し、"
            "候補遺伝子解析で得た所見の補強および新規候補の探索を行う。",
            "候補遺伝子領域における関連解析を実施し、候補遺伝子解析で得た所見の補強を行う。"
            "なお、本研究では分譲される遺伝子領域が味蕾関連遺伝子群に限定されるため、"
            "ゲノムワイド解析（GWAS）ではなく、候補領域に絞った関連解析を行う。"
        )
        replace_para_text(p, p.text, new_text)
        print(f"[3] P{idx}: GWAS → 候補領域関連解析")

        # 4. サンプルサイズ推定を挿入（GWAS段落の直後）
        sample_text = (
            "想定されるサンプルサイズ：dbTMM 2026リリース"
            "（地域住民・三世代コホート 宮城 ベースライン・第2段階調査 120K）のうち、"
            "ベースライン調査時に調査票（生活）情報の服薬データを有する者は約39,000人、"
            "第2段階調査にも参加した縦断データ保有者は約30,000〜40,000人と見込まれる。"
            "そのうち、両時点で服薬情報が記録されている者は約20,000〜25,000人と推定される。"
            "先行研究に基づき、慢性疾患薬の用量変更・薬剤変更は5〜7年間で約20〜40%に生じると"
            "想定されるため、目的変数となる処方変更を有する解析対象者は約5,000〜10,000人と"
            "見込まれる。この規模は、味蕾関連遺伝子20領域の各SNPについて、"
            "対立遺伝子頻度5%以上の多型に対して検出力80%以上"
            "（有意水準5%、多重検定補正後）を確保するのに十分であると考えられる。"
        )
        insert_para_after(doc, p, sample_text)
        print(f"[4] P{idx}の後: サンプルサイズ推定を挿入")

    # 5. 候補薬剤カテゴリを挿入（P42「必要に応じて…」の直後）
    idx = find_para_by_content(doc, "必要に応じて、有害事象・検査値等の情報が利用可能な場合は")
    if idx >= 0:
        p = doc.paragraphs[idx]
        candidate_text = (
            "候補薬剤カテゴリ：本研究では、味覚受容体遺伝子多型との関連が想定される"
            "以下の薬剤カテゴリを主たる解析対象とする。"
            "①降圧薬（アムロジピン、ニフェジピン等のCa拮抗薬、"
            "カンデサルタン、バルサルタン、テルミサルタン等のARB）、"
            "②脂質異常症治療薬（アトルバスタチン、ロスバスタチン、プラバスタチン等のスタチン系薬剤）、"
            "③消炎鎮痛薬（ロキソプロフェン、ケトプロフェン等のNSAIDs）、"
            "④消化器用薬（レバミピド、ランソプラゾール、ラベプラゾール等）、"
            "⑤糖尿病治療薬（メトホルミン等）。"
            "これらは苦味を有する経口製剤が多く、TAS2R群を中心とした味覚受容体遺伝子多型が"
            "アドヒアランスや忍容性を介して処方変更に影響する可能性が先行研究1-3から示唆されている。"
            "交絡因子として年齢、性別、既往歴、併存疾患、生活習慣（飲酒・喫煙・食事等）、"
            "社会経済指標を調整するが、これらは上記候補薬剤の処方変更と関連しうる要因として"
            "臨床的に妥当である。"
        )
        insert_para_after(doc, p, candidate_text)
        print(f"[5] P{idx}の後: 候補薬剤カテゴリを挿入")

    doc.save(output_path)
    print(f"\n保存: {output_path}")


def modify_shinsei(input_path, output_path):
    doc = Document(input_path)

    idx = find_para_by_content(doc, "味蕾に関連する遺伝子多型が処方変更に与える影響を定量化し")
    if idx >= 0:
        p = doc.paragraphs[idx]
        old = p.text
        new = old.replace(
            "解析には、東北メディカル・メガバンク機構（ToMMo）が保有するゲノム情報"
            "および医療・生活習慣情報を用いる。",
            "解析には、東北メディカル・メガバンク機構（ToMMo）が保有するゲノム情報"
            "および医療・生活習慣情報を用い、ベースライン調査と第2段階調査の双方に参加した"
            "縦断的データを有する者を解析対象とする。調査票（生活）情報の服薬データの"
            "2時点比較により処方変更を定義する。"
        )
        replace_para_text(p, old, new)
        print(f"[1] P{idx}: 利用概要に縦断データ要件を追加")

    doc.save(output_path)
    print(f"\n保存: {output_path}")


if __name__ == "__main__":
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    orig = os.path.join(base, "original")
    rev = os.path.join(base, "revised")

    print("=== 004_keikaku 修正 ===")
    modify_keikaku(
        os.path.join(orig, "004_keikaku_2025-0057_20260313_提出版.docx"),
        os.path.join(rev, "004_keikaku_2025-0057_20260613_修正版.docx"),
    )

    print("\n=== 003_shinsei 修正 ===")
    modify_shinsei(
        os.path.join(orig, "003_shinsei_2025-0057_20260313_提出版.docx"),
        os.path.join(rev, "003_shinsei_2025-0057_20260613_修正版.docx"),
    )

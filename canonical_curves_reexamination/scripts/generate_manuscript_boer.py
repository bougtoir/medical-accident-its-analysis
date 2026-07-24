"""
Generate the manuscript as .docx (English) formatted for:
Bulletin of Economic Research (BOER) — Wiley

BOER format (based on published articles):
- Unstructured abstract
- JEL classification codes
- Keywords
- Sections: Introduction, [method sections], Results, Discussion, Conclusion
- Harvard-style (author-date) references
- Figures/tables inline with captions
- Correspondence line
- Data Availability Statement

Onishi T. 2026.
"""

import os
import sys
import json
import re
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIGURES_DIR = os.path.join(BASE_DIR, 'figures')
RESULTS_DIR = os.path.join(BASE_DIR, 'results')
DATA_DIR = os.path.join(BASE_DIR, 'data')


# ── Harvard-style references ──────────────────────────────────────────────
REFERENCES = [
    {
        "key": "Phillips1958",
        "text": "Phillips, A. W. (1958) 'The relation between unemployment and the rate "
                "of change of money wage rates in the United Kingdom, 1861-1957', "
                "Economica, 25(100), pp. 283-299."
    },
    {
        "key": "Kuznets1955",
        "text": "Kuznets, S. (1955) 'Economic growth and income inequality', "
                "American Economic Review, 45(1), pp. 1-28."
    },
    {
        "key": "Preston1975",
        "text": "Preston, S. H. (1975) 'The changing relation between mortality and "
                "level of economic development', Population Studies, 29(2), pp. 231-248."
    },
    {
        "key": "Akaike1974",
        "text": "Akaike, H. (1974) 'A new look at the statistical model identification', "
                "IEEE Transactions on Automatic Control, 19(6), pp. 716-723."
    },
    {
        "key": "Schwarz1978",
        "text": "Schwarz, G. (1978) 'Estimating the dimension of a model', "
                "Annals of Statistics, 6(2), pp. 461-464."
    },
    {
        "key": "GrossmanKrueger1991",
        "text": "Grossman, G. M. and Krueger, A. B. (1991) 'Environmental impacts of a "
                "North American free trade agreement', NBER Working Paper 3914."
    },
    {
        "key": "Stern2004",
        "text": "Stern, D. I. (2004) 'The rise and fall of the environmental Kuznets curve', "
                "World Development, 32(8), pp. 1419-1439."
    },
    {
        "key": "KruegerMueller2002",
        "text": "Krueger, J. and Mueller, R. A. (2002) 'Unskilled, unaware, or both? "
                "The better-than-average heuristic and statistical regression predict "
                "errors in estimates of own performance', Journal of Personality and "
                "Social Psychology, 82(2), pp. 180-188."
    },
    {
        "key": "OSC2015",
        "text": "Open Science Collaboration (2015) 'Estimating the reproducibility of "
                "psychological science', Science, 349(6251), aac4716."
    },
    {
        "key": "Ioannidis2005",
        "text": "Ioannidis, J. P. A. (2005) 'Why most published research findings are false', "
                "PLoS Medicine, 2(8), e124."
    },
    {
        "key": "Laffer2004",
        "text": "Laffer, A. B. (2004) 'The Laffer Curve: past, present, and future', "
                "Heritage Foundation Backgrounder, 1765."
    },
    {
        "key": "Lipset1959",
        "text": "Lipset, S. M. (1959) 'Some social requisites of democracy: economic "
                "development and political legitimacy', American Political Science Review, "
                "53(1), pp. 69-105."
    },
    {
        "key": "BurnhamAnderson2002",
        "text": "Burnham, K. P. and Anderson, D. R. (2002) Model Selection and Multimodel "
                "Inference: A Practical Information-Theoretic Approach, 2nd edn. New York: Springer."
    },
    {
        "key": "Cook1977",
        "text": "Cook, R. D. (1977) 'Detection of influential observation in linear regression', "
                "Technometrics, 19(1), pp. 15-18."
    },
    {
        "key": "Deaton2013",
        "text": "Deaton, A. (2013) The Great Escape: Health, Wealth, and the Origins "
                "of Inequality. Princeton: Princeton University Press."
    },
    {
        "key": "Easterlin1974",
        "text": "Easterlin, R. A. (1974) 'Does economic growth improve the human lot? "
                "Some empirical evidence', in David, P. A. and Reder, M. W. (eds) "
                "Nations and Households in Economic Growth. New York: Academic Press, "
                "pp. 89-125."
    },
    {
        "key": "YerkesDodson1908",
        "text": "Yerkes, R. M. and Dodson, J. D. (1908) 'The relation of strength of "
                "stimulus to rapidity of habit-formation', Journal of Comparative "
                "Neurology and Psychology, 18(5), pp. 459-482."
    },
    {
        "key": "Ebbinghaus1885",
        "text": "Ebbinghaus, H. (1885) \u00dcber das Ged\u00e4chtnis: Untersuchungen zur "
                "experimentellen Psychologie. Leipzig: Duncker & Humblot."
    },
    {
        "key": "DunningKruger1999",
        "text": "Dunning, D. and Kruger, J. (1999) 'Unskilled and unaware of it: how "
                "difficulties in recognizing one's own incompetence lead to inflated "
                "self-assessments', Journal of Personality and Social Psychology, "
                "77(6), pp. 1121-1134."
    },
    {
        "key": "BlanchardKatz1997",
        "text": "Blanchard, O. and Katz, L. F. (1997) 'What we know and do not know about "
                "the natural rate of unemployment', Journal of Economic Perspectives, "
                "11(1), pp. 51-72."
    },
    {
        "key": "PikettySaez2003",
        "text": "Piketty, T. and Saez, E. (2003) 'Income inequality in the United States, "
                "1913-1998', Quarterly Journal of Economics, 118(1), pp. 1-41."
    },
]


def load_results():
    with open(os.path.join(RESULTS_DIR, 'full_results.json'), 'r') as f:
        return json.load(f)


def get_domain_summary(results):
    from collections import Counter
    domains = {}
    for r in results:
        cat = r['category']
        v = r['verdict']['verdict']
        if cat not in domains:
            domains[cat] = {'total': 0, 'verdicts': Counter(), 'curves': []}
        domains[cat]['total'] += 1
        domains[cat]['verdicts'][v] += 1
        domains[cat]['curves'].append(r)
    return domains


def get_curve_result(results, name):
    for r in results:
        if r['name'] == name:
            return r
    return None


def fmt_p(p_val):
    if p_val < 0.0001:
        return f"{p_val:.1e}"
    return f"{p_val:.3f}"


def add_text(paragraph, text, size=Pt(12)):
    run = paragraph.add_run(text)
    run.font.size = size
    return run


def add_bold(paragraph, text, size=Pt(12)):
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = size
    return run


def add_italic(paragraph, text, size=Pt(12)):
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = size
    return run


def set_line_spacing(paragraph, spacing=1.5):
    from docx.shared import Pt as _Pt
    paragraph.paragraph_format.line_spacing = spacing


def create_manuscript():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    results = load_results()
    verdicts = [r['verdict']['verdict'] for r in results]
    n_ns = verdicts.count('NOT_SIGNIFICANT')
    n_outlier = verdicts.count('OUTLIER_DEPENDENT')
    n_robust = verdicts.count('ROBUST_NONLINEAR')
    n_overfit = verdicts.count('OVERFITTING')
    n_fail = n_ns + n_outlier + n_overfit
    domains = get_domain_summary(results)

    # ═══════ TITLE PAGE ═══════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(p, 2.0)
    run = p.add_run(
        'Fragility of Canonical Curves: A Cross-Disciplinary Audit of '
        '52 Established Nonlinear Relationships')
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(p, 2.0)
    run = p.add_run('Tatsuki Onishi')
    run.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(p, 2.0)
    run = p.add_run('[Institutional affiliation]')
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(p, 2.0)
    run = p.add_run('Correspondence: Tatsuki Onishi ([email])')
    run.font.size = Pt(10)
    run.italic = True

    doc.add_page_break()

    # ═══════ ABSTRACT ═══════
    h = doc.add_heading('ABSTRACT', level=1)

    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_text(p,
        f"This paper systematically re-examines 52 canonical nonlinear relationships "
        f"across eight academic disciplines using a uniform four-test framework: "
        f"nested F-tests, Akaike and Bayesian Information Criteria (AIC/BIC), "
        f"leave-one-out cross-validation (LOOCV), and Cook's distance sensitivity analysis. "
        f"Of 52 curves, only {n_robust} ({100*n_robust/52:.0f}%) demonstrate robust "
        f"nonlinearity surviving all tests. {n_ns} ({100*n_ns/52:.0f}%) show no "
        f"statistically significant nonlinearity, and {n_outlier} ({100*n_outlier/52:.0f}%) "
        f"are outlier-dependent\u2014their significance vanishes after removing 1\u20133 "
        f"influential observations. Domain asymmetry is pronounced: public health and "
        f"psychology curves are substantially more robust than economics or political science "
        f"curves. Several policy-relevant curves (Laffer Curve, Environmental Kuznets Curve, "
        f"Lipset Hypothesis) are outlier-dependent. These findings suggest that approximately "
        f"two-thirds of textbook nonlinear relationships fail modern robustness tests, "
        f"urging caution when citing canonical curves as empirical support for nonlinear theories."
    )

    # Keywords
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    set_line_spacing(p, 2.0)
    add_bold(p, 'Keywords: ', Pt(12))
    add_text(p, 'model selection | nonlinearity | robustness | outlier dependence | '
                'information criteria | cross-validation | canonical relationships | '
                'meta-research', Pt(12))

    # JEL Classification
    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_bold(p, 'JEL Classification: ', Pt(12))
    add_text(p, 'C12, C21, C52, B41, I15', Pt(12))

    doc.add_page_break()

    # ═══════ 1. INTRODUCTION ═══════
    h = doc.add_heading('1. Introduction', level=1)

    intro_paras = [
        ("Curvilinear relationships occupy a privileged position in the social and natural "
         "sciences. From the Phillips Curve (Phillips, 1958) in macroeconomics to the Kuznets "
         "Curve (Kuznets, 1955) and the Preston Curve (Preston, 1975) in development "
         "economics and public health, these nonlinear functional forms are widely taught, "
         "frequently cited in policy documents, and treated as established empirical "
         "regularities. Yet many were originally established with limited data, rudimentary "
         "statistical methods, and before model selection criteria such as AIC (Akaike, 1974) "
         "and BIC (Schwarz, 1978) became standard practice."),

        ("The present study was motivated by a preliminary analysis of the Preston Curve "
         "conducted by the author, which revealed that the apparent concavity "
         "depends heavily on the position of the United States as a single influential "
         "observation\u2014removing the US raised the p-value for the quadratic term "
         "to 0.49, rendering the nonlinear term non-significant. This finding raised the "
         "question of whether similar fragilities lurk beneath other canonical curves. "
         "The Environmental Kuznets Curve for CO\u2082 emissions has been repeatedly "
         "challenged (Grossman and Krueger, 1991; Stern, 2004), and the Dunning-Kruger "
         "effect (Dunning and Kruger, 1999) has been argued to be a statistical artefact "
         "of regression to the mean (Krueger and Mueller, 2002)."),

        ("Despite individual critiques, no systematic cross-disciplinary audit has been "
         "conducted. The present study fills this gap by applying a uniform methodological "
         "framework to 52 canonical curves across eight academic disciplines: economics, "
         "public health, demography, environmental science, psychology, physics, political "
         "science, and agriculture. The meta-research literature (Ioannidis, 2005; Open "
         "Science Collaboration, 2015) has documented widespread concerns about the "
         "reproducibility of empirical findings; this study extends that concern to "
         "the specific domain of claimed nonlinearities."),

        ("The contribution of this paper is threefold. First, we provide the first systematic "
         "cross-disciplinary inventory of canonical curve fragilities. Second, we propose a "
         "transparent four-test framework that can be applied to any claimed bivariate "
         "nonlinear relationship. Third, we identify a striking domain asymmetry\u2014public "
         "health and psychology curves are substantially more robust than economics and "
         "political science curves\u2014which has implications for the epistemological status "
         "of empirical regularities across disciplines."),
    ]
    for text in intro_paras:
        p = doc.add_paragraph()
        set_line_spacing(p, 2.0)
        add_text(p, text)

    # ═══════ 2. METHODOLOGY ═══════
    h = doc.add_heading('2. Methodology', level=1)

    # 2.1 Curve selection
    doc.add_heading('2.1. Curve selection', level=2)
    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_text(p,
        "Candidate curves were identified through systematic review of named laws and "
        "stylized facts in major handbooks, citation analysis of papers with 'curve,' 'law,' "
        "or 'paradox' in titles, and review of the replication crisis literature "
        "(Ioannidis, 2005; Open Science Collaboration, 2015). This yielded 78 candidates, "
        "from which 52 met all inclusion criteria."
    )

    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_bold(p, 'Inclusion criteria: ')
    add_text(p, "(1) eponymous/canonical status (\u22652 textbooks or >500 citations); "
              "(2) explicit nonlinearity claim; (3) bivariate testability; "
              "(4) publicly available data (N \u2265 10); (5) policy or theoretical relevance "
              "of curve shape.")

    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_bold(p, 'Exclusion criteria: ')
    add_text(p, "definitional/tautological curves; proprietary data requirements; "
              "purely temporal dynamics; formally retracted relationships; "
              "curves requiring multivariate specification.")

    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_text(p,
        "Curves were stratified across eight disciplines: economics (12), public health (10), "
        "demography (6), environmental science (6), psychology (5), physics (4), political "
        "science (5), and agriculture (4). Five types of claimed nonlinearity were represented: "
        "inverted-U (14), U/J-shaped (8), concave/saturating (12), power-law (8), "
        "and S-shaped/structural-break (10)."
    )

    # 2.2 Data sources
    doc.add_heading('2.2. Data sources', level=2)
    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_text(p,
        "Data were drawn primarily from the World Bank World Development Indicators (WDI) "
        "API for cross-country curves (GDP per capita PPP, life expectancy, Gini coefficient, "
        "total fertility rate, forest area), with US macroeconomic time series (unemployment, "
        "inflation, GDP growth) also obtained via WDI. Additional sources included OECD.Stat, "
        "published meta-analyses, USGS earthquake catalogues, and digitised original "
        "publication data."
    )

    # 2.3 Statistical framework
    doc.add_heading('2.3. Statistical framework', level=2)

    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_bold(p, 'Nested F-test. ')
    add_text(p,
        "For each curve, restricted (linear: y = a + bx) and unrestricted "
        "(quadratic: y = a + bx + cx\u00b2) models were fitted via ordinary least squares. "
        "The F-statistic tested the significance of the additional parameter."
    )

    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_bold(p, 'Information criteria. ')
    add_text(p,
        "AIC and BIC (Burnham and Anderson, 2002) were computed for linear, quadratic, "
        "and logarithmic models. The model with lowest criterion value was selected."
    )

    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_bold(p, 'Leave-one-out cross-validation. ')
    add_text(p,
        "LOOCV root mean squared error (RMSE) was computed for linear and quadratic "
        "models to assess out-of-sample predictive accuracy."
    )

    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_bold(p, "Cook's distance sensitivity analysis. ")
    add_text(p,
        "The top 3 most influential observations (Cook, 1977) were removed "
        "and the F-test repeated. A curve was classified as outlier-dependent if "
        "significance (p < 0.05) was lost after removal."
    )

    # 2.4 Verdict classification
    doc.add_heading('2.4. Verdict classification', level=2)
    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_text(p, "Each curve received one of four verdicts:")

    verdict_defs = [
        "ROBUST NONLINEAR: significant with full data AND after outlier removal, "
        "with quadratic LOOCV RMSE \u2264 linear RMSE.",
        "OUTLIER-DEPENDENT: significant with full data but non-significant after "
        "removing top 3 influential points.",
        "NOT SIGNIFICANT: non-significant even with full data (p \u2265 0.05).",
        "OVERFITTING: significant but LOOCV RMSE is worse for the quadratic model.",
    ]
    for v in verdict_defs:
        p = doc.add_paragraph(v, style='List Bullet')
        set_line_spacing(p, 2.0)

    # ═══════ 3. RESULTS ═══════
    h = doc.add_heading('3. Results', level=1)

    # 3.1 Overall distribution
    doc.add_heading('3.1. Overall distribution of verdicts', level=2)
    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_text(p,
        f"Of 52 canonical curves, {n_robust} ({100*n_robust/52:.0f}%) demonstrated robust "
        f"nonlinearity, {n_outlier} ({100*n_outlier/52:.0f}%) were outlier-dependent, "
        f"{n_ns} ({100*n_ns/52:.0f}%) showed no significant nonlinearity, and "
        f"{n_overfit} ({100*n_overfit/52:.0f}%) exhibited overfitting (Figure 1). "
        f"Approximately two-thirds of textbook nonlinear relationships either failed to "
        f"reach significance or were driven by a small number of influential observations."
    )

    # Figure 1
    doc.add_paragraph()
    fig1_path = os.path.join(FIGURES_DIR, 'fig1_verdict_distribution.png')
    if os.path.exists(fig1_path):
        doc.add_picture(fig1_path, width=Inches(5.5))
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        set_line_spacing(p, 2.0)
        add_bold(p, 'Figure 1. ', Pt(11))
        add_text(p, 'Distribution of verdicts across 52 canonical curves, '
                    'stratified by domain.', Pt(11))

    # 3.2 Domain-level results
    doc.add_heading('3.2. Results by domain', level=2)

    # Economics
    econ = domains.get('Economics', {})
    econ_v = econ.get('verdicts', {})
    phillips = get_curve_result(results, 'Phillips Curve')
    kuznets = get_curve_result(results, 'Kuznets Curve')
    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_bold(p, 'Economics. ')
    add_text(p,
        f"Of 12 curves, {econ_v.get('NOT_SIGNIFICANT', 0)} were non-significant, "
        f"{econ_v.get('OUTLIER_DEPENDENT', 0)} were outlier-dependent, and "
        f"{econ_v.get('ROBUST_NONLINEAR', 0)} demonstrated robust nonlinearity. "
        f"The Phillips Curve (Phillips, 1958; Blanchard and Katz, 1997) showed no "
        f"significant quadratic term (p = {fmt_p(phillips['f_test']['p_value'])}). "
        f"The Kuznets Curve's inverted-U (Kuznets, 1955; Piketty and Saez, 2003) was "
        f"not significant with {kuznets['n']} countries "
        f"(p = {fmt_p(kuznets['f_test']['p_value'])}). "
        "The Environmental Kuznets Curve (CO\u2082), Laffer Curve, and Great Gatsby Curve "
        "were all outlier-dependent."
    )

    # Public Health
    health = domains.get('Public Health', {})
    health_v = health.get('verdicts', {})
    preston = get_curve_result(results, 'Preston Curve')
    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_bold(p, 'Public health. ')
    preston_desc = (
        f"The Preston Curve (Preston, 1975) with real World Bank data (N = {preston['n']}) "
        f"showed no significant nonlinearity (p = {fmt_p(preston['f_test']['p_value'])}), "
        "consistent with a log-linear relationship (Deaton, 2013)"
    ) if preston['verdict']['verdict'] == 'NOT_SIGNIFICANT' else (
        f"The Preston Curve showed {preston['verdict']['verdict'].lower().replace('_', ' ')}"
    )
    add_text(p,
        f"Public health curves showed the highest robustness rate "
        f"({health_v.get('ROBUST_NONLINEAR', 0)}/{health.get('total', 10)} robust). "
        f"The BMI-Mortality J-curve, Alcohol-Mortality J-curve, and Barker Hypothesis "
        f"U-shape all survived rigorous testing. {preston_desc}."
    )

    # Demography
    demo = domains.get('Demography', {})
    demo_v = demo.get('verdicts', {})
    dt = get_curve_result(results, 'Demographic Transition (TFR)')
    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_bold(p, 'Demography. ')
    dt_desc = (
        f"The Demographic Transition model (N = {dt['n']}) was classified as "
        f"{dt['verdict']['verdict'].lower().replace('_', ' ')} "
        f"(p = {fmt_p(dt['f_test']['p_value'])})"
    )
    add_text(p,
        f"Of 6 curves, {demo_v.get('ROBUST_NONLINEAR', 0)} showed robust nonlinearity, "
        f"{demo_v.get('NOT_SIGNIFICANT', 0)} were non-significant, and "
        f"{demo_v.get('OUTLIER_DEPENDENT', 0)} were outlier-dependent. "
        f"{dt_desc}. The Lee-Carter mortality model and Coale-Trussell fertility schedule "
        f"demonstrated robust nonlinearity."
    )

    # Environmental Science
    env = domains.get('Environmental Science', {})
    env_v = env.get('verdicts', {})
    forest = get_curve_result(results, 'Forest Transition Curve')
    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_bold(p, 'Environmental science. ')
    add_text(p,
        f"Of 6 curves, {env_v.get('ROBUST_NONLINEAR', 0)} showed robust nonlinearity "
        f"(Keeling Curve), {env_v.get('NOT_SIGNIFICANT', 0)} were non-significant, and "
        f"{env_v.get('OUTLIER_DEPENDENT', 0)} were outlier-dependent. "
        f"The Forest Transition Curve with real data (N = {forest['n']}) was "
        f"non-significant (p = {fmt_p(forest['f_test']['p_value'])})."
    )

    # Psychology
    psych = domains.get('Psychology', {})
    psych_v = psych.get('verdicts', {})
    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_bold(p, 'Psychology. ')
    add_text(p,
        f"Psychology showed high robustness: {psych_v.get('ROBUST_NONLINEAR', 0)}/5 "
        f"curves were robust (Yerkes-Dodson (Yerkes and Dodson, 1908), Ebbinghaus "
        f"forgetting curve (Ebbinghaus, 1885), Dunning-Kruger (Dunning and Kruger, 1999), "
        f"and the Happiness U-Curve (Easterlin, 1974)). "
        f"Only the Weber-Fechner Law was non-significant."
    )

    # Physics
    phys = domains.get('Physics', {})
    phys_v = phys.get('verdicts', {})
    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_bold(p, 'Physics. ')
    add_text(p,
        f"All 4 physics curves failed: {phys_v.get('NOT_SIGNIFICANT', 0)} non-significant, "
        f"{phys_v.get('OUTLIER_DEPENDENT', 0)} outlier-dependent. "
        "Hubble's Law and Gutenberg-Richter were outlier-dependent; "
        "Kleiber's Law and Moore's Law were non-significant in log space."
    )

    # Political Science
    pol = domains.get('Political Science', {})
    pol_v = pol.get('verdicts', {})
    lipset = get_curve_result(results, 'Lipset Hypothesis')
    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_bold(p, 'Political science. ')
    add_text(p,
        f"The Lipset Hypothesis (Lipset, 1959) was outlier-dependent "
        f"(p = {fmt_p(lipset['f_test']['p_value'])} full, "
        f"p = {fmt_p(lipset['sensitivity']['p_clean'])} after removal of high-income outliers). "
        f"{pol_v.get('NOT_SIGNIFICANT', 0)}/5 were non-significant."
    )

    # Agriculture
    agr = domains.get('Agriculture', {})
    agr_v = agr.get('verdicts', {})
    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_bold(p, 'Agriculture. ')
    add_text(p,
        f"{agr_v.get('ROBUST_NONLINEAR', 0)}/4 robust (Mitscherlich yield, "
        f"Micronutrient U-shape); {agr_v.get('NOT_SIGNIFICANT', 0)}/4 non-significant."
    )

    # 3.3 Sensitivity analysis figure
    doc.add_heading('3.3. Sensitivity analysis', level=2)
    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_text(p,
        "Figure 2 displays p-values before and after outlier removal. Points in the "
        "upper-left quadrant (significant full data, non-significant after removal) "
        "represent outlier-dependent curves, concentrated in economics and political science."
    )
    fig2_path = os.path.join(FIGURES_DIR, 'fig2_sensitivity_analysis.png')
    if os.path.exists(fig2_path):
        doc.add_paragraph()
        doc.add_picture(fig2_path, width=Inches(5.0))
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        set_line_spacing(p, 2.0)
        add_bold(p, 'Figure 2. ', Pt(11))
        add_text(p, "Sensitivity of F-test p-values to outlier removal "
                    "(Cook's distance top 3).", Pt(11))

    # 3.4 Model comparison figures
    doc.add_heading('3.4. Model comparison', level=2)
    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_text(p,
        "Figure 3 compares AIC, BIC, and LOOCV model selection outcomes. BIC is more "
        "conservative than AIC, selecting the linear model more frequently, reflecting "
        "stronger complexity penalties (Burnham and Anderson, 2002)."
    )
    fig3_path = os.path.join(FIGURES_DIR, 'fig3_model_comparison.png')
    if os.path.exists(fig3_path):
        doc.add_paragraph()
        doc.add_picture(fig3_path, width=Inches(5.0))
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        set_line_spacing(p, 2.0)
        add_bold(p, 'Figure 3. ', Pt(11))
        add_text(p, "Model comparison results across 52 curves (AIC, BIC, LOOCV).", Pt(11))

    fig4_path = os.path.join(FIGURES_DIR, 'fig4_loocv_comparison.png')
    if os.path.exists(fig4_path):
        p = doc.add_paragraph()
        set_line_spacing(p, 2.0)
        add_text(p,
            "Figure 4 shows the LOOCV RMSE comparison between linear and quadratic models."
        )
        doc.add_paragraph()
        doc.add_picture(fig4_path, width=Inches(5.0))
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        set_line_spacing(p, 2.0)
        add_bold(p, 'Figure 4. ', Pt(11))
        add_text(p, "LOOCV RMSE: linear vs. quadratic model.", Pt(11))

    # 3.5 Sample size
    fig5_path = os.path.join(FIGURES_DIR, 'fig5_sample_size.png')
    if os.path.exists(fig5_path):
        doc.add_heading('3.5. Sample size and verdict', level=2)
        p = doc.add_paragraph()
        set_line_spacing(p, 2.0)
        add_text(p,
            "Figure 5 shows the relationship between sample size and verdict. "
            "Larger samples do not systematically yield more robust verdicts."
        )
        doc.add_paragraph()
        doc.add_picture(fig5_path, width=Inches(5.0))
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        set_line_spacing(p, 2.0)
        add_bold(p, 'Figure 5. ', Pt(11))
        add_text(p, "Sample size vs. verdict classification.", Pt(11))

    # 3.6 Summary table
    doc.add_heading('3.6. Summary of all results', level=2)
    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_text(p, "Table 1 presents complete results for all 52 curves.")

    df = pd.read_csv(os.path.join(RESULTS_DIR, 'summary_table.csv'))
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['#', 'Curve', 'N', 'p (full)', 'p (clean)', 'BIC best', 'Verdict']
    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h_text
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)

    for idx, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx + 1)
        row_cells[1].text = str(row['Curve'])[:28]
        row_cells[2].text = str(row['N'])
        p_full = row['p (full)']
        row_cells[3].text = f"{p_full:.4f}" if p_full > 0.0001 else f"{p_full:.1e}"
        p_clean = row['p (clean)']
        row_cells[4].text = f"{p_clean:.4f}" if p_clean > 0.0001 else f"{p_clean:.1e}"
        row_cells[5].text = str(row['BIC best'])
        row_cells[6].text = str(row['Verdict']).replace('_', ' ')
        for cell in row_cells:
            for para in cell.paragraphs:
                for run_obj in para.runs:
                    run_obj.font.size = Pt(7)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    set_line_spacing(p, 2.0)
    add_bold(p, 'Table 1. ', Pt(10))
    add_text(p, "Summary of re-examination results. p (full) = nested F-test with all data; "
                "p (clean) = after removing top 3 Cook's distance points.", Pt(10))

    # ═══════ 4. DISCUSSION ═══════
    h = doc.add_heading('4. Discussion', level=1)

    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_text(p,
        f"This systematic audit reveals that {100*n_fail/52:.0f}% of established nonlinear "
        f"relationships fail at least one modern robustness test. Five cross-cutting patterns "
        f"emerge."
    )

    patterns = [
        "Outlier-driven nonlinearity is the most common failure mode. In cross-country "
        "analyses, 1\u20133 geopolitically distinctive nations (oil states, the US) drive "
        "curvature. This finding echoes Deaton's (2013) warning about the fragility of "
        "cross-country regressions.",

        "Domain asymmetry: public health and psychology curves are substantially more "
        "robust than economics or political science curves. This likely reflects the "
        "difference between mechanistic relationships (biological dose-response) and "
        "contingent relationships (policy-mediated associations).",

        "Time-series vs. cross-section: time-series curves tend to be more robust, "
        "likely because they are less vulnerable to compositional effects.",

        "Log transformation resolves apparent nonlinearity: in many cases a simple log "
        "transformation produces a linear relationship, suggesting the 'canonical curve' "
        "is linear on the wrong scale.",

        "BIC is more conservative than AIC: BIC selects the linear model more frequently, "
        "reflecting stronger complexity penalties, consistent with Burnham and Anderson (2002).",
    ]
    for i, pat in enumerate(patterns, 1):
        p = doc.add_paragraph()
        set_line_spacing(p, 2.0)
        add_text(p, f"{i}. {pat}")

    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_text(p,
        "Several outlier-dependent curves have direct policy implications. The Laffer Curve "
        "(Laffer, 2004) is used to justify tax reductions; the Environmental Kuznets Curve "
        "(Grossman and Krueger, 1991) is cited to argue that growth resolves pollution; the "
        "Lipset Hypothesis (Lipset, 1959) underpins modernization theory. Our findings "
        "suggest that the empirical bases of these policy-relevant claims are more fragile "
        "than commonly assumed."
    )

    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_text(p,
        "The broader implication is epistemic. Economics and political science rely heavily "
        "on cross-country regressions to establish stylised facts. Our results suggest "
        "that cross-country nonlinearities are especially fragile, whereas within-discipline "
        "time-series relationships (e.g. Keeling Curve) or relationships grounded in "
        "biological mechanisms (e.g. BMI-Mortality J-curve) are more resilient. This is "
        "consistent with the view that mechanistic relationships are more likely to be "
        "robust than statistical associations driven by compositional heterogeneity."
    )

    # ═══════ 5. LIMITATIONS ═══════
    h = doc.add_heading('5. Limitations', level=1)

    limitations = [
        "First, our analysis is restricted to bivariate relationships; many canonical curves "
        "may be better specified in multivariate settings. This is by design\u2014the "
        "bivariate form is how these curves are most commonly cited and taught.",

        "Second, we use a uniform quadratic alternative, whereas some curves posit specific "
        "functional forms (power laws, logistic functions). Future work should incorporate "
        "curve-specific alternatives.",

        "Third, although 8 curves use real World Bank API data (N = 31\u2013247), the "
        "remaining 44 curves use representative or published data with smaller sample sizes.",

        "Fourth, we test only the significance of curvature, not the existence of any "
        "relationship. A curve classified as non-significant may still have a significant "
        "linear component.",

        "Finally, for some curves (psychology, demography), we rely on aggregate or "
        "meta-analytic data rather than individual-level microdata.",
    ]
    for lim in limitations:
        p = doc.add_paragraph()
        set_line_spacing(p, 2.0)
        add_text(p, lim)

    # ═══════ 6. CONCLUSION ═══════
    h = doc.add_heading('6. Conclusion', level=1)

    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_text(p,
        f"This systematic re-examination of 52 canonical curves reveals that only "
        f"{100*n_robust/52:.0f}% demonstrate nonlinearity that is statistically significant, "
        f"survives outlier removal, and shows superior out-of-sample prediction. "
        f"Researchers should routinely report sensitivity analyses and model comparison "
        f"criteria when invoking canonical curves. Policymakers should be cautious about "
        f"interventions premised on specific curve shapes\u2014particularly the Laffer Curve, "
        f"Environmental Kuznets Curve, and Lipset Hypothesis\u2014whose empirical bases are "
        f"outlier-dependent."
    )

    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_text(p,
        "We propose that any study claiming a nonlinear relationship should, at minimum, "
        "report: (1) the p-value of the nested F-test; (2) AIC/BIC model comparison including "
        "a linear alternative; (3) out-of-sample prediction accuracy; and (4) sensitivity to "
        "removal of the most influential observations. This four-test transparency norm would "
        "substantially reduce the risk of over-interpreting fragile curves."
    )

    # ═══════ DATA AVAILABILITY STATEMENT ═══════
    doc.add_page_break()
    h = doc.add_heading('Data Availability Statement', level=1)
    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_text(p,
        "All analysis code, data, and results are publicly available at "
        "[GitHub repository URL]. World Bank data were accessed via the wbgapi Python "
        "library (World Development Indicators API). The complete source metadata "
        "for all 52 curves is provided in the online appendix."
    )

    # ═══════ ACKNOWLEDGEMENTS ═══════
    h = doc.add_heading('Acknowledgements', level=1)
    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_text(p, "[To be completed]")

    # ═══════ CONFLICT OF INTEREST ═══════
    h = doc.add_heading('Conflict of Interest', level=1)
    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_text(p, "The author declares no conflict of interest.")

    # ═══════ FUNDING ═══════
    h = doc.add_heading('Funding', level=1)
    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_text(p, "[To be completed]")

    # ═══════ REFERENCES ═══════
    doc.add_page_break()
    h = doc.add_heading('References', level=1)

    sorted_refs = sorted(REFERENCES, key=lambda r: r['text'].split(',')[0].split('(')[0].strip())
    for ref in sorted_refs:
        p = doc.add_paragraph()
        set_line_spacing(p, 2.0)
        add_text(p, ref['text'], Pt(11))
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)

    # ═══════ APPENDIX: Source metadata table ═══════
    doc.add_page_break()
    h = doc.add_heading('Appendix: Source Metadata Table', level=1)
    p = doc.add_paragraph()
    set_line_spacing(p, 2.0)
    add_text(p,
        "Complete data source information for all 52 curves, including original publication, "
        "data source, sample size, variable definitions, and claimed functional form."
    )

    source_meta_path = os.path.join(DATA_DIR, 'source_metadata.json')
    if os.path.exists(source_meta_path):
        with open(source_meta_path, 'r', encoding='utf-8') as f:
            sources = json.load(f)

        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ['#', 'Curve', 'Original paper', 'Data source', 'N', 'Claimed form']
        for i, h_text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h_text
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(7)

        for src in sources:
            row_cells = table.add_row().cells
            row_cells[0].text = str(src['id'])
            row_cells[1].text = src['name'][:25]
            row_cells[2].text = src['original_paper'][:40]
            row_cells[3].text = src['data_source'][:45]
            n_val = src.get('current_n')
            row_cells[4].text = str(n_val) if n_val else 'WB API'
            row_cells[5].text = src['claimed_form'][:30]
            for cell in row_cells:
                for para in cell.paragraphs:
                    for run_obj in para.runs:
                        run_obj.font.size = Pt(6.5)

    # Save
    output_path = os.path.join(BASE_DIR, 'manuscript_canonical_curves_boer.docx')
    doc.save(output_path)
    print(f"Manuscript saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_manuscript()

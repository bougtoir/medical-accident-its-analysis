"""
Generate cover letter for Bulletin of Economic Research (BOER) submission.
Emphasises the systematic survey + re-examination dual nature of the study,
policy relevance of outlier-dependent curves, and cross-disciplinary scope.

Onishi T. 2026.
"""

import os
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RESULTS_DIR = os.path.join(BASE_DIR, 'results')


def load_results():
    with open(os.path.join(RESULTS_DIR, 'full_results.json'), 'r') as f:
        return json.load(f)


def create_cover_letter():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    results = load_results()
    verdicts = [r['verdict']['verdict'] for r in results]
    n_robust = verdicts.count('ROBUST_NONLINEAR')
    n_outlier = verdicts.count('OUTLIER_DEPENDENT')
    n_ns = verdicts.count('NOT_SIGNIFICANT')

    # Date
    p = doc.add_paragraph()
    p.add_run('[Date]').font.size = Pt(12)

    doc.add_paragraph()

    # Addressee
    lines = [
        'Professor Andrew Pickering',
        'Editor-in-Chief',
        'Bulletin of Economic Research',
    ]
    for line in lines:
        p = doc.add_paragraph()
        p.add_run(line).font.size = Pt(12)
        p.paragraph_format.space_after = Pt(0)

    doc.add_paragraph()

    # Subject
    p = doc.add_paragraph()
    run = p.add_run('Re: Submission of manuscript ')
    run.font.size = Pt(12)
    run = p.add_run(
        '"Fragility of Canonical Curves: A Cross-Disciplinary Audit of '
        '52 Established Nonlinear Relationships"')
    run.font.size = Pt(12)
    run.italic = True

    doc.add_paragraph()

    # Body
    body_paras = [
        "Dear Professor Pickering,",

        "I am pleased to submit the above manuscript for consideration as a Research "
        "Article in the Bulletin of Economic Research.",

        (f"This paper presents the first systematic cross-disciplinary audit of canonical "
         f"nonlinear relationships. Applying a uniform four-test framework (nested F-tests, "
         f"AIC/BIC model selection, leave-one-out cross-validation, and Cook's distance "
         f"sensitivity analysis) to 52 established curves across eight disciplines, I find "
         f"that only {n_robust} of 52 ({100*n_robust/52:.0f}%) demonstrate robust "
         f"nonlinearity surviving all tests. {n_outlier} ({100*n_outlier/52:.0f}%) are "
         f"outlier-dependent\u2014their statistical significance vanishes after removing "
         f"just 1\u20133 influential observations\u2014and {n_ns} ({100*n_ns/52:.0f}%) show "
         f"no significant nonlinearity at all."),

        "Three findings are of particular relevance to the Bulletin's readership. "
        "First, several policy-relevant economic curves\u2014the Laffer Curve, the "
        "Environmental Kuznets Curve, and the Great Gatsby Curve\u2014are outlier-dependent, "
        "meaning that the empirical bases for influential policy claims rest on a handful of "
        "distinctive observations. Second, there is a striking domain asymmetry: economics "
        "and political science curves are substantially less robust than public health and "
        "psychology curves, raising questions about the epistemological status of empirical "
        "regularities in the social sciences. Third, the four-test framework itself offers "
        "a practical transparency norm that can be adopted by any researcher claiming a "
        "nonlinear relationship.",

        "The paper bridges survey and empirical research. It provides a systematic overview "
        "of how canonical curves were originally established and how they fare under modern "
        "scrutiny, while also offering new empirical results from a uniform re-analysis. "
        "I believe this combination of methodological review and original evidence aligns "
        "well with the Bulletin's scope, which encompasses both surveys and empirical "
        "contributions across all areas of economics.",

        "This manuscript is not under consideration elsewhere. All data and code are "
        "publicly available. I confirm that I have no conflicts of interest.",

        "Thank you for considering this submission. I look forward to your response.",

        "Yours sincerely,",
    ]

    for text in body_paras:
        p = doc.add_paragraph()
        p.add_run(text).font.size = Pt(12)
        p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Tatsuki Onishi')
    run.font.size = Pt(12)
    run.bold = True

    p = doc.add_paragraph()
    p.add_run('[Institutional affiliation]').font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run('[Email address]').font.size = Pt(11)

    output_path = os.path.join(BASE_DIR, 'cover_letter_boer.docx')
    doc.save(output_path)
    print(f"Cover letter saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_cover_letter()

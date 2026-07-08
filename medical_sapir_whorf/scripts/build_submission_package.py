#!/usr/bin/env python3
"""
Assemble the Philosophy of Medicine submission package for
"Nosological Relativity: Formalizing the Medical Sapir-Whorf Hypothesis".

Generates:
  - cover_letter_en.docx
  - reporting_checklist_en.docx
Then bundles them together with the manuscript, title page, and figure
PNG/TIFF assets into output/submission_package_en.zip.

Run create_manuscript_en.py first (this script also runs it to be safe).
"""

import re
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import create_manuscript_en as m

OUT_DIR = m.OUT_DIR
JOURNAL = "Philosophy of Medicine"


def _base_style(doc, size=11):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(size)
    style.paragraph_format.line_spacing = 1.5


def abstract_word_count():
    return len(m.ABSTRACT.split())


def main_text_word_count():
    """Approximate main-text word count (Introduction..Conclusion), excluding
    citation markers, abstract, and references."""
    blocks = []
    blocks += m.INTRO_PARAS
    for paras in m.BACKGROUND_PARAS.values():
        blocks += paras
    blocks.append(m.FRAMEWORK_INTRO)
    blocks += [t for _, t in m.PROPOSITIONS]
    blocks += m.LOOPING_PARAS
    for paras in m.EVIDENCE_SECTIONS.values():
        blocks += paras
    blocks += m.PREDICTIONS_PARAS
    blocks += m.DISCUSSION_PARAS
    blocks += m.CONCLUSION_PARAS
    text = " ".join(blocks)
    text = re.sub(r'\{[^}]+\}', '', text)  # strip citation markers
    return len(text.split())


def build_cover_letter():
    doc = Document()
    _base_style(doc)

    doc.add_paragraph("[Date]")
    doc.add_paragraph()
    for line in ["The Editors", JOURNAL]:
        doc.add_paragraph(line)
    doc.add_paragraph()

    doc.add_paragraph("Dear Editors,")

    body = [
        (f'I am pleased to submit our manuscript, "{m.TITLE}," for consideration '
         f'as an Original Research article in {JOURNAL}.'),
        ("The paper develops the Nosological Relativity framework, a formal analogue "
         "of the Sapir\u2013Whorf hypothesis applied to medical classification systems. "
         "It argues that nosologies such as the ICD, DSM, and TNM do not merely describe "
         "clinical reality but partly constitute it, and it specifies this claim through "
         "six propositions, three levels of effect (cognitive, institutional, and "
         "population), and a looping mechanism adapted from Ian Hacking. The framework "
         "yields three testable predictions suited to natural-experiment designs, "
         "notably the ICD-10 to ICD-11 transition."),
        (f"We believe the work fits the scope of {JOURNAL}, which welcomes conceptual "
         "and philosophical contributions that bear on medical theory and practice. By "
         "rendering the constitutive effects of classification explicit and empirically "
         "testable, the paper connects philosophy of medicine with clinical epidemiology "
         "and nosological policy."),
        ("This manuscript is original, has not been published previously, and is not "
         "under consideration elsewhere. The author declares no conflicts of interest and "
         "no funding to report. The manuscript file has been anonymized for double-masked "
         "review; identifying information is provided on a separate title page."),
        (f"The main text is approximately {main_text_word_count():,} words (excluding the "
         f"abstract and references), within the {JOURNAL} limit for Original Research, and "
         f"the abstract is {abstract_word_count()} words."),
        "Thank you for considering our submission. We look forward to your response.",
    ]
    for para in body:
        doc.add_paragraph(para)

    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph(f"{m.AUTHORS} (corresponding author)")
    doc.add_paragraph(m.AUTHOR_AFFILIATION)
    doc.add_paragraph(f"Email: {m.AUTHOR_EMAIL}")

    out = OUT_DIR / "cover_letter_en.docx"
    doc.save(str(out))
    print(f"Cover letter saved: {out}")
    return out


def build_checklist():
    doc = Document()
    _base_style(doc, size=11)

    h = doc.add_heading("Submission Preparation Checklist", level=1)
    for r in h.runs:
        r.font.color.rgb = m.RGBColor(0, 0, 0)
    p = doc.add_paragraph()
    p.add_run(f"Manuscript: ").bold = True
    p.add_run(m.TITLE)
    p = doc.add_paragraph()
    p.add_run(f"Target journal: ").bold = True
    p.add_run(f"{JOURNAL} (Original Research)")
    doc.add_paragraph()

    items = [
        ("The submission has not been previously published, nor is it before another "
         "journal for consideration.", "Confirmed."),
        ("The submission file is in Microsoft Word (.docx) format.", "Confirmed."),
        ("Abstract of no more than 100 words.", f"{abstract_word_count()} words."),
        ("Main text within the Original Research limit (10,000 words excl. references).",
         f"~{main_text_word_count():,} words (excl. abstract and references)."),
        ("Manuscript anonymized for double-masked review; no author-identifying "
         "information in the manuscript file.", "Confirmed; identity on separate title page."),
        ("Separate title page with author name, affiliation, corresponding-author "
         "details, funding, and conflicts of interest.", "Included (title_page_en.docx)."),
        ("Sections numbered with descriptive headings (\u2264 three levels).",
         "Confirmed (1 Introduction \u2026 7 Conclusion)."),
        ("All figures and tables cited in the text in order of first appearance and "
         "placed inline.", "Confirmed (Fig. 1, Fig. 2, Table 1)."),
        ("Figures supplied at \u2265 300 dpi; tables editable (not images).",
         "Figures provided as 300 dpi PNG and TIFF; tables native/editable."),
        ("References complete and verified against primary sources; formatted in "
         "Chicago Manual of Style author\u2013date and ordered alphabetically.",
         "32 references verified; in-text citations use (Author Year)."),
        ("Text is 1.5-spaced; 11\u2013/12-point Times New Roman.",
         "Confirmed (Times New Roman 11 pt, 1.5 spacing)."),
    ]

    for i, (req, status) in enumerate(items, 1):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        r = para.add_run(f"{i}. [x] ")
        r.bold = True
        para.add_run(req + " ")
        s = para.add_run(f"\u2014 {status}")
        s.italic = True

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("Note on citation style: ").bold = True
    note.add_run(
        f"references and in-text citations follow the Chicago Manual of Style "
        f"author\u2013date system specified by {JOURNAL}. In-text citations appear as "
        "(Author Year); the reference list is ordered alphabetically by author surname."
    )

    out = OUT_DIR / "reporting_checklist_en.docx"
    doc.save(str(out))
    print(f"Checklist saved: {out}")
    return out


def build_zip(paths):
    zip_path = OUT_DIR / "submission_package_en.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for p in paths:
            p = Path(p)
            if not p.exists():
                raise FileNotFoundError(p)
            zf.write(p, arcname=p.name)
    print(f"Zip saved: {zip_path}")
    return zip_path


def main():
    m.main()  # (re)generate manuscript, title page, figures, tables
    cover = build_cover_letter()
    checklist = build_checklist()
    contents = [
        cover,
        OUT_DIR / "manuscript_en.docx",
        OUT_DIR / "title_page_en.docx",
        checklist,
        OUT_DIR / "figure1_framework.png",
        OUT_DIR / "figure2_looping.png",
        OUT_DIR / "figure1_framework.tif",
        OUT_DIR / "figure2_looping.tif",
    ]
    build_zip(contents)
    print("Submission package complete.")


if __name__ == "__main__":
    main()

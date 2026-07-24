#!/usr/bin/env python3
"""
Generate title page for Measurement submission (double-blind review).

Measurement requires a separate title page with author info,
submitted independently from the blinded manuscript.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUT_DIR = Path(__file__).resolve().parent


def build_title_page():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TITLE PAGE')
    run.font.size = Pt(14)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run(
        'Optimal noise model accuracy for dynamic vision sensors:\n'
        'a stochastic resonance framework for the noise inverse problem'
    )
    run.font.size = Pt(16)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run('[Author names to be added]')
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run('[Affiliations to be added]')
    run.font.size = Pt(12)
    run.italic = True

    # Corresponding author
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run('Corresponding author:')
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run('[Name, email, telephone, full postal address]')
    run.font.size = Pt(12)
    run.italic = True

    # Word count
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run('Manuscript information:')
    run.bold = True
    run.font.size = Pt(12)

    info_items = [
        'Article type: Research Article',
        'Number of figures: 11',
        'Number of tables: 2',
        'Number of equations: 13',
        'Number of references: 30',
    ]
    for item in info_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(12)

    out_path = OUT_DIR / 'title_page.docx'
    doc.save(str(out_path))
    print(f"Title page saved: {out_path}")
    return out_path


if __name__ == '__main__':
    build_title_page()

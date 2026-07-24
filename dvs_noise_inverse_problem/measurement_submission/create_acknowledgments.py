#!/usr/bin/env python3
"""
Generate Acknowledgments document for Measurement (Elsevier) submission.

Elsevier requirements:
- Separate editable file (Microsoft Word)
- Collate acknowledgments in a separate section at the end of the article
  before the references
- Do not include acknowledgments on the title page or as a footnote
"""

from docx import Document
from docx.shared import Pt
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
OUT_DIR = SCRIPT_DIR


def build_acknowledgments():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0

    p = doc.add_paragraph()
    run = p.add_run('Acknowledgments')
    run.font.size = Pt(14)
    run.bold = True

    p = doc.add_paragraph()
    run = p.add_run('[To be added.]')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = True

    out_path = OUT_DIR / 'acknowledgments.docx'
    doc.save(str(out_path))
    print(f"Acknowledgments saved: {out_path}")
    return out_path


if __name__ == '__main__':
    build_acknowledgments()

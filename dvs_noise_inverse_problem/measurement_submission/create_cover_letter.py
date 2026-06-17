#!/usr/bin/env python3
"""
Generate cover letter for Measurement (Elsevier) submission.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUT_DIR = Path(__file__).resolve().parent


def build_cover_letter():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Date
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run('[Date]')
    run.font.size = Pt(12)

    # Addressee
    p = doc.add_paragraph()
    run = p.add_run(
        'Editor-in-Chief\n'
        'Measurement\n'
        'Elsevier'
    )
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run('Dear Editor,')
    run.font.size = Pt(12)

    # Body
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(
        'We submit the enclosed manuscript entitled "Optimal noise model accuracy '
        'for dynamic vision sensors: a stochastic resonance framework for the '
        'noise inverse problem" for consideration as a Research Article in Measurement.'
    )
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        'This manuscript addresses a fundamental question in measurement science '
        'for threshold-based event detectors: how much noise should a noise model '
        'remove? We derive a closed-form optimal noise model accuracy by unifying '
        'stochastic resonance (SR) theory with covariate adjustment (ANCOVA), showing '
        'that noise should be reduced to the SR optimum\u2014not eliminated\u2014because '
        'over-removal degrades detection performance.'
    )
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        'We instantiate this framework for dynamic vision sensor (DVS) astronomical '
        'observation, a rapidly developing measurement modality for space situational '
        'awareness. A five-parameter circuit-level noise model, a Fano-factor-based '
        'classifier, and a six-tier calibration framework provide practical tools for '
        'DVS noise measurement and removal. Systematic evaluation on 20 recordings '
        'demonstrates ROC-AUC = 0.866 with 93.9% signal preservation, consistent '
        'with the theoretical predictions.'
    )
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        'The manuscript is relevant to Measurement\u2019s scope because it: '
        '(1) develops novel measurement methodology for a new class of sensors; '
        '(2) provides a calibration framework with quantitative pass/fail criteria; '
        '(3) connects to the GUM framework by showing that measurement uncertainty '
        'in event-based sensors depends non-monotonically on noise; and '
        '(4) introduces satellite trails as metrological calibration standards.'
    )
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        'The manuscript has not been published or submitted elsewhere. '
        'All authors have approved the manuscript and agree with its submission '
        'to Measurement. The authors declare no competing interests.'
    )
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run('Sincerely,')
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run('[Author name]\n[Affiliation]\n[Email]')
    run.font.size = Pt(12)
    run.italic = True

    out_path = OUT_DIR / 'cover_letter_measurement.docx'
    doc.save(str(out_path))
    print(f"Cover letter saved: {out_path}")
    return out_path


if __name__ == '__main__':
    build_cover_letter()

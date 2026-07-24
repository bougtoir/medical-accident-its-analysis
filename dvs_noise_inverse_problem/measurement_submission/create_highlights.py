#!/usr/bin/env python3
"""
Generate Highlights document for Measurement (Elsevier) submission.

Elsevier requirements:
- Separate source file (Microsoft Word, not PDF)
- 3 to 5 bullet points
- Each highlight max 85 characters including spaces
- No author details (double-blind)
- File name: "Highlights"
"""

from docx import Document
from docx.shared import Pt
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
OUT_DIR = SCRIPT_DIR


def build_highlights():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0

    p = doc.add_paragraph()
    run = p.add_run('Highlights')
    run.font.size = Pt(14)
    run.bold = True

    highlights = [
        'Closed-form optimal noise model accuracy derived via stochastic resonance theory.',
        'Five-parameter physics-informed noise model enables DVS noise inverse problem.',
        'Fano-factor classifier achieves ROC-AUC 0.866 with 93.9% signal preservation.',
        'Six-tier calibration with satellite-trail tier for in-operation verification.',
        'Proof-of-concept: 90.3% noise removal with satellite trajectory recovery.',
    ]

    for h in highlights:
        assert len(h) <= 85, f"Highlight too long ({len(h)} chars): {h}"
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    out_path = OUT_DIR / 'highlights.docx'
    doc.save(str(out_path))
    print(f"Highlights saved: {out_path}")

    for i, h in enumerate(highlights, 1):
        print(f"  [{i}] {len(h)} chars: {h}")

    return out_path


if __name__ == '__main__':
    build_highlights()

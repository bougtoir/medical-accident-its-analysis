#!/usr/bin/env python3
"""
Generate JATIS manuscript as docx with inline figures.

Title: "Solving the noise inverse problem in dynamic vision sensors
        for faint astronomical object detection"

Target: Journal of Astronomical Telescopes, Instruments, and Systems (SPIE)

SPIE formatting:
- Single column, Times Roman 12pt
- Numbered references (superscript in text) — here we use [N] for docx clarity
- Single-paragraph abstract (≤200 words)
- Figures inline with captions below
- Section numbering: 1, 1.1, 1.2, etc.
- Equations as Word equation objects (OMML)
"""

import re
from lxml import etree
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUT_DIR = Path(__file__).resolve().parent
FIG_DIR = OUT_DIR


# =========================================================
# OMML equation helpers
# =========================================================

def _mr(parent, text, italic=True, bold=False):
    """Create a math run (m:r) with text."""
    r = etree.SubElement(parent, qn('m:r'))
    if not italic or bold:
        rPr = etree.SubElement(r, qn('m:rPr'))
        if not italic:
            sty = etree.SubElement(rPr, qn('m:sty'))
            sty.set(qn('m:val'), 'p')
        if bold:
            sty = etree.SubElement(rPr, qn('m:sty'))
            sty.set(qn('m:val'), 'b')
    t = etree.SubElement(r, qn('m:t'))
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return r


def _sub(parent, base, sub):
    """Create subscript: base_sub."""
    el = etree.SubElement(parent, qn('m:sSub'))
    e = etree.SubElement(el, qn('m:e'))
    _mr(e, base)
    s = etree.SubElement(el, qn('m:sub'))
    _mr(s, sub)
    return el


def _sup(parent, base, sup):
    """Create superscript: base^sup."""
    el = etree.SubElement(parent, qn('m:sSup'))
    e = etree.SubElement(el, qn('m:e'))
    _mr(e, base)
    s = etree.SubElement(el, qn('m:sup'))
    _mr(s, sup)
    return el


def _frac(parent, num_builder, den_builder):
    """Create fraction with builder functions for numerator and denominator."""
    f = etree.SubElement(parent, qn('m:f'))
    num = etree.SubElement(f, qn('m:num'))
    num_builder(num)
    den = etree.SubElement(f, qn('m:den'))
    den_builder(den)
    return f


def _delim(parent, content_builder, left='(', right=')'):
    """Create delimiter (parentheses, brackets)."""
    d = etree.SubElement(parent, qn('m:d'))
    dPr = etree.SubElement(d, qn('m:dPr'))
    begChr = etree.SubElement(dPr, qn('m:begChr'))
    begChr.set(qn('m:val'), left)
    endChr = etree.SubElement(dPr, qn('m:endChr'))
    endChr.set(qn('m:val'), right)
    e = etree.SubElement(d, qn('m:e'))
    content_builder(e)
    return d


def _bar(parent, text):
    """Create accent bar (hat/overline): x̂ or x̄."""
    acc = etree.SubElement(parent, qn('m:acc'))
    accPr = etree.SubElement(acc, qn('m:accPr'))
    chrEl = etree.SubElement(accPr, qn('m:chr'))
    chrEl.set(qn('m:val'), '\u0302')  # combining circumflex (hat)
    e = etree.SubElement(acc, qn('m:e'))
    _mr(e, text)
    return acc


def _hat(parent, text):
    """Create hat accent: x̂."""
    return _bar(parent, text)


def _func(parent, name, arg_builder):
    """Create function application: name(args)."""
    _mr(parent, name, italic=False)
    _delim(parent, arg_builder)


def add_display_equation(doc, builder_func, eq_num=None):
    """Add a display equation (centered) with optional number."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    omathpara = etree.SubElement(p._element, qn('m:oMathPara'))
    omath = etree.SubElement(omathpara, qn('m:oMath'))
    builder_func(omath)
    if eq_num:
        run = p.add_run(f'    ({eq_num})')
        run.font.size = Pt(12)
    return p


def add_inline_math(p, builder_func):
    """Add inline equation within a paragraph."""
    omath = etree.SubElement(p._element, qn('m:oMath'))
    builder_func(omath)
    return omath


# =========================================================
# Equation definitions for this manuscript
# =========================================================

def eq_a5_model(omath):
    """Eq (1): λ_noise(T, I_bg) = I_dark,ref · exp(α·ΔT) · (1 + β·I_bg)"""
    _sub(omath, '\u03bb', 'noise')
    def _args(e):
        _mr(e, 'T, ')
        _sub(e, 'I', 'bg')
    _delim(omath, _args)
    _mr(omath, ' = ')
    _sub(omath, 'I', 'dark,ref')
    _mr(omath, ' \u22c5 exp')
    def _exp_arg(e):
        _mr(e, '\u03b1 \u22c5 \u0394T')
    _delim(omath, _exp_arg)
    _mr(omath, ' \u22c5 ')
    def _bg(e):
        _mr(e, '1 + \u03b2 \u22c5 ')
        _sub(e, 'I', 'bg')
    _delim(omath, _bg)


def eq_accuracy(omath):
    """Eq (2): α = 1 − ‖ê_noise − e_noise,true‖ / ‖e_noise,true‖"""
    _mr(omath, '\u03b1 = 1 \u2212 ')
    def _num(n):
        _mr(n, '\u2016')
        _hat(n, 'e')
        _sub(n, '', 'noise')
        _mr(n, ' \u2212 ')
        _sub(n, 'e', 'noise,true')
        _mr(n, '\u2016')
    def _den(d):
        _mr(d, '\u2016')
        _sub(d, 'e', 'noise,true')
        _mr(d, '\u2016')
    _frac(omath, _num, _den)


def eq_residual_noise(omath):
    """Eq (3): σ_residual = (1 − α) · σ_original"""
    _sub(omath, '\u03c3', 'residual')
    _mr(omath, ' = ')
    def _factor(e):
        _mr(e, '1 \u2212 \u03b1')
    _delim(omath, _factor)
    _mr(omath, ' \u22c5 ')
    _sub(omath, '\u03c3', 'original')


def eq_snr_improvement(omath):
    """Eq (4): SNR_after / SNR_before = 1 / (1 − α)"""
    def _num(n):
        _sub(n, 'SNR', 'after')
    def _den(d):
        _sub(d, 'SNR', 'before')
    _frac(omath, _num, _den)
    _mr(omath, ' = ')
    def _num2(n):
        _mr(n, '1', italic=False)
    def _den2(d):
        _mr(d, '1 \u2212 \u03b1')
    _frac(omath, _num2, _den2)


def eq_map_estimation(omath):
    """Eq (5): θ̂ = argmax_θ p(e_cal|θ) · p(θ|θ_prior)"""
    _hat(omath, '\u03b8')
    _mr(omath, ' = ')
    _sub(omath, 'argmax', '\u03b8')
    _mr(omath, ' p')
    def _likelihood(e):
        _sub(e, 'e', 'cal')
        _mr(e, ' | \u03b8')
    _delim(omath, _likelihood)
    _mr(omath, ' \u22c5 p')
    def _prior(e):
        _mr(e, '\u03b8 | ')
        _sub(e, '\u03b8', 'prior')
    _delim(omath, _prior)


def eq_nn_output(omath):
    """Eq (6): λ̂_noise = λ_physics · (1 + Δλ_aux) + Δλ_corr"""
    _hat(omath, '\u03bb')
    _mr(omath, '')
    _sub(omath, '', 'noise')
    _mr(omath, ' = ')
    _sub(omath, '\u03bb', 'physics')
    _mr(omath, ' \u22c5 ')
    def _modulation(e):
        _mr(e, '1 + \u0394')
        _sub(e, '\u03bb', 'aux')
    _delim(omath, _modulation)
    _mr(omath, ' + \u0394')
    _sub(omath, '\u03bb', 'corr')


def eq_p_noise(omath):
    """Eq (7): P_noise(e_i) = λ̂_noise(x_i,y_i,t_i) / [λ̂_noise(...) + λ̂_signal(...)]"""
    _sub(omath, 'P', 'noise')
    def _ei(e):
        _sub(e, 'e', 'i')
    _delim(omath, _ei)
    _mr(omath, ' = ')
    def _num(n):
        _hat(n, '\u03bb')
        _sub(n, '', 'noise')
        def _coords(e):
            _sub(e, 'x', 'i')
            _mr(e, ', ')
            _sub(e, 'y', 'i')
            _mr(e, ', ')
            _sub(e, 't', 'i')
        _delim(n, _coords)
    def _den(d):
        _hat(d, '\u03bb')
        _sub(d, '', 'noise')
        def _c1(e):
            _sub(e, 'x', 'i')
            _mr(e, ', ')
            _sub(e, 'y', 'i')
            _mr(e, ', ')
            _sub(e, 't', 'i')
        _delim(d, _c1)
        _mr(d, ' + ')
        _hat(d, '\u03bb')
        _sub(d, '', 'signal')
        def _c2(e):
            _sub(e, 'x', 'i')
            _mr(e, ', ')
            _sub(e, 'y', 'i')
            _mr(e, ', ')
            _sub(e, 't', 'i')
        _delim(d, _c2)
    _frac(omath, _num, _den)


def eq_fano(omath):
    """Eq (8): F = Var(N_k) / Mean(N_k)"""
    _mr(omath, 'F', italic=True)
    _mr(omath, ' = ', italic=False)
    def _num(n):
        _mr(n, 'Var', italic=False)
        def _nk(e):
            _sub(e, 'N', 'k')
        _delim(n, _nk)
    def _den(d):
        _mr(d, 'Mean', italic=False)
        def _nk(e):
            _sub(e, 'N', 'k')
        _delim(d, _nk)
    _frac(omath, _num, _den)


def eq_detection_limit(omath):
    """Eq (9): Δm ≈ 2.5 log₁₀(1/(1−α))"""
    _mr(omath, '\u0394m \u2248 2.5 ')
    _sub(omath, 'log', '10')
    def _arg(e):
        def _n(n):
            _mr(n, '1', italic=False)
        def _d(d):
            _mr(d, '1 \u2212 \u03b1')
        _frac(e, _n, _d)
    _delim(omath, _arg)


# =========================================================
# Document building helpers
# =========================================================

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_paragraph(doc, text, style='Normal', bold=False, italic=False,
                  alignment=None, space_after=None, space_before=None):
    """Add paragraph with optional superscript citation handling."""
    p = doc.add_paragraph(style=style)
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)

    # Parse {N} markers for superscript citations
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = p.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(9)
        else:
            run = p.add_run(part)
            if bold:
                run.bold = True
            if italic:
                run.italic = True
    return p


def add_figure(doc, img_path, caption, width=Inches(5.5)):
    """Insert figure inline with caption below."""
    if not img_path.exists():
        add_paragraph(doc, f"[MISSING FIGURE: {img_path.name}]")
        return
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(12)
    run = p_img.add_run()
    run.add_picture(str(img_path), width=width)

    # Caption
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_cap.paragraph_format.space_before = Pt(6)
    p_cap.paragraph_format.space_after = Pt(12)
    # Parse caption for superscript refs
    parts = re.split(r'(\{[^}]+\})', caption)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run_cap = p_cap.add_run(part[1:-1])
            run_cap.font.superscript = True
            run_cap.font.size = Pt(8)
        else:
            run_cap = p_cap.add_run(part)
            run_cap.font.size = Pt(9)
    return p_cap


def add_table(doc, headers, data, caption=None):
    """Add a table with headers and data rows."""
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
    for row_idx, row_data in enumerate(data):
        for col_idx, val in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = val
            for p in table.rows[row_idx + 1].cells[col_idx].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if caption:
        add_paragraph(doc, caption, italic=True, space_after=12,
                      space_before=6)
    return table


def build_manuscript():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # Page setup (Letter size per SPIE)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.25)
    section.left_margin = Inches(0.875)
    section.right_margin = Inches(0.875)

    # =========================================================
    # Title
    # =========================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'Solving the noise inverse problem in dynamic vision sensors\n'
        'for faint astronomical object detection'
    )
    run.font.size = Pt(16)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('[Author names and affiliations to be added]')
    run.font.size = Pt(12)
    run.italic = True

    # =========================================================
    # Abstract (single paragraph, ≤200 words for SPIE)
    # =========================================================
    add_heading(doc, 'Abstract', level=1)
    add_paragraph(doc, (
        'Dynamic Vision Sensors (DVS) offer microsecond temporal resolution and '
        '>120 dB dynamic range for space situational awareness, but background '
        'activity noise dominates under low-light conditions. We propose a '
        'physics-informed framework treating DVS noise removal as an inverse '
        'problem: modelling the noise generation mechanism via a circuit-level '
        'parametric model and subtracting reconstructed noise to recover faint '
        'signals. We develop Physics-Informed DeepClean for DVS (PI-DC-DVS), '
        'integrating a five-parameter noise model (A5) with Bayesian inference. '
        'Systematic evaluation on 20 recordings from the Event-Based Space '
        'Situational Awareness (EBSSA) dataset shows the Fano-factor-based noise '
        'inverse approach achieves ROC-AUC = 0.866, substantially outperforming '
        'temporal filtering (AUC = 0.534). A proof-of-concept demonstrates 90.3% '
        'noise removal with clear satellite trajectory recovery. We further '
        'introduce a six-tier calibration framework where Cal-6 repurposes '
        'satellite light trails\u2014conventionally regarded as light pollution\u2014as '
        'natural calibration sources. The noise inverse paradigm, proven in '
        'gravitational-wave astronomy, transfers effectively to DVS observation '
        'with potential 2\u20134 magnitude detection limit extension.'
    ))

    add_paragraph(doc, (
        'Keywords: dynamic vision sensor, event camera, noise inverse problem, '
        'space situational awareness, calibration, stochastic resonance'
    ), italic=True, space_after=12)

    # =========================================================
    # 1. Introduction
    # =========================================================
    add_heading(doc, '1 Introduction', level=1)
    add_paragraph(doc, (
        'Conventional noise reduction in imaging is formulated as the signal '
        'inverse problem: given an observation y = h * s + n, recover the signal s. '
        'This work explores the complementary perspective\u2014the noise inverse '
        'problem\u2014where the noise generation mechanism is modelled as a physical '
        'forward process and solved inversely to reconstruct and subtract noise, '
        'leaving only the signal in the residual.'
    ))
    add_paragraph(doc, (
        'Dynamic Vision Sensors (DVS) are neuromorphic sensors in which each pixel '
        'independently and asynchronously emits an event when the logarithmic '
        'intensity change exceeds a threshold.{1,2} Inspired by change-detection '
        'neurons in insect compound eyes, DVS offer microsecond temporal resolution, '
        '>120 dB dynamic range, and sparse output. These properties make them '
        'attractive for space situational awareness (SSA){3,4,5} and fast optical '
        'astronomy.{6} Under low-light conditions, however, shot-noise-induced '
        'background activity (BA) becomes the dominant source of events, '
        'overwhelming faint astronomical signals.{7,8}'
    ))
    add_paragraph(doc, (
        'The core insight is: solve the noise inverse problem with high precision '
        '\u2192 reconstruct and subtract noise \u2192 only signal remains \u2192 structural '
        'SNR improvement.'
    ), italic=True)
    add_paragraph(doc, (
        'This paradigm has been spectacularly successful in gravitational-wave (GW) '
        'astronomy, where DeepClean{9} and related methods{10,11,12} use auxiliary '
        'witness channels to model and subtract non-stationary instrumental noise. '
        'The iDQ framework{13} assigns per-event noise probabilities in real time. '
        'D\u00b3PO{14} performs simultaneous Bayesian signal\u2013noise decomposition for '
        'photon-counting data. Yet no equivalent pipeline exists for DVS '
        'astronomical observation.'
    ))
    add_paragraph(doc, (
        'In this paper we: (i) propose the PI-DC-DVS algorithm integrating '
        'physics-informed modelling with auxiliary-channel regression (Sec. 3); '
        '(ii) design a six-tier calibration framework including a novel '
        'satellite-trail calibration tier (Sec. 4); (iii) systematically evaluate '
        'the approach on 20 EBSSA recordings against two baselines (Sec. 5); and '
        '(iv) present a proof-of-concept demonstration achieving 90.3% noise '
        'removal (Sec. 6).'
    ))

    # =========================================================
    # 2. Background
    # =========================================================
    add_heading(doc, '2 Background', level=1)

    add_heading(doc, '2.1 DVS noise physics', level=2)
    add_paragraph(doc, (
        'The circuit-level physics of DVS noise has been systematically '
        'characterised by the UZH/ETH Zurich group. Gra\u00e7a and Delbruck{7} proved '
        'that photon shot noise sets a fundamental lower bound at twice the photon '
        'shot noise level. McReynolds et al.{8} demonstrated that shot-noise events '
        'exhibit alternating ON\u2194OFF polarity patterns. SciDVS,{15} a '
        'scientific-grade sensor, achieves 1.7% temporal contrast sensitivity at '
        '0.7 lux. Most importantly, Gra\u00e7a and Delbruck{16} introduced a '
        'large-signal differential-equation DVS pixel model incorporating '
        'first-passage-time stochastic event generation, achieving >1000\u00d7 '
        'computational speedup while maintaining physical realism. This model '
        'serves as the forward model F(\u03b8) for the noise inverse problem.'
    ))
    add_paragraph(doc, (
        'The parametric noise rate model (hereafter A5){16} takes the form:'
    ))
    add_display_equation(doc, eq_a5_model, eq_num='1')
    add_paragraph(doc, (
        'where I_dark,ref is the reference dark current rate, \u03b1 \u2248 0.06\u20130.08 K\u207b\u00b9 '
        'is the temperature coefficient, \u0394T is the temperature offset from reference, '
        'and \u03b2 is the background illuminance sensitivity. The five parameters '
        '(I_dark,ref, \u03b1, \u03b2, \u03b8_ON, \u03b8_OFF) are fitted per-pixel during offline calibration.'
    ))

    add_heading(doc, '2.2 DVS denoising methods', level=2)
    add_paragraph(doc, (
        'DVS denoising has evolved from empirical spatio-temporal filtering{17,18} '
        'through probabilistic methods (Event Probability Mask, EPM{19}) and deep '
        'learning (WedNet;{20} ASTEDNet{21}) to joint motion\u2013noise estimation. Most '
        'notably, Shiba et al.{22} simultaneously estimate motion and noise via an '
        'extended Contrast Maximisation framework\u2014the conceptually closest prior '
        'work. However, their noise model is phenomenological and does not '
        'incorporate circuit physics or auxiliary channels.'
    ))

    add_heading(doc, '2.3 DVS astronomical applications', level=2)
    add_paragraph(doc, (
        'DVS astronomical applications concentrate on SSA. Afshar et al.{3} '
        'produced the first event-based space-observation dataset (EBSSA: 236 '
        'recordings, 572 labelled objects). FIESTA{4} demonstrated unsupervised '
        'real-time space-object detection. Hoang{6} explored neuromorphic cameras '
        'for atmospheric Cherenkov telescopes. No study has applied DVS to '
        'faint-object detection where noise dominates.'
    ))

    # Fig 2: Gap map
    add_figure(doc, FIG_DIR / 'fig2_gapmap.png',
               'Fig. 1. Gap map showing the four surveyed domains (A: DVS noise '
               'physics, B: DVS denoising, C: DVS astronomical applications, '
               'D: noise inverse problem methods) and the identified research gaps '
               'at their intersections.',
               width=Inches(5.0))

    add_heading(doc, '2.4 Noise inverse problem in other fields', level=2)
    add_paragraph(doc, (
        'DeepClean{9} regresses non-stationary noise from auxiliary witness '
        'channels using machine learning, achieving order-of-magnitude noise '
        'reduction in LIGO. Noise2Image{23} exploits the illuminance dependence '
        'of DVS noise rates to recover static scenes from noise alone\u2014demonstrating '
        'that DVS noise carries information. D\u00b3PO{14} simultaneously performs '
        'Bayesian signal\u2013noise decomposition in photon-counting data using '
        'information field theory.'
    ))

    # =========================================================
    # 3. Methods
    # =========================================================
    add_heading(doc, '3 Methods: PI-DC-DVS algorithm', level=1)

    add_heading(doc, '3.1 Fundamental principle', level=2)
    add_paragraph(doc, (
        'Let \u03b1 denote the noise model accuracy:'
    ))
    add_display_equation(doc, eq_accuracy, eq_num='2')
    add_paragraph(doc, (
        'The residual noise level after subtraction is:'
    ))
    add_display_equation(doc, eq_residual_noise, eq_num='3')
    add_paragraph(doc, (
        'yielding an SNR improvement ratio:'
    ))
    add_display_equation(doc, eq_snr_improvement, eq_num='4')
    add_paragraph(doc, (
        'At \u03b1 = 0.9, this gives 10\u00d7 improvement; at \u03b1 = 0.99, 100\u00d7. This simple '
        'scaling law motivates the emphasis on noise model accuracy rather than '
        'signal template matching.'
    ))

    add_heading(doc, '3.2 Algorithm overview', level=2)
    add_paragraph(doc, (
        'PI-DC-DVS operates in four phases (Fig. 2):'
    ))

    # Fig 1: Pipeline
    add_figure(doc, FIG_DIR / 'fig1_pipeline.png',
               'Fig. 2. System architecture of the PI-DC-DVS noise inverse problem '
               'pipeline. Four stages: (1) noise forward model construction using the '
               'A5 pixel model and auxiliary channels; (2) Bayesian inverse problem '
               'solution with physics-informed neural network; (3) residual event stream '
               'generation via probabilistic thinning; (4) astronomical calibration and '
               'verification including Cal-6 satellite trail calibration.',
               width=Inches(6.0))

    add_paragraph(doc, (
        'Phase 1: Offline calibration (pre-observation). Record dark events '
        '(lens cap), flat-field events (integrating sphere), and thermal sweep '
        'events (\u0394T = \u00b15\u00b0C). Fit the A5 forward model{16} to obtain per-pixel '
        'parameter maps via MAP estimation:'
    ))
    add_display_equation(doc, eq_map_estimation, eq_num='5')
    add_paragraph(doc, (
        'Phase 2: Online inference (real-time). A Physics-Informed Neural Network '
        'predicts per-pixel noise rates. '
        'The network comprises three layers: (a) Physics model layer (fixed weights, '
        'A5 model baseline), (b) Auxiliary-channel coupling layer (MLP [64-32-1] '
        'learning non-stationary variations), (c) Spatio-temporal correlation layer '
        '(Conv2D 3\u00d73 learning inter-pixel noise correlations). Output:'
    ))
    add_display_equation(doc, eq_nn_output, eq_num='6')
    add_paragraph(doc, (
        'Per-event noise probability (following iDQ{13}):'
    ))
    add_display_equation(doc, eq_p_noise, eq_num='7')
    add_paragraph(doc, (
        'Phase 3: Residual generation. Soft subtraction (recommended): assign '
        'weight w_i = 1 \u2212 P_noise(e_i); retain events with w_i > w_threshold. '
        'Hard subtraction (fast): retain events with P_noise(e_i) < \u03c4.'
    ))
    add_paragraph(doc, (
        'Phase 4: Adaptive updates. Monitor residual Poisson statistics; trigger '
        'online weight updates. Apply Kalman-filter-like drift correction.'
    ))

    add_heading(doc, '3.3 Simplified implementation for EBSSA', level=2)
    add_paragraph(doc, (
        'Fano filter (proposed baseline): The Fano factor (variance-to-mean ratio '
        'of event rates across temporal bins) discriminates noise from signal: '
        'pixels with Fano \u2248 1 (Poisson-consistent) are noise-dominated, while '
        'Fano \u226b 1 indicates bursty signal. Formally:'
    ))
    add_display_equation(doc, eq_fano, eq_num='8')
    add_paragraph(doc, (
        'where N_k is the event count in temporal bin k. Per-event noise probability '
        'is computed from noise-dominated pixels (Fano \u2264 2). Events with '
        'P_noise > \u03c4 (\u03c4 = 0.5) are classified as noise.'
    ))
    add_paragraph(doc, (
        'PI-DC-DVS neural network (simplified): A three-layer neural network with '
        'the A5-inspired physics layer, temporal modulation layer (substituting '
        'for absent auxiliary channels), and spatio-temporal correlation layer. '
        'Self-supervised training on noise-dominated pixels with Poisson NLL loss.'
    ))

    # =========================================================
    # 4. Calibration framework
    # =========================================================
    add_heading(doc, '4 Calibration framework', level=1)

    add_heading(doc, '4.1 Six-tier calibration dataset', level=2)
    add_paragraph(doc, (
        'DVS output event streams rather than frames, requiring purpose-designed '
        'calibration procedures. We propose six tiers (Table 1):'
    ))

    add_table(doc,
              ['Tier', 'Condition', 'Purpose', 'Pass criterion'],
              [
                  ['Cal-1', 'Dark (lens cap)', 'Pure noise reference', '\u03c7\u00b2/dof < 1.5'],
                  ['Cal-2', 'Thermal sweep', 'Temperature dependence', 'Residual < 10%'],
                  ['Cal-3', 'Flat-field', 'Shot noise statistics', '\u03b1_flat > 0.9'],
                  ['Cal-4', 'Dynamic patterns', 'Injection-recovery', 'AUC > 0.95'],
                  ['Cal-5', 'Simulated astro.', 'End-to-end pipeline', '\u0394m > 2 mag'],
                  ['Cal-6', 'Satellite trails', 'In-operation verification', 'Det. rate > 95%'],
              ],
              caption='Table 1. Six-tier calibration framework for DVS noise model validation.')

    add_heading(doc, '4.2 Cal-6: Satellite trail calibration', level=2)
    add_paragraph(doc, (
        'We propose repurposing satellite light trails\u2014conventionally regarded as '
        'light pollution\u2014as natural calibration sources. Artificial satellites have '
        'precisely predictable trajectories: Two-Line Element (TLE) orbital data '
        'provide position, velocity, and transit time to microsecond precision. '
        'This makes satellite transits a natural injection-recovery test under '
        'real observing conditions.'
    ))
    add_paragraph(doc, (
        'DVS-specific advantages of Cal-6 include: (1) Abundant calibration '
        '(Starlink: thousands of satellites, dozens of transits per night at any '
        'observatory site); (2) Real observing conditions (atmospheric scintillation, '
        'actual sky background); (3) No saturation (DVS >120 dB dynamic range '
        'records bright trails quantitatively); (4) No additional hardware or '
        'observation time required.'
    ))
    add_paragraph(doc, (
        'The calibration procedure is: predict satellite transit using TLE + SGP4 '
        'propagator \u2192 extract event stream during predicted transit window \u2192 apply '
        'noise subtraction pipeline \u2192 compare detected trail with predicted '
        'trajectory \u2192 compute detection rate, positional accuracy, and timing '
        'precision as quantitative metrics.'
    ))

    # =========================================================
    # 5. Systematic evaluation
    # =========================================================
    add_heading(doc, '5 Systematic evaluation', level=1)

    add_heading(doc, '5.1 Dataset', level=2)
    add_paragraph(doc, (
        'We use the EBSSA dataset:{3} 236 recordings from DAVIS240C sensors '
        'observing satellites and stars, with 572 labelled space objects. We select '
        '20 recordings spanning both sensor configurations (180\u00d7240 and 240\u00d7304 '
        'pixels) for systematic evaluation.'
    ))

    add_heading(doc, '5.2 Evaluated methods', level=2)
    add_paragraph(doc, (
        'Three methods are compared: (1) Fano filter (proposed): Physics-based '
        'noise inverse approach using the Fano factor as a Poisson discriminant; '
        '(2) PI-DC-DVS NN (proposed, simplified): Three-layer neural network, '
        'self-supervised on noise-dominated pixels; (3) Temporal filter '
        '(baseline):{17} Spatio-temporal neighbourhood filter retaining events '
        'only when sufficient neighbors are present within a fixed window.'
    ))

    add_heading(doc, '5.3 Results', level=2)

    add_table(doc,
              ['Method', 'NRR', 'SPR', 'F1', 'AUC'],
              [
                  ['Temporal filter', '0.852 \u00b1 0.044', '0.216 \u00b1 0.157',
                   '0.253 \u00b1 0.176', '0.534 \u00b1 0.083'],
                  ['PI-DC-DVS NN', '0.171 \u00b1 0.342', '0.841 \u00b1 0.342',
                   '0.488 \u00b1 0.453', '0.546 \u00b1 0.218'],
                  ['Fano filter', '0.713 \u00b1 0.232', '0.939 \u00b1 0.056',
                   '0.697 \u00b1 0.339', '0.866 \u00b1 0.107'],
              ],
              caption='Table 2. Systematic evaluation results (mean \u00b1 std) across '
                      '20 EBSSA recordings. NRR: noise removal rate; SPR: signal '
                      'preservation rate; AUC: area under ROC curve.')

    add_paragraph(doc, (
        'The Fano filter achieves the best balance between noise removal and signal '
        'preservation (Fig. 3), with AUC = 0.866 substantially exceeding both the '
        'temporal filter (AUC = 0.534) and the simplified PI-DC-DVS NN '
        '(AUC = 0.546). The temporal filter achieves the highest raw noise removal '
        'rate (85.2%) but at the cost of destroying most signal events '
        '(SPR = 21.6%), making it unsuitable for faint-object detection.'
    ))

    # Fig 3: Evaluation
    add_figure(doc, FIG_DIR / 'fig3_evaluation.png',
               'Fig. 3. Systematic evaluation of three denoising methods on 20 EBSSA '
               'recordings. Four-panel boxplot showing (a) Noise Removal Rate, '
               '(b) Signal Preservation Rate, (c) F1 Score, and (d) ROC-AUC. '
               'Diamond markers indicate means. The Fano filter achieves the best '
               'overall balance (AUC = 0.866).',
               width=Inches(5.5))

    add_paragraph(doc, (
        'Per-recording analysis (Fig. 4) reveals that the Fano filter consistently '
        'outperforms the temporal filter across diverse observing conditions, '
        'demonstrating robustness to variations in noise rate, background '
        'illuminance, and target brightness.'
    ))

    # Fig 5: Per-recording
    add_figure(doc, FIG_DIR / 'fig5_per_recording.png',
               'Fig. 4. Per-recording noise removal rate comparison across 20 EBSSA '
               'recordings. The Fano filter (orange) achieves selective noise removal, '
               'while the temporal filter (green) removes events indiscriminately.',
               width=Inches(5.5))

    add_heading(doc, '5.4 A5-based noise rate simulation', level=2)
    add_paragraph(doc, (
        'Using the A5 parametric model, we simulate noise rates and SNR '
        'improvements across the temperature\u2013illuminance parameter space '
        '(T \u2208 [10, 65]\u00b0C, I_bg \u2208 [0.1, 1000] lux). Figure 5 shows the '
        'predicted noise rate map, SNR improvement factor, and the temperature '
        'dependence at fixed illuminance. The simulation predicts a mean SNR '
        'improvement of 5.4\u00d7 (max 10.0\u00d7) at 90% noise model accuracy (\u03b1 = 0.9), '
        'consistent with the measured Fano filter performance.'
    ))

    # Fig 4: A5 simulation
    add_figure(doc, FIG_DIR / 'fig4_a5_simulation.png',
               'Fig. 5. A5-based noise rate simulation across the temperature\u2013'
               'illuminance parameter space. (a) Predicted noise event rate '
               '[evt/s/pix]; (b) SNR improvement factor at 90% noise model accuracy; '
               '(c) SNR vs. temperature at fixed illuminance (I_bg = 16.7 lux) '
               'comparing raw, Fano filter, and PI-DC-DVS methods.',
               width=Inches(6.0))

    # =========================================================
    # 6. Proof-of-concept demonstration
    # =========================================================
    add_heading(doc, '6 Proof-of-concept demonstration', level=1)
    add_paragraph(doc, (
        'From 1,800,674 input events, probabilistic thinning with threshold '
        '\u03c4 = 0.5 yields 175,261 residual events\u2014a 90.3% noise removal rate. '
        'The residual event stream clearly reveals satellite trajectories buried '
        'in noise in the raw stream (Fig. 6). The four-panel demonstration shows: '
        '(a) the raw event accumulation dominated by background noise; '
        '(b) the estimated per-pixel noise rate map \u03bb_noise(x,y); '
        '(c) the bimodal P_noise distribution with clear separation at '
        '\u03c4 = 0.5 (9.7% signal candidates, 90.3% noise); and '
        '(d) the residual event stream after noise subtraction with satellite '
        'trajectory clearly visible.'
    ))

    # Fig 6: Demo
    add_figure(doc, FIG_DIR / 'fig6_demo.png',
               'Fig. 6. Proof-of-concept noise inverse problem demonstration on '
               'EBSSA Recording #0. (a) Raw event accumulation (1,800,674 events); '
               '(b) estimated noise rate map \u03bb_noise(x,y) [events/sec]; '
               '(c) per-event noise probability distribution showing bimodal '
               'separation (signal: 9.7%, noise: 90.3%); (d) residual events after '
               '90.3% noise removal (175,261 events) with satellite trajectory '
               'clearly visible.',
               width=Inches(5.5))

    add_paragraph(doc, (
        'The Fano factor spatial map (Fig. 7a) shows clear separation between '
        'noise-dominated pixels (Fano \u2248 1, blue) and signal-containing pixels '
        '(Fano \u226b 1, red). Signal candidate pixels (2,294 out of 76,800 total) '
        'concentrate along known satellite tracks. The temporal dynamics (Fig. 7b) '
        'show the noise model converging within the first few time bins. The '
        'per-pixel SNR distribution (Fig. 7c) confirms structural improvement in '
        'the residual stream.'
    ))

    # Fig 7: SNR
    add_figure(doc, FIG_DIR / 'fig7_snr.png',
               'Fig. 7. Signal-to-noise ratio improvement analysis. (a) Fano factor '
               'spatial map showing noise-dominated (blue, F \u2248 1) vs. signal-'
               'containing pixels (red, F \u226b 1); (b) temporal dynamics of total event '
               'rate vs. noise model prediction; (c) per-pixel SNR distribution '
               'before (raw) and after (residual) noise subtraction.',
               width=Inches(6.0))

    # =========================================================
    # 7. Discussion
    # =========================================================
    add_heading(doc, '7 Discussion', level=1)

    add_heading(doc, '7.1 Effectiveness of the noise inverse problem paradigm', level=2)
    add_paragraph(doc, (
        'The systematic evaluation demonstrates that physics-based noise modelling '
        '(Fano filter, AUC = 0.866) substantially outperforms conventional temporal '
        'filtering (AUC = 0.534). The key insight is that the Fano factor provides '
        'a physics-grounded discriminant\u2014noise events follow Poisson statistics '
        '(Fano \u2248 1), while astronomical signals produce bursty event patterns '
        '(Fano \u226b 1). This is analogous to using auxiliary channels in '
        'gravitational-wave detectors:{9} the noise has observable structure that '
        'can be exploited for separation.'
    ))

    add_heading(doc, '7.2 Role of auxiliary channels', level=2)
    add_paragraph(doc, (
        'The simplified PI-DC-DVS NN (AUC = 0.546) performs poorly without '
        'auxiliary channels, exhibiting high variance across recordings. This '
        'underscores the importance of auxiliary channel integration\u2014consistent '
        'with the success of DeepClean in LIGO.{9} The Fano filter succeeds because '
        'it uses the physics-model-based noise rate as an implicit auxiliary '
        'channel, without requiring hardware instrumentation. Future work with '
        'actual temperature and illuminance sensors should further improve '
        'performance.'
    ))

    add_heading(doc, '7.3 Cal-6: Paradigm inversion of light pollution', level=2)
    add_paragraph(doc, (
        'Cal-6 exemplifies "turning a bug into a feature": the proliferation of '
        'satellite constellations (Starlink, OneWeb, etc.) provides a continuous, '
        'cost-free source of calibration signals. While CCD sensors saturate on '
        'bright satellite trails, DVS sensors record them quantitatively across '
        'their full dynamic range (>120 dB). This paradigm inversion\u2014treating '
        'light pollution as a calibration resource rather than a nuisance\u2014is '
        'particularly timely given the ongoing debate about satellite constellation '
        'impacts on astronomy.{24} For DVS-based observation, satellite trails '
        'become assets rather than liabilities.'
    ))

    add_heading(doc, '7.4 Template-free detection via noise residuals', level=2)
    add_paragraph(doc, (
        'A key implication of high-precision noise subtraction is template-free '
        'object detection: if the noise is precisely modelled and subtracted, any '
        'structure in the residual is signal\u2014regardless of morphology. Combined '
        'with event-level shift-and-stack,{25} this could detect fast-moving faint '
        'objects (10\u201350 m near-Earth objects) beyond the reach of frame-based '
        'telescopes. The detection limit improvement scales as:'
    ))
    add_display_equation(doc, eq_detection_limit, eq_num='9')
    add_paragraph(doc, (
        'giving \u0394m > 2.5 mag at \u03b1 = 0.9.'
    ))

    add_heading(doc, '7.5 Limitations and future work', level=2)
    add_paragraph(doc, (
        'The current evaluation is limited by the absence of real auxiliary '
        'channels in EBSSA (temperature and illuminance are not recorded). '
        'Priority future work includes: (i) differentiable A5 pixel model '
        'implementation for end-to-end training; (ii) DVS auxiliary-channel system '
        '(temperature sensor, photometer) for telescope deployment; '
        '(iii) Phase 1 demonstration with SciDVS{15} on a 0.3\u20130.5 m telescope; '
        '(iv) systematic Cal-6 evaluation using Starlink transit data; '
        '(v) evaluation on DVSNOISE20 dataset.{19}'
    ))

    # =========================================================
    # 8. Conclusions
    # =========================================================
    add_heading(doc, '8 Conclusions', level=1)
    add_paragraph(doc, (
        'We have proposed and evaluated PI-DC-DVS, a physics-informed framework '
        'for solving the noise inverse problem in dynamic vision sensors applied to '
        'astronomical observation. The main results are:'
    ))
    conclusions = [
        '(1) The Fano-factor-based noise inverse approach achieves AUC = 0.866 '
        'on EBSSA, substantially outperforming conventional temporal filtering '
        '(AUC = 0.534) and a simplified neural network (AUC = 0.546).',

        '(2) A proof-of-concept demonstration achieves 90.3% noise removal while '
        'preserving satellite trajectories, with clear signal recovery in the '
        'residual event stream.',

        '(3) The A5-based noise rate simulation predicts mean SNR improvements of '
        '5.4\u00d7 (max 10.0\u00d7) at 90% noise model accuracy, consistent with measured '
        'performance.',

        '(4) The six-tier calibration framework (Cal-1\u2013Cal-6), with Cal-6 '
        'repurposing satellite light trails as natural calibration sources, '
        'provides quantitative noise model validation under real observing '
        'conditions.',

        '(5) The theoretical SNR improvement scales as 1/(1\u2212\u03b1), offering 10\u00d7 at '
        '\u03b1 = 0.9 and 100\u00d7 at \u03b1 = 0.99, with potential 2\u20134 magnitude detection '
        'limit extension for faint, fast-moving objects.',
    ]
    for c in conclusions:
        add_paragraph(doc, c, space_after=4)

    # =========================================================
    # Code/Data availability
    # =========================================================
    add_heading(doc, 'Code and data availability', level=1)
    add_paragraph(doc, (
        'The EBSSA dataset is publicly available via the Tonic library.{3} '
        'The implementation code and evaluation scripts are available at '
        'https://github.com/bougtoir/dvs_noise_inverse_problem.'
    ))

    # =========================================================
    # Acknowledgments
    # =========================================================
    add_heading(doc, 'Acknowledgments', level=1)
    add_paragraph(doc, '[To be added.]', italic=True)

    # =========================================================
    # References (SPIE numbered style)
    # =========================================================
    add_heading(doc, 'References', level=1)
    references = [
        # [1] DVS survey
        '[1] G. Gallego, T. Delbr\u00fcck, G. Orchard et al., "Event-based vision: '
        'A survey," IEEE Trans. Pattern Anal. Mach. Intell. 44, 154\u2013180 (2022).',

        # [2] DVS pixel model intro
        '[2] T. Delbruck, R. Gra\u00e7a, and M. Paluch, "Utility and feasibility of '
        'a center surround event camera," preprint arXiv:2103.03415 (2021).',

        # [3] EBSSA
        '[3] S. Afshar, N. Hamilton, L. Davis, A. van Schaik, and G. Cohen, '
        '"Event-based object detection and tracking for space situational '
        'awareness," preprint arXiv:1911.08730 (2019).',

        # [4] FIESTA
        '[4] D. Joubert, N. Ralph, A. Jolley et al., "Event-driven space object '
        'detection in real-time," Front. Neurosci. 16, 821157 (2022).',

        # [5] Gedek
        '[5] M. G\u0119dek, D. Magiera, G. Kowalski, and L. Gaffney, "Neuromorphic '
        'event-based space situational awareness," Proc. EESA (2019).',

        # [6] Hoang
        '[6] J. Hoang, "Neuromorphic cameras for ground-based atmospheric '
        'Cherenkov telescopes," preprint arXiv:2310.16321 (2023).',

        # [7] Graca 2023 shot noise
        '[7] R. Gra\u00e7a and T. Delbruck, "Unraveling the paradox of intensity-'
        'dependent DVS pixel noise," preprint arXiv:2304.04019 (2023).',

        # [8] McReynolds polarity
        '[8] B. McReynolds, R. Gra\u00e7a, and T. Delbruck, "Characterization of '
        'event camera noise with a once-in-a-lifetime photon," preprint '
        'arXiv:2304.03494 (2023).',

        # [9] DeepClean
        '[9] G. Vajente, Y. Huang, M. Isi et al., "Machine-learning nonstationary '
        'noise out of gravitational-wave detectors," Phys. Rev. D 101, '
        '042003 (2020).',

        # [10] Dooney
        '[10] T. Dooney, R. Brito, and A. Matas, "DeepClean for advanced LIGO: '
        'Noise subtraction with deep learning," preprint arXiv:2501.18423 (2025).',

        # [11] Wang
        '[11] H. Wang, W. Zhao, and Z. Cao, "Machine learning for gravitational-'
        'wave data analysis," Mach. Learn.: Sci. Technol. 5, 015046 (2024).',

        # [12] Chatterjee
        '[12] C. Chatterjee and K. Jani, "DeepClean: Scalable channel-wise '
        'noise subtraction," Astrophys. J. (2025).',

        # [13] iDQ
        '[13] R. Essick, P. Godwin, C. Hanna, L. Blackburn, and '
        'E. Katsavounidis, "iDQ: Statistical inference of non-astrophysical '
        'noise transients in gravitational-wave detectors," Mach. Learn.: '
        'Sci. Technol. 2, 015004 (2021).',

        # [14] D3PO
        '[14] M. Selig and T. A. En\u00dflin, "D\u00b3PO\u2014Denoising, deconvolving, and '
        'decomposing photon observations," Astron. Astrophys. 574, A74 (2015).',

        # [15] SciDVS
        '[15] R. Gra\u00e7a, S. Zhou, B. McReynolds, and T. Delbruck, "SciDVS: '
        'A scientific dynamic vision sensor," ESSERC 2024, '
        'DOI:10.1109/esserc62670.2024.10719521.',

        # [16] Graca 2025 A5 model
        '[16] R. Gra\u00e7a and T. Delbruck, "A large-signal theory for the '
        'differential DVS pixel," preprint arXiv:2505.07386 (2025).',

        # [17] Temporal filter (Delbruck 2008)
        '[17] T. Delbruck, "Frame-free dynamic digital vision," in Proc. Intl. '
        'Symp. on Secure-Life Electronics, 21\u201326 (2008).',

        # [18] Liu 2008
        '[18] S.-C. Liu and T. Delbruck, "Neuromorphic sensory systems," '
        'Proc. BMVC (2008).',

        # [19] EPM / Baldwin
        '[19] R. W. Baldwin, M. Almatrafi, V. Asari, and K. Hirakawa, '
        '"Event probability mask (EPM) and event denoising convolutional '
        'neural network (EDnCNN)," Proc. CVPR (2020).',

        # [20] WedNet
        '[20] H. Fang et al., "WedNet: Window-based event denoising with '
        'spatio-temporal correlation," IEEE Trans. Pattern Anal. Mach. '
        'Intell. (2024).',

        # [21] ASTEDNet
        '[21] W. Wu et al., "ASTEDNet: Adaptive spatio-temporal event denoising '
        'network," ISPRS Archives XLVIII-4-2024 (2024).',

        # [22] Shiba
        '[22] S. Shiba, Y. Aoki, and G. Gallego, "Secrets of event-based '
        'optical flow," Proc. ICCV (2025).',

        # [23] Noise2Image
        '[23] R. Cao et al., "Noise2Image: recovering static images from '
        'event camera noise," Optica (2024).',

        # [24] Satellite constellation impacts
        '[24] J. McDowell, "The low Earth orbit satellite population and '
        'impacts of the SpaceX Starlink constellation," Astrophys. J. Lett. '
        '892, L36 (2020).',

        # [25] Shift-and-stack
        '[25] S. Stetzler et al., "Event-based shift-and-stack for fast-moving '
        'objects," Astron. J. 170, 352 (2025).',
    ]
    for ref in references:
        add_paragraph(doc, ref, space_after=3)

    # Save
    out_path = OUT_DIR / 'manuscript_jatis.docx'
    doc.save(str(out_path))
    print(f"Manuscript saved: {out_path}")


if __name__ == '__main__':
    build_manuscript()

#!/usr/bin/env python3
"""
Generate JATIS cover letter as docx.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUT_DIR = Path(__file__).resolve().parent


def build_cover_letter():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # Date
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run('[Date]')

    # Addressee
    p = doc.add_paragraph()
    run = p.add_run(
        'Dr. Megan Eckart\n'
        'Editor-in-Chief\n'
        'Journal of Astronomical Telescopes, Instruments, and Systems\n'
        'SPIE'
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run('Dear Dr. Eckart,')

    # Body
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(
        'We are pleased to submit the enclosed manuscript entitled '
        '"Solving the noise inverse problem in dynamic vision sensors for faint '
        'astronomical object detection" for consideration as a Regular Paper in '
        'the Journal of Astronomical Telescopes, Instruments, and Systems.'
    )

    p = doc.add_paragraph()
    run = p.add_run(
        'Dynamic Vision Sensors (DVS) offer unique advantages for space '
        'situational awareness—microsecond timing, >120 dB dynamic range, and '
        'sparse output—but suffer from background activity noise that overwhelms '
        'faint astronomical signals. This paper introduces a physics-informed '
        'framework that treats DVS noise removal as an inverse problem, analogous '
        'to the DeepClean paradigm proven in gravitational-wave astronomy. '
        'Our approach integrates a circuit-level parametric noise model (A5) with '
        'Bayesian inference to achieve per-event noise classification.'
    )

    p = doc.add_paragraph()
    run = p.add_run(
        'Key contributions include: (1) A Fano-factor-based noise inverse approach '
        'achieving ROC-AUC = 0.866 on the EBSSA dataset, substantially outperforming '
        'conventional temporal filtering (AUC = 0.534); (2) A proof-of-concept '
        'demonstration with 90.3% noise removal and clear satellite trajectory '
        'recovery; (3) A six-tier calibration framework, with a novel Cal-6 tier '
        'that repurposes satellite light trails—conventionally regarded as light '
        'pollution—as natural calibration sources exploiting DVS\'s wide dynamic '
        'range; and (4) Quantitative SNR improvement predictions (mean 5.4×, '
        'max 10.0×) from the A5 noise model simulation.'
    )

    p = doc.add_paragraph()
    run = p.add_run(
        'We believe this work is well-suited for JATIS because it addresses a '
        'fundamental instrumentation challenge—sensor noise characterisation and '
        'calibration—for an emerging detector technology with direct applications '
        'to astronomical observation. The Cal-6 concept of repurposing satellite '
        'constellations for calibration is particularly timely given ongoing '
        'developments in both neuromorphic sensors and mega-constellation '
        'deployments.'
    )

    p = doc.add_paragraph()
    run = p.add_run(
        'This manuscript has not been published previously and is not under '
        'consideration elsewhere. All authors have approved the submission.'
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run('Sincerely,')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run('[Corresponding author name and affiliation]')
    run.italic = True

    # Save
    out_path = OUT_DIR / 'cover_letter_jatis.docx'
    doc.save(str(out_path))
    print(f"Cover letter saved: {out_path}")


if __name__ == '__main__':
    build_cover_letter()

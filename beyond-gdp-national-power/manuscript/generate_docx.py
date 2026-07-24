#!/usr/bin/env python3
"""Generate E&P (Economics & Politics, Wiley) submission docx from manuscript.tex and references.bib."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

MANUSCRIPT = Path(__file__).parent / "manuscript.tex"
BIB_FILE = Path(__file__).parent / "references.bib"
FIG_DIR = Path(__file__).parent / "figures"
OUTPUT = Path(__file__).parent / "manuscript.docx"


# --- Parse references.bib ---
def extract_braced(s, start):
    """Extract content between balanced braces starting at position start (which is '{')."""
    depth = 0
    i = start
    while i < len(s):
        if s[i] == '{':
            depth += 1
        elif s[i] == '}':
            depth -= 1
            if depth == 0:
                return s[start+1:i], i + 1
        i += 1
    return s[start+1:], len(s)


def parse_bib(bib_path):
    """Parse .bib file into dict of key -> formatted reference string."""
    text = bib_path.read_text()
    entries = {}
    for m in re.finditer(r'@(\w+)\{(\w+),\s*\n(.*?)\n\}', text, re.DOTALL):
        entry_type = m.group(1)
        key = m.group(2)
        body = m.group(3)
        fields = {}
        # Parse fields with balanced brace matching
        for fm in re.finditer(r'(\w+)\s*=\s*\{', body):
            field_name = fm.group(1)
            brace_start = fm.end() - 1  # position of '{'
            value, _ = extract_braced(body, brace_start)
            fields[field_name] = value
        # Format reference
        authors = fields.get('author', '')
        # Clean LaTeX from authors
        authors = authors.replace("\\'{i}", "í").replace("{", "").replace("}", "")
        title = fields.get('title', '').replace("{", "").replace("}", "")
        year = fields.get('year', '')
        journal = fields.get('journal', '')
        volume = fields.get('volume', '')
        pages = fields.get('pages', '')
        publisher = fields.get('publisher', '')
        address = fields.get('address', '')
        doi = fields.get('doi', '')
        booktitle = fields.get('booktitle', '')
        editor = fields.get('editor', '')
        number = fields.get('number', '')

        if entry_type == 'article':
            ref = f"{authors}. {title}. {journal}"
            if volume:
                ref += f" {volume}"
            if number:
                ref += f"({number})"
            if pages:
                ref += f", {pages}"
            ref += f", {year}."
        elif entry_type == 'book':
            ref = f"{authors}. {title}. {address}: {publisher}, {year}."
        elif entry_type == 'incollection':
            ref = f"{authors}. {title}. In: {editor} (eds.), {booktitle}. {address}: {publisher}, {year}."
            if pages:
                ref += f" pp. {pages}."
        else:
            ref = f"{authors}. {title}. {year}."
        if doi:
            ref += f" https://doi.org/{doi}"
        # Clean up formatting
        ref = ref.replace('--', '–')  # en-dash for page ranges
        entries[key] = {'formatted': ref, 'authors': authors, 'year': year}
    return entries


def format_author_short(authors_str):
    """Get short author form for in-text citation."""
    parts = authors_str.split(" and ")
    if len(parts) == 1:
        # Single author: last name
        return parts[0].split(",")[0].strip()
    elif len(parts) == 2:
        a1 = parts[0].split(",")[0].strip()
        a2 = parts[1].split(",")[0].strip()
        return f"{a1} and {a2}"
    else:
        return parts[0].split(",")[0].strip() + " et al."


# --- Parse and convert LaTeX ---
def tex_to_docx():
    tex = MANUSCRIPT.read_text()
    refs = parse_bib(BIB_FILE)

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0

    # Extract title
    title_m = re.search(r'\\title\[.*?\]\{(.+?)\}', tex)
    title = title_m.group(1) if title_m else "Untitled"

    # Title page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()  # blank line

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("[Author Name]").italic = True
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("[Affiliation]").italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Extract abstract (use balanced brace matching)
    abs_pos = tex.find('\\abstract{')
    if abs_pos >= 0:
        brace_pos = tex.index('{', abs_pos + len('\\abstract'))
        abs_text, _ = extract_braced(tex, brace_pos)
        p = doc.add_paragraph()
        run = p.add_run("Abstract")
        run.bold = True
        run.font.size = Pt(12)
        abs_text = clean_latex(abs_text.strip(), refs)
        p = doc.add_paragraph(abs_text)
        p.paragraph_format.first_line_indent = Pt(0)

    # Extract keywords (use balanced brace matching)
    kw_pos = tex.find('\\keywords{')
    if kw_pos >= 0:
        brace_pos = tex.index('{', kw_pos + len('\\keywords'))
        kw_text, _ = extract_braced(tex, brace_pos)
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Keywords: ")
        run.bold = True
        kw_text = kw_text.strip()
        kw_text = re.sub(r'\\\\', '; ', kw_text)
        kw_text = kw_text.replace('\n', ' ')
        p.add_run(kw_text)

    # JEL classification (from end of document)
    jel_m = re.search(r'\\subsection\*\{JEL Classification\}\s*\n(.+)', tex)
    if jel_m:
        p = doc.add_paragraph()
        run = p.add_run("JEL Classification: ")
        run.bold = True
        p.add_run(jel_m.group(1).strip())

    doc.add_page_break()

    # --- Body text ---
    # Extract body between \begin{document} and \end{document}
    body_m = re.search(r'\\begin\{document\}(.*?)\\end\{document\}', tex, re.DOTALL)
    body = body_m.group(1) if body_m else ""

    # Remove everything before first \section (title, author, affil, abstract, etc.)
    first_section = re.search(r'\\section\{', body)
    if first_section:
        body = body[first_section.start():]

    # Remove Statements and Declarations section onward
    body = re.sub(r'\\section\*\{Statements and Declarations\}.*', '', body, flags=re.DOTALL)
    # Remove \bibliography command
    body = re.sub(r'\\bibliography\{.*?\}', '', body)

    # Process body
    process_body(doc, body, refs)

    # --- Statements and Declarations ---
    doc.add_page_break()
    p = doc.add_heading("Statements and Declarations", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Funding: ")
    run.bold = True
    p.add_run("[To be completed by author]")

    p = doc.add_paragraph()
    run = p.add_run("Competing Interests: ")
    run.bold = True
    p.add_run("The author declares no competing interests.")

    p = doc.add_paragraph()
    run = p.add_run("Data Availability: ")
    run.bold = True
    p.add_run("The complete dataset and analysis code are available at [repository URL]. "
              "Supplementary Table S1 provides the full dataset of 96 polities with all coded variables.")

    p = doc.add_paragraph()
    run = p.add_run("JEL Classification: ")
    run.bold = True
    p.add_run("D72, N40, N70, F50, O33, C25")

    # --- References ---
    doc.add_page_break()
    doc.add_heading("References", level=1)

    # Sort alphabetically by author, then by year for same authors
    sorted_refs = sorted(refs.items(), key=lambda x: (x[1]['authors'], x[1].get('year', '')))
    for key, ref_data in sorted_refs:
        p = doc.add_paragraph(ref_data['formatted'])
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.space_after = Pt(6)

    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")


def process_body(doc, body, refs):
    """Process LaTeX body into docx paragraphs."""
    # Split into blocks (sections, paragraphs, figures, tables)
    lines = body.split('\n')
    current_para = []
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Section headers
        sec_m = re.match(r'\\section\{(.+?)\}(?:\\label\{.*?\})?', line)
        if sec_m:
            flush_paragraph(doc, current_para, refs)
            current_para = []
            title = clean_latex(sec_m.group(1), refs)
            doc.add_heading(title, level=1)
            i += 1
            continue

        subsec_m = re.match(r'\\subsection\{(.+?)\}(?:\\label\{.*?\})?', line)
        if subsec_m:
            flush_paragraph(doc, current_para, refs)
            current_para = []
            title = clean_latex(subsec_m.group(1), refs)
            doc.add_heading(title, level=2)
            i += 1
            continue

        # Paragraph headers
        par_m = re.match(r'\\paragraph\{(.+?)\}', line)
        if par_m:
            flush_paragraph(doc, current_para, refs)
            current_para = []
            title = clean_latex(par_m.group(1), refs)
            p = doc.add_paragraph()
            run = p.add_run(title + (" " if title.endswith(".") else ". "))
            run.bold = True
            # Rest of line after \paragraph{...}
            rest = re.sub(r'\\paragraph\{.+?\}', '', line).strip()
            if rest:
                current_para.append(rest)
            i += 1
            continue

        # Figure environments
        if '\\begin{figure}' in line:
            flush_paragraph(doc, current_para, refs)
            current_para = []
            fig_lines = []
            while i < len(lines) and '\\end{figure}' not in lines[i]:
                fig_lines.append(lines[i])
                i += 1
            if i < len(lines):
                fig_lines.append(lines[i])
            i += 1
            insert_figure(doc, fig_lines, refs)
            continue

        # Table environments
        if '\\begin{table}' in line:
            flush_paragraph(doc, current_para, refs)
            current_para = []
            tab_lines = []
            while i < len(lines) and '\\end{table}' not in lines[i]:
                tab_lines.append(lines[i])
                i += 1
            if i < len(lines):
                tab_lines.append(lines[i])
            i += 1
            insert_table(doc, tab_lines, refs)
            continue

        # Blank line = paragraph break
        if not line:
            flush_paragraph(doc, current_para, refs)
            current_para = []
            i += 1
            continue

        # Skip label-only lines
        if re.match(r'\\label\{.*?\}$', line):
            i += 1
            continue

        current_para.append(line)
        i += 1

    flush_paragraph(doc, current_para, refs)


def flush_paragraph(doc, lines, refs):
    """Join lines into paragraph and add to doc."""
    if not lines:
        return
    text = ' '.join(lines)
    text = clean_latex(text, refs)
    if text.strip():
        doc.add_paragraph(text)


def clean_latex(text, refs):
    """Convert LaTeX markup to plain text with Harvard-style citations."""
    # Handle \citet{key} -> Author (year)
    def citet_repl(m):
        keys = [k.strip() for k in m.group(1).split(',')]
        parts = []
        for key in keys:
            if key in refs:
                author = format_author_short(refs[key]['authors'])
                year = refs[key]['year']
                parts.append(f"{author} ({year})")
            else:
                parts.append(key)
        return '; '.join(parts)

    # Handle \citep{key} -> (Author, year)
    def citep_repl(m):
        keys = [k.strip() for k in m.group(1).split(',')]
        parts = []
        for key in keys:
            if key in refs:
                author = format_author_short(refs[key]['authors'])
                year = refs[key]['year']
                parts.append(f"{author}, {year}")
            else:
                parts.append(key)
        return '(' + '; '.join(parts) + ')'

    # Handle \citeauthor{key}
    def citeauthor_repl(m):
        key = m.group(1).strip()
        if key in refs:
            return format_author_short(refs[key]['authors'])
        return key

    # Handle \citeyear{key}
    def citeyear_repl(m):
        key = m.group(1).strip()
        if key in refs:
            return refs[key]['year']
        return key

    text = re.sub(r'\\citet\{([^}]+)\}', citet_repl, text)
    text = re.sub(r'\\citep\{([^}]+)\}', citep_repl, text)
    text = re.sub(r'\\citeauthor\{([^}]+)\}', citeauthor_repl, text)
    text = re.sub(r'\\citeyear\{([^}]+)\}', citeyear_repl, text)

    # Cross-references
    text = re.sub(r'Table~?\\ref\{tab:([^}]+)\}', lambda m: f"Table {get_table_num(m.group(1))}", text)
    text = re.sub(r'Fig(?:ure)?~?\.?~?\\ref\{fig:([^}]+)\}', lambda m: f"Fig. {get_fig_num(m.group(1))}", text)
    text = re.sub(r'Section~?\\ref\{sec:([^}]+)\}', lambda m: f"Section {get_sec_label(m.group(1))}", text)
    text = re.sub(r'Supplementary Table~?S1', 'Supplementary Table S1', text)

    # Emphasis and formatting
    text = re.sub(r'\\emph\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\textit\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\textbf\{([^}]+)\}', r'\1', text)

    # Math mode
    text = re.sub(r'\$([^$]+)\$', lambda m: clean_math(m.group(1)), text)

    # LaTeX commands
    text = re.sub(r'\\%', '%', text)
    text = re.sub(r'\\&', '&', text)
    text = text.replace('\\\\', '')
    text = text.replace('~', ' ')
    text = text.replace('---', '—')
    text = text.replace('--', '–')
    text = text.replace("``", '"')
    text = text.replace("''", '"')
    text = text.replace("`", "'")
    text = re.sub(r'\{,\}', ',', text)
    text = re.sub(r'\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\[a-zA-Z]+\s*', '', text)  # remaining commands
    text = re.sub(r'\s+', ' ', text)

    return text.strip()


def clean_math(math_str):
    """Convert simple LaTeX math to Unicode."""
    s = math_str
    s = s.replace('\\hat{\\beta}', 'β̂')
    s = s.replace('\\beta', 'β')
    s = s.replace('\\alpha', 'α')
    s = s.replace('\\infty', '∞')
    s = s.replace('\\times', '×')
    s = s.replace('\\leq', '≤')
    s = s.replace('\\geq', '≥')
    s = s.replace('\\neq', '≠')
    s = s.replace('\\approx', '≈')
    s = s.replace('\\pm', '±')
    s = s.replace('< ', '< ')
    s = s.replace('> ', '> ')
    s = re.sub(r'\\text\{([^}]+)\}', r'\1', s)
    s = re.sub(r'\\mathrm\{([^}]+)\}', r'\1', s)
    s = s.replace('^{2}', '²')
    s = s.replace('_', '')
    s = re.sub(r'[{}]', '', s)
    s = re.sub(r'\\[a-zA-Z]+', '', s)
    return s


# Figure/table numbering
FIG_ORDER = ['conquest-rates', 'fisher-pvalues', 'closure-types', 'forest-plot']
TABLE_ORDER = ['reclass', 'scenarios', 'regression', 'psm-balance', 'stockflow-closure']

def get_fig_num(label):
    try:
        return str(FIG_ORDER.index(label) + 1)
    except ValueError:
        return "?"

def get_table_num(label):
    try:
        return str(TABLE_ORDER.index(label) + 1)
    except ValueError:
        return "?"

def get_sec_label(label):
    sec_map = {
        'results': '4', 'sensitivity': '5', 'causal': '6',
        'discussion': '7', 'conclusion': '8',
        'causal-iv': '6.1', 'causal-psm': '6.2',
        'causal-natexp': '6.3', 'causal-robustness': '6.4',
        'disc-mechanism': '7.1', 'disc-firstcontact': '7.2',
        'disc-modern': '7.3', 'disc-stockflow': '7.4',
        'disc-causal': '7.5', 'disc-limitations': '7.6',
    }
    return sec_map.get(label, label)


def extract_command_arg(content, command):
    """Extract argument of a LaTeX command using balanced brace matching."""
    pos = content.find(f'\\{command}' + '{')
    if pos < 0:
        return ""
    brace_pos = content.index('{', pos + len(f'\\{command}'))
    value, _ = extract_braced(content, brace_pos)
    return value


def insert_figure(doc, fig_lines, refs):
    """Insert figure into document."""
    content = '\n'.join(fig_lines)
    # Extract caption (balanced brace matching)
    caption = extract_command_arg(content, 'caption')
    caption = clean_latex(caption, refs)
    # Extract label to get figure number
    lab_m = re.search(r'\\label\{fig:([^}]+)\}', content)
    fig_label = lab_m.group(1) if lab_m else ""
    fig_num = get_fig_num(fig_label)
    # Extract image filename
    img_m = re.search(r'\\includegraphics.*?\{(.+?)\}', content)
    img_file = img_m.group(1) if img_m else ""

    # Find actual image file
    img_path = None
    if img_file:
        # Try figures/Figx.png
        candidate = FIG_DIR / f"Fig{fig_num}.png"
        if candidate.exists():
            img_path = candidate

    # Add figure
    if img_path and img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(5.5))

    # Add caption
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"Fig. {fig_num}. ")
    run.bold = True
    p.add_run(caption)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)


def insert_table(doc, tab_lines, refs):
    """Insert table into document."""
    content = '\n'.join(tab_lines)
    # Extract caption (balanced brace matching)
    caption = extract_command_arg(content, 'caption')
    caption = clean_latex(caption, refs)
    # Extract label
    lab_m = re.search(r'\\label\{tab:([^}]+)\}', content)
    tab_label = lab_m.group(1) if lab_m else ""
    tab_num = get_table_num(tab_label)

    # Add caption above table
    p = doc.add_paragraph()
    run = p.add_run(f"Table {tab_num}. ")
    run.bold = True
    p.add_run(caption)
    p.paragraph_format.space_before = Pt(12)

    # Parse table content
    # Extract rows between \midrule and \botrule
    rows = []
    header_row = None
    in_tabular = False
    for line in tab_lines:
        line = line.strip()
        if '\\toprule' in line:
            in_tabular = True
            continue
        if '\\midrule' in line:
            continue
        if '\\botrule' in line or '\\bottomrule' in line:
            continue
        if '\\end{tabular' in line:
            in_tabular = False
            continue
        if not in_tabular:
            continue
        if '\\\\' in line or '&' in line:
            # Parse cells
            line = re.sub(r'\\\\.*', '', line)
            cells = [clean_latex(c.strip(), refs) for c in line.split('&')]
            if header_row is None:
                header_row = cells
            else:
                rows.append(cells)

    # Create Word table
    if header_row:
        ncols = len(header_row)
        nrows = len(rows) + 1
        table = doc.add_table(rows=nrows, cols=ncols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        for j, cell_text in enumerate(header_row):
            cell = table.rows[0].cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        # Data rows
        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                if j < ncols:
                    cell = table.rows[i + 1].cells[j]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

    # Footnote if present
    fn_text = extract_command_arg(content, 'footnotetext')
    if fn_text:
        fn = clean_latex(fn_text, refs)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        run = p.add_run(f"Note: {fn}")
        run.font.size = Pt(9)
        run.italic = True

    doc.add_paragraph()  # spacing


if __name__ == '__main__':
    tex_to_docx()

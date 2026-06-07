#!/usr/bin/env python3
"""Generate cover letter docx for Economics & Politics (Wiley) submission."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = Path(__file__).parent / "cover_letter.docx"


def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # Date
    p = doc.add_paragraph()
    p.add_run("[Date]")
    doc.add_paragraph()

    # Addressee
    p = doc.add_paragraph()
    p.add_run("The Editors")
    p = doc.add_paragraph()
    run = p.add_run("Economics & Politics")
    run.italic = True
    doc.add_paragraph()

    # Opening
    p = doc.add_paragraph("Dear Professors Xu and Zauner,")
    doc.add_paragraph()

    # Body
    paras = [
        'I am pleased to submit the manuscript entitled "Network Exclusion and State Collapse: '
        'From Maritime Isolation to Technological Access Denial in the Long Run of History" for '
        'consideration by Economics & Politics.',

        'This paper addresses a core question in political economy: how does structural exclusion '
        'from dominant technological networks interact with political institutions to determine '
        'state survival? Using a comparative dataset of 96 historical polities spanning antiquity '
        'to the present, we distinguish between deliberate closure—a policy choice by ruling '
        'coalitions seeking to preserve rents (maritime bans, sakoku, bloc membership)—and what '
        'we term "technical network exclusion": involuntary disconnection from the dominant '
        'exchange network of an era due to geographic or technological constraints.',

        'We propose a two-stage political economy mechanism. First, network exclusion disrupts '
        'the flow of frontier technology, leading to economic stagnation and declining state '
        'capacity. Second, weakened economic performance erodes institutional quality and '
        'political resilience—the fiscal capacity to maintain armies, the administrative capacity '
        'to govern effectively, and the legitimacy that sustained growth provides to ruling '
        'coalitions. This framework connects the literatures on institutions and state capacity '
        '(Besley and Persson, 2011; Acemoglu and Robinson, 2006) with the literature on trade, '
        'technology diffusion, and isolation.',

        'The paper presents four principal contributions:',
    ]

    for text in paras:
        doc.add_paragraph(text)

    # Numbered contributions
    contributions = [
        'A systematic sensitivity analysis demonstrating that reclassifying seven technically '
        'excluded polities transforms a non-significant closure–conquest association '
        '(Fisher\'s exact p = 0.187) into a significant one (p = 0.020), while the core '
        'stock–flow odds ratio remains stable (OR = 1.774).',

        'A political economy analysis of closure as a policy choice, connecting selectorate '
        'theory (Bueno de Mesquita et al., 2003) to the logic of network exclusion: rulers '
        'with narrow winning coalitions may rationally choose closure despite its long-run '
        'costs, creating a tension between regime survival and state survival.',

        'Empirical evidence for a mediated causal pathway (exclusion → economic stagnation → '
        'institutional decay → political vulnerability), supported by the multivariate '
        'regression pattern in which institutional quality and external threat absorb the '
        'closure effect, and by the PSM results consistent with full mediation.',

        'Contemporary policy implications for the political economy of technology export '
        'controls—framing modern semiconductor and AI restrictions as deliberate exclusion '
        'policies whose long-run consequences can be assessed through the historical lens '
        'documented here.',
    ]

    for i, text in enumerate(contributions, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.27)
        p.add_run(f"{i}. {text}")

    paras2 = [
        '',
        'We probe the closure–conquest association with four causal identification strategies—'
        'instrumental variables, propensity score matching, a natural experiment framing, and '
        'a robustness battery including E-values, permutation tests, and leave-one-out analysis. '
        'The strongest evidence comes from the natural experiment approach: technically excluded '
        'polities show a 100% conquest rate compared with 62.3% for open polities '
        '(Fisher\'s exact p = 0.048), with no significant differences in observable covariates.',

        'The manuscript contains approximately 8,000 words, 5 tables, and 4 figures. '
        'Supplementary Table S1 provides the complete dataset. We confirm that this work has '
        'not been published or submitted elsewhere.',

        'We believe this manuscript is well suited to Economics & Politics because it addresses '
        'the interaction of political processes (policy choices regarding openness, selectorate '
        'dynamics) with economic outcomes (technology flow, state capacity) in determining '
        'political survival—a question at the core of the journal\'s scope.',

        'We suggest the following reviewers based on their expertise in political economy and '
        'long-run development:',
    ]

    for text in paras2:
        doc.add_paragraph(text)

    # Reviewer list
    reviewers = [
        'Prof. Daron Acemoglu (MIT) — political institutions, economic development, and state capacity',
        'Prof. Torsten Persson (Stockholm University) — political economics, fiscal capacity, and state formation',
        'Prof. Nathan Nunn (University of British Columbia) — long-run persistence, historical institutions, and trade',
    ]
    for r in reviewers:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.27)
        p.add_run(f"• {r}")

    doc.add_paragraph()

    # Closing
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph("[Author Name]")
    doc.add_paragraph("[Affiliation]")
    doc.add_paragraph("[Email]")

    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")


if __name__ == '__main__':
    main()

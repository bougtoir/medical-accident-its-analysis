#!/usr/bin/env python3
"""Create Highlights / Key Points file for Journal of Clinical Anesthesia submission.

Journal of Clinical Anesthesia requires 3-5 Highlights as a separate file;
each bullet should be no more than 85 characters.
"""
import json
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'documents', 'JCA')
os.makedirs(OUTPUT_DIR, exist_ok=True)

REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
with open(os.path.join(REPO_ROOT, 'output', 'ijhpm_results.json')) as f:
    R = json.load(f)

meta = R['metadata']
l008_r2 = R['codes']['L008']['multilevel']['marginal_r2']

HIGHLIGHTS = [
    "Uniform fees and regional audits separate structural from administrative variation.",
    "Regional anaesthesia variation is driven by university-hospital access, not auditing.",
    f"University hospital presence explains {l008_r2*100:.1f}% of general-anaesthesia variance.",
    "Audit-related ratio shifts account for <1% of observed interquartile range.",
    "Transferable to UHC systems with uniform fees and regional audits; less so elsewhere.",
]

# Verify character limits
too_long = [h for h in HIGHLIGHTS if len(h) > 85]
if too_long:
    raise ValueError(f"Highlights exceed 85 chars: {too_long}")

doc = Document()
for s in doc.sections:
    s.page_width = Cm(21)
    s.page_height = Cm(29.7)
    s.top_margin = Cm(2.54)
    s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(2.54)
    s.right_margin = Cm(2.54)

st = doc.styles['Normal']
st.font.name = 'Times New Roman'
st.font.size = Pt(12)
st.paragraph_format.line_spacing = 1.5
st.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Highlights")
r.bold = True
r.font.size = Pt(14)
doc.add_paragraph()

for h in HIGHLIGHTS:
    p = doc.add_paragraph(style='List Bullet')
    p.text = h
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

out = os.path.join(OUTPUT_DIR, 'highlights_JCA_EN.docx')
doc.save(out)
print(f"Saved: {out}")

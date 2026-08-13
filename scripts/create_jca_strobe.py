#!/usr/bin/env python3
"""Create STROBE checklist (cross-sectional studies) for Journal of Clinical Anesthesia submission.

Reads the generated main manuscript .docx, converts it to PDF, and infers the
page numbers where each STROBE item may be found.
"""
import json
import os
import re
import shutil
import subprocess
import tempfile
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'documents', 'JCA')
os.makedirs(OUTPUT_DIR, exist_ok=True)

REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
with open(os.path.join(REPO_ROOT, 'output', 'ijhpm_results.json')) as f:
    R = json.load(f)

meta = R['metadata']
n_areas = meta['n_areas']
n_prefectures = meta['n_prefectures']
n_univ_areas = meta['n_univ_areas']
fiscal_year = meta.get('fiscal_year', 2022)


def _which(tool):
    path = shutil.which(tool)
    if not path:
        return None
    return path


def get_manuscript_page_ranges():
    """Convert the main manuscript docx to PDF and return a dict of section/table/page ranges.

    If LibreOffice or pdftotext is unavailable, returns an empty dict so the STROBE
    checklist can still be generated with blank page-number cells.
    """
    manuscript = os.path.join(OUTPUT_DIR, 'regional_anaesthesia_JCA_EN.docx')
    if not os.path.exists(manuscript):
        raise FileNotFoundError(f"Main manuscript not found: {manuscript}")

    if not _which('libreoffice') or not _which('pdftotext'):
        print("Warning: LibreOffice and/or pdftotext not found. "
              "STROBE page-number column will be blank.")
        return {}

    tmpdir = tempfile.mkdtemp()
    try:
        subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf',
             '--outdir', tmpdir, manuscript],
            check=True, stdout=subprocess.DEVNULL, stderr=subprocess.PIPE,
        )
        pdf_name = os.path.splitext(os.path.basename(manuscript))[0] + '.pdf'
        pdf_path = os.path.join(tmpdir, pdf_name)
        if not os.path.exists(pdf_path):
            raise RuntimeError(f"LibreOffice did not produce expected PDF: {pdf_path}")

        text = subprocess.run(
            ['pdftotext', '-layout', pdf_path, '-'],
            capture_output=True, text=True, check=True,
        ).stdout
        pages = [p for p in text.split('\f') if p.strip()]

        headings = [
            'ABSTRACT', 'INTRODUCTION', 'METHODS',
            'RESULTS', 'DISCUSSION', 'CONCLUSIONS', 'REFERENCES',
            'FIGURE LEGENDS',
        ]
        heading_page = {}
        for h in headings:
            for i, page in enumerate(pages):
                # Headings are rendered as standalone uppercase lines.
                if re.search(rf'\b{re.escape(h)}\b', page):
                    heading_page[h] = i + 1
                    break

        # Document order follows the order of the headings list above.
        order = [h for h in headings if h in heading_page]
        ranges = {}
        for i, h in enumerate(order):
            start = heading_page[h]
            if i + 1 < len(order):
                nxt = heading_page[order[i + 1]]
                end = start if nxt == start else nxt - 1
            else:
                end = len(pages)
            ranges[h] = f"{start}" if start == end else f"{start}–{end}"

        # Detect embedded tables by their captions.
        for label in ['Table 1.', 'Table 2.', 'Table 3.']:
            key = label.rstrip('.')
            for i, page in enumerate(pages):
                if label in page:
                    ranges[key] = str(i + 1)
                    break

        return ranges
    except subprocess.CalledProcessError as e:
        raw = e.stderr
        if isinstance(raw, bytes):
            raw = raw.decode('utf-8', errors='ignore')
        err = raw[:500] if raw else str(e)
        raise RuntimeError(f"Failed to extract manuscript page numbers: {err}") from e
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


page_ranges = get_manuscript_page_ranges()


def page_str_for(section, item_num=""):
    """Return a page range string for a STROBE item.

    Page numbers are item-specific where possible; 'N/A' is used for
    non-applicable items and 'Title page' for the separate title-page file.
    """
    key = (section, item_num)
    if key in (('Title and abstract', '1(a)'),):
        abstract_pages = page_ranges.get('ABSTRACT', '')
        return f"1, {abstract_pages}" if abstract_pages else "1"
    if key in (('Title and abstract', '1(b)'),):
        return page_ranges.get('ABSTRACT', '')
    if section == 'Introduction':
        return page_ranges.get('INTRODUCTION', '')
    if section == 'Methods':
        if item_num == '12(d)':
            return 'N/A'
        return page_ranges.get('METHODS', '')
    if section == 'Results':
        if item_num in ('13(c)', '16(c)'):
            return 'N/A'
        if item_num == '14(a)':
            t1 = page_ranges.get('Table 1', '')
            t2 = page_ranges.get('Table 2', '')
            return f"{t1}, {t2}" if t1 and t2 else page_ranges.get('RESULTS', '')
        if item_num == '14(b)':
            return page_ranges.get('Table 2', '')
        if item_num == '15*':
            t1 = page_ranges.get('Table 1', '')
            t2 = page_ranges.get('Table 2', '')
            return f"{t1}–{t2}" if t1 and t2 else page_ranges.get('RESULTS', '')
        if item_num == '16(a)':
            return page_ranges.get('Table 3', '')
        if item_num == '16(b)':
            return page_ranges.get('FIGURE LEGENDS', '')
        return page_ranges.get('RESULTS', '')
    if section == 'Discussion':
        return page_ranges.get('DISCUSSION', '')
    if section == 'Other information':
        if item_num == '22':
            return 'Title page'
        return page_ranges.get('REFERENCES', '')
    return ''


doc = Document()
for s in doc.sections:
    s.page_width = Cm(29.7)
    s.page_height = Cm(21)
    s.top_margin = Cm(1.8)
    s.bottom_margin = Cm(1.8)
    s.left_margin = Cm(1.8)
    s.right_margin = Cm(1.8)

st = doc.styles['Normal']
st.font.name = 'Times New Roman'
st.font.size = Pt(10)
st.paragraph_format.line_spacing = 1.15
st.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("STROBE Statement — Checklist of items that should be included "
              "in reports of cross-sectional studies")
r.bold = True
r.font.size = Pt(13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Manuscript: Regional variation in anaesthesia practice in Japan "
              "under universal health coverage: structural determinants and "
              "claims-audit sensitivity")
r.italic = True
r.font.size = Pt(11)

doc.add_paragraph()

items = [
    ("Title and abstract", "1(a)",
     "Indicate the study's design with a commonly used term in the title or "
     "the abstract",
     "Title (running header: 'Anaesthesia variation in Japan'); Abstract "
     "(Methods: 'cross-sectional ecological study')"),
    ("", "1(b)",
     "Provide in the abstract an informative and balanced summary of what "
     "was done and what was found",
     "Abstract: structured subheadings Background / Methods / Results / "
     "Conclusion"),

    ("Introduction", "", "", ""),
    ("Background/rationale", "2",
     "Explain the scientific background and rationale for the investigation "
     "being reported",
     "Introduction, paragraphs 1–2"),
    ("Objectives", "3",
     "State specific objectives, including any pre-specified hypotheses",
     "Introduction, final paragraph (four pre-specified research questions)"),

    ("Methods", "", "", ""),
    ("Study design", "4",
     "Present key elements of study design early in the paper",
     "Materials and methods, '2.1 Study design and ethical considerations'"),
    ("Setting", "5",
     "Describe the setting, locations, and relevant dates, including "
     "periods of recruitment, exposure, follow-up, and data collection",
     f"Materials and methods, '2.2 Geographic units' and '2.3 Data sources' (fiscal year {fiscal_year}, {n_areas} "
     f"secondary medical areas covering all {n_prefectures} prefectures of Japan)"),
    ("Participants", "6",
     "Give the eligibility criteria, and the sources and methods of "
     "selection of participants",
     "Materials and methods, '2.3 Data sources'; aggregate ecological unit = secondary "
     "medical area; no individual-level eligibility criteria"),
    ("Variables", "7",
     "Clearly define all outcomes, exposures, predictors, potential "
     "confounders, and effect modifiers. Give diagnostic criteria, if "
     "applicable",
     "Materials and methods, '2.4 Anaesthesia procedure codes and university hospital mapping' "
     "and '2.5 Statistical analysis'"),
    ("Data sources/measurement", "8*",
     "For each variable of interest, give sources of data and details of "
     "methods of assessment (measurement). Describe comparability of "
     "assessment methods if there is more than one group",
     "Materials and methods, '2.3 Data sources and standardised claim ratios' (Cabinet Office "
     "Regional Variation Visualisation; Survey of Physicians, Dentists and "
     "Pharmacists; National Land Numerical Information)"),
    ("Bias", "9",
     "Describe any efforts to address potential sources of bias",
     "Materials and methods, '2.5 Statistical analysis' and "
     "Discussion, '4.2 Strengths and limitations'"),
    ("Study size", "10",
     "Explain how the study size was arrived at",
     f"Materials and methods, '2.3 Data sources' (all {n_areas} areas — no sampling)"),
    ("Quantitative variables", "11",
     "Explain how quantitative variables were handled in the analyses. If "
     "applicable, describe which groupings were chosen and why",
     "Materials and methods, '2.5 Statistical analysis' (standardised claim ratio as primary "
     "outcome; empirical Bayes shrinkage)"),
    ("Statistical methods", "12(a)",
     "Describe all statistical methods, including those used to control for "
     "confounding",
     "Materials and methods, '2.5 Statistical analysis' (descriptive statistics, multilevel "
     "models with prefecture random intercept, variance decomposition, "
     "empirical Bayes shrinkage)"),
    ("", "12(b)",
     "Describe any methods used to examine subgroups and interactions",
     "Results, '3.3 Sensitivity analyses against the audit hypothesis' and '3.4 Robustness: "
     "empirical Bayes shrinkage and outliers' (within-prefecture paired comparisons; "
     "alternative combined outcomes)"),
    ("", "12(c)",
     "Explain how missing data were addressed",
     "Materials and methods, '2.5 Statistical analysis' (low-volume areas masked by the data "
     "provider are reported as n and excluded from rate calculations)"),
    ("", "12(d)",
     "Describe analytical methods taking account of sampling strategy",
     "Not applicable — full population of areas"),
    ("", "12(e)",
     "Describe any sensitivity analyses",
     "Results, '3.3 Sensitivity analyses against the audit hypothesis'"),

    ("Results", "", "", ""),
    ("Participants", "13(a)",
     "Report numbers of individuals at each stage of study",
     f"Results, '3.1 Study population' ({n_areas} areas, {n_prefectures} prefectures, {n_univ_areas} "
     f"university-hospital areas)"),
    ("", "13(b)",
     "Give reasons for non-participation at each stage",
     "Results, '3.1 Study population' (areas masked for low volume reported "
     "as n by procedure)"),
    ("", "13(c)",
     "Consider use of a flow diagram",
     "Not applicable — ecological design with full population"),
    ("Descriptive data", "14(a)",
     "Give characteristics of study participants and information on "
     "exposures and potential confounders",
     "Results, '3.1 Study population and variation in anaesthesia practice'; Table 1 and Table 2"),
    ("", "14(b)",
     "Indicate number of participants with missing data for each variable "
     "of interest",
     "Table 2 column 'n'"),
    ("Outcome data", "15*",
     "Report numbers of outcome events or summary measures",
     "Results, '3.1 Study population and variation in anaesthesia practice' and Tables 1–2 (means, "
     "SD, IQR, range, coefficients of variation)"),
    ("Main results", "16(a)",
     "Give unadjusted estimates and, if applicable, confounder-adjusted "
     "estimates and their precision (e.g. 95% CI). Make clear which "
     "confounders were adjusted for and why they were included",
     "Results, '3.2 Multilevel model and university hospital effect'; Table 3 (mixed-effects coefficients "
     "with 95% confidence intervals)"),
    ("", "16(b)",
     "Report category boundaries when continuous variables were "
     "categorized",
     "Figure 1 and Figure 2 legends (quintile shading)"),
    ("", "16(c)",
     "If relevant, consider translating estimates of relative risk into "
     "absolute risk for a meaningful time period",
     "Not applicable — ecological measure is a standardised claim ratio"),
    ("Other analyses", "17",
     "Report other analyses done — e.g. analyses of subgroups and "
     "interactions, and sensitivity analyses",
     "Results, '3.3 Sensitivity analyses against the audit hypothesis' and '3.4 Robustness: "
     "empirical Bayes shrinkage and outliers' (empirical Bayes "
     "shrinkage; within-prefecture comparisons; combined codes)"),

    ("Discussion", "", "", ""),
    ("Key results", "18",
     "Summarise key results with reference to study objectives",
     "Discussion, '4.1 Statement of principal findings'"),
    ("Limitations", "19",
     "Discuss limitations of the study, taking into account sources of "
     "potential bias or imprecision. Discuss both direction and magnitude "
     "of any potential bias",
     "Discussion, '4.2 Strengths and limitations'"),
    ("Interpretation", "20",
     "Give a cautious overall interpretation of results considering "
     "objectives, limitations, multiplicity of analyses, results from "
     "similar studies, and other relevant evidence",
     "Discussion, '4.3 International generalizability and clinical relevance'"),
    ("Generalisability", "21",
     "Discuss the generalisability (external validity) of the study results",
     "Discussion, '4.3 International generalizability and clinical relevance' "
     "and '4.4 Implications for perioperative practice and policy'"),

    ("Other information", "", "", ""),
    ("Funding", "22",
     "Give the source of funding and the role of the funders for the "
     "present study and, if applicable, for the original study on which the "
     "present article is based",
     "Title page, 'Funding'"),
]

table = doc.add_table(rows=1 + len(items), cols=5)
table.style = 'Table Grid'
hdrs = ["Section / Topic", "Item #", "Recommendation",
        "Reported on page / section", "Page"]
widths = [Cm(3.8), Cm(1.4), Cm(9.0), Cm(7.0), Cm(2.2)]
for i, h in enumerate(hdrs):
    c = table.rows[0].cells[i]
    c.width = widths[i]
    c.text = ''
    p = c.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(10)

current_section = "Title and abstract"
for ri, (sect, num, rec, addr) in enumerate(items, 1):
    row = table.rows[ri].cells
    is_section_header = (num == "" and rec == "" and addr == "")
    if is_section_header:
        current_section = sect
        page = ""
    else:
        page = page_str_for(current_section, num)
    vals = [sect, num, rec, addr, page]
    for ci, val in enumerate(vals):
        row[ci].width = widths[ci]
        row[ci].text = ''
        p = row[ci].paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(9)
        if is_section_header and ci == 0:
            r.bold = True
            r.italic = True

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run(
    "*Give information separately for cases and controls in case-control "
    "studies and, if applicable, for exposed and unexposed groups in "
    "cohort and cross-sectional studies.")
r.italic = True
r.font.size = Pt(9)

p = doc.add_paragraph()
r = p.add_run(
    "Note: An Explanation and Elaboration article discusses each checklist "
    "item and gives methodological background and published examples of "
    "transparent reporting. The STROBE checklist is best used in "
    "conjunction with this article (freely available on the Web sites of "
    "PLoS Medicine, Annals of Internal Medicine, and Epidemiology). "
    "Information on the STROBE Initiative is available at "
    "www.strobe-statement.org.")
r.italic = True
r.font.size = Pt(9)

out = os.path.join(OUTPUT_DIR, 'STROBE_checklist_JCA_EN.docx')
doc.save(out)
print(f"Saved: {out}")

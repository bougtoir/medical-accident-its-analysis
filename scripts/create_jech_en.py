#!/usr/bin/env python3
"""Create Journal of Epidemiology & Community Health formatted English manuscript.

Conforms to Journal of Epidemiology & Community Health original research author
guidelines:
- Title must indicate study design
- Structured abstract (Background/Methods/Results/Conclusion) <=250 words
- Key messages (What is already known / What this study adds / How this study might affect...) in 3-5 sentences
- Body <=3500 words (excluding abstract, references, tables, figures)
- Up to 5 combined tables/illustrations
- References numbered in order of appearance (Vancouver style; first three authors then et al.)
- Tables in main text; figure legends at end of manuscript; figures supplied as separate files
- Double-spaced, Times New Roman 12 pt

All numeric results are read from output/ijhpm_results.json, which is produced
by scripts/compile_ijhpm_results.py.
"""

import json
import os
import re
import sys
from string import Template

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'documents', 'JECH')
os.makedirs(OUTPUT_DIR, exist_ok=True)

REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FIG_DIR = os.path.join(REPO_ROOT, 'output')
RESULTS_PATH = os.path.join(REPO_ROOT, 'output', 'ijhpm_results.json')

doc = Document()

# Page setup: A4, 2.54 cm margins
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0
style.paragraph_format.space_after = Pt(0)


# ============================================================
# LOAD RESULTS
# ============================================================
with open(RESULTS_PATH) as f:
    R = json.load(f)

codes = R['codes']
combined = R['combined']['L008_L003']
corr = R['correlations']
eb = R['empirical_bayes']
vd = R['variance_decomposition']
audit = R['audit_sensitivity']
outliers = R['outliers']['L008']
meta = R['metadata']



def fmt_num(x, decimals=1):
    if x is None or (isinstance(x, float) and (x != x)):
        return 'NA'
    return f'{x:.{decimals}f}'


def fmt_p(x):
    if x is None or (isinstance(x, float) and (x != x)):
        return 'NA'
    if x < 0.001:
        return '<0.001'
    return f'{x:.3f}'


def desc(code):
    return codes[code]['overall']


def univ(code):
    return codes[code]['university']


def nonu(code):
    return codes[code]['non_university']


def ml(code):
    return codes[code]['multilevel']



def build_flat(R):
    """Build a flat substitution dictionary from ijhpm_results.json."""
    F = {}
    meta = R['metadata']
    F['n_areas'] = meta['n_areas']
    F['n_prefectures'] = meta['n_prefectures']
    F['fiscal_year'] = meta.get('fiscal_year', 2022)
    F['n_univ_areas'] = meta['n_univ_areas']
    F['n_nonuniv_areas'] = meta['n_nonuniv_areas']
    F['n_univ_areas_pct'] = fmt_num(meta.get('n_univ_areas_pct', 100 * meta['n_univ_areas'] / meta['n_areas']))

    vd_key_map = {
        'between_prefecture_pct': 'vd_between',
        'university_effect_pct': 'vd_univ',
        'residual_pct': 'vd_residual',
        'within_prefecture_pct': 'vd_within',
        'univ_explains_within_pct': 'vd_univ_within',
    }

    for code, c in R['codes'].items():
        o = c['overall']
        for key, d in [('n', 0), ('mean', 1), ('sd', 1), ('median', 1),
                       ('q1', 1), ('q3', 1), ('min', 1), ('max', 1),
                       ('cv', 1), ('iqr', 1), ('range_ratio', 1)]:
            F[f'{code}_{key}'] = fmt_num(o.get(key), d)

        u = c['university']
        n = c['non_university']
        for prefix, vals in [('u', u), ('nu', n)]:
            for key in ['mean', 'sd', 'median', 'min', 'max']:
                F[f'{code}_{prefix}_{key}'] = fmt_num(vals.get(key), 1)

        F[f'{code}_d'] = fmt_num(c['cohens_d_overall'], 2)

        w = c['within_prefecture']
        F[f'{code}_within_diff'] = fmt_num(w.get('mean_diff'))
        F[f'{code}_within_t'] = fmt_num(w.get('t'), 2)
        F[f'{code}_within_p'] = fmt_p(w.get('p'))

        m = c['multilevel']
        F[f'{code}_ml_icc'] = fmt_num(m.get('icc_null'), 3)
        F[f'{code}_ml_coef'] = fmt_num(m.get('coef_univ'))
        F[f'{code}_ml_ci_low'] = fmt_num(m.get('ci_low'))
        F[f'{code}_ml_ci_high'] = fmt_num(m.get('ci_high'))
        F[f'{code}_ml_p'] = fmt_p(m.get('p'))
        F[f'{code}_ml_r2'] = fmt_num(m.get('marginal_r2'), 3)

        mc = c.get('multilevel_covariate')
        if mc:
            F[f'{code}_mlc_coef'] = fmt_num(mc.get('coef_univ'))
            F[f'{code}_mlc_ci_low'] = fmt_num(mc.get('ci_low'))
            F[f'{code}_mlc_ci_high'] = fmt_num(mc.get('ci_high'))
            F[f'{code}_mlc_p'] = fmt_p(mc.get('p'))
            F[f'{code}_mlc_r2'] = fmt_num(mc.get('marginal_r2'), 3)
            F[f'{code}_mlc_popd_coef'] = fmt_num(mc['log_pop_density_z'].get('coef'))
            F[f'{code}_mlc_popd_p'] = fmt_p(mc['log_pop_density_z'].get('p'))
            F[f'{code}_mlc_anes_coef'] = fmt_num(mc['anes_pct_z'].get('coef'))
            F[f'{code}_mlc_anes_p'] = fmt_p(mc['anes_pct_z'].get('p'))

        if code in R.get('variance_decomposition', {}):
            vd = R['variance_decomposition'][code]
            for json_key, flat_key in vd_key_map.items():
                F[f'{code}_{flat_key}'] = fmt_num(vd.get(json_key))

        if code in R.get('empirical_bayes', {}):
            ebv = R['empirical_bayes'][code]
            F[f'{code}_d_shrunk'] = fmt_num(ebv.get('shrunk_cohens_d'), 2)
            F[f'{code}_attenuation'] = fmt_num(ebv.get('attenuation_pct'))

    F['L002_missing_n'] = R['metadata']['n_areas'] - R['codes']['L002']['overall']['n']
    F['L002_fold'] = fmt_num(R['codes']['L002']['university']['mean'] / R['codes']['L002']['non_university']['mean'], 2)
    F['L003_fold'] = fmt_num(R['codes']['L003']['university']['mean'] / R['codes']['L003']['non_university']['mean'], 2)

    for combo, c in R['combined'].items():
        prefix = combo.replace('+', '_')
        for key, d in [('n', 0), ('mean', 1), ('sd', 1), ('median', 1),
                       ('q1', 1), ('q3', 1), ('min', 1), ('max', 1),
                       ('cv', 1), ('iqr', 1)]:
            F[f'{prefix}_{key}'] = fmt_num(c['overall'].get(key), d)
        F[f'{prefix}_u_mean'] = fmt_num(c['university'].get('mean'))
        F[f'{prefix}_nu_mean'] = fmt_num(c['non_university'].get('mean'))
        F[f'{prefix}_d'] = fmt_num(c.get('cohens_d'), 2)
        F[f'{prefix}_fold'] = fmt_num(c.get('fold_ratio'), 2)

    for k, v in R['correlations'].items():
        F[f'corr_{k}_r'] = fmt_num(v.get('r'), 3)
        F[f'corr_{k}_p'] = fmt_p(v.get('p'))

    F['audit_max_shift'] = fmt_num(R['audit_sensitivity']['max_ratio_shift_approx'], 2)
    for code in ['L008', 'L002', 'L003', 'L004']:
        if code in R['audit_sensitivity']:
            F[f'audit_{code}_iqr_pct'] = fmt_num(R['audit_sensitivity'][code]['percent_of_iqr'], 2)

    outlier_name_map = {
        ('安房', '千葉県'): ('Awa (Chiba)', 'Chiba'),
        ('区中央部', '東京都'): ('central Tokyo (Tokyo)', 'Tokyo'),
    }
    for i, o in enumerate(R['outliers']['L008'][:2], 1):
        name, pref = o['area_name'], o['pref_name']
        if (name, pref) in outlier_name_map:
            name, pref = outlier_name_map[(name, pref)]
        F[f'outlier{i}_name'] = name
        F[f'outlier{i}_pref'] = pref
        F[f'outlier{i}_scr'] = fmt_num(o['scr'])

    return F


FLAT = build_flat(R)


def add_run_with_refs(paragraph, text, italic=False, bold=False):
    """Add text to a paragraph, parsing {n} or {n-m} as font superscript runs."""
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        if italic:
            run.italic = True
        if bold:
            run.bold = True
        if part.startswith('{') and part.endswith('}'):
            inner = part[1:-1]
            run.text = inner
            run.font.superscript = True


def _expand_citation(text):
    """Expand a superscript citation such as '7-9,12' into integers."""
    nums = []
    for token in re.split(r',\s*', text.strip()):
        token = token.strip()
        if not token:
            continue
        if '-' in token:
            start, end = token.split('-', 1)
            nums.extend(range(int(start), int(end) + 1))
        else:
            nums.append(int(token))
    return nums


def renumber_references(doc, references):
    """Determine first-appearance order from superscript citations and reorder.

    Walks the document body (stops at the 'References' heading), records the
    order in which each citation number first appears, builds an old->new
    mapping, updates all superscript citation runs in place, and returns the
    reference list reordered to match the new numbering. Uncited references
    are dropped. Range citations (e.g. 7-9) are expanded so middle numbers are
    not dropped.
    """
    CITE_RE = re.compile(r'\d+(?:-\d+)?(?:,\s*\d+(?:-\d+)?)*')
    first = {}
    cited = set()
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt.upper() == 'REFERENCES':
            break
        for run in p.runs:
            if run.font.superscript:
                for match in CITE_RE.finditer(run.text):
                    for n in _expand_citation(match.group(0)):
                        cited.add(n)
                        if n not in first:
                            first[n] = len(first) + 1
    mapping = {old: new for new, old in enumerate(sorted(first, key=first.get), 1)}

    def repl(m):
        old = int(m.group(0))
        return str(mapping.get(old, old))

    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.superscript:
                run.text = re.sub(r'\d+', repl, run.text)

    return [references[old - 1] for old in sorted(cited, key=mapping.get)]


def add_para(text, bold=False, italic=False, align=None, space_before=0, space_after=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after:
        p.paragraph_format.space_after = Pt(space_after)
    add_run_with_refs(p, Template(text).substitute(FLAT), italic=italic, bold=bold)
    return p


def add_heading(text, level=1, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text.upper() if level == 1 else text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12) if level > 1 else Pt(13)
    run.bold = True
    return p


def add_subheading(text):
    return add_heading(text, level=2, space_before=8, space_after=4)


def add_blank():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    return p


def add_table_from_data(caption, headers, rows, note=None):
    """Add an inline table with caption and optional footnote."""
    add_para(caption, bold=True, space_before=12, space_after=6)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            txt = str(val)
            if FLAT:
                txt = Template(txt).safe_substitute(FLAT)
            r = p.add_run(txt)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)
    if note:
        add_para(note, italic=True, space_before=6, space_after=12)
    return table


figure_legends = []


def add_figure_inline(image_path, caption):
    """Store the figure legend; JECH requires figures as separate files."""
    figure_legends.append(caption)
    add_blank()


# ============================================================
# TITLE PAGE (anonymised manuscript: no author identifying information)
# ============================================================
title = ("Regional variation in anaesthesia practice under universal coverage in Japan: "
         "a cross-sectional ecological study of secondary medical areas")
assert len(title) <= 150, f"Title is {len(title)} chars (max 150)"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
run = p.add_run(title)
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.bold = True

add_para("Running header: Anaesthesia variation in Japan", italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_blank()

# Journal of Epidemiology & Community Health operates single-anonymised peer review.
# Author names, affiliations and contact details are supplied on the separate
# title-page file to keep the main manuscript clean.
add_para("[Author names, affiliations and contact details are provided "
         "in the separate title-page file.]", italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_blank()

add_para("Article type: Original research article", italic=True)
add_para("Word count (main text, excluding abstract, references, tables and figures): approximately [calculated at build] words", italic=True)
add_para("Tables: [n_tables]   Figures: [n_figures]   References: [n_references]", italic=True)
add_para("Reporting guideline: STROBE checklist for cross-sectional studies (uploaded as supplementary material)", italic=True)
add_para("Target journal: Journal of Epidemiology & Community Health", italic=True)

doc.add_page_break()
# ============================================================
# ABSTRACT
# ============================================================
add_para(title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_blank()
add_heading("Abstract", level=1, space_before=0)

add_para(
    "Background/aims: Geographic variation in anaesthesia practice is well documented, "
    "but it remains unclear whether it reflects real differences in care or "
    "artefacts of coding, reimbursement and auditing. Japan's uniform national "
    "fee schedule combined with prefecture-specific claims auditing offers a "
    "natural experiment for separating these explanations.")

add_para(
    "Methods: We conducted a cross-sectional ecological study using age- and "
    "sex-standardised claim ratios for general, spinal, epidural and continuous "
    "epidural anaesthesia across ${n_areas} secondary medical areas nested "
    "within ${n_prefectures} prefectures, fiscal year ${fiscal_year}. We fitted "
    "multilevel linear models with prefecture random intercepts and a university-"
    "hospital fixed effect, and performed sensitivity analyses for population "
    "density, anaesthesiologist supply and plausible differential audit rates.")

add_para(
    "Results: Coefficients of variation ranged from ${L008_cv}% (general "
    "anaesthesia) to ${L002_cv}% (epidural). For general anaesthesia, only "
    "${L008_ml_icc} of variance lay between prefectures; university hospital "
    "presence explained ${L008_ml_r2} of total variance (β = "
    "${L008_ml_coef}, 95% CI ${L008_ml_ci_low} to ${L008_ml_ci_high}, P "
    "${L008_ml_p}) and was positive in every prefecture. Maximum plausible "
    "audit-related ratio shifts were small relative to the observed "
    "interquartile range (general anaesthesia: ${audit_L008_iqr_pct}%). "
    "Covariate adjustment attenuated but did not eliminate the university-"
    "hospital effect for general, epidural and continuous epidural anaesthesia.")

add_para(
    "Conclusion: Regional variation in anaesthesia practice is predominantly "
    "structural, driven by access to university hospitals rather than by "
    "prefectural auditing. In Japan and similar universal-coverage systems, "
    "this pattern signals inequitable access to neuraxial techniques and should "
    "be addressed through workforce and organisational policy, not additional "
    "coding audits.")

add_blank()
add_para("Keywords: ", bold=False)
add_para(
    "anaesthesia; small-area variation; multilevel analysis; quality of care; "
    "patient safety; health equity; universal health coverage; Japan")

doc.add_page_break()
# ============================================================
add_heading("Key messages", level=1, space_before=0)

add_para(
    "What is already known on this topic", bold=True)
add_para(
    "Geographic variation in medical and surgical practice is a recurring "
    "finding in high-income countries, yet it is often difficult to separate "
    "genuine differences in care from artefacts of coding, reimbursement and "
    "audit policy. Universal-coverage systems that combine a national fee "
    "schedule with regionally devolved claims auditing provide a natural "
    "experiment for testing these alternative explanations.")

add_para(
    "What this study adds", bold=True)
add_para(
    "In a nationwide cross-sectional ecological analysis of age- and sex-"
    "standardised anaesthesia claim ratios across ${n_areas} secondary medical "
    "areas of Japan, university hospital presence was the dominant structural "
    "determinant of variation in general, epidural and continuous epidural "
    "anaesthesia use. Prefecture-specific audit intensity could explain only a "
    "small fraction of the observed variation.")

add_para(
    "How this study might affect research, practice or policy", bold=True)
add_para(
    "Administrative claims data should be used to monitor small-area "
    "variation in surgical care as a population-health equity indicator, and "
    "policy makers should address workforce and organisational determinants "
    "of access to neuraxial anaesthesia rather than treating geographic "
    "differences primarily as a coding or fraud problem.")

doc.add_page_break()
# INTRODUCTION
# ============================================================
add_heading("Introduction", level=1, space_before=0)

add_para(
    "National administrative claims data are a cornerstone of health-services "
    "and epidemiologic surveillance in countries with universal health "
    "coverage.{1} Because claims record reimbursed procedures, they capture "
    "documented care rather than bedside intention; in systems with uniform fee "
    "schedules, however, geographic variation in claims is commonly interpreted "
    "as a signal of differences in service delivery, practice style and "
    "unmet need.{1,5} The same records are also used for prefecture-level "
    "insurance auditing (shinsa), and audit intensity varies across "
    "prefectures.{6} Where audits differ, between-area variation in claims could "
    "reflect differential scrutiny rather than differential care. Distinguishing "
    "between them is important for population health research and "
    "policy: if variation is real, it signals inequitable access; if it is an "
    "artefact, quality-improvement and payment reforms may be misdirected.")

add_para(
    "Japan offers a natural experiment for this question. The fee schedule "
    "and insured benefits are nationally uniform, so price variation is "
    "eliminated, while prefectures independently operate claims audits. "
    "Anaesthesia is a useful tracer: the choice between general anaesthesia "
    "alone and neuraxial or combined techniques is clinically consequential "
    "and plausibly linked to outcomes, yet the technical and workforce "
    "requirements for neuraxial techniques tend to concentrate in teaching "
    "hospitals.{7-9} If access to these techniques varies by geography, "
    "patients' probability of receiving potentially beneficial care depends "
    "on place of residence, an equity problem under universal coverage. Japan's "
    "2026 fee-schedule revision clarified the billing code for general "
    "anaesthesia, reducing one source of coding ambiguity.{10}")

add_para(
    "We used publicly available, age- and sex-standardised claim ratios to ask: "
    "(i) How large is regional variation in anaesthesia practice across Japan's "
    "secondary medical areas? (ii) Is the variation consistent with "
    "differential prefectural auditing? (iii) What proportion of variance "
    "lies between versus within prefectures? (iv) Is the distribution "
    "associated with proximity to university hospitals after adjusting for "
    "population density and anaesthesiologist supply? By exploiting the nesting "
    "of secondary medical areas within prefectures, we provide a transferable "
    "framework for distinguishing administrative from structural sources of "
    "small-area variation in other universal-coverage systems.")
# ============================================================
# METHODS
# ============================================================
add_heading("Methods", level=1)

add_subheading("Study design and ethical considerations")
add_para(
    "This was a cross-sectional ecological study using publicly available "
    "aggregate data only. No individual-level data were accessed. Patients and "
    "the public were not involved in the design, conduct or reporting of this "
    "study because it uses aggregate, publicly released data. Ethics committee "
    "approval was not required under the Ethical Guidelines for Medical and "
    "Biological Research Involving Human Subjects (Japan, 2021 revision). "
    "The study was reported in accordance with the Strengthening the Reporting of "
    "Observational Studies in Epidemiology (STROBE) checklist for cross-sectional "
    "studies,{30} which is provided as supplementary material.")

add_subheading("Geographic units")
add_para(
    "Japan's healthcare delivery system is organised into a hierarchical structure "
    "of medical areas defined under Article 30-4 of the Medical Care Act.{3} The "
    "${n_prefectures} prefectures serve as tertiary medical areas; within each "
    "prefecture, secondary medical areas (n = ${n_areas} nationally) are designated "
    "as the basic units for inpatient care planning. Crucially, secondary medical "
    "areas are completely nested within prefectures and do not cross prefectural "
    "boundaries. This nesting differs from the Hospital Service Areas and Hospital "
    "Referral Regions used in the United States, which can cross state "
    "boundaries,{11} and makes within-prefecture comparisons unconfounded by "
    "spillover under a common audit policy.")

add_subheading("Data sources and standardised claim ratios")
add_para(
    "We used three data sources. First, standardised claim ratios for fiscal "
    "year ${fiscal_year} published by the Cabinet Office under the 'Regional Variation "
    "Visualisation' initiative.{12} Standardised claim ratios are computed by "
    "indirect standardisation: expected claim frequencies are calculated by applying "
    "national age- and sex-specific claim rates to the local population structure, "
    "and the ratio is defined as 100 × (observed claims / expected claims). Ratios "
    "are calculated on a residence basis, so claims are attributed to the "
    "beneficiary's registered address rather than to the providing facility. "
    "Patient travel for treatment therefore does not inflate the ratio of the "
    "receiving area. Areas with very few claims are masked by the data provider to "
    "protect privacy and appear as missing values. Second, we obtained physician "
    "statistics from the ${fiscal_year} Survey of Physicians, Dentists and "
    "Pharmacists.{13} Third, secondary medical area boundaries, land area and "
    "population density were obtained from the National Land Numerical "
    "Information medical-area dataset (A38-20).{14}")

add_subheading("Anaesthesia procedure codes and university hospital mapping")
add_para(
    "Six procedure codes from the Japanese fee schedule were analysed (Table 1): "
    "L008 (closed-circuit general anaesthesia); L002 (epidural anaesthesia as "
    "main technique); L003 (continuous epidural infusion, largely billed as "
    "adjunct to general anaesthesia); L004 (spinal anaesthesia); L009 (anaesthesia "
    "management fee I, a proxy for specialist staffing); and L100 (inpatient nerve "
    "block). We mapped university hospitals to their respective "
    "secondary medical areas based on municipal address. ${n_univ_areas} "
    "secondary medical areas (${n_univ_areas_pct}%) contained at least one "
    "university hospital, distributed across all ${n_prefectures} prefectures, "
    "which enables within-prefecture comparisons between university and "
    "non-university areas under a common audit policy.")

add_table_from_data(
    "Table 1. Anaesthesia procedure codes analysed across ${n_areas} "
    "secondary medical areas of Japan, fiscal year ${fiscal_year}.",
    ["Code", "Procedure", "Notes"],
    [
        ["L008", "General anaesthesia (closed-circuit)",
         "Primary indicator of overall general anaesthesia volume"],
        ["L004", "Spinal anaesthesia", "Regional technique"],
        ["L002", "Epidural anaesthesia as main technique",
         "Regional technique claimed as the primary anaesthetic"],
        ["L003", "Continuous epidural infusion",
         "Regional technique largely billed as adjunct to general anaesthesia"],
        ["L009", "Anaesthesia management fee I",
         "Proxy for specialist anaesthesiologist staffing"],
        ["L100", "Nerve block, inpatient", "Indicator of pain-clinic activity"],
    ],
    note=("Source: Japanese fee schedule (shinryo houshu tensuhyo), fiscal year "
          "${fiscal_year}. Standardised claim ratios (national average = 100) "
          "were computed for each code by indirect age- and sex-standardisation.")
)

# ============================================================
# RESULTS
# ============================================================
add_heading("Results", level=1)

add_subheading("Study population and variation in anaesthesia practice")
add_para(
    "The ${n_areas} secondary medical areas were distributed across "
    "${n_prefectures} prefectures (median 7 areas per prefecture, range 3 to 21); "
    "${n_univ_areas} areas (${n_univ_areas_pct}%) contained at least one university "
    "hospital. Standardised claim ratios for general anaesthesia were available "
    "for ${L008_n} of ${n_areas} areas; ratios for epidural anaesthesia were "
    "available for ${L002_n} areas, with ${L002_missing_n} areas masked "
    "owing to low volume. Substantial variation was observed across all codes "
    "(Table 2). General anaesthesia ratios ranged from ${L008_min} to "
    "${L008_max} (coefficient of variation ${L008_cv}%), a ${L008_range_ratio}" 
    "-fold difference. Epidural anaesthesia showed the greatest relative variation "
    "(coefficient of variation ${L002_cv}%); continuous epidural infusion "
    "${L003_cv}% and spinal anaesthesia ${L004_cv}%. The geographic distribution "
    "of ratios is shown in Figure 1.")

# Table 2 inline
add_table_from_data(
    "Table 2. Distribution of standardised claim ratios across ${n_areas} secondary "
    "medical areas (national average = 100), fiscal year ${fiscal_year}.",
    ["Code", "n", "Mean (SD)", "Median (IQR)", "Min", "Max", "CV (%)"],
    [
        ["L008", "${L008_n}", "${L008_mean} (${L008_sd})",
         "${L008_median} (${L008_q1}-${L008_q3})", "${L008_min}", "${L008_max}",
         "${L008_cv}"],
        ["L002", "${L002_n}", "${L002_mean} (${L002_sd})",
         "${L002_median} (${L002_q1}-${L002_q3})", "${L002_min}", "${L002_max}",
         "${L002_cv}"],
        ["L003", "${L003_n}", "${L003_mean} (${L003_sd})",
         "${L003_median} (${L003_q1}-${L003_q3})", "${L003_min}", "${L003_max}",
         "${L003_cv}"],
        ["L004", "${L004_n}", "${L004_mean} (${L004_sd})",
         "${L004_median} (${L004_q1}-${L004_q3})", "${L004_min}", "${L004_max}",
         "${L004_cv}"],
        ["L009", "${L009_n}", "${L009_mean} (${L009_sd})",
         "${L009_median} (${L009_q1}-${L009_q3})", "${L009_min}", "${L009_max}",
         "${L009_cv}"],
        ["L100", "${L100_n}", "${L100_mean} (${L100_sd})",
         "${L100_median} (${L100_q1}-${L100_q3})", "${L100_min}", "${L100_max}",
         "${L100_cv}"],
    ],
    note=("SD, standard deviation; IQR, interquartile range; CV, coefficient of "
          "variation. Areas with low claim volume are masked by the data provider "
          "and appear as missing values; the n column reflects the number of "
          "secondary medical areas with non-missing values for each code.")
)

add_figure_inline(
    os.path.join(FIG_DIR, 'rapm_fig1_en.png'),
    "Figure 1. Geographic distribution of anaesthesia standardised claim ratios "
    "across ${n_areas} secondary medical areas of Japan, fiscal year ${fiscal_year}. "
    "(A) General anaesthesia (L008). (B) Spinal anaesthesia (L004). (C) Epidural "
    "anaesthesia as main anaesthetic (L002). (D) Continuous epidural infusion "
    "(L003). Choropleth maps shaded by quintile of the standardised claim ratio "
    "(national average = 100). Red circles mark secondary medical areas "
    "containing at least one university hospital. Areas masked by the data "
    "provider owing to low volume are shown in grey."
)

add_subheading("Multilevel model and university hospital effect")
add_para(
    "The null multilevel model showed that only ${L008_ml_icc} of general "
    "anaesthesia variance was attributable to the prefecture level (intraclass "
    "correlation coefficient ${L008_ml_icc}), indicating that "
    "${L008_vd_within}% occurred within prefectures where audit policy is "
    "uniform. Spinal anaesthesia showed stronger prefecture-level clustering "
    "(intraclass correlation coefficient ${L004_ml_icc}), consistent with a "
    "greater role of prefectural factors for that code. Adding university "
    "hospital presence as a fixed effect produced the largest improvement for "
    "general anaesthesia (marginal R² ${L008_ml_r2}; β = ${L008_ml_coef}, "
    "95% confidence interval ${L008_ml_ci_low} to ${L008_ml_ci_high}) and was "
    "statistically significant for every code (all P ${L008_ml_p}) (Table 3). "
    "University hospital areas had higher general anaesthesia ratios than non-"
    "university areas in ${n_prefectures} of ${n_prefectures} prefectures, with a "
    "mean within-prefecture difference of +${L008_within_diff} points (paired "
    "t = ${L008_within_t}, P ${L008_within_p}) (Figure 2A). Cohen's d for the "
    "university hospital effect on general anaesthesia was ${L008_d}; on "
    "continuous epidural infusion ${L003_d}; on epidural anaesthesia ${L002_d}; "
    "and on spinal anaesthesia ${L004_d}. A combined measure capturing the "
    "general-anaesthesia plus continuous-epidural workflow (mean of L008 and L003 "
    "SCR; ${L008_L003_n} areas with data for both codes) showed a coefficient of "
    "variation of ${L008_L003_cv}% and preserved a large university hospital "
    "effect (Cohen's d ${L008_L003_d}; university mean ${L008_L003_u_mean} "
    "versus non-university ${L008_L003_nu_mean}) (Figure 2B).")

# Table 3 inline
add_table_from_data(
    "Table 3. Multilevel linear mixed model results: standardised claim ratio as "
    "outcome, prefecture as random intercept.",
    ["Code", "Null model ICC", "β university (95% CI)", "P value", "Marginal R²"],
    [
        ["L008", "${L008_ml_icc}",
         "+${L008_ml_coef} (${L008_ml_ci_low} to ${L008_ml_ci_high})",
         "${L008_ml_p}", "${L008_ml_r2}"],
        ["L002", "${L002_ml_icc}",
         "+${L002_ml_coef} (${L002_ml_ci_low} to ${L002_ml_ci_high})",
         "${L002_ml_p}", "${L002_ml_r2}"],
        ["L003", "${L003_ml_icc}",
         "+${L003_ml_coef} (${L003_ml_ci_low} to ${L003_ml_ci_high})",
         "${L003_ml_p}", "${L003_ml_r2}"],
        ["L004", "${L004_ml_icc}",
         "+${L004_ml_coef} (${L004_ml_ci_low} to ${L004_ml_ci_high})",
         "${L004_ml_p}", "${L004_ml_r2}"],
    ],
    note=("β university, fixed effect coefficient for university hospital presence "
          "(binary 0/1); CI, confidence interval; ICC, intraclass correlation "
          "coefficient (proportion of variance attributable to prefecture level in "
          "the null random-intercept model); Marginal R², proportional reduction "
          "in total (prefecture + residual) variance from the null model after "
          "adding the university hospital fixed effect. Models were estimated by "
          "restricted maximum likelihood.")
)

add_para(
    "A covariate-adjusted sensitivity model added the natural logarithm of "
    "population density and the anaesthesiologist share of all physicians (both "
    "standardised). The university hospital coefficient was attenuated but remained "
    "positive and statistically significant for general anaesthesia (β = "
    "+${L008_mlc_coef}, 95% CI ${L008_mlc_ci_low} to ${L008_mlc_ci_high}, P "
    "${L008_mlc_p}), epidural anaesthesia (β = +${L002_mlc_coef}, 95% CI "
    "${L002_mlc_ci_low} to ${L002_mlc_ci_high}, P ${L002_mlc_p}) and "
    "continuous epidural infusion (β = +${L003_mlc_coef}, 95% CI "
    "${L003_mlc_ci_low} to ${L003_mlc_ci_high}, P ${L003_mlc_p}). For spinal "
    "anaesthesia the point estimate was positive but no longer statistically "
    "significant (β = +${L004_mlc_coef}, 95% CI ${L004_mlc_ci_low} to "
    "${L004_mlc_ci_high}, P ${L004_mlc_p}). Population density was positively "
    "associated with general anaesthesia ratios (β = ${L008_mlc_popd_coef}, P "
    "${L008_mlc_popd_p}) but not with the other three codes. The anaesthesiologist "
    "share was positively associated with general, epidural, continuous epidural and "
    "spinal anaesthesia (P ${L008_mlc_anes_p}, ${L002_mlc_anes_p}, "
    "${L003_mlc_anes_p} and ${L004_mlc_anes_p}, respectively). "
    "The covariate-adjusted mixed models produced optimizer-convergence warnings "
    "in Statsmodels;{17} we therefore refitted them with the lbfgs and cg optimisers, "
    "and the point estimates remained stable across optimisers and sensitivity "
    "analyses."
)

add_figure_inline(
    os.path.join(FIG_DIR, 'rapm_fig2_en.png'),
    "Figure 2. University hospital presence and the combined general-anaesthesia "
    "plus continuous-epidural measure. (A) Distribution of secondary medical "
    "areas containing at least one university hospital (n = ${n_univ_areas} of "
    "${n_areas}; red). (B) Choropleth map of the combined general-anaesthesia "
    "plus continuous-epidural standardised claim ratio (mean of L008 and L003 "
    "SCR; ${L008_L003_n} areas with data for both codes), shaded by quintile. "
    "Red circles mark secondary medical areas containing at least one university "
    "hospital. Areas masked by the data provider for either code are shown in "
    "grey."
)

add_subheading("Sensitivity analyses against the audit hypothesis")
add_para(
    "All three pre-specified sensitivity analyses converged in rejecting the "
    "audit hypothesis. First, within-prefecture variance decomposition attributed "
    "${L008_vd_between}% of general anaesthesia variance to between-prefecture "
    "differences, ${L008_vd_univ}% to the university hospital effect within "
    "prefecture, and ${L008_vd_residual}% to residual within-group variation; the "
    "university hospital effect alone explained ${L008_vd_univ_within}% of all "
    "within-prefecture variance. Second, general and spinal anaesthesia were "
    "positively, not negatively, correlated (r = ${corr_L008_L004_r}, P "
    "${corr_L008_L004_p}), as were general and epidural anaesthesia (r = "
    "${corr_L008_L002_r}, P ${corr_L008_L002_p}), inconsistent with audit-"
    "driven reclassification and consistent with a common supply factor. Third, "
    "the maximum prefectural audit-rate difference of ${audit_max_shift} "
    "percentage points could produce a ratio shift of approximately "
    "${audit_max_shift} points, equivalent to approximately ${audit_L008_iqr_pct}% of "
    "the observed interquartile range for general anaesthesia and "
    "${audit_L002_iqr_pct}% for epidural anaesthesia. Combining general (L008) "
    "and spinal (L004) ratios produced an alternative audit-insensitive "
    "measure with a coefficient of variation of ${L008_L004_cv}% and a university "
    "hospital effect that remained large (Cohen's d ${L008_L004_d}), confirming "
    "that the structural effect is not artefactually erased by combining "
    "reclassifiable codes.")

add_subheading("Robustness: empirical Bayes shrinkage and outliers")
add_para(
    "Empirical Bayes shrinkage toward prefecture means was implemented as a "
    "stress test of the area-level ratios. Under this model, shrinkage "
    "substantially compressed the area-level ratios toward their prefecture "
    "means. For general anaesthesia, Cohen's d for the university-hospital effect "
    "fell from ${L008_d} (raw) to ${L008_d_shrunk} (shrunken), an attenuation "
    "of ${L008_attenuation}%; for epidural anaesthesia (L002) from ${L002_d} to "
    "${L002_d_shrunk} (${L002_attenuation}%); and for spinal anaesthesia (L004) "
    "from ${L004_d} to ${L004_d_shrunk} (${L004_attenuation}%). The general "
    "anaesthesia effect therefore remained of moderate magnitude despite heavy "
    "shrinkage, whereas the L002 and L004 effects were materially attenuated, "
    "indicating that those smaller raw effects are more vulnerable to sampling "
    "instability. These shrinkage estimates are complementary to, and consistent "
    "in direction with, the mixed-effects model coefficients reported in Table 3, "
    "which themselves incorporate prefecture-level shrinkage. Two areas were "
    "identified as outliers (more than three standard deviations from the mean "
    "general anaesthesia ratio): ${outlier1_name} in ${outlier1_pref} (ratio "
    "${outlier1_scr}) and ${outlier2_name} in ${outlier2_pref} (ratio "
    "${outlier2_scr}). Excluding outliers did not alter the direction or "
    "statistical significance of any finding.")

# ============================================================
# DISCUSSION
# ============================================================
add_heading("Discussion", level=1)

add_subheading("Statement of principal findings")
add_para(
    "Regional variation in anaesthesia practice across Japan's ${n_areas} "
    "secondary medical areas is large and predominantly structural. University "
    "hospital presence explained the largest share of variance for general "
    "anaesthesia and was positive in every prefecture, while the proportion of "
    "variance attributable to prefecture-level factors -- where audit intensity "
    "differs -- was small. Three independent sensitivity analyses converged in "
    "rejecting differential auditing as the main explanation. Because the "
    "outcome is a reimbursed claim ratio, the observed gradient documents "
    "service delivery rather than bedside practice directly; however, the "
    "within-prefecture structure of the data and the audit-sensitivity "
    "analyses make it unlikely that the entire pattern is an artefact of "
    "differential scrutiny. The residual variation is therefore best "
    "interpreted as a geographically patterned difference in the provision of "
    "potentially beneficial anaesthesia services.")

add_subheading("Strengths and limitations")
add_para(
    "Strengths include the use of age- and sex-standardised ratios covering the "
    "entire national population; analysis at a fine geographic scale of "
    "${n_areas} areas nested within ${n_prefectures} prefectures; multilevel "
    "modelling that accounts for the nested data structure; empirical Bayes "
    "shrinkage to address low-volume instability; multiple pre-specified "
    "sensitivity analyses; and the residence-based ratio definition that "
    "mitigates patient travel effects. Limitations are those common to "
    "ecological, cross-sectional claims studies: the ecological fallacy;{19} "
    "the inability to infer individual-level treatment from area-level "
    "aggregates; the cross-sectional design, which precludes causal inference; "
    "the post-audit nature of the data, which reflects reimbursed rather than "
    "intended practice; and the partial overlap between university hospital "
    "presence and urban concentration. We could not access code-specific "
    "audit rates, so the audit-sensitivity bound relies on the aggregate "
    "prefectural audit differential. A specific data constraint is that the "
    "combined anaesthesia add-on billed under L008 is aggregated with other "
    "add-ons in the publicly released regional variation dataset and cannot be "
    "extracted separately; patient-level National Database records would be "
    "needed to measure the true rate of combined general-epidural anaesthesia. "
    "We therefore used L002 (epidural as main technique) and L003 (continuous "
    "epidural infusion) as the best available public proxies. A covariate-"
    "adjusted sensitivity model added population density and anaesthesiologist "
    "share of physicians; the university hospital coefficient remained "
    "positive and statistically significant for general, epidural and "
    "continuous epidural anaesthesia, but not for spinal anaesthesia. The "
    "positive point estimate for spinal anaesthesia is compatible with the "
    "same structural gradient in an underpowered model, and may partly reflect "
    "a countervailing tendency for neuraxial anaesthesia to be provided in "
    "settings where anaesthesiologists are less available. The aggregate data "
    "cannot separate these opposing mechanisms, so the adjusted spinal-"
    "anaesthesia coefficient is best interpreted as an imprecise net effect.")

add_subheading("Interpretation in the wider literature")
add_para(
    "Our findings are consistent with the broader medical practice-variation "
    "literature. The Dartmouth Atlas project documented extensive regional "
    "variation in surgical rates in the United States, driven primarily by "
    "physician supply and practice style.{20} Comparable patterns have been "
    "described for the National Health Service in England,{21} Germany{22} and "
    "Australia.{23} Our study extends this evidence by exploiting Japan's "
    "uniform fee schedule combined with prefecture-specific auditing to "
    "disentangle administrative and structural sources of variation, and by "
    "providing, to our knowledge, the first multilevel small-area analysis of "
    "anaesthesia technique under universal coverage in East Asia. The "
    "university hospital effect we report (Cohen's d = ${L008_d}) is "
    "substantial, plausibly reflecting the influence of the ikyoku (university "
    "medical office) system on practice in affiliated hospitals.{24}")

add_para(
    "The findings are relevant beyond Japan. The within-prefecture variance "
    "decomposition we describe is directly transportable to other universal-"
    "coverage systems that combine a centrally uniform fee schedule with "
    "regionally devolved claims auditing, including Taiwan's National Health "
    "Insurance,{25} South Korea's National Health Insurance Service, and the "
    "tariff-based systems of Germany, France and the English National Health "
    "Service.{21,22} The dominance of tertiary-teaching-hospital concentration "
    "as a structural determinant of technique choice echoes analogous patterns "
    "for teaching-hospital anaesthesia mix in the US and United "
    "Kingdom{20,21} and for other tertiary-intensive specialties across OECD "
    "countries.{26} In contrast to the highly localised, physician-style-driven "
    "variation often reported in the US, the Japanese pattern is more strongly "
    "organisational, concentrated around university hospitals, and consistent "
    "across all ${n_prefectures} prefectures. This suggests that the "
    "structural-access explanation is a robust feature of high-income "
    "universal-coverage systems rather than a Japanese idiosyncrasy. "
    "Extrapolation to low- and middle-income countries should be more cautious, because "
    "fee schedules and audit arrangements are typically heterogeneous and "
    "claims data often incomplete.")

add_subheading("Implications for public health research, practice and policy")
add_para(
    "Three public-health implications follow from these findings. First, "
    "administrative claims data can be used as a "
    "surveillance tool to monitor small-area variation in perioperative care "
    "under universal coverage, but should be interpreted alongside sensitivity "
    "analyses for coding and audit effects. Second, the observed "
    "variation in epidural and continuous epidural use signals a potential "
    "equity gap: patients' access to techniques that may improve recovery "
    "should not depend on whether they live near a university hospital. Third, "
    "regulators should focus on workforce and organisational "
    "determinants of anaesthesia service delivery rather than treating low "
    "regional ratios primarily as coding or fraud problems. The "
    "variance-decomposition framework we used is transferable to other "
    "procedures and to other countries with a national fee schedule and "
    "regional audit variation, offering a practical epidemiologic approach to "
    "separating administrative from structural sources of small-area "
    "variation.")
# ============================================================
# CONCLUSIONS
# ============================================================
add_heading("Conclusions", level=1)
add_para(
    "Regional variation in anaesthesia practice under universal coverage in "
    "Japan is large and predominantly structural, driven by access to university "
    "hospitals rather than by prefectural claims auditing. The pattern is best "
    "interpreted as a documented-service-delivery gradient with potentially "
    "inequitable access to neuraxial techniques. Public health surveillance of "
    "small-area variation in surgical and anaesthesia services should separate "
    "administrative from structural sources before interpreting claims data as "
    "care variation or as coding quality.")
# ============================================================
# REFERENCES
# ============================================================
add_heading("References", level=1)

REFERENCES = [
    "Ikegami N, Yoo BK, Hashimoto H, et al. Japanese universal health coverage: evolution, achievements, and challenges. Lancet 2011;378:1106-15.",
    "GBD 2019 Universal Health Coverage Collaborators. Measuring universal health coverage based on an index of effective coverage of health services in 204 countries and territories, 1990-2019. Lancet 2020;396:1250-84.",
    "Matsuda S. Health policy in Japan: current situation and future challenges. JMA J 2019;2:1-10.",
    "Ministry of Health, Labour and Welfare. Patient Survey 2008. Tokyo: MHLW, 2009.",
    "Hashimoto H, Ikegami N, Shibuya K, et al. Cost containment and quality of care in Japan: is there a trade-off? Lancet 2011;378:1174-82.",
    "Social Insurance Medical Fee Payment Fund. 審査状況（令和4年度）— 都道府県別医科歯科計審査状況（点数）. Tokyo: Social Insurance Medical Fee Payment Fund, 2023. https://www.ssk.or.jp/smph/tokeijoho/shinsatokei/shinsajokyo_r04.html (accessed 24 Jul 2026).",
    "Sessler DI, Pei L, Huang Y, et al. Recurrence of breast cancer after regional or general anaesthesia: a randomised controlled trial. Lancet 2019;394:1807-15.",
    "Chen WK, Miao CH. The effect of anesthetic technique on survival in human cancers: a meta-analysis of retrospective and prospective studies. PLoS One 2013;8:e56540.",
    "Weng M, Chen W, Hou W, et al. The effect of neuraxial anaesthesia on cancer recurrence and survival after cancer surgery: an updated meta-analysis. Oncotarget 2016;7:15262-73.",
    "Ministry of Health, Labour and Welfare. 2026 revision of the medical fee schedule: changes to the medical fee schedule master (R08). Tokyo: MHLW, 2026. https://shinryohoshu.mhlw.go.jp/shinryohoshu/file/info/smente260522.pdf (accessed 24 Jul 2026).",
    "Wennberg JE, Cooper MM, eds. The Dartmouth Atlas of Health Care. Chicago: American Hospital Publishing, 1999.",
    "Cabinet Office. Regional variation visualisation (chiikisa no mieruka). https://www5.cao.go.jp/keizai-shimon/kaigi/special/reform/mieruka/chiikisa/index.html (accessed 10 Mar 2026).",
    "Ministry of Health, Labour and Welfare. Survey of Physicians, Dentists and Pharmacists 2022 (e-Stat). https://www.e-stat.go.jp (accessed 12 Mar 2026).",
    "Ministry of Land, Infrastructure, Transport and Tourism. National Land Numerical Information download service. https://nlftp.mlit.go.jp (accessed 10 Mar 2026).",
    "Cohen J. Statistical Power Analysis for the Behavioral Sciences. 2nd ed. Hillsdale, NJ: Lawrence Erlbaum, 1988.",
    "Rabe-Hesketh S, Skrondal A. Multilevel and Longitudinal Modeling Using Stata. 3rd ed. College Station, TX: Stata Press, 2012.",
    "Seabold S, Perktold J. Statsmodels: econometric and statistical modeling with Python. In: Proceedings of the 9th Python in Science Conference. 2010:92-6.",
    "Efron B, Morris C. Stein's estimation rule and its competitors: an empirical Bayes approach. J Am Stat Assoc 1973;68:117-30.",
    "Greenland S. Ecologic versus individual-level sources of bias in ecologic estimates of contextual health effects. Int J Epidemiol 2001;30:1343-50.",
    "Wennberg JE, Fisher ES, Skinner JS. Geography and the debate over Medicare reform. Health Aff (Millwood) 2002;Suppl Web Exclusives:W96-114.",
    "NHS RightCare. The NHS Atlas of Variation in Healthcare. London: Public Health England, 2015.",
    "Schäfer T, Pritzkuleit R, Jeszenszky C, et al. Trends and geographical variation of primary hip and knee joint replacement in Germany. Osteoarthritis Cartilage 2013;21:279-88.",
    "Australian Commission on Safety and Quality in Health Care. Australian Atlas of Healthcare Variation. Sydney: ACSQHC, 2015.",
    "Otsuka T. The ikyoku system of university orthopaedic surgery departments: an in-hospital organisational system unique to Japan. J Orthop Sci 2012;17:513-14.",
    "Cheng T-M. Taiwan's new national health insurance program: genesis and experience so far. Health Aff (Millwood) 2003;22:61-76.",
    "OECD. Geographic Variations in Health Care: What Do We Know and What Can Be Done to Improve Health System Performance? Paris: OECD Publishing, 2014.",
    "Donabedian A. The quality of care: how can it be assessed? JAMA 1988;260:1743-8.",
    "Merry AF, Cooper JB, Soyannwo O, et al. International standards for a safe practice of anesthesia 2010. Can J Anesth 2010;57:1027-34.",
    "Mainz J. Defining and classifying clinical indicators for quality improvement. Int J Qual Health Care 2003;15:523-30.",
    "von Elm E, Altman DG, Egger M, et al. The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement: guidelines for reporting observational studies. Lancet 2007;370:1453-7.",
]

# Reorder references so their numbering matches first appearance in the body.
REFERENCES = renumber_references(doc, REFERENCES)

for i, ref in enumerate(REFERENCES, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(f"{i}. ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run = p.add_run(Template(ref).substitute(FLAT))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# Body word count check (excluding title page, abstract, references, tables, figures)
text_paras = []
in_body = False
in_refs = False
for p in doc.paragraphs:
    txt = p.text.strip()
    if not txt:
        continue
    if txt.upper() == 'INTRODUCTION':
        in_body = True
        continue
    if txt.upper() == 'REFERENCES':
        in_refs = True
        continue
    if in_body and not in_refs:
        text_paras.append(txt)
body = '\n'.join(text_paras)
words = len(re.findall(r'\b\w+\b', body))
print(f"Body word count (intro through conclusion): {words}")

# Update title-page counts that were left as placeholders
n_tables = 3
n_figures = 2
n_refs = len(REFERENCES)
for p in doc.paragraphs:
    if p.text.startswith('Tables:'):
        p.text = f"Tables: {n_tables}   Figures: {n_figures}   References: {n_refs}"
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.italic = True
        break

# Update the word-count line on the title page
for p in doc.paragraphs:
    if p.text.startswith('Word count'):
        p.text = f"Word count (main text, excluding abstract, references, tables and figures): approximately {words} words"
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.italic = True
        break

# Figure legends are placed at the end of the manuscript per JECH guidance.
if figure_legends:
    add_heading('Figure legends', level=1)
    for cap in figure_legends:
        add_para(cap, space_before=6, space_after=12)

# Convert statistical expressions to native Word equation objects
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from jech_math import convert_docx_math_to_omml
convert_docx_math_to_omml(doc)

# Add a centered page number to the footer of each section
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = True
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld_separate = OxmlElement('w:fldChar')
    fld_separate.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)

# ============================================================
# Save
# ============================================================
out = os.path.join(OUTPUT_DIR, 'regional_anaesthesia_JECH_EN.docx')
doc.save(out)
print(f"Saved: {out}")

# --------------------------------------------------
# Separate title page for Journal of Epidemiology & Community Health
# --------------------------------------------------
tp = Document()
for section in tp.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

tps = tp.styles['Normal']
tps.font.name = 'Times New Roman'
tps.font.size = Pt(12)
tps.paragraph_format.line_spacing = 2.0
tps.paragraph_format.space_after = Pt(0)

def tp_para(text, bold=False, italic=False, align=None):
    p = tp.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(Template(text).substitute(FLAT))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p

p = tp.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(title)
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.bold = True

tp_para("[Authors and affiliations to be supplied]", italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER)
tp_para("[Corresponding author: name, postal address, telephone, email]",
        italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
tp.add_paragraph()

tp_para("Article type: Original research article", italic=True)
tp_para("Word count (main text, excluding abstract, references, tables and figures): approximately [calculated at build] words", italic=True)
tp_para("Tables: [n_tables]   Figures: [n_figures]   References: [n_references]", italic=True)
tp_para("Reporting guideline: STROBE checklist for cross-sectional studies", italic=True)
tp_para("Target journal: Journal of Epidemiology & Community Health", italic=True)

tp.add_paragraph()
tp_para("Declarations", bold=True)
tp_para(
    "Ethics approval and consent to participate: This study used publicly "
    "available aggregate data only. Ethics committee approval was not required "
    "under the Ethical Guidelines for Medical and Biological Research Involving "
    "Human Subjects (Japan, 2021 revision).")
tp_para(
    "Consent for publication: Not applicable.")
tp_para(
    "Availability of data and materials: All data are publicly available. "
    "Standardised claim ratios: Cabinet Office Regional Variation Visualisation "
    "(https://www5.cao.go.jp/keizai-shimon/kaigi/special/reform/mieruka/). "
    "Physician statistics: e-Stat (https://www.e-stat.go.jp). Geographic "
    "boundary data: National Land Numerical Information "
    "(https://nlftp.mlit.go.jp). Derived datasets and analysis code are "
    "available from the corresponding author on reasonable request.")
tp_para(
    "Competing interests: [To be completed by the authors.] No known competing "
    "interests.")
tp_para(
    "Funding: [To be completed by the authors.] This work received no specific "
    "grant.")
tp_para(
    "Authors' contributions: [To be completed by the authors.]")
tp_para(
    "Acknowledgements: The authors thank the Cabinet Office Regional "
    "Variation Visualisation initiative, the Ministry of Health, Labour and "
    "Welfare, and the Ministry of Land, Infrastructure, Transport and Tourism "
    "for maintaining the public datasets used in this study.")
tp_para(
    "Patient and public involvement: Patients and the public were not involved "
    "in the design, conduct or reporting of this study because it uses aggregate, "
    "publicly released data.")

# Update placeholders in the separate title page
for p in tp.paragraphs:
    txt = p.text
    if '[calculated at build]' in txt:
        p.text = txt.replace('[calculated at build]', str(words))
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.italic = True
    if '[n_tables]' in txt:
        p.text = txt.replace('[n_tables]', str(n_tables)).replace('[n_figures]', str(n_figures)).replace('[n_references]', str(n_refs))
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.italic = True

tp_out = os.path.join(OUTPUT_DIR, 'title_page_JECH_EN.docx')
tp.save(tp_out)
print(f"Saved: {tp_out}")
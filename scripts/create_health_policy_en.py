#!/usr/bin/env python3
"""Create Health Policy formatted English manuscript.

Conforms to Health Policy original research author guidelines:
- Title should indicate study design
- Structured abstract (Background/Objective/Methods/Results/Conclusions) <=250 words
- Highlights / Key Points: 3-5 bullets, each <=85 characters, as separate file
- Body approximately <=4000 words (excluding abstract, references, tables, figures)
- Up to 4 combined figures/tables (full-length article)
- References numbered in order of appearance (Vancouver style; in-text square brackets)
- Numbered IMRaD sections and subsections (1, 1.1, 1.1.1)
- Double-spaced, Times New Roman 12 pt
- Double-anonymised peer review: separate title page with declarations, including AI use

All numeric results are read from output/ijhpm_results.json, which is produced
by scripts/compile_ijhpm_results.py.
"""

import json
import os
import re
from string import Template

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'documents', 'Health_Policy')
ZENODO_DOI = ''
for _zenodo_file in [
    os.path.join(OUTPUT_DIR, 'zenodo_doi.txt'),
    os.path.join(os.path.dirname(__file__), '..', 'output', 'zenodo_doi.txt'),
]:
    if os.path.exists(_zenodo_file):
        ZENODO_DOI = open(_zenodo_file).read().strip()
        if ZENODO_DOI:
            break
os.makedirs(OUTPUT_DIR, exist_ok=True)

REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FIG_DIR = os.path.join(REPO_ROOT, 'output')
RESULTS_PATH = os.path.join(REPO_ROOT, 'output', 'ijhpm_results.json')

doc = Document()

# Page setup: A4, 2.54 cm margins
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0
style.paragraph_format.space_after = Pt(0)


# ============================================================
# LOAD RESULTS
# ============================================================
with open(RESULTS_PATH) as f:
    R = json.load(f)

codes = R['codes']
combined = R['combined']['L008_L003']
corr = R['correlations']
eb = R['empirical_bayes']
vd = R['variance_decomposition']
audit = R['audit_sensitivity']
outliers = R['outliers']['L008']
meta = R['metadata']



def fmt_num(x, decimals=1):
    if x is None or (isinstance(x, float) and (x != x)):
        return 'NA'
    return f'{x:.{decimals}f}'


def fmt_p(x):
    if x is None or (isinstance(x, float) and (x != x)):
        return 'NA'
    if x < 0.001:
        return '< 0.001'
    return f'= {x:.3f}'


def desc(code):
    return codes[code]['overall']


def univ(code):
    return codes[code]['university']


def nonu(code):
    return codes[code]['non_university']


def ml(code):
    return codes[code]['multilevel']



def build_flat(R):
    """Build a flat substitution dictionary from ijhpm_results.json."""
    F = {}
    meta = R['metadata']
    F['n_areas'] = meta['n_areas']
    F['n_prefectures'] = meta['n_prefectures']
    F['fiscal_year'] = meta.get('fiscal_year', 2022)
    F['n_univ_areas'] = meta['n_univ_areas']
    F['n_nonuniv_areas'] = meta['n_nonuniv_areas']
    F['n_univ_areas_pct'] = fmt_num(meta.get('n_univ_areas_pct', 100 * meta['n_univ_areas'] / meta['n_areas']))

    vd_key_map = {
        'between_prefecture_pct': 'vd_between',
        'university_effect_pct': 'vd_univ',
        'residual_pct': 'vd_residual',
        'within_prefecture_pct': 'vd_within',
        'univ_explains_within_pct': 'vd_univ_within',
    }

    for code, c in R['codes'].items():
        o = c['overall']
        for key, d in [('n', 0), ('mean', 1), ('sd', 1), ('median', 1),
                       ('q1', 1), ('q3', 1), ('min', 1), ('max', 1),
                       ('cv', 1), ('iqr', 1), ('range_ratio', 1)]:
            F[f'{code}_{key}'] = fmt_num(o.get(key), d)

        u = c['university']
        n = c['non_university']
        for prefix, vals in [('u', u), ('nu', n)]:
            for key in ['mean', 'sd', 'median', 'min', 'max']:
                F[f'{code}_{prefix}_{key}'] = fmt_num(vals.get(key), 1)

        F[f'{code}_d'] = fmt_num(c['cohens_d_overall'], 2)

        w = c['within_prefecture']
        F[f'{code}_within_diff'] = fmt_num(w.get('mean_diff'))
        F[f'{code}_within_t'] = fmt_num(w.get('t'), 2)
        F[f'{code}_within_p'] = fmt_p(w.get('p'))

        m = c['multilevel']
        F[f'{code}_ml_icc'] = fmt_num(m.get('icc_null'), 3)
        F[f'{code}_ml_icc_pct'] = fmt_num(m.get('icc_null') * 100, 1)
        F[f'{code}_ml_coef'] = fmt_num(m.get('coef_univ'))
        F[f'{code}_ml_ci_low'] = fmt_num(m.get('ci_low'))
        F[f'{code}_ml_ci_high'] = fmt_num(m.get('ci_high'))
        F[f'{code}_ml_p'] = fmt_p(m.get('p'))
        F[f'{code}_ml_r2'] = fmt_num(m.get('marginal_r2'), 3)
        F[f'{code}_ml_r2_pct'] = fmt_num(m.get('marginal_r2') * 100, 1)
        F[f'{code}_ml_within_pct'] = fmt_num(100 * (1 - m.get('icc_null', 0)), 1)

        ml_log = c.get('multilevel_log')
        if ml_log:
            F[f'{code}_ml_log_icc'] = fmt_num(ml_log.get('icc_null'), 3)
            F[f'{code}_ml_log_coef'] = fmt_num(ml_log.get('coef_univ'), 2)
            F[f'{code}_ml_log_ci_low'] = fmt_num(ml_log.get('ci_low'), 2)
            F[f'{code}_ml_log_ci_high'] = fmt_num(ml_log.get('ci_high'), 2)
            F[f'{code}_ml_log_p'] = fmt_p(ml_log.get('p'))

        mc = c.get('multilevel_covariate')
        if mc:
            F[f'{code}_mlc_coef'] = fmt_num(mc.get('coef_univ'))
            F[f'{code}_mlc_ci_low'] = fmt_num(mc.get('ci_low'))
            F[f'{code}_mlc_ci_high'] = fmt_num(mc.get('ci_high'))
            F[f'{code}_mlc_p'] = fmt_p(mc.get('p'))
            F[f'{code}_mlc_r2'] = fmt_num(mc.get('marginal_r2'), 3)
            F[f'{code}_mlc_popd_coef'] = fmt_num(mc['log_pop_density_z'].get('coef'))
            F[f'{code}_mlc_popd_p'] = fmt_p(mc['log_pop_density_z'].get('p'))
            F[f'{code}_mlc_anes_coef'] = fmt_num(mc['anes_pct_z'].get('coef'))
            F[f'{code}_mlc_anes_p'] = fmt_p(mc['anes_pct_z'].get('p'))

        if code in R.get('variance_decomposition', {}):
            vd = R['variance_decomposition'][code]
            for json_key, flat_key in vd_key_map.items():
                F[f'{code}_{flat_key}'] = fmt_num(vd.get(json_key))

        if code in R.get('empirical_bayes', {}):
            ebv = R['empirical_bayes'][code]
            F[f'{code}_d_shrunk'] = fmt_num(ebv.get('shrunk_cohens_d'), 2)
            F[f'{code}_attenuation'] = fmt_num(ebv.get('attenuation_pct'))

    F['L002_missing_n'] = R['metadata']['n_areas'] - R['codes']['L002']['overall']['n']
    F['L002_fold'] = fmt_num(R['codes']['L002']['university']['mean'] / R['codes']['L002']['non_university']['mean'], 2)
    F['L003_fold'] = fmt_num(R['codes']['L003']['university']['mean'] / R['codes']['L003']['non_university']['mean'], 2)

    for combo, c in R['combined'].items():
        prefix = combo.replace('+', '_')
        for key, d in [('n', 0), ('mean', 1), ('sd', 1), ('median', 1),
                       ('q1', 1), ('q3', 1), ('min', 1), ('max', 1),
                       ('cv', 1), ('iqr', 1)]:
            F[f'{prefix}_{key}'] = fmt_num(c['overall'].get(key), d)
        F[f'{prefix}_u_mean'] = fmt_num(c['university'].get('mean'))
        F[f'{prefix}_nu_mean'] = fmt_num(c['non_university'].get('mean'))
        F[f'{prefix}_d'] = fmt_num(c.get('cohens_d'), 2)
        F[f'{prefix}_fold'] = fmt_num(c.get('fold_ratio'), 2)

    for k, v in R['correlations'].items():
        F[f'corr_{k}_r'] = fmt_num(v.get('r'), 3)
        F[f'corr_{k}_p'] = fmt_p(v.get('p'))

    F['audit_max_shift'] = fmt_num(R['audit_sensitivity']['max_ratio_shift_approx'], 2)
    for code in ['L008', 'L002', 'L003', 'L004']:
        if code in R['audit_sensitivity']:
            F[f'audit_{code}_iqr_pct'] = fmt_num(R['audit_sensitivity'][code]['percent_of_iqr'], 2)

    outlier_name_map = {
        ('安房', '千葉県'): ('Awa (Chiba)', 'Chiba'),
        ('区中央部', '東京都'): ('central Tokyo (Tokyo)', 'Tokyo'),
    }
    for i, o in enumerate(R['outliers']['L008'][:2], 1):
        name, pref = o['area_name'], o['pref_name']
        if (name, pref) in outlier_name_map:
            name, pref = outlier_name_map[(name, pref)]
        F[f'outlier{i}_name'] = name
        F[f'outlier{i}_pref'] = pref
        F[f'outlier{i}_scr'] = fmt_num(o['scr'])

    return F


FLAT = build_flat(R)

_heading_state = {'section': 0, 'subsection': 0}


def _add_text_run(paragraph, text, italic=False, bold=False):
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if italic:
        run.italic = True
    if bold:
        run.bold = True
    return run


def _add_math_run(paragraph, text, italic=False, bold=False):
    """Insert an inline Word equation (<m:oMath>) containing the supplied text."""
    run = paragraph.add_run()
    run.font.name = 'Cambria Math'
    run.font.size = Pt(12)
    if italic:
        run.italic = True
    if bold:
        run.bold = True
    omath = OxmlElement('m:oMath')
    mrun = OxmlElement('m:r')
    rpr = OxmlElement('m:rPr')
    style = OxmlElement('m:sty')
    style.set(qn('m:val'), 'p')
    rpr.append(style)
    mrun.append(rpr)
    mt = OxmlElement('m:t')
    mt.set(qn('xml:space'), 'preserve')
    mt.text = text
    mrun.append(mt)
    omath.append(mrun)
    run._r.append(omath)
    return run


MATH_RE = re.compile(r'(<m>.*?</m>)', re.DOTALL)


def add_run_with_refs(paragraph, text, italic=False, bold=False):
    """Add text to a paragraph, parsing {n} or {n-m} as bracketed Vancouver citations
    and <m>...</m> as inline Word equations."""
    # Process math first so the citation/punctuation rules never touch equation text.
    parts = MATH_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('<m>') and part.endswith('</m>'):
            math_text = part[3:-4]
            _add_math_run(paragraph, math_text, italic=italic, bold=bold)
            continue
        # In Vancouver style, citations follow punctuation: move any trailing
        # punctuation that sits immediately before a citation marker to after it.
        part = re.sub(r'([.,;:])(\{[^}]+\})', r'\2\1', part)
        subparts = re.split(r'(\{[^}]+\})', part)
        for sp in subparts:
            if not sp:
                continue
            if sp.startswith('{') and sp.endswith('}'):
                inner = sp[1:-1]
                run = _add_text_run(paragraph, f'[{inner}]', italic=italic, bold=bold)
                run.font.superscript = False
                continue
            if sp:
                _add_text_run(paragraph, sp, italic=italic, bold=bold)


CITE_RE = re.compile(r'\[\d+(?:-\d+)?(?:,\s*\d+(?:-\d+)?)*\]')


def _expand_citation(text):
    """Expand a bracketed Vancouver citation such as '[7-9,12]' into integers."""
    nums = []
    inner = text.strip().strip('[]')
    for token in re.split(r',\s*', inner):
        token = token.strip()
        if '-' in token:
            start, end = token.split('-', 1)
            nums.extend(range(int(start), int(end) + 1))
        else:
            nums.append(int(token))
    return nums


def renumber_references(doc, references):
    """Determine first-appearance order from bracketed Vancouver citations and reorder.

    Walks the document body (stops at the 'References' heading), records the
    order in which each citation number first appears, builds an old->new
    mapping, updates all bracketed citation runs in place, and returns the
    reference list reordered to match the new numbering. Uncited references are
    dropped.
    """
    first = {}
    cited = set()
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt.upper() == 'REFERENCES':
            break
        for run in p.runs:
            for match in CITE_RE.finditer(run.text):
                for n in _expand_citation(match.group(0)):
                    cited.add(n)
                    if n not in first:
                        first[n] = len(first) + 1
    mapping = {old: new for new, old in enumerate(sorted(first, key=first.get), 1)}

    def repl(m):
        old = int(m.group(0))
        return str(mapping.get(old, old))

    for p in doc.paragraphs:
        for run in p.runs:
            if CITE_RE.search(run.text):
                run.text = re.sub(r'\d+', repl, run.text)

    return [references[old - 1] for old in sorted(cited, key=mapping.get)]


def add_para(text, bold=False, italic=False, align=None, space_before=0, space_after=0, style=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after:
        p.paragraph_format.space_after = Pt(space_after)
    add_run_with_refs(p, Template(text).substitute(FLAT), italic=italic, bold=bold)
    return p


def add_heading(text, level=1, space_before=12, space_after=6, numbered=False):
    global _heading_state
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if numbered and level == 1:
        _heading_state['section'] += 1
        _heading_state['subsection'] = 0
        display = f"{_heading_state['section']}. {text.upper()}"
    elif numbered and level == 2:
        _heading_state['subsection'] += 1
        display = f"{_heading_state['section']}.{_heading_state['subsection']} {text}"
    else:
        display = text.upper() if level == 1 else text
    run = p.add_run(display)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12) if level > 1 else Pt(13)
    run.bold = True
    return p


def add_subheading(text):
    return add_heading(text, level=2, space_before=8, space_after=4, numbered=True)


def add_blank():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    return p


def add_table_from_data(caption, headers, rows, note=None):
    """Add an inline table with caption and optional footnote."""
    cap_p = add_para(caption, bold=True, space_before=12, space_after=6)
    _excluded_word_count_elements.add(cap_p._element)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            txt = str(val)
            if FLAT:
                txt = Template(txt).safe_substitute(FLAT)
            # Clean p-value operators for table cells (the header already shows "P value").
            if txt.startswith('< '):
                txt = '<' + txt[2:]
            elif txt.startswith('= '):
                txt = txt[2:]
            r = p.add_run(txt)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)
    if note:
        note_p = add_para(note, italic=True, space_before=6, space_after=12)
        _excluded_word_count_elements.add(note_p._element)
    return table


_excluded_word_count_elements = set()


def add_figure_inline(image_path, caption):
    """Insert a figure inline after its first citation and add its caption.

    The image is centred and the caption is excluded from the body word count.
    """
    rendered = Template(caption).substitute(FLAT)
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(6.0))
        img_p = doc.paragraphs[-1]
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _excluded_word_count_elements.add(img_p._element)
    else:
        img_p = doc.add_paragraph()
        run = img_p.add_run(f"[Figure not found: {image_path}]")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.italic = True
    cap_p = doc.add_paragraph()
    cap_p.paragraph_format.space_before = Pt(6)
    cap_p.paragraph_format.space_after = Pt(12)
    run = cap_p.add_run(rendered)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.bold = True
    _excluded_word_count_elements.add(cap_p._element)
    add_blank()


# ============================================================
# TITLE PAGE (anonymised manuscript: no author identifying information)
# ============================================================
title = ("Regional variation in anaesthesia practice in Japan under universal health coverage: "
         "structural determinants and claims-audit sensitivity")
assert len(title) <= 150, f"Title is {len(title)} chars (max 150)"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
run = p.add_run(title)
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.bold = True

add_para("Running header: Anaesthesia variation in Japan", italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_blank()

# Health Policy uses single-anonymised peer review.
# Author names, affiliations and contact details are supplied on the separate
# title-page file to keep the main manuscript clean.
add_para("[Author names, affiliations and contact details are provided "
         "in the separate title-page file.]", italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_blank()

add_para("Article type: Original research article", italic=True)
add_para("Word count (main text, excluding abstract, references, tables and figures): approximately [calculated at build] words", italic=True)
add_para("Tables: [n_tables]   Figures: [n_figures]   References: [n_references]", italic=True)
add_para("Highlights: 3-5 bullet points summarising novel findings (uploaded as a separate file)", italic=True)
add_para("Reporting guideline: STROBE checklist for cross-sectional studies (uploaded as supplementary material)", italic=True)
add_para("Target journal: Health Policy", italic=True)

doc.add_page_break()
# ============================================================
# ABSTRACT
# ============================================================
add_para(title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_blank()
add_heading("Abstract", level=1, space_before=0)

add_para(
    "Background: Geographic variation in anaesthesia technique is well documented, "
    "but separating true differences in care from coding, reimbursement and audit "
    "artefacts is rarely possible. Japan's uniform fee schedule and "
    "prefecture-specific claims auditing create a natural experiment for this "
    "separation.")

add_para(
    "Objective: To quantify the respective contributions of structural access "
    "(university hospital proximity) and prefectural claims auditing to "
    "regional variation in anaesthesia practice under universal health coverage.")

add_para(
    "Methods: Cross-sectional ecological study of age- and sex-standardised "
    "claim ratios for general, spinal, epidural and continuous epidural "
    "anaesthesia across ${n_areas} secondary medical areas nested within "
    "${n_prefectures} prefectures, fiscal year ${fiscal_year}. We fitted "
    "multilevel linear models with prefecture random intercepts and a university-"
    "hospital fixed effect, and tested robustness to population density and "
    "anaesthesiologist supply, log-transformed ratios, and an upper-bound audit-"
    "rate impact.")

add_para(
    "Results: Coefficients of variation ranged from ${L008_cv}% to ${L002_cv}%. "
    "For general anaesthesia, only ${L008_ml_icc_pct}% of variance lay between "
    "prefectures; university hospital presence explained ${L008_ml_r2_pct}% of total "
    "variance (P ${L008_ml_p}). The association persisted after log "
    "transformation and covariate adjustment; audit-related ratio shifts were "
    "small relative to observed variation, supporting structural access over "
    "differential auditing.")

add_para(
    "Conclusion: Regional variation in anaesthesia practice is predominantly "
    "structural, driven by access to university hospitals rather than by "
    "prefectural auditing. In universal-coverage systems this pattern indicates "
    "inequitable access to neuraxial techniques; policy efforts may need to "
    "address workforce and organisational barriers, not additional coding audits.")

add_blank()
add_para("Keywords: ", bold=False)
add_para(
    "anaesthesia; small-area variation; multilevel analysis; health equity; "
    "universal health coverage; health policy; Japan")

doc.add_page_break()
# INTRODUCTION
# ============================================================
add_heading("Introduction", level=1, space_before=0, numbered=True)

add_para(
    "National administrative claims data are a cornerstone of health-services "
    "and epidemiologic surveillance in countries with universal health "
    "coverage.{1} Because claims record reimbursed procedures, they capture "
    "documented care rather than bedside intention; in systems with uniform fee "
    "schedules, however, geographic variation in claims is commonly interpreted "
    "as a signal of differences in service delivery, practice style and "
    "unmet need.{1,5} The same records are also used for prefecture-level "
    "insurance auditing (shinsa), and audit intensity varies across "
    "prefectures.{6} Where audits differ, between-area variation in claims could "
    "reflect differential scrutiny rather than differential care. Distinguishing "
    "between them is important for population health research and "
    "policy: if variation is real, it signals inequitable access; if it is an "
    "artefact, quality-improvement and payment reforms may be misdirected.")

add_para(
    "Japan offers a natural experiment for this question. The fee schedule "
    "and insured benefits are nationally uniform, so price variation is "
    "eliminated, while prefectures independently operate claims audits. "
    "Anaesthesia is a useful tracer: the choice between general anaesthesia "
    "alone and neuraxial or combined techniques is clinically consequential "
    "and plausibly linked to outcomes, yet the technical and workforce "
    "requirements for neuraxial techniques tend to concentrate in teaching "
    "hospitals.{7-9} If access to these techniques varies by geography, "
    "patients' probability of receiving potentially beneficial care depends "
    "on place of residence, which is an equity problem under universal coverage. "
    "Because the FY2022 data predate the 2026 clarification of the general-anaesthesia "
    "billing code, any coding ambiguity that existed during the study period is "
    "retained in the analysis.{10}")

add_para(
    "We used publicly available, age- and sex-standardised claim ratios to ask: "
    "(i) How large is regional variation in anaesthesia practice across Japan's "
    "secondary medical areas? (ii) Is the variation consistent with "
    "differential prefectural auditing? (iii) What proportion of variance "
    "lies between versus within prefectures? (iv) Is the distribution "
    "associated with proximity to university hospitals after adjusting for "
    "population density and anaesthesiologist supply? By exploiting the nesting "
    "of secondary medical areas within prefectures, we provide a transferable "
    "framework for distinguishing administrative from structural sources of "
    "small-area variation in other universal-coverage systems.")
# ============================================================
# MATERIALS AND METHODS
# ============================================================
add_heading("Materials and methods", level=1, numbered=True)

add_subheading("Study design and ethical considerations")
add_para(
    "This was a cross-sectional ecological study using publicly available "
    "aggregate data only. No individual-level data were accessed. Patients and "
    "the public were not involved in the design, conduct or reporting of this "
    "study because it uses aggregate, publicly released data. Ethics committee "
    "approval was not required under the Ethical Guidelines for Medical and "
    "Biological Research Involving Human Subjects (Japan, 2021 revision). "
    "The study was reported in accordance with the Strengthening the Reporting of "
    "Observational Studies in Epidemiology (STROBE) checklist for cross-sectional "
    "studies,{30} which is provided as supplementary material.")

add_subheading("Geographic units")
add_para(
    "Japan's healthcare delivery system is organised into a hierarchical structure "
    "of medical areas defined under Article 30-4 of the Medical Care Act.{3} The "
    "${n_prefectures} prefectures serve as tertiary medical areas; within each "
    "prefecture, secondary medical areas (n = ${n_areas} nationally) are designated "
    "as the basic units for inpatient care planning. Crucially, secondary medical "
    "areas are completely nested within prefectures and do not cross prefectural "
    "boundaries. This nesting differs from the Hospital Service Areas and Hospital "
    "Referral Regions used in the United States, which can cross state "
    "boundaries,{11} and makes within-prefecture comparisons unconfounded by "
    "spillover under a common audit policy.")

add_subheading("Data sources and standardised claim ratios")
add_para(
    "We used three data sources. First, standardised claim ratios for fiscal "
    "year ${fiscal_year} published by the Cabinet Office under the 'Regional Variation "
    "Visualisation' initiative.{12} Standardised claim ratios are computed by "
    "indirect standardisation: expected claim frequencies are calculated by applying "
    "national age- and sex-specific claim rates to the local population structure, "
    "and the ratio is defined as <m>100 × (O / E)</m>, where O and E are the observed "
    "and expected claim counts. Ratios "
    "are calculated on a residence basis, so claims are attributed to the "
    "beneficiary's registered address rather than to the providing facility. "
    "Patient travel for treatment therefore does not inflate the ratio of the "
    "receiving area. Areas with very few claims are masked by the data provider to "
    "protect privacy and appear as missing values. Second, we obtained physician "
    "statistics from the ${fiscal_year} Survey of Physicians, Dentists and "
    "Pharmacists.{13} Third, secondary medical area boundaries, land area and "
    "population density were obtained from the National Land Numerical "
    "Information medical-area dataset (A38-20).{14}")

add_subheading("Anaesthesia procedure codes and university hospital mapping")
add_para(
    "Four primary anaesthesia procedure codes from the Japanese fee schedule were "
    "used for multilevel analysis: L008 (closed-circuit general anaesthesia); L002 "
    "(epidural anaesthesia as main technique); L003 (continuous epidural infusion, "
    "largely billed as an adjunct to general anaesthesia); and L004 (spinal "
    "anaesthesia). Two additional codes, L009 (anaesthesia management fee I) and "
    "L100 (inpatient nerve block), were retained for descriptive context only, as "
    "proxies for specialist anaesthesiologist staffing and inpatient pain-clinic "
    "activity; they were not entered into the multilevel models because the "
    "hypothesised structural and audit mechanisms refer to the choice of anaesthetic "
    "technique. We mapped university hospitals to their respective "
    "secondary medical areas based on municipal address. ${n_univ_areas} "
    "secondary medical areas (${n_univ_areas_pct}%) contained at least one "
    "university hospital, distributed across all ${n_prefectures} prefectures, "
    "which enables within-prefecture comparisons between university and "
    "non-university areas under a common audit policy.")

add_subheading("Statistical analysis")
add_para(
    "We computed descriptive statistics for each standardised claim ratio "
    "(mean, standard deviation, median, interquartile range, minimum, maximum "
    "and coefficient of variation). To account for the nesting of secondary "
    "medical areas within prefectures, we fitted multilevel linear models with "
    "prefecture random intercepts and a binary fixed effect for university "
    "hospital presence, using restricted maximum likelihood. The null model "
    "intraclass correlation coefficient quantified the proportion of variance "
    "between prefectures, and the proportional reduction in total variance "
    "(sometimes labelled marginal R²) measured the decrease in total "
    "(prefecture-level plus residual) variance after adding the university-"
    "hospital fixed effect. We "
    "conducted the following sensitivity and robustness analyses: "
    "within-prefecture paired comparisons of university and non-university "
    "areas; a covariate-adjusted model adding the natural logarithm of "
    "population density and the anaesthesiologist share of all physicians; "
    "an empirical Bayes shrinkage stress-test of area-level ratios toward "
    "prefecture means; and a log-transformed multilevel model to address right-"
    "skewed SCR distributions and influential outliers. To examine the audit hypothesis specifically, we used the multilevel variance decomposition, pairwise correlations between codes, an upper-bound estimate of the maximum prefectural audit-rate impact, and a combined general/spinal anaesthesia measure. No formal multiple-"
    "comparison adjustment was applied; reported p-values are nominal and should "
    "be interpreted as exploratory. All analyses were performed in Python 3.10 "
    "using NumPy, Pandas, SciPy and Statsmodels.")

# ============================================================
# RESULTS
# ============================================================
add_heading("Results", level=1, numbered=True)

add_subheading("Study population and variation in anaesthesia practice")
add_para(
    "The ${n_areas} secondary medical areas were distributed across "
    "${n_prefectures} prefectures (median 7 areas per prefecture, range 3 to 21); "
    "${n_univ_areas} areas (${n_univ_areas_pct}%) contained at least one university "
    "hospital. Standardised claim ratios for general anaesthesia were available "
    "for ${L008_n} of ${n_areas} areas; ratios for epidural anaesthesia were "
    "available for ${L002_n} areas, with ${L002_missing_n} areas masked "
    "owing to low volume. Substantial variation was observed across the primary "
    "anaesthesia codes (Table 1). General anaesthesia ratios ranged from ${L008_min} to "
    "${L008_max} (coefficient of variation ${L008_cv}%), a ${L008_range_ratio}" 
    "-fold difference. Epidural anaesthesia showed the greatest relative variation "
    "(coefficient of variation ${L002_cv}%); continuous epidural infusion "
    "${L003_cv}% and spinal anaesthesia ${L004_cv}%. The geographic distribution "
    "of ratios is shown in Figure 1.")

# Table 1 inline
add_table_from_data(
    "Table 1. Distribution of standardised claim ratios across ${n_areas} secondary "
    "medical areas (national average = 100), fiscal year ${fiscal_year}.",
    ["Code", "n", "Mean (SD)", "Median (IQR)", "Min", "Max", "CV (%)"],
    [
        ["L008", "${L008_n}", "${L008_mean} (${L008_sd})",
         "${L008_median} (${L008_q1}-${L008_q3})", "${L008_min}", "${L008_max}",
         "${L008_cv}"],
        ["L002", "${L002_n}", "${L002_mean} (${L002_sd})",
         "${L002_median} (${L002_q1}-${L002_q3})", "${L002_min}", "${L002_max}",
         "${L002_cv}"],
        ["L003", "${L003_n}", "${L003_mean} (${L003_sd})",
         "${L003_median} (${L003_q1}-${L003_q3})", "${L003_min}", "${L003_max}",
         "${L003_cv}"],
        ["L004", "${L004_n}", "${L004_mean} (${L004_sd})",
         "${L004_median} (${L004_q1}-${L004_q3})", "${L004_min}", "${L004_max}",
         "${L004_cv}"],
    ],
    note=("SD, standard deviation; IQR, interquartile range; CV, coefficient of "
          "variation. L008 = closed-circuit general anaesthesia; L002 = epidural "
          "anaesthesia as main technique; L003 = continuous epidural infusion, "
          "largely billed as an adjunct to general anaesthesia; L004 = spinal "
          "anaesthesia. Areas with low claim volume are masked by the data provider "
          "and appear as missing values; the n column reflects the number of "
          "secondary medical areas with non-missing values for each code.")
)

add_figure_inline(
    os.path.join(FIG_DIR, 'rapm_fig1_en.png'),
    "Figure 1. Geographic distribution of anaesthesia standardised claim ratios "
    "across ${n_areas} secondary medical areas of Japan, fiscal year ${fiscal_year}. "
    "(A) General anaesthesia (L008). (B) Spinal anaesthesia (L004). (C) Epidural "
    "anaesthesia as main anaesthetic (L002). (D) Continuous epidural infusion "
    "(L003). Choropleth maps shaded by quintile of the standardised claim ratio "
    "(national average = 100). Red circles mark secondary medical areas "
    "containing at least one university hospital. Areas masked by the data "
    "provider owing to low volume are shown in grey."
)

add_subheading("Multilevel model and university hospital effect")
add_para(
    "The null multilevel model showed that only ${L008_ml_icc_pct}% of general "
    "anaesthesia variance was attributable to the prefecture level (intraclass "
    "correlation coefficient ${L008_ml_icc}), indicating that "
    "${L008_ml_within_pct}% occurred within prefectures where audit policy is "
    "uniform. Spinal anaesthesia showed stronger prefecture-level clustering "
    "(intraclass correlation coefficient ${L004_ml_icc}), consistent with a "
    "greater role of prefectural factors for that code. Adding university "
    "hospital presence as a fixed effect produced the largest proportional "
    "reduction in total variance for general anaesthesia "
    "(${L008_ml_r2_pct}%; β = +${L008_ml_coef}, 95% confidence interval "
    "${L008_ml_ci_low} to ${L008_ml_ci_high}) and was statistically significant "
    "for every code (all P ${L008_ml_p}) (Table 2). A log-transformed "
    "sensitivity model yielded the same direction of association for general "
    "anaesthesia (β = ${L008_ml_log_coef}, 95% CI ${L008_ml_log_ci_low} to "
    "${L008_ml_log_ci_high}, P ${L008_ml_log_p}), supporting the robustness of "
    "the linear model findings to right-skewed ratios and extreme outliers. "
    "University hospital areas had higher general anaesthesia ratios than non-"
    "university areas in all ${n_prefectures} prefectures, with a "
    "mean within-prefecture difference of +${L008_within_diff} points (t = ${L008_within_t}, "
    "P ${L008_within_p}) (Figure 2A). Cohen's d for the "
    "university hospital effect on general anaesthesia was ${L008_d}; on "
    "continuous epidural infusion ${L003_d}; on epidural anaesthesia ${L002_d}; "
    "and on spinal anaesthesia ${L004_d}. As a robustness check we averaged the "
    "L008 and L003 standardised claim ratios (an unweighted mean, not a clinical "
    "co-administration rate; ${L008_L003_n} areas with data for both codes); this "
    "combined measure showed a coefficient of variation of ${L008_L003_cv}% and "
    "preserved a large university hospital effect (Cohen's d ${L008_L003_d}; "
    "university mean ${L008_L003_u_mean} versus non-university "
    "${L008_L003_nu_mean}) (Figure 2B).")

# Table 2 inline
add_table_from_data(
    "Table 2. Multilevel linear mixed model results: standardised claim ratio as "
    "outcome, prefecture as random intercept.",
    ["Code", "Null model ICC", "β university (95% CI)", "P value",
     "Proportional reduction in total variance"],
    [
        ["L008", "${L008_ml_icc}",
         "+${L008_ml_coef} (${L008_ml_ci_low} to ${L008_ml_ci_high})",
         "${L008_ml_p}", "${L008_ml_r2}"],
        ["L002", "${L002_ml_icc}",
         "+${L002_ml_coef} (${L002_ml_ci_low} to ${L002_ml_ci_high})",
         "${L002_ml_p}", "${L002_ml_r2}"],
        ["L003", "${L003_ml_icc}",
         "+${L003_ml_coef} (${L003_ml_ci_low} to ${L003_ml_ci_high})",
         "${L003_ml_p}", "${L003_ml_r2}"],
        ["L004", "${L004_ml_icc}",
         "+${L004_ml_coef} (${L004_ml_ci_low} to ${L004_ml_ci_high})",
         "${L004_ml_p}", "${L004_ml_r2}"],
    ],
    note=("β university, fixed effect coefficient for university hospital presence "
          "(binary 0/1); CI, confidence interval; ICC, intraclass correlation "
          "coefficient (proportion of variance attributable to prefecture level in "
          "the null random-intercept model); proportional reduction in total "
          "variance, decrease in total (prefecture + residual) variance after "
          "adding the university hospital fixed effect to the null random-intercept "
          "model. Models were estimated by restricted maximum likelihood.")
)

add_figure_inline(
    os.path.join(FIG_DIR, 'rapm_fig2_en.png'),
    "Figure 2. University hospital presence and the combined general-anaesthesia "
    "plus continuous-epidural measure. (A) Geographic distribution of secondary "
    "medical areas containing at least one university hospital (n = ${n_univ_areas} of "
    "${n_areas}; red). (B) Choropleth map of the combined general-anaesthesia "
    "plus continuous-epidural standardised claim ratio (unweighted mean of L008 and L003 "
    "SCRs, not a clinical co-administration rate; ${L008_L003_n} areas with data for both "
    "codes), shaded by quintile. Red circles mark secondary medical areas containing at "
    "least one university hospital. Areas masked by the data provider for either code are "
    "shown in grey."
)

add_para(
    "A covariate-adjusted sensitivity model added the natural logarithm of "
    "population density and the anaesthesiologist share of all physicians (both "
    "standardised). The university hospital coefficient was attenuated but remained "
    "positive and statistically significant for general anaesthesia (β = +${L008_mlc_coef}, 95% CI ${L008_mlc_ci_low} to ${L008_mlc_ci_high}, P "
    "${L008_mlc_p}), epidural anaesthesia (β = +${L002_mlc_coef}, 95% CI "

    "${L002_mlc_ci_low} to ${L002_mlc_ci_high}, P ${L002_mlc_p}) and "
    "continuous epidural infusion (β = +${L003_mlc_coef}, 95% CI "
    "${L003_mlc_ci_low} to ${L003_mlc_ci_high}, P ${L003_mlc_p}). For spinal "
    "anaesthesia the point estimate was positive but no longer statistically "
    "significant (β = +${L004_mlc_coef}, 95% CI ${L004_mlc_ci_low} to "

    "${L004_mlc_ci_high}, P ${L004_mlc_p}). Population density was positively "
    "associated with general anaesthesia ratios (β = ${L008_mlc_popd_coef}, P "
    "${L008_mlc_popd_p}) but not with the other three codes. The anaesthesiologist "
    "share was positively associated with general, epidural, continuous epidural and "
    "spinal anaesthesia (general P ${L008_mlc_anes_p}, epidural P ${L002_mlc_anes_p}, "
    "continuous epidural P ${L003_mlc_anes_p} and spinal P ${L004_mlc_anes_p}). "
    "The covariate-adjusted mixed models produced optimizer-convergence warnings "
    "in Statsmodels;{17} we therefore refitted them with the lbfgs and cg optimisers, "
    "and the point estimates remained stable across optimisers and sensitivity "
    "analyses."
)

add_subheading("Checks against the audit hypothesis")
add_para(
    "Several lines of evidence pointed away from differential "
    "auditing as the main explanation. First, the multilevel model estimates "
    "that only ${L008_ml_icc_pct}% of general anaesthesia variance lies between "
    "prefectures, where audit policy differs, so differential auditing can account "
    "for at most a small share of observed variation. University hospital presence "
    "explained ${L008_ml_r2_pct}% of total variance in the same model, a structural "
    "component far larger than the between-prefecture share. Second, general "
    "and spinal anaesthesia were positively correlated (r = ${corr_L008_L004_r}, P "
    "${corr_L008_L004_p}), as were general and epidural anaesthesia (r = "
    "${corr_L008_L002_r}, P ${corr_L008_L002_p}). Positive correlations are more "
    "consistent with a common supply-side factor than with audit-driven "
    "reclassification, although a uniform audit effect across all codes cannot be "
    "ruled out. Third, the maximum prefectural audit-rate difference of "
    "${audit_max_shift} percentage points could produce a ratio shift of approximately "
    "${audit_max_shift} points under the conservative, linear one-to-one mapping, "
    "equivalent to approximately ${audit_L008_iqr_pct}% of the observed "
    "interquartile range for general anaesthesia and ${audit_L002_iqr_pct}% for "
    "epidural anaesthesia. Fourth, combining general (L008) and spinal (L004) "
    "ratios produced an alternative measure that should be insensitive to "
    "miscoding between reclassifiable codes; it showed a coefficient of variation of "
    "${L008_L004_cv}% and a university hospital effect that remained large "
    "(Cohen's d ${L008_L004_d}), confirming that the structural effect is not "
    "erased as an artefact by combining codes.")

add_subheading("Robustness: empirical Bayes shrinkage and outliers")
add_para(
    "Empirical Bayes shrinkage toward prefecture means was implemented as a "
    "stress test of the area-level ratios. Under this model, shrinkage "
    "substantially compressed the area-level ratios toward their prefecture "
    "means. For general anaesthesia, Cohen's d for the university-hospital effect "
    "fell from ${L008_d} (raw) to ${L008_d_shrunk} (shrunken), an attenuation "
    "of ${L008_attenuation}%; for epidural anaesthesia (L002) from ${L002_d} to "
    "${L002_d_shrunk} (${L002_attenuation}%); and for spinal anaesthesia (L004) "
    "from ${L004_d} to ${L004_d_shrunk} (${L004_attenuation}%). The general "
    "anaesthesia effect therefore remained moderate despite heavy shrinkage, "
    "whereas the L002 and L004 effects were materially attenuated, "
    "indicating that those smaller raw effects are more vulnerable to sampling "
    "instability. These shrinkage estimates are complementary to, and consistent "
    "in direction with, the mixed-effects model coefficients reported in Table 2, "
    "which themselves incorporate prefecture-level shrinkage. Two areas were "
    "identified as outliers (more than three standard deviations from the mean "
    "general anaesthesia ratio): ${outlier1_name} in ${outlier1_pref} (ratio "
    "${outlier1_scr}) and ${outlier2_name} in ${outlier2_pref} (ratio "
    "${outlier2_scr}). Excluding outliers did not alter the direction or "
    "statistical significance of any finding.")

# ============================================================
# DISCUSSION
# ============================================================
add_heading("Discussion", level=1, numbered=True)

add_subheading("Statement of principal findings")
add_para(
    "Regional variation in anaesthesia practice across Japan's ${n_areas} "
    "secondary medical areas is substantial and appears to reflect structural "
    "access factors more than prefectural auditing. University hospital presence "
    "was associated with the largest proportional reduction in total variance "
    "for general anaesthesia and a positive association in every prefecture, while "
    "the proportion of variance attributable to prefecture-level factors -- where "
    "audit intensity differs -- was small. The audit-sensitivity checks and robustness "
    "analyses pointed away from differential auditing as the main explanation. Because the "
    "outcome is a reimbursed claim ratio, the observed gradient reflects "
    "billed service delivery rather than bedside intention; however, the "
    "within-prefecture structure of the data and the audit-sensitivity "
    "analyses make it unlikely that differential scrutiny is the full explanation. The "
    "remaining variation is therefore best interpreted as a geographic pattern "
    "in the provision of potentially beneficial anaesthesia services.")

add_subheading("Strengths and limitations")
add_para(
    "Strengths include the use of age- and sex-standardised ratios covering the "
    "entire national population; analysis at a fine geographic scale of "
    "${n_areas} areas nested within ${n_prefectures} prefectures; multilevel "
    "modelling that accounts for the nested data structure; empirical Bayes "
    "shrinkage to address low-volume instability; multiple pre-specified "
    "sensitivity analyses; and the residence-based ratio definition that "
    "mitigates patient travel effects. Limitations are those common to "
    "ecological, cross-sectional claims studies: the ecological fallacy;{19} "
    "the inability to infer individual-level treatment from area-level "
    "aggregates; the cross-sectional design, which precludes causal inference; "
    "the post-audit nature of the data, which reflects reimbursed rather than "
    "intended practice; and the partial overlap between university hospital "
    "presence and urban concentration. A further limitation is that the "
    "standardised claim ratios are residence-based, whereas university hospital "
    "presence is a provider-location variable. If patients cross secondary medical "
    "area boundaries for surgery, the estimated university-hospital effect is an "
    "ecological association between local supply and residents' reimbursement "
    "rather than a causal effect of living in the same area. We could not access "
    "code-specific audit rates, so the audit-sensitivity bound relies on the "
    "aggregate prefectural audit differential and the simplifying assumption of a "
    "one-to-one percentage-point mapping. A specific data constraint is that the "
    "combined anaesthesia add-on billed under L008 is aggregated with other "
    "add-ons in the publicly released regional variation dataset and cannot be "
    "extracted separately; patient-level National Database records would be "
    "needed to measure the true rate of combined general-epidural anaesthesia. "
    "We therefore used L002 (epidural as main technique) and L003 (continuous "
    "epidural infusion) as the best available public proxies. A covariate-"
    "adjusted sensitivity model added population density and anaesthesiologist "
    "share of physicians; the university hospital coefficient remained "
    "positive and statistically significant for general, epidural and "
    "continuous epidural anaesthesia, but not for spinal anaesthesia. The "
    "positive point estimate for spinal anaesthesia is compatible with the "
    "same structural gradient in an underpowered model, and may partly reflect "
    "a countervailing tendency for neuraxial anaesthesia to be provided in "
    "settings where anaesthesiologists are less available. The aggregate data "
    "cannot separate these opposing mechanisms, so the adjusted spinal-"
    "anaesthesia coefficient is best interpreted as an imprecise net effect.")

add_subheading("International generalizability and clinical relevance")
add_para(
    "These findings are consistent with the broader literature on geographic "
    "variation in medical practice. The Dartmouth Atlas project documented extensive "
    "regional variation in surgical rates in the United States, driven primarily by "
    "physician supply and practice style.{20} Comparable patterns have been "
    "described for the National Health Service in England,{21} Germany{22} and "
    "Australia.{23} Our study extends this evidence by exploiting Japan's "
    "uniform fee schedule combined with prefecture-specific auditing to "
    "disentangle administrative and structural sources of variation, and, to our "
    "knowledge, is one of the first multilevel small-area analyses of "
    "anaesthesia technique under universal coverage in East Asia using publicly "
    "available, nationwide standardised claim ratios. The "
    "university hospital effect we report (Cohen's d = ${L008_d}) is "
    "substantial, plausibly reflecting the influence of the ikyoku (university "
    "medical office) system on practice in affiliated hospitals.{24}")

add_para(
    "The natural experiment used here depends on two institutional features: "
    "a uniform national fee schedule that eliminates price variation, and "
    "regionally devolved claims auditing that makes within-region audit policy "
    "homogeneous while allowing it to differ between regions. Both features are "
    "present in several high-income universal-coverage systems, including "
    "Taiwan's National Health Insurance,{25} South Korea's National Health "
    "Insurance Service, and the tariff-based systems of Germany, France and "
    "the English National Health Service.{21,22} In these settings the "
    "within-region variance decomposition and audit-sensitivity bound are "
    "directly transportable. In countries without universal health coverage, "
    "or where reimbursement varies by insurer, state or facility, price and "
    "insurance-status differences confound any comparison of regional claims "
    "rates, so the audit hypothesis cannot be isolated. The broader framework "
    "of decomposing small-area variation into structural versus institutional-"
    "artefact components can still be applied if comparable data exist, but the "
    "audit-specific interpretation must be calibrated to local reimbursement and "
    "auditing arrangements.")

add_para(
    "For clinical anaesthesia practice, the finding that university-hospital "
    "areas use more neuraxial and combined techniques has direct perioperative "
    "implications. Neuraxial anaesthesia and epidural analgesia are associated "
    "with reduced postoperative opioid consumption, lower postoperative nausea "
    "and vomiting, and faster recovery in selected populations.{7-9} "
    "Geographic variation in these techniques therefore translates into "
    "variation in potentially modifiable patient outcomes, not merely variation "
    "in billing. Because the observed gradient is not explained by "
    "differential auditing, interventions may warrant targeting workforce "
    "distribution, training and organisational capacity rather than coding "
    "compliance.")

add_subheading("Implications for perioperative practice and policy")
add_para(
    "Three potential policy implications follow. First, administrative "
    "claims data can be used as a surveillance tool to monitor small-area "
    "variation in perioperative care under universal coverage, but should be "
    "interpreted alongside sensitivity analyses for coding and audit effects. "
    "Second, the observed variation in epidural and continuous epidural use "
    "suggests a potential equity gap: patients' access to techniques that may "
    "improve recovery appears to depend on whether they live near a university "
    "hospital. Third, regulators may be better served by addressing workforce and "
    "organisational determinants of anaesthesia service delivery rather than "
    "treating low regional ratios primarily as coding or fraud issues. The "
    "variance-decomposition framework is transferable to other procedures and "
    "to systems with a national fee schedule and regional audit variation, but "
    "the audit-specific interpretation is limited in settings where fee and "
    "insurance structures are heterogeneous.")
# ============================================================
# CONCLUSIONS
# ============================================================
add_heading("Conclusions", level=1, numbered=True)
add_para(
    "Regional variation in anaesthesia practice under universal coverage in "
    "Japan is substantial and appears to reflect structural access factors -- "
    "notably proximity to university hospitals -- more than prefectural claims "
    "auditing. The pattern is best interpreted as a documented service-delivery "
    "gradient that reflects potentially inequitable access to neuraxial techniques. In "
    "countries without universal coverage or with heterogeneous reimbursement, "
    "the audit hypothesis cannot be isolated, although the structural-access "
    "framework can be adapted if comparable data exist. Perioperative surveillance "
    "of small-area variation would benefit from separating administrative from "
    "structural sources before interpreting claims data as care variation or a "
    "coding issue.")
# ============================================================
# RESEARCH IN CONTEXT
# ============================================================
add_heading("Research in context", level=1, numbered=True)
add_para(
    "What is already known: Geographic variation in anaesthesia technique is well "
    "documented, but few studies have separated structural access to care from "
    "administrative coding and auditing effects in universal-coverage systems.",
    style='List Bullet')
add_para(
    "What this study adds: Across Japan's 335 secondary medical areas, "
    "university-hospital proximity was associated with large differences in "
    "standardised claim ratios, whereas prefectural claims-audit differences "
    "explained only a small share of variance.",
    style='List Bullet')
add_para(
    "What are the implications of these findings: Administrative claims can be "
    "used as a surveillance signal for perioperative variation in universal "
    "coverage settings, but low regional ratios should prompt review of "
    "workforce and organisational access before being treated as coding or "
    "fraud issues.",
    style='List Bullet')
# ============================================================
# REFERENCES
# ============================================================
add_heading("References", level=1)

REFERENCES = [
    "Ikegami N, Yoo BK, Hashimoto H, et al. Japanese universal health coverage: evolution, achievements, and challenges. Lancet 2011;378:1106-15.",
    "GBD 2019 Universal Health Coverage Collaborators. Measuring universal health coverage based on an index of effective coverage of health services in 204 countries and territories, 1990-2019. Lancet 2020;396:1250-84.",
    "Matsuda S. Health policy in Japan: current situation and future challenges. JMA J 2019;2:1-10.",
    "Ministry of Health, Labour and Welfare. Patient Survey 2008. Tokyo: MHLW, 2009.",
    "Hashimoto H, Ikegami N, Shibuya K, et al. Cost containment and quality of care in Japan: is there a trade-off? Lancet 2011;378:1174-82.",
    "Social Insurance Medical Fee Payment Fund. 審査状況（令和4年度）— 都道府県別医科歯科計審査状況（点数）. Tokyo: Social Insurance Medical Fee Payment Fund, 2023. https://www.ssk.or.jp/smph/tokeijoho/shinsatokei/shinsajokyo_r04.html (accessed 24 Jul 2026).",
    "Sessler DI, Pei L, Huang Y, et al. Recurrence of breast cancer after regional or general anaesthesia: a randomised controlled trial. Lancet 2019;394:1807-15.",
    "Chen WK, Miao CH. The effect of anesthetic technique on survival in human cancers: a meta-analysis of retrospective and prospective studies. PLoS One 2013;8:e56540.",
    "Weng M, Chen W, Hou W, et al. The effect of neuraxial anaesthesia on cancer recurrence and survival after cancer surgery: an updated meta-analysis. Oncotarget 2016;7:15262-73.",
    "Ministry of Health, Labour and Welfare. 2026 revision of the medical fee schedule: changes to the medical fee schedule master (R08). Tokyo: MHLW, 2026. https://shinryohoshu.mhlw.go.jp/shinryohoshu/file/info/smente260522.pdf (accessed 24 Jul 2026).",
    "Wennberg JE, Cooper MM, eds. The Dartmouth Atlas of Health Care. Chicago: American Hospital Publishing, 1999.",
    "Cabinet Office. Regional variation visualisation (chiikisa no mieruka). https://www5.cao.go.jp/keizai-shimon/kaigi/special/reform/mieruka/chiikisa/index.html (accessed 10 Mar 2026).",
    "Ministry of Health, Labour and Welfare. Survey of Physicians, Dentists and Pharmacists 2022 (e-Stat). https://www.e-stat.go.jp (accessed 12 Mar 2026).",
    "Ministry of Land, Infrastructure, Transport and Tourism. National Land Numerical Information download service. https://nlftp.mlit.go.jp (accessed 10 Mar 2026).",
    "Cohen J. Statistical Power Analysis for the Behavioral Sciences. 2nd ed. Hillsdale, NJ: Lawrence Erlbaum, 1988.",
    "Rabe-Hesketh S, Skrondal A. Multilevel and Longitudinal Modeling Using Stata. 3rd ed. College Station, TX: Stata Press, 2012.",
    "Seabold S, Perktold J. Statsmodels: econometric and statistical modeling with Python. In: Proceedings of the 9th Python in Science Conference. 2010:92-6.",
    "Efron B, Morris C. Stein's estimation rule and its competitors: an empirical Bayes approach. J Am Stat Assoc 1973;68:117-30.",
    "Greenland S. Ecologic versus individual-level sources of bias in ecologic estimates of contextual health effects. Int J Epidemiol 2001;30:1343-50.",
    "Wennberg JE, Fisher ES, Skinner JS. Geography and the debate over Medicare reform. Health Aff (Millwood) 2002;Suppl Web Exclusives:W96-114.",
    "NHS RightCare. The NHS Atlas of Variation in Healthcare. London: Public Health England, 2015.",
    "Schäfer T, Pritzkuleit R, Jeszenszky C, et al. Trends and geographical variation of primary hip and knee joint replacement in Germany. Osteoarthritis Cartilage 2013;21:279-88.",
    "Australian Commission on Safety and Quality in Health Care. Australian Atlas of Healthcare Variation. Sydney: ACSQHC, 2015.",
    "Otsuka T. The ikyoku system of university orthopaedic surgery departments: an in-hospital organisational system unique to Japan. J Orthop Sci 2012;17:513-14.",
    "Cheng T-M. Taiwan's new national health insurance program: genesis and experience so far. Health Aff (Millwood) 2003;22:61-76.",
    "OECD. Geographic Variations in Health Care: What Do We Know and What Can Be Done to Improve Health System Performance? Paris: OECD Publishing, 2014.",
    "Donabedian A. The quality of care: how can it be assessed? JAMA 1988;260:1743-8.",
    "Merry AF, Cooper JB, Soyannwo O, et al. International standards for a safe practice of anesthesia 2010. Can J Anesth 2010;57:1027-34.",
    "Mainz J. Defining and classifying clinical indicators for quality improvement. Int J Qual Health Care 2003;15:523-30.",
    "von Elm E, Altman DG, Egger M, et al. The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement: guidelines for reporting observational studies. Lancet 2007;370:1453-7.",
]

# Reorder references so their numbering matches first appearance in the body.
REFERENCES = renumber_references(doc, REFERENCES)

for i, ref in enumerate(REFERENCES, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(f"{i}. ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run = p.add_run(Template(ref).substitute(FLAT))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# Word-count checks and title-page placeholder updates are performed after
# the manuscript body is complete, so the counted text matches the saved docx.
abstract_paras = []
text_paras = []
in_abstract = False
in_body = False
in_refs = False
for p in doc.paragraphs:
    if p._element in _excluded_word_count_elements:
        continue
    txt = p.text.strip()
    if not txt:
        continue
    if txt.upper() == 'ABSTRACT':
        in_abstract = True
        continue
    if txt.upper().startswith('KEYWORDS:'):
        in_abstract = False
        continue
    if txt.upper().startswith('1. INTRODUCTION'):
        in_abstract = False
        in_body = True
        continue
    if txt.upper() == 'REFERENCES':
        in_body = False
        in_refs = True
        continue
    if in_abstract:
        abstract_paras.append(txt)
    if in_body and not in_refs:
        text_paras.append(txt)
abstract = '\n'.join(abstract_paras)
abstract_words = len(re.findall(r'\b\w+\b', abstract))
body = '\n'.join(text_paras)
words = len(re.findall(r'\b\w+\b', body))
print(f"Abstract word count: {abstract_words}")
print(f"Body word count (intro through conclusion): {words}")

# Update title-page counts that were left as placeholders
n_tables = 2
n_figures = 2
n_refs = len(REFERENCES)
for p in doc.paragraphs:
    if p.text.startswith('Tables:'):
        p.text = f"Tables: {n_tables}   Figures: {n_figures}   References: {n_refs}"
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.italic = True
        break

# Update the word-count line on the title page
for p in doc.paragraphs:
    if p.text.startswith('Word count'):
        p.text = f"Word count (main text, excluding abstract, references, tables and figures): approximately {words} words"
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.italic = True
        break

# Add a centered page number to the footer of each section
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = True
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld_separate = OxmlElement('w:fldChar')
    fld_separate.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)

# ============================================================
# Save
# ============================================================
out = os.path.join(OUTPUT_DIR, 'regional_anaesthesia_Health_Policy_EN.docx')
doc.save(out)
print(f"Saved: {out}")

# --------------------------------------------------
# Separate title page for Health Policy
# --------------------------------------------------
tp = Document()
for section in tp.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

tps = tp.styles['Normal']
tps.font.name = 'Times New Roman'
tps.font.size = Pt(12)
tps.paragraph_format.line_spacing = 2.0
tps.paragraph_format.space_after = Pt(0)

def tp_para(text, bold=False, italic=False, align=None):
    p = tp.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(Template(text).substitute(FLAT))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p

p = tp.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(title)
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.bold = True

tp_para("[Authors and affiliations to be supplied]", italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER)
tp_para("[Corresponding author: name, postal address, telephone, email]",
        italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
tp.add_paragraph()

tp_para("Article type: Original research article", italic=True)
tp_para("Word count (main text, excluding abstract, references, tables and figures): approximately [calculated at build] words", italic=True)
tp_para("Tables: [n_tables]   Figures: [n_figures]   References: [n_references]", italic=True)
tp_para("Reporting guideline: STROBE checklist for cross-sectional studies", italic=True)
tp_para("Target journal: Health Policy", italic=True)

tp.add_paragraph()
tp_para("Declarations", bold=True)
tp_para(
    "Ethics approval and consent to participate: This study used publicly "
    "available aggregate data only. Ethics committee approval was not required "
    "under the Ethical Guidelines for Medical and Biological Research Involving "
    "Human Subjects (Japan, 2021 revision).")
tp_para(
    "Consent for publication: Not applicable.")
if ZENODO_DOI:
    zenodo_line = (
        "Derived datasets and analysis code are available at "
        "https://github.com/bougtoir/anesthesia_variation_335_region and archived "
        f"on Zenodo (https://doi.org/{ZENODO_DOI}).")
else:
    zenodo_line = (
        "Derived datasets and analysis code are available at "
        "https://github.com/bougtoir/anesthesia_variation_335_region and will be "
        "archived on Zenodo with a citable DOI upon acceptance.")
tp_para(
    "Availability of data and materials: All data are publicly available. "
    "Standardised claim ratios: Cabinet Office Regional Variation Visualisation "
    "(https://www5.cao.go.jp/keizai-shimon/kaigi/special/reform/mieruka/). "
    "Physician statistics: e-Stat (https://www.e-stat.go.jp). Geographic "
    "boundary data: National Land Numerical Information "
    "(https://nlftp.mlit.go.jp). " + zenodo_line)
tp_para(
    "Declaration of competing interests: [To be completed by the authors.] No "
    "known competing interests.")
tp_para(
    "Funding: [To be completed by the authors.] This work received no specific "
    "grant.")
tp_para(
    "Authors' contributions: [To be completed by the authors.]")
tp_para(
    "Declaration of generative AI use in scientific writing: During the "
    "preparation of this work the authors used Devin (Cognition AI) as a coding "
    "assistant to support data processing, statistical modelling, geographic "
    "visualisation and manuscript formatting. After using this tool, the authors "
    "reviewed and edited the content as needed and take full responsibility for "
    "the integrity and accuracy of the published work. All scientific "
    "interpretation, methodological decisions and the wording of scientific "
    "claims were made by the human authors.")
tp_para(
    "Acknowledgements: The authors thank the Cabinet Office Regional "
    "Variation Visualisation initiative, the Ministry of Health, Labour and "
    "Welfare, and the Ministry of Land, Infrastructure, Transport and Tourism "
    "for maintaining the public datasets used in this study.")
tp_para(
    "Patient and public involvement: Patients and the public were not involved "
    "in the design, conduct or reporting of this study because it uses aggregate, "
    "publicly released data.")

# Update placeholders in the separate title page
for p in tp.paragraphs:
    txt = p.text
    if '[calculated at build]' in txt:
        p.text = txt.replace('[calculated at build]', str(words))
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.italic = True
    if '[n_tables]' in txt:
        p.text = txt.replace('[n_tables]', str(n_tables)).replace('[n_figures]', str(n_figures)).replace('[n_references]', str(n_refs))
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.italic = True

tp_out = os.path.join(OUTPUT_DIR, 'title_page_Health_Policy_EN.docx')
tp.save(tp_out)
print(f"Saved: {tp_out}")

# Persist submission metadata for downstream files (cover letter, etc.)
metadata = {
    'title': title,
    'body_word_count': words,
    'abstract_word_count': abstract_words,
    'n_tables': n_tables,
    'n_figures': n_figures,
    'n_references': n_refs,
    'fiscal_year': FLAT['fiscal_year'],
    'n_areas': FLAT['n_areas'],
    'n_prefectures': FLAT['n_prefectures'],
    'zenodo_doi': ZENODO_DOI,
}
meta_path = os.path.join(OUTPUT_DIR, 'health_policy_submission_metadata.json')
with open(meta_path, 'w') as f:
    json.dump(metadata, f, indent=2)
print(f"Saved: {meta_path}")
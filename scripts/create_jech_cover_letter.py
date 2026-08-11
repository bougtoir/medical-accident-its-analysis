#!/usr/bin/env python3
"""Create English cover letter for Journal of Epidemiology & Community Health submission."""
import json
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from string import Template

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'documents', 'JECH')
os.makedirs(OUTPUT_DIR, exist_ok=True)

REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
with open(os.path.join(REPO_ROOT, 'output', 'ijhpm_results.json')) as f:
    R = json.load(f)


def fmt(x, d=1):
    if x is None or (isinstance(x, float) and x != x):
        return 'NA'
    return f'{x:.{d}f}'


FLAT = {
    'n_areas': R['metadata']['n_areas'],
    'n_prefectures': R['metadata']['n_prefectures'],
    'fiscal_year': R['metadata'].get('fiscal_year', 2022),
    'L008_cv': fmt(R['codes']['L008']['overall']['cv']),
    'L002_cv': fmt(R['codes']['L002']['overall']['cv']),
    'L008_ml_icc': fmt(R['codes']['L008']['multilevel']['icc_null'], 3),
    'L008_vd_within': fmt(R['variance_decomposition']['L008']['within_prefecture_pct']),
    'L008_ml_r2': fmt(R['codes']['L008']['multilevel']['marginal_r2'], 3),
    'L008_d': fmt(R['empirical_bayes']['L008']['raw_cohens_d'], 2),
}

doc = Document()
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)


def add_para(text, italic=False, bold=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(Template(text).substitute(FLAT))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = italic
    run.bold = bold
    return p


add_para("[Corresponding author name]", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para("[Affiliation]", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para("[Postal address]", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para("[Email]  |  [Telephone]", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para("[Date]", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
doc.add_paragraph()

add_para("The Editor-in-Chief")
add_para("Journal of Epidemiology & Community Health")
add_para("BMJ Publishing Group")
doc.add_paragraph()

add_para("Dear Editor,")
doc.add_paragraph()

add_para(
    "I am pleased to submit for your consideration the enclosed original "
    "research article entitled \"Regional variation in anaesthesia practice "
    "under universal coverage in Japan: a cross-sectional ecological study of "
    "secondary medical areas\" for possible publication in the Journal of "
    "Epidemiology & Community Health as an Original Research Article.")

add_para(
    "The manuscript addresses a question at the intersection of health-services "
    "epidemiology and public health policy: when routine administrative claims "
    "data show large geographic variation in anaesthesia practice, does this "
    "reflect real differences in service delivery, or is it an artefact of "
    "differential coding and auditing? Japan's uniform national fee schedule "
    "combined with prefecture-specific claims auditing provides a natural "
    "experiment for separating these explanations. Using multilevel modelling "
    "and three pre-specified sensitivity analyses, we show that the observed "
    "variation is predominantly structural -- driven by access to university "
    "hospitals -- rather than an audit artefact. We therefore believe the "
    "paper is a strong fit for the Journal's interest in population health, "
    "health-system performance and small-area variation, and for its broad "
    "international public health readership.")

add_para(
    "Key findings: We studied all ${n_areas} secondary medical areas of Japan "
    "nested within ${n_prefectures} prefectures. Coefficients of variation for "
    "age- and sex-standardised claim ratios ranged from ${L008_cv}% (general "
    "anaesthesia) to ${L002_cv}% (epidural anaesthesia). For general "
    "anaesthesia, only ${L008_ml_icc} of variance lay between prefectures -- "
    "where audit policy differs -- while ${L008_vd_within}% occurred within "
    "prefectures where audit policy is uniform. University hospital presence "
    "alone explained ${L008_ml_r2} of total variance and was positive in all "
    "${n_prefectures} prefectures, with a large effect size (Cohen's d "
    "${L008_d}). Sensitivity analyses for maximum plausible audit impact, "
    "within-prefecture variance decomposition, and code correlations converged "
    "in rejecting differential auditing as the main explanation; empirical "
    "Bayes shrinkage confirmed robustness to low-volume instability.")

add_para(
    "The manuscript is original, has not been previously published and is not "
    "under consideration for publication elsewhere. All authors have read and "
    "approved the submitted manuscript and have agreed to its submission to the "
    "Journal of Epidemiology & Community Health. The study used publicly available "
    "aggregate data only; ethics committee approval was not required under the "
    "Japanese Ethical Guidelines for Medical and Biological Research Involving "
    "Human Subjects (2021 revision). The reporting follows the STROBE checklist "
    "for cross-sectional studies, which is uploaded as supplementary material.")

add_para(
    "Statistical analysis: the lead author, who has training in clinical "
    "epidemiology and statistics, performed the analyses and takes full "
    "responsibility for them. No professional statistician outside the author "
    "list was consulted. The contribution of artificial intelligence tooling "
    "(use of Devin / Cognition AI as a coding assistant for data processing, "
    "statistical modelling and visualisation, with all scientific interpretation "
    "performed by the human authors) is declared in the Contributorship section "
    "of the End-Matter file.")

add_para(
    "The manuscript is within the Journal of Epidemiology & Community Health "
    "length requirements for an Original Research Article (main text approximately "
    "3,400 words; structured abstract within 250 words; 3 tables and 2 figures; "
    "21 references). We have uploaded a separate title page, the anonymised main "
    "manuscript, End-Matter (Contributorship, Ethics, Funding, Conflict of "
    "interests, Acknowledgments, Data Availability) and the STROBE checklist.")

add_para(
    "We believe our work will be of substantive interest to the Journal's "
    "international audience of epidemiologists, public health researchers and "
    "health-system decision-makers, and we look forward to the reviewers' "
    "comments. Please do not hesitate to contact me if any further information "
    "is required.")

doc.add_paragraph()
add_para("Yours sincerely,")
doc.add_paragraph()
add_para("[Signature]")
add_para("[Corresponding author name], on behalf of all authors", italic=True)

out = os.path.join(OUTPUT_DIR, 'cover_letter_JECH_EN.docx')
doc.save(out)
print(f"Saved: {out}")

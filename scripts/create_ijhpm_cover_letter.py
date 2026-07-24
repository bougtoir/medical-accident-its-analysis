#!/usr/bin/env python3
"""Create English cover letter for IJHPM submission."""
import json
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from string import Template

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'documents', 'IJHPM')
os.makedirs(OUTPUT_DIR, exist_ok=True)

REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
with open(os.path.join(REPO_ROOT, 'output', 'ijhpm_results.json')) as f:
    R = json.load(f)

def fmt(x, d=1):
    if x is None or (isinstance(x, float) and x != x):
        return 'NA'
    return f'{x:.{d}f}'

FLAT = {
    'n_areas': R['metadata']['n_areas'],
    'n_prefectures': R['metadata']['n_prefectures'],
    'fiscal_year': R['metadata'].get('fiscal_year', 2022),
    'L008_cv': fmt(R['codes']['L008']['overall']['cv']),
    'L002_cv': fmt(R['codes']['L002']['overall']['cv']),
    'L008_ml_icc': fmt(R['codes']['L008']['multilevel']['icc_null'], 3),
    'L008_vd_within': fmt(R['variance_decomposition']['L008']['within_prefecture_pct']),
    'L008_ml_r2': fmt(R['codes']['L008']['multilevel']['marginal_r2'], 3),
    'L008_d': fmt(R['empirical_bayes']['L008']['raw_cohens_d'], 2),
}

doc = Document()
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)


def add_para(text, italic=False, bold=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(Template(text).substitute(FLAT))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = italic
    run.bold = bold
    return p


add_para("[Corresponding author name]", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para("[Affiliation]", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para("[Postal address]", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para("[Email]  |  [Telephone]", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para("[Date]", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
doc.add_paragraph()

add_para("The Editor-in-Chief")
add_para("International Journal of Health Policy and Management")
add_para("Kerman University of Medical Sciences")
doc.add_paragraph()

add_para("Dear Editor,")
doc.add_paragraph()

add_para(
    "I am pleased to submit for your consideration the enclosed original "
    "research article entitled \"Regional variation in anaesthesia practice "
    "in Japan: structural determinant or claims-audit artefact?\" for "
    "possible publication in the International Journal of Health Policy and "
    "Management as an Original Research Article.")

add_para(
    "We believe that our manuscript is well aligned with IJHPM's focus on "
    "health policy, health system governance and the translation of empirical "
    "evidence into policy action. The study uses Japan's single national fee "
    "schedule and prefecture-specific insurance auditing as a natural "
    "experiment to answer a question that is central to any centrally "
    "governed, regionally administered health system: when claims-based "
    "variation is observed, does it reflect structural determinants of care or "
    "artefacts of differential auditing? Our within-prefecture variance "
    "decomposition offers a transferable governance framework for distinguishing "
    "administrative from structural sources of small-area variation, and our "
    "findings carry direct implications for regulation, quality improvement and "
    "equity policy under universal coverage.")

add_para(
    "Key findings: We studied all ${n_areas} secondary medical areas of Japan. "
    "Coefficients of variation across areas ranged from ${L008_cv}% (general "
    "anaesthesia) to ${L002_cv}% (epidural anaesthesia). Multilevel modelling "
    "showed that only ${L008_ml_icc} of general anaesthesia variance lay between "
    "prefectures — where audit policy differs — while ${L008_vd_within}% occurred within "
    "prefectures where audit policy is uniform. University hospital presence "
    "alone explained ${L008_ml_r2} of total variance and was positive in all ${n_prefectures} "
    "prefectures, with a large effect size (Cohen's d ${L008_d}). Three pre-"
    "specified sensitivity analyses converged in rejecting differential auditing "
    "as a plausible explanation, and empirical Bayes shrinkage confirmed that "
    "the findings are robust to low-volume instability. We conclude that the "
    "observed variation is predominantly structural and institutional, not "
    "administrative, and identify the marked variation in epidural and "
    "continuous epidural infusion as a modifiable equity concern under "
    "universal coverage.")

add_para(
    "The manuscript is original, has not been previously published and is not "
    "under consideration for publication elsewhere. All authors have read and "
    "approved the submitted manuscript and have agreed to its submission to "
    "IJHPM. The study used publicly available aggregate data only; ethics "
    "committee approval was not required under the Japanese Ethical "
    "Guidelines for Medical and Biological Research Involving Human Subjects "
    "(2021 revision). The reporting follows the STROBE checklist for "
    "cross-sectional studies, which is uploaded as supplementary material.")

add_para(
    "Statistical analysis: the lead author, who has training in clinical "
    "epidemiology and statistics, performed the analyses and takes full "
    "responsibility for them. No professional statistician outside the "
    "author list was consulted. The contribution of artificial intelligence "
    "tooling (use of Devin / Cognition AI as a coding assistant for data "
    "processing, modelling and visualisation, with all scientific "
    "interpretation performed by the human authors) is declared in the "
    "Contributorship section of the End-Matter file.")

add_para(
    "The manuscript is within the IJHPM length requirements for an Original "
    "Research Article (main text approximately 3,400 words; structured "
    "abstract within 300 words; 3 tables, 2 figures; 30 references). We "
    "have uploaded the title page, anonymised main manuscript, End-Matter "
    "(Contributorship, Ethics, Funding, Conflict of interests, "
    "Acknowledgments, Data Availability) and STROBE checklist as separate "
    "files in accordance with the journal's submission requirements.")

add_para(
    "We believe our work will be of substantive interest to IJHPM's "
    "international audience of health policy researchers and decision-makers, "
    "and we look forward to the reviewers' comments. Please do not hesitate "
    "to contact me if any further information is required.")

doc.add_paragraph()
add_para("Yours sincerely,")
doc.add_paragraph()
add_para("[Signature]")
add_para("[Corresponding author name], on behalf of all authors", italic=True)

out = os.path.join(OUTPUT_DIR, 'cover_letter_IJHPM_EN.docx')
doc.save(out)
print(f"Saved: {out}")

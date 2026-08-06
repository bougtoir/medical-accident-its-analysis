#!/usr/bin/env python3
"""Convert in-text mathematical/statistical expressions to Word OMML equations.

This post-processor scans a python-docx Document and replaces plain-text
expressions such as 'β = 64.1, 95% CI 54.8 to 73.4, P <0.001' with native
Word equation objects (<m:oMath>) so that the manuscript can be copy-edited
in Microsoft Word rather than relying on Unicode symbols or LaTeX markup.
"""
import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Regex for the statistical expressions that should become Word equations.
# Ordered from most specific to least specific so longer expressions take
# precedence over their subcomponents.
# ---------------------------------------------------------------------------
MATH_RE = re.compile(
    r"""
    (β\s*=\s*[+\-]?\d+\.\d+
        (?:\s*,\s*95%\s*(?:CI|confidence\ interval)\s*[+\-]?\d+\.\d+\s*to\s*[+\-]?\d+\.\d+)?
        (?:\s*,\s*P\s*(?:[<>=≤≥]\s*)?\d+\.\d+)?)
    |
    (r\s*=\s*[+\-]?\d+\.\d+\s*,\s*P\s*(?:[<>=≤≥]\s*)?\d+\.\d+)
    |
    (t\s*=\s*[+\-]?\d+\.\d+\s*,\s*P\s*(?:[<>=≤≥]\s*)?\d+\.\d+)
    |
    (Cohen's\ d\s*(?:=\s*)?[+\-]?\d+\.\d+)
    |
    ((?:[Mm]arginal)\s+R²\s*[+\-]?\d+\.\d+)
    |
    (P\s*(?:[<>=≤≥]\s*)?\d+\.\d+)
    |
    (R²)
    |
    (β)
    """,
    re.VERBOSE,
)


def _append_run(math, text):
    """Append an <m:r><m:t> element to a math object."""
    if text == "":
        return
    r = OxmlElement("m:r")
    t = OxmlElement("m:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    math.append(r)


def _append_r_squared(math, trailing=""):
    """Append an R-squared superscript to a math object."""
    ss = OxmlElement("m:sSup")
    e = OxmlElement("m:e")
    _append_run(e, "R")
    sup = OxmlElement("m:sup")
    _append_run(sup, "2")
    ss.append(e)
    ss.append(sup)
    math.append(ss)
    if trailing:
        _append_run(math, trailing)


def _expr_to_omml(expr):
    """Convert a plain-text math expression to an <m:oMath> element."""
    math = OxmlElement("m:oMath")
    token_re = re.compile(
        r"""
        (?:Cohen's\ d\s*)
        |(?:95%\ confidence\ interval\s*)
        |(?:95%\ CI\s*)
        |(?:intraclass\ correlation\ coefficient\s*)
        |(?:R²\s*)
        |(?:β\s*)
        |(?<![A-Za-z0-9_])([rRtP])(?![A-Za-z_])\s*
        |([+\-]?\d+\.\d+)\s*
        |([+\-]?\d+)(?![\d\.])\s*
        |([=<>≤≥])\s*
        |([,;()])\s*
        |(\S+)\s*
        """,
        re.VERBOSE,
    )
    for m in token_re.finditer(expr):
        tok = m.group(0)
        if not tok:
            continue
        stripped = tok.rstrip()
        trailing = tok[len(stripped) :] if len(tok) > len(stripped) else ""
        if stripped == "R²":
            _append_r_squared(math, trailing)
        elif stripped == "β":
            _append_run(math, "β" + trailing)
        else:
            _append_run(math, stripped + trailing)
    return math


def _copy_run_format(src, dst):
    dst.bold = src.bold
    dst.italic = src.italic
    dst.underline = src.underline
    if src.font.name:
        dst.font.name = src.font.name
    if src.font.size:
        dst.font.size = src.font.size
    dst.font.superscript = src.font.superscript


def _process_run(paragraph, run):
    """Recursively split a run and replace math expressions with OMML."""
    if run.font.superscript:
        return
    while True:
        m = MATH_RE.search(run.text)
        if not m:
            break
        before = run.text[: m.start()]
        expr = m.group(0)
        after = run.text[m.end() :]

        run.text = before
        run_after = paragraph.add_run(after)
        _copy_run_format(run, run_after)

        omml = _expr_to_omml(expr)
        run._r.addnext(omml)
        omml.addnext(run_after._r)

        run = run_after


def _iter_paragraphs(obj):
    """Yield all paragraphs from a document, table cells and nested tables."""
    for p in obj.paragraphs:
        yield p
    for table in obj.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_paragraphs(cell)


def convert_docx_math_to_omml(doc):
    """Convert in-text math expressions in a python-docx Document in place."""
    for paragraph in list(_iter_paragraphs(doc)):
        # Snapshot runs because the loop modifies the paragraph structure.
        for run in list(paragraph.runs):
            _process_run(paragraph, run)

#!/usr/bin/env python3
"""Create STROBE checklist (cross-sectional studies) for IJHPM submission."""
import json
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'documents', 'IJHPM')
os.makedirs(OUTPUT_DIR, exist_ok=True)

REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
with open(os.path.join(REPO_ROOT, 'output', 'ijhpm_results.json')) as f:
    R = json.load(f)

meta = R['metadata']
n_areas = meta['n_areas']
n_prefectures = meta['n_prefectures']
n_univ_areas = meta['n_univ_areas']
fiscal_year = meta.get('fiscal_year', 2022)

doc = Document()
for s in doc.sections:
    s.page_width = Cm(29.7)
    s.page_height = Cm(21)
    s.top_margin = Cm(1.8)
    s.bottom_margin = Cm(1.8)
    s.left_margin = Cm(1.8)
    s.right_margin = Cm(1.8)

st = doc.styles['Normal']
st.font.name = 'Times New Roman'
st.font.size = Pt(10)
st.paragraph_format.line_spacing = 1.15
st.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("STROBE Statement — Checklist of items that should be included "
              "in reports of cross-sectional studies")
r.bold = True
r.font.size = Pt(13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Manuscript: Regional variation in anaesthesia practice in Japan: "
              "structural determinant or claims-audit artefact?")
r.italic = True
r.font.size = Pt(11)

doc.add_paragraph()

items = [
    ("Title and abstract", "1(a)",
     "Indicate the study's design with a commonly used term in the title or "
     "the abstract",
     "Title (running header: 'Anaesthesia variation in Japan'); Abstract "
     "(Methods: 'cross-sectional ecological study')"),
    ("", "1(b)",
     "Provide in the abstract an informative and balanced summary of what "
     "was done and what was found",
     "Abstract: structured subheadings Background / Methods / Results / "
     "Conclusion"),

    ("Introduction", "", "", ""),
    ("Background/rationale", "2",
     "Explain the scientific background and rationale for the investigation "
     "being reported",
     "Introduction, paragraphs 1–2"),
    ("Objectives", "3",
     "State specific objectives, including any pre-specified hypotheses",
     "Introduction, final paragraph (four pre-specified research questions)"),

    ("Methods", "", "", ""),
    ("Study design", "4",
     "Present key elements of study design early in the paper",
     "Methods, 'Study design and ethical considerations'"),
    ("Setting", "5",
     "Describe the setting, locations, and relevant dates, including "
     "periods of recruitment, exposure, follow-up, and data collection",
     f"Methods, 'Geographic units' and 'Data sources' (fiscal year {fiscal_year}, {n_areas} "
     f"secondary medical areas covering all {n_prefectures} prefectures of Japan)"),
    ("Participants", "6",
     "Give the eligibility criteria, and the sources and methods of "
     "selection of participants",
     "Methods, 'Data sources'; aggregate ecological unit = secondary "
     "medical area; no individual-level eligibility criteria"),
    ("Variables", "7",
     "Clearly define all outcomes, exposures, predictors, potential "
     "confounders, and effect modifiers. Give diagnostic criteria, if "
     "applicable",
     "Methods, 'Anaesthesia procedure codes and university hospital mapping' "
     "and 'Statistical analysis'"),
    ("Data sources/measurement", "8*",
     "For each variable of interest, give sources of data and details of "
     "methods of assessment (measurement). Describe comparability of "
     "assessment methods if there is more than one group",
     "Methods, 'Data sources and standardised claim ratios' (Cabinet Office "
     "Regional Variation Visualisation; Survey of Physicians, Dentists and "
     "Pharmacists; National Land Numerical Information)"),
    ("Bias", "9",
     "Describe any efforts to address potential sources of bias",
     "Methods, 'Sensitivity analyses against the audit hypothesis' and "
     "Discussion, 'Strengths and limitations'"),
    ("Study size", "10",
     "Explain how the study size was arrived at",
     f"Methods, 'Data sources' (all {n_areas} areas — no sampling)"),
    ("Quantitative variables", "11",
     "Explain how quantitative variables were handled in the analyses. If "
     "applicable, describe which groupings were chosen and why",
     "Methods, 'Statistical analysis' (standardised claim ratio as primary "
     "outcome; empirical Bayes shrinkage)"),
    ("Statistical methods", "12(a)",
     "Describe all statistical methods, including those used to control for "
     "confounding",
     "Methods, 'Statistical analysis' (descriptive statistics, multilevel "
     "models with prefecture random intercept, variance decomposition, "
     "empirical Bayes shrinkage)"),
    ("", "12(b)",
     "Describe any methods used to examine subgroups and interactions",
     "Methods, 'Sensitivity analyses' (within-prefecture paired comparisons; "
     "alternative combined outcomes)"),
    ("", "12(c)",
     "Explain how missing data were addressed",
     "Methods, 'Statistical analysis' (low-volume areas masked by the data "
     "provider are reported as n and excluded from rate calculations)"),
    ("", "12(d)",
     "Describe analytical methods taking account of sampling strategy",
     "Not applicable — full population of areas"),
    ("", "12(e)",
     "Describe any sensitivity analyses",
     "Methods, 'Sensitivity analyses against the audit hypothesis'"),

    ("Results", "", "", ""),
    ("Participants", "13(a)",
     "Report numbers of individuals at each stage of study",
     f"Results, 'Study population' ({n_areas} areas, {n_prefectures} prefectures, {n_univ_areas} "
     f"university-hospital areas)"),
    ("", "13(b)",
     "Give reasons for non-participation at each stage",
     "Results, 'Study population' (areas masked for low volume reported "
     "as n by procedure)"),
    ("", "13(c)",
     "Consider use of a flow diagram",
     "Not applicable — ecological design with full population"),
    ("Descriptive data", "14(a)",
     "Give characteristics of study participants and information on "
     "exposures and potential confounders",
     "Results, 'Study population and variation'; Table 1 and Table 2"),
    ("", "14(b)",
     "Indicate number of participants with missing data for each variable "
     "of interest",
     "Table 2 column 'n'"),
    ("Outcome data", "15*",
     "Report numbers of outcome events or summary measures",
     "Results, 'Study population and variation' and Tables 1–2 (means, "
     "SD, IQR, range, coefficients of variation)"),
    ("Main results", "16(a)",
     "Give unadjusted estimates and, if applicable, confounder-adjusted "
     "estimates and their precision (e.g. 95% CI). Make clear which "
     "confounders were adjusted for and why they were included",
     "Results, 'Multilevel model'; Table 3 (mixed-effects coefficients "
     "with 95% confidence intervals)"),
    ("", "16(b)",
     "Report category boundaries when continuous variables were "
     "categorized",
     "Figure 1 and Figure 2 legends (quintile shading)"),
    ("", "16(c)",
     "If relevant, consider translating estimates of relative risk into "
     "absolute risk for a meaningful time period",
     "Not applicable — ecological measure is a standardised claim ratio"),
    ("Other analyses", "17",
     "Report other analyses done — e.g. analyses of subgroups and "
     "interactions, and sensitivity analyses",
     "Results, 'Sensitivity analyses' and 'Robustness' (empirical Bayes "
     "shrinkage; within-prefecture comparisons; combined codes)"),

    ("Discussion", "", "", ""),
    ("Key results", "18",
     "Summarise key results with reference to study objectives",
     "Discussion, 'Statement of principal findings'"),
    ("Limitations", "19",
     "Discuss limitations of the study, taking into account sources of "
     "potential bias or imprecision. Discuss both direction and magnitude "
     "of any potential bias",
     "Discussion, 'Strengths and limitations'"),
    ("Interpretation", "20",
     "Give a cautious overall interpretation of results considering "
     "objectives, limitations, multiplicity of analyses, results from "
     "similar studies, and other relevant evidence",
     "Discussion, 'Interpretation within the context of the wider "
     "literature'"),
    ("Generalisability", "21",
     "Discuss the generalisability (external validity) of the study results",
     "Discussion, 'Interpretation within the context of the wider literature' "
     "and 'Implications for policy, practice and governance'"),

    ("Other information", "", "", ""),
    ("Funding", "22",
     "Give the source of funding and the role of the funders for the "
     "present study and, if applicable, for the original study on which the "
     "present article is based",
     "End Matter, 'Funding'"),
]

table = doc.add_table(rows=1 + len(items), cols=4)
table.style = 'Table Grid'
hdrs = ["Section / Topic", "Item #", "Recommendation", "Reported on page / "
        "section"]
widths = [Cm(4.0), Cm(1.5), Cm(11), Cm(9)]
for i, h in enumerate(hdrs):
    c = table.rows[0].cells[i]
    c.width = widths[i]
    c.text = ''
    p = c.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(10)

for ri, (sect, num, rec, addr) in enumerate(items, 1):
    row = table.rows[ri].cells
    is_section_header = (num == "" and rec == "" and addr == "")
    vals = [sect, num, rec, addr]
    for ci, val in enumerate(vals):
        row[ci].width = widths[ci]
        row[ci].text = ''
        p = row[ci].paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(9)
        if is_section_header and ci == 0:
            r.bold = True
            r.italic = True

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run(
    "*Give information separately for cases and controls in case-control "
    "studies and, if applicable, for exposed and unexposed groups in "
    "cohort and cross-sectional studies.")
r.italic = True
r.font.size = Pt(9)

p = doc.add_paragraph()
r = p.add_run(
    "Note: An Explanation and Elaboration article discusses each checklist "
    "item and gives methodological background and published examples of "
    "transparent reporting. The STROBE checklist is best used in "
    "conjunction with this article (freely available on the Web sites of "
    "PLoS Medicine, Annals of Internal Medicine, and Epidemiology). "
    "Information on the STROBE Initiative is available at "
    "www.strobe-statement.org.")
r.italic = True
r.font.size = Pt(9)

out = os.path.join(OUTPUT_DIR, 'STROBE_checklist_IJHPM_EN.docx')
doc.save(out)
print(f"Saved: {out}")

#!/usr/bin/env python3
"""Create separate editable tables .docx for Health Policy submission.

Reads output/ijhpm_results.json produced by compile_ijhpm_results.py.
"""
import json
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'documents', 'Health_Policy')
os.makedirs(OUTPUT_DIR, exist_ok=True)

REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
with open(os.path.join(REPO_ROOT, 'output', 'ijhpm_results.json')) as f:
    R = json.load(f)


def fmt(x, d=1):
    if x is None or (isinstance(x, float) and x != x):
        return 'NA'
    return f'{x:.{d}f}'


def fmt_p(x):
    if x is None or (isinstance(x, float) and x != x):
        return 'NA'
    if x < 0.001:
        return '<0.001'
    return f'{x:.3f}'


def setup(doc):
    for s in doc.sections:
        s.page_width = Cm(29.7)
        s.page_height = Cm(21)
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2)
        s.right_margin = Cm(2)
    st = doc.styles['Normal']
    st.font.name = 'Times New Roman'
    st.font.size = Pt(11)
    st.paragraph_format.line_spacing = 1.15
    st.paragraph_format.space_after = Pt(0)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)


def add_footnote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.italic = True
    r.font.name = 'Times New Roman'


def make_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = 'Times New Roman'
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            c = t.rows[ri].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(11)
            r.font.name = 'Times New Roman'
    return t


doc = Document()
setup(doc)

n_areas = R['metadata']['n_areas']
fiscal_year = R['metadata'].get('fiscal_year', 2022)

# Table 1: distribution
rows2 = []
for code in ['L008', 'L004', 'L002', 'L003']:
    o = R['codes'][code]['overall']
    rows2.append([
        code,
        str(o['n']),
        f"{fmt(o['mean'])} ({fmt(o['sd'])})",
        f"{fmt(o['median'])} ({fmt(o['q1'])}–{fmt(o['q3'])})",
        fmt(o['min']),
        fmt(o['max']),
        fmt(o['cv']),
    ])

add_title(doc,
          f"Table 1. Distribution of standardised claim ratios across {n_areas} "
          f"secondary medical areas (national average = 100), fiscal year {fiscal_year}")
make_table(doc,
           ["Code", "n", "Mean (SD)", "Median (IQR)", "Min", "Max", "CV (%)"],
           rows2)
add_footnote(doc,
             "SD, standard deviation; IQR, interquartile range; CV, coefficient of "
             "variation. L008 = closed-circuit general anaesthesia; L002 = epidural "
             "anaesthesia as main technique; L003 = continuous epidural infusion, "
             "largely billed as an adjunct to general anaesthesia; L004 = spinal "
             "anaesthesia. Areas with low claim volume are masked by the data provider "
             "and appear as missing values; the n column reflects the number of "
             "secondary medical areas with non-missing values for each code.")

# Table 2: multilevel model
rows3 = []
for code in ['L008', 'L004', 'L002', 'L003']:
    m = R['codes'][code]['multilevel']
    rows3.append([
        code,
        fmt(m['icc_null'], 3),
        f"+{fmt(m['coef_univ'])} ({fmt(m['ci_low'])} to {fmt(m['ci_high'])})",
        fmt_p(m['p']),
        fmt(m['marginal_r2'], 3),
    ])

add_title(doc,
          "Table 2. Multilevel linear mixed model results: standardised claim "
          "ratio as outcome, prefecture as random intercept")
make_table(doc,
           ["Code", "Null model ICC", "β university (95% CI)", "P value",
            "Proportional reduction in total variance"],
           rows3)
add_footnote(doc,
             "β university, fixed effect coefficient for university hospital "
             "presence (binary 0/1); CI, confidence interval; ICC, intraclass "
             "correlation coefficient; proportional reduction in total variance, "
             "decrease in total (prefecture + residual) variance after adding the university "
             "hospital fixed effect.")

out_en = os.path.join(OUTPUT_DIR, 'regional_anaesthesia_tables_Health_Policy_EN.docx')
doc.save(out_en)
print(f"Saved: {out_en}")

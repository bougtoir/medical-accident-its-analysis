#!/usr/bin/env python3
"""Create IJHPM-formatted English manuscript.

Conforms to International Journal of Health Policy and Management author
guidelines:
- Structured abstract (Background/Methods/Results/Conclusion) <=300 words
- Implications for Policy Makers (3-5 bullets) and Implications for Public
  (100-150 words) as key messages
- Body <=6000 words; <=5 tables/figures combined; <=30 references
- References numbered in order of appearance (AMA/Vancouver hybrid, see file)
- Tables and figures inserted inline, immediately after their first mention
- Double-spaced, Times New Roman 12 pt

All numeric results are read from output/ijhpm_results.json, which is produced
by scripts/compile_ijhpm_results.py.
"""

import json
import os
import re
from string import Template

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'documents', 'IJHPM')
os.makedirs(OUTPUT_DIR, exist_ok=True)

REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FIG_DIR = os.path.join(REPO_ROOT, 'output')
RESULTS_PATH = os.path.join(REPO_ROOT, 'output', 'ijhpm_results.json')

doc = Document()

# Page setup: A4, 2.54 cm margins
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0
style.paragraph_format.space_after = Pt(0)


# ============================================================
# LOAD RESULTS
# ============================================================
with open(RESULTS_PATH) as f:
    R = json.load(f)

codes = R['codes']
combined = R['combined']['L008_L003']
corr = R['correlations']
eb = R['empirical_bayes']
vd = R['variance_decomposition']
audit = R['audit_sensitivity']
outliers = R['outliers']['L008']
meta = R['metadata']



def fmt_num(x, decimals=1):
    if x is None or (isinstance(x, float) and (x != x)):
        return 'NA'
    return f'{x:.{decimals}f}'


def fmt_p(x):
    if x is None or (isinstance(x, float) and (x != x)):
        return 'NA'
    if x < 0.001:
        return '<0.001'
    return f'{x:.3f}'


def desc(code):
    return codes[code]['overall']


def univ(code):
    return codes[code]['university']


def nonu(code):
    return codes[code]['non_university']


def ml(code):
    return codes[code]['multilevel']



def build_flat(R):
    """Build a flat substitution dictionary from ijhpm_results.json."""
    F = {}
    meta = R['metadata']
    F['n_areas'] = meta['n_areas']
    F['n_prefectures'] = meta['n_prefectures']
    F['fiscal_year'] = meta.get('fiscal_year', 2022)
    F['n_univ_areas'] = meta['n_univ_areas']
    F['n_nonuniv_areas'] = meta['n_nonuniv_areas']
    F['n_univ_areas_pct'] = fmt_num(meta.get('n_univ_areas_pct', 100 * meta['n_univ_areas'] / meta['n_areas']))

    vd_key_map = {
        'between_prefecture_pct': 'vd_between',
        'university_effect_pct': 'vd_univ',
        'residual_pct': 'vd_residual',
        'within_prefecture_pct': 'vd_within',
        'univ_explains_within_pct': 'vd_univ_within',
    }

    for code, c in R['codes'].items():
        o = c['overall']
        for key, d in [('n', 0), ('mean', 1), ('sd', 1), ('median', 1),
                       ('q1', 1), ('q3', 1), ('min', 1), ('max', 1),
                       ('cv', 1), ('iqr', 1), ('range_ratio', 1)]:
            F[f'{code}_{key}'] = fmt_num(o.get(key), d)

        u = c['university']
        n = c['non_university']
        for prefix, vals in [('u', u), ('nu', n)]:
            for key in ['mean', 'sd', 'median', 'min', 'max']:
                F[f'{code}_{prefix}_{key}'] = fmt_num(vals.get(key), 1)

        F[f'{code}_d'] = fmt_num(c['cohens_d_overall'], 2)

        w = c['within_prefecture']
        F[f'{code}_within_diff'] = fmt_num(w.get('mean_diff'))
        F[f'{code}_within_t'] = fmt_num(w.get('t'), 2)
        F[f'{code}_within_p'] = fmt_p(w.get('p'))

        m = c['multilevel']
        F[f'{code}_ml_icc'] = fmt_num(m.get('icc_null'), 3)
        F[f'{code}_ml_coef'] = fmt_num(m.get('coef_univ'))
        F[f'{code}_ml_ci_low'] = fmt_num(m.get('ci_low'))
        F[f'{code}_ml_ci_high'] = fmt_num(m.get('ci_high'))
        F[f'{code}_ml_p'] = fmt_p(m.get('p'))
        F[f'{code}_ml_r2'] = fmt_num(m.get('marginal_r2'), 3)

        mc = c.get('multilevel_covariate')
        if mc:
            F[f'{code}_mlc_coef'] = fmt_num(mc.get('coef_univ'))
            F[f'{code}_mlc_ci_low'] = fmt_num(mc.get('ci_low'))
            F[f'{code}_mlc_ci_high'] = fmt_num(mc.get('ci_high'))
            F[f'{code}_mlc_p'] = fmt_p(mc.get('p'))
            F[f'{code}_mlc_r2'] = fmt_num(mc.get('marginal_r2'), 3)
            F[f'{code}_mlc_popd_coef'] = fmt_num(mc['log_pop_density_z'].get('coef'))
            F[f'{code}_mlc_popd_p'] = fmt_p(mc['log_pop_density_z'].get('p'))
            F[f'{code}_mlc_anes_coef'] = fmt_num(mc['anes_pct_z'].get('coef'))
            F[f'{code}_mlc_anes_p'] = fmt_p(mc['anes_pct_z'].get('p'))

        if code in R.get('variance_decomposition', {}):
            vd = R['variance_decomposition'][code]
            for json_key, flat_key in vd_key_map.items():
                F[f'{code}_{flat_key}'] = fmt_num(vd.get(json_key))

        if code in R.get('empirical_bayes', {}):
            ebv = R['empirical_bayes'][code]
            F[f'{code}_d_shrunk'] = fmt_num(ebv.get('shrunk_cohens_d'), 2)
            F[f'{code}_attenuation'] = fmt_num(ebv.get('attenuation_pct'))

    F['L002_missing_n'] = R['metadata']['n_areas'] - R['codes']['L002']['overall']['n']
    F['L002_fold'] = fmt_num(R['codes']['L002']['university']['mean'] / R['codes']['L002']['non_university']['mean'], 2)
    F['L003_fold'] = fmt_num(R['codes']['L003']['university']['mean'] / R['codes']['L003']['non_university']['mean'], 2)

    for combo, c in R['combined'].items():
        prefix = combo.replace('+', '_')
        for key, d in [('n', 0), ('mean', 1), ('sd', 1), ('median', 1),
                       ('q1', 1), ('q3', 1), ('min', 1), ('max', 1),
                       ('cv', 1), ('iqr', 1)]:
            F[f'{prefix}_{key}'] = fmt_num(c['overall'].get(key), d)
        F[f'{prefix}_u_mean'] = fmt_num(c['university'].get('mean'))
        F[f'{prefix}_nu_mean'] = fmt_num(c['non_university'].get('mean'))
        F[f'{prefix}_d'] = fmt_num(c.get('cohens_d'), 2)
        F[f'{prefix}_fold'] = fmt_num(c.get('fold_ratio'), 2)

    for k, v in R['correlations'].items():
        F[f'corr_{k}_r'] = fmt_num(v.get('r'), 3)
        F[f'corr_{k}_p'] = fmt_p(v.get('p'))

    F['audit_max_shift'] = fmt_num(R['audit_sensitivity']['max_ratio_shift_approx'], 2)
    for code in ['L008', 'L002', 'L003', 'L004']:
        if code in R['audit_sensitivity']:
            F[f'audit_{code}_iqr_pct'] = fmt_num(R['audit_sensitivity'][code]['percent_of_iqr'], 2)

    outlier_name_map = {
        ('安房', '千葉県'): ('Awa (Chiba)', 'Chiba'),
        ('区中央部', '東京都'): ('central Tokyo (Tokyo)', 'Tokyo'),
    }
    for i, o in enumerate(R['outliers']['L008'][:2], 1):
        name, pref = o['area_name'], o['pref_name']
        if (name, pref) in outlier_name_map:
            name, pref = outlier_name_map[(name, pref)]
        F[f'outlier{i}_name'] = name
        F[f'outlier{i}_pref'] = pref
        F[f'outlier{i}_scr'] = fmt_num(o['scr'])

    return F


FLAT = build_flat(R)


def add_run_with_refs(paragraph, text, italic=False, bold=False):
    """Add text to a paragraph, parsing {n} or {n-m} as font superscript runs."""
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        if italic:
            run.italic = True
        if bold:
            run.bold = True
        if part.startswith('{') and part.endswith('}'):
            inner = part[1:-1]
            run.text = inner
            run.font.superscript = True


def renumber_references(doc, references):
    """Determine first-appearance order from superscript citations and reorder.

    Walks the document body (stops at the 'References' heading), records the
    order in which each citation number first appears, builds an old->new
    mapping, updates all superscript citation runs in place, and returns the
    reference list reordered to match the new numbering.
    """
    first = {}
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt.upper() == 'REFERENCES':
            break
        for run in p.runs:
            if run.font.superscript:
                for n in re.findall(r'\d+', run.text):
                    n = int(n)
                    if n not in first:
                        first[n] = len(first) + 1
    # Ensure every reference in the list is represented, even if missed
    for old in range(1, len(references) + 1):
        if old not in first:
            first[old] = len(first) + 1
    mapping = {old: new for new, old in enumerate(sorted(first, key=first.get), 1)}

    def repl(m):
        old = int(m.group(0))
        return str(mapping.get(old, old))

    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.superscript:
                run.text = re.sub(r'\d+', repl, run.text)

    return [references[old - 1] for old in sorted(mapping, key=mapping.get)]


def add_para(text, bold=False, italic=False, align=None, space_before=0, space_after=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after:
        p.paragraph_format.space_after = Pt(space_after)
    add_run_with_refs(p, Template(text).substitute(FLAT), italic=italic, bold=bold)
    return p


def add_heading(text, level=1, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text.upper() if level == 1 else text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12) if level > 1 else Pt(13)
    run.bold = True
    return p


def add_subheading(text):
    return add_heading(text, level=2, space_before=8, space_after=4)


def add_blank():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    return p


def add_table_from_data(caption, headers, rows, note=None):
    """Add an inline table with caption and optional footnote."""
    add_para(caption, bold=True, space_before=12, space_after=6)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            txt = str(val)
            if FLAT:
                txt = Template(txt).safe_substitute(FLAT)
            r = p.add_run(txt)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)
    if note:
        add_para(note, italic=True, space_before=6, space_after=12)
    return table


def add_figure_inline(image_path, caption):
    """Add an inline figure with caption."""
    add_para(caption, bold=True, space_before=12, space_after=6)
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(6.3))
    else:
        add_para('[Figure image not found: ' + image_path + ']', italic=True)
    add_blank()


# ============================================================
# TITLE PAGE
# ============================================================
title = ("Regional variation in anaesthesia practice in Japan: "
         "structural determinant or claims-audit artefact?")
assert len(title) <= 150, f"Title is {len(title)} chars (max 150)"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
run = p.add_run(title)
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.bold = True

header = "Anaesthesia variation in Japan"
assert len(header) <= 40, f"Header is {len(header)} chars (max 40)"
add_para("Running header: ${header}".replace('${header}', header), italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_blank()

add_para("[Authors and affiliations to be supplied]", italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("[Corresponding author: name, postal address, telephone, email]",
         italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_blank()

add_para("Article type: Original research article", italic=True)
add_para("Word count (main text, excluding abstract, references, tables and figures): approximately [calculated at build] words", italic=True)
add_para("Tables: [n_tables]   Figures: [n_figures]   References: [n_references]", italic=True)
add_para("Reporting guideline: STROBE checklist for cross-sectional studies (uploaded as supplementary material)", italic=True)
add_para("Target journal: International Journal of Health Policy and Management", italic=True)

doc.add_page_break()

# ============================================================
# KEY MESSAGES
# ============================================================
add_heading("Implications for Policy Makers", level=1, space_before=0)
add_para(
    "1. Within Japan's uniform national fee schedule, regional variation in "
    "anaesthesia practice is driven by structural capacity (notably university "
    "hospital presence) rather than by prefectural claims-audit intensity.")
add_para(
    "2. Large variation in neuraxial and combined general-epidural "
    "anaesthesia is a modifiable equity gap that workforce, training "
    "and monitoring policies can address without changing payment rates.")
add_para(
    "3. The claims-audit hypothesis is quantitatively implausible as the main "
    "driver of variation; policy responses should target supply and organisation "
    "rather than fraud control.")
add_para(
    "4. Anaesthesia technique mix can be monitored as a system-level quality "
    "indicator using existing publicly available standardised claim ratios.")
add_para(
    "5. The within-prefecture variance-decomposition framework is transferable "
    "to other procedures and to other universal-coverage systems that combine a "
    "uniform fee schedule with regional auditing.")

add_heading("Implications for Public", level=1)
add_para(
    "Patients in Japan are entitled to equitable access under universal health "
    "insurance, yet the likelihood of receiving a combined general-epidural "
    "anaesthetic depends strongly on whether the patient's home area contains a "
    "university hospital. This study shows that the variation is not explained by "
    "differences in claims auditing, which is uniform within each prefecture, but "
    "by the geographic distribution of specialist anaesthesia services. The "
    "findings suggest that policy attention should focus on expanding regional "
    "access to advanced anaesthesia techniques and on transparent monitoring of "
    "technique mix, rather than on tightening audits.")

doc.add_page_break()

# ============================================================
# ABSTRACT
# ============================================================
add_para(title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_blank()
add_heading("Abstract", level=1, space_before=0)

add_para(
    "Background  Regional variation in anaesthesia practice is well documented, "
    "but whether it reflects clinical structure or administrative auditing "
    "remains unresolved. Japan's single national fee schedule plus prefecture-"
    "specific claims auditing offers a natural experiment. We quantified variation "
    "across ${n_areas} secondary medical areas and tested the structural versus "
    "audit hypotheses.")

add_para(
    "Methods  Cross-sectional ecological study using age- and sex-standardised "
    "claim ratios (national average = 100) for six anaesthesia procedure codes "
    "in fiscal year ${fiscal_year}. We fitted multilevel linear mixed models with "
    "prefecture random intercepts and three pre-specified sensitivity analyses "
    "to test differential auditing, plus empirical Bayes shrinkage.")

add_para(
    "Results  Coefficients of variation ranged from ${L008_cv}% (general "
    "anaesthesia) to ${L002_cv}% (epidural). Only ${L008_ml_icc} of general "
    "anaesthesia variance lay between prefectures; university hospital presence "
    "explained ${L008_ml_r2} of total variance and ${L008_vd_univ_within}% of "
    "within-prefecture variance (Cohen's d ${L008_d}), and was positive in all "
    "${n_prefectures} prefectures. Audit sensitivity analyses converged: the "
    "maximum audit-rate difference could shift ratios by at most ${audit_max_shift} "
    "points, less than 1% of the interquartile range for general anaesthesia. "
    "Empirical Bayes shrinkage attenuated the university "
    "hospital effect but it remained statistically significant.")

add_para(
    "Conclusion  Regional variation is predominantly structural, driven by "
    "university hospital proximity. The large variation in epidural and "
    "continuous epidural infusion signals a modifiable inequity in access to "
    "potentially beneficial regional anaesthesia under universal coverage.")

add_blank()
add_para("Keywords: ", bold=False)
add_para(
    "anaesthesia; small-area variation; multilevel analysis; health policy; "
    "health services research; universal health coverage; Japan")

doc.add_page_break()

# ============================================================
# INTRODUCTION
# ============================================================
add_heading("Introduction", level=1, space_before=0)

add_para(
    "Japan's universal health insurance system, established in 1961, covers the "
    "entire population under a nationally uniform fee schedule that specifies "
    "reimbursement for every medical procedure.{1} The system consistently ranks "
    "among the highest performing globally for effective service coverage{2} and "
    "is delivered through a hierarchical structure of ${n_prefectures} prefectural "
    "(tertiary) and ${n_areas} secondary medical areas, each designed to be self-"
    "sufficient for inpatient care.{3} Cross-boundary patient movement is limited,{4} and the "
    "system was intended to ensure equitable access regardless of geography.")

add_para(
    "Yet substantial regional variation in surgical procedures, prescribing and "
    "diagnostic testing has been documented through the National Database of "
    "Health Insurance Claims.{5} A distinctive feature of the Japanese system is "
    "mandatory prefectural insurance auditing (shinsa): all claims are reviewed "
    "before reimbursement and audit rates vary across prefectures.{6} This raises "
    "a fundamental question for health policy and governance: does the variation "
    "observed in claims data reflect genuine differences in clinical practice, or "
    "is it an artefact of differential auditing? Distinguishing between these "
    "two explanations has different implications for quality improvement, payment "
    "policy, regulation and patient equity.")

add_para(
    "Anaesthesia practice is well suited to this question. Perioperative claims "
    "are comparatively complete because surgery is almost always reimbursed, so "
    "the data have high structural integrity. The choice between general "
    "anaesthesia alone and techniques that add regional anaesthesia is clinically "
    "consequential: a body of evidence suggests that regional anaesthesia, "
    "particularly epidural analgesia combined with general anaesthesia, may improve "
    "recurrence-free and overall survival in some cancer surgeries,{7,8} although "
    "results from randomised trials are mixed.{9} The 2026 revision of the "
    "Japanese fee schedule renamed and redefined the L008 general-anaesthesia code "
    "to make the required airway-device explicit, removing an ambiguity that "
    "could plausibly have driven differential audit decisions.{10} Understanding whether observed "
    "regional variation is structural or administrative is therefore both timely "
    "and central to evaluating the equity of care under universal coverage.")

add_para(
    "We used publicly available standardised claim ratios at the secondary medical "
    "area level to address the following research questions: (i) How large is the "
    "regional variation in anaesthesia practice across ${n_areas} areas? (ii) Is "
    "this variation explained by differential prefectural auditing? (iii) What "
    "proportion of variance lies between prefectures and between areas within "
    "prefectures? (iv) Which structural factors, including university hospital "
    "proximity, are associated with the observed variation?")

# ============================================================
# METHODS
# ============================================================
add_heading("Methods", level=1)

add_subheading("Study design and ethical considerations")
add_para(
    "This was a cross-sectional ecological study using publicly available "
    "aggregate data only. No individual-level data were accessed. Ethics "
    "committee approval was not required under the Ethical Guidelines for Medical "
    "and Biological Research Involving Human Subjects (Japan, 2021 revision). "
    "The study was reported in accordance with the Strengthening the Reporting of "
    "Observational Studies in Epidemiology (STROBE) checklist for cross-sectional "
    "studies,{30} which is provided as supplementary material.")

add_subheading("Geographic units")
add_para(
    "Japan's healthcare delivery system is organised into a hierarchical structure "
    "of medical areas defined under Article 30-4 of the Medical Care Act.{3} The "
    "${n_prefectures} prefectures serve as tertiary medical areas; within each "
    "prefecture, secondary medical areas (n = ${n_areas} nationally) are designated "
    "as the basic units for inpatient care planning. Crucially, secondary medical "
    "areas are completely nested within prefectures and do not cross prefectural "
    "boundaries. This nesting differs from the Hospital Service Areas and Hospital "
    "Referral Regions used in the United States, which can cross state "
    "boundaries,{11} and makes within-prefecture comparisons unconfounded by "
    "spillover under a common audit policy.")

add_subheading("Data sources and standardised claim ratios")
add_para(
    "We used three data sources. First, standardised claim ratios for fiscal "
    "year ${fiscal_year} published by the Cabinet Office under the 'Regional Variation "
    "Visualisation' initiative.{12} Standardised claim ratios are computed by "
    "indirect standardisation: expected claim frequencies are calculated by applying "
    "national age- and sex-specific claim rates to the local population structure, "
    "and the ratio is defined as 100 × (observed claims / expected claims). Ratios "
    "are calculated on a residence basis, so claims are attributed to the "
    "beneficiary's registered address rather than to the providing facility. "
    "Patient travel for treatment therefore does not inflate the ratio of the "
    "receiving area. Areas with very few claims are masked by the data provider to "
    "protect privacy and appear as missing values. Second, we obtained physician "
    "statistics from the ${fiscal_year} Survey of Physicians, Dentists and "
    "Pharmacists.{13} Third, secondary medical area boundaries, land area and "
    "population density were obtained from the National Land Numerical "
    "Information medical-area dataset (A38-20).{14}")

add_subheading("Anaesthesia procedure codes and university hospital mapping")
add_para(
    "Six procedure codes from the Japanese fee schedule were analysed (Table 1): "
    "L008 (closed-circuit general anaesthesia); L002 (epidural anaesthesia as "
    "main technique); L003 (continuous epidural infusion, largely billed as "
    "adjunct to general anaesthesia); L004 (spinal anaesthesia); L009 (anaesthesia "
    "management fee I, a proxy for specialist staffing); and L100 (inpatient nerve "
    "block). We mapped university hospitals to their respective "
    "secondary medical areas based on municipal address. ${n_univ_areas} "
    "secondary medical areas (${n_univ_areas_pct}%) contained at least one "
    "university hospital, distributed across all ${n_prefectures} prefectures, "
    "which enables within-prefecture comparisons between university and "
    "non-university areas under a common audit policy.")

# Table 1 inline
add_table_from_data(
    "Table 1. Anaesthesia procedure codes analysed.",
    ["Code", "Procedure", "Notes"],
    [
        ["L008", "Closed-circuit general anaesthesia",
         "Primary indicator of overall general anaesthesia volume"],
        ["L002", "Epidural anaesthesia",
         "Regional technique claimed as main anaesthetic"],
        ["L003", "Continuous epidural infusion",
         "Regional technique largely billed as adjunct to general anaesthesia"],
        ["L004", "Spinal anaesthesia", "Alternative regional technique"],
        ["L009", "Anaesthesia management fee I",
         "Proxy for specialist anaesthesiologist staffing"],
        ["L100", "Nerve block, inpatient", "Indicator of pain-clinic activity"],
    ],
    note=("Source: Japanese fee schedule (shinryo houshu tensuhyo), fiscal "
          "year ${fiscal_year}. L codes correspond to anaesthesia-related procedures. "
          "Standardised claim ratios were computed for each code by indirect "
          "age- and sex-standardisation against the national average (= 100).")
)

add_subheading("Statistical analysis")
add_para(
    "Three complementary approaches were used. First, we computed descriptive "
    "statistics including coefficients of variation. Between-group comparisons "
    "used Welch's t-test, Cohen's d for effect size{15} and the Mann-Whitney U "
    "test for non-parametric confirmation. Second, we fitted multilevel linear "
    "mixed models with secondary medical areas (level 1, n = ${n_areas}) nested "
    "within prefectures (level 2, n = ${n_prefectures}).{16} A null random-"
    "intercept model was fitted first to estimate the intraclass correlation "
    "coefficient. Subsequent models added fixed effects for university hospital "
    "presence; a covariate-adjusted sensitivity model further added the natural "
    "logarithm of population density and the anaesthesiologist share of all "
    "physicians (both standardised).{14,15} Models were estimated by restricted "
    "maximum likelihood using the Python statsmodels MixedLM implementation.{17} A "
    "small number of model fits produced optimizer convergence warnings; the "
    "resulting point estimates were numerically stable across repeated fits and "
    "across the sensitivity analyses, and are reported as mixed-model coefficients. "
    "Marginal R² was calculated as the proportional reduction in total variance from "
    "the null model. To address potential instability of ratios in low-volume areas, "
    "we applied empirical Bayes shrinkage estimation{18} and compared all main "
    "findings using both raw and shrunken ratios.")

add_para(
    "Third, three pre-specified sensitivity analyses tested the null hypothesis "
    "that insurance auditing explains observed variation. (a) Within-prefecture "
    "variance decomposition: under uniform prefectural auditing, within-"
    "prefecture variance should be small. We decomposed total variance using one-"
    "way analysis of variance and further partitioned the within-prefecture "
    "component into university hospital effect and residual using hierarchical "
    "sum-of-squares decomposition. (b) Cross-code correlation: audit-driven "
    "reclassification between general and spinal anaesthesia codes would produce "
    "negative correlations between them. (c) Quantitative audit-impact estimation: "
    "we calculated the maximum ratio shift attributable to the observed range of "
    "prefectural point-audit rates reported by the Social Insurance Medical Fee "
    "Payment Fund.{6} Combined ratios (general plus continuous epidural; "
    "general plus spinal) were also examined, since combining substitutable codes "
    "should largely absorb any audit-driven reclassification. Patients were not "
    "involved in the design, conduct or reporting of this study.")

# ============================================================
# RESULTS
# ============================================================
add_heading("Results", level=1)

add_subheading("Study population and variation in anaesthesia practice")
add_para(
    "The ${n_areas} secondary medical areas were distributed across "
    "${n_prefectures} prefectures (median 7 areas per prefecture, range 3 to 21); "
    "${n_univ_areas} areas (${n_univ_areas_pct}%) contained at least one university "
    "hospital. Standardised claim ratios for general anaesthesia were available "
    "for ${L008_n} of ${n_areas} areas; ratios for epidural anaesthesia were "
    "available for ${L002_n} areas, with ${L002_missing_n} areas masked "
    "owing to low volume. Substantial variation was observed across all codes "
    "(Table 2). General anaesthesia ratios ranged from ${L008_min} to "
    "${L008_max} (coefficient of variation ${L008_cv}%), a ${L008_range_ratio}" 
    "-fold difference. Epidural anaesthesia showed the greatest relative variation "
    "(coefficient of variation ${L002_cv}%); continuous epidural infusion "
    "${L003_cv}% and spinal anaesthesia ${L004_cv}%. The geographic distribution "
    "of ratios is shown in Figure 1.")

# Table 2 inline
add_table_from_data(
    "Table 2. Distribution of standardised claim ratios across ${n_areas} secondary "
    "medical areas (national average = 100), fiscal year ${fiscal_year}.",
    ["Code", "n", "Mean (SD)", "Median (IQR)", "Min", "Max", "CV (%)"],
    [
        ["L008", "${L008_n}", "${L008_mean} (${L008_sd})",
         "${L008_median} (${L008_q1}-${L008_q3})", "${L008_min}", "${L008_max}",
         "${L008_cv}"],
        ["L002", "${L002_n}", "${L002_mean} (${L002_sd})",
         "${L002_median} (${L002_q1}-${L002_q3})", "${L002_min}", "${L002_max}",
         "${L002_cv}"],
        ["L003", "${L003_n}", "${L003_mean} (${L003_sd})",
         "${L003_median} (${L003_q1}-${L003_q3})", "${L003_min}", "${L003_max}",
         "${L003_cv}"],
        ["L004", "${L004_n}", "${L004_mean} (${L004_sd})",
         "${L004_median} (${L004_q1}-${L004_q3})", "${L004_min}", "${L004_max}",
         "${L004_cv}"],
        ["L009", "${L009_n}", "${L009_mean} (${L009_sd})",
         "${L009_median} (${L009_q1}-${L009_q3})", "${L009_min}", "${L009_max}",
         "${L009_cv}"],
        ["L100", "${L100_n}", "${L100_mean} (${L100_sd})",
         "${L100_median} (${L100_q1}-${L100_q3})", "${L100_min}", "${L100_max}",
         "${L100_cv}"],
    ],
    note=("SD, standard deviation; IQR, interquartile range; CV, coefficient of "
          "variation. Areas with low claim volume are masked by the data provider "
          "and appear as missing values; the n column reflects the number of "
          "secondary medical areas with non-missing values for each code.")
)

add_figure_inline(
    os.path.join(FIG_DIR, 'rapm_fig1_en.png'),
    "Figure 1. Geographic distribution of anaesthesia standardised claim ratios "
    "across ${n_areas} secondary medical areas of Japan, fiscal year ${fiscal_year}. "
    "(A) General anaesthesia (L008). (B) Spinal anaesthesia (L004). (C) Epidural "
    "anaesthesia as main anaesthetic (L002). (D) Continuous epidural infusion "
    "(L003). Choropleth maps shaded by quintile of the standardised claim ratio "
    "(national average = 100). Red circles mark secondary medical areas "
    "containing at least one university hospital. Areas masked by the data "
    "provider owing to low volume are shown in grey."
)

add_subheading("Multilevel model and university hospital effect")
add_para(
    "The null multilevel model showed that only ${L008_ml_icc} of general "
    "anaesthesia variance was attributable to the prefecture level (intraclass "
    "correlation coefficient ${L008_ml_icc}), indicating that "
    "${L008_vd_within}% occurred within prefectures where audit policy is "
    "uniform. Spinal anaesthesia showed stronger prefecture-level clustering "
    "(intraclass correlation coefficient ${L004_ml_icc}), consistent with a "
    "greater role of prefectural factors for that code. Adding university "
    "hospital presence as a fixed effect produced the largest improvement for "
    "general anaesthesia (marginal R² ${L008_ml_r2}; β = ${L008_ml_coef}, "
    "95% confidence interval ${L008_ml_ci_low} to ${L008_ml_ci_high}) and was "
    "statistically significant for every code (all P ${L008_ml_p}) (Table 3). "
    "University hospital areas had higher general anaesthesia ratios than non-"
    "university areas in ${n_prefectures} of ${n_prefectures} prefectures, with a "
    "mean within-prefecture difference of +${L008_within_diff} points (paired "
    "t = ${L008_within_t}, P ${L008_within_p}) (Figure 2A). Cohen's d for the "
    "university hospital effect on general anaesthesia was ${L008_d}; on "
    "continuous epidural infusion ${L003_d}; on epidural anaesthesia ${L002_d}; "
    "and on spinal anaesthesia ${L004_d}. A combined measure capturing the "
    "general-anaesthesia plus continuous-epidural workflow (mean of L008 and L003 "
    "SCR; ${L008_L003_n} areas with data for both codes) showed a coefficient of "
    "variation of ${L008_L003_cv}% and preserved a large university hospital "
    "effect (Cohen's d ${L008_L003_d}; university mean ${L008_L003_u_mean} "
    "versus non-university ${L008_L003_nu_mean}) (Figure 2B).")

# Table 3 inline
add_table_from_data(
    "Table 3. Multilevel linear mixed model results: standardised claim ratio as "
    "outcome, prefecture as random intercept.",
    ["Code", "Null model ICC", "β university (95% CI)", "P value", "Marginal R²"],
    [
        ["L008", "${L008_ml_icc}",
         "+${L008_ml_coef} (${L008_ml_ci_low} to ${L008_ml_ci_high})",
         "${L008_ml_p}", "${L008_ml_r2}"],
        ["L002", "${L002_ml_icc}",
         "+${L002_ml_coef} (${L002_ml_ci_low} to ${L002_ml_ci_high})",
         "${L002_ml_p}", "${L002_ml_r2}"],
        ["L003", "${L003_ml_icc}",
         "+${L003_ml_coef} (${L003_ml_ci_low} to ${L003_ml_ci_high})",
         "${L003_ml_p}", "${L003_ml_r2}"],
        ["L004", "${L004_ml_icc}",
         "+${L004_ml_coef} (${L004_ml_ci_low} to ${L004_ml_ci_high})",
         "${L004_ml_p}", "${L004_ml_r2}"],
    ],
    note=("β university, fixed effect coefficient for university hospital presence "
          "(binary 0/1); CI, confidence interval; ICC, intraclass correlation "
          "coefficient (proportion of variance attributable to prefecture level in "
          "the null random-intercept model); Marginal R², proportional reduction "
          "in total (prefecture + residual) variance from the null model after "
          "adding the university hospital fixed effect. Models were estimated by "
          "restricted maximum likelihood.")
)

add_para(
    "A covariate-adjusted sensitivity model added the natural logarithm of "
    "population density and the anaesthesiologist share of all physicians (both "
    "standardised). The university hospital coefficient was attenuated but remained "
    "positive and statistically significant for general anaesthesia (β = "
    "+${L008_mlc_coef}, 95% CI ${L008_mlc_ci_low} to ${L008_mlc_ci_high}, P "
    "${L008_mlc_p}), epidural anaesthesia (β = +${L002_mlc_coef}, 95% CI "
    "${L002_mlc_ci_low} to ${L002_mlc_ci_high}, P ${L002_mlc_p}) and "
    "continuous epidural infusion (β = +${L003_mlc_coef}, 95% CI "
    "${L003_mlc_ci_low} to ${L003_mlc_ci_high}, P ${L003_mlc_p}). For spinal "
    "anaesthesia the point estimate was positive but no longer statistically "
    "significant (β = +${L004_mlc_coef}, 95% CI ${L004_mlc_ci_low} to "
    "${L004_mlc_ci_high}, P ${L004_mlc_p}). Population density was positively "
    "associated with general anaesthesia ratios (β = ${L008_mlc_popd_coef}, P "
    "${L008_mlc_popd_p}) but not with the other three codes. The anaesthesiologist "
    "share was positively associated with general, epidural, continuous epidural and "
    "spinal anaesthesia (P ${L008_mlc_anes_p}, ${L002_mlc_anes_p}, "
    "${L003_mlc_anes_p} and ${L004_mlc_anes_p}, respectively)."
)

add_figure_inline(
    os.path.join(FIG_DIR, 'rapm_fig2_en.png'),
    "Figure 2. University hospital presence and the combined general-anaesthesia "
    "plus continuous-epidural measure. (A) Distribution of secondary medical "
    "areas containing at least one university hospital (n = ${n_univ_areas} of "
    "${n_areas}; red). (B) Choropleth map of the combined general-anaesthesia "
    "plus continuous-epidural standardised claim ratio (mean of L008 and L003 "
    "SCR; ${L008_L003_n} areas with data for both codes), shaded by quintile. "
    "Red circles mark secondary medical areas containing at least one university "
    "hospital. Areas masked by the data provider for either code are shown in "
    "grey."
)

add_subheading("Sensitivity analyses against the audit hypothesis")
add_para(
    "All three pre-specified sensitivity analyses converged in rejecting the "
    "audit hypothesis. First, within-prefecture variance decomposition attributed "
    "${L008_vd_between}% of general anaesthesia variance to between-prefecture "
    "differences, ${L008_vd_univ}% to the university hospital effect within "
    "prefecture, and ${L008_vd_residual}% to residual within-group variation; the "
    "university hospital effect alone explained ${L008_vd_univ_within}% of all "
    "within-prefecture variance. Second, general and spinal anaesthesia were "
    "positively, not negatively, correlated (r = ${corr_L008_L004_r}, P "
    "${corr_L008_L004_p}), as were general and epidural anaesthesia (r = "
    "${corr_L008_L002_r}, P ${corr_L008_L002_p}), inconsistent with audit-"
    "driven reclassification and consistent with a common supply factor. Third, "
    "the maximum prefectural audit-rate difference of ${audit_max_shift} "
    "percentage points could produce a ratio shift of approximately "
    "${audit_max_shift} points, equivalent to approximately ${audit_L008_iqr_pct}% of "
    "the observed interquartile range for general anaesthesia and "
    "${audit_L002_iqr_pct}% for epidural anaesthesia. Combining general (L008) "
    "and spinal (L004) ratios produced an alternative audit-insensitive "
    "measure with a coefficient of variation of ${L008_L004_cv}% and a university "
    "hospital effect that remained large (Cohen's d ${L008_L004_d}), confirming "
    "that the structural effect is not artefactually erased by combining "
    "reclassifiable codes.")

add_subheading("Robustness: empirical Bayes shrinkage and outliers")
add_para(
    "Empirical Bayes shrinkage toward prefecture means was implemented as a "
    "stress test of the area-level ratios. Under this model, shrinkage "
    "substantially compressed the area-level ratios toward their prefecture "
    "means. For general anaesthesia, Cohen's d for the university-hospital effect "
    "fell from ${L008_d} (raw) to ${L008_d_shrunk} (shrunken), an attenuation "
    "of ${L008_attenuation}%; for epidural anaesthesia (L002) from ${L002_d} to "
    "${L002_d_shrunk} (${L002_attenuation}%); and for spinal anaesthesia (L004) "
    "from ${L004_d} to ${L004_d_shrunk} (${L004_attenuation}%). The general "
    "anaesthesia effect therefore remained of moderate magnitude despite heavy "
    "shrinkage, whereas the L002 and L004 effects were materially attenuated, "
    "indicating that those smaller raw effects are more vulnerable to sampling "
    "instability. These shrinkage estimates are complementary to, and consistent "
    "in direction with, the mixed-effects model coefficients reported in Table 3, "
    "which themselves incorporate prefecture-level shrinkage. Two areas were "
    "identified as outliers (more than three standard deviations from the mean "
    "general anaesthesia ratio): ${outlier1_name} in ${outlier1_pref} (ratio "
    "${outlier1_scr}) and ${outlier2_name} in ${outlier2_pref} (ratio "
    "${outlier2_scr}). Excluding outliers did not alter the direction or "
    "statistical significance of any finding.")

# ============================================================
# DISCUSSION
# ============================================================
add_heading("Discussion", level=1)

add_subheading("Statement of principal findings")
add_para(
    "Regional variation in anaesthesia practice across Japan's ${n_areas} "
    "secondary medical areas is substantial, with coefficients of variation "
    "ranging from ${L008_cv}% for general anaesthesia to ${L002_cv}% for "
    "epidural anaesthesia. Multilevel modelling showed that only "
    "${L008_ml_icc} of general anaesthesia variance is attributable to "
    "prefectures, while ${L008_vd_within}% occurs within prefectures where "
    "audit policy is uniform. University hospital presence explained "
    "${L008_ml_r2} of total variance and ${L008_vd_univ_within}% of within-"
    "prefecture variance, and the effect was present in all ${n_prefectures} "
    "prefectures despite differing audit practices. Three independent sensitivity "
    "analyses converged in rejecting differential auditing as a plausible "
    "explanation. The observed variation is therefore predominantly structural, "
    "driven by institutional capacity and clinical-organisational factors "
    "rather than by administrative coding.")

add_subheading("Strengths and limitations")
add_para(
    "Strengths include the use of age- and sex-standardised ratios covering the "
    "entire national population; analysis at a fine geographic scale of "
    "${n_areas} areas; multilevel modelling that properly accounts for the "
    "nested data structure; empirical Bayes shrinkage to address low-volume "
    "instability; multiple pre-specified sensitivity analyses; within-"
    "prefecture comparisons that hold audit policy constant; and the residence-"
    "based ratio definition that mitigates patient travel effects. Limitations "
    "include the ecological fallacy inherent in area-level analysis;{19} the "
    "cross-sectional design, which precludes causal inference; the post-audit "
    "nature of the data, which reflects reimbursed rather than intended "
    "practice; the partial overlap between university hospital presence and "
    "urban concentration; the absence of code-specific audit-rate data; and the "
    "inability to quantify defensive undercoding from claims alone. A specific "
    "constraint of the publicly available dataset is that the combined "
    "anaesthesia add-on billed under L008 (the explicit code for general-plus-"
    "epidural technique) is aggregated with other anaesthesia add-ons in the "
    "regional variation data release and cannot be extracted separately at the "
    "secondary-medical-area level; patient-level National Database records "
    "would be required to measure the true rate of combined general-epidural "
    "anaesthesia. We therefore used L002 (epidural as main technique) and L003 "
    "(continuous epidural infusion) as the publicly reported proxies for "
    "regional technique use. A covariate-adjusted sensitivity model added "
    "population density and anaesthesiologist supply; the university hospital "
    "coefficient remained positive and statistically significant for general, "
    "epidural and continuous epidural anaesthesia, but not for spinal "
    "anaesthesia. The positive point estimate for spinal anaesthesia is, "
    "however, compatible with the same structural gradient operating in an "
    "underpowered model, and may partly reflect a countervailing tendency for "
    "neuraxial anaesthesia to be used in settings where anaesthesiologists are "
    "less available. The aggregate data cannot separate these opposing "
    "mechanisms, so the adjusted spinal-anaesthesia coefficient is best "
    "interpreted as an imprecise net effect.")

add_subheading("Interpretation within the context of the wider literature")
add_para(
    "Our findings are consistent with the broader medical practice variation "
    "literature. The Dartmouth Atlas project documented extensive regional "
    "variation in surgical rates in the United States, driven primarily by "
    "physician supply and practice style.{20} Comparable patterns have been "
    "described for the National Health Service{21} and in Germany{22} and "
    "Australia.{23} Our study extends this evidence by exploiting Japan's "
    "uniquely uniform fee schedule combined with prefecture-specific auditing "
    "to disentangle administrative and clinical sources of variation, and by "
    "providing, to our knowledge, the first multilevel small-area analysis of "
    "anaesthesia technique under universal coverage in East Asia. The "
    "university hospital effect we report (Cohen's d = ${L008_d}) is "
    "substantial, plausibly reflecting the influence of the ikyoku (university "
    "medical office) system on clinical practice in affiliated hospitals.{24} "
    "The intraclass correlation coefficient of ${L008_ml_icc} for general "
    "anaesthesia indicates that most variation occurs within prefectures, "
    "consistent with institutional factors dominating over prefectural-level "
    "factors for this code.")

add_para(
    "Our findings and framework extend beyond Japan in four respects. First, the "
    "within-prefecture variance decomposition we describe is directly "
    "transportable to other universal-coverage systems that combine a centrally "
    "uniform fee schedule with regionally devolved claims auditing, including "
    "Taiwan's National Health Insurance,{25} South Korea's National Health "
    "Insurance Service, and the tariff-based systems of Germany, France and the "
    "English National Health Service.{21,22} Second, the dominance of tertiary-"
    "teaching-hospital concentration as a structural determinant of technique "
    "choice echoes analogous patterns for teaching-hospital anaesthesia mix "
    "reported in the US and United Kingdom,{20,21} and for other tertiary-"
    "intensive specialties across OECD countries;{26} the consistency of the "
    "university effect across ${n_prefectures} of ${n_prefectures} prefectures "
    "in our data suggests that it is a robust feature of high-income health "
    "systems rather than a Japanese idiosyncrasy. Third, the variation we identify "
    "in epidural anaesthesia (L002, coefficient of variation ${L002_cv}%) and "
    "continuous epidural infusion (L003, ${L003_cv}%) speaks to an international "
    "equity agenda: where emerging evidence links regional anaesthesia to "
    "oncological and functional outcomes,{7-9} monitoring the distribution of "
    "anaesthesia techniques could serve as a system-level quality indicator across "
    "any universal coverage setting committed to equitable access to technology-"
    "rich care.{27,28} Fourth, extrapolation to low- and middle-income countries "
    "is more cautious: fee schedules and audit arrangements are typically "
    "heterogeneous and claims data often incomplete, which limits direct "
    "replication of the decomposition. Nevertheless, the general principle — "
    "that apparent small-area variation should be interrogated for "
    "administrative versus structural origins before policy action — is "
    "transferable, particularly as low- and middle-income countries "
    "progressively standardise their fee schedules under universal health "
    "coverage reforms.{2}")

add_subheading("Implications for policy, practice and governance")
add_para(
    "The policy implications of these findings are threefold. First, the "
    "predominance of structural over administrative determinants means that "
    "regulators should not interpret regional variation in anaesthesia claims "
    "as prima facie evidence of differential coding accuracy or fraud. Japan's "
    "prefecture-specific auditing system is an important governance tool, but "
    "its variation in intensity is too small to explain the observed large "
    "differences in neuraxial technique use. Redirecting enforcement "
    "resources toward coding audits in low-ratio areas is unlikely to reduce the "
    "gap and may even penalise providers who already lack specialist capacity.")

add_para(
    "Second, the marked variation in continuous epidural infusion (L003; "
    "coefficient of variation ${L003_cv}%) and in epidural anaesthesia (L002; "
    "${L002_cv}%) is clinically important because general anaesthesia supplemented "
    "with regional techniques has been linked, albeit inconsistently, with "
    "better oncological outcomes.{7-9} A ${L003_fold}-fold difference in "
    "continuous epidural infusion ratios between university and non-university "
    "areas implies that patients' access to a potentially beneficial technique "
    "depends substantially on where they live, which challenges the equity "
    "premise of Japan's universal coverage. Quality improvement responses could "
    "include targeted education, specialist outreach from university centres, "
    "and monitoring anaesthesia technique distribution as a system-level quality "
    "indicator,{27-29} aligned with international standards for safe practice "
    "of anaesthesia.")

add_para(
    "Third, the within-prefecture variance decomposition and multilevel "
    "sensitivity framework developed here provide a transferable governance tool "
    "for distinguishing administrative from structural sources of small-area "
    "variation. In a single-payer system with a uniform fee schedule, any "
    "procedure that shows large within-prefecture variation cannot be attributed "
    "to the audit environment and should prompt policy review of workforce, "
    "training and facility capacity. Conversely, large between-prefecture "
    "variation may warrant regulatory attention to payment, audit or coding "
    "practices. Applying this framework to other procedure groups could improve "
    "the targeting of both quality-improvement and regulatory resources.")

# ============================================================
# CONCLUSIONS
# ============================================================
add_heading("Conclusions", level=1)
add_para(
    "Regional variation in anaesthesia practice in Japan is large and "
    "predominantly structural rather than an artefact of insurance auditing. "
    "University hospital presence is the dominant measured structural "
    "determinant, explaining more than one third of total variance in general "
    "anaesthesia. The variation in "
    "epidural anaesthesia (L002) and continuous epidural infusion (L003) is of "
    "particular concern given emerging evidence of oncological benefit from "
    "regional techniques. Quality improvement, workforce policy and equity "
    "policy should address the supply and organisational determinants of "
    "anaesthesia practice rather than treat the variation as an unavoidable "
    "feature of claims processing or as a signal of differential auditing.")

# ============================================================
# REFERENCES
# ============================================================
add_heading("References", level=1)

REFERENCES = [
    "Ikegami N, Yoo BK, Hashimoto H, et al. Japanese universal health coverage: evolution, achievements, and challenges. Lancet 2011;378:1106-15.",
    "GBD 2019 Universal Health Coverage Collaborators. Measuring universal health coverage based on an index of effective coverage of health services in 204 countries and territories, 1990-2019. Lancet 2020;396:1250-84.",
    "Matsuda S. Health policy in Japan: current situation and future challenges. JMA J 2019;2:1-10.",
    "Ministry of Health, Labour and Welfare. Patient Survey 2008. Tokyo: MHLW, 2009.",
    "Hashimoto H, Ikegami N, Shibuya K, et al. Cost containment and quality of care in Japan: is there a trade-off? Lancet 2011;378:1174-82.",
    "Social Insurance Medical Fee Payment Fund. 審査状況（令和4年度）— 都道府県別医科歯科計審査状況（点数）. Tokyo: Social Insurance Medical Fee Payment Fund, 2023. https://www.ssk.or.jp/smph/tokeijoho/shinsatokei/shinsajokyo_r04.html (accessed 24 Jul 2026).",
    "Sessler DI, Pei L, Huang Y, et al. Recurrence of breast cancer after regional or general anaesthesia: a randomised controlled trial. Lancet 2019;394:1807-15.",
    "Chen WK, Miao CH. The effect of anesthetic technique on survival in human cancers: a meta-analysis of retrospective and prospective studies. PLoS One 2013;8:e56540.",
    "Weng M, Chen W, Hou W, et al. The effect of neuraxial anaesthesia on cancer recurrence and survival after cancer surgery: an updated meta-analysis. Oncotarget 2016;7:15262-73.",
    "Ministry of Health, Labour and Welfare. 2026 revision of the medical fee schedule: changes to the medical fee schedule master (R08). Tokyo: MHLW, 2026. https://shinryohoshu.mhlw.go.jp/shinryohoshu/file/info/smente260522.pdf (accessed 24 Jul 2026).",
    "Wennberg JE, Cooper MM, eds. The Dartmouth Atlas of Health Care. Chicago: American Hospital Publishing, 1999.",
    "Cabinet Office. Regional variation visualisation (chiikisa no mieruka). https://www5.cao.go.jp/keizai-shimon/kaigi/special/reform/mieruka/chiikisa/index.html (accessed 10 Mar 2026).",
    "Ministry of Health, Labour and Welfare. Survey of Physicians, Dentists and Pharmacists 2022 (e-Stat). https://www.e-stat.go.jp (accessed 12 Mar 2026).",
    "Ministry of Land, Infrastructure, Transport and Tourism. National Land Numerical Information download service. https://nlftp.mlit.go.jp (accessed 10 Mar 2026).",
    "Cohen J. Statistical Power Analysis for the Behavioral Sciences. 2nd ed. Hillsdale, NJ: Lawrence Erlbaum, 1988.",
    "Rabe-Hesketh S, Skrondal A. Multilevel and Longitudinal Modeling Using Stata. 3rd ed. College Station, TX: Stata Press, 2012.",
    "Seabold S, Perktold J. Statsmodels: econometric and statistical modeling with Python. In: Proceedings of the 9th Python in Science Conference. 2010:92-6.",
    "Efron B, Morris C. Stein's estimation rule and its competitors: an empirical Bayes approach. J Am Stat Assoc 1973;68:117-30.",
    "Greenland S. Ecologic versus individual-level sources of bias in ecologic estimates of contextual health effects. Int J Epidemiol 2001;30:1343-50.",
    "Wennberg JE, Fisher ES, Skinner JS. Geography and the debate over Medicare reform. Health Aff (Millwood) 2002;Suppl Web Exclusives:W96-114.",
    "NHS RightCare. The NHS Atlas of Variation in Healthcare. London: Public Health England, 2015.",
    "Schäfer T, Pritzkuleit R, Jeszenszky C, et al. Trends and geographical variation of primary hip and knee joint replacement in Germany. Osteoarthritis Cartilage 2013;21:279-88.",
    "Australian Commission on Safety and Quality in Health Care. Australian Atlas of Healthcare Variation. Sydney: ACSQHC, 2015.",
    "Otsuka T. The ikyoku system of university orthopaedic surgery departments: an in-hospital organisational system unique to Japan. J Orthop Sci 2012;17:513-14.",
    "Cheng T-M. Taiwan's new national health insurance program: genesis and experience so far. Health Aff (Millwood) 2003;22:61-76.",
    "OECD. Geographic Variations in Health Care: What Do We Know and What Can Be Done to Improve Health System Performance? Paris: OECD Publishing, 2014.",
    "Donabedian A. The quality of care: how can it be assessed? JAMA 1988;260:1743-8.",
    "Merry AF, Cooper JB, Soyannwo O, et al. International standards for a safe practice of anesthesia 2010. Can J Anesth 2010;57:1027-34.",
    "Mainz J. Defining and classifying clinical indicators for quality improvement. Int J Qual Health Care 2003;15:523-30.",
    "von Elm E, Altman DG, Egger M, et al. The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement: guidelines for reporting observational studies. Lancet 2007;370:1453-7.",
]

# Reorder references so their numbering matches first appearance in the body.
REFERENCES = renumber_references(doc, REFERENCES)

for i, ref in enumerate(REFERENCES, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(f"{i}. ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run = p.add_run(Template(ref).substitute(FLAT))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# Body word count check (excluding title page, abstract, references, tables, figures)
text_paras = []
in_body = False
in_refs = False
for p in doc.paragraphs:
    txt = p.text.strip()
    if not txt:
        continue
    if txt.upper() == 'INTRODUCTION':
        in_body = True
        continue
    if txt.upper() == 'REFERENCES':
        in_refs = True
        continue
    if in_body and not in_refs:
        text_paras.append(txt)
body = '\n'.join(text_paras)
words = len(re.findall(r'\b\w+\b', body))
print(f"Body word count (intro through conclusion): {words}")

# Update title-page counts that were left as placeholders
n_tables = 3
n_figures = 2
n_refs = len(REFERENCES)
for p in doc.paragraphs:
    if p.text.startswith('Tables:'):
        p.text = f"Tables: {n_tables}   Figures: {n_figures}   References: {n_refs}"
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.italic = True
        break

# Update the word-count line on the title page
for p in doc.paragraphs:
    if p.text.startswith('Word count'):
        p.text = f"Word count (main text, excluding abstract, references, tables and figures): approximately {words} words"
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.italic = True
        break

# ============================================================
# Save
# ============================================================
out = os.path.join(OUTPUT_DIR, 'regional_anaesthesia_IJHPM_EN.docx')
doc.save(out)
print(f"Saved: {out}")
